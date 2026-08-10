# -*- coding: utf-8 -*-
# =====================================================================
# EXAME INTELIGENTE 4.0 - VERSAO WEB (Streamlit)
# Porta do app_desktop.py para funcionar como site no navegador,
# reutilizando os mesmos arquivos JSON do desktop.
# =====================================================================
import streamlit as st
import json
import os
import base64
import random
import io
import calendar
import re
import csv
import hashlib
import unicodedata
import hmac
import shutil
import pathlib
import html
from urllib.parse import unquote

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import streamlit.components.v1 as components
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml import parse_xml
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

try:
    import psycopg2
    from psycopg2 import pool as psycopg2_pool
    _TEM_PSYCOPG2 = True
except Exception:
    psycopg2 = None
    _TEM_PSYCOPG2 = False

# =====================================================================
# 0. INICIALIZACAO: garante que rode via `streamlit run`
# (evita o modo "bare" ao executar com `python app_web.py` pelo VS Code)
# =====================================================================
import sys
if __name__ == "__main__" and not st.runtime.exists():
    try:
        subprocess_ok = True
        import subprocess
    except Exception:
        subprocess_ok = False
    print("Executando via Streamlit...")
    if subprocess_ok:
        subprocess.run([sys.executable, "-m", "streamlit", "run", __file__])
    else:
        print("Nao foi possivel iniciar o Streamlit automaticamente.")
        print('Execute manualmente no terminal: streamlit run app_web.py')
    sys.exit(0)

# =====================================================================
# 0. CONFIGURACAO INICIAL
# =====================================================================
st.set_page_config(
    page_title="Exame Inteligente 4.0 - Web",
    layout="wide",
    initial_sidebar_state="expanded"
)

ARQUIVO_BANCO = "banco_questoes.json"
ARQUIVO_CONFIG = "configuracoes.json"
ARQUIVO_PLANOS = "planos_aula.json"
ARQUIVO_GRADE = "grade_horaria.json"
ARQUIVO_ANOTACOES = "anotacoes.json"
ARQUIVO_TURMAS = "turmas_alunos.json"
ARQUIVO_CONFIG_GRADE = "config_grade.json"
ARQUIVO_AVALIACOES = "avaliacoes.json"
ARQUIVO_MAPEAMENTO = "mapeamento_sala.json"
PASTA_IMAGENS = "imagens_apoio"
ARQUIVO_USUARIOS = "usuarios.json"
ARQUIVO_VERIFICACOES = "verificacoes_email.json"
PASTA_DADOS = "dados_usuarios"

if not os.path.exists(PASTA_IMAGENS):
    os.makedirs(PASTA_IMAGENS)

CONFIG_PADRAO = {
    "escola": "CEJA Profa Cecy Cialdini",
    "professor": "Marcio Cordeiro",
    "fonte": "Arial",
    "tamanho_fonte": 11,
    "margem_cm": 1.5,
    "espacamento_pt": 0,
    "usar_duas_colunas": True,
    "mostrar_aluno": True,
    "mostrar_turma": True,
    "mostrar_data": True,
    "aparencia": "System",
    "tema_visual": "",
    "cor_tema": "blue",
    "cor_principal": "#1f538d",
    "cor_secundaria": "#14375e",
    "cor_borda_card": "",
    "cor_fundo": "#2b2b2b",
    "provedor_ia": "Google Gemini"
}

DESCRITORES_SAEB = [
    "D1 - Identificar a localizacao/movimentacao de objeto em mapas, croquis e outras representacoes graficas.",
    "D2 - Identificar propriedades comuns e diferencas entre figuras bidimensionais e tridimensionais, relacionando-as com as suas planificacoes.",
    "D3 - Identificar propriedades de triangulos pela comparacao de medidas de lados e angulos.",
    "D4 - Identificar relacao entre quadrilateros por meio de suas propriedades.",
    "D5 - Reconhecer a conservacao ou modificacao de medidas dos lados, do perimetro, da area em ampliacao e/ou reducao de figuras poligonais usando malhas quadriculadas.",
    "D6 - Reconhecer angulos como mudanca de direcao ou giros, identificando angulos retos e nao-retos.",
    "D7 - Reconhecer que as imagens de uma figura construida por uma transformacao homotetica sao semelhantes, identificando propriedades e/ou medidas que se modificam ou nao se alteram.",
    "D8 - Resolver problema utilizando propriedades dos poligonos (soma de seus angulos internos, numero de diagonais, calculo da medida de cada angulo interno nos poligonos regulares).",
    "D9 - Interpretar informacoes apresentadas por meio de coordenadas cartesianas.",
    "D10 - Utilizar relacoes metricas do triangulo retangulo para resolver problemas significativos.",
    "D11 - Reconhecer circulo/circunferencia, seus elementos e algumas de suas relacoes.",
    "D12 - Resolver problema envolvendo o calculo de perimetro de figuras planas.",
    "D13 - Resolver problema envolvendo o calculo de area de figuras planas.",
    "D14 - Resolver problema envolvendo nocoes de volume.",
    "D15 - Resolver problema utilizando relacoes entre diferentes unidades de medida.",
    "D16 - Identificar a localizacao de numeros inteiros na reta numerica.",
    "D17 - Identificar a localizacao de numeros racionais na reta numerica.",
    "D18 - Efetuar calculos com numeros inteiros, envolvendo as operacoes (adicao, subtracao, multiplicacao, divisao, potenciacao).",
    "D19 - Resolver problema com numeros naturais, envolvendo diferentes significados das operacoes (adicao, subtracao, multiplicacao, divisao, potenciacao).",
    "D20 - Resolver problema com numeros inteiros envolvendo as operacoes (adicao, subtracao, multiplicacao, divisao, potenciacao).",
    "D21 - Reconhecer as diferentes representacoes de um numero racional.",
    "D22 - Identificar fracao como representacao que pode estar associada a diferentes significados.",
    "D23 - Identificar fracoes equivalentes.",
    "D24 - Reconhecer as representacoes decimais dos numeros racionais como uma extensao do sistema de numeracao decimal, identificando a existencia de ordens como decimos, centesimos e milesimos.",
    "D25 - Efetuar calculos que envolvam operacoes com numeros racionais (adicao, subtracao, multiplicacao, divisao, potenciacao).",
    "D26 - Resolver problema com numeros racionais envolvendo as operacoes (adicao, subtracao, multiplicacao, divisao, potenciacao).",
    "D27 - Efetuar calculos simples com valores aproximados de radicais.",
    "D28 - Resolver problema que envolva porcentagem.",
    "D29 - Resolver problema que envolva variacao proporcional, direta ou inversa, entre grandezas.",
    "D30 - Calcular o valor numerico de uma expressao algebrica.",
    "D31 - Resolver problema que envolva equacao do 2o grau.",
    "D32 - Identificar a expressao algebrica que expressa uma regularidade observada em sequencias de numeros ou figuras (padroes).",
    "D33 - Identificar uma equacao ou inequacao do 1o grau que expressa um problema.",
    "D34 - Identificar um sistema de equacoes do 1o grau que expressa um problema.",
    "D35 - Identificar a relacao entre as representacoes algebrica e geometrica de um sistema de equacoes do 1o grau.",
    "D36 - Resolver problema envolvendo informacoes apresentadas em tabelas e/ou graficos.",
    "D37 - Associar informacoes apresentadas em listas e/ou tabelas simples aos graficos que as representam e vice-versa."
]

DESCRITORES_MAT = {
    "D1": "Localizacao/movimentacao de objeto em mapas/croquis.",
    "D2": "Propriedades e planificacoes de figuras 2D e 3D.",
    "D3": "Propriedades de triangulos (lados e angulos).",
    "D4": "Relacao entre quadrilateros.",
    "D5": "Conservacao/modificacao em figuras poligonais (malhas).",
    "D6": "Angulos: mudanca de direcao, giros, retos e nao retos.",
    "D7": "Transformacao homotetica e semelhanca.",
    "D8": "Propriedades dos poligonos (soma de angulos, diagonais).",
    "D9": "Informacoes apresentadas em coordenadas cartesianas.",
    "D10": "Relacoes metricas do triangulo retangulo.",
    "D11": "Circulo/circunferencia e suas relacoes.",
    "D12": "Calculo de perimetro de figuras planas.",
    "D13": "Calculo de area de figuras planas.",
    "D14": "Nocoes de volume.",
    "D15": "Relacoes entre diferentes unidades de medida.",
    "D16": "Numeros inteiros na reta numerica.",
    "D17": "Numeros racionais na reta numerica.",
    "D18": "Calculos com numeros inteiros (4 operacoes).",
    "D19": "Problemas com numeros naturais (4 operacoes e potencias).",
    "D20": "Problemas com numeros inteiros (4 operacoes e potencias).",
    "D21": "Diferentes representacoes de um numero racional.",
    "D22": "Fracao com diferentes significados.",
    "D23": "Fracoes equivalentes.",
    "D24": "Representacoes decimais (decimos, centesimos, milesimos).",
    "D25": "Calculos com numeros racionais.",
    "D26": "Problemas com numeros racionais.",
    "D27": "Calculos simples com valores aproximados de radicais.",
    "D28": "Problemas com porcentagem.",
    "D29": "Variacao proporcional, direta ou inversa.",
    "D30": "Valor numerico de uma expressao algebrica.",
    "D31": "Equacao do 2o grau em problemas.",
    "D32": "Expressao algebrica em sequencias numericas (padroes).",
    "D33": "Equacao ou inequacao do 1o grau num problema.",
    "D34": "Sistema de equacoes do 1o grau num problema.",
    "D35": "Representacoes algebrica e geometrica de sistema do 1o grau.",
    "D36": "Informacoes apresentadas em tabelas e/ou graficos.",
    "D37": "Associacao entre listas/tabelas e graficos."
}

def esc(s):
    import html
    return html.escape(str(s), quote=True)

# =====================================================================
# 0A. BANCO EXTERNO (Supabase/Postgres) - persistencia na nuvem
#     Quando BANCO_URL estiver configurado (.streamlit/secrets.toml ou
#     variavel de ambiente), todos os dados passam a viver no Postgres.
#     As funcoes carregar_* / salvar_* continuam iguais para o resto do app.
# =====================================================================
def _obter_banco_url():
    env = os.environ.get("BANCO_URL", "").strip()
    if env:
        if env.lower() in ("0", "off", "false", "nenhum"):
            return ""
        return env
    try:
        if "BANCO_URL" in st.secrets:
            return str(st.secrets["BANCO_URL"]).strip()
    except Exception:
        pass
    return ""

BANCO_CONFIGURADO = bool(_obter_banco_url())
BANCO_ERRO = ""
if BANCO_CONFIGURADO and not _TEM_PSYCOPG2:
    BANCO_ERRO = "psycopg2 nao instalado"
BANCO_ATIVO = BANCO_CONFIGURADO and _TEM_PSYCOPG2

def _conectar_banco():
    url = _obter_banco_url()
    if "?" in url:
        url = url.split("?", 1)[0]
    conn = psycopg2.connect(url, sslmode="require", connect_timeout=8,
                            keepalives=1, keepalives_idle=20, keepalives_interval=10,
                            keepalives_count=3)
    conn.autocommit = True
    return conn

@st.cache_resource(show_spinner=False)
def _pool_conexoes():
    url = _obter_banco_url()
    if "?" in url:
        url = url.split("?", 1)[0]
    return psycopg2_pool.ThreadedConnectionPool(
        minconn=1, maxconn=4, dsn=url, sslmode="require",
        connect_timeout=8, keepalives=1, keepalives_idle=20,
        keepalives_interval=10, keepalives_count=3)

def _com_conexao(fn):
    # Reutiliza conexoes do pool (uma unica criacao no processo) e, se a
    # conexao caiu (ex.: banco dormiu), descarta o pool e tenta uma vez.
    tentativas = 2
    for i in range(tentativas):
        pool = _pool_conexoes()
        conn = None
        try:
            conn = pool.getconn()
            conn.autocommit = True
            with conn.cursor() as cur:
                resultado = fn(conn, cur)
            pool.putconn(conn)
            conn = None
            return resultado
        except Exception:
            if conn is not None:
                try:
                    pool.putconn(conn, close=True)
                except Exception:
                    pass
            try:
                pool.closeall()
            except Exception:
                pass
            _pool_conexoes.clear()
            if i + 1 >= tentativas:
                raise

_TABELAS_OK = False

def _garantir_tabelas():
    global _TABELAS_OK
    if _TABELAS_OK:
        return
    def _criar(conn, cur):
        cur.execute("""
            CREATE TABLE IF NOT EXISTS app_dados (
                chave text PRIMARY KEY,
                valor jsonb NOT NULL,
                atualizado_em timestamptz NOT NULL DEFAULT now()
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS app_imagens (
                chave text PRIMARY KEY,
                dados bytea NOT NULL,
                atualizado_em timestamptz NOT NULL DEFAULT now()
            )
        """)
    _com_conexao(_criar)
    _TABELAS_OK = True

def db_get_json(chave):
    try:
        _garantir_tabelas()
        def _get(conn, cur):
            cur.execute("SELECT valor FROM app_dados WHERE chave=%s", (chave,))
            linha = cur.fetchone()
            return linha[0] if linha else None
        return _com_conexao(_get)
    except Exception:
        return None

def db_set_json(chave, dados):
    try:
        _garantir_tabelas()
        texto = json.dumps(dados, ensure_ascii=False)
        def _set(conn, cur):
            cur.execute(
                "INSERT INTO app_dados (chave, valor, atualizado_em) "
                "VALUES (%s, %s::jsonb, now()) "
                "ON CONFLICT (chave) DO UPDATE "
                "SET valor=EXCLUDED.valor, atualizado_em=now()",
                (chave, texto))
        _com_conexao(_set)
    except Exception as e:
        print("[BANCO] Erro ao salvar", chave, ":", e)
        return str(e)
    return None

def db_get_bytes(chave):
    try:
        _garantir_tabelas()
        def _get(conn, cur):
            cur.execute("SELECT dados FROM app_imagens WHERE chave=%s", (chave,))
            linha = cur.fetchone()
            return bytes(linha[0]) if linha else None
        return _com_conexao(_get)
    except Exception:
        return None

def db_set_bytes(chave, dados):
    try:
        _garantir_tabelas()
        def _set(conn, cur):
            cur.execute(
                "INSERT INTO app_imagens (chave, dados, atualizado_em) "
                "VALUES (%s, %s, now()) "
                "ON CONFLICT (chave) DO UPDATE "
                "SET dados=EXCLUDED.dados, atualizado_em=now()",
                (chave, psycopg2.Binary(bytes(dados))))
        _com_conexao(_set)
    except Exception as e:
        print("[BANCO] Erro ao salvar imagem", chave, ":", e)

def _chave_de_caminho(caminho):
    if not caminho:
        return None
    norm = str(caminho).replace("\\", "/")
    partes = [p for p in norm.split("/") if p]
    if partes == [ARQUIVO_USUARIOS]:
        return "g::" + ARQUIVO_USUARIOS
    if partes == [ARQUIVO_VERIFICACOES]:
        return "g::" + ARQUIVO_VERIFICACOES
    if len(partes) >= 3 and partes[0] == PASTA_DADOS:
        return "u::" + partes[1] + "::" + "/".join(partes[2:])
    return None

def imagem_existe(caminho):
    chave = _chave_de_caminho(caminho)
    if BANCO_ATIVO and chave:
        return db_get_bytes(chave) is not None
    return bool(caminho) and os.path.exists(caminho)

def imagem_bytes(caminho):
    chave = _chave_de_caminho(caminho)
    if BANCO_ATIVO and chave:
        dados = db_get_bytes(chave)
        if dados is not None:
            return dados
    if caminho and os.path.exists(caminho):
        with open(caminho, "rb") as f:
            return f.read()
    return b""

def salvar_imagem_usuario(nome_arquivo, dados):
    caminho = os.path.join(pasta_imagens(), nome_arquivo)
    chave = _chave_de_caminho(caminho)
    if BANCO_ATIVO and chave:
        db_set_bytes(chave, dados)
    else:
        with open(caminho, "wb") as f:
            f.write(bytes(dados))
    return caminho

def _testar_banco():
    global BANCO_ERRO
    try:
        _garantir_tabelas()
        BANCO_ERRO = ""
        return True
    except Exception as e:
        BANCO_ERRO = str(e)
        print("[BANCO] Nao foi possivel conectar:", e)
        return False

def banco_falho():
    return BANCO_CONFIGURADO and (not BANCO_ATIVO or BANCO_ERRO)

def _tentar_religar_banco():
    global BANCO_ATIVO, BANCO_ERRO
    if not BANCO_CONFIGURADO:
        return True
    if BANCO_ATIVO and not BANCO_ERRO:
        return True
    if _testar_banco():
        BANCO_ATIVO = True
        _carregar_json_cache.clear()
        return True
    return False

def aviso_banco():
    if banco_falho():
        st.warning(
            "O banco de dados externo nao esta acessivel no momento. "
            "Por isso, as contas e os dados salvos podem nao aparecer "
            "e o login pode falhar. Verifique a configuracao do banco "
            "(variavel BANCO_URL) e tente novamente em instantes."
        )

if BANCO_ATIVO and not _testar_banco():
    print("[BANCO] Desativando banco externo - usando arquivos locais.")
    BANCO_ATIVO = False

# =====================================================================
# 1. UTILITARIOS DE DADOS (JSON) - mesmos arquivos do desktop
#    Desde o port web, cada conta tem a propria pasta em dados_usuarios/
# =====================================================================
def usuario_atual():
    return st.session_state.get("usuario", None)

def pasta_usuario(usuario=None):
    usuario = usuario or usuario_atual()
    if not usuario:
        return None
    pasta = os.path.join(PASTA_DADOS, str(usuario))
    os.makedirs(pasta, exist_ok=True)
    return pasta

def caminho_usuario(nome_arquivo):
    pasta = pasta_usuario()
    if not pasta:
        return None
    return os.path.join(pasta, nome_arquivo)

def pasta_imagens():
    pasta = pasta_usuario()
    if not pasta:
        return PASTA_IMAGENS
    destino = os.path.join(pasta, PASTA_IMAGENS)
    os.makedirs(destino, exist_ok=True)
    return destino

@st.cache_data(show_spinner=False, ttl=120)
def _carregar_json_cache(caminho):
    chave = _chave_de_caminho(caminho) if BANCO_ATIVO else None
    if chave:
        dados = db_get_json(chave)
        if dados is not None:
            return dados
    if caminho and os.path.exists(caminho):
        try:
            with open(caminho, "r", encoding="utf-8") as arquivo:
                return json.load(arquivo)
        except Exception:
            return None
    return None

def carregar_json(caminho, padrao):
    dados = _carregar_json_cache(caminho)
    if dados is not None:
        return dados
    return padrao

def salvar_json(caminho, dados):
    if not caminho:
        return
    _carregar_json_cache.clear(caminho)
    if BANCO_ATIVO:
        chave = _chave_de_caminho(caminho)
        if chave:
            erro = db_set_json(chave, dados)
            if erro:
                st.error(f"Nao foi possivel salvar no banco de dados "
                         f"({chave}): {erro}")
            return
    with open(caminho, "w", encoding="utf-8") as arquivo:
        json.dump(dados, arquivo, ensure_ascii=False, indent=4)

# ---- Chave de API da IA (salva por usuario, para nao digitar toda vez) ----
def carregar_chave_ia():
    caminho = caminho_usuario("chave_ia.json")
    dados = carregar_json(caminho, {})
    if isinstance(dados, dict):
        return str(dados.get("chave", "") or "")
    return ""

def salvar_chave_ia(chave):
    caminho = caminho_usuario("chave_ia.json")
    if caminho:
        salvar_json(caminho, {"chave": str(chave or "").strip()})

# ---- Contas de usuario (login local em JSON) ----
def carregar_usuarios():
    dados = carregar_json(ARQUIVO_USUARIOS, [])
    return dados if isinstance(dados, list) else []

def salvar_usuarios(usuarios):
    salvar_json(ARQUIVO_USUARIOS, usuarios)

def hash_senha(senha, salt_hex=None):
    salt = bytes.fromhex(salt_hex) if salt_hex else os.urandom(16)
    digest = hashlib.pbkdf2_hmac("sha256", senha.encode("utf-8"), salt, 200000)
    return salt.hex(), digest.hex()

def usuario_existe(usuario):
    alvo = (usuario or "").lower()
    return any(u.get("usuario", "").lower() == alvo for u in carregar_usuarios())

def email_existe(email):
    alvo = (email or "").strip().lower()
    return any(str(u.get("email", "")).lower() == alvo for u in carregar_usuarios())

def _login_alvo(login):
    login = (login or "").strip()
    return login.lower(), login

def autenticar_usuario(login, senha):
    alvo = (login or "").strip().lower()
    for u in carregar_usuarios():
        ident = (str(u.get("usuario", "")) or "").lower()
        email_u = (str(u.get("email", "")) or "").lower()
        if ident == alvo or (alvo and email_u == alvo):
            salt, digest = u.get("senha_hash", ["", ""])
            if not salt:
                return None
            _, novo_digest = hash_senha(senha, salt)
            if hmac.compare_digest(novo_digest, digest):
                return u
            return None
    return None

def conta_atual(usuario=None):
    usuario = usuario or usuario_atual()
    if not usuario:
        return None
    for u in carregar_usuarios():
        if u.get("usuario", "").lower() == usuario.lower():
            return u
    return None

def migrar_dados_primeira_conta(usuario):
    pasta = pasta_usuario(usuario)
    if not pasta:
        return
    for nome_arquivo in [ARQUIVO_BANCO, ARQUIVO_CONFIG, ARQUIVO_PLANOS,
                         ARQUIVO_GRADE, ARQUIVO_ANOTACOES, ARQUIVO_TURMAS,
                         ARQUIVO_CONFIG_GRADE, ARQUIVO_AVALIACOES]:
        origem = nome_arquivo
        destino = os.path.join(pasta, nome_arquivo)
        if os.path.exists(origem) and not os.path.exists(destino):
            try:
                shutil.copy2(origem, destino)
            except Exception:
                pass
    origem_imgs = PASTA_IMAGENS
    destino_imgs = os.path.join(pasta, PASTA_IMAGENS)
    if os.path.isdir(origem_imgs) and not os.path.isdir(destino_imgs):
        try:
            shutil.copytree(origem_imgs, destino_imgs)
        except Exception:
            pass
    _ajustar_imagens_banco(pasta)

def _ajustar_imagens_banco(pasta):
    caminho_banco = os.path.join(pasta, ARQUIVO_BANCO)
    if not os.path.exists(caminho_banco):
        return
    try:
        with open(caminho_banco, "r", encoding="utf-8") as f:
            banco = json.load(f)
        prefixo = os.path.join(PASTA_DADOS, os.path.basename(pasta), PASTA_IMAGENS)
        mudou = False
        for q in banco:
            img = q.get("imagem", "")
            if img and img.startswith(PASTA_IMAGENS + os.sep):
                q["imagem"] = os.path.join(prefixo, os.path.basename(img))
                mudou = True
        if mudou:
            salvar_json(caminho_banco, banco)
    except Exception:
        pass

def criar_conta(nome, usuario, senha, email="", email_verificado=False):
    usuarios = carregar_usuarios()
    primeira = len(usuarios) == 0
    salt, digest = hash_senha(senha)
    usuarios.append({
        "nome": nome.strip(),
        "usuario": usuario.strip(),
        "email": (email or "").strip().lower(),
        "email_verificado": bool(email_verificado),
        "senha_hash": [salt, digest],
        "cookie_token": os.urandom(16).hex(),
        "onboarding_pendente": bool(not nome.strip()),
        "criado_em": datetime.now().strftime("%d/%m/%Y %H:%M")
    })
    salvar_usuarios(usuarios)
    if primeira:
        migrar_dados_primeira_conta(usuario.strip())
    return usuario.strip()

def conta_precisa_onboarding():
    conta = conta_atual()
    return bool(conta and conta.get("onboarding_pendente"))

def concluir_onboarding(nome, escola):
    conta = conta_atual()
    if conta:
        usuarios = carregar_usuarios()
        alvo = str(conta.get("usuario", "")).lower()
        for u in usuarios:
            if str(u.get("usuario", "")).lower() == alvo:
                u["nome"] = (nome or "").strip()
                u["onboarding_pendente"] = False
                break
        salvar_usuarios(usuarios)
    config = carregar_config()
    config["professor"] = (nome or "").strip()
    config["escola"] = (escola or "").strip()
    salvar_config(config)

# ---- Verificacao de email (codigo enviado por SMTP) ----
def _smtp_config():
    """Configuracoes SMTP vindas de st.secrets['smtp'] (host, port, user, password, from, starttls)."""
    try:
        s = dict(st.secrets.get("smtp", {}) or {})
        return s if s.get("host") and s.get("user") else {}
    except Exception:
        return {}

def _smtp_disponivel():
    cfg = _smtp_config()
    return bool(cfg.get("host") and cfg.get("user"))

def _enviar_email(destino, assunto, corpo_html):
    cfg = _smtp_config()
    if not cfg:
        return False
    try:
        msg = MIMEMultipart("alternative")
        msg["Subject"] = assunto
        msg["From"] = cfg.get("from") or cfg.get("user")
        msg["To"] = destino
        msg.attach(MIMEText(corpo_html, "html", "utf-8"))
        porta = int(cfg.get("port", 587))
        servidor = smtplib.SMTP(cfg["host"], porta, timeout=30)
        if str(cfg.get("starttls", "true")).lower() != "false":
            servidor.starttls()
        servidor.login(cfg["user"], cfg["password"])
        servidor.sendmail(msg["From"], [destino], msg.as_string())
        servidor.quit()
        return True
    except Exception as e:
        print("[EMAIL] Falha ao enviar para", destino, ":", e)
        return False

def _novo_codigo_verificacao(email, valido_por_min=15):
    codigo = f"{random.randint(100000, 999999)}"
    registros = carregar_json(ARQUIVO_VERIFICACOES, {})
    if not isinstance(registros, dict):
        registros = {}
    alvo = (email or "").strip().lower()
    registros[alvo] = {
        "codigo": codigo,
        "expira_em": (datetime.now() + timedelta(minutes=valido_por_min)).isoformat(),
        "tentativas": 0,
        "criado_em": datetime.now().isoformat(),
    }
    salvar_json(ARQUIVO_VERIFICACOES, registros)
    return codigo

def _codigo_verificacao_atual(email):
    registros = carregar_json(ARQUIVO_VERIFICACOES, {})
    if not isinstance(registros, dict):
        return None
    reg = registros.get((email or "").strip().lower())
    if not reg:
        return None
    try:
        if datetime.now() > datetime.fromisoformat(reg.get("expira_em", "")):
            return None
    except Exception:
        pass
    return reg

def _validar_codigo(email, codigo_digitado):
    reg = _codigo_verificacao_atual(email)
    if not reg:
        return False, "Codigo expirado ou inexistente. Solicite um novo."
    registros = carregar_json(ARQUIVO_VERIFICACOES, {})
    alvo = (email or "").strip().lower()
    if int(reg.get("tentativas", 0)) >= 5:
        registros.pop(alvo, None)
        salvar_json(ARQUIVO_VERIFICACOES, registros)
        return False, "Muitas tentativas. Solicite um novo codigo."
    if hmac.compare_digest(str(reg.get("codigo", "")), str(codigo_digitado or "").strip()):
        registros.pop(alvo, None)
        salvar_json(ARQUIVO_VERIFICACOES, registros)
        return True, "Email verificado com sucesso."
    registros[alvo]["tentativas"] = int(reg.get("tentativas", 0)) + 1
    salvar_json(ARQUIVO_VERIFICACOES, registros)
    restantes = 5 - registros[alvo]["tentativas"]
    return False, f"Codigo incorreto. Tentativas restantes: {restantes}."

def _enviar_codigo_verificacao(email):
    """Gera, envia por email e retorna (ok, msg, codigo_visivel)."""
    import html as _html
    codigo = _novo_codigo_verificacao(email)
    corpo = (
        f"<div style='font-family:Arial,sans-serif;max-width:480px;margin:auto;"
        f"border:1px solid #d3d8e0;border-radius:12px;padding:24px'>"
        f"<h2 style='color:#1f538d;margin-top:0'>Exame Inteligente</h2>"
        f"<p>Seu codigo de verificacao para criar a conta:</p>"
        f"<p style='font-size:28px;letter-spacing:6px;font-weight:bold;color:#1f538d'>"
        f"{_html.escape(codigo)}</p>"
        f"<p>O codigo expira em <b>15 minutos</b>.</p>"
        f"<p style='color:#888;font-size:12px'>Se voce nao solicitou este codigo, ignore este email.</p>"
        f"</div>"
    )
    ok = _enviar_email(email, "Seu codigo de verificacao - Exame Inteligente", corpo)
    if ok:
        return True, "Codigo enviado para o seu email.", None
    modo_teste = False
    try:
        modo_teste = bool(st.config.get_option("global.appTest"))
    except Exception:
        modo_teste = False
    if modo_teste:
        return True, (f"[MODO TESTE] Codigo nao enviado (SMTP ausente). "
                      f"Use o codigo {codigo} na tela."), codigo
    return False, "Nao foi possivel enviar o codigo. Verifique o email e tente novamente.", None

# ---- Lembrar da conta (cookie) ----
EI_COOKIE = "ei_usuario"

def garantir_cookie_token():
    conta = conta_atual()
    if not conta:
        return None, None
    if not conta.get("cookie_token"):
        usuarios = carregar_usuarios()
        alvo = str(conta.get("usuario", "")).lower()
        for u in usuarios:
            if str(u.get("usuario", "")).lower() == alvo:
                u["cookie_token"] = os.urandom(16).hex()
                break
        salvar_usuarios(usuarios)
        conta = conta_atual()
    return conta.get("usuario", ""), conta.get("cookie_token", "")

def _autenticar_valor(valor):
    if not isinstance(valor, str) or not valor:
        return False
    valor = unquote(valor)
    if "|" not in valor:
        return False
    usuario, token = valor.split("|", 1)
    conta = conta_atual(usuario)
    if conta and conta.get("cookie_token") and hmac.compare_digest(str(conta.get("cookie_token")), token):
        st.session_state["usuario"] = conta["usuario"]
        return True
    return False

def _autologin_por_cookie():
    if st.session_state.get("deslogado_manual"):
        return False
    try:
        valor = st.context.cookies.get(EI_COOKIE, "")
    except Exception:
        return False
    return _autenticar_valor(valor)

def _autologin_por_querystring():
    if st.session_state.get("deslogado_manual"):
        return False
    try:
        valor = st.query_params.get("ei_auth", "")
        ok = _autenticar_valor(valor)
        if ok:
            st.query_params.pop("ei_auth", None)
        return ok
    except Exception:
        return False

def garantir_usuario_teste():
    if not usuario_existe("_teste_"):
        criar_conta("Usuario de Teste", "_teste_", "teste123")
    return "_teste_"

# ---- Acesso por entidade (sempre na pasta da conta logada) ----
def carregar_banco():
    return carregar_json(caminho_usuario(ARQUIVO_BANCO), [])

def carregar_config():
    padrao = dict(CONFIG_PADRAO)
    conta = conta_atual()
    if conta and conta.get("nome"):
        padrao["professor"] = conta["nome"]
    else:
        padrao["professor"] = "Professor"
    padrao["escola"] = ""
    return carregar_json(caminho_usuario(ARQUIVO_CONFIG), padrao)

def carregar_planos():
    return carregar_json(caminho_usuario(ARQUIVO_PLANOS), {})

def _norm_dia(dia):
    return unicodedata.normalize("NFD", str(dia)).encode("ascii", "ignore").decode("ascii")

def carregar_grade():
    grade = carregar_json(caminho_usuario(ARQUIVO_GRADE), [])
    for item in grade:
        item["dia"] = _norm_dia(item.get("dia", ""))
    return grade

def carregar_anotacoes():
    return carregar_json(caminho_usuario(ARQUIVO_ANOTACOES), [])

def carregar_turmas():
    return carregar_json(caminho_usuario(ARQUIVO_TURMAS), {})

CABECALHOS_ALUNOS = {"nome", "nomes", "aluno", "alunos", "estudante", "estudantes", "professor"}

def normalizar_nome(nome):
    """Padroniza um nome: remove simbolos (so sobram letras e espacos),
    primeira letra de cada palavra em maiuscula e o resto em minuscula."""
    if nome is None:
        return ""
    texto = "".join(c if c.isalpha() or c == " " else " " for c in str(nome))
    palavras = texto.split()
    return " ".join(p[:1].upper() + p[1:].lower() for p in palavras)

def carregar_config_grade():
    try:
        return int(carregar_json(caminho_usuario(ARQUIVO_CONFIG_GRADE), {}).get("max_aulas", 6))
    except Exception:
        return 6

def carregar_avaliacoes():
    return carregar_json(caminho_usuario(ARQUIVO_AVALIACOES), [])

def salvar_banco(dados):
    salvar_json(caminho_usuario(ARQUIVO_BANCO), dados)

def salvar_config(dados):
    salvar_json(caminho_usuario(ARQUIVO_CONFIG), dados)

def salvar_planos(dados):
    salvar_json(caminho_usuario(ARQUIVO_PLANOS), dados)

def salvar_grade(dados):
    for item in dados:
        item["dia"] = _norm_dia(item.get("dia", ""))
    salvar_json(caminho_usuario(ARQUIVO_GRADE), dados)

def salvar_anotacoes(dados):
    salvar_json(caminho_usuario(ARQUIVO_ANOTACOES), dados)

def salvar_turmas(dados):
    salvar_json(caminho_usuario(ARQUIVO_TURMAS), dados)

def carregar_mapeamento():
    return carregar_json(caminho_usuario(ARQUIVO_MAPEAMENTO), {})

def salvar_mapeamento(dados):
    salvar_json(caminho_usuario(ARQUIVO_MAPEAMENTO), dados)

def salvar_config_grade(max_aulas):
    salvar_json(caminho_usuario(ARQUIVO_CONFIG_GRADE), {"max_aulas": max_aulas})

def carregar_horarios_aulas():
    cfg = carregar_config()
    return cfg.get("horarios_aulas") or {}

def salvar_horarios_aulas(dados):
    cfg = carregar_config()
    cfg["horarios_aulas"] = dados
    salvar_config(cfg)

def horario_padrao_aula(num):
    base = datetime.strptime("07:30", "%H:%M")
    inicio = (base + timedelta(minutes=(num - 1) * 60)).time()
    fim = (base + timedelta(minutes=(num - 1) * 60 + 50)).time()
    return inicio.strftime("%H:%M"), fim.strftime("%H:%M")

def hora_para_time(texto):
    try:
        return datetime.strptime(texto, "%H:%M").time()
    except Exception:
        return datetime.strptime("07:30", "%H:%M").time()

def salvar_avaliacoes(dados):
    salvar_json(caminho_usuario(ARQUIVO_AVALIACOES), dados)

def primeiro_nome_professor(config):
    nome = (config.get("professor") or "").strip() or "Professor"
    return nome.split()[0]

MESES_PT = ["", "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
            "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
DIAS_COMPLETOS = ["Segunda-feira", "Terca-feira", "Quarta-feira",
                  "Quinta-feira", "Sexta-feira", "Sabado", "Domingo"]
DIAS_UTEIS = ["Segunda-feira", "Terca-feira", "Quarta-feira",
              "Quinta-feira", "Sexta-feira"]
DIAS_CURTO = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sab", "Dom"]
# -*- coding: utf-8 -*-
_LOGO_ABA_B64_PADRAO = "iVBORw0KGgoAAAANSUhEUgAAAQQAAAD5CAYAAADBc0t5AAEAAElEQVR4nOxdB5wdVdU/9055/W2v6ckmIQkQIHSBANJEiogbRFTEhqJix25YsXz62RFRxIJdVgQUkKJAaNICJJBAQnrbZDfb97Up93y/c+6dty8hQfADTGCvLtl9b97MvJm5557yP/+/gLGxJw8BsFBA+xwBa/4hYfFPAwCBu9nWirV9ZHKpUKq3km5GZJumi1JxOvqFLAImRayqGePxGoGiAYVI0E5QoSsAhbBsBCERAUMhZEGgyiNiHiQMicDbhp5fBCGLEIttk+g9o0b6N4VDIyMOYCGbCtb1rvjF8G6/wfyFNv+7aDkCzEaADvUyXaux8RIM8VLsZGy8xAZg/t0SFjUiQGf43E2aUum2BZNKRZyLiXjWTqYP9kveDJVIZUQsPlFYVi1argDbBWE7AMICJSSAEACWBAEKEASaGy8U/x+loEMj2QQBQghtdegz9Fn6FQEkKlChD+jnQfhBCEKWQAUbrEKuC5XqV6H3L/CKXTJurw4eu+EZgCWDZHd2OP2FKGF5pz58Z7t6HgM3Nv4LY8wg/NfHQppxEmD5rgyA5U68YCama+YqXx4ikplZaDkHoBuvBicdBycG4DiA0uK5q5AWX0WrPM1u+p/iiU4D2QbQnwJRIL1ME5+mK81IUd6MXAZ6XwpUIUpJ9kPQ67yREJYwG5AVkQj0X4fnvQgCQN8D6ecQiyMDIKw1Vn5oMfqlZ6FQuNce7t5a3PqL9Tt+/YUS7gYJi0ABdLDj8gpd+LGxizFmEF75IQDa9bL7XAMQT0+8YGpJiEMxU388xLL7o+XuC4mMhW6SJ75etRUiCJ8mt6AZrKcQvSN5RpvZHlkBmsDaFtALbC0Uz29jExQbAXINJIBSirYXgkyBUuw1CCHpSPy7NC4Dkl2hY0oFkrwMOrpF75ARstGyLEl2DiQIzwPwCwBBvh9zheUy8B8Q+b77vPXL74Ohzr4drkD7tRZAJ0BnJ1m3MePwCo8xg/CKGoFuAbAoqHzDbT5nlnJTJ4BbdQjEEq/DZKoZk/VJIHffIZcflUAIlaLVnie8RCn01NR7Rp72PIkrppBZ8nk+a2vAAQFIXuDNZuwwkAHQTgoFDsjGhqe6wJDDC+0a0PyU2gaUTY2S2hcx+yLDoV0JBIkopI1oWRRaSCltgUJaaLlShgjgDQPmh7plIXefHN5+N27e9A9/65VPj1oAATD/y7bxHMbyDq/QGDMIL+tot3bhCVhW3YJjINt0CsYSx4ETn4exrIRkFtCiKS/Jp/dQoOSlnZdZ7QjoeN7MT2HRIl3297WzwLOR5jeCsCg3wLaA9smmQNLOLV74gcML/QgI8hdMkAGhtjRsJPh1SjvSKbDt0V4K7Y+cBvp2bBzIgaBHiXKeIhB8ynQE/XhJCjFoIwAFwkLyJPiItmvT9wA/ABjoDkVxZIkq5G+xerbeHqz9xv28PQ2yPsdeao0Zh5d/jBmElyUncLes9ATa2k6JrS1UHQEycZqKV78BEtWzMVUHEIvrZ15hwN41KprlPIN48nOsT2E6eQa8K1689erPyTgKAOg1RUsyCGlpy0GfsEGqECAoAQQ+QFgC4Re60S/1C6VKQlp9IMUGDP1uKz/QD0HJU5YbAFA+gE5BqkBnIKWk86HzDD0L7FhMOPG0cpNVwnUmIIgW8Lw4WjIr4ukWtFMueTcilgCwY4DCAmHR6YeAyJFKQCaG6hlk+ci4ICVALUegdCybDNVQL6rB3mWi2HedzG3/s7/0G0/tEFJ0LhurVrxMY8wgvHzegIi3nnOEL2JvwkT2ZEzW7C+SdaA4DBAhT3Z6+AXodXZ0otNntSEg/0BH+maXbAMUSIt+yGBIsKSkqB1CBcIfBpHvK4JSG9ArPSUxHJRSPai2bVgjktXDQVathaU/6YaXYTQ0zE8PtBw5GYcG6lA6KVHdOAMtZ19AmcV4YrJw3JnKdpMiWWuDkwC0yfbR9wpDk5Eg90ZIxwLhxISQMQvJoI10hzA0dL/o3/one8MTNxW7frNhh2pF54JdVGLGxn86xgzC/29QoGsBLOKHml5IjD9+XBjUnRO41e+ARPoATNQCUjXAtij7j8iJf5rItk7T8ZynJZhjAhQclXNWHzn9p0t/oc76oY3CEULaYBE0ID8EoV/YBBishtzIP4QKloe9yx6G2MQ+6Loqv9tTLhcfFkiYP/s/ewYa5yDMXoZw6aW0ypczm7u9TsmTm6F13xqZqjlKuPF9wHGPFIn0eCuRGhemagHJo0CfQw6JISU/EaVN+3YtKy6wVADV2zUkhvuvS25b9eOhFf/7aPn7tP/Jgs4FY0nIl2CMGYT/V5JwNDeQGHf64aFIvztwUu2QaqxGJ0mrIIXROu1Pk55db12p0xl5Dg04P2gKATok0IdQYDmBAHBQWFIEPkhvBMDL9Vil3EMiKC2CwsjDqVjw+C6BQeRady8TPHE7O82L176MdX+KdS4V0L5cUJEA2tuBj09jUccOidRotLWdkt0KUw4M6iYdoRKp+SKROVIl67KYqALpWhRihKhCxdeMyp224wgKaXo3IeT7bpXdG6/0Hv/i38z1GjMML8EYMwj/D0PQ0NCQHobDzgoS2bejmzkBEw3ShAQ+J8K0yy9A2uQICKToQFi88nOGUAgVJQFoogopQ50/tCywbCn9ImCuD8EvPWGXhm+VEu52Nq16bGTkpu07nhYKmH+ppSc/xdeXkve9J5XsTC3iUgHzQfJ5XttOpcodzzFxRqvbdsDrMF1zjEpUnyzSmemYbuBci4AgRMCAip3KcmxpuxJzgyD7tt5vbV97RfHhz5MZCvhSvoU9hrFQ4j8YYwbhxYUGvNIlG+Y3e5h4PzjZd0C6oS2IV1NIEHCJUKHNSz3nACirT+U8/p2rhpwSIFdY/8ZJNTISaNkA0rZk4IMo9AOq8FEx3HOzGNp+S7DtDw/vcDbtJl/BC//eXK9HAe0LJEA7wOx2hA6ubUbDtWd95nCoajlKpLKnYzJ1GNZOEOi6INHzQBFqyqHQKybzQyD6ti7GzRt/6D9xye/LhuHLX5bQMVayfDFjzCC8CI8gOfHoFr9U88HQSb9fJRubwE3RHA+Bkl9UJhQWhwAaHUBQYcss/pagEEFwNYBrhrSNokqAsGwbwAJRGgFZGFgr/KFOOz/UecnG3zzWEZXdeCy0oX05QufL6fbvCQaiUwOfKlZ4ekjj+3/ysCDZ9DaVqT9D1LRMxnQtw7AFqlJIMZgdd2UhB2L7uodF/5ZveQ9dch1/eKwq8aLGmEHY7aBV2BiC5CHNpUT9ByBedSEmG5qVk6BL52v0LpUL2AugRgFd0NfQAXYQzHv6dy4pWkpISWVCB0IBVqHbE37uZhnkfl2X7769q+um/I7n8JptCBLQTsa4HaDznDByglqgJdk7+/wToHbCeVjdcArWTcgqiyoulGSRKG3XxeIIyJ5N98r1Ty8sLbvsrnKT1W5yGWNjdIwZhF0bAnbDqybNr87l0h8BN/mhMNXQhDbhBqAEqMi/l4T6MzEAJw21QaCrqrFE2jDwewTGUSBdeq5t6ecBC0ObwcvfiOhdDet/8Xj58PzgjuH6nzPKPQ+jkzre/O5JavK+b1apuvNVzYS5oqoKQIYlMrhgxRJisCfErav+YK14oqO06cerxsKIfz/GDMKuwgMhwKk+/v1BrO6zmKqfwjBiAJ+BQxFckI0BTXYDGOaXefLragH3HdDHpALhkF2woJQDOdz9rPBzVySt/t8Mbbq9bycMw16cD3jFhoD2a3cIK+YD2P864BtnWXWNHwvqxh8ZZhvIWQiQy6G2Az2btouNK7/xukc++sNFlF/gMGKsTLmrMWYQdgoP4jXzjgpF7aVBuvX1GM8SbNYHDDkBUG4HjtoD2QhEjcQcGmhvgBA31PBjOT4ojIniCIhc91OgCt9vUKv+uG3b0tzocV+zIcFLMBZKrlrc0xFEvpm775dODeonvh9qWt8I9RMsIUMPUdjS9yX2brrfXbf0k4WlCx/i2/YWNgxj1YiK8Vo3CJp/ADpUctKpzd5w/jKM1b5XJVsAJXiAAYUGGsSjGQS0UdAwYX319N8RaF+HCdKhh8wWhUGAXM+zQhW+PSEx9Ov16xcVdw5L/qvf/tU0eNUfzTXYsy85VLXudwnUNp2N1c2Uvy0KkHHV3+1j16pv491f+hrAttyYt7DjeA0bhFGvwKk+/J2h2/A/mKhrwXgypN5CNgZlj0DjiXiy8zC/G+MgJCEQgXgBQiEtSwS+lMNbNlvF3q+X3L5fQddikyicb1eiGsfGy3Fb2y2YfW25hBk76LLjVc2Ez8qWKSeGqRQI2/VEKXBh28on1erHLwqWf/O+8uc6d0VI89oar1GDwBMzyLTuU1coNl2hEs3nqESW3tBeQZQP4KqfJgsxfQbcA8w0IppDiBv5KJOFVgxkGDgy1zViFXqvyqSsb2/fcFuXPt6YR/BfSULCpUCGge5c7MCvvDVoafsSNLfNlsm0D6Ac3L45lJtXf7v63p9ctg2W5mD+XTYsOu41XYl4rRmEaIlXsZpDTlFW9RVBqmUqum4Aoa8ThIQZ2GH9Jk8Ay56CJg/CijDB9onYSFJ4UOi50ZIDX/C33b9Mf3bMI9gjQoky6KkpZR/5xY9B08RLZNOErLLj7C1Ym5Yslpufel9+6Tcf56apDsZ5vCa9uNeQQcAyMtCqP/Ybwk5/MkzU24hQAAgpV8BMIho0YCgFCWMgND2YSRgQrkgbDI0+tKzQt0R++xqR7/u83//PP+ljjRmCPdIw/HkBMb5AZsp7Z3gzj/qKqplwDtY0EYIcVNemXHzrhs/lHvjA5ewN4pflazHZK15L3zM1fv4cNdR9pZ8Yd5Ry00wSUqYbjfKGmglIUwxxZ7LxDghiXOYnpPSBpauLXv7nMZn7zPCWf/aOVQ32kpKlqSy4B//v6dA06WvYPHU/jKd8GB525KZl13h3XvB+Dh91efI1lVd4LRkETE48bp7Xu/kY4dh9EIYCfAIM8sIfQoiCy4a0raJigAESKSL14PdkCAqtECXYsUAixu1Eck2u97E7d05Sjo09fSyUTG3Pk318Inb0F74atkz5GNa1CkmhY8/af8rVj7+vtPRra19reYVX0CDwhDGsfwtf4fhsufmeL8eEnecApBHg2L3EveyIftlLzvflDiPOISoniB36jeOD+rYrROu0fSDuAm5Zv1msefS84ImFi15rRuHlHhW+9tjYAxOsr/UheMIDQBUcVWMfc/Wv5FsXofzQSrTPumU4fswP38lbtSNTVMGrfLzcX5AudEBHsZIzjxVCHB6Ceh0EXpq7fXSfPLPr6X5hIt3jPmF23CMGQU3HQ4k+JYmBXMf2oU7uWVKhUkwFqlUFNO041wypt4hIwpkH1EKhQoYdIxF90ruGS5SYwwWgRLA4PmAKY6T2e0YXEAYZ6Tj6rCJqY2ps4keEWh01nRkzlBKVOSsWGApBFkGh/ShmKpKkd2C6FfkMJApFDEZKIlGrMdQZLEG8IBolrYlSBTGjagI1Qe/SNQvp3FAhoSIVWIIok5l3CYnFlFqviD+AlRXo4qEllBQipuzMV2Bo6W1R+fVlfgb2jlGRL3AO/sbbg3H7f1eOn9EgR7YBrF3yVf+eD32pTALzKk42vlwGISIHVMmq/Q8shvhtdNPHoxujCl3U/FM+fGQbkJmAzeuaQzTaDf+pA/5RtmC9D/0Z/T6xCUcAIppQNFErPsNzq3Ib035gOMrLZ25siz6HCHZgUIlaA828rkGLO3wXA2dmZmM+h+g4bL12YDvW5xF9fnTuR+/pl2k+69/5+/NnyAaZ4/P+zD6ic9bZ0VEAFX8fT/9NpzTcVZDh4Flhft2YUXhO+zVI6BQhTL9kljXt0F/ihOmH0b2zVz/+0yP+ccGHuReCiXRfnUbh5TAIRiFIKDc183OBnfqCStWnwLbJ+iqBitZOszLrma+noqqo/0cPdKQopkuGzERGWxO7oHlN64LQlI74CiuqhLwH/YI+YPkNZjM0kyyaQdHboyVHfs0c2hQljfAJn4HxRka/uP67LJNm1I6Q3Ar9u96D6YaKzkTbrOh7l41ihaHQ3wOR3RPztXnyj8otjd7KipKpoWnU1oiuLxsoclNsMdgdqkLvGeCt+XvZkxsbUNkq3QT7p7qPuvgqOePgt0k7DtampbcV/9V5DvR3Dr5akY0vtUHQZAACQis155eYrHuXSmbICISIxCQUzbeKyWl+1X/pKUWTLDITenE2+gDlzxMPYfSKnpZaw8R0I/Nsq5jFfGCaD9GspihEH06/Iiq2ML2M5OLrSWX4v0YlScqnX3EF9Qlpt8bstTxjmU3dxD288hOGgeYnf1O9nTmhyFaxOSlrLDIgShuB8nXSFiLybcr+C1Z8N+01mU/QdnxaSghh+YiWFMNbfOn1vWnMU9jFaKcJ/2fujXCP/PHCYMLMhXbNOIHrn3jAv+9nZ8AwlZlffZ7CS5xYaidQT2ilZ/5YpRvepZK1JAZKT6Jp7+XKPZglm51i1hnUUECtRsjv8kNu/oiUR0YNR+VENyag/Jt5TU9sngbMX2hIgVnr2GxpJl40VUdtT2Q6WKtM7y/SMTOKiXpbPnWO2GnCjpIhaGuk/zZWTX9X/faouJF+ywgraYdDfyayaGWzYyRYjMJSxWlFV4JPK7pG5vpGuIqySTV2z9yPECHTHEO39kYrNukUk0vQSs1jA/TqrwQhF70HLuqAZx95a9i/aRCnHHKkddR7bwI4oQpEh9IQ6VfPeCk9BM70OZmZ71Lx+l+GyWpPqJCJBPhZjhZYfmB5vpkln9BBxCAUJRFNGpG9Y+PqsoiZCRi0ApCGEOvZwOk8vbYTs7HktGN5lTQnoAMAk5hg9cLo42XdA0pBUNuyNhGoyNOJPq3/ZUkiGqyZoC2KzmBGc5GTFHxSkc3SEYmJEjigMLbJGDHGQGlzYAyV+Va0jaXPU093Ojf9p1Zs0DZT2w59RKItLX9hvjCkrkKsThYqnRjRzlf5DCiJKcRwT6C8rjdDfvPNY+HDLoYpO8aP/N4RQdP030PrnMly7SP3ef/qPA36rh2KDC28CoZ4Sfczcb9qawCXq+rJjQBBNAPL6QL9VGvPVgfxFoUTDgaEKA051tBuMXnrJjlHOqLRpN0hxwCjXoFxJ/QSWNmRWLF5pYsR7UvXH8pBB9EijlqQnfZREeyUc54VXdFlhUOeqdHnK5J7rHtWmejTndL6pSgUqUwG0qdpwVZGsq3iO5tEa/m8K78nf7fIizWFHJtDLC5UcEgRGR4d2SgAV8jhzb7Ibzk79LaNGYXnMQrZtvdPy816w3X2pHlzxcaH/lq8sf0saEcBndz4ttcbhZfKRaSnPbB78xeozPhmEKIEyJpgJvet64r6USQxUdIrIAahogWFvq0S8G6R638KHenSlKQamgTB56aA0vX0KDuGgcCjhiLCEpNQsdL6wrQVopSW7nklXXS2C1Qr5H3xfxh1qF8wM4X2bRMVGn0m5F8kSgio5kdyAJZFm5tJoz+qFEpbSqU4U0i+Bn8ZUiDSEkrErEruBaGd6fhUhgxDKcGlVxQdh0WWo5IH7RRtpQR9hkh+bBJ2NSSsOubXiRmDomSlI6of8jWgHUkLXa68EuLSaMLxXVGyELrZ40W24SjQEqvWTtaThgCCbGZaHED1FwtiZ4fehpvGPIWdBgGT5i+0hxZ1rE4N95xUQPt6OemwM9xTf/dDr1N8mHEKVJ3Yy8dLZRDC9muvta5771fa0U6SYqilA1ctNxAlw3QATT5saIuhrQURhv8TpOp+DFsXbecNAvZh9Q53dZDdvPZ8n3m+1/+T7fh4pX//+crz2tV7u933f3Bu4c6JzoohxMwsBrVHg2PSOJECa5Rr5N+Vll5KN9kY4nWWJReEhXU3UkmeqePGhh7E59h+rZXrXNANT058oxT4FzX5dR+yjv3VcNgpPvdqQDS+FCEDP4rpdHNDXjSsVdWTUpywGgUHmK14nVMSQ2ENrN6Ghd43BUHfQ+ZdvdCPjZdycCnRSu97dZhpOh8cJ2THpbJOyqNcOUFgJXgpxOCWUAYDbwkLG8kojJUknzOi6kJdxn3Tr29U9dOPgydv+GTw0CXfhXk/dWDxhXutEX0pJiE/YIENSbTcKMGl39FxNAXLJklvIYx0+3aw7UxtDGa7owEx18HHfl7ya0BhhEZEjhZnyvmICMmkIyNO4qgQq8ZJZVd1Wu6EM8x+tCT02DDDVBdE33Dmhneeib3P3CNmHPcd64DLTmBjQDiGvXS8ZKtyzHUtDg90pKurbrr6ZkiJrUD6OaIX+26hMPywbgpabuBzY+NlG5R+0dVVwwjNAAWduYycBapNMEk0J1strgJVj7NUvOo6KzHpTSZs2Gsf8pdldHQo4kzoFX3D4aKvno65nrtF25E3OHM/f0AUWsBeOF6ym6xU0ua+AS4HUn1Nuwo6IU8pL2GJYm9fIu1+Z2SYXl4cvKAW1Ugw9OUcpDXIxzkWoLHH/L677Rr0e7yd+R3uBv5s9DqJi3TfLaDxWGS9NdqOXwf9+3P2W7Ev+gx9tjwq9r3Lc+pB6GzfrZoTpS1ZS1KXcHXtcbQkyxVMk0gw/RV8CzWSqXqipQY2XmvBhHYTPozlFHYYHQq+vJDk4oaCv5x6mvvmv90D0w67pW7TkTN7Z7fn9kbg0ksx2bhZIN08Z1auGHscq8e5QgWUeOfaGlfXpBWCChw5sP4vwfDTZ+/UlPDcQe7YmCbfixzPefh0DiE182pVM+E9KGOkK2EWALIDGgmlwQkVqYRRUCMZDQVgCTm0GUVpZEFYXPMXEpoZgznvPK61QJwTplNvrPdO/shiR217NHf9O8/eG5OML6EbSN6/ozXNdeHc4PdMoUv5oAL/ERPHPo9B0MbA3eezM1TduBNRFVuU7QYCuCyoC5Ga8ZiE1FGRcKoWUeIaIG9DbX+Wo40SVwAFSps0GCOwUOTNmVWVSE9ozxZBikPqP0RFsmDUFVk+Hv3NQIUIC2io1qgli9Zf1m5TYAtFX51rmxoGRLrP3OskCVxBsbyuDHKDJdUupQpZ2003LQHKQH8GqOpJ1UwMhZKO4uNTRdRSBDRyqfaJfrApWPfkX2Btx7ayQ1Y5FB2BT1N3Nmlktf4GERwKKqENhPDiRk7+HiACxEwrgNr8ewumvTUsrL5Bh3uL99rE2Us/FoSk8TDSuWB7uv+Nb/Abpy2LnfqzC0u3HPfTvY116SUzCLYVc3iSaE+UIlIDp9EIPBkEIIMc9xr+O88g8bpvXehlJ31fZZviQJynBgrIu9biqfwA81Vm9WTtDHOPsZZa19vRe6azMYwAzoYw1TQJ6Hm9AxiaG6vLDY+8lYEARH0BpjdAwyXNFNSlP5NTNfsbrfDp1ymJz/sqT1vT1qVRhbomyFAJM7WjBK2BDmgAggFjWQJCWsCDAKxMzedCC0+BVWL5c7wraUmGH1CWoAzNjJAhUddXGZug0VoaUsZt38RCSSEfZsa5YmjLny239azQW/y3serDToMm/fy77JE7j1suT/rrW53s1D8mj/rGo/nOBYsNceteETq8pIkiAweOklUMMdR+Avf4UwLbVPDnU1F3xw+bBzl5wGfm+rVTr1BV9RIg8LgMZtoVuT/REsqghg2AMaTJZIDIdDjTWs3JM9M/TCwK2hzwCi9EWOm+0Gd00wOBmitbDPnl0TaFCFPMAGKNX+YEnaDjmy4B/aaOwXXrsvbK2cBok6GXa2Z4NrVZ01TB2X7tG0ToRt3ERMck5BV9KXY7mFGBzIsUtghw6pwJstT/C7UKjoEOrgrsEAoag1euI3DMoM01HU3fm8gwMNFk+RqMYiglhljdKnAI/2wBnB16Wwi8NJZTeA546S5b3X7cn9QZtxykmvb9NQiYA8uZKWw3SJFXaZVBiIAmUwSLNQ9ZtJZGK+bz2B8S8qR0dlXLOzHbYAkVlFChA0z2gTb/S1NUgcU7EmgTbQk51hQgkHiaQUya7XlFtLW3gvR+9CNRceelLWhC8LYcW9PstQUoW9Df5rhCoEWrJB2LXudUPaJNNVQJyiJoIu2P2FX49WgbijH4nJEEXyQSgQntQzMw6HPjzyjeF/UbCASbPiOAvht9ls4LNSyTkJ28HaVrmeGVJqMFKnQhKAVQ23pgbNL5rTqPQGpUemhEJI2oy9z0WZX7ucqw7lHMtL6hEeZaN34AYRhCxKrxNiTq/hyLtZxsqg9jJcnKsei4EBbeZft/PfUzwve7Uif/9DL2HiI9yj18vGQn6flUTCAnlHxe47/zY8aEJLrF7wU4TYgiQwBhpEnD84aeTKm9A1rXEGmiCcGhMb1mURROVAj0L7cVcs8Rv8ihOrMj0d+KP649Db0rnmjaKybnIfqsUW1me6Hrp9F2RtWV5zUt58zrTtsxfFlYhLim06DPWkjsSIRsps9ImsuW7uQSll7n+XvZ9HmhpAWEiabt0LIALUfy58neUL807VvydhLpx7IBbJu2ZeVpDENVGt723AQWfXPKO2gWJR3XKKQT1m2TdG9UZVuoaQ3VfFOsYc+NXvo9G5WPYabV9eKNN1hu86ljRuE5A6Hjbs6/1PY9fW4QyDfHTvp1GxuFvaAz8qU9wcgVGO1TLg+mF9ANT7Dbkh59JD/8N/DyEtwYzTEfhKSfgH9k9GPRjqgbKuAffl8EQgp636ffQVj6dUuGIC2zHxGioL/NNpLSd5YPUgQoRYiWxfvW++N/fQBJ+/WRPmNJn/bFvws6tgyIW0Dw8cyx6fzA8gX/bt6TViD4u0Coz1HQ+/q78H5EdK4+WJaP5nfmLeD9iXD0XGWAlk3b0d8B2I4HsZQlh3sfh75bt5pqQ/k6c4aF77KhdiqnNbi/TCcZy3nIyhs22gwatVCbYEY3UGRbXRWvv8FyG88cMwo7jw4FCzrl1vu+0CNE+EEMh7/ABnf5nJe/hL6n5BAoZABwjDta2X+nWUD0c/Q8OzAWNOz44s0i9aPfQ0Pb20QsCUiUa+URJflpgY66B8mJ1ojcaAE0CTND1RYl9czvOuPOCcdKr7l8ztwzZPozy8nMchxe/lYVjZaaAq4M0dY5DL21ztnpSxIlJCMehIp9RUlL481rKielMx5RQaYi/IoAHgxCxNARm1ZvlFvWXRRqVGHUQ8lD003qQ5bP2iQ8RotAO1MvmebyHcPeysZVg2icYCtBNTdrQeh1jcGcn5NkXGgX7/jAPYn535yXOva75+UWffK3ezpd/0uGQ8jUzZw5EsSXQPVEF1VQZjbgB8u2AlkYcLF75SeV6v7u82SoI34wK3bA5y/wknWHoBNjJ5dDWII+c96QfGi9UFFrL5RpiYRQ3EpoTkxq9jKa4ErauiuSOyooWCC9noqvwByIFu3PMBJQrjB6PWISoEbGUFiUxaDSJFl929LEqkHI4T1JO/Ak5FFhATl5IBACZFQgz2xqcSTpON2UyfUNLgqE1O6o6GtqzlZt8nQLpwoph0npTiXCQEr018r1T/0yv/zbW3cqOxocwoxfqOyEC9CJByRrrz/J36jcA82nZ9KgZR6WCr6pslnYkcaK0jE6ozu0BW0Yfqs3uPbPY9WHXQjDPH6vXVXj/GCw1P1lWPrbHvMevqo9BKQyOS1YOsGvAXC00IR6nuqkuu7rfb7dmOsUlp74+tUAcPWurhy+iM7EXf39Qs0z/qedh//Be7va5oV0co4OChWeW9rStUO+L8zmrAschpRBU1tyVSHiqipzuhnSFuNwRIniCIxujqPIQodY1SKCYfEHt2p86A1uun7MKJQHQucyBLi8BHPf9pNx1dUHbwa4xVTUXt0GoSKFoB1gw/rFRcNo0eLl80VY1l3BfMfGc8eiYwnRtMsrG1LMYVZ8fY8iVyCKAOh1U5GkNGxEwcQ5hnJwot3AKHYq36UoJRGGmG22wmH1JysZnBnmt44Rt5aHxoQMLvn9Ewecu3DTZrqCezAK96XLIQS5EDARcftFhGWGHZ2A9OSqRjnMXeAQdhy4N6G79uiBUnHUQ+ETDeZwi6io9StR6SFiczSbGc63iGw2CgC5hGMyM2xBmNeFeOdVptmCwP+zlbTaw/zmW8Y8hR2GWPSHju3w2qkyxHQJyxCj6Fp8hI9lwJuJ4MfGKzkE8ypGoKyIHla/ZTKZGvdFLoG0FFguMTVRkoWiDRuFdJBfYGwHqWRb9DII2/Bgch2X9m0x3WNVS0LF626wMk2nmTzRWJekHntkiLDzeOluluEOLzMfG1QvkSIyasDQj2mT8Lzewdh4CQdCSA0mBqlIPRJlPmcd1Uk71KTzaIFXBOH3gVBeiH5hFTUyAAoPUYUGKBWDWKwBpVuLTgaAEr7SJop9Yo2zyPoIgo5mGi3My+ssW54R9nfd9vy9K2PjVWkQPD+gB2NUNUgzp+m6GcMU9woD+eobBBPTzRIVnWbcDBEQlyP4ni28YQB/ZBkEwSIZjNwisy1bJjVNXr561a2lSoQCpRcS4w4e7/eunQHx1hNVDl4vbPtgSFQB2i63gCAEhJTyIVMv1FDxOqtOnRv2MnHrmFF4LRmEGJTQG1VW0HltWoq4zY+QfVwbGEsSvmJD52ksUIqU7g2YglolueNJ+AVb5LYVBRb/BsL+5fjaQ+9av/6aIlc2tm+GVTtEu1EPB+3l0Y0AsBFKvf8kMRP3H4+dEA4PfEzZiVMwUQ0iHkOhAlfkBylN6SDCZTXzpt7Tv3jN0N6C538tj5eu7IhEmByhfHSVnb0D0+gUdTOMjVd4EJc7I6BtH1AmRGkEoDTQB8r/ve2UrvT6Vy2nKbp++OnoeUCAhajRjvTrpQKWL9eGvHO2ub8d5u9O6la7DYS4zU03vjEcHv465J1pAtRDEoO/u0nxUGbeIYu7brqpMGYI9o7xEiIVvXJrffTSqDaB/pfIAl6q442NFzZUSK48ClEcTGCxr0cEhctjtRN/Udh852bPcCQBtAuATsNruVDCfJBwN4bs4XXsdFcpV3QsSFh0KZU6NbcFInrD226eNP/8f65fdE2TELCesOXeIMDITeRQmAdgIYj5JpG9aDlgOxFLLQPROAews113VIzd11cJUjFW1TjVU03LsXayC4oeNUvTqTGA0ApkccgJu5/5FKju74yVo16RwUjFmDPxZhmzXq+s+PfStrq8t3fFFvN+hP0elcZeSDXyUTxDdvZ7asPa8ePRLzlgYyCGB7tyS3/SXT7CjuQfrNxlfo9IcGgoraj877FV7deCNXsZYEfHmCu59+cQ0JUeU26Vn4VI/jDSSQZb+bhX8Unt3YMnth13/2DHa74wuP2RJ0yG0H6OBARPbBGSN+DO+MQ+QWPzOZipO2k4lpgkLKsOpXQENVshDljTj15tFfpvjPX84xfDnQt6K4wC87dUsGjr/S8ECR30+0K5z/m/PDBWHRyfqMFGKWUDqrBf+WJbTCaWrHqy+HDngi29+nTAGvMY9nYPIdswrYRNT0Lt1IQIPYUEwjepbZSWJ4sDMdj61KcU9P9/PAShe/1NDDvaUznmau72ekXXhmnPgudcKzOhU/t+rqnU3PZ1lcqei7VNCeG6nAbi8jHdS0tSF7YlpAMyDAD7Nm7E7pXvD2+/5Nbd0IQZjwPUpPbGc6sn2p9KNuL+1Y22bceoShmA41IrCBU4FHgjois/hDcV+8KfPPi/2x/jU2sHq/MFeBZjY09tfzbcRoZyzHCJcHZaUwpJ5z/MKlKHWOTiMuwzWpEiPT0BMJ+MzFiOYsdB18ZiaXPNgbijMSBqr84FoXvUd9oLE/Z9NBw/892qsTWGtixh6AfUyaULlwoxIBlODFToB4ElvLBx6gSccOBfU8dfcqImAOF7NDrmg0XGoOUtDV+p38/5fcNseVCmWUqwQ0+FyiN8QxiAF4boOWkIaiZgS8tsfF/TPvjw67/dcPXUt09tJGNA3sJ/9syMjf+uLkNEWRqJjmvCnTJNinYkX7Q9MPEHt4uG0A6W2zx7tm2PO9y2xx/i1EzdNzv78Fr9oLOcOVGgjz1AO45QS5vvNGhV7xBKHvatL/mts69VrZPHoxX6bAMEumhJC6J2S+pBsZikxqJeTUItYljyVO0EpzjhsKur5p5ZDZ3XMoOr3jdYsAiChnNbLm06JP6lqol2iQAK6IPAQBALFvW0O6jAViE4yhNWKQ8KfQhqm2ycsp98z6z5Qw8e+PnmgzsXQDh/4RjacS+kUCtzGI7yFBhKgNGF+8UQ9baXk17x6rb5ds28X1h3HL4kVLUPYe3kf2HNxIeU1fRIcbt4wq0/5EG3du4XkhOPbtHGg5lpdvAWFi4EaR6sMS/CuPjxo390nphy0FdEbbUvBQnSSlsQExPhRtgQELsbM0xpwlomtDUGwnYcUCUPG2dPLM5dcA4vB/Pv0p5cJ4SNbxi/X81E+wuZBtvHEmq6J8Ix6sdOM23yU8Nd4yCJHYoYX3wh0BeliVOdKdP3CW85+euTjljUAcGYp7CXJRWLWCSqrfJkY+o0Q+NLDTHUYaMjxhc0mESipmX2xGE//d2SjJ8tYhkAN0HyyMRwVNLE6yBDVRwPiBPQKx4WlkY+6DQecmmwveNqg5EWCymOvZSzDlpR+TU8Fi5cKFtbT7dmzBjG98o/TNhYjP/AqkkSQ4NEl9VgMQw56WOWClLD1kQO5l5WcLazzZXKjWMYa2oHgJ/CsccqaARB2jR2q78w0xKzFQUGAolIggkxjTyU5mHQbbHlXbNXKQWxy7thUfoNrbLBSo78fcHVk8+4dsGGe97/0/c6V1141Rj9+96QVMxmG6YNq4YnsbYtAUhlx4pdE9VYYcDB7hWfBNj+fAQpxjPoDO36WccoTP8OM+PGg+X4YAkSe4kLxtuHBhYtAV2bFNiJB74glUpgPgey2PvbqqPWXdR3U+9wRMBwwIWHzm2evGleWIzfdEfHmu7RPuBX80AB7Z2SlKjw7kuJk2LUIE7/wS+geeIFkHBKYMVjEJcIaQcg44BlWcRxI5A9hQrPL6ohRJ4CGVgnLu3C9q0Nq3+5f9dNVzG+sWr+pMm1hwZPNbfZCRmE5HQIx0bKSyLZF8tinmxGqpAUBnHCxG3yEgLIhcOiFIagMKTXfTcBTnEEurcNwRsWXwiPIZJE2oJRHqixsWd6CKWIKy3iC9P8Y6xGwC1OvMaQVsrzDm0MqveZj5i5DbMTYiihKJUXlyMDDpSGl1uidJs/OLSEELmWY8Wwqn4+QuwNyq2qR8dWMpEM0Iq/ve+O+FRU7zzxqC9dsV8sWfpEbcvW07MNmFj+j8y7AeCX7Z3tsnMPprL6fwzNJdHJv4ZR/V+IDvjun5dO3a+tefbf7199msqFb581rSmU0rZKnsBntxdhybPDsGRTDvvJZ2hNEc8rhKEhTdA3UHsKWoeGUCb6v8Kq7i1W1QAAGwSnEY5IVlkpCMEDKZkDD4ndikMG+gBxTfEDIixhoy092Di0HRTWwaS6w9TU6vFyfM0+6IKQPcPPBlvl6sY51dV3fXHlMR8UYsHvaX8UQnQ3gGjsGQM17UkGIVozAEpeAG7UWxtR8+nuZ8pD0UsGpPC8nkY83johgORvSBIOJJSsUjEOw12P28H2745vcjtXrV5VitYGRc7j9g0/T8yc0Iq9g+9DyH4q9BvSYSHl18yAI/c77fIl9VP6JlbVWy6JHq1anHmsa0XtnfR4d2pk3qtpCACi+l4QRiXAqknnV+cmHrB/Yrg05R9/eveb9xmfPS6djGWOn9tAF3uX5dr1W0bE72/bCt+9eYvYngjBakpgSKJcukxkhHypOcKknwzZHHiUVtbDTYdzY3HKP2jxiOhxIFSKpmYgsSyFCcuGodIQDIeAZ0z/lGif9X5sq5tGu1A7PV8BIGQLYe53z/Qce+zPHnvbF79z8rruyghwDNT03zcIUY1bCyiVBn1wx1N/PMeCWn3ECL5S9zOHpcSJ8Lz7U4FbcyXWTZkAll2Q+YGEm+v6VfZg+PC2OzbkVg3zdpYpM5o26nYsrOjcIizoQPeUnydbh/+y/4LN86Yd1u1ZINpiKS8c6YuFqx6u+tHjP85+CeAu2surLFyIdB0XhC1wWnL7/gecgU76jTk7fnwwAK0nHz0LDp3RwHdLKfAprAclbFundEQYBhD6RHwvcVJrGj53QRsuOLlFXPA/z8C9m3LCmpDCkHHORgpOhwtG2wqR5eYG8qNaELKQFE5Gp5mZ/5YIVsoUrtR8KVxhwYbBbao2dgBcfeovYU7TvvRRLAYFNj0OcWAyES5CQNyRlggs6YiZ9Ye9b7/a951+yPc/90BbTeOqvKcWDWwSj3Uu6NlaBjUtGMMuvNIGgSd61X5vnhoWrOGRVZ09CHEKDwUFiSx2WPYG9EKhjcRuadgZ9hqrmnG8n2p+I1hW0SoMJ6yRTTcUB5ddULgDAGbPdmH5cvIJjJsfcSp0iugheN2nVh3YMKNQk0yE4A24UjT53ppH6p31D9ZdvOH2+39UPvdXlTHgmDqcP3++ff+Ww97Tnan7hKpqmgHJDPWXAGweDC9+z8GMIiwFAWk7y7hjCw8U9vcWmAIpVe1CKqX1Vjw/AD9EMa01BXf88AB4x1eWY+fTw2BNSIM2CjSi2ytIWo6IZRFio+zYiCLGDoTmxLFJMkPLXFEvNoqY5cDW4W04PnuM+N1b/opJNyXyfp6NQNxOsNkYKg1izsuBZUnRmGzm1zwsUQjqn9D29uYbn/3xmxvH5aE0LC+pr1dd9T9vvL5rbfjDzgW9Kwx59Y50b2PjZTMI/CTUTD0hm9u4+muxZP2H+cVEMVJu0h5lJHLItAgShCL/gPPMuxqaWdVKXQiJGsBA2ZDv3hzLDr+7OAAC5s2zYfFirwIKW3YpqYrQsWAhnvql334lO9X7EqHqVE4qOwn2E3+o85++LisstDZoozFfGLzCzkMCzN+pBNuIAN27CHF29Xr02u7eo1H5evRa5eis6Ct4oYMSsAtCe9pHj7i3r/kHOKH1EEwk6QZ5UioIB0pi5oSMPGB2neUHJO0gwbGk+MUNz8KV16/B1f2KGaLrm+J4+MS0OPuoJjjr9RPAdQALxUC4lo2/+fwsse2SJ+GevgBltQuKTQuXC0wFkWjUBHpcotSDcrxay5LVcY1wpMaiSLBEPsxjzJ6Kv3jTX0TSTUHez2PSSfL+/vL0H+DOtX/FdQNPAGI/6+HOrD1KLDzmSsgk6rEU5sW46vHBfnWnhKsHfiVa0w0QZP3mqka8KJkN35H+UVWHEIPfiUrNY30RL7tBaGegUG5w+LJAJsd7W/7ZyytFARVkNA0CSReNKidS9VlR6REQqEV616FCVdWk6hG0jkISJyoWbFD5Xw1u2NDPCMTFi/zUvsc35Z7KD0HHgwVjFLD9WpAdC0T4us//4aqqWcH7hAg9JNKPmAwfumpcfu09jXGnkaoO3ed2LoC/AizaVdnTEHcs2hNyCi+cL2D+QhsWdQT21I9fGFZPuhzrWh2wlA+Bb6EEhyu8QQizpozDTDwmRgohpBMS/nn/ZnjPFx4QMKcJoSrGmILBEuDqJYP4u3u7xUk3bYEffHw27DMxI4qlUMRjNvz8E9PxwI8vlbkaV+tnspKVpr6n31WIDgT58rUNJYSk0mGRjB1JaZsiI3kjtg2wpndI/Oj0a6AmWYV5Ly+TbhLW9a/Hi295l1w3cJ9oSjuYiSWF4wiMSQGPbPoLdtyVxu+deo3wDCb+wIlHWYuf/Dk0pDwoFZCKTUFDo5Wpa8ZvV/2q+vglDyTO7+jo2j5mFF5Wg8CxaphtOLCtFG/8kLTsX6u+Mqcvo0zMrY9yjZF7bpyHcFdZfbrB4cjI8CFQ29oKwvLAz0lb4N983smicPpZc97aOHvdlXY79G1dOfszKzqW//mUj7TFOhesKh3++amXjZ9dep8FmJe2iA302pjrSZ2x9r70+U6z/1YlY6AKibnzz58fX3TNouJOk46NUaL5gENCLzySiuCWLQsUlCiLpOBJWo0IzBXxCfAjTjyxDkF5iTIsDIWyYyHh70ijkcA1lEANQ9aKBEtYfsi8x7Re2kprSEqwSIGJBHClkmEYCouY5fzhx/PbFz/2Au+DNgaTPv4x1bzP9zCbVRAWAwikRaqX5YoA3ZCYEAQQJf51+r6/uH4liNYMOtkY+MRoQ5vGLZDpmMBxGXH71gIe+eGHxE3/cxAcObsW8qUQ28an4eOnj8PLbu8WdmsKiBuLACFac5srDixMW3FhWW7aqL1ESg9gCxu25/phXvPJcFLbqaIYFCHpJnHN9rXinD+dAOnUFty/uRWo7EiVCEp4eKBgZl2L+NeW68Wq/s9gWw1zMqiJVTMsRcR8xL8jUIRK2MoDlYjJoG2OPDURL905c9+qMzsuHlw7ZhReNoOgiTIKIvM/IjNOiuH166NnoKCLCJWFhMgqRArQz7tnGUu1BNLhyjZ6gz0J2PZ0Qcs8Q7o+95mWfUS146jqdHXu943jpmy49VurHj7uc/suqJ0+/EVLqBKJnw52O9aaRxIfeOpXy291xzv7h0HjW8GJBRiKiYv+uLUZANZVGATOW8TrDzjXK4jfYLyBILngG3pyDnvokS6nPYwgUkDqJ0ZXlkStifolguuwj8FyBzwFPLMPZjNkBSrzNz3wLOuuq7M0j0GJwKqad1I4uPiu0STh7sKEjkBO+si7VeO072Eq5UHg2WyJmDrdXGejvKQKikWqqD7gQwgPrukHzDgQFJWWwOXzJd4ERVVhdBri0B+XcMH/LBMPX3kEJmP8iIhzj6+Dr926jaoDpvSo5Ru0mgMRucYq+dlNuyvrahL9KjPoERJ6MB/AgmPeyW9TxrEUevDxv78TUvEtoindCMO+p9MPpEttOmYDhcJxCzhQGuRShwJleaHGJ2l2Ps01T7et4KPtoihNmCr3AwF/bTuv7XUdHauGdgo1x8ZLYBAYI2C5494QuvVno10HEowIDWMPSCC4rDRskAem5ESSRSa3uLudy1h6KgibZhWZfb9voM/XKu0A2WanT9pBWMqLUk2rSnq50hUHvGvKRZlJQ5fHEqhowSoWHGfDU/Lip361/KcLEeT/TrGf8EuUzibxZIwD5Hb+rvz8eiq8CLOTLHRTwxB6TiQioQnDjD4B18pMyU2vqYSiUUZWRqP5uMc70jtgnVm+MrrBS4P+NKUcHYBoB/VHtcR93AMrkRLbez8GAHdFhve5QxsKZ8JFc4J08+WYTYcQFm0uD+j5Y1iUQaEXCJF1xOJl23DF2j4xc0ot/O5vK2FtrwdyUkq3lZT5LyMdORB+AdHOJnHl+mH48V/Ww+feMU34PorJTXHct8mBpYUQZMIWmqjRaMWFnoXF0SqDCEpCiKQ+gH6VL2JJlcCyq+DwCUfx6hGzXPWrR38hVw3cD/s1T4CcVwSbLAEZWqPByxaTQE20JzJazBQrxQObO8G1FJCOLbN5c5Sqr2gQQqyQR79pEu6734nbrl2Vmnd6+/LFqnOMwu0lNQj6AYjXfUnVtlEuCYQKBvQ7CEkAmY8mvBFlMFgELQ5ErbTPa6A1kA55MvESXHYp+rbEf107I3+8dHzHKwhV02odbCX9u1PVkPR94Qvbjm1bKa98/MrVl7d9pC3WIVaVqqaEqqiPJ6jVFkqlnQ/I5+rI4FdhvuuoUDVkkLUSA81YHtXdTUJd6xRUOsCsZWQMBfdtjEqaRLKIUWWO0X1acX4UpmH0D7hHIBYT+W5Q4D/8/LeADYUKYtkfYOO4JCMxhHAqDkgnpaWyCDPkCNxUCsXr3nML7tuahXvX5QGb07qqTzle1sWmfED58zzzMAiFqEuIP9+/DT/xtsl8CROOjVOb4mLps0UhkpEnEjkJNFF1zlcPhjcacDLJdAuwhYSRUg5n1h4OTdlGVMrnBombn/4DNqZTohCQ7q0OOplYp6zUyULekI7Z4mePfRPb6q6Bezf8DW9ffbWYWFMDvlLcfqElI3T4QN8kDMGxAyi1zZInvyG/+rOdF8FlYyXJl8wgaDhxLDvjFD8z/QhI1pYgtz2mhrr6dtyunDswoq/8q1lxn590VwX5NSaFHQonVm3XTJ3q9695sr293er8Ofw227Kkffy+wRtVEUuWlE51rUpiSXl23HJ7Vssn77+m+lN8wy91+SSG+0dmQJqiBEXYiBBi1QpKkaweD12K2/70zxN1c7rF0LOvAynjAIp7tDWa3zYVVOYg148q+da86LMIIzUAck6dAmlJsmYcibA4ayiEsslhJ0F2S9isC4+AgWkTkMSHLGyL4IAlLBWWqMLqn+tT2xWCUt8Da/wnTlaZlteDxABCZQlLS9yXVZ11mMJc1+TdiyoXekOERVsLAK0pJrIqyzJQEYCNQWR7tctD3gMmbHy6p4S9/SXRWk+rPTGuS2Nwou0j4KKlWPcyupfmrmsEE2v08BUl7FpdspGvlJQS1/dvEKsGlsKEmjQEIdkmk3Aq98hqSUAvDLAmlsXNhdvh/L/tJ4KwF5uzcWKOZ/54+p5c12SQg+6eoVP0fekkXRG0tgZfOPKbmWs724dXjuUTXhKDMJtdZVVX/ymoI++AAOgOgJutzIjzzY+exzLK1SiJsZSHjO3CJlBfTCd5nhsNzDkUsUxCefmZAPBU57JlFuJyv/XY+e/K1m54sHa8muYVRFEoS9quJXLbcLBnk/VO6Fqc7+wEC5Y3aGfYTk4GSmnT1LOtDdX7NvT0PvAsn2flOdOpFXqX/Q0E/E2/ZMTJdjrLF45xjoxi2Tg+5/O73FeZn3Z3QxOcirh7CRBKmGjqdEBjRNYqlLHLSs2StDWRbBaXDFmDiSDE5AaYz+ggnD0DfR6KygQIDhEWgNiyzUNtEJTwvBC4UMS07sYLpFtGYZCry4Z8jpbFALVIES7S/LWkgICEgM1Jdo/0YBDmQIhqrWStgx29dmhkIyMXyIkKFUJdvAqEyIEtM+RLUnhZEYNq/0J/d23nSBHYK6JqHi9jQ0P4BRDwTv6+HS/4Zr7mxgtof+YkFsaz0w6HZONxGEuHQlC3vANCujneZLT5JbIH7EAbtdCyavCuBeE1hLh6/PhHZOhtkipwwE6CZVun8s1dPie89FIQXYsWbd+yPPGW7eutDWCJuIih64PlbFsjP7Xkx88+wa3Nnbp8OG8e9dq7bxLURRMElJZ6oveBB4Z3lko3QycYSZWIlndyBv5fP9E+XuS+NMELPF/uwG0+b7ZKpI9G6SuhQvYNIilng/qJvp5mvWbPm5tQqfKh13+O3KIZTX+Z3xkqwBaD/08Xg3IzS7bkeX+BUri6p4TgSIU0O/VR+Bgc78MoMEnTro4mZinvqkIFyVgKV2xfCnmPHxtc378CA+XrFIbJQhk0kalUaupuOn86YKDonFxFdkm7QtQUERkPTu6adludnuGrQv5cAGFVFbQffVnTFO56pQTj2NjleAEXhgEzCG7jJ1V6EvUcByBtnUaXqHMIAJAve4pal0WDV8xNMWotXKF+7uAJ2bvigWHw8rdCKS/Rsj1AeW5NU9MccpM7OuZLuolLrl7+xLK/Nx65aYnzo6Gt7lPbVuC3H/z+2quJaot65gHm0aRSTz0z/QNWrHo6CFnEUlG46F33b74vLdpBBRPTf+Fnl4ApPYgFmeujtWdhus4RqpKakjJvfJH175qtKnJSTDgQeQ1mU5La1MBBvomj+zL9J7TWKoGQsuHGu7ohVwjE7xZ1w9M9HghXCkIblsNC8g44YiqNGnu0eS9sLfjpoQIGYMyOQ39xI/zswStwoDiMf13RKZJxhwIvjhJQUe7DiIBpxGPUCRH1Q7C2pE7dCkWSYGVINMMi6bAaKGWWJq5khIEIG5qteGNDQIsMgd7HDMJ/GDJw4O/WHjjLd2pPDYUVCikcISwKf30rFh8IK5OKlP0h1XCd8uY0vY7A6SGiB4RS7bsaOmwAKFwlCr3vxkQGw2xjfKi/9LO2U7Ydt+rWRSXomG+3tzeKzs7OzZvuhI8ALHQBOriS3smPH3MG+unpc2YVBrNfF8kEAfRjotS7pmqifQMjHl+M578njUXLNZITnH25fMnJT508Y39dlw4rK7vaV7N2oDHQr7G9LvPjk79gYjuK+HQ6j94PfQWQjYu/rS3AjA8+jFs8EFCfopPQ5Ck63We8c7RAeRWTLDDhS2SoNBLDDxBbMjVwzRPfgOuf+Tnkw22yLlkFnvIpnDD+fzk7ywnGciXHRB9aTDyq5kQ8PDqhyLqBBrGsi6Om8IUgYjZiqiY4EQCuINr3V/L27U3j31hKDef1/dhHQqc2IahJPpklP0wiBkUJQX+0JT8abBDYqJse6AoPlovdETvBzoOTaFYwtOZRWey5SpZyMeWkimFm4hFrHjzgxtq2tvG0gnYSFdi8ec7s9tlkDIJ575/ngJhvsniL/cSEfQ8uDGVvVKm6tJJWIEuD0g63d2xbujS3m3BhbxikmRASuYlIJKewTQtD8uj1MhitqFFGrnyFCVBlYGFMVGQ8A911bFZ4Xb43S/Kojo6eXZTNAWiIwxaCODalmdpotHoclVG4EiNhaFSVS4iQFmY9WfVUjhZ5ESCIukxGSLtPVifS5PPrB0XHGLyEoFA6R6DdGG0YdBTEoA4k9KNeatjBYZdAv8Lbawy9brTlfaGyAhWKRNKaogFtukvrv3Av92qDwLj/qkmTqoWTOQssByGeFmBRFp/+rzy7t6hzCCaMMxh3/UJF2+soLc7zCjxyfjhl4+esgTXPiNJwXLh2UWUnnjzYW704npr2iXRbWwMsXuwv71xONS61+KrFvqBznD5xqpOd+51ioeY+lWyYjkqUZDEfl/ltVxWHN/96J82AvXJ0dNyRUkJO0k6WycVHQEC9YkZr6WhhZ1Tq+bnbsYGOAEzlJmWNshpNRXA/GsFIINDO/w7742H8wVhYfi1UVFDh/CQ7LrRKcx7AKNF7AVUGYhhQfpHjDUoGEKJSMGaLsAV8epzhQCT4AddtNOyBcxK0H/Y6qcKg4dTI9SRzHIKUKkbMM2EXqBLRQsL0sHG4lU6NmbTGxosKGWgSBfne1ALIVDUTkg6ydRZDSglAxOBVHTHQyNOEixwAvWJVPjtRd9zzUaix+ejvXzPo1k56sxjecjNmxk+R8XQBndb6IJ/5jr+98CmrtvpeNyGeFkFxHUjX9kN5yvBA4kTIZrPgUkhjFYU3HBdB942HfbTlQ4s6zn0e1N/eMIxVTY+Lc5mG8VECpGWBcC0G5tBllTb1MUqGXNAkiUSVjKkedbr10l1ZAhHSEgi25K5F/nwEttJvk6Oh/9D1WArTK24urdo0J0fLjvoN7XBwtkKnJUS5PKgB7qMk3bTQS4AYMy9yvRRCEXBnlGmVA1TUR82oSw0IMUCqKB9KJ+xaFkqpBAlW08cc4aJj2xDQBWHfJLTR7R+Tp3+e8XwXh6JFUHbN+9DOANgOYKYORGFIMBKEMktVFoLmKKBhvLxRf8GsPPoFhqFI4st7vkHvWl7f+qdjVXA8Dvm/C3H8kRjLQJioK0FCtYAKFhQItkpoAWlTcEi19QAllITnx0RuuyXzW6+cfspBH1vU0RkALHrhDUN75LiUzz9RPd4toGVbroQwhyLsIcVmREjYhPqBsETwYSmgNgWyLk5pBmIjMRX9MkohytiZ1Z4+IoXalmOwEiXkoTkNImGPUmJykKCBx7xHxmJwbBJ5FJoZC0fLjhBQ3ZIKLWQBNFCKdsg5Qt3rpEFoGmst4hIg7+dU19AQNWeLtGtjYyqJSimBROjKfgQHAPqjHPZQOwnhDYSISRAjfj/2BIRZdNGVVWDbBLIeEAlLYbVrQW2iCkdk4AxAZa5jbLxAg6BBMInM9EO9ePUBaMkQEhlLWK5ORGmzHQ5tWl0JTxuNYHklKj9NZXwsKOqB+7eDvA5ZGly/bvbs2cet2rDqY8rNfhhj9ROUm6AaN4CdML4yNVIWgSXNS3lbeP1LXTXwlcLwhuuWd67mM9m7jQGNS0mLAkRfXygzVSrcMgC1rsQ3nDBNnnTsdBg/IUuruxzs98SKZ3vw2r+vgsdX9QmYUI0yZgGt+EAeAI/oH11eJC9DdRfw/MOaYcEbJ8HGzXn4xg0bYf1AEWQ1AX/K7SgmkxehzaL96eZ2s+HooByFNj5sfjTiwdgNY5M4MhFkDBzYNtiPtfFp4qNHvB9m1E3H3z/5S1iy9XpoyVSjj9w7JYhTwfTNMeiInCAHLAjRx1X9gzCt6hg4ddrZMLN2P2xKtHFGZXNuJT7T/xA8uq1TrBt4AsMCoTDpGa58bMfGCzAIunc/gPS5Kt5gQ+j7kKml9I3x0vRkbwKJ20Y/pP3UMkLR+I36Hc7wvAi/nb3N5cs5V/CtbHb81QUZni085whlu0cJhQ2Muw3RkzLsFn7xPuGom/ZrT92++KoNvqnblyPhvXvoSxiLbS7ke+qK7zn/oKovfPQomDKhprJjU/9+/GT45HsPgd9cvww//c0HsC+VFFZNAgiDoIs+UuMQDIZJBSGkYghfvXgOjq+JAxwMcPLhjeLYTy2GDcUQZdzwV5SjB104Mq3Pxt4bfgRyyc1gemsqQJusBj0zZIJMw4XxOcixsWCgOALV9my47u13QFVCA93GZaeKBX+6WdsZw6tgcbxBKGhyUBS40oJCkBN9xQx+9JCr4A1TF3ATlz5bOhcBLekJeHDT68XZ0y9W967/k/jzvT8N+wceYQTUzhq2Y0OP3fATLAph3mlJJeOnMQY/REukqgBCutAWWNRpIgRug+rKz0fw2QipVM4um4r1iy3+Rk6rNTS0qc/f/uTPgr7H3t06BebWwZpZTdlN+2eSj8w667hpB/iDyz/obX/mZkoy6twH5wxeNXecMmP9/f8QH/3g6+yrv316pTEob0I/+bzPkL53v2UO3v/Ht4hZMQHhYA6kU84jRDdB9wxYAkrDAXR35SiZJ3K5kJuYvvmuSQDbi0JEnysv7bp4bMJBgxOm0BwtF0erDJr5QCcMdfHA6HfptL8OM6iPRAjoHSrAJ4/7IhuDkdKQoOzBhoEN4AfBaJkiotWhijYZGmGJYmkYCmGD+tEpt8Eb24gzrYjFIGqwonVO08PRfxJWGk6a+h7x7fa/4cK3fEJLEF+68JW5ea8CD4H5Cdw1G48L3PFtIGQITkxALA2CYnfDn63zwCsqwTQGhMS/6mqRfp40DzPXy4MXq93EaCZ9Y+dbgD1y04MPUpfSVt1oKaCzs9PSGISpyqAe9+pqwnNHO+ULCf8xMHVG05Kv//iu4+97bIvKhVIGvoLqdAyP2K8Vzn7DTDFrah1SFF0sBbDPtBr4xzVn4uELrhebij5IN0bXXmPKGdZHcFMpglDgb27dAAfNrhOOeRqO378Osom1OEh5OJ0CNPmCCJJqYM+6uscztdIJ58w//fDsBSTiZotKAJamLyDoAn2yGPiiNtUEh088nPfi2kSybcEfnrhaxB2SeaJjUHKD+rcNkBElqDDA7cUEXHHKjWJGzVwoBXkVs11JFGzrBp/B5T33i835DVDw+yHlJDFuZ2FCZrLav/Yw+c5DvzNyvvguXvVoF4PY/gs3dA8dnCbYfVIxVPFzMV5P9yCEVJWDjitESGTruuzMaeEdh/ERDXg5gq1oOgSzQD33Qy9gGHqzCiRf2Uk2eB1YTD9Rh+SrxSBEwrahmPqRI51Ew9c+etndh3A7ZjJlgWAhA84H3LJoPXz9J4/A+948S3zzs8dBPOZAruBDa3NKXPKe/cVHLn8UralJnU/gJk7t8ytfCWiOwx/v34qXXRRg2uU3hB2zwaVkZVTWNPczAgJFNAhRAwFHAtTXFQ2lG6gMFQKnIpEwrsynEjHmEAcDIaETnCakj1mWBaXAxxU9j0FVLCk8wiKSipw+CGOkXcuGzUN9+Lb9FsI+9XOhGAxB3M6KQlCE7z/8Cbh/86+FlINAuWbXYjo38n7oUqkqJyO29uQ/2/yR9GUXHnxVD7XJd4yezmtxCMOARospz5udDQKj+erqjsz0hfaJSL6mX7ChYQJni6kgzNMbQ+51qryQaQA5QpaCq8r81BkgkK5B0lKhKB3+ok42muCLVENDQ3oQm+YHhXCa5VqNHKUGXp+U1qbqmHdPd/faKJ0RGZ3nWv+9R5wlMqFKTHjPZWFm/GeDRJUtbIWWw5gkzVorqdtPCNEosABC/PBPz+DSZdvFTb9pR8eywPcV7De7gWIDziGOJgMjaLDiZqXhgZwo5DxIx5M6/qfNnEq0owEh8R2JGJcZxKpdCOa7iLqjSAmOfBFuA9ftKzrAiHj5daYJBVBJsHt4OwyWhqE2UW8Sjkrw8U3akc9Tv8Mw6yD0hWPVwJnTLwA/LJExgO5cl/jsXQugq3AftNXWCRsaKSigKiw9cdQcRaGIpURRTZse/0hNtTp76KDUJR0i9zuT/t5bnouXYpgscTsD3uiHLkC65cgjSiXv3J0MAqH+FoXDpb4TwJ3SAEIGEAYWJKuMUKuR29GxXRn1QmNE/1OhyzUKqNdIdRovOGKIfIAwWT1uro817+vDzOnKTUyEJMmUR2GsgiD0YLtfGpB1tXc5mLvS63vmjgguGxmFsqy4we3t4SuCvlTzj7WsdRf9JKyb/h4gIhTfC1CBFXgUtlEviSQhO/04B4RBEOjMbIC7l3XDOy65Vfz58tOZ4HT1hiE9eQ3IUM9OqucTxsxBDEPRXOViKhOj7nMy6VKEClUhRFJxKnsHptSg6/8Mhih3GOh7MWrsSZ8p6n4yQ6cgDcI44jqwLRtD5YmVW1eqKdVTgOjg405cNCanwNb8esjEq9DHgHonuP2cCpeEdiBPg7BNjhWDDUOr4KO3nQm2tRLaqlrAC0oQSJ/RFvwAGB+GzJUtY1J4GLQ2YmttVvz2rb/KHCnE8IeZqNf0SsBrwhugedEJdXUzM7miemPoVF1QgJqTFHWS7viZRTzXQiXaMValA0HbEZDM6sCQTC7FgmTimcNkFxewouJt6Jc1slY3x72QC15+AK2qOR1+rPoSTDfGQvL9SC1WoaVVHkzLvPA9UpKDUJ4VFgfeZNclr7PU1k8U+7dsnDdvnvPoo4sDwb01ALPbG9LLO3uM7dqDw4T25UI+WH9t2NB2FmSSPgQF4mI0MhdRprZCb1EnaYSX89GaWgvX37sB3vC+6+DEQyeL/71ulYDmBOoWQYKQSIrlhcqHqAoBwOoBfMcH98Oka8lCKYREzIIlXUUxXESUdVKwGI5WwjDwJhMNRsxRfNY7O346gcg4A/YYuVGBfAj2FBlqaCGESoqYA3j3mtvFyfucXH46Ljz0Y3DBn24FgX2YSLgiFUsKuvsBUs8bJbVH4Cv3nQ8nTjlH/P7p76OUa0Rdsh5yvgc2l1g1qw3DHDne0FUO1otSYBfyQsUSUk3f37ronGsytR3nD79tIbxajQJV3JaPegPUWVwzc788JNsHRfV5qqFpqkpU657U3jWBvXO4kB0/vjY3UnsCONTI4gnhJhDiKSDNRKCAjtxY3XVbiXajkAFGyh5pxKChV6XIH1OaM/H5Bj/xLfNOS2zf2P1HFa89HZPZADAsOsWROBRzDqK/XSnoQRWGllQNwnabIF4FELM9UZ1RYaHqLWrAOax6snv64sWLl9CpzLlg1mmts0Yurh+fn1k9aZ+3PvDtZ/61ZxJlkAXvCK173/lDNX7mWRC3PCjlHU0gFVG5RbCOkLCizNfGv7PzJkXoKZTjquD2J/vE7Q/2AYzPEtm1Roq5AsPeYQHDCqY2xGH/6dXi8DdMgve8dTqVJtHiFiSF3/nFM+D15oUcn0TLCalbUDC3emQIyL0ow6aJw4zCggqkImFW9exmcBQ/DVwv1DBn6kUgnhYEHyw7hN8//jN8x7x3w+ymOYJW+BOmnww/ffMNcPMzN4hnR5aotdufhYRTEK1VtUAaM9lYUmzNPyJ+vPgurE9mRTpeA0XfQ0tfJ7JCjKAi5Tg+D9KVNngI9m0FylIRSW/SmzpbvvXMnybDji35d7Z3gux8deSgKrwBjdKtqppbnRdVZ6AVe9uQEzsWs80xZcUBLNuHsKRdf+Hs4CFwvO4PJI+BVE0Dof+g5EtMNIIgMpSovzxawnfVqDTai0pbMl+xDkS1Z6tTYM87CL8S9q5Ze1WYnXI6xpMFKJUS1kiPJb3+31sWXlM7o3rZdyc8uPWcP0M48fCpjds2wL7h8Mh5WKp+O6aycSeBBc+fOGFgW/a6OW+rv6Rpv/4LaxpHTqqb7EGxPwb5tU41HWj58j0Ny24YkepOe5OqnfgRjEkfvCKRTmhSgKicr6+xDod4IhorTJONwIq0ka+UVZ+WooUgJIzhkcrzBawdwTOPHo8fPGe2OOrgRky5TjmEIsYii0xLiPDpc6eobGaz9eeHB8HPxlA2xZFo1ZAl+RiJRMSHUQ+E9iQrBsHFqArJ3gFFD3T+XDDUzZExaUPv8DYoeg4ePP4kcdZR7VCTrMVQBWDZUvjKx1Nnnwn044UePrVtCf7hyV+Jf6z5jWxMhVBjZTEmkpDNVKFSJCwTUDqFgxleq7g0yZGGBkczUFJJynVGFLQEWfB9cBwB3qy5znmAVY92Lhj8/l5OsyYrcwP0QqJ23mGBss8ZcarOUsm6yZioISNAVykADAT4JHBGuT+hyIJWGIR2APFn8GXijeikNRQdlSXSNZp3gisLIdlajUfWWPfy/B7hKrFRCtYZYabaMN21pll3l3wI0fG10Gtm9oUq2XweJNI5KORScnjzChtH3lscXHkfbZV/EGDBg/oTa/61phsA7gQBd2abZl5ZLHhXFIfGzaubUPDnvn/LtLrJxeuSqYBcZNj0pF3ccE+m46m/PPl34lbo7NijbjrdRAWT5sdRtH4Ls/UIYVEn5jThVIT+NC3O5Ti+wmU3eUit1ixZGCEMUMZsqQaKMAlB/eBbx8KZJ0zRBiAIIV8g8hhEy2Y0sEYCSoHzD2oS8w9qUp9dPSS/es060fnYAIpp6SgPYHquDR4hqjw4ozgETbZvmo5ouaameCWYgo3M2qruLpjfdhp+7KjPiEMnzys/QzT5ybug9aToF5gexZK2OKj1EKCfx7vejV+680KxeXgJjq9pgGJQoi4O4nenp1JXNgwDMz+AZLpsvv1kQw2OXjd3Uk8knYtXEnYmI4OGZrXwwE/XXdu5oLdrL2Nofk5uINO6T52Xi71ZycQ5Xrz+aJVqdjGWoG0p5KYVg24EB1PG2zNzF8sGga1K+1varev+0XU02i5xB1K7GECmmon9dbYwAhzpxoSdsEYmUxBlpKOlSxNysle7Wz4EPSESdRNaPUh/DVN1gfAKcTG4aaOTHDqxuGX1Rs0oxDmOcpNuuRIxrc0eWrXiUZhw5FmzT1799IFnbcvIUIZ+0cb+Pg+3PFVz5+oHk5f1P7j4fvMQ72E3m5O5gSw2vQkbW6YzhFA/x8w9G3HOVATrZXCPvgMRP2K5aYmjNC74F3wxBQO8+7dniYmtGVEoUtsiMjGJTQWLChxi9G/JIzl2hP2nZeHar+wPv/zrRnHh1evAn5yhx0jjk6KiBRsSuRMamIADUR5SNynTY0eXfXNvP1xy3P/AR4+9mA1TKSwJpRQknIRwmarefD9OGenfvbAoQgzFgS3z8A/ti/DCv50ttg3fBU2pOoIu81JEl4ksgeZn1ay43IxPpCvcWRFRsnEIYXgT9HUrFTFsHmdVN08pfRwAPt0+h0OHvc4bcJITD1JO8/k5L3s21NSPw3gVIOmJCOmDCugSMB3taJ7YPCvlxZ1kPXhwvRv/fu+GGYCxiUjJw4B2YAEkdEIxUvYcLUKVqcYrTrCsB6LB6qYRzQBOCSuLz9dZ6QXVF2JNYx0tFmJgq+Oq7e/Nb9m4EdraYvPPW8QApUWaD68Mjm5vB+jsXFUCODJzyie2/bSprZguDVuBcCxyhOwlv6/bsP7Op0+h70s0a4sEMyPtwjtZIwHSO50f6UfOr/idxvydfn+xgxWidjJILO0mwEm9A+0ksZMwHG+0zVmvydzpz9Ew3UhTYeC/mQJfgUX5hPL/qIAPuGVEffObx7ExGMn7mE6yhqN4fNl2vPmhbbBpc05Q4SKVcFTL+JR4/QH14rD9avk6+L7GhF1wxgSor4/BOf+zEktTUppW3yj6lpOb3iirtWVr0LIpjnKF2HKE2NizHX949i9wwYHniFLgcQ9z3InR3Rdr+9aoB9bfA091LREFf0QG4EE2Xg0nTj8NjptKvCYgin5OZNyE+unp18Hb/3ysyPsrMOEmKZ8iRh9HPeu1e8qdmOZVfYkp80JU75xgNCbCD8BOuIi1tfiW/T/ZdGnngm25PbREbbwB4tek3EAn1NRMrSpg+oxQJc9XidpjwkyLA06C7hHl3mhht/hf/dhEdPuGv6iMP6d/OQFjDMLdHJOWCv5hmGgk38IHDAmdIiCWBPRI4CA6Jamh5dosVD7YmsdvtJchYubR4YP49w1NaDlvBjeJyi+6Mui/JT+08Xb2DFYtKhlDUHFhuKedEoNhy2nz6g86fuv1dRP8owpD4NuucPrWAvzr6tag/9lsa2KcdXbhvUuvX9TRjpqZaed9sYXdTQix6Hn+3vm9FzwqS58sjQewfyNaiaNIboUa+c09Yqgw5XA4RW5ZEYViJGc62mtkuKki7AcXF4sBZBtj4tSTpnIvKhmDTVtH4KMd/8IbnxrBMCYlJAg6RDGlrsV84Tdr8HVTM7DwPW1w4iFN4IcKC14oTz+yUV19USDOu3I1WjOyGsWuCRj19KugVCSIEnMQcEQPQJN+Y98G+MjRn2FjQOGAkBbG7ZhYvnU5fnfRd+DBzbcCwiA4DpUeLS4OeOjhX5/9Ee7XcKL46ok/wmm1bdxunXFTOKfhSHhg8+OQdrNCgSIuWa1UQzo4lO0gA0rt4OwJk0fAvRNIj/ZgMQfZWDqSCeH/hj6EtbX25Gxj6WC6se0L9qgEYxQWlL2BVGrK/qFb865hK3VmmGyaiok6YFJhytQiefc0tw0+jE2bgYHo+TnK0l1xBLp2xiCwK07AoWPRTmqbQXc1mQRwYwAlime1B4uCu6IjrpPKab4D6EXnwKLeFw5Q+Jbs4ssyeUkyOfHAohWfCcIKRHHIEcHItfrEF6kDP3nEpNrqjR8r5nHb/b868IfQdVOesAUdX4Fw4nn7tOx3aN8tDePVAX4BCm4cE9s3Wncs+m5rulCsPsweJ2RpULwVOoB4FXfmY+BLZWWnn6Xy3pEc/QYBpeyUkKQkoGSowoATKNKxKWsniC1UBRQcE/4NheVQoG9QtUoKeg/M56VNmT8E6bIZZQp3J3FzUFz2wKhRYLcPYuP2neO7qSxlv0Zr/nTGEnBgBKriFgz2IUB9xhiDskte5hIzHzPiMFRVsCA/4MM3L38Ev/iBg8XTy3vh7I/eCasDCaKtQRCrOjc00MNiWxyaUHrw/mEfTvrM47DwnKlw6YUzWFat4IXibae04l8f7RN/WjkCVlMGQp+o0rgcDeDFRq9qBR+RJS0YKgzBzNq5+OkTLhF+6DFrt2u54qoHfgbfvOcSkUr52FJdIyyrTnOtMdkeAqnpWUKJNYN3wjs7jxcfO+YbOKvuIHhg/R1w97o/wPjqeqZopwIIy0wwOMYwrxnktC456ievFPqQL2WhJX00bh25G+rSOmaiCx0qoVIZiVXZ4BSaEN0XES4e9pRBC1nY0jIvOZAvvFFZ2fMCN3uin2xOqniGKaqYhTukBCHJgHE1MIqJ9PpQKYtQuTpraED5OFEOgVd2FQYzhCTOPv1oiFiSRA15apdVifSlppbjnRuImE+R+f526JE1lJj/prWphGoSStcBvxTKsFhy05nHfN29hnUtvf9TPdF+qyghHP3+pWfc+8N5b5h90eJcS+G05KzDl99YMy48oFgMC7GEndj0dPjQipvUmQGWLnMa3CMQbRB5awrJpS9atIPeARsimZz4IeU2/gjSVfrcTWOgccoNhF9/vZ0g0+bbmfYAvrjRIs2eFBP+liFZXKeVACObPw3OAcfC8BNkFMj94w1ChNmK3V+N+OUTdG0Ie4bxza+fBt+77Axx023P4Ic67gA5rlrQ/ddw8UhW1yzLXDkkI0R5G4Agk4DLrl4qrl+0Edb3ejAcT6A9Li2Cgg8BQ/kYWqhxxAyFBrASMcDZCdHRuQ6UH8JXPjwLiyV96RaeP0ndcMmT0mMNDXNvbQuwwkMw1Q+WhLVtoQb6hsXlZ35ZxKw4ibuSniNccdcV+JV/fgzmTG6lSEcEtERDiLYhQqOcA1GdULFiXHUjhGEev/vAewWiizYEorkmabRmNVGKAUZELhLnQLSbpNss6XHsLQT46XnX4CGtx8Plj35IPNJzpWhKVgMRthHTq5sAYSdxEn2Fxp49JVxot1LjN81CiL15m2+fi9mp+2CylvlJtDcQGOpCE8xrrgj6oAnn+Fd9nSLc+OjaHU3sqONIy+jRFomaw8YJJzmz3N1MD1Uqq4ks+MNRQpt0tih7syN0udwTF21XXsFMmMIv66d4V8NyE9MI7CQsaWFQGrJLKzZFXyff780IS0DPTL51BhxxxLu3f7PjOAhmve7Jb9SPDw7x8qroxGRiy7Ow9OHfZt+46cFNBTst13Ncy4zA9tSHH17TGCH5y+dLI7TfB1XN1K8xjIl0AZOZPCQyeUyk88C/V+Uhns1DMlGAeCqP8VQe6L1E1vyk85ig1+gnmcdUNi+S2QJ9HhPJAiQSeYjTTzIH6ZohSDU4FnjM/qvP5W7+RRUHanhCRrEwP9AIUCjBW08/ACfWZ+D0E2dATY3NkzQqLFTcggg7qF+IGC0lgDWtGp4qIAzXZ0ifQQTFoLyIVCb/9BmRYhp1Hil09q/Gy/62CX564zqIxywoeiHMmpCBY+ZUKRxhYlQ+dX2+Fd2ObFhIpUPiUH5IHDjuMDhxzklcFSBjcNvTt4nL/vkpmDNhIvXLCS/QVOx0p8qZAG3fOOyksqIQLozP1MGkbEqMI10KIorX65P5prr5kgyBLmkR0CHylC0c8UegOXmgOKT1eD7PaTVzRYnp4bQ2RUS/UpO2aMmF2cv2CIMgJ81POsHIwBcwN9wBVeP2wZrWEgj0IfRCQN/IgVXAy8sOY+XqVIavRY5k1JdinhXzCFDBJnqsfK9vMsYm1YJwQs1ZS8Efr1jRyqfhsnTjydAQa9KOw/DxRcHJ6PNhrvfuGqD1N7FjpPxsykIqHByEgHFQgJDrtzc5rndA4AP54EHrDPuC4760b1DfMnRBWAwDJybs7k12fu2KqnOHn3mi18Bpu4lfjPNxCmMF5bm7Oq604ScwuOlKjGUyTD0WgWsj+1pu+derZ8QvWtHeY+6DYe6IWOKIsqjckqAfWCFHAAo9JSGKN5hPG5EoVpbR5zd6i5QKlIS4BDfl8Ib12YRoziSxPxeisGl5528QGbloKpl/uK2Ib33ohShjMT1hOFsTnddo0ULL15mORi4JCAjzIYhpafGZa9aI0w6rx+YGVnCCk+ek4Y7VwwBZnaBkp8arKDMoitlJmMXGvpFh+PARp4EFUlHLZt7Li6/fvhCb6jLCp/ujsYuan8GcvF7QTVOmeXIoR0Ds8/pZp6YnRhjp666bIE1jZvlJjFKyHLGWggAmZ5rKE4QcKH1/OZzVyg9IEbIgjB10fKUc+Pw3DYNav+iaEiC+LVE1+V7R9+xlULNPLQkYG09cjyiNGvFblkmxzIrBcZhx2iNTUH7II60bJrKjkGGe5G5BmT4GbUq0sFSXxX4Wu7AmRKiwNqRHSChYgDYLYFWUtScKnmiClPsYeERRLxWid1PwQ1XaTKrqEU1PTS2IPsPpPNwV+0NpIDwjng5RhVJa6UDWtOU+xOW4QIb5vLS7V1sXrPrlE8vbTmmLrbp1VcmFcDDPGThLgQrjgFZFkMuDr2iYW/uTmJRD/sC6fSwhPHZ3VUCxPyUODJjKFsq2UfO/0Y/pnKXttGajTv4pKUHaxncPqchhoHyUiqU/bWFl6m4pDKx9NJr0owwFpuJD11+X+jXG1wtR+bQvgETcxVltDfDM/euEqM0YkRZDfLJDXiFSPtV+u1GgIxzJ6ANgxHTMudMiaiwas2XyTKPKs+1IHFQAV96wWXz1AzN5N+MyLomfRGS7AgLPJEDNU0yZFsK8SCVskYCZjbOFH/qCiiD/WHkPrux+Ukwb1wpe4BGPo8YraMiFZlw1tpcDWZ1+MhIQuiGLHElbU7lpbMaoDSA7of/UqyEDIejCkjcwr+WN5YsUhj73YZNmL4ULmpcJiADWxHh7TJWB721BwI/duqp/2NuXfDNMt7wpzI6jSIFCBjPpzGJLBcHokpTrvmXpz6jh0PiOlSS9OoSwARZH+YODKPkVSWtwSiCWBKFCbkGNxH7LBy8vkZWj4s5FLm3U8VqRiXzOF6aJOdS/AuOttML6wk5UF71JUwHXP8Gx//V3/6l+/MxTph2tzkdUpTAgMC5l/xCFYzk9z1i/f/yHq37LZcWOcSHAKlEa9MZhnKqgjLvNx1KJfHEXzFm0i+LQqt/v6k7sYLuoY+LfbrTj38/Zno7fuya6iOZajGgTGnhFzlNqE0Rc44z0oPBgw/r+aMET737LgfCXG55C2VQFquRTT99ouS0SkyW9dOpWZFFoaoKKWlPJiJkmJ5MjjiJugg7oFYbXaf0EUa7URxSNSfHbB3uw431t3KJ8/eJeBXGL5eP1WqP7E3e+q6gIdRhCwc+hQ2hXAHHj8ushnrAx4EmoVyndKGdqhYaRzSTG2UeIGB012ILkLKXuYoy+gAhFyFRruudSr4WRzqOFJVWEMJwgXjfuTRhCwNX27sIA69oazVHTX4ngFdWeya+GYHvbl64EAWfFw7bzoDjynbB2WhPacR+UT0WZMrth1Eo0+nukWl6u/+kvG3WVakFmthxkPMO2k0+OrX2g92DQMrracpDpdpPcs2D0zo1coNFrY3d+1Y45hPJ/De5jtPum3Fm/y69K63gyuyJQ4TBgkFBu1vYK248AgCWLFo2IhQuF6OhY+IFUw28mNswMjpMCSgS8lK6wezZgbtvKzCW8VlxKV6RRI/GUdQjaccLkUjC5/NdXfWPLggULoqu20/ErcQi7KyXuCnNQiVN4IYMmP3E3VJ4DEbsspjrSGvBGAOyaSHKaJNMEpG24Z/Ea+Oh7jhSFgo9vPGGGfPd5c/EXf1oq3H3HQVBkZbayX2a5FgREZz7ogUE7SojZCJkkWnFbMIJRj1E568pLUfY0tOPOFKmuRTyLeMEPnhYttTH4yxN9CM1VwPtibV3cgQ+BYCv0R6hCVZWpkl+74+uwtn8VDJUG4c41N4uqTA36YWieKl6DNCMzryVGBd4sXEpxkzdXDBz6TSoolEZgEHwyFNIRApMJiZl4CiyUEGgRG/4GtpAYdxCWdQ+LDx1wOdbEaqllWliWjUt6boeUHRMBARPY8aCQBCFfVFvprNphjyo70iBPnNPsxaFVv0vUpu6W25/+XpAa367SDYTOIrfdgNOi5MkO+aHIKOihEStRPsBsJhW7qV2PrahGIEAQZZsZ+m5SxI7uc+aD6C52aiLWicbn5BAMTQ4fW6fXjXHiMJF/39lr58HGwxte86yUVWsgI/cXNgEr3AUg4ErAqaqDJozoKC656ZQ3zYUVf2qYhKe4roLB7QBdK+2Ln+1csnnBArCgk8sz2DC7Id23zT6TAVaeBwK9lQsWLAgjePRzT+H5cAjR2J2h+I+xCNGxdcl3ZOUDEmNKJauoFsZuvKI0RDoJ9z7wLPT0jUB1NiF8L4ArvvEm2T3gqZtuXi5gSqOUqThIS2LghSLYMIxEnnpwWz00NKQ565MfDsWSrhGxfWsOxKQsCS0yptdEOhGGwSBJdNmSKqych6ApSqLzjUnxm0f6SVxRQEuKED5ltCQSWYtTIfZqZiTpzMZlXAz428VX7/gyuJbApuoGEIEkJIuJ701oazAE2hxpKCltojFXCDELoHt4qyh4MZzWcICoz9aBBS4Sr+Kg1w9rB5ZixiqJpmwGHMul2hh6KicWbx0Rb5/Vge2zz5d+WADHSuCS7gfl5qGHYGptigwTn4PFOUqB0rdZnrD77j2t14WHsebz7ULfos0gYIFb3P8Tyh/5dpAZL6hGW1ZXpqHtgekliRJiRoSnnImtQBOjgS4HxeRcEa9KM4cJK2FQzz3dNIt8Rr2DKBSIclHPXfNN3GbEQEykwA8fGZqocvE8SEUZDv9RlUbmQjxbUk72GCs96Q3hcOffoYNl2rDv4VuH7nq4/bR5H3jiQ8kq3Ld/q7do2TVrflfuTWhri8GqVaX+TcmPqFS2WViyBN5IzFLDN+rZvucUliuGTgoOjtsMbv4pALEfCFrmQpvsqxV3oGfjAHzvp/fA1z93KuRyRYgnXLzx528TP7r6AfjBLx6FNRt6QHkW1tcl4Lxz9hUXnDNXzZ1Rv8MxtvXk4Ls/XSy+9ddVKKbUMi26fqei3EBhA3cMmjyGrl/rqR1IsOoSDFBlgRUNLaZ+F/1sVPQyaNUE3esQhoFwpAOT6kljRgk/8FFJRegH3anFQY0pk5N+i6RGbm5NRIK/MRRDCLG2rx9PansXnH/AB2DflgOUNCdKR/PRl8u2PQHXLvu5+tfWGwXAVsZ0xq3x8NlDfiTeNPN89BWD6zjQvubJr2E2GVDIoh9qIUibQpZyKEq5+AMAg3tQ2XFXg9nDJGC78HKd343X1jysiv2dKlXXpOVtKMtSzvuZexjR5dKVKWce9TCqWfTDBkEVhidjMkEPRkDi6Rzq2y4IJwZYMgkjWkXYx9MNTDrbs8OIoKxR0lOrc5gYV0czkUGohP/y4HUqGMn/zLK7PwZOqk5kmxD7ildlG8LjhnoWr+JMXns7QmdnuPgn8MPyJ3UjCg2bjEFV3aTj8qrxyyKR9UToO1DsW1mXKd7SNVJO4u1pA3Uvw60lCE//E3iF/SEWizLAzGMgW2vhf6+4C056/T5w7KFTMV8oiUTMFRe/90h87/mHqkX3rZeDQyV57NGTsLk2xQ9BqUiJ0ShfpKC+Lo7f/OIxsG6gCNcu7kOrxRXU2ViuLESkJ1EGqoxkK0NUdbhRhlZEsaeJ5P0KWLpuyzbIIIYJY7Hks0q8lo/kbifTDsNNiRGnErOs6UZq3ULjWrbYPtKHb5nzcVh4/Nf5nVJQFKhL2Bz3UqX6gOZDgH62DC/ER7ruEZ4qieMmnoT1yWYOExSEGLOScO2KH8OK/ptgTgOFDxqareGMQub61XBvt3MvHWMPKTs+32Anir5AAmS/H0s6prOURtTPoX+HyvqVXqd3lEjQ/T10N9ggiGzrYZyEpzvFpPc+iFjM5AkqSPCYa4KouSM5nh1GhIdmcdcIpMO/RnwIu6dQM19uS68M0l/E4sDPVLq2ANUTxo8MyDtjTVXvLG1bdjd08gpvtZ3SVo5XVnWsIgoPcopUqvWwt4/k8XLINMZQWCWR75eODC/r6urK79lybseSjD2otP1LMbTls9jQluIZoxdPakzDoDoL57/71+K26z8I+0xvosYP9AIFsZgt3nDcNNoJ3498IRBEEhKLl691dJ945r7ppGlw7b09UTO6+de4kqYFSX+qokRc5h+KSllcGjBJyMjtrAwHVcVjw0ZAlzI5UUgegMkPsLngtlqtRlWmX4k4XNllhQAS4v0HfTDaOatIVyRm+SENlE/S8qI10yLOzJxTzoAXghy40gFHJsXfnv0FXLPsYtFWl4USJSDMiVhShLYN9mA//u3Br2/aTC3QHXt+CzQhXHH27NnOyu7wz2EqVQeUleYKId+aHUl0IvTBaCahEqHIskzkO5iQwWvDFCEUzfUlxRzH1ftkTn+z/nO2yJDvoqRs7I4XrZym0jkubX8qDFLZQ9hl3M2xhT+y8uqYjSco2z1HxbIFq2bCBG+4+w6n5uArRaxwpbdt2dOrbl1VPi6dit06fa7Kpz5VCFNvh+osKCmLopiLy1z39aefvO8fOjuf3oONAQ1qVGm3YENnFzSc9TMoFj4BiTgV9ricoyhLmEnAhuESHHXGlfD1z58iLjjvUEy62i4GPiOuwXVtSCb0aw89tgVvfXCDWLMhJ9JVDuw3rRZqknG47BfLEJqSggUUo0RTeYYyVWYEmyZaLD07NcmIloKL1qWoGEUBPoWZlbl5hhfxomP8DL0cRIeLeq80p3oZWaSjDEPezzaFQwaqZvt4+SPfEG+f+wHozW+Hp7Y/Lp7d/gQXOVJOSsyffArMn3gyxB2Nk/CZ5gnAkY5I2CkIwgB//Pg34OZnvwTT67M6oOVHkx0xTMRQbN8CQfcG+b/88O6RkeWuO2RXbRQXhw2T9wFJHY2hswN4ZkcYhUkemlaXMijNvG7EMu3xh7cntixfqxOKEND00rXwWFwgR3qmDcFkf6jMZUz8LnQaKynaDepltOuR9v/vXHYuFGUd5729AxvqRdWE16t4yofqcRCUih8Rxdz7ZE3Vw7bK/1N5Ki/iiWmhCg4JS8kDMJ6xIBErESJHjAzGrfyWv0wffvrczs7le2KYsIvB/ewSE95XxeDqszC+7xTOLCu0SHkJ/QBkJga9YQwu/PxNeNXvHxbvOPtAfMPr98HW8bWQTtiQKwZw9+2r8PvXPA7/WNpNLFJMk6TXBrItDkBDSkDaRaoms0cQwXd0PKhnuV5YRqtGFfoaOj1cBr3oQTM5VhkyGCIlSWhFxJD4G017ZIR+iErfnMNkqVdT/KoEhCnuuMS6WBpuffbXePMzfxAIeSDatUQchSuJWxHgznU/x9bMLHjLrIvghKlnYlW8hvfdm+vGR7vuFdeuuBz6vAegraEOA+aFJdZHnbkk/FWAwt62Rf1kUcfgE3sJQQohXMNYrHlykBz3BUw1hOCPUB9DVFbQ2BI9DESD76u+q9I2pUqGPEchAxPf2N3rnp0Cdnay/ijdFnOniO5Mj1HsscG68w+1OOzY3GRcyWj+lT06Q5rErua/o2HnD/T0LB+BlnlnWAMbvoPJ2g9gLAvoxHx0XUuE4TF+EB7DnG0U4VCHjgQOG4RXjEFhUGFx5KtB4ZkvLd/xRPb0gdyGvqGjP9Zy5vle7+p7VMNM0mIzJGQWMSGBcCyQkxvF4i0lWPy1O+Ez37kHJk6sginjqmDd5jysfHYEoDYBYlK1kJbLVUcO9agTjgwLdcTS3OVnJ4oxoypDRRPMqFKrMeearW00fNRZx3LuqFRBkGIgh7pyYDTkWejVkOdotgedjOK0pMk0mWK1geXrKgZrRyA2pmrBtojmLU1symgT2QqFI9SOYSHkvbXwkyc+Ar9+6lJoSI9nx6a/tAFC0Y9NKUtMq6oj0BFjD0axuxjEksLe+Kza2LWt4TMLFw7KjvY9Ms+00yBZAhGE8UN/GNbOqAZqazXhepl+AMphQ7kMEDlogmj5YgkgxBt3RmowGi/80u7dRIKJKcPbV8mbSYyUo/EGOQtMTEe3jyElu7hwUdhiYpeIQUn33wHDSf790I9F1+J8OLLsg05x2+utoXX/tHNdQgYlKjsDLRGQiAEkCK/mg/BzjsxvF1au655E0Pt6zC/9Ei89e48x2CF0KHbdeK81sO7tVv8GJZwYkwJprBgX9ETohygzcbSnNkGpoVY+ux3F7Y/1wsqhAGRbrbDqEgJ9BSGFEhxOhBh6AYSlkMlI9OqhBddGkahGlSkCKjGURFOTRgpN5deYi0EnPSNEILijMQP9Sfk6munEVFSe5frHgFnodSqK0hpEiUeWbi+rVmskrnYbyKgFGAo/RFEKFJYCFNTfQJgDL1RQ8kMRszMwJdso6pNKeMEqEapV0JxEMTVbD3FRhUUWhB6thlsog1hK2j2bsG/rs/LchztWDZmT38OfFyYKCmLpuReq6umngyX0ah9d5EgsL0KqmtiM753tKCvXJZ2+R39rDa7aTt2RwqKwgOCZAW0d2qHMViN1OEZgzWhJsGxTpeE8ooFO080x4POdOmsqi5w7ZKtpZpbFX1+Q2Gu0X/oysjS48k6iSUsmJx0YDPcco2KpI0K/NFU7OiKUrtwmguBeO1T/zBeefcLXZ7SH5wyebxAmot3yhzp/F8ukNns98EVR1XwMOiQ9QkMjzGhiq6LH2FuRcYSwYjoNQOzKhFAkz6CszGhIWsseXQVktYz0jQRdTV8DryamNTYqW3GiMfL2cMefUrFs7BUTcuklnmtckR4k31NNXxYhE43gn/ZBdM6pguqM8UJGwJ5ti1lptPAz2TWmZrLIK1GUKCR8J2EN+DMBsSQEIa1lo50+muZVQRzs7Vtw2da17tvuvqx76V6SSJRkDDJ1k2fm3YZvqXR9CGGBeaJ2ur8V09Fce+GEdpi3ncHV1xQGl70rY43MKHjD38f0hDeEiSyZaJ8ukw1VjVM4tuSHwfTyspG2qFhAqV+dgynj4nXYQYXbHU/AnC8PE5eYnAVZA8pXcVPrCx8MdovS2/n8+scB4HEowA8qj6hy+t+KnNarQL2JjML7ndLmq+629784CAFvI0kCndE33eimvkwuOVI+gHF7TAxAYxSbBhUFAq3mbqadIk52nqv6zkfNa7xwmxkf3ccKZVeGvURT2Ax+8EZ7GRh5HZU8NaMZZ7ANDDZKNOvMJZ0JT3q9T42T0rGr9mE1aok/H7U0s0dB65duw6P9MSZH83IZqmf2L5gPtKxbrT8eChdkX1d4U9fD7lvv+M623F6SNxBGwUyUsOHnWD0tC6pEeVUTx5WnlmlwirxjtrQUW9nQs3JFYnDZxQWYbw9vX7RSiLWnJsL8u0qlxoWYGTdZSFGUylNtZQy8pufS0DDqayAmHZ7E/KNvFc/uyLWLziKyC0aG0bBnlS0V6/zRBhXglRc+jFdB4KRJVG+i5EZE3kgNVjH9Q67UHos1eJGDEJVX+cmW9nlhzvo72i5JOGumD+5HqCgZUWhUXqij7LJx83nOcCUgYhYlJhDDOChA5X3uFZQEWaf3KGpgZjGDTNU90roPm1tPTUdkGShtCkhkiCqFWnzdw6YVOTh5zdULnUswIBUTQhDnQST1YlIVhNDkX0mOpuTlIISSIoxc5J8gV0L00Xnqay9Etz6bLBmBnsgYRKCnCChHUQpFx46rmvsaPHvhXWB37h15A4u8Ayd7wLf9qpmvU8Im0UvOxJSZr9n6cmrFeHZ076mhzUHRty4EGHhfH8CQIUSiWpDI9y35VWLoX4dYvSuuF/ntcemXio2s8BwZGXYXRzU4yow8UT6AGHx0nvC5CcJR8GKFSpBefXRkuMuOqH83uMgFQMrO64sggNrUCOFO50ONTCX9Q+gtopnaIyGn/wGc+ZRYIVZ/NVY3pbUTXhZp1ZtE5UC98pNwK1iOTjLoXp3KxpZRg805vmIIckteTU/b4PQUhBoq6j5H0xMxCkoqlxii3IJJPZjjR54p5yVGfTTJ1PHGBpkZyk2k5cKF/ir8ktlXBHzVjU0W+EEIm/oLkHFnQt5Liv7csCDuAm6G5f1pJhld+dS5bIkWYw64tkHMKlqgRXdImZgYlbCCEqja8fbB05rVNcSrQXoMsBfkDdzaOW8OMpM+hsnqAJRHMX0UXdEol4JMuKAhyzIWyOF+2y5s+lqw/cl79WLD3pBJDM23R0ZGtvu9/3qzY4+cagulEgbBXAZzmHtmhFlMa60OETTDrja3DkKbBCAQYeUJlYsLpvShyThHmTte8IhCkpA8ACvjHB8qPM6ynVbpOOMUyKRQQbfwSiuEm77flsMP5rd3dumP7q5nYW8Y+txlQ/b9KtlwAJB2EiFKI7felPBHM/wIwrEFNXuDpxAa0wwS4RnDrX87OOzcM5wYKMBN3z9ZHD6vEZatGoB3fuFh8QxxH6QcnVUoiz6z+2lKU5qjvSJPW07a6qh/1CCQaLSxV5wmJBJOBh4ZRQfaM5NycqpLIlHuUI6RMIt8LFuK7QMl+NFbfo0n7vN6sT3fh+/787nYV3pCZGNpKpUwujFUKInvla4BVR1GSgUYHMxjYzYhkm4CFWvTMeaGeqJ09EudnyhkYQT9hgnyzDf/pOaczgX9f9qDwwaLjEF19Yz9h6zGX2P1eAVYiriOon93WH1H6Y9kIIp5x86vvL3U//ilgg3LzvOCYdC87Oe3Pvl36tFP6efLuJ6UiNJ9B1FmebQxwiSFmF2CDYP/3NW4nImIVhudhWaHgjQJefzbDkEd9QpQTmafd8q65kcx1XqLqJ76aVUz7TyVnXQsZiccGmYnnRbWTPuk79b8pSBallq1R1+RSEwfp780SVjtjaNTwbz3O+gkPgAkpBIQkRyR0pRXgiivpo1BTCJu64ejJlXDW06aAWK4wHoqZdFd/tFVBXKqVNeQeM9pU+G4Q5tBegoOnlkLH37rNMD+HFgx4wVwJE8aoIZ5q5wpNja/XHGOqgz0+ij/jGTNSWM/TDDPKagKUC3nPCkxwnBmIzuBQjiWJQZygzh/+knq1Nlv4K/Zkm7GD7/uM9A3UgALLNZ4iFYXbo2WQgyXclCbmAufOfJXELdm4Yg/SEpUOv1hBIYo10nGiZJjKgBJlTc7gx3EoWHCBgF71uDzaZg9Pz0sG36jatpSXDPWky/Kk0aLg/lEuYIUghKW7Ht6SzpefI++YMz4vat12bzWbknhxlv19eKHLlrlR0Eq0eoQFZ5ZCIT+JInf9WUYY/nf0aggqvDodLHuI3khsRo7rw0Ns9N2et8/BFVTroGa8ftipqEEyWyJ2tIIq4PoAEoH0E0qzNR7kG2sx3TtRaVU88MyO/c8I2FlyjF7y2B3Du3la+ZBsnqWjrg5g1bO/JVzrZRcJw9g+wi8fp96vPeGd2Hn907DI6fVAQwWWHxF+9VcWtSsJcw57uPZ8yfSwdirpn9LedKFRxShNhp2XCA4ApDImF2glkWm8NO67rRPaoKIGqDoFDXKevR7GDEnjmVNG5ymN0PSknUtl/LVXBGgMMASEl3poi0ccCwbtw0N4Vn7L9DnyGy/AI2pGtKLjhidTYuGpqGXYGFfbhguOOBzePqst+GFB34FevNUfOAqBUc03LPHCUwjLa+kFXgyzDTKmS0nDB5Lp9x+7Z4WOlDeAMK+ruIVYc30/VlNjWnVTW6IkzCRka508EmGy0ZrcDM4kD+3b/OTmwAWlMWPn4/IlXpg6/Rd4z1GXdJ00zUjL2d8GLSgc4+G2Jd4KnY72XTeQYt1GGWGUbTL8w4+6VTjlKYBz71eNUw/Ah23xGyyRQIdUanYWydCfxkExYJlOXXoxA9QdqaG9SdjbgHshlZZTP5WArapoaUde2P4oKprDgMnNSqFoa9f5HuZqJgjAMThHLzt7Vq3gGr/mRgRl4QVYCKDC2THj3x3hfmCp0F65vGfM6MacbgIysmC6i2hGhzhfVvUz6IsgghjEJcC6gkaXKn5XcYiA3M5lF/V/REMVkRgCin6w7UcyBdzsC03CFWxFIRSUXcj2kqIkiqIVMzFnMrh6bMWwKn7nIohhMLixDbAiJcHxc1IrCxNFBCUkjKoeAmlUHJRO1ShOKDxUEhZLVAMesGSCS3LwkbA5Et1Jw8GgVCpJGJto5oLALd1L9uTFg+TN0jP+ryfmfxOiMc9CD2HK0kR+DrCGPAg7KUp31iuEsPbbWto1UeLuafv0fvq3IUeyXOHLdBKKTpIRKoQJXJDX1AXpYa3RklGanwy4hcsA3ekC/AAMZlG9XGTQzCItzJnn6l10b7UbnsZ9MaT5seLA4PXYc3kI9BxSsIvxuRwVwClkd8KhVcFjU1Lcc0/BqPDJBsObSoUek6EXNcXINWwDyZqfExWo1DBpTYeOBwMd35Xhw9a9HLPHt18DZSIH6Q59o0di6wwQ/zoiSbtNQVhGAor7eDB+7fwZkRH7hG3YdSIOipToqVPaalMuPDzG1fCG06YznVMP1Bw0uGt4nsXHQAf/99/wfzXT4MzTp8Khx3YAPU1caIqgzXbi/jXe7vgp/f2YNAQ132Our9B75z8cqd6dDJx+oBDWolckaLmWQE9w9vE9PQM9Z4TvgSHTDlEVBHjMYIY8obhsfWL4fFVi6EfhsVP33ol2pZDrLrlJ+in/7oC0vEYhJQ+KPM+6gmgD5aAtBsn2ndMuAmRitcIL9gOSVdrWVA1g8mluWkXyWZS+GBZrhAu4lEA8K1jLwW1k/7Hf9UYxGvmnOMnp34NM4w3sDVSNDLyFQU+tvnsOwEI15cjA64ztPr7xdzTP4z29UKPTAdJmfZm3elomOzA91iNzQgFabos/ZvQzYzSgeaZcdj6wLBeMqpYI4M34tqxqXuYcJNTCeS27vZUtBCFGOj9ElZPex3EkkXwvLgY2rzVLva8rZRbfRdvNvIkH4H0DBA7Va77oW0A8Nu6mUfeOLx10/+GoXehyrQGkKoOMQz/1xYH3hMMdTy6N4GVRBDETfqc/mQFEpNDjNJ0usUYFTh+CHWE2oxSuSy6YnwtWjIj1TJJglAKrfoUXHfHavGT3y3BD5w3l1iPuM34Y+/ZHw/cv04cdUgLE6JWeH84bVIGTpzXoJRcAj++q0dYLXEMic2J8xihWagrcgiEIIroHgmlbEkYKQ6Lfev2VX/78I1Qlayq/Lp8tgeN2w/hyHfx16ebFKgQyCjQuOz2r+J9a24W01takBiaSaWa45CQZjk1MwWQsFJYE2cdX8ZKO0JiybRQR5dVU77xpdTkkUynSHT5fgPsMWM+T+BEYsLBntP087BqnIIwrzuJjUqPieBHiVJ5lQ2ILssXpbwrB575+2cGlnyyQ7x471giEnUO7ZjFbUwjkwT0CeupG5o0Lx9FF8S9anNzk7CsWCybYoZaGg2zTwrAdgqjse5oKlpjkuiX3RKksFhlLNs2TcRqPoGxdCCC0JZDG3utcORkbQzoQnGMbbAG/EVNvmO+3bvigWFv6NkPOMXuL1u5bTaCFUC6WgpLfqeCT253LqFxf2j/fIyKn/Lf5n0jWfecbct/V+xnV6/tvO2uBgfc+srtKIoxynBjTjsUCorMWMIDZ8+oFRo1GIVrUZ8CraPEdxOCnFyNH/72fXDVn58i6K8ilB95CvMPGYehT8zvjBeJnjS+ZtuGC7Bha559ylHSXaOuzDmE/A731jgRBiIkoVDycHrDDGGMwc7PARa8AhSDIgYqIHZXBiR1DW7BL9z4Ofjpg98QUxtboegxLwhP8tDgjSgMoO8Ql9XQnNW5kUApkfeHtYNCFQwOtRm/pQVg2aswkI6ADcMeEiq0c0WB5oGfmvQ3VT89SfQvJgU7WsXTifoIPq4zstIOhOe7sueJhzODS97awU8Ks3G9qOIeI1LKk4XCAQ7RyOxStYtTxoZSzRA4krFgKILjqngsFe2o6vVH+b2/WZkzTDl6RPyZnGRiD3I3F54RWEFgpd4DiXpaHUsw0h2TxZ4P+7k1SzUoaZHuaX3uwNHSCVrF4Wcvi4F1mHIyb8R4qhS62WPsmoOOCuAxU4N9jsWsSJ/vqu+18wX8/u/ee973n5vssWC7CAMKBiuv1yhGgHUetRaLXwrhsRXdYtoUkp0AePub94XLf/uwNtqRSK9uJNH3g+C9joAw7cJnO/6JRx82DmZOqGFRqWIxgHhcI6TXd43Ag4u3iTuX9sPq9cP4VG9ObAssEPVJLeMWnY1mUtcBuhkhk2dHJkOAHwRQm2mAG5bdIvovPwdnNM8EIQNoSjbBhIYJcETb4dBa3cr7I2UnWohscMTtz9wBX7/rf+DYfadyKCSYXNHQpxPFnCJuBBt6CkNi/8bjoSZeCwpCMVDsh6FiXjQkiRfNJN8MA4vpFONkI0PykTwnfeqXAuB/MWJgDzkx84RWr690s6rbpxmRBCtIvzpK3JXBh1EPiakIWaEIlGP1r1yRSK07tW+QwEe7eK5ewLCJuWy0jEGHotUBqcmebACZnlE8u2G3Yk5MO+7IZG3k+4nzfnixf9lvrxrhVjd+kKlqodnJNT8fVzx2YxAWBQsXLpSX/eDvb1LUV+EFMcx13enn1vzRuFC7MwaVw1RFEFQBPyvk1hNUbKqFdgqw0PtOALh317qOgPPmgfP06nGzlV1THdBJBwHaDlfLNRKrTArBIoABWAmHkEAkSQwiQBZ/Y64ZkhVzLWo9o02dkKpb3AgUWDLmQhBAqEol245Z1JoTjOTXQeHZzaNGiUVfwS50PRQkmy8CO0N5A+IrM8BbDgGidkQkeThSjlv86EZoP2VfKBYCcejcCXjeqfvi725eLtxZ48Ab8Ubl2ujOSSkI8Vqb9+De694OsybUaPAiyXDEbXxgabf4zi+ehjuf7BEDBGRMOAhxV0AmIbjEyZ2SPGOjxifBUOEdGJMUBqEAm2CCBj8ReB5kUhm8e909ctGau6hzEYjgNAg80ZxuhsOmHIofPekTYt7keUTPTrRo4vzDzochbxC+cdclsE9Lk/DoRCOiZmZZpgpFKLaPBPDeUy/WWFuwxMObH4TBQhdMSDeLEnrcAqEDhSjnoQ0qTyspSM/nPwHMvZRDEPVXW9vvYmu7e/8U1h84E5CalpBQVjpLwnG68QZ05EDzip5Ryvk69sCKnnR9/i39zwz3/n/CY5sa4zX1UkXZgnpK/QLXaAz4HBl1rMVHyTCEEEs4wYhnYq/321+xpC+S07ohHtFymW7zqPtd00Tv6sLzyX/tW785CFOtbWA7PuS2O1Dq/+1/QGLK+APf73hK4n73isA7AWMJEMNw0Fva263OTuYciAavmbGq6VOfeMbpVHbyICpl8he2GOG7M0uloZ/nxJ6xo0xJbAoqtE2MV0f9MQmliq/PipgyBghxCHQrOIATG5axQz+nBh6+wsi66evj+4+JUq+PiUwF+w0rvJmsslZ7C2lGNabg2pufwi9/8gTp2BYGgcKr/ud00b91AG55eDPKmS2UTtRuMvnPMQtxUz9c9Z03wuwZdYL4Bmi22LYlrvzVErj4x0sxqE0LaE2xlJzh0BEsRcfGIIJPG+RqGRezA0OKXsNNpiMiVQ+UgqpsLVOoU+8BkSgT8LWEBbhtwy3ipu/dDD9YcIV6x+vOQy/0ZKA8+MjRF4sN/ZvgL09fjhNrWoD0HRzShuS27lA8uWUjfPjwb+AREw6HUlhkubh/rPmryLgWeErxI8wkYLo/lx0nEoPV8jYCPR8hKFlb6HwXaMTiK51nMlm2DlzXNfsPYcuco8C2fQh92zSbGSraivyBDgNZ6Y7ERe2BlYNxseXM/mfWP/X/zZVJQ5w7SrlEB6RF0S9pHLRt+hnopLiZnAmxAGyS8RY6ZGjbILk8XfR6+FwYF6YbqLVx0F6GJF91NwMLuRaQjsNpn2A4dOOxFfodUmx+MeMmzjMoEbvTaAyG4CYn3XzXc6Tc+GkOMXmpqppyEFZPKkG6IYRMg4J0Y4jp5gAzzSFmmkJIN/vmd4XZ1gAzLSFmmxRkzTbpphD4fXp9XAhV40J6j/ZD+8NMC/+OmUYF2dYQshNDqJrsQUNbBpT3hfHjDyfFbUUPBbvNIw8sA2/oEWCdWV0zK9eey6IaFIwpsJIxWLu+V/74l/eh40gMQiWJNemvf7wAPvKOAwWs7IJg83ZUg0VhJ22hegYF4RDOPmmGKPk6B0zG4PKfPQ4XffNfoGbUgd2QoHYBFRYUKYYxlyL3IJTxKYZhh8MEo5ZU2vF2Vmxm2A00fTcZsSAIiM9QeL4PRWbFlqIx2wDN42rhomvfK36+6JeSsAqasCuAS0/+MjbGpkMhzAnHtiAf5mF9/zaxticHnzjiO/ip+Z8VpaAIZAxW9a4UD67/G9QnalmchY6vcQhcVdAN35SDCDmHjn5JQN9WdR+ddPcrX3aUzIFBHNiJmb9XLfufBbGkB8QFqpvIDC2hQXVpQ6DnFDEZSsu2B9cO2+HW04e71//L9PP8vwyaTY0tmtc/YsBmiyDAzyP4RS0xTZRU5LoaV4XvlJMAma1v5Zf8Jg1VcoSv2GgwYkTXz/mZ0ddZqdJuYxqZTWVCbrLxbQiKvQ70ryjyOzus6i9gsM4BOpnk6iAo0HnSkp/O+/275IAXIlgEweA7hGPHyniPyr4tPjnKppbfkJrnhb+X6Qqr2N4gfssyIuYzzGeruwJNJivQ+1Sqb1OTH8CmKGyYTyQPgdV0+i8wKB2JdlqBIDCKSeRFSUL27CwWa5HjanHht/8JJ86fKebOacVcvgRxNwY//Opp8I6z98ff/e1J8fC/uuDB1b0K8gXxmQvfpOMHhUBd1dffthIv/t5DYM9rgbAUUGc8fWf9PTk6MfeQv1LUEWlOl5quIAQ3rUTZR+DoVS/LGrhsVkFTB9HfhN/UXoYivFSIJMEyeeI4uOTGj8H+E+aIQ6YeyonGVCyNHz76s3Dh9efBhNpaaEm1iTfudzK+Ze65OLtxFhGqUt6NL/RX7vi0iFsehipNWnhaWUDzNbLnF8EviBo0BGEN9AR+ruTcQZ9d9Mo2xnGlDKAjFIkZvwua570VYlkP/DxBwMroyoiI0JAXaf0L6SqhfNseWJl3wq72fO+6e2HSpDgcyqF1mcmsvSJbRb//u0HbEp7N8KRH/jHpcFPY7AN4eYB4DbBUF3OyVnabWYC2nMkvOtt0/jHublFIj0Xc9OJX9NszjQJj3XZ9dSwrXYZlEyH4MJe+/uPhCMyROiwR80FYtGGwa2chCW7N9IeX/9yS+0/DHOUZfEqh6AyO5LY8wtNpPjH2lyPtc5rQtAuLJdB0UGqVgRrMBoCEGtJJeo5gDaepZjBi8Qu6Ll0yU/VBtfgRupHGgizS5yW3/MEaSXwmjM+cTgZCi+qUvYNRd51+taTIp1PqzHf+Utx+7XthxrQmKJV88EOAQw6cSD98toseWCv+cNsKOHi/8Uyp6DhMuyY++72HQUytJu5Gk32Mul8reqTKXP6GmJWNUlRsCMEjYV0zWBatoheLexlopaY9m8gnSuwz56qWWOKatQwtSFXF4au3fF385aK/oM1SeQDHzZgvvjj/G/DmA94M46vGI4nG0seLfhHjjiZd/fTfP42PbL0VZzQ0Sy/wCbpM/Q6GK1ifOv1Ot5jNFQh7aJu4+Zkfbntq4UKQHR2vmEEwF7BT2anpvw6b5r4N4hkf/IJrlLhHr7l+TiJwGrWshxgo2+l/psu1eheM9Kxj7wbWry/C+heX5t7VsE14O8q+y6dhAEi5YYDaJgAgBSCdzOIcLT3y5N0rOZ734h6GALeCKJUeArsEJLRWgVowBWmd4d3diWCx1I+MpRcB2HY2XzVuEgwO9muXil3pFwPuEcVCfiq6jRrcg8qLV9f6xYGh5xyWVqewf8nnaXEpv1YBBDX8v+U/xwGITaOVifL3aTa8oNvMVGoGENRpRZAh0x0C9B5t1xUlU4Qoqlz5q6kdlKS6OvNQV/sBMbDxH1g3kSw0VdF1TqHcU6LPgAQXZdaV60cKcMwbr8CffPtseNNpc3Ucjwh+KSAhFzj08Amw2UEkj5AmimtLeGjxFljZNQJyVqMggxAxcZW/GQOQoomvBRjZ2HMIUVF2rBwMg9EOZSSzUI52mHTLPE6G2kMz/OvLWQpCIMDSQ+vvh1XbVomZLTMgRCXqU43wtsPOg4nZCeCpEvqKcgkOkDHYPLhZfP62j+M9668XMxqaoeAFjFXgdhzuuzHwW3Ju+NRQyZgQvRtVKb9ZfIHsecelL/wJ+38Ofp7a20Fe//dZv1YN+74NY9UlCItEc15JRzqKPOJLQ9bMpqy0bfUsXRGr6j9jeO3alZPfuN8ZSviNwilQFt8lVIC0QZJjL10RYkC2OJTSlr4tSQUPSY3Pp+dByIAyQbRMqeKQR6LIji1o5aNsE2WeDBcrnwd9YGQ7gJypy/hMn1Z+V3e6x1OUQgRIdOm1I983AImJ0ReqcA/MvJKxXRKv8Wctu4vOBS0rBLcmBoWRg0nKDeDuF1k+4Uw9MXvuC24ChAqkENaKOVNP2rJ48VWj3++5N0lHKBWXYFe/b9rN66bNsjyiv0df38V22gDv4vvp5qywt+NOR510Lkrnt0G2yQYSnqHeYMITc3TJ7gkve6zfkE7BNjsGZ13UCW+94Ql817mHw8EHjoc6o9Xw0OMbYe2SDSgPmcqhAZUBbn5gvRBxWrmV0dk1ZCt8ZhGNtyHMY0PARsEYDkO2QFxG5WQU+14YKpsb7pUiiSeNZdepMe0rSIvxBGXnOGJeZ6IUgkurUN2/9l8ws2WGCJkRlrIpHB4rV8YoyYUrti6HPy+9Fv781DWIVh9Mrx0PPhWBmHKNs10atazrjMzwDDb4TkzKwgDafVvCjyy+YtuT7d1gdXa+IslETvhNmj8pfv3fa/4U1s88AxPpEgRFlxP50SlE2gl86alXk/rbnUAWcjG7b9Vj6fruM/qe3bSZnhF0fvXFpn1ShyA4jFTlXhQ3Is3SIaqO/sjxp39DTabFsp9a5Ilut23HYLA7BFtYsgCASb0m8h64lMSuy8jgaJMry3sZym1W2wgBXGdyzdT2qv7FVzGUWNRMHZTgBSHd0coGOUP7pVn2nyPUwi81Vzcu3pIfXgsqmEL5CZTiXAC4OirFvcDBmfoqqKoeEuJNQK10xRKtrssWL77Kf56+hgpkXuXhXuzvzzd2u91ujB1Brefbfv/tf4rF39SvBtWPVKp+OpFj6pkjPKbRj9hJhYXKD6g0CGJqA/zxX5vxj3f+XrQ0VsPB+zRAXSaJt/5jOXz5kuP1emNOp2trDqkTijG97INwnDpKlBJVSnTORBsL/ptbpXR1iiaqb8jr9FNUhlJFKBcGB1FeXIsEkpei2RW5HZlhcdpFYJ5eSX0OYtnmp9hCEX9iyk3gzY/fBjcv/Su0NrXgtqF1YnnPUlB2DluraoUt66Hk+9ylT+hmet6pKZLRL/S1HCEF5bQlut0b/NLIZvjYY9/edmV7+ytlDNr52Rs/e3btliecP4WN+5yAnEAskTGooLMro/mMfCL1djuBKBRjVv/T9ybU0jP7noX+efPmOYsXd/i2M76YTEEY+jIkz4AL5dRBStOVQiMiliHoENUELPIeDDEiU1oKahNknEoibSk/h6RfrrpMcGt4r8xpUR5tqNeQ5pSTnfoImmQRUVoT+rGqKfoa0084ayX4uVWGuN1MMv0wGdXuXQ06qr1p04MFUej7I5SGBUqnBLHq463MlNP0BCZg0r8dXPugiZRP1X9WJOoaIUQf/CFKJv3uBXwe/0s/zzOY9MUqdd1wu2pOHixz3ReKXM99YmQgJzzlsnEu00xw9plLg6oUgNWYBjGpAbqkEH97rEv86t5VYutgCClW6Kp0b6KISCMZR8NGUymIlIB45Y5mudFj0BVFzXzojkKXNY+rLoxpZmWj725cBe3CRQBKWv64a47gbvzZsv64WSMjt9JXRfjnxlvEP9f9BdfmlkJzXVJMrG1mGI0X+Logq7u3Be2T8iR0iUIPVH5Q9Q1uDLf0PK3+0r1Mvu6xb2+7gvIGr4wxmM88BFVNTZO7t6b/EdTvewImkh4ozy3zIZbzBaONbMwIZ8eUGBmIuUMrf9c0Y+nJg4OinzyN0047jc+bYOGWbuykJAz/SxeUfAqm3SNKSVJIYk4NtFj8WtH2gji4aV5ZQgkbPbQhRFsq39MqdvpmG5oZReqXAob7iJA1ArYYxCIZeEt7EERbW9tSo7/zQmt5Z4eHgTfMxXiDETXUTkycF8XSu8AW8B0P3fzP5UiPB9KWKtOi0Gn6abxm9kTNlsRGYVdQX6GtL5dvPCe774Iw3vxpdJO+UL4rC31LvvihM2/T2+1dXY96GG6Hhy8fUht+eRWu+vHRjsB5omflx0VhqMhIUj2L9CqvCQtF6AUCvZCSUGDVpUS8pVoQGeuG3hGTZNd7T2VjTBKgPQLGG+vGJe5wNsQHvMiW7YihQtIs6BBySCYqlZsUi7Pr9kJdnFJEq6rbkPWk1fQHjGky27HDTI6ClnYnp6OpqllXr/TZik3Dm+T42jrRmm2SKTcp/FBR2VKHK1FBrpye5/hWuXEphraVHul5PJj91PXePg8s7Dr7mau7FjOp6iuSRJzPvQlOatz+w/7ke/y6fQ5Ex/HBKzmG+FpnWHU4btwxbiBUKC1pb99su31PL/R6Fr9904OCvHmNlTA5D0ndSCS/ASEVTqjpjVRZ6akQQoRCWopo65khRnAegjwCutlh5ClQCoFRh0Dbhl6xr9xTHZW2IixCYQihNMLSuFpE0mDodRASQpxUkuIH8WubH+K4VtrOEkHlvigTpV7QYqi5Cwa719jewLet4oCDtuNjprW1hKnb7Ezb4doolG+gSYUzCQr3cYPoUDIz56Ohnf6dSjciCgdkvkdYwcAXOzq403EP63V/MYPPXxs+RPBW/GCF6vrd9y2veIthVA91ZjoSRDD5W934j2EQMMQZXUs89pSuCFlaVwPmz20QWPS0Xnd5aa9otdbuf5QO1HybvIDQ7CYtdbPkl0ahy7oUVUGjMbrYmKdIg+9CEnc1XgMZDL3C88lzB/frZ3J4A1Gl4bGNj4MlbfB8j5OiKpSoQtOxrymaI2+JMxu0BFKtJ/Dliqd+3r2t71+9w6wFSp7By8+OJMpdi/GGo5Q77p+qfp8JKDDgNmbDSlRmO4pEkpkMR5DasiX71o+4+WfeVRx45iuoaNHTsVrlQYh0m0MEksqW/H05QSsIxkrtSUwVpRGelEcg9kFaAdhIWGyJWCTdULgTZ3V8hMptWqG53JooWKglKAEM9gqIcVmHMhK66sYagFRdTAEmMjMgwiLQDS0NL0b6nA5udaxAFp7+//wGmSfttFPdDjmy7nYoDsXAsktYNX6f0Km/x0rP+YFTfcBcE3qYJ7dD1bYdmnUTU8+yMnPuwUTD91VVqwTLDmWh15HFrd/w8htu2ps6HZ9naMNHY/ZCl4yD8HM3CL9o4rEy7b0muzElT7PCC6Ijh0wcHn9qixgp+MKyLZqE4g3HTYX6dAxUzhckI8+jTMxlPIKIJTciba008GWygwqkIqmwGK5PSgeax67MzqqzScYpHl0kqRkWXDsJ3f094oS2E8TciQeIIAzBsRzYPLgV1vWtgbgbF76iIrCuWpETwx2L/MhTmVE7OQRRpLUyN4BipAdu5GtwEDhAXsHL7xlI/c+iIFY1/f1+esbtYcOseqaPBcq26gVZz40IdER/k3ZCLJABOFbfipXxYNUx+cF110Qhxy5XVcN3TLUc0q6hSJ+YSig3o6eqZpSiW8tpQELlWCQXaqayJUa3o7jNjjkDGoognovIIXvTtwWAxTWNclN5O2qXBoBk1RH8wrt+xU+ETFU/IUpDlBrlfIXejwa1KYNl2A2FGj9lyzuX+5n6dLsY2fB3KPTHhOUEUDNJYN2MiwMnuxhq5j0hUrPuE8lZd0L2gHv6e9UTYdWkv6jafY6GTDMJFYDM97pY6v+uP7KWyon0lO8FXAgvYiy/1OfuUA9ug1x3N7v0tLTqeh+NyPUsU+ER9FimYrBu0wDcfOsyKsuBVwqxrioOl19yNOCz24mcV9/iihVdkwmYxGVkIKJgnXMY5dfLA0MNwWA2PpMriH5GZR41f46Wh9FYBtKm7Mv3QTJIw3fP+RY/6AGSvCDAzU/eLPpy28ASMeqBoB4O8hBAsRKV7stjxajQvK6XRTmy3X+20Ofeys/q4tEL9DIOi77hQlwonOycb3vJKT8N66bHQXmEj9QWl/HUfKWi/I3uVad8Qa7Pln1PXNc0cf2R+Z71TwA+P5+BsHR5hxKHPOFZRsXQoPLfmo2AAv3IIOjg31DXcbXB3HcW18JwHftnfOc0lFW7m6z0AqJni1HKKGeaI4Oga9DSnQa1bRno0FQHDX1PLpUq2G7QaeZJieiB/+1gr6Jv1cNDar+aM+zhNd+XI1tsDH1bUQhR1YJYNWkuNuzzOmycfRzUTz8aa6ZMCVONRZRWCH7RtQc3KDmy5rPY9+gnTdPW7njk9uJB/uV8a2TrL3tEceQ7ImCNDaLXGs3I0dCUZ0aPgXkWQNRn4LKf3o/EamzZUviegreesQ9c8aWjIHyqC5Tvgx2jp4T2Ff1EbSjlbKExCNTWGACQ90EBfTRCS7ArXz4VjXCOcEz63EZJtHjVijm4dbALM0ECr/vItTClaSq3QjvSgpFSDq6463KozdZAqRRqRvqAiKFo8muCqCAyEGQYmCdBBqW8EoVB1bGpc1OB6v6vwHNg07dPpZoav15z8y1B7exPIkHVQ+IQiqTWzGViVTTTsCYtn9Zra/uzwh5e9pWw98m3bHlki2lSen5yE/YxKM/PPpaR8Iy0dcprvI78ygyJEQuiMGGFNiAcXsmw1L8MFZGhMAWN6Ww0N8x2APu3IoY+EfjpIrHOH5g6dKhEpr7Wqj7t4Eg7YQtiAVTpYeqWrIwRIufwBTQsaV930aLAz635uB0MnGT1P3O/HNhgycKgzQ9gdHz2wDwQpeG4GNokxeAztzj+liODkXXfZBDPC8rk762DEY1SDa/+GQxtWgtIQo62V16xKZtX6d4rRWQoIKsSsOzZXvjOVQ+AY0sC/UAQKrjonfPw2u+/ESfnihCsJfwJHSMiaTH7ikqMzNNocg06cUxUTeUzK3c2mTxn2ZkwxkE3V2pAE6HCc4U8bN20RRzXerS49/N3iyOmHYpe4PHuKWfw5Rs6YM3As5BwScmZEg/GNBm7VNnmQbkIIRXVQN2R9eG1y7/f/ftXoLQojGcQJGrbDi+60+7za6adjDHXB1XS8dNzn0OjRe8GJFHobF++PuatOd3ve3YhIufG6DPPc86mUTvChdP05N8NKjhSUYx+Kv1/ulx6CmkebWNI6HU7zA0NCId6FaKtK3D4TgxhpA8gPwRgJ2k1KNOmGh8jxFSNgw0T22AN3AmT9rWEEL4Vm/wssBFJKkBPOzJRheuFOe/mKChLfU/dAULeYdfsd6jq3XqMlaw7NAR7qr7QDrWobAF/6H4r7vzDH1qxuKCfuldDzuCFicMOdvRjdsoC0bfhz5htmgQxopOQoQESaTyIdiq1urLngzW+Fr747bvwoP1b4KSj2qBE/QsBivaT2vD4w8dDx3ceFD/+51rE5rRGL+qFQj/QjEUoowtMopFp/suhJncPGYgCo9yJhifkmrVukuSmWQtsKmSNDOOMqolw2du+CKfOewN/vugVwbZtTib++PYr8cf3/AgmjR+PpWLABS6yCXwiBhtVhjtIQt1KSwXo5rrwjnhxwvsAuwi2+zJSoLRrbIuA0K2aeXHJafyWqpoUA6KaUb6tKUeiFK2RVmNjKkP6knKkR8qhtddXOesu6unPbzWl8xcc2miZFNPGyd4BC67x+h7J5QhJ1QfNahhxJBvYcQUwQOd6bWfcNCrvUoDj6nqoEfvgZ8gSEBQB+rcCtM5CKNCqbwpWZJLIX3PjgNmawwHgKkhN5unuJuy7S6H3CQRieNcpZw2Jd4yV3AGYtLvBSxNPblRh0LfkYQB4GErrdrC1WsoLQGmOz8gQvtqNQUX1YaGEjR2PJia1H17qy/8QE9VvxlS9jQ7B1ahXnucLZaj5E7yeE5liaxWe84Fr4cdffyOcewZxjJLbjSKbcOGHHceqQYXi13etA7sljQGtt1otKUIHRM8UszoL5YtKXQZ6FKkfixlMmNKEsloEX9RQRBZxohaTgJWk4NqP/x6mNUzCgl8UCcsVcVcDYL/yp6+Jb9z5LRzfOhHCIs11ak7QdVVD7UkrXWgR6NcCG31h/197XwJnZ1We/57zLXeZfSaTyUp2QsJO2AWDWhW1dQ9WrdXaulSrtXWlasPYVlqtrdU/tlJR3JVxFxCrqIMoAoadJJAFsk0ymf3u91vO+f/e9z3n+77JRoAEBHJ+DLlz595vP+95l+d9nvJIWIkmov/Y8oXzP0YT9bMHXJ2PVBXBQfLSrsVdHdXReZ8N2xa8QRe6kJYKDZBHl51lDRiHaXgEqTkprHv+xK6SG439Q6304BUjvFCSl/FoDgI5IRGZSARnlEAweQHj7zvULIKNyqh7ScQQZJMckutLNBLoYiJuWc4+bel20PWdhvmMHa+s/pHjaxjejp0wacBn6sJkkwmPUFiFWgKw3pxMs34nBJMlgtkaaUj8caWdp4+K44AQe4lU2763ln93AFb6tpUUnlGDjUJt28CeeOgrlzjNyVOd0tCHZWn3OlGdxBwyItsZGGJur44UCN8Rkz1t8Lr3/Qhe/pffEL+8aQvUgpBaqMu1QGzZOg5Iv07VCcaSmNwB5RWwPMZlSgoZ8LanrMtErGrAR5Y/haELzHfE7j43P+EjO1kaxY+LgpfXY9Vx/aNbrxd/9E8vgn/++b/rWbPmCnQ2rWxIsksJkfBx2kk3mlJuZWt9aPyeyufqDzTO2fKFXWvJGEzLYx3RIen/YjDKt8569tTE4lvDnhVv0MXuCDTOzERXxATsZu6hgXY8R9YnPW/igRtavF3nVUcevIK4Xh61JulaPhDDOmKJ0U3dgrUUWVibssNWeDsNDxgyyjPZcltp6W6/7hsTonjyELTqxdQfSubAZiKwnJQXsPshowBkfDTcjGlyQmicbuk6EUozlwL0bwBY69Yb/TuF03o/xPF5phBqw0rxKDwEe+EVJVYEQHfX0vZSoW+lrk72OrEXqdb2ivTbR4I3rXoQ+vsDE1cl0qbwjBkGp6BBhDu+cj8AYGf/x51Fb3+urlf+Hlq6X6Jb2rj4Ty48EgGg4pMAWNQrfnj7LvjhTd+GZcd1wqK5M2H7eENsnGwyXVoYc4ugAc9Nd72thqSezo+CTgOXKZmjF/MGJPuB0BnD94FwIjQWwtUv/8zr4ORZy+nwHhrfCVsnd4DbnhMzZ83TYRBS5pz6axmvA9InmgS3MRzVopL6oa5F34Jx+M3QjbvHTK+vAwP8dB6Fi+3iKo4ox8s/tewfwuK8y3THHFdL1ExoYuu6xY7ariSOobCCENVcOTVUdqOxS4OJjVeMj6bbe/SHYXMIaXEYHTOjq2aLmtR4ayWz2VUnupy0OY8/arW/WcMNaTViigNNHo7bUTlp5BUEjO4CUZ7QOpcHESHfDA37TwTFTg96Z50Dm2AjLHjYhW1EAXszqOZ54PoKsRho2ikefZSx2cqVK/1NQ4XXRm7+tRPSPQWUOwvycxiLj+i2ZlCH/775AdF11g/dYOp7YfXBe8wGHhOn3FN4mCuCCalfSdCDcbz1f34BAL9wZ772o3HXvH49cz4u2VjHc7XDeh+ou+bM6UJ/EjbVItj04KiAoq9FT15o7J9OIA422OQ23CQ6I4/TAe0nfLsAcYBVTvMwUv8xAW/YCTUMgQQvjnUulxcVEcHPt/2eZNuLxVbonTmbEhAqCAzNASUdyN1Gjp1wLHLD0fAb9a3xv4zfvGd9sl80BCtBQ/9RCRkl/V9A1Na9cPk/fbrnSt09/9k6365BhwoUkhWjNBZaLUs2hElDHWGsLCrD0i3t+Lmfq723OvbQPZnw9jGWQtFD6CdCK0uMTrbBlhhN/2FaHDQa7tZ+mK58/gy1m2IdkvmxpKNvV3FttXbb8cSMI2HsDoYEcRP0ni0Ax58JGnsDECupTY4B73y+HUS++Eca4GpoqdEklLr6E9Uce6/25qDbQdnItLnpEb0DMga5lhXPe2Co8EnVOut07RVsPwVeYAwoiYVDgPaFgtN01H5a1Gj5B0fkvhbr4Uuhunf4GWgUrLdgzpmYnSHa+81/AvnnvxdSfEPPmNdJIGL065lPH70AyhHIvNSiFfujsXsy4olsqsum85Fo120WyoQgxK/nB5XUSZAuc+1aVXLbVU2SCsyUQXkfvD1UNBLQnush6RW0FM1mROsY4vQZo4A0oqBFTorKlmYzGI7eMX790NV8ioYQBD2Co1dJcPEy4hXItx3/11XRd7nqnd8BLlJjBgz5sVSXrGmKszLGJmQnLHswtX2Hq0r/3CxvuTKcmuYVPG4PBt1v/Je7RqwJ4BwgEfFmPst3kw0DNX1ZjsnUz6ADg7g2cT/kmwCoHiyMpBvh2G2XiQew7X4BJ55n8hBoyxGzgNAF5ZBAe1vn87q6/qhjYv0AnXK84OzfiB33b4e4byG4mEw8XOQwAzHcluXvjdrm/rtqnYFH0IAwyoNqonFyIY5QpBY0Qqp9NBQegJdv6txsVxW7/0KWcs+RLXNfF+2985anonLTkRvmvDG/s+7KnzjyNS+JQVyve+a2sxybIVUwTiUpI+F9t6zf5FmSe28YXgwlHiRK0bbCAIGXocczrj3iAojzgFMJ6KmaFYvsgtEQ59f0wBk+WVJ55twTyYORnl3eEaUN5Vr5gfAVzXV7fwarwYWLCHV4NO+tY1FDXt8pJ+nI/WQzN/tifCYxBQtRlIYIPLfwxGMQHvZ4uLK8o+kEI1e5bfV/ru3YvhtWYJ4LwTq9CioVD1pbpxuEi8y/uysCHlin4aIsgG9wH4SlCRmMcrZhnbAAIZtKMH1TSU7Q5gJMEwJpgXPNiG4C8XphRbFrCvGSlP1ItBn4mwT9yrcCDD2oRb0M2kEOb1olLPGjJBRWx8xZpTnnnAMTP/8ZLL3Yh/XfabqtC3+kosq7Y6+HT0S6j6TuSEAMWVz6IdW+6HLV3huRkkjUzIvq3pIMy9cp4VynaqOT1MsWa092z79IBc2LRaFrRdwyA7TrNETHwoWqtPM6v+f884KxgQeeOspNR2lQ6/dbvWjoyt868KpXKaWu1zPmOVRGNjq8NuwkD4AUVGklMFUiUgYyj5PxHhOkNH4e30+7HUHo2KqEEzkv0qdT+gJ5khFaa5KUTF6C2GPLlIJFCY5DTLMfTjCZc3R1eyWsbqy8tHnH1C/greDBlRA+utz0oxoiSfJpEH7vqX8XxW0fUx2zW8HLY64ALSZ7JkzbyhAr4ZB2tVMbAaiOfKNNNy+dmHhge4D9iTjWr89mWvYf+57P4OAjhgxEIGoEcdgfIKp06l9GAIS1VQYKkbEdVOAhWDMbFWoQ5ZChxZtaN9WcrEFhXjFR+kkzBdzaOj4CMLQJxHGngK6UkLeL/UF6KFSs22ZI1TXrBQDwf5ZByRHqW6ox8i6Z72IlUKSNPvi0pBuQa195cVyce7lq7Q1FHICoTfiyOfXNoow+Wpq6e4v9sF0W4pGhHy29+OJLd/7mgTU6rFyu2ubN047XgPa5XfHE5u91w9LzxqG/moIsnqnjyhA9hXjdlTe6C/7qb+NG53/rYi4EBR7RzJHnT8VrboNm/AD7nSY3xoUsyiswbTE9bQbNKNIEkVBODOgdcIGQJIBsVwVlE9gekMlI+hnQElB5jLKHXCzFeCHvxMHehh9vC99OxmCVMQZHZwguJRI6MC60zzozmjHvP6LinAt1AfmEMXEYMiTfZufRsDluJLTyZG0SZHP8955T/kht/P6fNl944Rly5/JXgAsN8HG6NVqoLCBjKR3pMERTKOljD0HsKGKO4auBaRl8mct7seNATkD027GvD91IzVnkKaQKEiZ5yJLfxPpHjR22ncXU/Hn9ZmY/rgtAEsVhZoQ1MskgzJ59wZ6pbRtHIGou4M7GlFE0kYJ381pvvVuIZWcDVEumfmE7TXFLEkTPzJfOhL61w/dfVsPsQ+MMuF2u+9o9oMNTuCphrcx+VQbaybx58wq7q22fVi29SmKbSnlvTlT2fDCqbvhESdtwIkuYgkxxa2DzDQPYTfU1yDdukkHzx7pr8Sk6X2jojuNWlvWWD0MJPvgMASsdhqew1o229f+PdN/xCu0uegEyCIGKPQ44Lee/XSANdyOzqFsKfisiZELLTMrCDM3dRtRVS2lFfA45T20DTXq+DEkKp8Epx2TQ8CRLTaQmcVxTfmNL7bvlX+2+iozBuqNmDBx+PgajOXPm9Iw1Znww9LrfHbfPynEFARmwcQE1+XjK3rlo6lzRrHqyMvSwG4x9vPmeTV8K+3F6rZX12ve/COcvPhVZAbWnsOIIqLSH9hdTcyhBgnt1HA05HylfOa73RJRg/7xYke7v5IY9VwHAjUbUKLngiWyLka+x8BDTpM54aVOyTSgymRg1oeYki2FqCCRLtmHDdwIpxG1Gj8R2OmV6L2KgsGHbRoD6JALPM0BpClck8nXrrjnLhlf97Wray9JbPRjsj6SOvwTVUUPdTuQcB5O91kPjrZfo4ozlGnnpq5M5Wd3zOTQG3OCBx4qWG+Ni+wP2X0F8CY3d21V5ywuc8S0Pyij2tdcSxrmZf+O3nYxksAyLOfAQT6GfxznWM5SzOfpeWdqNIHvsEEqZjyhkwCUc/6Uf7hrCf7kvgv8lzgR+33bFJ8OwqxrIgsZNxciBEJM0O/5uJNkJpsOMbUoS1yqKsCJFeqxQzdlxmjsaY17Y8h469z8+KgbdSurFq9eudt3Ok98x3Fx4e9h+/Puj9vmY/EJjgOpJRnxDI1wyElJKGdVdZ2L7Dqey/QMdi0ZPbU5t+l9AY7Aan9d+pXOyqVtzkS54dcj5Qey5Qex4QeT4zchxAiVloEGGsRBBqEQQRTII8UfJIFZOEMciDKWoBFqGWjrMYnvifgxidPmY34oqHAxg5rlrG9kpKWioroyzx+la21rBwQQhmldhRRhcKe8XqpmpKJocMSnsIhedL6BRAnj4Xg3oPlFdyRSoyZOMtW7rAzFjLlKfaZj7Qbp5ve09XxfNyb2IeFTkfRwUlw/azf2V9otaRJEHpd1bwu7aexkXQRbxEbTtLYlKdTinS293KsMYrGqday/GUqJyExycYCXbj7fvjzjAz77vH+xzh/tzqP3s9/M4eR3QgK5xwp3X3OdUd1wtQnSBMVbGQhCWF7BjKNSmcUmQbpvCH5rV2FXEeSUyIvgezm66BRljhTgoQ6AYU/6AiFJIhMrSNiD5J3dECoQ1I5cBrnK46ThGKg8Zx2WQ0Uj8X+ODm3fCanCOcNtyAgbCZEWhc8lLb/5s7da4ddEVcdfiRcrLhUKHuLQ6gLBK6SssoSN2x6mXPLe862G/sm1tXuw4Mxr5/SfHb9tcslUd+NWv6HkWPtIZIVMRpt0lwqh8LZQHoHKo8KOV9oWOUSPMV1r4GL7Ra+Xg7x4o4UbKyYWx9MImsRsB3D+wDxJEMWOdIYkxsla2mdDAOc29YcOQKMFZ8lma51JbHMI6usiRqv4SwsnLoNBOjfIGUkJZHkOPKSDXqvUDdwKccCHXWImbgiTeycQDIhE7ZrwETnj9bBh8zm6coLt3rxsV7cd/C1T8LgnyQEgEyoT4bSuWhW7hDHALsW5MuSKuXU3U0o8KuIFGYa2sVvt/KXO9twoVngdeHi3os+z1O8CX9CoAbx3AjBYTUlSNS9sCIKr79vZixzdl1pLOoSRza/+eISfIjgOt7trsK0lv4D7NMdB3WjLHAz3LazD2QBke9yCFKFFsjH+iPLnrDXrGgkKiu8EJZeNPEkZt+r3m/EKqLUjPConDpiMW+AwZdSf0ErCSjY8uh7qcAxNYmKP3TIzLf7NKka7jiNFaqTgMX6jhnrCiMHgk9RAGiHA/17FwdezMeH8zP+MlqjAD/QVUTULqMU9Tbd7FoN+FKHKc6gSIxvgdQjQ/V5wdf2/q3ntNupAqWVaAOLsrA1mmPhKeR4wetLxIFPwb9A8qStF1ZzVt1CLA/J8p5uznhplLbXIF9jZx9oBdBNqMEaWgHjADUrQ5BCKzNLkE5pNgYBIbhDnL14s9w8OgoY/Ujlh33Ei6U96Sw4ahhwEmRgDyRZSMNydM1O1o2iPoXdgl+859tdr49c/C0vMlbF4n3Naeq+Ny6d0gjb739EFHFNcmF0PP7CKC7URYi3XcuNmwBB1EhPWQyk0Y294AcXSelr6QfmFe9/kvbRv77Y/KGY+ApdyOO3vhXWNT35VKnVBHt0wjtoFnZ41A+JZxOGUjbhBaxiTckpXRKN1z8w//FfF19jZyj4gVK2FYK0k4kHpIAgarG8JNcwyyxnVALRHx2RRV3X36R9XEXV/KtHY/hoEVlzXO1N6BrSI/6waI574SHCwzkSAMAt3NcWIS2BiBpJ+JlHs54GeOI2Mas2VHVj1IBF5YeIxWKzKWlEjAKiR1WZNwAvmbLLREzfqO0q5oRj8d3bh9d5pIe1wj00E4AIW2pedGTtsHo0Lny+PWPtDSwdAAZ5hDwbyL10O4MmxKqI01RVT7gSPqX7l09K4b+gWoqdFDGQIedMmQXS5RAmT2QW55xhlGvWckoUr0IURrqjAfa5qRqBpLN1oSfer+g+s3djmhIqLpT8CryffGgsFo4caKj8GYcoGJ7Q/rgwtEV5nJsenaUdFy2gYdNfvAxZtrQAucrODlAHkW46aGB2/TcO4fA0yOGAEJk3zUkdBeq9atfW8AgCvg9d0h9IMId/32brc49z7VDIxa9P6mXnq5TkUiFcqDqDGR73TWV9E8PWrlJqrtokjHwxEqT+Xz6Kl1OKUKKntkV1euj9TU36muZaeDXwy1Cgv8kJOcR/oc2clhnm5rIcxJT3cU0ijNIHcMOIzfT7SLTSBnsjwGGWqvukVf02Y5QqDiXux2iqmtn4RFZ3wPtq6bOgKVEyF19JW4WXkl5FvNQ8vQd2vYuPXZ0ILb8+Zj55MiQ6cAqhPpcRCNIn3XeI9aYDKNeTXNokd8R/waWT/pYkli78GHFlQlBF2LcCUQ8KvHAzAzbryZtF5xzmna63t/I9/5p7o4Q1IYgBBahaK4HqneICeobJRABpMPO07jq75T/3pl770PoB3rFwlWxuavDnF1YwmO0cHEiivOUZZh4wK/ZLgPieLSVKbGAxJaxbysbUSiYp5VAl9vF0guOwrHFQ4KrdGsJheDCFLMA8qT2DQqsfHm55bLEKYjFd2yCMnJ0RXiQYx6Uka/VFHtIvA640R51up2keqs0lBoF/DAbRrOfAGTwFsH2VIzqyiG7rlnwQkfeB709/+cMAlCNKFj0ZdE1JxxsNvqeFjgNVJnUTPMjW5tGLf5MQ4CWJnG/NiNpyIrcZXiuHFEUZ2nuHYoY46sGzZKZ0NgXePMjLfna39N6nLpptlWpHh2+51ktbXXzW7K8DxwFjfzNSMui8oaaBbiuApbpzIW6LGOa4jPqBg3b67Up4Z1vqWPaJXss8PPDT9TpMlBS501/tYFTkkJss0MisqY5nLwuXKnI18Qrk9Yzl6zgOKu0PZSqlg6YjKqwObKLXRBBh+LMZi+ehe6F54Txl3vjgtdl+iWXlcTCaEOuMmClmqJCmWyPhnJoDzo6vpXCt2VH05s3TrVnB5qcG/N4QzmOrcyeGYBYeNvyCpMS7jJGEkSj2DKO0oE0AUhSlQSZcaxkoSIEmASNpsaHmNuXGNhX96daWqmGYy0aYT2oriBLDOWIKnphO09PeXGILDV0UL8RkRTSkOXRaSl1A5kprQGPwcwvlvA9vsB5p+EJUjamwGtcCdc50yAGfPeBgA/g9PfHMHmnwi/49nfjDrFafG2m/CD+1nWsBmUddEYIMdviXpP7IWR+/eTWjrMIcJGMFu0Yf6GEqPllgVd9bEd9De7krFkWnv8r87U1rmOyD0XpKijQLnVI6AmkIQmFo0nY3UNBNxsgw0F2lhuPqVWY1TDwEQurotcHzZ9IyJVxaGElm31MRJJ+D/U9jNVf6umDii91AJRuFt66n0RYALr8cKy8RFcK8tD/WNw3Ju2QpvqA0muCpbWTFukrTYZTIKFxRq0LBOkYPIxxt5ZmdgEogDHTZBbkHpDyZJlHirTa0WpCjxbDh+UUOCoRnRnsGt816Ni0eAQU2aqUJDrO/EiHblvD/Pdr4oKM13w3AiUbkKMyTxM5MUgwhKIeuVeETcGZF59J5i8cwMK9zUmE2/AJLUfTehKfr6p65lgiTww2yCWKeRhlRYfHqzMcKWWBCxQ4sZBRnULATzAsIz4FF0gZ5x54OxSQ8bBZLysJpRJAJB2BVkmFmGgfRiDwG55nN/7O1Hzh4VQs23p0txhm07gbebbNNz3axBLTgddsXUNs6qp2CEfaM7SF3mL37YsGLhkE6xZ49QGfr3bqJvZ67CPJKXTwBSzFqoJubaWamn4FADYmgGKHOYY5JCh2LE8RkizRI+wOnLO7Lml7dP3zf9uv3ciFvBnfTOf35I/YU48Y/n5cX1iKFl9C11zNKwDWNd7H82ClW3dmrppyuNiZdseDbASewuhcMFuDetWQX3xkKDvAABux75eZ7fZNUevzGwfx3raVjd9LiiPi3BDSXgr2un3zQCwAvf5wG89cd+NVZRUN+MIZNzX0/LtuOL3sYrPJaE4ghEZm2ce5dSbQa4uWkjSvysMumOU20o3i+Qg1HTPiCNDus5+MvfTmj5IfiKNjgdHYqgSgMwHDdhCH0Dqs4HDrhhQWRrtljfn7BepWLwnznW8QHltSCgVCgwLGk1XxKHrhFUQYe0eraIbHKmvO3FM3rKOmb3N9qjeT7iEx3p1ca1mQB7V+HDVNN2mxnUySX4KJNlNor9ZGSzrR5JzgA4i3bJ9yo4G423uVVI2sBS7qLtD1RsGirI3gV0sBgdi8gjUv0KRnb2FdBH27q1CvucmCBuvAVnAHDCjH9OVgQRaoNgOsONB0CPbAYozWDo+IdXBDFEcwcyFxXjuynfCVnhPOulTcv99XFfIdcy/oxE1xnQcdYGLufbCCwHg+4+SO4G239FxXFcpVi8Bx1UQNqSIogcGBrjcdoC4j85wePhnVRRf3Db45UPuYP1BXie/rHvkg1x/OO9t2+9vmMGFI9uwRe6nlnH5t3FUfxe4LSZcQn5uq8uJvN3Gv6fHkzVbkzQJoZ3Rqc7kEKzaK2bKUD+Mwg5MpNHyZdSrbVMgzgZyWBk85wqtG5EQU9F9fIiHDI1scEcr+Mq1a/ytV216tYqK71Ai/yzdgikrFzSS/NTGPBmUR4SKfyMd57euDn5ee+cr7ob+ftQ2N7dtdcYbSME/j3lYkaMkV5RwkhiCAl5MiYqexZDp0pk2DpqtCPh0Kc10EPFjrPAatLmtbpnsNpKRcK8iFXSximFqkPh5xFsycwNnjAzFVQYYsBpbZpVw9KBqll4DLS2MQrOSS7Y+wQhFREpruHcQ4KI/F9BE2XV0D81yomNXe57Ssxa9sbDsbZ+sX7NmyLZc739G/NTVRgb3iLZTbhJR8+XaLUaQa31NwVv6z/XxTbsMbvxwQClYxQjqsf9W0drVq6VsiKCC4Nfv8Jdt/DVt7JMV/IMfR7AWz/gP2ZL/vZx8aBKKHQVMJZFfS7A5K9NlOBQ4pyoNlS8/fg031o1RAL9FJLXS2HFhpFEXo9AArTzyVR2NNXkeCIdRKJtgvAeK4DF/ITTk3ADCWKgIvUPKPx8qacpdFIvOOl5Plt6+5Yr1b9J+VxfuXsdRCJO7N0ihtzphddDJufcX4123TUxOTOEWKbpZ/xEHlr6LFWY6mgpaZ2uAi6y3cRjjIoCZn9NAi83+I2EGslU4LjPaDDwTySAfGNcGKc1q0E8UQtHFJv27JHgH2GufYdP+bMW02InHkgQR0uAdtOEYWW9TIkalBrp1EVEjcA2RUkFEmc+9DObhoIvrQu2XUVxBCIqbgI84w2wmDea2IoCWNoDNdwOcNQXU60A+h61+YAkyiPTMZZ3BvDPeAkJcBqvXujB4MDwBwzG9qPyluD71irizPVatczqbY/XPAohX8HYPSUctAFZ6AOsDr3jcGVG+58O60BmLCJWbJnb0dLbfsLueToCDjMeTrX+qDjrn5vqvbc4d95JV3uRYTjVKvkBGNBA6cJEmCxEysQp9XwPSsgUonmC/HmlkPQ1bugDaz9kMExvp3eKM9sujuOsKqKoYNOZs6wJ8X+RQ3SmHxiFSKmIRCIESlVXJ3OGI1+vu1m3CCXML3bHNG8lNOtA9EyvXrPG23bb+WcHk5CUwtfclUsmSjsT1oiXcDJWpXyu3dXjNf39gw8All9D38cFJk4PGqA7Q3x4H+rH/kH9FFAM3bLBXkNAK2KqTyT6RbUhTyrRYUwYHFazB0UQvj6E4jpk2ZOB9I9bPGARrOMiYJA+zze2iGBuJupuiBSrhkuabbVHB3cokZMBBYUfworM3edc/eIeOgrM0QUcIN2E2nlQgGYwd1ATc/xuAc16sYWoU5apNjYNKVhLXANW74O2tS//iisogjB2865AmunxZ/czrf+Csv1G39j1PecUmtM17uQDxeX3xknfAwECUNp/Yi4LW0rp46wO3sODcKDfru7p1Xhto3ZCNUt6NSp/cvXtT7Vgvw6Gf3eb267aaCbPfOMzmAXyeqYxdveeeveTs7zMOtv19B+XyDu2x6eFmM+e1toxHE1Of8nK5j1b3bh6lmZMpLA9ccksmv4CJcwRk9av8yjccF0LnW7SXL0lUP3VFTmHbP2Cpz5XktpOuosPGD1FVEh9uKpWiWqIjfTfSUuf06LYN8OvPfT2NBzIDu4Lt6mzAQhw+JT4p1wWo4YvgG0kugeoNJHJHveC4gu9zPRIPgTUtLRgR+0wt1AuTjDbjwzwUBmyYgIyoGZX+ZkCn+2KJHZx4Or/0WmhMngUtMxV1d9kOJ5O7SHTscWW45yaAU54tAAXpyaRYXA+mTsMY+pb2VRed9wHY/Nb3w5prkNrqoGMASVHkiW/Tk1vu0l1LWnW+LRBi3ludG7et9FqXfqRR3TwIen8vIdexYKFSrW+L8t1/q1t6CyCcpggaeac2dG3zfZuugP5nIlHKYY/peZ3HHjg9Erz88IfF0x6iIW3sRwQyuxtfN5NaFPW84DJqJ7+p4doKAWMSGs14iVhw8kd0vsglT98jOSNNoRA2f+IMwsybBOGnsgNcwue2bYVpFs8BiAhA93VYM4DJz2nHqolIyC7bvBabghKGSaYe63BGj6YNqmAyhhiLEQQKt/XJ/cLtpOxoeSaT8oWFxVl+GtqRdUyIkVmj9JVGmXhOcWiNDRU61vE+BoHLjyJXuA6iqY8CzMDqpSk/mRQoyzDxxx0fYHIPwP2/BnHmi0FPjgr2YUycpJSjfRWL7oVvgxXvvxKuWbMZLlsrMZFzgHtMOeZm6f4tTstxL4cJ+BF0LirqfGcdvNwFQaXwK+HP/A3oeND15RYhowklC/Nj5T47bNReINpmtel8AY+1Cc1qTtZ2/ba4cPEbpvrvnV5VODaywxh6AV7HiteHjeZZEDdcqjsYZ9I8zQoJgzjTJWPsbxaIi7DsFhLp3unzDPM1HPDGPhD0yKBduREOIXmcUUR5QPRBMCfFoSmLD1YcV3wrGt/8u0dIou5jvg6zIlBDVfBGBHlKnEuEVtM6qjCmZ4AQpa3wkIKMuIFrLgfWB+M40tp3QeYRXg+w8v79n6+En9rk1rICCQQOIAPDE4zWWbIFfPWYBI3WbvwEcdLRMCuqEXtNmsMM1QHna/myUIrAlBs4ezwtSqGSEoGZKCuIfWfTcgi0M67Nv+eVd4pP/fh+iMJTwfE4J0kbShD6RvA9BGjr0nDPLwSculoT0X6CisLeeiUgCrTund8mZi/7Fy3EJbDmmkMlbKgSEFcHbsxF8R9HOvyGals4SyGKsGMe1ryfJcLGsyKFCOMIqE/EK4Au9CglvDrEcUHWJ3KytOP77XPa3jR+9w+PQL3+aTv40Vy5pkXuuPdrkdP1MujuNGlrU73Kwuft1DPkGhmSBOMpZ+Zm8nv2M2L65yyYy04OA4xFEUICOlZH3+l2Lbs0mtj074e4h4/JyLtOU0Zx5LJ7gI46OtUWQm4AE6aYkqJICchkSyMWRepChA8hTtDLNPRPzylwGT05VFufT68BSxzb7RpwK5VfaJITuTSiGdG0Is3dgS6A9YEswIPLt6a+S4RUXMnAUjzZHSpAGtFwLgfTrh2uJB1gcq52cAUXUfUaiCpGfcl03iVYCltKwd6svIByGfTdNwK0d3HclIhSGEUOqWIxZ9mr3NM+upoSORg6HHRwebDZ3PVLTwRnOxMbvyWmHnZEs+RjfUXnWyM0AFDsUzrXjWIXSP0iZW2s4E5u2uOWN/19VN/4SpSDO2YMDjUQwAPKeWhdv26Z/zLdMTcAzwnAy4eQwypPMYS8+SmYf/E93/zNL0SQox9+H7/jF/h9/Iyfx9ch5PL8GS8X0rbdfMT7yIXg5yLw8V/6nbfneaHwCqHuXOBIr/WTLd2zn59MnSM0IhKxtCu/0RGyiqcC4Tro8NAPv0avF4GaJJ/sMs8oabDbUIAMwn7BlnbwM0YklzBrSSaR/sy4HW5cpooDslCTTCcjD0m2l/ABVBU44KKGFQVkyqdeMhboEpSENBC6SAlKrWCOIWL9S2QaEfgdeh8bWhVAGCFQSWLgv+/gaoPymt8VzanLdKHLSs9YTeBUqJJOKATo6AH4/c8BTrgAuQ35PZu+wE+hiE33bKlnH3853AXnH9C9OoBRaEwM7AABr/WcE/49Cip/KoV4iRDOEkCdW9tAJOWoiOq3Kt34Xq63+7rq1of2Zq76Mc/gwIPQfO29Jy6p1tQ7db41hrhhGs+M9ppdubNIbRvbGyClKUub9JilTk0wrvYZydxr0ydhXFsT1JowlD9A8DlcWaWMVEuvE42Pv5sQr0d02ElvQKK2tC4yP/Z4E1CWFThNTo5Jz91D2inkn+XPWzyxXVY5hGCPIL1a9g+MPky4aRHn6soD4jJsZs+wFDBfJYYKnPOgnIHBeiThi2WKx+Ik7zcWjuMEzeh3B3LfOeooPfSgUPXbCNqEcV0W4cdgJdvcg6EC98bfei1AezvDWW19hU0V9jhEevby8+RZ/3oJ5RAO6SXgMMQnyCAxtXGdLt39/vmdHWe0TW45viPcdlJnbcuJ/uRtS4vh5IqodPefqPIDX6puvQWNwdNc0/FIDCzzYjJu6tm60JNjPQ5zi6kVMdFuTDR9jSJ0clt5orMHShktS8VhXH/zHYYxpYhvTrIZLC6Hz6x3mtArk14tKjrFItKeUCI3ywD3UuKexzlI8YdxVoamiZZazcbK6mJmtDGpxTsp+OOznvS1oF4qvey/bL/njSWq6fwNAQ3mHoyJtOlCanA0TFX0r2EcoGvHRDNkL1HDHTc6E/N8mX1QG6lR3KYGCCO1wCkCtufEjcL7TPR/cZtIRBOLWLiOrJWjPaWx+hsPFs9zMjGc+jY0pxBwRgJ/5ggs4NLUVvHmxgBt3QAbbtWwe6uAQgvTddPHjC41XpzWDg0zj/tEF6zqYC8hS6pxwGEjJCppbds22JiEyW3jlT0bxoKxjQ2AbeXyRqOSy3yTz3iatEcx0Js07Mvmhtp/zDJiwImG8zRD8cDJsIQLlQUBTKu8zYtzEo2fSbMAJkQdiZNgicGN38kChEm6XCLOZpqTd2SMvIecJcm5Gnk6zHllQwYWSuTrYIt09juE5KTPGc8eYO3+IQPrrKeKBub6GO8ooW037RqWx8RQGRjCdO4GoNjBXIjpZToiykYyK1bWpUscI8k+2hP8l6DNCFPG99mpizGVirY3Vko4QlTHY90c1396+xWjDx7EIKzlsKGl9TuiOTKJdBVpjJDIP9huYAu51uDmAG4aAGqASk42qcIizZqCeSsXlJ795+9jL2HgcNl/iArbbEvu82ONAP4c8woOa3BYKHp6b4GwhN1JhGMzfbJJhiq954aMK+lrsJgU6wXu4/TaFpps24i1DbyMcFdl4sjZ3dpOKK79i2ZNSEfcy726RzCHEFFLd/IAZ8IgMIBe06vBHV2GMd7WHDPhA/5no26T9s8MbE2iS5tRYjeEhtZIGHNiXTErYMOWiPw00yKU9MHvuw/8jGHAi9FF4NfkuLBRoKYnzkvQtjl6iUKkP5Bxtazk6NbaW2/+z52Dq94K3kEmJJUFJYyu2y3D8o9EWMGyED44zIxkfESjw21DAwmt7QJ2PqjFfTcJ0dZBezX6dsZdjIXK52M996T3eSd/5GQYeE2MjU+P4l5m6srJzzEj8OgHJemCXXdtknF9QDYmHfD8ABwXBXAUIFmIxFXQiUDIgH/HciBEJJQjHezmDOnvqBuJv9NnXfsvvo/d9Zn3XQT34GdDJifF125k9kmfE46H/i9qxAdCOp6sDYfSLf1H5t4fkeGGsSQyR15SWSmS0/qay480i/B5ZaQ/1/XScAhfEx0cLrdG9fqyAxwft3FRC1gSMtC2IX2Nu0ftaNofrem0P6JGps9g8CIQEGq2j31BqbYj8hjQJEMIYcy1Q6JXwtAETTx1/3MLNG2dPoP2xlXlydgb317/u/u/MXrV6rXgrrsSUK76YIN6vxHTeBXU9r4BOttJbNvoNpjCY5YpDCd8hBRqoG+5VovFpzOkmdBa3OFNlyxqat13XF5PLPs03Kv/CPauTCqzR+qGHxuHNShX1LJkxd/Utm9tFXHzpcpvzZQT9/84pw8ZbZ96i7bAbR6Lfb+eJCUTLyBxKA1Fm/EbDMM781c4oj4RCt34q+rw8H1HvHRMufQErpNyO0D2wM1B21LptEqqhSKbyZ5WGfZ5hs1FmVZfTbafITRMdplcDfIhiLeAcESoIasPGDJoReEA62Sa9C7rbDDXjdkX6uHYeM71RdRoKL8+Gq29/8t7P43GYJCYopP254Mm9WT03ktuFp/84e8gCs4Dh6DM1k1Pm4JsfIXfQLw7Crrc/GOAi98gYGKvoRJIoiep0XVYdMpz3fMv/5to8NLPMoKRMefHxhM26N5N3XvdBEj5Mr9t6ctlOLlaRVE7NsghiFYSPbtGNQ1k9dIKwUiYxcIiGXF+YT4Q++k44EeeH4fKcDYzz0I++Bn+hHSpv4Yr7PTscTc0CTJgtksJx0Vg8C5w5NfCsS13HY3SceRiW7yVQLbT2HZ3mpI6XwPDcJXJndi/07DENQcZ3FKcwV7YMDqTh7MYDVbF4lSCVWU0eV0bn/Pktp3DZpNaYm8xF3wMcx36cpS2SyGFkkS4EXntStVsKr+6t/Hxe76052MohTfI6ld0MIcwCDhWI6owcgrLvxzXJ87TrT1cJsFHIFH6MRlNq/YU2wTjrwFOOlvDzAUCqmWWW+Oma4Bm4MS51ljMPe2f/MVvuyH4ziWbYO1BEYzHxtEbBrWidDD14A8AAH+SkbXQ+1rrg/3tUFb9YH87xHeODqgsCtPCPYpAkMxsAjFmGDGL0Zj8ga2fWD11EzZQ9GwMwgGASRAhkzXR2rNhcI0lsUAkGpnTY+/DfMYQRJBcLgKMBOEkdWIMeF+xVhLzjcTFaETxqL7JloEz/5wiQp7IuFlTXmOi9m/3XrX3w4lKdrr3R2rz5O7A9rkd34Tm7j1EwmlE4Ke7R4mzwFlVHIU2gF9+HRU7iePZ8kMbOikJQU3rWYs7wqVnf4nII9afuK/Pdmw8McMse6Q7+Dgp3o/UQDp9yi0dnQUiDBNQRZJUNFwmYAp2/NKUWWz5j/5qy6+GPi429uwAwCSLzzNJVUbxcN7AbN9gN7j0mIlHbN7CzDTMV4TIYpDNIdij5d2SzcEZRq4AHTL6W8yYj8fvSh02tFcdbnz83qv2fmhfY/CKTy6euWbtytZH8BDocJzxLbeVRHHF1dAofQgKPUhT7VtNGNvgncBRGIukId8mYHQ3wK0/FnD+qzWM7hGArhpXqDAZ6YIOQ1h0xrPEeZ/4mB645MPcIt3/mBlqjo3HM9YHcM01DlzyzllQyEvknC7W60iiKQXUVQ2ELkDewfex+l4r5J0CskLVG6pIz0FR1KCuilAVulBwRF2oWibyboEaofeJxB5q9Hn7jJltCVGXsWxthJXqulHQRFlydHJLNodg8FU02BPQtgPITBWWMuK/2VCHj4mZClm17CBDk/NucwNprd7sx9ZeDb7YVm9ofhhUEQMhKPbnroQMHwIPLJjEVFlgHU36Jno25GJQmk8hN3KjopzaePMjG76651+yxmDNGpC7FvYUtzxce3vRDz75SAaBz0uDyOfcK5r1kXeoQldrIgjLZnL6wsLSkKgnDsitKG7/qdZLTwNonQXQQCg0NbHwfY4CR+eLoVh6zgec0j/cFA/2//RYPuEJHxjCK697xeuiv7r8A7pjwVKb2Kp7EUejZn1soI9MDyn/V8dHtx3ZDqizlhApDaLqAwk+guNZnAHr9jVaMrDpntBxpuWGWf2EVi7SVYi8JHUYt1N9t2t253tG1g9WjoZRcH1HxUwfz4eCyyvNc5nJfViwhHVcsZBv8gxJX0YIQofcBEHApH1CBuyTSOjoM4lVm2Bnh4DTkklLtOk8sgaKsFvYkUlJD4CZRBGYDORcpKQunhGir6mlzEAU0Vy5UlWmArcx0XzP5m/u/S9jDMjbwDLjwJUQLnmT/ISU6o57rthZdw+3RFUfv3en17r8Kzqo/I3225HD3+WHg87EAlRSkVC+2UIXWwF+9g2A13wQgIShDLMmu1ESwprWPX1SLzvry4XSO8+oD7xm6Fg+4Qkb1F6cb1vyulAVvq47ZlMbcAI9SKZihnHevjbeLpfnLVW9xR5Nc8cz2DIe+8LReDlOWH111Kz/5cTOnYtWrlzzovXriQfjyHoJhL6lFyk4jggINCf8kglqtCzJSzBty9Y+EHMpZulY//BAgyDY1vHPBiR80olrkuwnrcQwFJwhEQapZMqb+/AhACpg2RXawBtiFr2JpSuc8nigJndV37TnuvEvw1pwSWoOx1vBwzLjgtf2/vW8BflXrDmx7f1/88W9BwY7HOgS0iVxGp8S1SGkoDK1pxT0nQGvMGWkjbFQ3AVDh9/8GKCzV0NI2AS+JPw9CUFdqXnL+8KlK5Fo4lg+4YkZZL57li9vC7X3z3H7fCQEwTZSbpUhIi7EvFHB3UBkEDLD1WxSZaTX9Lmkyk0sSlYoEicLQWHwuSB1Ry7gU7sN/WCpmsG0NCno87EotjV1vvO5D+2652VsTR4VVuWRB04JW9ZnK5bWQrXF1Rg4tc0ZYF2f/0tqebymHWLQ6m1g22xoDHbHbttsx/AeJOaCZRi5OznJNxw4h0CHZhtUmUgF9xkhiUt5IhwPpuov3c8Y4OsrIZy1ZtYbl53T9rnZc3JX/c0l6yur164+RMlk+iCb2Zza9rAKp66CsCoJjGLx3txZYQ8x/RdTnni0nTMA7vgZwNZ7ANq7BUSY1MmekXJABUG8+KyLxLM+9UkqQa5ee2QfgmNj30Fr4NRQ/WRV6F5Ej2kcIf+hqWYnlfYsZhendIoQTWpLFsJo1GY4r8QkjGkOO8EoJ7EzL4CGU5SY6ZG6Q+Lipv0OFcT6Ajg4F+bjGmbeWwifAQkpA30zNi3pX8Bv4Gcy4E3DSoLZuoPuw/JOJtPd9kJkts+2lI8j29NB146unzEJxCuaGabKgNhzjXkEwA5GvHphGDnu+K7G1sZU47lbvz36E1idGgMME/D1vNf3vXnxafmrZ/SIyrYHql/FHVwEg1xCPszBfl1eflpUtpeRWs3S8BrLmnmIqJJtsemc6kTk4s++LESzLMDPE3035xqoy0JD2PQ0ts4uOufv5Fn/8hZKLmKS8dg4qiMKHYfdfIYKcCOaUXlmYL1tbOIJwXkuc68tFMa85vliP2fVQ8xnrBJU8rtZD00fMn2H/kxxt5KegKhuJsHgkT1n/B8zixodOnNsis4Dr4ERMqDXxpsw58yNTqYxwOpcHniwzpEJr+z2+XXmehkSGTYWvH+aE6YDCcnT0TcIjFDLvoO0Mbl9WTkQNRsqN7G79rvShuaFD18zcjcZA8tlatCIfa/sfdO8E3JXLT3Rg/Hdzetv+Y/dG9dcA7K/n9IQhz3YvZ+4f4eMSp8TjSkUqotS9SFj2ZI4yKK0zBvY54CfuO5/NRRabcRg5MC4owWimtDtnbGef8oV7vHvexYZhUfsijw2HttYS/fF75w7JuM6ri0WTWTXcivomqFYzj4LxgOkJEJ2TLvv6Wu7XJgW3aRpav/tsuS4CoXSgRFTPcIDVZ15V+Zf0xWk0fhhyt4kGhlGbCdq2rFoKIrMhD54foOojowBSJo9TJLFSgDZQCENIyy4kD9vk48oX3OAKgM1fSihHU/oZjX0SrtrnwEx9NyxO8eGKIGYMQboGXS9ZOZblpyS/9KKE0VYGYvDiW3ik9mtPdq6M5nTuCP3KWjsGTPlFJsRYgotdm9MKoUqo9z1hRe4pVPDCOYTviuguxeZmVMAFiUhhYCgKqBvoRsvPm0AjvurRfCd18RMznpsHNnB/SrB7l9vhObUb0TYlALFbu0Ete3NDN9lZ5aNhLnHiT/IjzLfxwz1ke0dtqkus10rD8iZ9JRE3BohoZFMwxGVIeU67nfMwR5RPIIrbGuzmbSJfLpOJyL937aAm6SjncFJFw2e9sHTCPzw24m9L3THug1WQXufxjK2z6ZNSYPwTL4vYV3mIQQ0EXswNVSHYCx419APR/525wDUSSCXqwnCGoOel81624pzc1cuXSmaUSzc0kj4o9uv2PH7NWvAGbiEKw+PdqJhgkfCnrtGPF29XATjLjWnEE+LaWZjVje7CpiHgEglFehQQ9cMgLtvVPDgrRo60Sg08XPI32LWDOzYqkWw5LTZctl53wCtPViLnWTHjMJRGOiZKadQfI+s7J5CYlpA5k16zB3J7ED44+FrBzBKpPcdZA2imJ9q5sS8hQBZia8xg8+hIkGdKTfAn7XvI26YCMrJM0ReRfzB8iQKufnu1EOOL6ofCGu77j4qaEWviMQ6uHuraWK7PcGkOCznh21qZKghP6GcdKVKGjpWts//ACNGHkoSqEk7PykJm+TZ7FwxJQLmRjR21Ohqok/Cuu48BqY1N8UV1T72YLk2url08fYfDP8/8grwKFktm1Hl/RDNfuXsSxec4v3P0pNEqF0t6hUFk0P63zKdFDQeQ4xOsm9O0Nf4f3L39j8Dpw15F0niNgEppQrBprBrdQExVgsVdM4ScOM3AGbM1lDoFFCnbkpzaET25uqo2YTFp5/rNP/z6rhfvN4wNh9rgjqygzk0R++50+tafmFc3tHvSO9MkIguYPYPTSpM2ClLFGCGGo86a/FmMqEPP8ykDmbKcYaal0TJjGRIEoczFyi3FBsUP7biaqUjXZE62ih1/T8ble3XHjXochiYfAYxtRipTtY449KiAQ4QpMIuc0bShBVYLYUgT/qDDLQF7DijfAq9w2F3Us1Iio9pXTfB96R+mFkpzXXgXob+fmKUhlqjvjksqw8076n+AlaBBwOGNd/gDa65Bpz3/WD+ZzrmyXcsXanDOESWAuFPjURf//1nhm5fcwE4Axm26MeStNN0UJsHmqJlxftEZejnumMRdzomFOy2Zm0V5jLlTWriQL77PMCPP6/Fmg8K7eWwC9JynDAgW4U+6vHp5ee/zo3+rRoNXPJWNgprLMfjAe/BUUG2Pb0HPaThxAP3AsAr55zwvJ6KozpZRAEjCE8IGaNlEBpVmR0VM4+3EtBQAnIsfKS9WIrQQSAM53xCU58n0lDmHfSVI4N8iLzYoD2ucOXDWDURWicCbPMtVR+6ZW9SqT9K0OWINOUSHKKNFoQJczjhZ2N8tgkGwkyHZf9GfoMRVj/whTUQHWsLzXbT/ZG6q8mr8hlnEQtMNsLNhLQjPT2HQAdWvaX8pxR9YIjQb4wBJhIHIIKl3e3v+Z5/Tdci74ULFseBT5L0QpZH4+bOLQGVKVbeP32+PMYsPnVCOnF1w42i6HxbBN2v0bk2TDCyDHTCCG/+b0TokxPF0CFfBKhMCf2Tz2t4xbsBpgIuv9g4FC9es+4ozw/08Re+xYn+LYoHLnkHrNEWdnmg+/BMMAaP1ujtu4Klq14SOKfu5dDGG5GBCn+ekIFQxH2GKV8+HkWlRxi4djGXGbrvkgH/ht3IQpZxHhLdOoVCRJOaMRqcdESm0v2kTTKD2VLpFfXr2MStbY5Kic0MssvmYFL9VszfUIHGSF/q6TkE7kTH68UhgoDVnEjseOHsM7rnO190u71TFy6Kwq4O5ccRBK7nuGND4vMbPj+6CXMH/dzpmIzHk6xjbFpn1/uhunuCyCSMk5QhYrXi1Em2JtmrCjW0dWoY3i7gZ1ezhDyRUVj4m7koQcPTOTfQy8/5a+e0j30KBgSyNtsW7GT0rlzTih4EadUc5R/cx+PaT8rwm/5+ON/h1ylDsAnhH+G73FxG0uT0Y2TKp/Fk8kNrXbokdn6EH6vxKI7Az/Tn6ii3wge8Olna86TCoLmqwNUE0yCgp1cdLP8hzU7iRz942ZESkUk1gzAFCUFKkp1MkpkG02PxCkk+EnWjMLI5VI+PIg9hLQg0BjNfPvu13QudQa/bP3XebBX19sReIwDl5sEd26n27F4f/9PatSAHVu6/sDyeOj8jyIYGdoi20/p1bfzTUOyLDMuF7cQ2KsEJcMUyKLHrhW2o2O/w4J1CF78L8KxXIaoRk1ac/aVvo2Rcw1UthUCueNbfO7LfiwcuebdphMIKhECFnnK1frK84V+vEC0rmgLimhKuLwTp5rJIBN4ObAHDmAVpt7VCFIxSqDhCACuMHJFbgsS0yBtEZI3NqiuEfvFTg0nwBPSOKSMkn0BLrhCJgsI4iNIzCwPqbhIPvstRIPNcCoG8AeiHayEclN4iRW6arvz00IwQmG1LlL0RCMaFKUrW4YpDUEE8G4sHJ9eSyO9ZpwMfIqQtNa47vXKE1nWhw4rOd30iGrvr9kysTuGabF3wPhHDqcJxmpxDYLpkVgfkAj6erQEi4fUkKT7aLaqMUvLRPOFE4SMwhUgduvSKkvNxxJbNEY6u1cARE/nupZ+e2jY4edRDP2xu4uIG4wyYT0jQs5ZUAwibYBY02/WY6UPAI8QuBsqrHHikPU3J8omvmHs59ZczYYKZLJaJmrsREPsUYb71oBcE8wVmpe95SW9/53z/H7XnQlebChcvjN1yWWjP1yoIHHdiWH/swW/uHl1/DThgKgvTLs3ju7KcYFRnvOwKcfv3X6f9jrPB8xFIaQDx2QuCSRnL1WorjYg9aGrdPRvgzp8DFNsFnPZcgLEhlNRiEAh+ASFQzbqnCsVQrHj2u5zcf4Tx4N+/N80p9ENj27W3QO9Fl8ra6Ld0YXanwryEvcg4x5NKKM9xC8e3yxExwhFpT9JnYTpi7fdMHSTjmNgiCn3OSHVyqiorZsKfnE4Vynq/FuBHFifBxpiPGO+RLUsafdnnj+0qmyh7VHZ/6V2eXsmy/KdUBghroOujewAgaxD446F4o26bd5Kd/1xJNvmupKUh5d3N+mqWbnPa6dpmWHsq0waefwiiWYFaeWw7AHyR9TsPU4XpMQwXFDGMEdLangvn9ATEFDoYJUQWNWFRqkwVjcVRiWsVSLD9IIPVqGzzlKGXttnGjDgL3ZKsa0BeikL+AiSEFVNNV07W+Gin4xA4RBiAaO7Fc+cFHfH/tM/Kv6SpnWhGeyRPOjn0qjWppQexk5Pu8Ob4Du/mpf+7du0O2X/JgfMzjxcJyDNmsD922he/M65su1l3LkNhFlKSTR7ItOXBPBm2UwbLkbh2RwA9czTc8iMNxTaAZWcAjA2b+MxmYx1EM7o654d64Wl/78Sf0PHAJe+DtVpCv5mRI7/6qTd/1ZnR2Pj3Rcfsk7Wfb0IcerSGp90zuNJxQGKTSGaNMMebJEGSeSQOYM+TszPJIjYi/L995jDzjSdk+FmNg4TFi75khbanWwC7X5uYNn0sCV7AXtlpuJd9F1nuODefkU6svIIjmlPPJk7LREjXXG4BUxpFU5BX0XYuoe+S7Mmei9lREh6SwUpMhqXrSK9spkEqOTB8CjtDLXKeLj08D56QQRfanKsxbFRY0FxOpOqCfY+ph1hhKZN9se9H2Ltx0N2wnF3aQ2URkOxd4LBdSeyGsviaT262I8LIc4bLNWfP5LdzzconSe1k0CiYcxJRYYjQc3HPnzS61X93zivMbdbccE5v0z31jFBUsHjHcpW6MglQGg8/cPvgYDRzJmXvj4pBSEKHqDTwe1FwPwW17n+Alh6kpLFKcilQKUtDbYjpKC4zRN7Q2Qfwi28KyBU1zD0eYHwvakXaZY1psqPAVTkvlItPf6+rLuuI+sXbeFsYOvzKbe4Y3AKzVz1XlnZ+X7fOvgD8QkCdmWaNS7zfJLDBG4aeecIqzH9PDz1d13jem0xx8hyYzxvSWTMrErePJkoys825s04XXxH8Oj+NSVUm3YdFC6ZzjL5O7a02bEkCsUzexhqi5CYluBdKU0US3JwGP3ciXHfraQCwjpWcqMYdO7ncSBSHqPEYgY6MIEDG8vAKl5y+tYrpzqyXZVp47cdUpgqVponMTAhcx1HbnxDKLI9mvc3hWfMu6KBYfMbcN6P5SFUvWuWNLLMpgDCD6SF2ZNCNXC6gN5Jnmf9sbSzLJzkiAsdxIWxId7gcuKXy1wqV8f+Y+OXE/dOUs01vwrnnQmFzz8yPyR7vfW2z8lCfEtGSxaG74qRIVEqYX6IHI5KO640PhV+5/fK9N665JgUhHWgcIQQghw560Wv7RXXXXQJXZerOMheBseA2wZIxTrakYxIsODFQDu6GLwjYs4WboqLAftCadIGegsr5gVp8/l85qy7/BsBxeULerXknfs6B3etGVau8WNaGfyibNV9APkwSVwYjl9wcRm9naiK2YzO5WXYFSZt8Utw+M9sken/Jam4gq+bzdg92+iannfEmbJHNZqLT40xLV7wPRnTakREPTL6TVHnt5ywVuJmpdOxxrHNdjhDuK809tEeBH7ubLzUJf/P2GOdvDJ7dduJMmZ0k14l6htNoxtbtE1C/dR1Nj1EkISgHEqKbsjTxR23UDW950oNhG5D09JB2P/hwwnyUNYgHzSFAlAkn6Bmy14850w3NEXolKKOHf3CdnaOxv3HoGy27x89r/HjLX6IxSMBG6BWQlwBRx/OKp2+cPfOmwqLi+1r7cmFpBNQJKwL35FMaemqSvTmlhJK+cCZ2RbuGh+DvMZF4zZpDl3KPFCSYr876/sCV9T8Xpe1N09FmyBupoywh1kgmUhqjm6uLCDCpSNDlus8DDD+skMUZpeAMw4xlmgEImq7yvFAtPfM18vy/v7Zl5p/0mS5JLlsN31NVr33nGlHf/b+iPpYD8JmYOt1fBqOVUGSZTHPyQKQriQWTpEttyt+ftrImsrqJbBmr5mRLfMmppvtJ1qm0nyAhHMgE31ZtmTPi2IxvN7tPxcCUtuyxWI+M6MDNR+JYAvIUSe/lsAopy/Berea/NUqDIijb8+MAm3eVPkxW6YjnDOdK+dDT13y98DitwpO95yl9P4gYFY6c5titzalztj0RepwRNM2MobZr4ymYKgNf+rTPwE5ae1uspoQ9/BDbtw88sJPbJJUwoW2ejSS20+DIEFxHiEj57q4J5W8YuqZ1rHROcP2W10/9YvsdZAgYgqwoV9APChmOOl/Ue6mc0XVz28LWM4Ujw8qI9s48syEWLQ706BhldakNGnMHjZqUe7fB3228Ymhs/YkJQ+RBx5HsESA9xrD6wL2yMfxBUd0rAbP8ZAhsRs8mZbIrjEGA2dIa/tXJAbR2A1x/pYCR7QBtvWgULByW/3VQjjzytNRNfdxJz6svf8HN7pJ3nmcaotgoXPm2KK7c/1ao7FwrqqMuE70ScQu7aVYfO8usm7TzWmSlKYtbpCX/zZxJRtkoGxzz9lLi2YTO2553EnnbT1iq35RWi7NMfI8SFzw5tsw3E5fdwILtYSS31nak2u+k3OkY47odK9wH9p7DHyU2HhGfu/TXojG+AaKI2JSMEbSeSXIASe0xuZ+2GpxoL/Hxs8KT+Z1qlWn84fhK1krCiStfSOT7jvbA9urU6FvjLNKOTtKTs+VH7mzkJieaaezNWk6FgyMVrX00QjhmQcTQQMbgu1ho8eXO0cDZNHR1Yffos4KfbHzN1P9tWZcxBDGsN6Z4EKJZL+058+fRnF/581o/3j2/WIgbIlYN7T33+aGeO09BpeSAg3lMuosqFp7j7dkcfeW2T+4aeKRQwY4j3DTEgKUo3PxfurbrWghrGIci9juZNkn8znGVdR3TBxmNA+ZTsdTT0i7g+s8DjGwR0DETwwfr0tuhQEW+hiBS81csVQtP/6Vc/rdvNgrTgvIKGhzV2Pwx0dj7Xre800FLwseUZJX2cfsyk96+O60hj94wxmE6vCJJGE4rECUhkv286fC0q7f1TtI8YRq+2FxEEnin1zGx9clET1k/COdiATDJQaSv0wxfrP0CQl/ekMlEShgcjFxPfFRUdmMyF3MI1jPJdETaSU+TI70eiaCguU5J/sHc2yTRjIk1LxZRw4fqnt+euGLRN83+jz4dPwrHMOzClB1tPkCxR2AnrsUeWGFbZg3neB8T4axhevBu3ARbQHwxGruqIOcjh4wrto003A27vlDcNnRGdMPGvyj/6uFbpxkCxBGa16tXL8j3vLDrnxte7ubc3MIFxTYnnBwDVczF8vkXN3S+oKFcJj5jmx+L8i2OM7xVbdx2e/U9hDm45NCegR1Ho4uQXeTW+C2y+vCwgBh1Iq1TbckkTQdd0i9v0Vv4faxls5AVisgieOmGqzTsfkBDZ4+CuMHfx152TvjgjZMQ1iM1Y76r559xlVx56YdY56FfmRDCVbWN/+HJ8iu9qa2jIlIuSC/IZL33mcH2JmYQItx5lklEWa0909WZkmAkvmUiucvzJLst6zLb/fOSwx9JQwaj4sPgmeSg2M2lB9WQbXD3i81b2CUvE+tacIzZocE9ACYXvYJWUHwV9C2eaagCKA8TlB76nqObl4n6aIwSDMl9y16PNObLvJ/kXzIZk9RCcEhBWA4lmlVXlrZt8H3nknXr1oVPiHeQZNLx8SFGKHNZbThKhsBcd/wb/VjAkiaNRSxbxyEIham+Q7Q/ozeL23JkAK4rRQyufHgscO/b+dXi9r1nRzduekvltj3r9zEEmpKGeCD9oNpf0P38u7rrg2Je+4dzM4u+VBCN7QZ3/pzAuei5dRJHa9aR7IqfPTyJYhuoyRGhRrbCG7dfNzWBocLhhmFHg4CEEowwsn6PyC15E1TyP9HtizFbjWiQjCuZ9Iin6Eab/OKVlrPleFFR5+HnXwZY/VoB81cAjA8bslY74Qz7bdzQurM30l7hcul9dJW6+/q/gsH+KQYx/cqtjw9+329bvl6Wd3xVFfvO0n4+JJhIAiZPMvQMH0kzdmkFwToQJr+ZfidLHMQ+m0Gzp4POhxHwtsxod2vP3qiAmw651Edh42WvWXaOZQoU2b3RsVmpDrON9Mt2FccqTgMKnZ1ycueLFcDVZisIQ5dRaWN/rveU30UN78cq1y60sNrmmbRPUje1xwLTs+p2n3wkDNYQMpYqduXUth+0zGj7i6ltd1sw0hNSZIjANyu34Tkg/QJbTqWcCZOjSlTFJkcKK1yYETPMTkgcxWqqoBG/fOAhpAy16wqoNApydLLsTlau8UqVz1Tv23UP+RZoCBAtaOHDlgAVk4arOxaG+dzlsif/p7I9B9LVgWqA26gL98yzGnr5ilCXJjhqJRpM44y4PkTNwPGHHtQfufv/7bwNVZkGLjGcCIcxjhYjEZ6gGze33OC4/uWi2n6pLvSGAFh9oKRMJqakS5e6mVzyYeQQDWMcWnsk/PJbWlz4SoBFp4Me282CGrQ1mmgcmcRNVJ8OxNwTX+24ueVi4sy/jgb7f8PErSN+UF7/QO/qNc+d+P3dV6hc159j5oXEZZK1KZP/M9WnaWXI5DPm2JJn3lQsiZh3muvPw1bcjBOUvJ9ES/uIK2cnvTE009Tz7LElxKb2O6nTdcCRVDjMwUjXRREl4ebeCQ395aQlmCen1xy556ey44SfgeO9WDj5SBO5buLpmZy7pQzOojbs/DZJVZtLwfccD+TUQ9rX45dNbds6aZ7DJ5Z+P3VwTMKbbhDfQJsL4d6alB/GYhboJqOAVQxSuBZRsP+o1Lvk5qGmmCx90dHBp4PbH3owSCc+aq5kMQVY9Y2Xvmtpbs+G0nt10f9Ay6xCh3ARZ62d5hS4hYKGZ7+gDjN6tZiccLSDGKpMflpIiITn+Dvuiq+9/RO7Pn64eYPsOJoUZWwUqhv+QWh3hXDzL9e5lhBZtm11yyyR9mm2w0yBTHqBVlahoKtH6N9+D+ncBaw4X8HEsEGPJNGIMRKBF+fygZh/0sm6peeXMv+hD6n+fhQNDWDlGn9kcAB7at7ouMu/BOXKPCfWkeNTYClRI9u0qJITZo7JvmbvB1t8FZZVceWgtm1FWQkJUln0G8vR4EPFf/F85XB6TiiF5Sj8Om7NyJGb76BUN+8wFuD5hmqOLqc9BlwTMABXSAPOXggBiJBYT6J+r6NIOBgl1BKzQ9caBUMp4qWLxDUgBNzkWh3wc1XFYiPZVZpmtFsffY/S4sKoc3ERBO5UO1ReTO5c1rpl8UiZkii9FStwCpGsT+VkdfjLtWDq7ifFGEANW7LZw7bUcRauqDNW2EI9EL0oFGXvrWi0EEgp74BWkWce0dQu4AdXr3a9yd1X5naPXF++9+EH6KlAENjKgdQjsJann0GwLat7Xj68pXKZM7/jVM+V4AgVQaS8WknAoiUBnHluE7ut9eQ4SpywR2PkXDD3GfpF4e/aqLc+dL94IxkIFqA9rNxBcuhwdAc5sa1LL+ip7andqrtOWKSl1Ye0yWrLP5/1GExEse+50MlLgMlRgOPPBDj7YoDxMeyJsC61hdQa5V7M+rhS1CoSRtZf749tfFdz+/e2MtnKehSzfXL1JA+xkP8BDTJrXnHBm+JC35dUW28AMUoOK8wN7eMtmdZgO/bxeLTrNkUQ5cXUQ9uKLfqsyp7NtqvyCZLwQ/bmgRhmv/wCmHXKryHfjr03TuICCDLuvKgYbYPUszHtyuwFIfu01n5ei+GNt8Odnzl/GoJ8/5FdUPhzaxJGI2g/v+dM2epeKrv8V3rdeeQljl0JolHFFhilzzyvAQuXRFCewnw7to6kgSDroYnYaxGwd5sob7lDX7TpS7vuJhakDM/B4Y6jTWJKK2pl880jxY4TXt0sb/913L4UyQ+MbzsNFGQcy2RmWxDI9AcMJ3znTAGb7tRQngQ47+UagrqARg2bfjLwWLyh0sEeGk0hxGkvjr3iGV6h933hA/1fp+0tvTgHm+vx0SDyPKzxh2UMaAk/2D0M69uudqU3U+jo3+LiTADph6BCZmA+cDXDJiGJlUkLT8lGJS8qe3ZIJ3ppZc9DI08E5uDAI5F6yKBV7akmmhNcHsXMvf0crVsy1uCg9IGP8Hyhgs5D3EZnn+uatCejMeh8zqwFYRR9EDpyf+nPLfrSFRGnjqRTmQCYf1ygzz6/LvychtK40GgJHEbAci8M5xGxi1FN7pXeng3OmzZ96eG7KW9gKdcf5XgiWI0R8OLWpgbv8IrL36Vr/lWqdUEIcRNptBALkGjWpCmoBNpq4tNsIw26cIGGjm6AkW0AN35ViGe/Uuu2NuRX4FMSJpmH33fQnoaedmUUzV45SxZmfk04H3i1nlj/t7D52u0pNdtRRsc9tQcGMk5U2fwJp7BoixPUPqlbehbpXCdiRzCvYCsjZCAMszqm4Pj3WDmissdxmqM/lbPmvrO5+bdbrEjMk3Ey+NDHGJsR859BKKdYE8BBh449KDrGf5R2XJRhcrTWrggjVzbGpkRU/pmQU98NuQchm2lNr9u+HsEgRN3nds+tS+fvgqL7l35na6dTdOM4UpETa6fRwH4/BWeeU4eVJ4dQrUiolgTlC7J5Loz08OhzeRGVyo7/8N3x++/53I4fIs36oCVKeQzjCSnzmEGxopNb+vG48/hLoTAzoMYjgsdmkl1pjZzNA7/HnI3cImjLWywkW0etQRBwwSu0bp8JUBqzXWaGMddgZm3pTkglmjUX9m4dFqWRD6vN/+8qPjxkd77kYMQrx0bG5W5f+YLuxvYH3haJwpt1rm2pwCXKQYn1FCZBDl7UABGWQEbBzQKiTwelLd81Fd4nyTPg43dnvPDCePZpN+liVyS0cjSFCgyOS/K3+Jw5XqyFcgW2fQZ1kPWRpgjrv3VqU98Htf3a5vZfP/QIO5SwGnEdvFrPumBWbyWM/xJavXc7fS2zoUB9gJGL/eJaQ7UEoq8v0uc9uwot7VpUJh1yeqlfmnxqbLSn6FioWCrXgyiMXH/LncF//e5f97wHPYPBx+gZPBkGQVADjfhOLIon/0B3HP8ycNvQHuYTKDEYHkYuOVpkXGoAkjKlLbBhH4GLar4S6lUNpz9fw4KVAqbGDHiGKCppq2wQqLMyJhZ7JXOiNAHuxEM/EsHODwWbvr1hWpx5bBxsJCv7vHnnFvZMjp3leLnnR2FjiZZyrnC8PhBQhzjYJIS+3w+qP67V99wxDZD2pBiDjEGY/cIL4xmn3qQL3ZHQkWTtGcoXKJCUW3ZJqy1uIuFkoGul20Q49Q23tvtXwZ6bNkzfHo59nhesGiDC0MTwM1bNmF0F8XbZlX+zMyM/z2t18fFuQKR9RBY2m6CdKJKnntUUK04OdbUKOqgL4XqU9uQ6lKUlpahZg3RlANLNPXxv9J2bLhtaYyoKj3tBeyINgt2fgN6VRVnLXQvdy1crp4Bq0m6SoKKSjmWwzqYUstp3mS1S2dFcKQy8Fp6q4dSLBNSmAJohw6FtmY9+CJDCzDWOq0Sj4omx4RKUxz6lNt/07wDramkYQVTlx8bh0JzRrZKw5J3/mTv9wndHA5dY6bGDfP5JNQh/fEE848Rf61xHRO2LRJGGRsB1JOahG+OxCKq/F9Hkj0R95/Xh0OBd6Tbw2fiVBLhIHeD5sEgQer/jhX0LG4H3FnDkX3ptLX1OO61QAYZQ0gGBXK9BBcSCpQGccV4D2lo1Rb2UtOG10HTL8LpI1LVKoL5VpB3P23F//JuJB6MX3rNkuK4vYzD+471CT7RBwEEztKenp20ymP/zuP2Es8HLhdSiTEw7+1QeOWSwnfZGPdKyYJqiP+N1BPVOlMYBumcJOOfFbCxqZRNC2EalTLcaw1MjiCIHqlUhp3bcJcKJS+MHr7whfYBWoqrvMcNw4GEm+mph+iCyE97AGxNqtj+AYQ3CnzxL9544qAqdCMDKYTVUNksKguqdEJevlWr8e9FDP75nejb71Qd7Fmx+IFmdO0+feUrgOm+FDu/1uqul0/EdfEwjqSOHlOg1iFoJoLM9gjPOb8D8xZGuV4UgTwEjaPaGKYwxk5zMQ4xoStyI73rb7lV33PGr8gvKN5bHEm6EIzCeDIOAg5ft1qW9DnT+UnUdf6KWTohCeQlqxuYTOFmTpZmy0nEGNWhpAcznEcFYK2lkhYAzXiCge7aG8nhmGwmy2ECOQ7zSCAGOhNa+KA0DTO7+kdPcszbcNmBWhmOG4TE8T3+AuRg2CE7fi58jZp/0C6UiBc3yfSJu/tSL6tc1Hh64OTVquMpc5BhP4ED1fKtxSTE7Portp/Q9J855b1et7sucnkJO+kiTB5GItOO4WjiehnoJwBcKTjqjAcef2sAUBtSrWOlkcpZpaUmT30zWSBdC6Xvejvui+9bfVH7unhsqI9ingBJsR+oKiSf75uQWrF4YjTV/obuWLFLSi0CZujDWXww9ftIvMB0bnEXNplJZFgWJKtP1GsCyVSCOPwN0rcrcCgli0PYUELyXcwwxtrIKLYLAFRPbGk7UvDqc2P1fsPubG/mYV7sHcRWPjafEQHe/X3lLLjlDCf/1ornnmx9+88/vmDahViPMHQ5+jxllCNZwHPeSk7uG91ZeBaDeKFucC+SMFhA+ZipVDLFGz184LmLpUB1e62UrGnDK2U1RKCooT3JC3UFcWQJs4pXPkJEm+C/p6EgL19v5gLrn3tuDi0e/P7o7gTofwfEkGoTUKLT0nXVisy5/Erctnq8dLxRaeUYDLmMQEjAGkxCmKDnOIRDKFJOMVEdmtBxWjqYmhe6dp+Hk1Rp8X2Aql9uyE8IPIzpr+t45OsGqhAONBoiRh+qyNvJFN1f+VPPBAZNVPuYxPL1GkhxUB/FsLH4gtn9vX7V0STNu/IUuOG/UHbl5siWPtQhcVbRUMfWISkeLoC409uPNXxiKU8+v6e7eGCpTUkQRkYGh1DUtYJn4P8me4ROKndeOKwLwZG7z76M77r2+8sLKusro0TAGfKJP+mCj4LUsOTGGrh+rzqWLwPEiAdo1iqCm+pDwMWZbZtJGGysuS9gSE3bgd9A816r8vZOeJWDWIg3VMnLhIeyYOyaT8CHpRDQYMIihGbgSQ46pbeM6mPyi78afb24Z2Jwe+7Ecw1NzHKRCsH8iFAeTYCO8eMW8Z8c5/RfQmnuF7sh1yKKDTmeo0BvQGg0BuG5MWLmoBjB7XgQnn1ODuccpqNUAGjUhSG8b4VwkaYlFtdQIMRKCOTiJKDunIy2lt/WO+I7fDsoXwuDu0ceKQnyKGATrig9Gud7nLAlqlRt01+KlIPMhaGUbaVK0WGo/05G9nERTblMFVmlYYqVRQq2iYe4ygJXPYiKMRpmZwK0cOrc8JxAw04uDPQIx1EZdqIwLObWz6uj4apH3rwg2fdWWoCTAGoRCH8MxPLWHyBiCpJ7fu3LlrEpcWRMXxCWyPXeB7PAAfIlJfyWVkkSTIIV0PaXDBkBY1jBzTihOP78OfcdFEAdCN2pIJIPCdpBGvGmOgPPj9CsJBSC3nHJ9iJuh4+3cEP/04d/WXrfr/0rj+ih5BtkL8Acy2FPIzzp9QVSVP447l5+sZa4pIPQsFH6/9kAuymTepw4103xn722mjRjDiGpJQK4N4MTzNfTO4/JkECmimmHKK8s/MB1+G4UKKpMxVEd9pHIQjcmajGoDvo4+W9s+sG76eRzzGp5iw1I9xfaNlSvX+Dubtzw3KHivifPOH4u2/AxddLG5IBIxgjAVla4kohhcgKiBShcA3d2RWHl2HRaeFOo40KJRldgioUlJw+BxuUtiegctgyQ5UaaUVH4B4mpZ+tvvVV+49RPL/hoXzCOdQPwDNwgp6KVYXDCrAZ0/1d0rTwE3H2gdIG48y0CcFZS17dJZ8kuDTcgaBMM2gFUIFIhBb2He8QKOPxPTPgqqFZPJSekM+TuWhhFbq5WGZgVEfUpBWHVFFIKojYYyrPyfkPqqYMW86+GGz1qCXMcIcx7zGp4insCaa65xfvrhD68KXfVKcPUrVFvueNWao4KVUDrUAkVxsIFBgPRR11aTJCkZgi40BA29cGWTHplqCfUqQbhuItlmGALZGpBdSAnxzGpFbdjKKQpR2ivkw3eqf7rzit3/SM/7ZYmi81G/MH9gwyAFW4+fIXThx9C+7Fydb8MOO+RSSOd4gj3IttqmUAMe1jmYRiyQfgYxCrmChiWrBMxeCBDUAJpNNBAZimxbcTJUYQhzVYGGegVEoxxDUJGiOeVAUAU3rN0lw/LXdSx/0Bz5v83TzwnHMePwJI79cgL2zfYZs1cFbe7Lhee+SLX4Z6ievBA+Pm46JDo/qhZwW41wUTeLDIGOGxp6Z4Ri+Zl1Pef4kFadepkVV7Bb0UqAMNqYySySbolMKowFmkgCLHbywt27TYTDm513r/vMjv85UgjER3OR/nDhsV2rOrxm/APVueii2O8MQAdoq61XkJYhU30TW55MU44JZce0kqXJNSDjbag1GoaZ8wSsOFsD9pHWy4Y63qAcs3kF22WFgmmYgAxqGiaGYijtERK5xnDXzVJZRLXvurrx7c6Zs389fM/XqtPzJSjYecw4PIFewDRw1DXXrHHe+oHbT2rEzZeonP+yuJhfpdsKjkC+Q2J6CGNUSAChJLUPYIrAxTyg0lFDiKghYGZfCCeeV4U5S7FSDlAv86PlnORGiAAAIQ9JREFUOIgmIi+AySEMp2+GPYK9A8vAw7p92s3rULjS37E+Hhna6Lx2w5eHbjS9CUll45lsEDK966uKshR8VbXOfyUUZ8YCGSIyFD0ZYun0/+yRGUyindj2j9nX5vSxg6RWYRGORScBHLcC8wmMY+Bww5gVQx2fZVIiaJkAKI0CTAwpCOuYOvbwLsvmBIjm1AMyCn6s4+oPoz9afAsMZLPaieeQ7ZU/Nh77SGif9yVd6Vq1qqO5d+85oPWL47xzYVz0zlDFIkDeB3CJhyoWMUo9azQCXPRGRiLkqFEawrIQnlbQd1yol51Zh9mLQ4gDgHoF1VtAux6Hlkb7lp6WpEMCPQNaQzJerNVrUqD9FhGFkePtuj++a3iD+rMN3xq+/0g0Kj3NDAINtu5SgpNb8jGdn/NR1TYfiSyQHJNX45TFOO1XsHJiaTXCUr/b7VoAk9GPxM5JqSGMARBT2tIlYOnpGnrmamjWJAQNU40wvRAkRsvKG4mj4vg8pyujAOM7QdRLEXhYN3JdCJpoGECE9fUyrv/E19Vre1rU3du33zuROVfaCsCaY97D4Q+xT0IwU28S0D5z5kKVd54bOc5zlONdGLcXFujWHCgXiaSRKBFizOBpJuwR6OIROMXB1gZk4wOI6loUPA0Llzdg4Vk16OyNQYWg62WmOnC5C4dWB5b65QWD5WEN9NhgkU3Oi/1bOkgdewXtVssubL9ff2XrLeE7RwZHKqtXgztoOiSf6PGHbhCmdch5+RNeq/zO/1Xti1vAcUMN2P+QbZvOpg3SDur0NBOQo+W3SbUVcJCEBGomNzU0AqG75wEsOxlFaNGDwKaoDMbBbCb513gOLkY1AqC8F2B4s4Z6mQQ7wcsJgf0ajRqIoAQirg0JiG8USv1EhrVbmlN3P7z/eZPoaSaJ8Ywe2RzAfgYAR9eJS+Y3Juqna99/YSzl6brorNK5nA95jxSftasiojhAQl6UtZY0hxMNIImyHQogqgotAyG7egJYcHJNz18ZEblSsyZQZU2g8B96D6hBjEdCfLyJaodRejMxAsuppTkDUxnXXk5GWoK/d4uqTmxXf3vbZ4a5Df8olxWfDgZhGp+C13HCqjjOD+j2pYt0rg2JVqxRsNwJtn0pg27MsoryIsATzLRapxTKllqPKSsbdZbcmrtEwIITsOcU6Klg/T9rFNIeiUy3NhoA/DJMjmgY3SGgMcHin8iTSE1vMfI3gsAUdWNyUsbBBggrP3OkuMsNy7fU69v3ZPQY92kkOii+/ukwMtZ7tTQGcb+kGt7u3t5iX8NvO70pvPO14zxH+85Klc91q0IOdI6IBIjT0jBqsKA6T2Kq9NMENvnjqKmFrgudd5WYvSjQ809twowFoZBS6WYFNYGEwPIhOooUs5J/ylTckhKIFnOMN5dvjeRowoDrib2WalZ+AVQcue62O4I7994j37zp+qG7TEnxSb+nTyWDkACYiidcOLu+q3w1FOa9QBd6kOLdNJqYSWoBjokyE37XpB2y3ZQ4eGKn1yE7Ca3Cb71K/NZw3AkAsxby3xp103uZuBe22cosOUYLBr0D3Hd9ihmepvZgZUmD5wuQXgzIyKWVR0YmaoJslAGCShlUsM6Jm7+TYWNQd7VtCXfdtekANJPWi7AnBU8xb2LfVR/HAV3lU045peWhsbGT4jg+Xkvn2XHOW6pychX4uTYyAHidXaISCSFSNPfpATB0iFTpM63EyIuLPLXY6RxXAUQI0NUTiwUnNfScFU3IdUYibjoQ1oXG/AFqHqBHYF1MCgdINd7wreKkzxS9LEUo2QMTupLD4KrIzTtuaUTC9vvCq+74RvweGBmpPFn5gqeBQcgkG/Va6RW+/Z9xy7x3q2If9uJzCGH7pK3GmS38Jh4bocBsqcBorNmOKEoJMJM6gZwYEiWEIzQKczRrALlWoRecpKB3DsvLNVE4xqhJJUCprJ5DzL4jGhTH16JZFrB3q9bjOzkEyfmYv2DxD+QeRBumtEe/BzXQzTrIqF7RUWOzFMGdEAR3FloLt8nG5IOl0o6JfbyIg2TZjWZjShO375f0UXymkBhHAOzFFmm7r4MdB2oZQEuxOMPr7JwdBvrcWMjlquBfGEsxExxnofAcgHwOQvQAHLIlyIUW03VGL97C1WweiZCqRtXIKFWomhS6oSHfEqk5SwO54LQQOmYHFDw2apIEwjClRBBjE/LT0kJXk5sP0OIIye+bB40sjzSPFwv/ESAWt6H9vIjj2PF3PxyP7N4Qv3fDl/d8lbb3j0eudfkZahBoGMEP0H7Xaa9WKv8/qmVej/KKgVAhZoywhyyBGyS1xmSyGuwCBZRJYwmpb7DnbzM/tlJBFQb2GKIARKOhSTxm/jKhe2YhLFrjcsIGxPSpJd2Z9igMR6TrC/IamhUFI9sljA9h0IqehMA4N0lIkIdptC+18rkejgw++FNBvOwYqOBu3agPeZ7zgC4P3yb9Ysnv7R6Z2n7Pw49ykst00j6WMW2iH3q/5pLgP7OOP37GZK22SJWCTq+9eE5QbyzRxdwKJZwTQMoitLZ4iuJ/XPnJQGMxkKw2i2gRh6Od+BmnD/114msR6IEh8FQ1hICK1K6voGd+JGavaKi+pQHkWgEbkAAhx2CKRlJSlpGXC8QXMYm9rV7To4ERA1UjbAExedbYJPEzgPouInJy4FSnpLPzfnXDQ3c3/mbklyNbnmh8wdPdIGTdzTjXu2pJ3Ai/EBfnXwT57ojkt/BxYHSiVdph7lWriEjrRiJsgn+3OkxJ2sGqw/Ng1XdyHwhr0EQFatAoXz97GUB3H2+72WBtSkvlnTgjJvuImAaUTXCMAUBDMzmE4YSA2iT3YZA3wUwZRkIc3QcWr0GT5WgEy1JvHcGosZ0uaAKmv6UOx6FZ2621rjq+9yDUqxuVEhNewR/xZLgtLI0NOfXxhoR25XbPjE5+1+trN/X3Iw32478hQsCz//Ef3fWf+1x+ZGTELRZnFHLd7TMi6S+Im/UZUbPeIVuKJyghT4JYt+ic7NOuN0dha6DngfZdUESrYnx0lEnXivwmcupMK4Al0koSymZe0aqMNhVZUUMBugkaAgDf16JrdgS9SwPoXRZAS6fSUSh1UMFuBMIYmAqBoglNuzYFbaTgtAJbLA/MBEYozUFeAvceWHEeepMMiUJsAUq9uN6eDXF1ZIf6yL1f2P1p/NCTWUV4OhuE6SHEmmscce1H14p830d16zzQqN2oI+MtJIFkYo+TIoTplrK4kWxWkv6aBTjsT14BEAQAYQBQ7EC0I/ZHYHCJtPAsjGKfpiRvkfEOWQkYIJfn4LM6CTC2DWBiN0CzLsBztPB8fMKoYZ6cBSQ4xryDJQJF+i/cH2a70Hio0EGSLUDqSCZ/wdAGHB2AiMJQR2FDqzAiGXahm1rDuBbxTqHUuIqhAXENveIYe3NlLoccgxGqkkj8PE1XIWKldRRUc6CUi0gc6UpPaJUHx+mSjtenVdylhcxrATnpODnhu3mc9HHeY1QetQSi2jfdhAioPcgqrishYuZBNrF4ooxr+trTq2kdKMSI0b0AARX+eKE9hM45McxYHEHX/EgUumIMLCCqCxGjZjC205LmEl47MgiA+m20jBgBLopB8HdCzVMUwvkDIvi1XoS5lVZ2SGvteCIWrvCrExp2bdTX7Vkff2jnz4bvo0XoCYIgP4MNQgavIABybSe+IBb5z6qWBccrvyUyzMtImpmZ41YAy/xq8Yz8u7EItt2a3rMUbOhDchIZ1xZa7c33oyZyYAldaNMwdyFQ4xQ+Sc06i6rSmmPBEQkSwqxDyOMAuJQJwKc0bAJMDbNhQFEazFV4rgbfJVeYz9SqDuGw/RaJJ2Jp+81retNJpNkTKXQ0GhyZJK48vpeqxvLXs3gumoVWIZ6WShaSSqos2bKvsa14rdhBUqiAZHgorJK35QJJ2bGsihKtxhkJFFzBTfpY4/qKhxpg1yomCQQUu2PoXhDoGctCaOuLhVvAz2kd1rGow9xa0iUtYebZEtaWmgBRUiGIOpPJDBH9KueHaOE3waBxEtlqMcCIaPykp2LHlW5Q8mD3g+HWiR36wxuvGfoWfuxotiwfyfF0MQgZYA9EHced3FUZVf+qcr1v1S2zMJlnOBvtg53RIkzCiuwTnbROm9eZqoTdUzK3M5LwhGYJAYlVIJfnikTvXADPYy8Cm6qm7cbSwHEakwleTK4C8Qy+BOyRgPERgLEhgMokbx9DDfzBQrg1MqxDyEtokr5IFKUz3aFGKzMrrkIGgxIgqE6WCtESwN6ELegf02Wz3zN6mPzdjKReUtZNw2MWpcziwtKFPjViqdS8/STG6KRHYAs3WgC2jtFl1NpvVdDap0XX/FC3HxdDsUcJx9UQNUDHgRBIikeTmraD1QF02BIoMb3m/gKTJBR8OjZfYAyEwRdQBYP/jnGDaZ+jJcKByMmjXIjjDm/S5Ymt0X/v+LX/r1Pbt08Q3yGOP2Cv4OlqEKb3QWDip3DCH2un+AldmLVCF2Ywnh0RjhZFQiNJPaar0jR6Z4uGNA8s2wALUDXaD/a1+RyumkjhhhUILyehZ7aC2fMFFFoFhIGGoGFy1vbps70XGGJQoswYh5gJXjxPg6sZA4EU82MjbBwwb4DehesISlaSTiSuwMb7IUNnpeyTc+VhOSP4g6k3JOm5NeArAxLnc2dtjOQ7Gc8iWdnN9bDPFZ6Pzeom19liNjKVGAajJyB08k1Q/hJX/9BqcMQq1xKJ9llCtM6NdKFHi0JPDG7euB4BCIWyqBiFOQpTNAmxLiMQ0SAgGMEm/1AIyQSNycQXpnKAOo6swI29SekpMGDV5oeFI7BNGZEO3vgOARM74DvjO5yPbrt228anklfwdDcI0xKOXYv/qKMyvONy5Xe+XRePE9orhqAx48TOKi9mVB2Y1hNhtMH5d6s/mS5eVt6TH2uLfqDXxNHI5U38VhwBBCHPlvZuLeYsFLqjCwEzGsJGCnIiw6DoiDhKRSFZq/NuVl80DDkPk5JY7aCOS5gYEzA2qqFSRTkiq5WERgLVQjkTaktwNkxilz0DAaIQwSogpy4RnRHKG9n13JTxUMaWpp6R20uMjt2PAXpljRHNNpYmpew/ZWjNX3H3odaADYOU3hTgFDTk2zS09MW6ZZYWxb4Ych2xdnN4V7WIAgkqRMFaXN4xD6mUlFz6pxwyY4So8MeXlwvObIPZI8BjMO/zd4TtRTCHjIaC0s28SpAirOmgR14EdFLKeyTs3RzfUdkjPrr1ul3XZ5KGT2hT0pEaT1eDMD3hiIjijhUXilD/iyrMvVChNqFwAm2SZPjJtEBoi5DslmamOg674iV6vxxRmpWQhtUJzGpUYnU61horE1EsoLUDYOYcDViy9D1B3gSFE6Q/baFu+DTyxDENMrw0YVocqeGwhOkC+HgKCoizq1wGmJzUUKoIqNXZg4gjnqiYfXelRu1bDsSTSW+MjjVtibeQmdAZl36aEbHXJglJDMiTIXvMXcMGkn4wIUo/AknH2GZgWgTTI21aF7pBFGYqne9RkO/WwivQRNca1RPQA8ApiJeI1YwEam07SGNKXj8ZKhsS0E3FUMHcITYOVE4kj4EuMRsCk90hfAFQhojKN5L1AHkdMAbF1cpxkbjf8SojEia2xRsnt+v/3PJQx9Wwfn1AaMOnUHjwDDQI03ML6O95xSVvjaH1w6o45zjIdyCAPRA6Qlam/S9HmitIJwOrh2PQvp+QH3+DoCj2d4ZLcte0Yi1LEBAEnDj0cwK6Z2qYMUtAWydLtWOOAFk3cCM4gQlhZzijGFrDv1NIi3E8GRGuVnhYsUBZsliSkQlCQL0JKFcApioayjWg/WITFzs8PDNsexDmJGzqPGuErFKIvSb0tvGEbYKT53762iYJ8cch1LdwPQFOi4Zchwa/Q+tctxL42slrcHIEBMJNaCyO6AgEJQ6xiIBJP/TdUTPWMA/xqm+R6eRgWR2DhF/X8A/gxOYORpNHoMtHuQCTILQlRsEYFCxG0j+cwxCug5IIoBsV6dXHAMYfinY2KuLTbtR25QM/eqD8h9CDcKTGM8Eg2IHTjZbCtjkn9DSmvPeD2/LOuHV2q3ZbEamutEKdvwRuyCVvg1bgBz0NF4yQhkzCfZz2ZhLYOCJTsLS63baQZgqeMWgkZIljTfmF3pkAvb0ArW28MiO+APG1uEHGWXFJk8IJTl6wwUAPhWYkGySjyYwhg/AkaF8SaAq5HyQW6IO6Jjh2rSmg2tSocKXDGHSkjLeCHBAgqWxqFZGtmhal5XnlFh7i+0Ej0FLmAJuDhJPXGie33yqE06K0V9DgFASAT/V+cJE610E3HBsINEInzMQnoAEXh9GVd3BJpoqAYRfiCMw4N/yZxBsgA6moU4lEkrmMyHbNJAVNPYgnPiUa+fsU2aAcEicLqNSIfVA+XVEV1LVqlEWuOiZhcnu0Qzf050XU+vkHr31w9KkeHjzTDcJ+ScfWrpUrg0BcGntdr9HFmZ52WzCC1UIrJL5ihIshvbTeY5I+JGORJNkomcfwOPNrmrDjTVrpGdtuwQEHNVeRli2WFgPUqtFatLYD9PYJ3derEaZLXI8RGQee7LiSI2G0TXeQcUAXnCoNvLJjcpB2SpaFUfz4Hk40FyczOjpYxlOM7adZqYQjsMuHypFCYlCPcgCEiyJ2L/KMkF4cm36QPAqTbK4T8+84bSWir6ktnCuYCNmM7HXiXAWH8Sa2N4VQ/J0QwUx7bVCCrPRpGhNMUQa/Zyc3ufmSPkegIhP3214CU7Zk42FjPQQaEfEV78dkPpAS3UUSFFeglpeqjIEojztuVAKoT8abIYDPLVsurr7uX7dPZDyCPzik4eMdz0SDMC3piK8KXaecE0biw7HT/ie6OAvALZj6YJTiFxgRhw5oMttNronQCETMmjBnIiGvKU9zyGFaJoz4HEMQmYaXNmX6HXBfGDBjPgHdfQTUd3ZJmNmnoKcDoOBzBQHdfggNtp6qEzZJyEhHCivM+1ZHl5ZKqkBwvO2gzDl9TwvEKKFxAKUZhxRR1k0KNAj4NxLNNpMO5xeecmRQvBiX08wXEiGFNJHpc3hCjPQzkntCxjh3k+Pkip9tF2XJVTooLBXipXL4mGkCk4NGoQGXD5LyIAOHyBhReMFemtUUp2MxuAO0B3h8mBnE72KJEkWrMdpCW1ufEqoyLkSjLNwY4SNN/XtX6S92tdW+9tsvjnFosBpc+BU9N08rQ/BMNwh2pOQaiG0pHn+RlvmPCK/jeaowA7RXjLjoiKEBLUbGH6DJbyuQNhE3LUxIku2cuDe1+umVN5MTSD9Pk4WFJZjUNdIQRaghwViGjjYBs2dq0dcOuuAxChGNQ8yHaXINXGmg0IFKg8ZYcIVBYM6BVntql6CKBa7ujMVRQsqAZhYXVyKBk51/R4OAE5pQ1GQcqFSHsQWV7HDiWy+CY3i6MAbeS2BL9Eao/kIeh0EgkzdDFyXpHzDtAPwdjYqfNBQ2HCU4UktUaFTSTfqDzod9AYoCbeXAXl/MU7iIGvcV9i/oyjioJvaWVkVONSUElUC5Em7Mee5/v2rpth8mLMdPU49g3/FMNwj7q/ZKCX770pfFceF9OjfjAp3vQX4swiKSmhM/a4nfnyIgjTFgEI7BIdoMnM3bZ0pyybClTSMSkzxvJmeRmCwlmOI3Asi5AD1tGmZ2C+hpRxow/j4aBsxHIEYXn13ELpBLbnMMuC2s0xkcgpnslA/ASaoRzofxtHlNky6mCc+fJfpxFEtOPAsHDQV5Afg7eyRMEmJKjPw+XSv289EocG6Ak4Ess2UNgfWfTLux7RewqZc0VDBGi1uRTeeJSTJidYGvHkGPMcQxngCzJDfLoMsjEDVL5DZ5iKJW9XAyDtQPpI6v3PKDvbckt+cZYgjsOGYQDpJ4JI+h9aQXaem9XcvixTrX7YPXgk9oyN6tRmoNAxGmxTFTnjNZhqQffhp1k7UcnC9PQU/4XaJ+Ndn9pPMqWf2TDEYE0MQeihgrCxpa82wYZrRraG8RUEDNXJwSWK8L2UiQt2IMElY7aKmmSY2UYdwIwQaBVntaXdkbIHJRLNKQsQClEN5rPAT6nUIOk/XnS8deAxUxMIGHPHW8Wks2ABg+cG8G6RMxdNh4IiZRyBgBircYNkzJhyTESLECBm3INGhUVkRCVAJyupQnEIhcbEwBVCe0alSkUoHwsOEU2yh0CPfISH9dxLlvbvnhlh3mOZCwBsQzyRDYccwgHCrxaK9Qx0mnCOW8WWj/9TrfPQML5iC8mFZe6lOwdPDWQzBbSVon6Zd9H61M7GD+bvHQSQrTdkpSmALTcAM2JY7eQIh6NobfEfkC2gsAXa0aeluF6MwLKOaMmlUsNBoHBEtRnoNAT2gQzIRWSsvITEDqMU5DhNQgWENBk568Bw4X2MvgCc8oBKJ3QAvHk5byBLyaa4neBhOYmUZOzpfapGgKHzZ9XNbNMjkE/N3CjaVHlQ6BDaSURA00Ulfo+riGZpl6GRQocN2cFDqIMWEwqWP1f7IWfe3Pzxm+blpYgONpUD58rOOYQXhkYBMF0PhbsXflrKghXxeDfpPwu09W+W5QbhEfzUCIyNHUT2D7oxmPkBbwrcueeZngBRMPgid71p+gfgIDkkowAWbhSsKQjEFSJmzA0AK/goSiLTmA7jyIzgLojoIQrS6VCVE8gBHbWgs0ErhYq5CMDMMHuS+JDAbCA83ETfIH1GxkjIOd2KavCpswGTphPmNKsQnHACUjKeeAKQDjIrAhYXRhykZEACOuNmikLqMQwMQPKtZChxrCmtaNika6Sh1WRRyHWsmccHwfHB0pUM2oBiHcBtXoe13C+d7dv9i5K7nNmCi8CNRTGVB0pMYxg3B4w/TYcQ/76tVr3d/d/uXnxk7xzVoWXqxy3W2QawctfVz2qLhOYPwsHNo0yiUoQB5puGB1KNkjtk1L5h4ZY2D/xklDu4nMdyxziDEQOAvRSEWxwRbg4SNsE1mHHBCtvhBtvoYOF0SrB6LggiwCiBy54pQI5RUak/0Rfd80L3KCEUMM4gcyuQSTmMRdkMcAlH+wZJUUqlA/ljUquG0kPjIsx4RAtH/jq84MNsQyh7xISsR1YkKGsIpGAADFUpTSBlIslXSw8117WJmIK6GGRnyXaKrvSR39cOdNe+9N7mjqDTzjwoJDjWMG4dENy2HI5BYCIN+5+Lgo8F4hnNwrY6flPJ3r9DjX4IWc2sIuHcYBmPy/Rfpkyw0ZoQfbTmTUqJOPUYu0yasjTiCjD0GT0PY9ZCodSdeg/dVUMdDroNCB2MjRI+Djwd4HpF/IO0IUPDQMWuaFEHkJIi+0aBECmYmJUMTDmj2t9trAj+lgHDQA7B2Qsja68iZVYjUNTdUQcxaknq5VqAigpBoYguElo4lPMGX8QY1EhC2T4YuNZfQ0MtJpxA04xMqAyAI0ECHoINqgIv0DF8T3LvzZuXcOiEQLYz9Z92Nj+jhmEB4fjiGVCkfj0LbobKX9FwvHe2ksC6crrwu01wIaNRsEIGELOrlUJEuIPzLSkzzJjb1IypBpjoEL6zZ8MFgDzlNYq5FpWpqWrMyUNE2WwzY9JTQRxpBgrwHOZSxp2m5kasCyfpIxRERRphX67qRXRWwLBmGI1Q0GFjFXAk1yUzix+7GwZtq0StDeSGyNwCcGJSFqULG2IpZ4MEnoIYBDaIhVjraLZrcRgxPp+yBSg0JHP/qTP5vzyyvftg7dtNQbWAn6WEjwyOOYQXj8Qxq68GTVWb12rXvHZ799ZhDGL1KQu1g5LSdqr9gifDQOHrJ0UIaOyP50zGQ9pvkv6aY0k5dvEGfsUwkw0z5MEEciG0pbsROGqIRn0ObguLpgCQ04Zs/kL6hXg40KGQnajHltwET2mbGhja1+WPow23PBbddM5mIl0DFvQNgHRh8iypHxAwKEh/gHq25k0ItaK/RApMuVHNOZ7FIet4E9GtGYA+oOiPWNUqifHb/YXz/45W3IjLhvXmCay3RsHHocMwhHWVYcn+S22cuPb5QaF2kovEQ77gnabTleOW0gHB+0l0ciQPSrI2rrQwPBiUNuDSZGPwNNtqoyOCto0hnmJwOUshDqxOOw1Qh+Py110nFZtifDvERZPOMdsP9i7FOGut6oEpk2aqwpYsehxWQY1CSVE03HpuG+5lPgqgV2J5pcB018gybE/AGWNNDrwNNyJIKrsesQe7UAdL0JqlG/zw31XcJTP2jrzP1m+/Xb9kyb6cfyAo97HDMIR2dkRFVMvsGMvuf/WUvp1t+sDAN4AQhxnnILq4TX0qe8VgHY7ytdhjkQpJBSaTjJHEMFR/k6ngSE/U9JUJL9JpgF0+xkGeizdGQmuWm7J+2MJnq2TFhh4ZYMB56Oh5DZ/ZjtUUhhYdOph8INRSkikk7Co7qnFi6ipOklFg4oaUA2phY0QMUjEOpfF3Lu73VD/2LVLHfjDTdsRtw2TMsJHPMEjtg4ZhCe2CoFwwOTdwX09sxoLYetp8f14ATtF5+vtFwi3OJK5ebzbCCYgVmTxBBNVEz3mz4dQ57CiTqTtLRal8YIZLENyV1PchMZ9qIsG5QFWdlaoTUixghNa402DFM2PHGtQUHwEYGJGKKIlQTqmZBEWkRlSzRySFKLJLCR3qWieKufE3fqenBb0ffuzE1s37p9GzSOeQFP3DhmEJ74YRAGlHeYJlNOfxQCvNY5y2IlF4IsrIYo6tN+8Syt1Wxw8h3g5XMaCfwcz/C92g4f23tN4USKAU68A8I32KSjIT+x5C62SmFR2InASdo0ZXkZrLEgcIANE2yyESQRsSDzE58MEyaQYxCBEzVANaIpIXXNj8LbQccP5XPePeXx8TsXze7b8eA6bineZziwerWAQbpWx/IBR3kcMwhP/kBlIwkwYCFL+6Pk1qxxYGCg2Lrg9Hm1icYpMgy7RCF3mm4Ey5XjtUhfzgKQfUprDwtwqOugkYtRuIYamAr8aYWAejGYBojJBQ0w0qppU2xveyuQ2ZTjfi2xQzNFDRJ4yUx2xCToKDaNnXEklN4T62inI2TZA3lHXG9uKLa7Q8HOyl2Vtj1VsQdq+4lOISFpP4UBAgb3r6EeG0d/HDMIf9DJScpB4O/RoW7grFWrZ5R27J7TLE8V/Lw7Q8nCiSpoLFSxwsxli8753Y5EriLRCho6QECL0hqzmaaL0xA9UDqQcdLcv805BO43gCa4uiEAKkLrcR1FIxCrcSllA0Uf/Jy7S2h9n6rV9kqt692z2oYevu+h4UPMZgmrVjnQuk7DIO3oGEAInvzx/wECHtfWVBEP4wAAAABJRU5ErkJggg=="
LOGO_INICIO_B64 = "iVBORw0KGgoAAAANSUhEUgAAAQQAAAEOCAYAAABvk2lsAAEAAElEQVR4nOS9BXxdRfo+PnPserypURegLVpgKdaWYmuwSIq7S4FdXNvA4iyyyOIuS4O70+KUFq1Qo96kaez6vcfm/3nmzNycpPcm3f2usL//7Ibc3pw5MmfmnVee93kpIYQwxiillKVSqX6madbqTHddzbXwN/yZUqratp2uqqpaTTZtFMek0+kt8vl8paZpjndKht9EURRDVdV4OBxeI4/tfoK2trYhuq5HHcfJBwghJqUq+juOE3Bdt71Pnz7r5T36+zHGlEQiMTKfz5NAAD0JDRDC8oRQRVG0aDTaSCltK9FXzefzoyzL0hljFmPMpZRq/CSU0mg0uoZSmijR10gmkyPk8eKZ8Mw4WGlvb18+ePDgbJGxwkGRZDI5WJyGMZazA4GAY5qmjkuXlZX9TCm15fHTpzOFkFlKff0sd9Gqc/oN7xesMoxoGyEEx1cTYpf/vDZV9fn3qdyStclRP6/PVTMrt3XcpgMa1yeYSt2+4Yg2Im25JO8Q4roucZlLFEqJbmgkGgoQZrrN7R3ZZf36R0nfMNmoEu2HQYPL2mvLqz//7W6VyVFDAiYhJE8IcQkhAULy+vLlCVcJsLDG1JaHH76rccaMGazY+0kmk6MppXrnV8zWdd2yLMtwXXdDRUVFqfejpVKpLTGFXNfF9V3MJfF+tFwut6KmpqbU+9HxfkzTVHRdt/FOcE1Ccq5lqcFYLLaKUprpNh8pPm/YsCEaDoeHiWsy+T3OgXtZsmTJ0p122snqdj1+Dy0tLWXBYHCY4zgZw3VtU1EwPxT+R8Yc8W6dYn0TiUQNpbQf5iGeFQOlKAofN13X84FAYFmR55R9+4i+ch6r+LtlWaqiKKnKysqVpERLp9MDbduOymeUE7rwtlzXtV3dtcVA8GYYBrVtu/Dv7nNcvCSq61hbhX9iUdviAUv15Q0LwnEcLnjwFtCwoFVVpfl8vse+hsFsQgKamCzUZMwm+CFEicfjPXWVzyyfVU4O/PBz9dDNdV3XUhT+rvlEkS+jmMDr3oQgYZRazNECVhD37U0cZRaZRaZPn67U14+lhEx16uspnyAhnZCtBs8gixY1D/jox8Y9FqxMjVvX6my9rjU1LJ4x+yayLJyzNZJ3XOIwhzguI9S1iWNbhOBHUQjB6/ZukfHHo3jEFCFUrVU1pfantnaiKZRoeuAgfX6GRNTV5L5XnFR5VN/Qt7Jy4fC+kcXjRwVWHrBL1YoRI/osIIS0E0IS9fX1bn39QnXixI/orFmTHN/E5eOK58LYiPHBeGGB2OXl5QXB173NmzePbrkl5AEXQkRMcsW2bdUwDEz0HucFJqJhGOiLfgzzCE+tql3ezyafXdflv1VVDYr5iwXMr4V/jx8/vsd3i3lhGIyZpmIEAlQhea9/jjHW0NDQUz9XVVU+WELYWxBKmNe5XI6JDa9oQx/TNF1N0/gawjgHAgF+qh7WrRwnCFosXH6fXQQCGk6k6zq1LEuTi8Q0TS4oehwJMWhCGnLJKn5bvoEvtVgK30upiMHAA4VCoZ4WGDNNqknhQynFlsulKxaoGNySzTTNwmTFOcTi5kKho6Ojp76FnQWLW/bvSfj5dhElGAzyhcLv0TT1ODMsyyIka1vq5JM3MtJQ7+JGHMbKvppvDnvj41W7L2tM7brTSbN3aU6ZI7KuoaVMQnImtnwMLy7rEMJynkBSKMMFXHzWFEI0jRKXUeJajFCiUlX1Vqy3BAiBALFcQlSFmI7imGbOJsSlcQLthERJSyhK1qdGhBd2/P6Z2Tky/Wkt068yuGx4/8i8UQOj72Wz7Is+VcrK2bMbPBmDMaibqYjx4b+FwHQMw3ByOeJSamitra0lB3h8cjxLuAm8Q0XMCcW36TiVlZVcUBRr8+bNI+PGjWOYw1IIYYOBdiTmZMm+rhAIouHepSDzv/uuk0E8tOM4LiF51zQN3LRuWYTqAaMwz3pbmHJzEufzb6o9rj2hAcnNmP87n89jodNAIOD22pcQbMT8opsIBDkQUs1zXTeHfweDwS6qTjHJKB4Kx/GXgIvhJ5VK9fhA+Xxe03Wdv3B5fYwRzkkI6bEvJCdUIzlhIFGFuscHuae+uFdxj9xEkTs8XoD3ckt3lY8thAHXEvACcJ5BgwaVFGI1NTVQ0/j18oy5+TxhVjodGDo0kiekrI3NrCNvfZrY48VZK/6w16kf/rqxLTO6xQxracsiFnZ6XIpaLlEVqJSEqt56w3rH0xDHUokNJQkvWyHYWDWqEsNQSCykEcXJEjOfz+u6SjRNIZRQYjs2cR1GHOYSTVUDilGhpvI2yTmUQE5ApjtmO8kQRjIOcTdmWHhFB9l23nq6bdmP+RMbZr+T3O3kjz8fN6jvW2cfNuSjHbeP/hBvmIq7IUtXtZLykJPVY7oV5DdPSDBIVNvOBMrLq0sL3UmEBfIBbEaFxSHGu/CeSnUdP348TSaT2Cj4hJdCAeeAYFi5sqQGTQYMGEBSqRTmM0wkPhnkDirN4E0mA2NcKKiqqti2YmgaFFVvk+Bz0zAclsv1uH4Kq7nzs9RK5Pwq2eRmLcaI95VmSm8aazqddoPBYF5uZsUEAqQLdlsMCJU31NtNib9z9dYQL8zy1HgWDod7vKlQKASTREpJXE9Kx17Vb/gPsDCEFJUv3zPoe+nv29G5tSS+5t/1IhAgiLjmBEVDTAgulDbDZOACp7nDopquWIP7ltukD6n59qfWsQ++sXSn+Wvt/VY1ZXdoSeZIJp3G4ieEug7f9SlRqOJQ5iklzLUthW/rWoQENEqCukmCqkYimtVWXR5sCyr20to+sXh1QF2dzDg/7Dimn1MVtNeub081bj+6NlRRHlUVxVYyGdfGumtuyVqmlasienDonB8bSSrH+oSjgV2Wr9lgODQ2bn2rFbFcfVDGdgh0iEwuQ1pSWdLimrEFjen9v14R3/+1OYszux33wVe7bDvwqRmnDX+NELKMEBJbubIjWtEv4CqmqUNoUxrA3OC2bg+Na39ic8jJyR4IBNSWlpZeTEkDc8OGpp3PexsUhIH4W7G5THH+pqYm+I+kPwjXUizLkotMXbBgwaYdxT6GPuGwZubzfKOR85lBGBiGYdXV1fU2H3HNLveDaxqGIS3pkuaGoijcB+C6Lhc8cnPdjE2RN8MwbNM0VelE4zcajUaVdDotnSA2Y4ZLqcntNtu29V5OXFiIphAeVLzATCbT48vL5XJwnEj1CJKYmw5+IdTdoSK/SyaTBrQZxljWL1mFUJFqa7EX4WnNwokD1cpzThKSSlnQWPhkxUAX65vP55UgpSznuo7cveSOVOwZ5T3M+q5D23kU0YZvUdlCCAnf0bDq0Pe+bjxk4erkzo25qJq3bEJs+CNNm+ow/EOUUAaJ5zLHhsxXiErUgBEh5UGbVIastsoy/Yf+Fdr3Y0dUrxk3svbrw3fts4yECc5vl0cNN5OxiM0IeZZsdvtcfogEKUllXYwp5orx/Btrxy1f37btj6syw9ZuTI/f0G5t35HTattMRtozcdKeZOGlrebkOevdya9+3bhmh5Hlz1900KBnttqqYhEhpOLnjdloVUxVdZ2QeDwuza6iCyWZTHItCD4sqN26rjueKm7BvKM9vB9ojZjHLE/4nIIN7lgWtyLVAQMGFFskDP/p168fa29vhwrOpa4Q+ni3mP8O7PlSfdFMk1Asrnw+z01PbBzUNNW8ZQVCodAmHeWz+zZB7gT1bYzc51Jq/uN3eXk5aW1txXXdPMm7JM+FpkstSyFFrulvKhx1lqXmPW2Z4KH5BfHHdDoN137OJ5FoPs+Yqqo5TdNsNpOppI7QWbNmkUmTJkkHHG+pVMpFlMDnhOG7Jw1RGqIhLPJSGgZNp9N5x3GyUPXlro7zwNa2LAtSVu4khQUvPkObyZumaXvOHANaAstms7hn7pH+6KOPtEmTJvn7yskgzQT0c7FqXC5hA0xVXQtOTnldOIPq6uoKL3/NmjVGLBazXcNgcN3g5cObjUgBJp64f//z0nP/ulS767zR+UnbVziJRHrg+fcsPurjH9sPXbnRHtaeyBCCwIKahgdI8QQAhAzUAIcRpqpU01Wo/BVBmtyiJvj94D6RWaOHlH18zSmjfi6PaMvnZBzyini4I7qNLyF1Cu5/zJg+dOHYjayhrs4T3kUM4ukzCB07llD4vxpIA0k3TJXRA0v8fEW9HxLyhMXQOx76aegXy5smr94QntTY4W7TbtHK1qxFWlc6gxY0bvzjrG+bTt1+SOTlP9YNf2KPnfquJCSv/rw2YwQVTWXsIw32QRG7XtU0LY+Ii+M4RNO8vQtzpNsY++dD4Zl1XYdlRAxm2FmDuZgTUI2FQ1f1vR/qe8dM9HVw7VwuhznIv8Mcw3Hjx49Xul23MC/wXSKRsFkmYxuhkBNiITfjZphiGLqby8G3gev63kunI7u9vR27NHckir9DGGF6cZ+Jbw2QWbNmUaw//N64cSM2RSUSieRYNmtHg1HukuTPaRiObdtmT2sPfhw9EskW/BfZbHZoPp+HbRWIxWIItSXFCbpIpAVkgT6iY8RAUzFlKAXqCXcUqarK+xJCMn5JJtQm/ISSyeQQ6f3kL8owLPyG5C0rK1tJFYWHE+SAeaYbLuHGEonEAGHvF3YGLOIcy7nlwXIYhFJNksJECri+hJD+rutiy5WOPDmZnPLy8hU+LzbrFoJFvwphS8KHAuFD5XOXl5cvFy9LaCSeUMfnTCbT37KsMO63PU3V2gotHQqF4svXdux409+bTpn9Q/uvV2eMilw2TYiT5poFoRr1LFGXMddyiG1qRDFILBgmfUNW26Da8OzxoyrfO/L3Y97ccRhZoyvUxa7fObngxFvACJtRUJMKrv4SO/BmNjxk4R8zZsyg9fX4xKMgBQedoRGStxhdupQMvP+l734z5+fErxvj+T0bU6Q6nc0RoofIoJib32Vk7J0zfj/8roOn9H8/mYn3WbnOGVhdppqeszHH7XSMNckRJ1gRXNVDmLuv4zjl4jvpu+LvGIuqsrJyDe4V4VDffOTvK5PJDEBfv40t/UCmaSbuuuuu9TNmzMB3OG9hd8bneDw+QvgUuCkt5xNMC8dxEuFwuFHs6AUBJ+YIQuQjRB9Mbu60l/39IVip+biuN5+gnSGMKkxSfjj+g00IjtNYLNZECIkXm8ciRD5cfMcMZljM4A5eOF2D2Wy2sV+/fhnZFw84SvxDy+fza/v06ZMsNitwro0bNw4xDIPbKjLWiQeCTRiNRldQStMl+oZSqdRIeTwfLU8b4I4XLhAoIm9d+sgHiiaTyUEiPCIHUTr/3LKysmWlYrtNTU19o9EoFjb3h3C/BqVwbGBR50TforHdZDJZyxirFvYr3MQQmgZUV03TzPLy8iWkRGtvbx+atkhlSNWsqqqolUhY1dPuX370Zwvbj1/drkYsK4encChxKXNNBVMScgaBQuIomqETUmnk7MHVkS+3HVn9zF/OG/fuwAq6PJP3S+mZKpSWmTPr3P/jgv8/NYzX1KkNSkPDAkZIfWERhAxC1q9iw6bdP+83i9YlT1zVnBrfkta4Tj24glk7bVU985rTx943dnjlqvb2VKVpWiQUUrngxlhD6FZUVCwrdj08bzqdHuC6bpXYUQtaAuambdvZqqqqYsKEN+Bt8G7hsIQ5IndH27b1YDDYDsxMMXwDWkc2O9zJZELw6ot5DJPEUlU1lFNy7X0iJTEz0B5Gi4iJ9FFwLVUIiQ1lZWUbe8BkbAXzRdd1PvflGKGv4zhNlZWVHSX6BlKp1CixecMcgVVgBwJYe4re3Bxf48fM8N1e3BSDB9w/6N0HA8JARhCgkgFcIhfM+vXrN2tS+gaRvzzsDL118dv7wpbjEw/aRk8dQ6EQf1kybINr6oqi5fJ5OxqN+k2bTZrvhXEnI4On2LLgHDIE+KpYH4qNZUMmGdxyQAw6yqA/3bto6hvfJP+wMmFUmBmONnCoSiH+ufpIPbOAMBpSwsGI0jdoNm03IvpC3d79nzj991vM+TrrkIcvxdmnK4TMQPRAqiQO1PqeA6v//ibmiRgPhDr4qNKsOYNU9qcrKCH3uIw9cE/DsimvzG48cnmje8jPrdno6rkdR3+76tsD99mx/LEHLt72CUKizQuXbwzUllMKDaG3kLFwnhWcff5IgFDzN2lyXosFyE0DuXtjDuu63qPjHC1AiMu3U3FtXBdqLMJ8oWCPIXKsH6jwfM2hifHjGkMvzmgecRN+LX5drCH/xlyqLV26FH6RQhQCmyqeOZ9H/+JeXNxQFyBSD4MBFCGXcPBICqcHN+56AoqsX78eTg4spoJXX7xwpbewIpwlUhr7TAapxvYoTHyTStprFtRShCpzhLjBHvrGYjGaSqW4SgfHEL7L67qlwFFT9LqM0kmzVDJ7sj1jBsvc/tyKY56ctfG0+U1OH8ux4e3wBAHsV9clXPlxbYe5VAsHwmRwZWDxhNGRJx+54lePBVS67uXrxX3XzVSYpwW4hNT/1wVAzw1WD//QqbZObcCkhe/h7aBG3v70m9wNM56cc/p3yxJHru4gfR+ZlZn29aLP9j/rd/1vPvXQLV+Nx+PlyZwTrYoFe/SsS8dhJBLRgUj0R3BKhQd96nTBxBDan8Q60N6uKRzg0iorrJ1AIKCRTM+DgzUj5EAhiiGc/PCB9RgN8AsPHG94YWsKp2BFRUVJYTJq1CiSSCT4upPgPXEeoB/dQYMGdTleOkf8PyVbnosEwgAPxg35HEG0X79+JR8IDjohjeWLKKD1TJP7JEq26upqbioUkaJuLseBOCWbjDr4JCk0Eo5gC/QSo0UsGioaOiOiwJ8VHlavXxep7MGLKaOzJ9sz32/cZcrlC/4+46W2K75dm+9juRAE/Doq40EIb3dnTKXhQFTbqm/5T1P3GHDqosd3G//EldtcRyldZ7ozVeLtXIzweP5/zyT4vzR+3w1TgZiiZPp0JWfPVHfaNvjTO7fu9cdXrtxmxylbBab3D9lN361TRl/88PKHDrv0y3vLw0b//jXlydVNbdjxS85HRVFUXdfxjrg/C3MLTm04lsPhcKl5ISNQvI+Yw/7ws5LNZnt7Jmyi3DktzRUxtygJ99hVzl/un5NALSxuEcno6R3z3bDgB4YA5JGTQri8RxCdBPuJZ3Vw30KL2uS6/Oaw4HDB9vZ2SCy5A3f58S6a5yfM4SYCAZyYe/eFlORCr9iP7yVwNKBPKPhxA5tczz8hpKmBh3F1LxqB+4antVRfH0qOL2QB/+RmDxyppZ5TfHbhtJGS1TTh9zJsRGDEy/SOn/iRBngxvMDH3b5mxsVPb/jsw0WZCYlc1qGawghFOFVMRM/5RfRAubpFudF02C7lMxY9vuuEJy7Z6iH4XxwyE2YEhyxDhSD/zzTKSD38C1MdwLItVqeOHz9k/ft37nFNw0VbT/j1tuoDlFrWi3Nzh407/vNnHn9t2V6jhvRvwRqAsMVSKDEvONpTVdWEpmkpCHkI77a2Nq4WF/kpOJwxIbi3WFE0EeKmmBdYKKXWAH6EtujBg70NERONb2piwyu5foSGjMbnveY4kD5wtEuhVrQvonpiDvMNNWAYLM/RiDgnc1tbW3u6LkcZy3wOWDdwKpJcDsJU+uQKP3yX5iNLqQY4qNiNNpmM3DYD8kpRNKhqtm0zW7FdnemWUEV43LfYdMCDpNPpgFDPkKOgC+GgBgIBvJdNEmN8aqfUJnC8B0kGXFmhgKKakydPLmZyFEKXbW1tUPG44MNugp0FziNVVXEvJePRCOXIFy/sUzw3Hypd17PyfjVK7Gdnb9ht90uX3PJ9U2C3dNwhVLFdaHM86ABsAn8rlsNoUK0JK+74kerDT5676/S+tbTxictxFmgEdbAjOLrv/+WGvAfvE6PuxFnqhAn9VwZUcvqdzy1++tG31s74fo02+fy/rX5w1tzm8Yyx86Sz2Tc/+G+BlxE5IZTb8ERRYCRzS7AnrSqTyQBBCG2xi3ZqWRb8F9J/VrR/LpcDsAo7CwNkzLYsbkJIU7aHvkhCgvDhGxFHNXp5N45wVMt7LoY3sLPZuAoYBTc9MDEVJYD5DQg7tOge1m0hKQyhUz5mJnERLaQ0qK5cuZINGzasMxKzcePGGJyJ8XhcKy8v1zKZDJwkXJ2WTjW5i5umma+urlZk4pANeKHwH/CVJhY6+jqO46FAVNVwHMfEOcvLy514PM7V7fLycnhd+U4s+8nQokgc4i/LcTosxqIF0JNt21y7wIACTRgMBsNCCiLawYWMhHLGYrEcQCY4TtprNTU1uC7fEeQ13ZjrKEml8LJyuZxaXV0Ni9CCDwPXwvmErcZ+amlhURou36I6bFz16NJjn/y04/JVSd0gzAWOQGUOfiFHwCWUuYAUUiMQoVv2Uecds9fAS648qt8HHPtWN1MV/oH/h7SBf6x5mtwMighFWVglh1721Znvztt4fWvCqNh2C/ODTx7c43zDiK797rs1kZEjKy2pZYbDVjYe9zbOiooK+JrcasZY3AMUwUQN+pF7YmHw9xuJRLKI+2MCoC9aR0eHDD0yTdMC/oQsGeGKRCKIyuUABOro6ODXRYvH41yMY25iPvocnvIZuTYbjUZNrBt5z7gm+mB+qaqKDSfUPUVAatPRaD7X0aFKByT3SuJZRIKYKqDahQiN6E5jsRjWHPOvnZqaGtba2qpUV1dj4wvYth1yHMdG3lCX/QihHMZYrd/JKJKEdKhkoVAIcfuiLZPJDBHpz7yvP2VVVdVkOBwuljotX8YITdPKkLEFXwOkPeCp4uU3RyKR9aX6xuPxLYFcg3ofAMhJsXWkb+PWY7HYekppvNRETCaTW5N83mXC2SmdLrZtQwNaXioEK/qPPfaWRbe89r3563iyg1BNdwgF0EZm90AqWg5xXbVvSHX33rbqtmeu2PIqSmmOawQE4KD//wqC7g2mQX09EFmEvfbx8tF/fnz1LYsaIwcOqLbWPHTRoGMnbDNo6cp1HX37VOjUplQpD4d/opSmip0LYLSdd955nIgeFEKZ+BtQtx0dHT8NGzYMUOhN2rp168Ll5eU8zdKXVMRzKrCoKyoqfuwe5pYNyOX29vYRwukovvJsV/Svrq5eWErzQJjbcZwhqqpiIyo4SYVGagWDwZ9KjV1jY2NtNBqtlQl+PlwOtPF4KBRa2dPa03Vd4m08uwUCDL9FOJCnLAvPJ5w03PGRTqcLdnOxn2w2i50+L3wKhXxxaff01BcgJRyKRY3jgUMHchHngbT136P/B/Yoz50gSHJkeCKmM34O9LOQsSiPk7+9H6bMmzdPgzOHwR5DhpFnfwJyatuqmkskEvy4mTOZit/8WhOBqiPkg9cXDNnpTz899fRc+ut4JutQnYdjVf4eYCUgx1bRXC1YrY4dGFly3YmDD37jujEXUTrDnDj9I2369AVs+vQZhfvq6ae3Y+Qz9Xaenvw7v4Qm0ryZS2aqv91rxJKFz+170CG/Cp7Z3GHVHjpjzdsPPP/TAUMHVljrmnO66rrZZtK8id9J/kyaNAlaYcZxHOBiCqo0FhbmKNC6pcZZ0zQeOZAhTKwBsQ6AfjUXLFigbtJXzM3W1lYsQFP2l+sJvzVNw7pQS10XYXAggsVw8EimQN/Cec7RhkWelX8XDof5GgVtAfA13M/muuCdwJze5JpCRhG59rDW8AMoAe0OxiGE9BE3xdV5DATsLThuBMlJ0dbc3Nw/GAyWy34+TD/gx8nq6uqSBCmZTGaw4zgV/qiF1EwopRuj0WiTvEd/Px/Yg+MhpNcWDwkUVkdHR/MFF1wQb2hopoTM3sTXkGXZkUECbyscjAhF5qwgCUJSYuHDqdUlkKRSQm595qdf3/Vu8qGfM9EBhGUd6jiqt5F4wsDjHAi44XBA2XGI/sgn14z8Uykt5ZfS6urq1IaGht7wIP+x5vFBLKSENDhPvzZ/l2ueaH8ykSKjT/2tcWH92bu8vGBpa2BsJL+SDhxYNNDnI7CRfgYZSeDRrkAgABBddw2BijkUTqfTHKznA8BJzIobi8WAUC1KkCLIVYYIE4Gn4vszNHsC0THGKtPp9GCJz5EELVgD0HJ6AtGBXAXrVl5HmNuwKgKqqnb0ptnDZOADQKkHEJcNyU3JZFJ6TnlWGnIJ8O9cLldqN+EDGYlE1Fwuh6QgjkQU/gCOcyiW1NE9tgtTQwgfmdRUQN721FeYFQWwkqsj5Y+Q1qYULS+POpjomkqJZbPq9sb26Lq1rRx/oITDdN1PzRWaphqMcIAV4oPJXC6vGkTHc8fmfPJDNp8z7eaE7f5qdDX9y7upiX95o+PRtSlbo6rpMEQQMN84NMcmisqI6zpuv3LHPXhs/uZ7/zTs/u++XFk5e87SsjLbdTVNUUz4I3JZknEZC7iMKZEwpfkcZYEgw2/cuzco/P/A0lJd4yYQb3mFUhwCvQvTHWMXMAL8c2GiuJ6diWPhTQ4Gg4SfN+f9zRV/32bbcWqwgqwTE1xmq/5iHI8TJ36kHf37cXMYY3vvdtLHf33gHevW1fGv7Ucv3/mZexqaS4KIli5dSvv37y+FAd+gJKkNsAAwR8Wh/g2Kyf5yYYn5q3a3y7tfT8JdamtrsakWEIj+82NN9MCPgobrwP/BryMEitTSN4eLROJ0JJSAa/m9ZTsKvwMSBPk5ugiETCbDeRCgbmFNgzVGqlyC1KNY4w+M5JFueQxyUAppyDJvvHvDhHUcBy9Jwis5IYbo0+POBXop2IXyxSGzbOXK9dmtttoK2X7VZ0+7oe7Hn1b+dtvdj97eNs0+pu0AHUiBDrQdm2dWejcMKAHChLBiGVGJqxIkfxGFEVVjKtFYm6mH2zMWUVjeZfx6OmFwOEPzdi3CsJ4AnQ2a5IMvkycOe8Q4n2ugqoHliylJOWKZMMIjkd476BzGAtSCdf5THEJxH4UwOj/EAy3iv51j7kuXobh1zhLE4xz81XB7hj8yDg1oxC0vq/jp3D/+5S/3/vXCv9vOXloxTeq/1WbPnmzD8QpsRiRADp181pfXvPeVfcd5f/mx7KlrdrwWvhjGijtl/cJALAye/Absf1MToP89pj8Tn6nR/dybLDA5ryX5jQhLS4i/ZIhyO/PjNm3JZJLjdVQVARDcqlEAXwmEcMkmtALJIYL3J/kiehXyfoAUz3YUg+dfyBK1xZ0q8sFKpH0WWnV1NUf24TioLHgIkT4KG0fO5FLn8MIpPmYZkbhRyCorlf4JQSSdoI7j6K6mhbbaaquOU6bdcMJnXy+8uHFj66BEOo2d0Vv0GDcITSlneJCbCoFQGCZvNeIHjmkhJLgkoQq0AIUKXBP+zajOF51cmi1pS2lhTl/CYMIif1HrPIdY7AXIYUFAymuKz37BWVRwlBrGTQapsx+RzykS0x3gNZt2+nld87OHHHVVn4Ynr7mLkYm/KKEAYJNnQhDyxu27Xn3MFV+teesL+4Htj/yoefb9e95PKRy0XZGjo0aNQnIantBmCEeKxQV7Gh59ofWi+QeS+YWJyBOQ3/XGBVLQAgRoTUa1CqAhbFb+RKvufeU6k2YNzFY4y4GdyROihHpIfxbgOSCB5VqFts2FBMiHehlhuWHzzRE3rYG/bvz48VB3Kl3XLZfppXIw4HQzTTPTp0+fjfgCQAk0pBXLvplMpq9t2xEMJKSkcNRx4aBpWjocDjeXuqN0Ot0HfTtnrUdAgrRiRVGSsVisXV5HpF4XxiSbzfbHB4AsNI1VRCIVFXsfdN5FcxetOSCZhoPZdvji5WQFfni1bwF211r45o2cI46IkTyEfPkL0h9CiQtx7OugFLI7BXrXt83755XYxcUtFHZ0vNvOKdnlV2EBywMEJWKXc/g7yPP5+Vq6pj/6BYzLHJdURkLqlPHjTn2h4aaHGPuFCQVpL0+apQY+nWxPu+nbs1/6OHf3gLLsYV88s/cLR1+1IvjYjKFQteWDafFcvL+bdZ0gY24GzLaRiCHmpJvJZJobGxstzCd5fjm/CCEhZNcaBg7Pa74dnjsZKyoqEPHqLlD4ddvb20PhcLg2m83a3VR3hMehCSMTsjvbE38R8Xg8FgwG4QdAFjA0DVV85o6/UCjUOGvWLOab+4W+iLDH4/EKHIcUfDjiAbDCmtB1Pdd97fmeFWuvRpCscjIXxFVHS/XCNM31ItS2ifMPjpp4PA624C5fS3tHZkqWcP4h/XmoPwwDlQgagYiVwlGTL5XtmEgkBkq/hDAhOPtNNyZb/L1it71PfPWLH5fsQbSgRVVN4ynJwAMouCeKjEJG+Hfdtln/ouNgIuz43uKljHPxSUnROTxcIBT5vnA+1ycQfEIHcqnL3tRN+HdqJFJYiRN7JkKX4edvTmg5nsYiDTTfffF/83EoPB8FryJXl/CsjFlZtzpiqPvuttO5M5+56S73FygUvDZdC6j19mX3zj+94b32+46YEjvlyjO2e3nxypY+Q/sHSDptwdOfraioWN6TAzwQCMT8SFm5gGOxWJJSulbMp000046OjpGapvGNT4bqZHIVpbSlmANcogjz+TyYoLm/y++nALu4rutNsVis6KY5d+5cfdSoUcNFZnEBxg+EYyaTMSorK1dTSktlOxqCvVrCnOVk5JnGIku54JyFBOSIP8Eg3GsYyofIKiRa4CI1NTVdeAW6JzeFQiFoCiAVpaAl8iWk9HpJGBLAGHBGGJFXoLu6lefM4J5vX9cU9+AjLrxr7qLlexBFsSgikPADcXPAdZhpqUTFXOG8I50QAG+ZyeHz7G2/1iD/iKiYfwfnFry3IIsNWsE1IY4ReqG3GMU5vTfjHdgpi3BuZD9yv4F3HQKAkzBIuABQfOtd3Is4P+U+KdrZHz/yWVwbiElPd1EDhGhhQgCWA7JVVZXWZMp578vv/npI3eXZ52de/xD5pZkPvNXbeWe69uczxt0/7fYF9ksftz44ZviqFYfsN6S5sTEeKisLWAIvV7KFQiGE8rht7ss2JFDRfcSvfjqzQvMxoBZyKAQUnmsBxZrkHRabmSQFllq4q+se1Rq/aJFFPX78eBfIWZ+tz52iCJFjTfUyYDLvAloL34t9woHAZ+I/WPoQYBYUFnUp8IS4YQnT5EANIf165BIE4AgoLY7hNgwAegvsxr3ZZ2jwRmBS27YdVlU4fF07b+StAAm4DQ0NPGPyhhsfP/Avj7x4tOVQm6qazhV8rrG7TKdM7dOnOlMVi30bMNhGD0zMNQYv7MpXlPfofKPg5oXnIxTPTFz8Q5CIcwIDDztJCeey8x6hi0OXJzJ5i5VzLzlgSBZ/U0HICfS17Ceuwl2VClGFZcOp1AUm3GVgTvIAoVRVbX4MN1F4GhSlrq1Ylu24uA5IlxR8qRBFxaPgN3ykkDWgMac0l3drvlnetKuJiyA2jcP0sNKaNp0Pvvnxwbqjr4y8+Ox1dzruL1Mo2Owj7a4/jn34+Kt/iPzlkdUvTPlVsK5//1hHW8q1Qx6upWQTiNYC1F5A0y0ICJGNWJICQPyNO5AURUHiA78Wsh17owoUyXXwlcn8BxGW5D6Dnnx0MkGPa8UeezVwM95EBkqzh74ybCojclhABWazbizTXCBIbgFIR6WnwRDpzxSmmBIkPC4iLtTjogYQRGaXmSIUKR6MOzN66tvS0gIIMiQ6PLA6IYZKaR5gJCWbZc499yyAdDX2+f05lze1ZQjVDcocoMz5QnODuqrsvsPYR++55Zprd9i+7wo8f2HhFrSA7tbDphZFsdbTcYJfkRdGKXEAL5biY4gSwmjTQzFJQSOWz1vAecBnksCrEJMI7xBxXT0cDi6F87Qz7sCFCb+HnNcX7wlqVeVzz7z561MuuXFX0wU/ADcfhPKrKu3ptPPhnO/uOKTuMuf5mdff/Uv0KRAy2bbZdO3xa7b965QT50QPOfvnGz94asKRVjruuCooYHtuwnPfxROjM8ZS2ayc08VfLIJKZqEIEeZySCww0w31nMLMtQgahGSXfTmHYi/CwLtBb9Ly6/IYeR4bpWGZnOO3x8la4Gr0f8lBSJS6AwYM2EQgyPglsPw9plEK8lQWDBYosQH0c7xgQOnwBmjJE4kEBw1JIxdZWrkc79ujF1R4UKVzxvLmM/d3KJS6yuzZ9fbLDXtss2zV+p2Rh0kp9kSGndCJhELq9luPvvzjt+65YavX7uKnK+6K/6W0krKRb+5f/vhjsH9ZddhxXOSHcBVQ17WAYVAVhBfZLM/gZMX6LlmypJoYRmBNY2NiyoQJwdXr1kWlCQM6Z+lvYAiNUkVpjced9+Z8fddhR16iv/jczbc7znidsbm4Zu+S8j/W6m17x/v1dx/Z5fqDTv9uyHEXfn/hE7dud97CFe2gzivZ/NmvwL1IFdqknJylZ7r0HHMDQeLgNyDUcny5s0uJ9rQxepE0rCKPDcNT5cTO3RtBuH+nFznXap7lHU/b7bl1ExjSxOEkLd3NHJnAwWG7yHaEdgAuOh+0tZBG6ePE5xLLq5pUcFJ0X2wFH7l4aEWQq/CbYswjRIVgEOcroCa7XZd3R+QDUs1xdK7uYEBaW1PcZPnwi3m/29iRATmpI7Q5lziOOm7kFvM/ffvemyzb5TBkca/u/9qPtFPHDBliV1XFsjU1Zdm+fSudfv2qlJqaclZWVobvU74iNZu0QYMG2X3Ky/WthwwBlZMeChmMCOZ4yqs6dUZA+PCqutKRzjofzvnhtqlHX3UeJfN4YhH5hTU27zSbTfxIe+PhHU7PpGnw1oeX1Y0ZVukAct7t0MJ88tRvU2Es54DiQkKMiees5jwMkpjUPx/lJaGq8xRdAW0W/gap9frnrn8NSD7OQuEVP8ZGvKZN+nb7zIFKMpwpsxh91yhcS/ZduXIlX7uSmQwkyvgt6QjXrOHg48K65QMguN89OAulDEgxQTVV+E54JXnuuLwwL71mFlCF3cMp8jNbs8YzhWFyyHoLCIkIqizuJPTFVPn1oM7gd3V1tXSicE0iGOTlvKBu6YTkA2VlEdbU3DHOgpngITa5qRArLyfjtxl7t6oqbl3dTCqw8v/TTTp0MRGkyin+hBcux637jxxbCFX4YCpB/57PO1kCYmrB7cG1BTA5Sa0Br10PKK3JtPPWJ3PuOPyoKy/QNdX2HI2/nBwIHiGZNcmxbJdO2m/on5avdurm/9Q2eOrUTRKQuqXYBwACVRnLFzJ7qY+2X6bG+9eB+DePKMAB6cMZcGVN/PaPeWENIBIm1g+nCsCXQvhgLRSwNt37yrVne2hCHgQQ33PT28fl4L/PQt+hQ4fyzRPwZF+UQbEVJdCtNilvcITwTKja2lqgpYItLS3IfOLwY6RDckKUXA4kj3Yul1sHswKOCMA0+bAGCGtuboZ0jKZSqSocnyZpGmYea002m1VDoaSTzWZXx2prQSrRZTK1t7cjFlrZ0tJiBAIBTmaK7zOZDJea8Xg8X15evqq5uRlZiLw8Vb9+/fg5vvtuhZ5IpGsn7X+qxSEG3LWOsCLVghpzdthh1Fz4DDwC0P/dJidyLBaLNzc3FzD48l3Izyj6Ceo3cAVIVZZ7l2MxJURIvKOjI96UajL69u1rbtzYMUTTAoSY2S4IST9LOPcladAUMs57c7659bBjrgw890T99a47CRPJ/iWNT10dU6cdUt56xf3rbn7ps/yRjKVWrF6dpRUVBqe58vxPnnYQDoc7mpub43IOowG9qCgKra2txfhtISuRCaCPGgwGObYgEok0CkYt6YRXMDfRV9M0FX39oD5foSOkXa9tbm7G2nGxZrDmZMEZlHZEtjHmvbfukg4hUdwL1qVbGYvJfIQC+KW52YtS1tbW6oI4thDjlnUdEol1jmkGUNyWVldX+4Uk7z9o0KBQc3NzRSgU4n/T+vbtW0ghTaVSMUVRKqPRKEAZmGz8xKJ6TbZv3749gYsqcjm3PKK5ZoiEaN7MK9Fo1DYMomcydrpv374c1FSstbe399E0Leqr3KSGUQLHNPEiN4DMqXjPiRpjswZSgqwvjBNCiiKAwFxm6EGhUtX/TwsE2UTOQVEPOlTcXC5Xi8I1SGqBYNY5O47h0kRCIWVl6ysrKwvv+qxzr893CVkW4qed8trbdxhVdENpjSftt2d/dd1hR17e8fyz19/r/sKiDw0NlJsJU6fST2fO3jjk+2W5fbcbGfmooyMPvAF28wDmk9CE2/r27dsF9yLb6tWrQ5WVlWBz5lxqnlshCJAcTFY878ZS6c/gFkGSEZiRRd4ExwqIeY1DNvTt27dU5KI8mUxWwK8B+eW64RwKToXDYckTsrTUs4vkJmh+BQejRP2GQpWJ8vJIySKara2tg4LBYJTzwUHd75YW6ToOT9MEiSNP4UTsHxcJBAL4dxffgt9WyRDialqhvqNLRf0FyyrQmMnwpL+v7M8HUNpHPtMX4UqP9diXqivvoa5uEs6X5XPanwcAZy/ASL2UY/tfbD2kMUveSCbG0c4JNh7MfOxEOG7i9Onc72JwB7IfIt0NDSmCLwh8IhGDGobansk673/1zT1Tj75iGqWzbTJ+fI/VvP7TDWYCfEVTJ/Z5ekC5/mU8bkIwSocc0oFF4JgzYRWbhxQRMaQQe+UIPco9H50emqQ6KzSZwozPoOiXfgUhDLhJgoIxvvmP1iX9Gan6ArDHiwfhojGDAW8DNuic7xoFP0Nn2nuswK0oiw/B3wYINDgii6xZ/9wpvPCYYTiFWKywX2ggwHMKdFtVQ3BC2rYdEDkFfvt+E1slFBZQvkCAl6/SPS9ugemoW58u/X0Vagv8dBhUvMBUKrUJxkHew88/N+J7RC4KIJIC5Af/7g3F7d3cP8Uj8N/66elRBOEoyGKk45XbrPgNViyMWe3CheJdMI0IYp9OmSDmS/c9THARUlVT2pIZ590vv/vr1COvOJnOm2d5PoVfTgPJCn736VMGNZnjNoSw5DY3PKMrV6702+ld5rHjONxB6IXHTVlekKO9pA3ePdKCHAW5fmQEQ/gTCn6eHMsxVMLq5ldgsi/neISHM+fxJ3g8noXUbUceJ+/X3xdRaJkHoSgW96+J987N/iJr1v8dbzgWJetLhUlkaWoegegWsy3aaIYPPtR8/HBPFUIb2PkRIpODUKK7fCEyqYqDLyQzbemrzsN/ApoKail+BnEBDwOgoYBgTw2Rh+KOuP+5nxkNCyRwBaof1xQkmAWCvaOjo4vXXdM8WriCyeCZB11Q1p14Su8XdzSqutKWTDvvzfn2obojLzudcrPhF6UpSOwcMpkS4vk11DyEVhzpLJVWtA0aNMi/kKWzT5aV5//s5cIFTU3svpx4CP+uI6XTHeELQvk4hNWDYnFzpz01UUCpV6HLoW+Ok3NdXeZRiNB8z5EhQUxQ8HV0vRCQVsEgV1vgRZVVabpBLbs3CeCXsMpCMoispBwOe/zUpQBPsiqSOKbA2iQ9t6UfZzz+ozPqbiLYekNhi3uBwEMgdlCRKMnmNqmG/TvMEz9wxt/8Jb3wORWLGE3nr2vnXJ0CxeZ/J91j0UCLdV5EQquFhuC9UK4V+BRqKR441ANCwf3g6+/vO+yIy62Gv1//iIgC/SJ8CnKOtbW1xSsE8aHBDMdAocxeivt0a9KUKEQiSC8NRK2apklHoYetR0iyk3G5aEM0DZwjwvbnwgjm+uYsag+PwsscMNPEs3PGTlkBu0fQYC4cprpic3wF/t01/TkcxgziJoNY2BI8USilXaT5VSBJUupEdZ07By1FgTNHDkzRBYeBQIlulmcO07wbQ7aXqPJbsoLzvHmL8Z1l2zxeVvCLSS2hZJs+naOBz7p79V7bXLT6dl1zRjHbogTkqDzPAVnPnDRd+CM82cZDcvzMQna4XsYjX2wuKrN5x3qQaaczuamw4kTmJH4436XIURDf8QWI/lR8lr4QX8qzcJeIgLpXA9DMWsm9Dr2voaIiclNHR0c/QjRXIHARbkIqeqEmZkNDg7f2OUmKuD+/UVbQGCSKU/yHP4MALyFKrGqkNZF2Pvz6m4ePOupS7e9/v/EB14WmADPil9FQoqy9vR2aAneyUkUxkhZq7XgLpkRjPo1VJjihjyaIVYptaEzu8qLegciPyDMwG6MKsws8UTC4iVDwzWswi0EZNy3LG3s5/5F7UWIj9d6OV6cRhANaMAj+R+/lif5ubxqCbXobWgCBe/gLUe5J5JBztUOka3KJJcqkS4457mXxNwAfhg4dilp76JuF8xEEE3xAPIcA99BiJ165ciXBsd3PkclkYDMlc0rODegBWdAFSVfAi/Priuu4khEHv7/77jtoEBnmMNAce4uqVywio6SespUrU/33va3lpaWpQBUx495kd5RuCwD5C8hBcLsuZr443U5hgfyBwrH4jN+yv8QX+bIjZT9+LM4lhYcUPp3CoLBohaDqmk4thJaTi61oDJ538rS/r3v4riPu/vlnJK3GOBs1OP4cx+HswiAjbmhooHV1dfYFF/+l8xq8zow3NPK0hUQrn79W1mnjBeigwKlANCbcd7/85v5jjrmaPvHENfcLPoBfDB2b4zitwWCwSmSKuHo+r5aXl/uLwLBuvxXbtnmNUqD4JKLQdV2sFcy34OrVq51BgwZtstAA/Udf9AmFQkou55BAQEPpamJ7dUAgmNwFCxaQsWPHogsHDkFAgY0cUSTGDNNx8rzeo6AAwDlt3xrwrx+6dOlSKYwQ0nBYJgvMAO+r6zo4QmB2S6KjQj+5FjOZjKPD1DAME3EV4AtG9uvXz0mlUkiFRBpl0dCiKBrpL9jKTYrKykoHFNCxWAwFW4tS0TQ2NkZSqdTo6upqE9fxeTq546ulpWVJKRbcJUuWlFFKt6yurrbziQRyIdT+/ftzrzpyS8Lh4PwJE4/zMvtETlLBJi7R4DabvdzZocXUqhSatmmAqsR2Cce2YZ4UiExEEpNILfb+JjkSeK0mAjcMiJY4AQtnRRJCRZqcfE0XrKrCOblmIH97bp4unn5YQQUciUS5Sm2DQw9E9iXnYYhYadMka1oyk8qigVuSaXN1sdCU67pD9txzT5ys1XtUDdO/8Gx885cp0yLtk/nuSWBJO/8Na0TXSUsq47z++bz76o68Qnlx5vV/cxzum/pFRHhqamqQ94GfQmtraxvS0dEBhnCJNeDJTcQj6UnU1NSUZClOJBJjqqqqAvlkPm9Ri/toZDFgFgq1VkYiRfkL/cVeBw8ejDXDw4KynqrjOOvKy8uLhhYRhkylUlvX1NQoqVTKP64MVHG2ba+prKwsmu4NYdDR0bGVqEfKZxo0B1wXWIpoNLqMRiKFclWFKkr4jBoG8uZLnJyHB2VM0FfboEf7G2XefEy2HPHot7uHDh1adADxG5qLgGz6nUHc8eOoDvyJkqO+S1QTlkanp6vrrWHW7zZY/7GPYaZdR9EQnXRcQl3boa5tU9fBb4e6ltP5ndX9x6GO+OzYlvjOoq5jURf/5t+Zvu/EZ/E3h382qWPlqGvlqWvnxW8ch+/Fd/xH9MdnE8eZ3m/T6+tYjhYr72MMHlz9eS5vk7o6ziLUZRwxKKqqAqkIm1p3qa1QLUCoGnAZ0WziMpdZecryKcpyScqsPAeAw+qgKujbPOElMkV9+6lKaCDqtqVM8/25P9571DGXnwq8FIhbyS+0AQYPzclXOVo61ahk9+phDWAOm0gqEjyLPDKGSJys5Faib2csJ1BARRbCkqSXviJkyvOBZIhR8jzK+hAl+sJZKXlEuDkDdmd/xqP/YLkweUfkMvQwjnLQeAqlj1jCTyjZW+OzSUhUqSX0mu0lNRLeyQe3DIhEKQcalTT9OtEKSOkqdgsMsepRo+iaix5de+J7S5x7Mzm9ClouzPeCBuAljIqd2Ud04jcbuA/BIZTzrcA8ULn9X1DFuX/Bb2Z4PyBc6XIefx+/icAfw3es1Fqkf5D7KxRXCUSs0VtUPPvQTQfd+vDNDm2YCedE16cW4TdV7BThnXcc1/7k61+7zHW1cCRCYkGDVEYUR1Uc+HvUTNYOJPKamshkiZVJYBgcoukKtgGvpitmgOWyvKVqIUMvr64ihpsibYnEeY7jPoZqQ+SX2+Csg9OtEOIWNrwKn8zmnAAYA1mXVFQxZ+FwuLe+MAFsgxkyekH80YieOoqUackdyoXB5jg54Q7o37+/REtKNmYpzDZphRRMKe3EjRZ9MJkQIXZoKeF4qbOeHshHkCIdNRK5xQcEN91Dk4ANjr6CdBbUbDqIMLx0ZnG/Ps7CnrLARC1GToTLGPtsziISdXN5ZinIqoSbJM/TS6XDJBALyIpsXhO6Sg5SCAd6qa3esbxTgAQpofLvfkic+DMPjBv8HwJMmS8dDYNPFrLFO4FROFZRCGe7GDJqoL3FQLqc3idHbNNziGpCOVVV7ZUrG6OHH/7bpc+8/smT6zYkdtt+zPCGbUaN/PDA3+yxdtiwGtyvOuu1z9053y0eM3/56v0X/rTksJ/XNVVlLIdQPeh5Na0sDapMHT58wMaxW45+a+zWI98eOqBmQVmZDvu7t/z+/3aDb0zWGeWkvphTuq4Dat9zR0EyArvf59RDuQLwEvTk0ZeYG045IPEKYi5zaHMvfWUET95GgbqgpwY/guQrEYStMi/IKkqQ4qdplrjqnprMCpMgDPHd5kwA5iu5JgEXfFBgFpTq1NHRAVJGFFuVuAgMvkWp6aKiCr8FX9a2pA/rjc+Av0xeSJSWrAr1v9VQtAYknsWFufSYI5RVVUUc8Fi+8uSNdxJCrggG9XX5vE3O37TbT6pCXrQddv2p59ww7Z1Pvj5rTXMqBOk0sE9F8rhD9n36+mvPfyAY0L5tMDv9iHUz69Tp0zEJpxdAO+QX1GAaV1RUSKwGJx2BELMsC7t8b2sAOTWqXxXHD6jMertuIEA12/YyHv1szOI+esb5bMppIGNDmxPy5v46EZpEhiYXBkIR6PJuMBjy5vydi2oJPngkr6hEAkQBiELEvLvjDPh50NasWcO93cLvwIWQZGSGZJVQygJJqW9sRO07fo++VG0sZMBEYUrIEZPZu5ufiyc0hRkzfskcCZvTZohaBvUljxDjB1WE0YCiMos5Gzd2qOtb0pGbH/q07Lxjd+3ieJPNcccjBL1KUciFX375w9vHnnntA5qqh5657+oztt1hyx82bNiQyY90DNKHuGS2h0NomCqLvtTjvvBBmc6mk3oqC73+d5q0sYWvjEO8pfkrKh3p4t+cAsDf1X8eTdNkWjovrS6/64Eunc9x00S81jOvuSAxGCNQ/jwBIY8r3KePgEcKDv/5EGd0LavAkL7JHJbfSUekLGQkPuOaKpCZ/nXL4Zw+DLxUTYqOp5RIBTveJIpgmdQQxvBXkfWfB1dMpVIgWOJZlNLkEOQMUlJLc4X4fyOjDPTuUqqJVG3Or6/rLowhwkADVkDQ+NmLe2/d0lT/n22Cv8J1XcPRiWm1ZrTAwFqd9OkzUN1u64GRUaPWb9HUSrfRwoGqdCq3sSMVmnvpERUtIZ2mVqxhke+WzCvbeedtWl597qZ6O2+rY8aMWrtgwdJQaECijSzktD3a1+u+HrQ08fXw2kjVmPXtTT/TYGTJsVueujGoRFL1lAsGOp1Np/8twSDnWDabRThOzmO+SYEp3LZtjgoE3LfUnBCpxPDFcO1WmBCACgcpDXObvshmyuc3iq6K9eaZcJ4tibJzQQFm2mTtCaEgodDA5sDBiD+5pklktfBSgD82f/58d8iQIdLE534IMDfpehScEFwgdIEwgxFZFnRIJpMR0KE7TsAKBrlqr2cyGRdIQzxANBoFIzPXEpqampR+/frxk6BARTgcLkP5NODoxbkh+bC78zLwkUgkIe0V9EPKqDgIoY8o+iJ2KwSTjXoLXkFpNV1WVgZWWLp+/XpezAIMTOvXr1eWLFliTJo0iU2ccvKdXy5ZdbzpwJHlqgAY9a0M2nf/+dI96ur2/kogav6fX/Q9NcZQl3KSsWBBs5LL1eYn76lYr8xqnPL3T5J1Xy9P7t6cMLfM0aCuhsLEZSpBGKR/LLd6+yH6m387eeDj4fLyZYuWNw0a0r8simKmK1asWDt69GikAodu++qKU5e0flWXNFu3s1gmogUUks9Z2DbNsFLWWKb1nTOmZvxjF0+47c2MnSR1M4naMPU/j1WAmwDlArLZbALIQNDz1dTU8PvA3OzXrx+SjABAKvOVcSsQApmmiUzSJMxXpOIL+jFuo8N/ACeZ67pRUcaAz29gebAOYIog7Rpp1pj/4jfSoFWkQcOZlM1mo9L8zmazHJAkfsOUQcYvZyrrdg6EQoOO40RwXQg7+bhI+QYOobq6Gn0Lse/169fztTBgwAB8F25rawsahsEp+aAhFLwoqVQK8YuKQMAxPSe6C4AD1CJIJgRON5Qa7Gw2W+uvIgtiVVG0FYu6vXudRH+Lx+MDQYsNZJMQTsifAHsuiCgghUulbyJ0EvXKLwkPvudB+F+3Af7ljdLJogAlYWz58vLJF6+89ui7W07YmGUxAD09BctyST7NOEyTEX1j0h7xU6s67bOflu93yX7BQ089bEx7U1Mb0/WsO3r06KZF6xaNvHvBnx7b6CyekEklOd4C7EvgHERmC5TirNM6pI2tHLI2/mPd0W/s/MHk/kddetS2p8wlM4lK/vNCAVOsLJPJtNbU1BT1Hm7YsAECrwa5RjI8DsAdLypsGEY0Gl3fU/pzKBSqFmzIHIqKSAZ2ddu2seDW9+/fv+jGxBhDyjV+wG0WAjs4PiOagUWN6lWlHkrQFlQgZwjH+zR5mEDxntYeaAsCgUC1hFlvgjsQVWN4opHPE+qK3bvYg8hMPMR1gXLk/WVOuPA3SLulEIv1p2SK1GekbHpVbL2bQ8opEJNdKNb8ra5uOpeYnRzlndml/79WB4o0YBM0lbD7nl02ZdwNqa/fXs6mNeZZzFbyDtWoSzWdIXGJgpGYOQolFqMKc00rby1Px7a88QP37TVN8YH9+oUSmKCAj9z61Tkvrs3Pm5BPpi2V6a5CNIaC7qCF5rw/lsoUW3ehmVq5jNuU/3HKS6tvmX3z55efr07llF6bS8X/f24yd0XTtOSgQYMKFZH9f2fgfvB4RbuYNIFAgCIdWpitJdOfa2pqMJcBlebpzxJnEA6HZXm1QpSg+3Xj8biHNhR2vlyHWD8CPVw0hVmsH6S5y7KJnJhFrD9kdElHZGHt+W6d+yZ8eIzSbMmwObBDS29mOp0uGlf2pVEijstrQ/oASDz2KYpT+J2FXVIyoYXALkNsWOxi/AWIASlpb4r05yBMi0KtRPHDr/RP5OBtbqrx/1SbPl1paJjq3PXkun1ver3jlQVt9igz325ToHmJhsA7xxd4OpZXBwLARdhsSKam1LTW5Mq2+OODq04iJLS2trZ2bf0H5x6xwVw4lsVVkzBVx9GAbvLePBTMvW74t0JVqui6RpW85mSyLeEvWv5++6XvH32xRgwX5sN/YgiknZzL5RJgCiqS/kvwb5iwPopybFQ54BOQkyMXU3d7vb6+np+rpaWFDyFMbUFzx8PxwBB0u50u/5brQKRrB32heS5UgsGgP+fI/8PPJXAQhYwUnr5d+ppd/HzwmYhjucQsJhD4gFiWhRAVB0IUewj/88i/w0krEYl+iSgSPko22Fmi4EQBpIF89M2kqA7pmqJJGK8nEMTvfyDNBlJeVOD1Cas6SaQhB+Z/Tkhwctn6ejZz5vLBt7/V1LAi5UYosxzKbJUhmZVoFATkXg1L9BDmAxewqsfIzFzFynWwLCHbCjs2vyG5uL+iArit8JpYBeg1zAbAt/Gbe24UxhSbOkYeIoNobtDNxTvs5bnPbrrzi8uOgS8BYcr/1HhUVlZma2tri80Mhv/AN+Dzxstd3vtOEJVu0lEmvXWC7fybmS/xpHQrLwdLvhfO9wH2CjVTeuqL6+RFfdMiEYkeG9Zm1Iv08XXfRSAIAcFfDtQeXdclNFMHpVlPJxbl4vlA+uva4RxIfNqMe5PqjjRZNvfBNEUBtraQw+v72ayrcqcQwna6RnH/EcYYnEpBShrgv2AShsuFxP+Ye6J+YQMP/978bvJvy7OxchR5gV7PbIdWBjVl20qrtTwQgVAoVK0SYSC/rIdQoFx+EFIGP5NKiA16T1lf2Ivke4JBZk0rSPLQHRpSajLlyqC0HiMq0W0SMMJK0mx1v2h5/e75G+aNbJja4P4HxxVzuiQKzPHCcNyE9YOBeMUxAT4r1nyhQh7GFGp+L3wenS0ej/NFKfvJCEgul+NVq3vqC7Mg6ksNkPfhJ2jpad3iYlwr9/HCF1QNVFCWTMwgeETdOZw4l8uVHAspaSA4Ap1ssFzFMs2Uv0RcSdCMdOCgwVTpxvBcFBcR9dKf4fSxOZW4yM4rmA69N06RHzA09577X534hyNn3LHPQdNm77L3cZ/95g9nvHfSWddftXTp0kENDQ2cXpsxVlvgYPkf8Rsoz091zr9x8WHLM5HfOEreIfARKGE6okpvnzYpcsb3922/1S619jWqGsZaRu411woE9FNoDKIKNl+zOWxlfVRF0/ycvbw8XSEjkidLMUYtWhbu33Hl3s/uf/X4V7YZ5O78mBJUFUexmGoH3YS2vvzhedddA4E1nUz/jwiE7qZCsabYdiDombocjcvnZiCASV0UHSjPKRnCfVQBkmwF3n6sh5LpzzjQcRwA8CTRCtfOYS6A8Kan9GeYzCavL9FZPV2G9nvbVHGvCOHzkCalBtRygAH4y2hKpfIhQjpQo54r4IrC0zAhDMDrBgJKhPxwrEifBCzZq9Liupaqqh0Z0wQpK+/nnVcj6XQcklKQs3SFNOM3QpWqqsIbasKmCQaDaZTuBk4BpbsZY8gbcX1hE2XAgAHk/TlzcM604/CMTx/GX1ygZx+CVJNCU4+6+N4Zt9x3QtIMEkNziUGzZFHWJOzbDXt8Mnfp2Ref+eJZl79z9PbtxoaT7/rhzwdN2/bKuQRIvPpfRkZfKVOhvn4Wry+wy6lzr2zPBRg1sEhDbHAsv+Gqo/r/7rR9B8y7+08q2ebkHwLElrkgcj7JhS3UYchmBfDn9hQh/SPM4eCPwpES/+HxknuAGge5oIpmbxnYrjHUl64I0uiJZ7y8n7ucfHGSnXcdl1nuz86Cw95d8MbN9Q31P8xkM9WpdOq/PfIAn9v69et1zCH/I6xfv54OGDBA6ejoSFuOg3A7N3slGlHM6aDfFFizZg3nYdQ0DX/DJtYuQn5clZfrgNcUJSS8bt06Jqslyfnvu6/WcDjMSx7i39hVBVcInI2hlStX4rx00CDw+XRpCE+i0DIPO4bDYfRDuBJ98+vWrQtD8/F3wDlwfaxruAc0zWbY8zVUZZZMR0o2u6aiBLMyijQRQrYESAgqTf/+/ZHCqZSVlfE69pFIZEmp9GcInWQyOQyfpQqDwYrFQFJL2MpIZPk4SouqcVDf29vbR4RCoYxpJhXw6eOa6XQ6MLrfIBoOB+btuuextsL5BWSp9V42ca8aIp5D32Pfk56d+92ig6r6VNl/mLjNcyOGDntjQL9yd+mSVZXfL15++FdfrZj0yXb3vdA39RNhMY0My+wImu25dWMJ5Qx5v7jGKKmDCURhD7orVqS2W52mYwlLEWYZtDJqKAeP0849Ye8B81pa0gOPvXXFs5+scPZ0nRSiDCrD65Galne+gumgQbcwBmzMWsrGk/++R54AJi8DvsK9wuk4PEuWKkxj7dbamtM/2/WDZ77+26lH7Xzme7cf+MIfj39+zwkt0QVbayxo5miHMafxzSMC9eHvptZPBdpZmU6mC+Tlv3hkBCIvmUyWh0KhgaCrB1xZw2rwOECUbDYbr6ysLJrCLKs/C8IVjratqKiAMx3YBT2bzWJBryrVF5XWy8rKUI2Z304sFuMqmEihbqysrCyadg2hnkgkRlRVVXEQoaAQcF3dNe20HaiqqlobDod5xZXubcmSJYH+/fuP8m2A3B+SSqW08vJyTazbQgiWqzaowJTPeyGXUjRn4sYKOz/UGPRDYEAgpwq1ALr3gxQFMEMwODsiEmGCqo2fRxCeFHt5QtqahORQ7YmBpNp0XeaoqpbJIEKJJLxomW1DqdEKjjHMab0Ey2rd1KnKCwp1Tjyr/oLv5i86qLY62nTbdeeddPIJf3grkRBjQwnpW9b3vuNu2OfLNTXvjdf0cmV4bo/rL9v/tpfhU+iE5xYckv+RBo92KYnnaQXUpQ3EcTOZQTe91PG7kx5ef2SbrahUU23mUm1QoHXeHedNePWO85i+53nz7pnTEtvTdDtsAjWXZ1VyGqZOB6NoHp2aS/lGZ2kQAAhSdPXXFCgZvf/yYI9N3I360iHvND76wldLP5lAKV1wxyczbvoktfYxM51VnLztLk5+c8Jf5l6SGNv/Vy9OGfj7n+pJ/b8VvBSLxcDLwVPyVZ4O4zF8a5pWsNVLrQPs1rYNUnJvYQUC0CA0XvtB+guK9cWijsfjsswAz2fgOAVmIAUAQDzbnzbQ7RzSdJFOR+6jCGjMRXihp3tGnlAymZTrlte0RCKepjl5LyWjayiU2+25HJeO1CovL13kUjSRuizCiYWUaAbUVE9bswBM8FRpCAOv6hNPb3B7Sm7id8qvBxZZQix89tCM5uDBJJHL2YPGHHDzvnbuS0LDuleYgQMT4dLe9Fxi0HAP5WN2mXqe5bjOEUf+4cK6Q/Z/ix8wkWgTJ03UPqv/MnfGk3U3zsu9vL1CdW1xwwCyZkWVoeouaMi6nPPfsZv9w61upgqtAE7RE29cftnY85eetj5r9IlbCPyBt8ohuqKQLaoiL4GA9LI7F03+KV5+kGlnbKoYGi91Lb2BaP7P3hcEbiuvnKlNbGa7TDC5ewGGzr3Ai0xwtD/OoSi5gLVRXxl7btE90xljh2fbyIdffPRqU8pa3U9zAu6G7M9932l65LpPG1++8sIPD33+uJFXXjNu8Lhl/2rwkm9uc1SezEOQWiulFIWLChWbSpxDLkoVJRSEn5Ev2F5U0wIBq8wjwDnyLO+iirmgHuhkD+96fck9UjDPvdAizsfcjo6OkhcFu1jfvn0lJyS0efgIHdPEm/TIX+U1CgIBtjoShmpK1LfvdmOSUZZLNIEz6K1foZKNDKtIyGcpYk7fgDi6q1umaaGUABedqBnRknbM6v6DWs69u/3CNalIBRDhjKDMtMN5xUulP4ukFfbCCx/turEj3Xfw4P5Lbp5+9kxl8jcn1JRVb3Hh9jffQKmau+6D80790X3zEqbaZMjGyV9/+G589IChLSeZplsP5JeUxrqukRdfnD2utbVFMQwukPiWQ2nAtYhF4O/EdVVVY8QyvUgoqs4FlM7viUncvMYUhOcUSQJqEdXVeFzcUWx4e2msqpodc/h+3+a99OfOFVjHVKWBOg1vrdtxl3MWPbQwEdohzSuUpeFb4f5Z5hpKWYiRPbYpX/kOJWTeWuuApK0woqIwLk7iq8gn6drkVfj3DrFdytJ5b/kf//fdPPJWyQJVwIeK18altpAWNtXMrMsag6unALkerqZrznz998vizrp+UEo1ojOSspwOd20oZTUee83cRb+7bvbFl02feMf99kzzX45onDVrlj1u3DgLKdAisahQbj0cDvdYAV0USeXaAWjS5ZwEMhe0Zb1dm7OYo4C01LgtC1TLyD7t0amKnT2dtqhhFGw4HpoU916yryhOY8Enx6u7ArPTWcdjk37cayow2gqw3b09jy/OWUiYACoqnU6X3CWdQYMYSyScAAnQPMtzr6ikpu6tlDwaGAWUTlCbywzm5tc7FulPar/+YfWxqVQCefoU9i8PP/KEkOKiZtYsL1D2+VcLylQtzDa2paAZaOvz3z66krSTaz5KDr7znWvffa/tb3drfbOkb2bcO3cf/8wZb95x8McbGhv7PvHES3BkZqZOncrRab879MI7ptU/cJZl5QHHUcXq4oxNHpeJdz987Hmpdo/6TGC5uhZK8W8xfiIUz7nnqMyl+/32jIdffeHOcyidYfOKVEjhrqfObY8s2ffKho4XF6XVKKFpiyqqhrojXneQIhmKw7J2dYyuKosoZF1Losa0YpSoAi4iFj4HfVLwqygoWu67vkNcTgXH54yiaF6RWH/peRmz5M8KgcB5Zz1tQVU0mlHioad+eBgh3aaLPjhivWaphFpwdKJCGNE0wNFSrpNQ1lXOUZ+/78zXD6q553cvXefW2Spp+NcJhUmTJsH5llEUK2SavEamrI+o9kZ5jhAgDoXa7au1KIkze2yoyoRZyVjOZqZXaZcG+M7A8yBKC6JZlOV2dZBkyVHmSPDzrgdbGhiGktcW2cScVEVo5wqMjFyOZ2sqyET2H6/5F3ZNDSDchXhqd0AFz7YScGSpHXjSNJ9HwlHhYbqnQYMgLUEIzbpZq5g/AmqNr08XOwrHBqGVMGbDUMMgJpNU2WqrGmfmGz/vvbw5uyWpjTHWFFBgTHilJmzOhmxZpcO3mqY70HnLIgYyS/LDnO3vnr/xg3O+aH391JATO9WKpsmA5PYL7x/34eHcj8Fslbo5dcCAwRFwEgr24lhLe/LglStWU0WzVL6AeHhOcBQKAqbOYRQLrvvN8Lkod2ifEOBk0g7fecF97lomGdy3Yk8vflJv1c1kasNU6sx8c+0Olz+38aWlHWaEBnSHEU33dmssWC58KHFMZoQ0LRw0BkDQ9CkPGjTlQQr5FeXGLrlm5H9lXhiYoZhKyyKKlcuEybEv7FiIJvDF79tsJCak87wQL4A0a1TTIzxEd8l7RykKxU4LLQMgR840jX1a00nAdXIpd3X4yz9f+t4J5g37PnrL1dMdpf5fFNXB9VOplIiy8c3QAwKZpqMHAoivF9txC8lB0BJAJkxNalnMo0j2IwW79SG+fh5LsRFgSPPhrE1A+HqCZRMotWRUImQSNWlCdQ03H8wRYqlqyPScNNwEb2lpKbhy5T34z5VKpQAdkOTFoOhHFFXN5zlewrcHcZZN3iSqsJiaVFDdfXwGnTnd8CGIwfDFVLtsGwhRVvfvryqdiCvOX4e/Aa8watQoP810FzuqsZHRyjBhru7kDRritFVp0430IUR9/cPG37cFooRU9XXJN66KOAgo0qHA4BahznV/q7Nn1YMojUzZd+dVjza8y4hK9wbQ5tpfPznt7IaDjBWRT09LOWtIX3fMunNG3HkoHUHjDz777tZtcae6b2110z77jOe8AagoTSntuOLaR+oGDqw9IZtJuI7NcE2CCHLBV+MZ3bLcOuM1axRCdI8Ji2+lDFg/MLkLSncBevC2VtEfiQihcIRN3nPCg5TSNDAGDVNnMMY2xnY6remppWk9QjXLgZbiCRFZ3l3SzubdnKWTTxamQ7bNyNghsRVfNWVp1nHEGPGwLQwURNJ927xfS3FJIu1yOPupL00GWL4bMt63BHBaoQRBq3aITcJ6WccRWx8RP4IdETrplX22t6yMojCDUEV1kBbFkdT8lgF11KmZzTmLrc9vfui7O784eftzPoXz9l/or5HJPwpQfho0VsdhNBKRgKTuxxfwAsL/xZ82oAdsWUcxg7omXed/YR0gjyKbzbp5FNExTQ1eyUAwCGCQjXydXC7XhWq92xpygEdAvVUwbNJOTgOJV+BrqVh/CX+WWcrSdM/nPdPdz+KMPmBL/lnQO+OAaHt7ew1wBHCucHSRhymAs8WKxWLLJW00TgBnIHb3UaNGIWuqKpPJcIADUpcDgQAnY0VaM2Kh4UBgheAX45Rp6IPPa9asUaNRUtne3h6WpggSqfDwXsGYRGrZmjUrQVsN+upIZKySi7ZgUm7x7eL0ZGfEEKJlDIWXZoAwkPTnGAdPC+s2Dbxput/kHb8d0j/y/Xfzl2x/8tnXTnOZdT2l+ukHPrltKFae2290fr/9t91q25/Ky8PkiUdfvCZpa4HtBg1+2tBVeHBU5AbgdNddddLnkUjwcy6KsIt63vXCxfzzAu/GsyAErFe4U7wd1X9MZ6hPquAKVfj5X3zK89c0jKljKp3qHnX51GsWJCrHENpqc60bwlBqGgUmKa41OnlaaSxfn9whmWdPTJ3Q8ejr3y87fVWS1FJiu8x2lfJISIkFFLI27aOhLzyKRhRUFxCOBuTDumo3Yho+43yDjJ1fgZ/EdYwyQx8cHvUK+Emac82Do6F+86qd0UpKbR9GYknFzltEMQ1IQ6+QgMOoYuskpTWRz1bMvJsx9ivgVEpFsv7RFolEWtavXx8X6fSFUoLrCdHa2toGw8nOy5uJUgR4ZoGNaWlqarJ9xMDwHfAP0WgU6dWcfADYHcx/USgW6wiO+3VQ0TH3xRqgq1evxr9Bhx5C9Wf0AYZAyp5QKASWaFJVVbVcqP9YB478jL6JRCLW3t4+VJRM4H1FLVZwidjt7e3LBg0axOerWK8MawnnUFW1MpvNYu1xxyMWcN6f/ox0YrAUyWgGUqAFiQMghyWxm21tbViklYIdFmhDfizGJJ/Pd0QikZJQx46OjoEYT03Tcj4CCtjoOkJB48aN86U/Y+fsk7v2np+2bUq4fZT+VQ5dZKkELNocKuG92x6BitOn88y362968MZlP6/++6tvfnzVLX99xo5FteveOWn+cTkzX4mK06i3ecwJF13//OuzDxuwxaDWM0879faPXr0TRTaZYALi1GXpdKG69D+KtvNP7N768jU2ffp0wZBE3W/mxLesu2vlmTmz3aUa3INYiL7UExR5gRmlhFgsVGbsPDD77m7DQ09/QCg7eApZsfuF37+xehE7kdhptyJIlMN3Clx0xO41s057ZO3TS9PaaGrnOh0MuIBGSSyk2NlckBz55E4MtPUy7cHbjoTvRsQeVQ0IPuawiKnHnNHrjtj6Cgy8SYmyLBqIHZHIdVQ98PmdI1Ykvz59uTrn2CRtNRQr4EL6cWp7PIypOh2RNdvd9vnlBxNC/o68B3/I919dSXv16tWqXllZ5mXic7ZjzEfTtxbWDBs2rOj1N2zYYIAXhA+FYDoCdQDC8mI9rBk9enRRYdbW1hZyHKdc4CGCcAYi4xElGbCmKKVFcQZoiUQCPgi+bgFcEunWHN/gOE568ODBJWkL2traIM1i4hoFbkNp+zuCJw6OE0kGyRd4bx5UwZvgyNiu9C9ojpNXvfTPkk3EdjkttijphgHnYUl53YKPoY4obCYjvzn+3d+3YwcsDzMbXneoyZ3WEFfRXVaiQk99PacJv2b6Oc8dd+LlOze88dUF02968s/jdzt2v1122fr9h596PXH6hTf0/9Xkk343f/6CsRVlEfsPv5ly2tTfjWuCAOiqttZzGjZ88lXJoZsjAPy8j8XUzGLpuxBEdXVjYUGQO99vOmG1HQ4QNQljBQy2QqUosENzV2REZezwHdi1z1y4Vf2HGCoWr77rOfPwO99rPIAAbk+pGosazg2nbfXeTiOC34889dv2pUmTY/Q8+QrfiHfuZNaFU9U8/tm9obaIe/Me2SNpR9TChTFKlRDc9qoSzo1euWvF1IO27Tdiw8K1C6vn5z8bUZkdiPc87/Tdz/9aIdrXd3/054c+y7/8UrOytJ/mBlyKrCmAI4lOTCfDFjR/eTJjrIHO+LcR3VCBeZFlCGWpAQn/xcJUixWikXOztrYWvgm+ZoByBGBJ+MP4uRsaGpRSfeU/pfMPiYIiLArMAuLzvNpYsb7i/gp1FwSx6mY3CANpSnTRqSVpqo+yWSZLyDBjyZbJZKAe4SN/6ILTMBCQA9xTk9eVYGM+GLqug4nJ7Y4hgP+zqbV9P6eyllCHqEzWki3orf6q28Ub8hNwr088et2Fp5x569J3P/7i6h8Wrtjrx8Vr96LGV4TZJmF2hgwbNvTH04458KoLLjge6q4CAdD9XD34XXp85t7qT5a6BswVaE+TL1r421zeZERhiqfei3wOiSFwHDdkBNQDtzavfPzcLa+Dw//yv/101g5nrr9yZTbUvwPFhFSo9QZbm9LViZfMe3nwcZ81frws9ysKZzjSHaH2i3M63vTAvAD9oMrdDjCVxNvl6HFqEzUYobob3Rgzytb3U0a8WjfgnL/ttNNOjYCgn/jGfq8llbUTSIvKnl52y0/XzPrjw1dNvO1OSumXy5evPnDG3MNfa3Z/qtWVEDyqCnOJYudc2mE07/H+6veHkHry87/Yl1B4H0SwCKVSKV64RQgDy6uI7iUbzZs3r+S7BQuYGgpxPVXkB/jB86iYVbKvpGLzJURJ21/eW099+XpFCoCvFiSHBPgqq/f03Exeu4tAELYPWF6gPkiqaJkO2uOJkWxg53IWFQKA+x5M08AE3QyC/i71FhDJCASYTc0Aqt0U1DpBhspufG719uuTgT5km34usx24qEUFIywIL7S3mUvNzeVMeu8d597vuOyNaRfcsc/qdRt2dxSjPzPtn7baasC8v/z57DfBOvNLKVHm21Eq13fYQz3/LP4iEpLEZ/7+FF3duia74JkrdrrnmSuYeuLNiy988At648YsQOsmL+rgeTTh1HDIjy0qzjeUR8nBmuTPSaPMKxhFUEW6AswHvJuMjnJzxXLcQCik7LrFH6ZfvvutT6JoF6bUTeQxsnppx8jjnp9yf6OzYIJiMoc5RE1mN26dDK299azXf9qTMXY8pfTrJz976OhX2299N+e0EiWvEygbxFZcK5gKvr+wYWtC6M9AMv57h5h1Sbv3SIT5QmXjx/MCw12aH/OCcKCv6HHhfJtTu8R3jMQY8ERBQW/WY1fPVs5TV2e2YnEPtuR87DlTEn6HbBYaCffPFAQCniGdTqtBhQewzbzHQilJKCVhgy8U0uVBKPew2nDUOjmo/UBE8VwXcNFlvb2jFK+hGDCeUNX5XcDNs5xbXg6zylPJFy5soJGQShbOXXVQB9Z9n74usWXlZ1kSTYJkNluzhB8SUnUtIeSxgKE+BidRKp0hb7qE3HbdOThmE1Xvv9WkUHz70w3b5h01CmUMjgKPRU4UffHKrrlquEwJBBKvlhu047lZrft/sMi8cWM65yiGRl3ER7viHDjHMv9A8TepcYmScVwoALRQwf/luA5qKXv5IyLECseuoUTJr7c8/C1K6Qq8ty/XfLndKz8/9esLFuz7pzhZ3YdmkDAHrkyeOe1m0gl7bfSHg85785AHGGPHhLXIB6e9/ZsXVtqf1rlp6jCXqAhPEsVRgoY2QSHaG+4/QnTRQ+seYmeMcd5OXddtOLRl1q+3QQWcHOEp/ptiBeh0ytgMaAjePztLFPjTn7uavl0bPy9qoAKAJDMk5aIWm5EMx8tzFe4lmUxyYZLPEzdggP+QlyyQ+nJ3GID/HhCLtEMe+y6vr1JIfy50QEEUj0UWN+QCnokFnsvlZApmkTA678+DEqFQyJD4Ap7/zB8OxI+lbT8pCSm1pHdVpZTHh3Wob/L+oCqnMnZoydL1k82oTmkw4AXaZUFVXmqsNwWraBN+gDo1bzpKIpmhrktUURKNY8fJL6zZtlvB9AglnCyqMyJRcKoSprq2SVa0qCecefvSv13wxIp71ybzjOoqdRkPWYitXRSn9RRAoW6K8xXG1ftROMUJQSGWJEKOgJZBKEhBrFODJvLt5PoPzvn7ha8e1XD66wd8ddtX5835seWNG+L51X2UvOpShUsVUVgb4EzdMNtce0Xyx6lPfnH7AVknQ7es2vUpw6ogtgWgOiEqTGLFJTknVekDqP2fWxHWJIZMRMxhJDoJqK83B/J5TbVUnsK86YKu598BVSrCgJwxSZL9iFR+rItN+vrXVDAYDKAamThHoS4hEpF8ff3+Bv7jCz+ieFFYwqvBviSEWrFr8nMAVCPo2IFV4NWfI8hDADgD6EFN01KqqiIuyskdNS2sWFbeDAQCJkgk8dBubS3rJ25MsNUq6XQa6c9x4UvgA5HOZrmgQAgE5JXwfPJZ512LM3Th39Aq+JZBQtwhg+o5mhZilpWDh9bBdcHSOqq6Ov/x/MQOjR3p0WzQEEYd2Ali3CTPKvnnJIIYJH9yitPQMJX8UptG9BbmplyicLIzEVaB2Q36M6jy4PrOkyZT73/H5/kz8vCvcn2Qo7cKUYgCf4RoBaXBy43rtGTBlei9L2aoYefIp3b1pjtX/DyTAcBx1VFJUmkcvsjaMByhfTtjE+qqDvgaXcIKddJ4V5cS0JAgLzOnJtlnTe8eQQh9bcKAXy+f9fPTKaa4UZjxFGUQFIUEjIgpvJz/5ybS/rkGDBZiwUBMBety0nUzTFWDOYT90NIe+7jSTtolM3KhJmoLIUoNIawdEdZsNomQI/xpKDmKELxwMsIfEYEWAeejn1EMvgcuNTQtwd3Ank9O5lRwPxoiXnLY1q8nbMAAT2rLvpZlpQSegb8t+A4ExwJfe7W1tRK8xE/b1NTE1x8CAbZtp8DwjLWL9OfB4XCYq0dW1FpbRss2lhhAkDeMhupkJJMswRiPv0YiERCpGN3TKLv1xaiOkIUqpb9AsitFo9HlPVSdjlBKRwRSnESyY9bsZXt3uFQnfWocZtkqUYPCt+0Lhv/j8oBnLNbXL6SE8AEGfZq4/syCVO8pE/Q/1WbMIAwRzwPG18wPPLMqTRi8/qBAowrsgC413V0gI7MsT2FLQnQAceQFJwsIRbQuSU3+aSP+I6IXePZokDqWy/2VFhySkvAacpkjmyEcbNW1M5A2GjhbFUZcZEwUUItSgHiX5Z25RyJDMsMNNUC27zdmucbURi1Ah6sKcbhrjypKjuW/AsDpX9FyudxAhPnwORaLZZFSrChKIBgMtodCoZ976Dc6ZbUZrutR/nESE0oVJDQHFKU9HImsLt03vmUkEgmmUm0OIbyoC/e3QejEYrEmSmmp9GeY89tJ/1osRlgq5VWdBpNZNBpdh2I6JfpCU9kymUz6/SLgSkBET1ZtL6x5Tv5gGMJp2F4gGZXMsl0IR1VVzfIQiAdfdlgOKckc/giaZ1asj/jNM8t8zhYpIblW4u/T/QfaC7QWJ8uPi34/f8MOKaYTEo4xgpRnxL4x8QuaAvlnmgqvNadMI7NtTEKFIgrho1DzokcFf8Z/q0nbj8RIfHCV3qhrMU3XwjpRIzzWi2JCHO+MxhGbvBKtxrB6ha7Oi9sUTAGpLXiCwnNNSnOiK/oQ3kdZZ1tVdEqg1Im/dSpo3BepECAjXKYCCuEV0PUo2rmv0qsKSpBHASMEIUo1yEgoGioDAIujqGxSzTRHVYLEUCvUgK5UKrtvccBcXAV8Cf+XJgS7Al+BMAF4aM80TWinjn8N+NmOO9nFkY/EKdZ5uA4aroD0u8X6dn5ncEZxxnQe4vM5L+EHUEv1FRXWOK2bIETlgoRv4pYFQpcu67bIffParVhr4je/f/wb1AT+44Vt49EwCxuImz7dKynJiCSkSpoQRUe2V+cuL3MiuvSRE3j16tVUkEhKBVRCLSWksmS6aWNjI6eyC4c5v2Nw9fq2bdxICIgRheUtT0X2IL9ipxOTtJCH02sDabgzf37zlo8+8/JBS5asHtXakQjXVETSO2w//MurLj0Z6cLtClWJw+wYpRTFav67zXsX9q6jq843dGvboGrpLWlnr7UJukt7LlqeNJGIlPclE4j8CvFPjhWQxa58WIguFhc3Q3zaFo/edGoV8jPPYfBy5zwVwY/o9/YC8Q9PKHh0bB6EQWZMw52gBnTq2mQJ7uuLNV8Etu6z2x1pMvpXCZZqsc3MV1XBoR0HjqzjO+j/NeQIPMD+++8Pn5e/UpgCHlGwhXVfA7LhWKjyIkRnS85Qae8DzETpJJWQ2ZuoMWL9QNjwvAW/LibLx5W6riRs8BWLLXAvAHMQjUZJD32FT47CaSirXXOvLAQacpD8ffhCRn04BJ6qq6t73P2EY4Q7CyHfZZl2SCpUcirVTyZQyNRpOb6+Qell180rffuWx79ZnqlqSWmDSXUlI1RDnEvkmRUYPr2dzwMQl8qsLjTs/DNnztROO2f6lb+rO+5P69udcDBaRVRFJebSRvLhFz+c+uq7n0yffs4rJyb2e3q3M2ZNOeXh72485uTtL/0UxUz/Vck2/3Dja5CyWJC8lcwRzuOAifbWV639Pvk+s9cH85InfNti7mMrFmEYJ94KkG5utfEsK9cr7e4nOBEMq90uKKMQnY8rfZL82AJ7HYSCfLHogwUnoj4+IQHdAXsylwQOVZnDiJNW6ZDa7d/POW9jjsQnDZ5yLc5h8XpnHM4d0FXDBqMSm+GjffwnWl1dHTz6QPNJjAxnIKKbQewrcwcAEcbCAsQ5k2F2R0e7UzU8nIMwYIz1//HH1blttx3SpWISAazYK80G7cIErF/u9ptxXZkB572hTgej359RrPEXLoQBd1bCd6C7upnLe74Of+tCGtne3t7jPQlvvIwK4LsCAkiWdSvW1KEqpUmPjVaywYoqOghRFmU+lBITDsvGxlajvJwEZ3+2eqe2XDBAyqqBdOfecJnuLHoV0oZ7q/6Me3/zjdecqcddct9b731+CiV5svOO49+tral9kSmsycyZI5ev/PnQuXNWTtCH3v3BYONn0pZOkeHarqMJIZ8uHPsfYwnu0urqmNrQQJ0Z9688dtay8GUjalvunbpT5SwUDjp495q1eYc8w9azxYMv/m7fNRnTpaovJEkYCVFGkVIft5FqwHNp5cruSo6CVtjCxRJnjCQzHMzGjn1mPx6alEqGjMXL7MfONE/u4BToSWwjlJCAqWhBRbHzLjHjDlJJtUHONguunHzjS+hwy+wZf7po9knh2tCAj7bts9uGz1e/tfsJ70y+9rzX6p77y2///kc6lmdR/F/CwFwz9RUoITIb0RU53j34ivjCBIIWQnjduhYyYsQW6aqqEOD1g4867vLzxk449szhg2reYoydRCn4NzxoO0oLYGHCVGG5nGYoCkqPcZO5rKysR4EgY4cCLCWrTnNzp7W1teRcRO5RTU0N7yPWOtdumMEs6oGtSLFy8EhCkou8xya0F1l4FaAJ1HBQZL24Ym0QGUSTJCnDd1wad8ucLNk6OjrUvIv0b6IvWrFxuyz8BoGgZOMReTs8tVy+Kyhv3sZV4oE4BVpDg3PK2fWHPzrzw1Mi0bLk5ReedOG5px/6DKU05SP4uO3EO45/N731h3uDRHp4crdrL//dLY9w7cBH2iH9Ct0qBhdtKI/u+1zy+KJl1FFwpZ46T76wZtf6NxMPLUu7xjfNxl0fLFjNahtWrdnv8u+/qQ6Tr3e5es4+G3IhQlUQzgmBqVA3rLjOb8cFLjxp934Lz35k9VM/J51aqsptnj9JQTaIB5Mf+AJ3XAXPjHdRpoPZRQoC7sORN8n9lr7tm2MiPMGAFM6gSWsCQ1cP1IffkTKTI9aqKw5XtZB9/l5/PjmoB9Z/+O1HO3/R+NpNrLyDLGkPTf9y1Qu2qae1rJsgcbL2/CveOOWHG3/76KPOdOv/QnIL3gNG8nlHCYX0Tjq0QEBV1R54Cfj3sMHNtmQyVFNeTkeM2KKtubl51J8uu/eyOd8vPmbtxo4+2WyeVJUF+nojWchzEd1RsdkCG07niCmKDh9CTzcstBeJL5AnQ2Q2WF1dXVI796J33uHCLQAWARuRkJCigA+hy/HcZMAP8MyxWEzWRCi2uLn978dTe6wxHCoJyGdPL4ehsq6qqlyvFxRQnGYaQmHBggVFsQ347TgAy/Db0tY2poc7zCIUNckcW+TQO51qrp9shCIMxmd79xPDhsQE0LefeMLVedslv9l3l4vPO+OwB847g5Dp908PN7zfYC9sWGhe/PxJR6SHf/4rO0OU5W/Vktiaobqi2aS+vmvSb295CP7WmRTV8/HiuE4jnBPDYshnxPY48+snlrVSQzGydiKl0oRL1JUdZDDZqAwOKewPWTByoy43yCHgCeT2gaZE9Xzq7nNGvzmkumLZ+HPmzFxllU9z3CQqOHk7hQc+LMw5EVcUI6oS1+G7p5wzAvUiHsSnZEgZ5PWGFgfPovcoBg2Tw0adffSRW536qUMc0pprvuvjn7/Udxo0YT5bzUInzv3NzbnABqImST5nZw3XdTSa1VxV1VnaTSoLsnPu/mrDFx/vVDv+ZzaDK9L/jOnAMQYsEChwHAqflu6z64ue9+efG+nYscOdsrIyq6Wlo/bsP9Wf8vncH05Ys6G9Ig/LWDXymk4MM5/BpskZKH3vmpsHyOk0BZeBWAOwwktVjpZ9cc9+cQ0zAFXmuMZWqi8yOUGM7BckAiPhWoqid2dw5unP8DSiYzKZLOvo6OgrIZS+ikvAGWSR/ixTl0XKNCdQxUkzmUxVPB5HkRM+yIiDcikUCgHFaFZWVi4R1O1M9hUNn6va29tDMn1TSGv+okJ2KJPJkMU8dtrmDobzGl7AAphGesq7jUVJbefqq1HJyH3j3TkT1jc2jelfo2149sFrGiYfE50cNcqjx+529msuccgN751/0NeJ1x7P5FL6gPgOP779Sm7gwBHNJ5uW+2fAceUOgqTM7xas3MXKZHTTdDhNlfRdcC+PSMHmRHqc10DzdG/boYwpKijWwLMkqT5RQY8QSw8Gy9LbbTdybmGTnorEmKnOWTcuOeuHtsAoSlO26+ggS8R4eNxMTpZlQTlCXZRmU5AOLnZr+FvcnKHG/vJi+7Y5Z+aK3UbGnliWNM/ZkAkQynIgQhCsqIL6TC553ACgCx7DAr6FBpUyHSuv8CimjyCls4JRwZQrfO15IFk+Y9GPF79ZtzGXbnvqh1uHfLLkw+bfb3vwPJCq3r7sumMT5csnkVZmMwc1pPFaVfiMFewL1NLsdLQ5fPdnl/6RHvLRORRU+P9cXCkXi8VWyL5r1qxRRHkBBWsCKcyZTMZ2gkGnljHWblnBlpakNmrU4EQoaKz+8KM5uzz82AvnzJ6z4PerNqYqrHwOppBNVQPhVY3XZmUuFvmWy5atyVdUhLKo4pxKpdZi41y6dClPgfZdGxB9HDRcRh4kjQCwBe3t7UplZeVyoCj9G68oRwAS1Qj6Iu3AlwLNNfhMJpNvibYstZZaPOnJx39AsXYVRans6OgYINYc93gW0pIZYyp4CZDxKNKQMQG4cMjn82Y0Gi2Z/pzNZvuhai2nxvDSLnnmpNAMUHOhZPpzNpuFEIrZtm0h7RnhENu2uQOnOZs2Bw+OZuPNzaNSKTYY+TZUM2TavdcKFQLEJOa/wOW36bXqFi7k9Olffv59P1SQtGz7baRifNzx5IeObpKbP7rkxK3UXRY9s/GymRl9g97f3mXVw6e+deYHd/3hyeamtoEPPfQOnjE7depUrGF1vwPPu+d3R15wLAWEmDkggi0wBHSBSnVxanifQSXnHYRQHPKIuM8VcUIgAK0jj726/unH62/iwq2hDoK2Zvwp885O5imjqAMgORT4CvTcTrwuo9zhvcEROA3CUlZQeevrlrqgPvXFW6eRubtf+u1VqZbKP6cTzS6htgObQEQNCucu2GXMKwodCVI3kyfkqKenCF5WcQ9dGRULj1nY0vDZoYqdt8mP5hfnHvnkjidZWi4a0EPZ296rr/vTvjPeuGTuAS8ubPzwlMbggp1pSgXLi+pSXlLaux2EelzG2tymoz9c8emtk4dNWPnPYEOElpsrhXtxHCcm0vXt1nQa5dSxILOLFi0btN+B02YcecpVx6/b2BaxMe/0kEN1QwFpM7e+FZWbag4ctoS0O6oNyjiA8rB413bPWJSto6MjYNt2SNO0rGVZhuM4mGdYB8gednqq4JxKpcpt244qipIV1AG8vqpASLrD6LCSay+TydSoqhqBXwMFkjaxl0TlWq7Ky9gnFiaQTCUGkPeFNqBpGrfF8L2kWlMU/AqXpKf2nYfjFPARKlU+QFxw4+ZyWS7t3pvbsVU6b4UJGMCBZRWx7QLDchfTV1yiaNQRgCOIBNcmShjQOgx8k5GPNGyMrKz7YuNzD87Jv5baqK80Kszhqw4des6v8XdGnAhI3DKZBFftYHYAF5FIp/dbuS4RABsBsbOFmoi+LCOfC97XfDuqZ+MLTZx5uRkBXQ2ksxnUUqRk/GkKmUetW55cNXltWh9E3JTDXCRASweeF1npguLjK9EbF+/qTHWtjLs2oU897cavn/zrBTu/+ckN29+w2wXfGEsy7MrWPPVo2HHv3HNcgCUVzi/ymXhzueDpLog9JLl0qgu4mMAoePeKCA5C4SnWFKWm4phaW+iLxGvPP/PFA7/aaaedfnj688cven79zR9l3BbKKbA4pT78EIhO8Aiaw4LZircXPwimq0cmzUKY7/+OVpJzs6mpiaBYCnwMKFLcp08fZ/78Zf3rb3v62K++W3bUhtZ4mZlNQNVzIJS9aJuH+OR6EXe3qGC4wgvN6ozBiagT09RnzZpVCLN3vy53LnjhQP53Ub2ZZwwLQF8xavdCX5ElWUifRh9QtAF13Mtzw38oq0V3zXYUrcCV6KsPJ0MexQ7mv8HsKggkJKmisIl45doeJbh0MOLFGgand1ICxGAm6axg++3CtsEmTq+6DoHNK2PcHtM3V2j5pOFwezRODrKJ3dDQMJUP0O67brvyvqffcbWAsg94Kh/Z8ssTz1273xYryz6fkG63KyL2kPghW5xb97sJhyy6/YGZ2yeSVlW/PpUt555blz/vPO6YhE0XP/6c6/ZX1fDxJmfWUXVeRNnHmgbKCY+7HB4PsI56f8D/MMqKqnEUvMVziwE4dBVNU2llTXX88mmn3MrTvSdOp4ZGyY/L26e0Zx3wtMnyKF1xA526eQFd2NWwd2l7ziF/n2s/et6fv9qfUvodXJXX3rf4g0+WJk9d3WqNbc1mK1rd2DCX2nxL96JKMA38dUu5wOegokKWpffRs9y6OCY79QZOoO/5NBSVGqA/QUU4p8NYF3yz6dlHGGN4F5+9/2LDJ2mlZS/iEEdRiMotGCFTIa5sO882snWTGWOPTp36r0GOysXWr18/J5FIYJHooVBIa25uHXb0SZe/9cOqNoVAAVaJSfWACtoPTyDKlHPPfEV6AVLnw5EQwm4woTmnh0mpM2nSpJL3WlFR4ba0tCjENAMiY5gLOdfN2YYR7TWFGTRugQC87Z6fghOxGkbXl1aiQV7B0QntvotAABU0PJ6y1oIv1FiUslnejPcLmG1elFKmXcrworTNStr14uZ5bFc4YRRqmsgkoem0t40uXZmLZHC7eEYMPs/JkcFw32QsXIIjGEuGjg44YMIPg/rGFixYtmabw0+6+qSZj1578ydPd/z2/sAB7ymhlh0Pqp521jF7nPb9tIo/kYYXP7047QS18UMGPxcOBRMgSWloqOdaz+N3X/ETIeQyj4aw28B0u7pUCiAo/H8DOY90wuE8aLZDyIuPT5e7AAhmtQP++M0Olp2n1PAKLncJt4ox4JNSVh/28sE7Q4C8ukqONeeDtTN/Ih+ddv28P91/2Y6Pa5R+HDDIx+k8i2Ju7nbeV699sdbdnlIHI+iFDfge0UlCw1ykOgk2aQFi8ghTPREhcEzyrjyFg4OSxOoG8asXnlTdtOK20lXjH/z01v1O3ePCF7as3vbdlub5e1l2nhGZzMrPD95JMDQSmmeZrbEPNTQgcvcvhZQzpAIjpIh5H4uFM+N32uZBEly6U2tHeqtU3o3EQR4DKSVhm9xO9MwafhdWluiqAuakMkVR0oggIBrX0+KMx+OgZ8M2Lc1yimyjHDOsLMvaQVK6RIJIZpK1IbpnV/YYxfNhgzyd0gfFLbAvC1Mhr2laGqQLOCm45YrAIr0rM4YkBi48ZNVc8SMlmyyPXfSmcrkceOd4nNQ0OfEl7BaUhKe27Xpp14HgLjk4w4WDq+DDLjgWC/fi2ymLNuQpYBe2D9x/j9sioQB5+8OvrzrprOuP3f+U2vYnD/1ilysOumPkSXtOewacAxMmn3DzV98sOnJg3/L4+Wcfd0s2l++y7HnVpokTNcedqNkOUf0/jlunOq787f0dny2763GMTNRcNlFjrPMcOKfMPIcAwotu7cgrfIuUFkjRoragOdIU5jCFOABuoR4Hytx5WY2egzHnNmasipkL6SPjz/3242OumXPaXc+sGi4HNQhxZOd86KPOt+gVdQI5EwDKgl2NSzNvPmIKwaPhMAvIPerYjktcaqNIfAGt0BmjwIHAnbFsPs2WxefvVq7UOEOqRy6D8mqbjKJAObwqnkruzVMIuLydxgopVC36FzYqEIhZTdOStm23PnLvVfd+//ETJzx+92V7/2HiDocduMdOL8YiZRwF52lEQvZ5TlmKcbdtsM+AFjCigyAVc7qhgVfiLqw3/+fy8nIJOJJ1IZycrnMzIERDBei0/N3tPKi9IGkEXF/1alkYptDXj7vwpRdwGDM+4kY7NT4UcwgEwHXCsd1AcwUCATg4gqFQKN9DNRvWkk6DLRdpn/I77jj3g5d66A/vPIpo4prwdhZ45CEOoFRve8hrjisrW/vXABcA3sRlhboCvcEBGoAyVWZcderj839avvM7s+ee9dxrHz4xeqejTpx63DWvx6Kxxb876pKjxux86MHLVqzbsbK8wvrVmCFHHrjPDuu6s/WIzyVUOlnhyV/pqVhFyNlFvsJ38jkQy56B0gqUMxzBfhfqkJcchPCrN74hVadbVztfpzMsuryDbm1TmF2e/lHwZAA5zGzWkUqzbzL67ks2sN0/WLoqf8vrPy+3zHx4VYIM5VkuLi++IM0Nzz/JX6hGVFUHyRm3710OHxdXAXWBodOoVZHNOmlYTEFiWAr6U5c6SKaH39mzJITNbSkKgggLNnx3xN1zbv7mnZ8ajs2bGUIZaMQ8U0M+AaYUDyZ37gr/0rZ+/Xqlujoadl0dxXhAHJJbt24d0ogDU/baaekffrvXnO13PeHkrOVSosBnhvC3SNPiuE84Sgxieux9hhZSqrDwkFE4dWqXQrZdGLM6OjoghLBWOA8C7H8dxyiWBp4Dvw+vmKbtYzjjPj+BKuY1NEqFMgUUG55LcJdyMx9qUay5uZl3SBqWWaaEUK0FmzMP+wF2HAgEUpZl5VtaWuAx5ceihgNOCicM6LQMx8krqtqGrKpIhHPvc20D4RawLjc2Nkb69euH43kNSSx8hCubm5tx83FkliCTEq7OUCgkUI1ZascCeZuxyr2OeW04cXJ8XjHsdnJCecPR+ZT8P14CTw9lGdA41VAkHDy77phLln/46dxLflq6cvLy1S2TFR0xd0pcM0NGjRw+f799dr3gjuvOexeApvp6z1T4jzWvMC23JvrXhnK0DSsjWxCEXlEVTmTmMhJQ+gdz8+b+bffdSTbbb9Ap3y1Ym3CiVMXG4e1fos6tSHpyKLGyTirrkFRaCxBFHeOpwviPALR4xClcRXZsnSUzTNGVgIuiKrwcJExd11OXXccFJFfZtd9+Mw4fdurMRem5gSWNSweuaJ8/Pk6bjkkHmkfl4xaheY0RzYUjwXZdTqeqqEQjCae57+MLZzwO3JsGQKTGCzYUbkMYQsy1XKLpQajW/A2XqtL1z7QBAwZYqVSqw7bz4Ahyq6urM+Xl5elly9aG+/fvn5i838kPvvvlwl/bDOg3hBQ0wpw81wpQb4E7ZbUQyZo89JXTCUm7rpsNBoP5tra2cpzTNE0HVaZbWlqAInRFCjMW7kbkSoTDYaw5DqOmNERjMdfG2quursYtwtdAfWsQ5gZg0M2hUAjszSwSQYKwRxdv2zZC5DEUYUKuUm0typAQhmvW1taCayEXCARahJuAaqlUamh5MJgDKUo+m18T6hMqVcEZqdIj+Qvw+nE1NhwO8xRm4BlKsTJDo6GUDkWxy3A4zJGKAHbADIHdhJiwPwXT395csiRASOUYi7lbCKOEEsvictQzHnyBbkFx7mlxbLOcz+lMjj72QP1tqVTq2WkX3PqbNRvadmZMqzQCgaYtRwz+7Pbrz3gDdRBAniJ4GP+zjRPLMlVTqHPUtQs/M4KB3fK5HCSe9/dCjgGkoE0CGoLiJEBCoY1lAbeNUBrzClVw09vbWiWiCIuZubyKBHcV8aMQCeUlmbo6ErkWpvCwIyrHu1xkSqS4tyQd16UBEmHHj5/27qjyUYvEE8Bx+QZj7Pbz3jz8tOXKjzcmSVxTLYUpYaI5OeywKg9pKtj7XQdFV4XqLNUAPC58QsJjZjEWCVQiRTmP3AauMv3LhptC1e7OUswVlDPPu/HUj+ctOiVv5W2iGhoPALgO7V9d2W6oTnBVU1uIhKII36qZHPeHraipqSmUQ4vH41tqrqsY0agNpmQs+kQigd0fO3lzJBJZV+yeoGHoemJUa2sr+EFsVJrGYk8kEtDasXk29unTZ32JvkBBjoRjE5pAMpnkQCngG1C9Xay9wrqFhgB+aJ6rVlNT06XiS/dBEYVeZeRBssFi1SkAG/XQlyugkgnGg0uDjiqI6IQs+lK0b3Nzs05IytnYkTWJFiBcO+DmDnYnPkO6aI4FqC6vDFvibrqNGXbfaDTaqCr0YULYwxKX8xYj5I4bzuRU64T8hzUDX6urQ3SEkAlbVb330fLWi9dmXM+TUsAAcJ+fgnGxHDZu1aqOPiOHVa7Y/4LPFi7a6A4WxRplTLBApCKECv5BUU7Nkz4ylCuO4VWbcKxDFIqiKt4ZLGZj7QqNw/tO1wySMDvon14+4dlDn9jnhfKy6nxtqP+sG/a7/TMvS1T5y/0f3Ln0jdYn/p7OJ0ID3CGPZ2jHr1oCa7eiDnKhET9RVMjxApUKj54UnKLEseFP0mi5VvsRT02fSZRiRtj/pTGfk1KaiG++98Wkk8+59vZ42nSpZnDbleXSdp/qMv3PF554BkL69Xc89tLq1pSKncrx6p360+XBVwrnMA8ye7QDBqIPiCxI/lJp43dX8fnxmuO4zFPt+XtDghQUa0QoeugrIxZI9y6seVHv0UtE9609CdnkUE4kSfirL/l/5s2bJ8OQPAohKaIlPRQKvXTv46eoEsxJ8rOOmnq4IZSDL3Y9+R2ntt5gkVyWqZ4da1Evj8HzJ3Tm8G8S6P9H5gDPgUdtNzgBYaa6vMDBdDFY/93qzg11HAJEzzy079x+WmYZt+Vd2zPcPYgTwBkoGeU2ZsOxO19eub3lMvV3u5Q/0jeiUuaqogyU51z0hAFzgEhkACPZnDOhqy+Y7/zCISm+50EbMd0cbNNeHlBnCokK8IpKOqz1Q9q0lX9akf3msi8b33zn98/sOu/6j644Elv56VOmvTq5/9Tf7DVk6j6vHvvNCVftc/+FkUAZASQceRWe8uJpA56ZIMvFQRunjBiuYjix1DHbX/g+Lj2zrucswX+m0a7CAAu19sLLb71vfUtLhGqck5Yyl9oV5RX6QftOuG7atMNnnnzyQa8fftDehw6oDNrETFGX8fWR8c9n+ORAySYzhDH3wSmCOpG+6xZdP5zv0KMnlNghOOK5tt3S0sJ66Mv3CxFNkCxNMnjgrl+/vsu65Q4/QZrQ4346fvx4xbZt7tUVyU3+GGdJnIKvoSYN1wZE06hVyLbaZPX6HCdmW8btFw4Y1TAV+CNygSBANwW1tosnmvdHZebNbWLwJL+9+PG48sh/u3GzoQHCt2OP4eV3l2s6BU9HJx9iZ2pz2iLKZ4sSZyMGfkbd2Fe3q3bviIardMg7yoCXNl1mZl0khFSHQ2TbAYHV2/YNNFGUKMQU4AVvUIFKpJMXaJU9tIcMOoANr0DLKNiP4MXgMQhXd0mS2Czu2E4677amVm/5UdNLz5zx4tFXqlQjf9zzolnT97r2gzZzY58nvvzbIVY+zfV+6g9LeuUkUV3NC5DgTxpz1CqF9q8e/vQug7ZfCnPh3/V+GCFUCAO61wFnPPPTmtYtiRq0IZawwwY0qu21y9jHH7zn6hmZTE4Zf9pp+p03nf/qkQfufUSfMh16d59ip4WZIF27UmMG81hXJGuRGeBtvNxUkCzOOBf6ahpAasXb0qVLC059X2Ii0grQF4DiLtflOd34gereCx8C4rOdtMsi1i3yrP0hxtKN59tw9haPSg2BVgFg2uRineOjJNJmX6qqQQmsp7bPN8Dnqh+qKyw+Lxy22WoCdoOJPNTHk0wUQiZqnCnpH1Q1/m2t4XAsC3r7paPu37a/+x1xA5qCBS55DngOmKMyJ+cu3ECmXHDHvENQJenteyZf8Ks+qTtrArrGHNSstZU+IapMHqq8eOLkfrt9/8g+W3//+L4jdhgQ+ZgoQUJd24FvgZtmBT+F/I9XCQoNFgp2cuQ7gWgVrh1PdgCX5SouYpMK1VRDVYJa2M06SXs5/ebaS18990yV1+5ltWe/dNRzX8dfP8nO2igpyU0aLxzlYbs82e85QplLXaY7aoz0bTt74vQrbeIVWf13Neqp0qTuhOkPfvbjyikO1WyKue5Yts7y2h7bj/jk3ZfvPEEUUWVz77/ffujhNyN33vynFw+Ysufh43fc/lrMeJ86znkOebE6xizDQxZyVV6Agkq2WbNmyZAiJ2URfAqowUBR67TXZ/EidoUkK9AOANdgWZZXg647HwIcDKhEQYJByV5UYow4zoCHJyQFFMAbpTgNfI2HH/mNGai34GENhKdfhf+hp+fJZh1i8bIq4tl5+KAr/VcnOEfam+QfaTJ64IYCOieNSCZnu7DbvQbB8F9wKHZphbqQuVc+WHVM8+MbPl/crpYpNA8+dEGb7tWbTjnUfWFe5q4nX1oMrspPVErOv/AvX7732cLM2fGMOvjgPWuvqz/jV69olGZuOYcFl3yxpjpn2THBaoKAReGqhbnAtQb/oHobj4dq9nEpSCtC6oyeP0cx3CDJ5hLu9/an1y9ds+pDQsiykGJ8oWtskmv5ylnK/hzkC6gDvyWXoTSZEta2Kp9w+s5lu7T8m4q1iDZRU9XZ9qHHXf2nD79efJJLiE0p0Zhr29TKadttN2bB+2/cP5XSB+j06V5FLfSC8xnr4cmHb+QzZ+YTNxN/4wvaMKiby7ku+BRMU1Z95nVUSS/NC8N7lPAwt3kujIdyLNkHSVSgaUcTZj6uZQAFDK0DiVX+43l9e0gaLGqfuNhkOSFFecSILWyYO8Kp6NW397jpeoNIcop2xKio6QkWTjJBAxpuzF+BVja/KugwCtpeQpAIiAZOBM9bUjAZvKXQyamIX/B49nBPhUspCnE+/WDeyNsfe23/ZavX79TS3qFtM3BgfOzIwe/fe8f574rEkk1w6P/p5vE7zlQPmjJkwf1PrTjiLx8kXljSjFyMFCMUXOUeKgFkL6s6aPTK59e9ffFffzz3pmnjntJV+oah0zfSORfvzq4/g/U5tf7T83c+/q2D1yXskU35QAWhFnZij/ZIME/5yU845sjH/iUjDF4ph05ONA8g2VnVjmdPAoBgqXbaaK144IfrTr55iwcufvv7t15btujrS5LBjSpxPO4GiWfgp1Jdomic30LRjZiydXjSzTdNvv95Nt399wmDiRM1+vFs+4+X/PWEB599/dZ41rWprqnEpQ51HW3rEYOXvvvWk7+hlDZBUqHGpuwqppsrNhB8dvzvrr29nTvjoVUj9Cgd7D6kYVGszqRJkxARINAuEEISwsDoVuCoJ5xPgdQYZgb4S0o9vlZWVgaeOopdWlXVqkQisQVUCk6VJqra4iETiUQuGCxfFQTJsU+NRjo0Wnl5OdIowQkvOd84KzP4EtLpdDIajS72VX12FyxYQMeOHUtXLV2FyrZ94vF4yDAMG44XeX7bVnAD69JZtioImHYK0gjRXc8J49m3Hqeo52OVvjFPa1CoLHVetPEewYDuHlh3wWVTp113aVvWLVM1nU/Gb5etI98uXnX2+5/NW3DxZXfOuOn6c1/wc0D+txqvOl3H1NOPoW/97enFk+59J/vUojZtlMNdft5ixhxTmMVWtbDw395NPvT5j6sOn/bnj2+7/fK9INzsOXOWDhp/3DtPLNyQm5RDjAkEM3CjSmeADwkqnZZcQ/F84d7gAQMmbPtOkjC/dtZZY7JQSNpRqJ232Nr4iskOs7Gh/HD/0vBKVacjSJLnS4GclYOueKADTkTdZSqNpkeV73nJ3fs9+Td3uqOR+n8R9XK3NnHidG327Hp7Rv39kx9+7rWH4skUoUZYBaiFuVTdeuiQZc/ec9nvqyJ0tRd5Kp65CG0ylUr1A0UaBkxVVdArIY157dChQwt9UIEZVc0xWIlEApXXh0E7xYIHutErvpLnJeNbW1tXIglp1KhRkk6dDzbWVG1tbQR9/RRvcAFgLbW1tZlVVVVLfDeHUH+BAh7ZjvF4fKCIOiCzhpfY5i3T1qYphiEoEzk5CeeJIsREVVngDUo6HtPpNEBMQekokSy2kEZIae6pcnQ2mw0Cj2RZFsKavDCMCEwgE1EN6cFUIGCIrQNQ17wXCuOYeK8OgQd258/rg8b2ZF7VKYb+krP73sfe/tp7H5+vBCJk+3FbvrfVkAEvV5SH1ja3J8d+v2jV0YsWLxv72PPxhpytIP54n2RbIv/N1kAdaApnHr3lHMbYXtuc/MWc+S3WFpQnjML1z9cWag6xZD5vf7kyuG+MxFcrlLztuMz41ZFvPPldPDzRJWmTwssHBK4D4Sk3mG4bjcQ6eBWRveaimK7fp+yFKz1XN+vCr1hIY0BNZ41QO2ANBlbCoMHM0a9Odtqzq6CqdhbgQ4U5LotcptlBOqn/kSddOrn+eVJH1H+XMCB1dershnr7xD9es909T7/S0NzarlDNYAR17y2mjhs5bP0jd172u+122XqxpwH0HIaORCIkn0jozDCcTCbDuUiHDh1aEu3b1tYW1DQtBLoArBevNolmIX0ZmZfDhpVOYWZtLJgP53nUAhE8UXjFBUEQNmX/Gu/e2tvbDV3Xw4CZY6lL3DSV6UlYyJIaHd8hj5sQg0udEueU38tzcWeJlFQCW92j6u5zVPLMMJzDSzCBypVUa2r0VtvK54mieVjdXNrziHMbFdSK0vfn39WQGFGi+jN3FjY4p0y78eivF60/3wiFUhefddR5P85+aL9H/nb5vX+99YJXn7r/6ht+mP3QzgfvN+HORCpFZr718c133/3KSJ72jPyFztdBRa7Bv/eHX7MzXjxzZp07fTpT5n+9Ts/aSq2EJ3vRAW8T4hwpVNcq9Hzm9KPG3wfN/YzLPzj2pw420bU7LATxMAf4++eIQ5nz0MlN2emu9dVu4OY9DPuugd9CN585UQghwvAAiC9ICfUyEJjJchS8ig4Ke3HwqczPEAAzC6zylNpOpmLmzJkqyFH/HY07kxsanMsuu2Pbd96Z+05zPFtNjZCXRWxl1XFD+6984C8XHLDLLlsv9hzNm7UhpFxd53TnYVVFISIs6E1Udbn24A+QiUbAC4BKQEb+JL17qb5xJY5K6dz8l2zDKHAkD+vpJgXrGf+MoEGXLRSqhOu6CG/wuCfetqaBy8CE47HUiWWMFSgq1MHjNyRJWMFK64skFANOSCp3hD8hITmBhWVZjuuSfCpF6ehQeWM2Z7YRVe/PCwqaOUrNHCFGkOee8xRj6QkXnnFMQKsoMglkIx6F2nZ7HX95OmeTA/ff46pbrj3zr6l0TiOkjtl2g0smTsQLygV05fxt9zgl+t3KjpNffH/WRbqmnG7V1/terBfrFdXSvWsLeEBnIo9vpAo+5663VPS1+fry6iy+IqcUfIz19e6wrZeOTGZJgNd4ZK6vxDyHNPPyWP2qQosPmtT/23BQJd+uaD0ykXFQ6kZhXuGiTqdsQaZ2Jk1x+jPu9ZF1GgTaEcVVBCqU+3IK9+txLHLzjaeje0IFfK5M5VTgimYbKCqCxdI3nmwPEnCmwkGBqnTyFjjURAH1PWnNbTSOmnq0Y5PStu8/28aPH6/Pnj3bOuHsa7Z55IV3393QkelLVOQocFJ5bavhQxe+/Ozdfxg5snYp6a4dSnnZLciFb1ChqaysDDibALwGzMti7M2nxaMB0Ky75QD1yLeIMZXrVUTwkBMk13CPDZEP+BDlJq6VKkuNz9AWDINZLOegEmSvu7zPXMj7eBR6fJhuUQiZvsmr0WbdBFXjAYcMIjFdZ4xkxTghNxhmQzDCPYLee+mWmce/23Q8pk+/mjuk3n9/3q82bGgZM7A2vOHFx2Y8Rh+vxw6Lfco7yezZNtRy8CccfvDvbl1654PHrVi1+mDTci4AEav0cq9bx8JTT/7Ttm6+1aiqLDOyOYeDnAxdU5CZBbSTQ6GRu0wxqIIyN4bh1dREQxzYNW1XMZA/A9GGFGgE7TRUR2LUdVlrvN3qN2CQ+cqzN3wDs40j6WbM4P0VKOHMYoQ7kCWIyNPReXqSpimGZs0PatTJ2WzAuEOeGkccSyFQh4WS7gUXOp2A/L++xe/VgHWJg4ikiEDgExYx51IXNFHc0wDCVbmsgacUyVGQETbQUrZKa0L9vgwqIXfmgie2MWlmMEVhJ11ClcX18QhIbo1oRNUMBiOI12f5lxZ+nqjNmzfbOvv8m3Z67vVZr7Ukc32pqjjYkAO6ru08btw3Lzx262/79o02QYuY3dDQ1VwpzefIUqk1qhoaGgooAam+9ljBGUWNU6mU1Aakhi1rnfYYYsWxADzBxJDkKHDiQ7CUYmjyNZgWsAp4pWg4Lvjehp0lnU7zGCd/IhH3BE8brxCTy8FALXlj2WzWVVXHyeW8/gFCQO/KQ43CdKBgGRbkoV0qIKEvroXJK0EXSImOBWLOymTKLCOkrCZm6EtbM1zd5GpoziS0WsgznvLts2fFRC4mHiVrzRdffLeFY7tM15RPCCEJ+BQ4XbOvNcys49vsBdN+v/G2v92zoa0tMbDh7c8GEkIWL1y4EM8U2OuAaQ1fz1/5G8dOeYsKSxq7ncixKJCIid1SMgd5i0YuXBnek0MiszexO4NwUSfhphw56qT6x8Kh4Ik4YvqMGZyItYxoq0ElQxwbxdpwMp8vAAwsOlFJ3kbKwOrV8RBhNMJ1d67v28VnNCczELGFguYFqLIt5QFRYEnqqCbfGWASDOad70G4v3gZCBeQpzyN0Epzh5H7P5Bnj5HZq97f1Y1kiJZWHUahrXbrrzKi6Torj1a1geeybmwdbSiaLfqPNS5QJ01S6ezZ9vU3PHzwXx974bGWeKqMaobDzAwL6Iq2x7Zbvfb+6/ccCxIc4TeyDSXIBSDIYfC+bDcvN1RHpTrP6LAZR97iQeCLC6CqEhac/9qbjDZjyHZkMBO8Q1ABuot2W1ijxfpCkGieqW6hkDMujSgGon/wI/S0bpHtKBqsAx6CKgTxk8kkf92WZYGFmRNOim0HteB6RIWhHDwyoDXNY1MWfPN8DGQGV1fsUOe50FdgG/i1vFR7BfehNOVS6lakpsXQ2FpClT48gR6CM5sv7D6dcfNi/ABdW23t2QzpxolEyqJaiGRMmlZVzOrDNtVkfJXKdE0DoINYaQ8QI6YlEokrDdjEAYgfj0tEQn/5ffqp0jbBOPlshS7FfHwoU0UnTNGIrhJimnk1nclScP0zNoNBtB5UN2zDVa8tb12btGu4mBEOPaGo83KPtqtUGSphplO+kVFlPaHGKGBmUPKxUGyF36fISxDhXF5mTT4PIM9wLRU8RgphnP5epDBzULyQRYpkWRKLxyHMckwrUBkwRkXG3XTCr46az5gbPfKlySdwM5kaCnO8SJEnSD3/JBz0LKVZWw4eMweXHLNgzL8ClejtUITYF1x+9wkPzHzz0caN7UQJhGw3n1Yqy8rU3caPfvTDNx86idJ7AVjTgFF5aNFNAz5c/OozphmvVhQVxplyWMN28KzyuM6BM0dj46SHvbgtJTZxXdVRt6+c+PHVe997e3N8jROk5VDreV2Tbs0TuW1tLO05BF048XGbcCbKSF+JnV6uW8Y0jddpME3Aom3wM3KvPKhNxdor2tra2kDxphrAVqLeLNIq8SDChoeqsUFIKg8awtVBCtgyXlBlPB6HeoMwCZ8aZWVlJJFIYPGmTFPJQdP1YvsKSkwhfoqqtrboi7xw1MKD5FIkMjKZTKYVRcmnUinXlzoN76waNrEUSJDq6kqiB3ZgNIdtg5B0FrsiUjUKsS+PtUZ6tlyfD4FuQqG2yy5jNzz6yqfUCAT2dRy3glKa6M68U1fXoDQ0EPfBJ96vSWbcvuWx8vhvpkzi2aB1pI6DhJ55+dPDli1dfrRl5rlCQsEAJPKFOJiXc2rxKY5yqF38Sj5CPMEwAhub05DAUw+AECZQFtZGWTSavvSC45/qWrKL0VhY69j7tPe/mb9B3RfZRgxEpLI4K3EU17JIR5qONx1WiZJ0U4578b2lHe5ovASkHXt30nlPEASeX88r9+ZxTAiB4BU34vNQ80xWkfPoUwh4hXnKVA0UqVwdYnnb1IKxkDE0sPXT99U9OeP+qU+Rqz46+4JkoHE4SYJuA6TTIirRqeC5LnWUEK2eX7fdsRsIOY7XquhGY/8PNfF+uaSpO+6qCx6e+e6t7cmUSwMhx80n9IG1fclv95ly2TOPXnljNpvjoKN6HlpUWGvHhrqU9vNEM24RNQB/hxceFdlAXGPwZL+HpbAVk3y3fja238aNiYQ2uEJV2traKuA8xPXLy8tdaA9oqAURj8VcLZ1uytCMGyYeBylK0guHvCvXD0oiYu35kMLcb6coSqv0xaELGJbRICw6Ojoq+fbvukzWcBDnApwgDX+OCboDaACGYWyBR8JuXFFRsbpU7UI44VKp1GiUdMcih28IpoQgdkAO9uLy8uLVn5csWQKJNUIQn5gCyKSAWgossS0tLUtGjx5dNCz5yQ8dlaiIM3JgjH7ycwKlgCkBdBtYhGyGkGCZgLWJSV2ULcnvzeON1tXt83X9bU8tXrKqccvT/viX4yOhwJ2el3a6W1c3ljY0TGX4KSuLsGcaXjs/kVWMLYdv8dqA/gG8RYXjAQghR/1hj/WqQm4pUJh30QaKeQ4776OAo/KTrRZgwv6ulGP6L7vwBPFPT2h5Astx9ti6+snPVjn7tWbwbmGueHwQXnWVrNuWoVvMuOvHiYSQV46e1P/OJTN/Pm1NwvH4ThEMFM7AAueI30ErPaRglPEq3gqtABuIgC10eiIISrmg+pGrOdSIaSSghkhZpjaxddW2t9150F1gkLae+OyJ3V5quWtGLpu1FBLQHKgQkjlbZlqiQm0oQAeFhz8LYFjdzDrMnX863IsohYAZ0+NPnfH4+59+c2xH1naIFnSZm9e3HDp446mHH3raxZcd/TIGBS+Fg46QLwHhitK1KFTDFNfOURUCgVe+FVokj8sJTiiqUkdXDVXV1SSlSkK+T6Q/q44aBM9TMpk0oc4LX5tOksk10bKyxmL3jntOJpNbq6oKVnRo73IicaRvJBIB9QAqR23S5s+fb8Risa3A3AzhAisADsdAIAB8kF5WVrbYz4he4LXHLi4vXuxn5cqVXLIJdta8DC+Cah2/kTVVqu+oUaNAksIzJOFbENcusMtCmnXvA6edt3PGcViaOta3Bp9tIrEJlkMc7LfyEWRdx87FpBf1IuCAOs7RsvduY24OqxZpePHNGZdcefdR5WVhGwlNWOyolYznPfqEK27++pvFZ1aWBZ0DfzfxumzOhBrZ5WU5vG4AU2yHZ0uKH/9n/79Z4Ycf7+/DP3fv533fPU0VzRNK05WLTt/m5dGVyo9Ui6oK4o6uTShUcZB2EIclHJ288fWq00OGwk46acLS3+5Se14fQ4P2AOlhEZRcp2BJxH7Pq4T5ksYEBoEC+iSK6vLnptwoLIQVEa20KLNsh+hKOD84vNX3A/WtXt+taspFN0y5Z8e//uHuelHPQv/N9r9Z3c8Z8QnRg7qr2S5eIS/Mh2UIuhRGXUtDVmNl06nbXPUk5FBDHWe5/ifbRA1sRdhlJxwwbebf3/j82I5khw1hGVRdfddtRn/18F9v3fPCSyAMeD6LzKvnDYMOxZah3jXP0OBSF+8DKinoZDyWXy9pG54cfoxmUAUKG69HyRivcsZ0r0oUjwjougkQoOM4aakpFqvgPG/ePFnhDH4G5B7lpK9NhCZLVo4G+E8cz5OhwF0qqlZzB6ScV4Xqz5Awug7KdMMBq0qfPn2K2hu4W5CiiFinIQWCLPGOghGlfAzr1q2joVAINE2oGN1ZW0rULgF0uUhfrh4uWrQRx2THDK/oCH3dyrKe2PLCWhAI3F4XyZeipLm3adFiQQbRGhx2WJ16/11XPrLfgaftNuvzb0+++Z5nnt5z/3MOGzSo/3OBSLjZzuW3nnDAWYfO+Wru3jCLJu8+7vIr/nj0t1iAftis3wfz727FnM3QZhD1uOmeReet/KDxw8Y0ligM24IurzJiuz816wecevm75905fcrdoC/vf/P79LnPNl6zOqdXp5DgyDd5CJGMB8DDFx6ypBsc0WuW64IQhVCeLOep+47tsFAkQo8df+IpZ+161tPynd5IbuX1Dqa9csZRL3774pxDdjhkwb0HP3vYaS8ddsuK8HfHOTnLVUxej4PvsLZlu4YR0Sb033f6uGHDmjgJyj9Rtq1Q9l37xJ7Z8OH4cbsf98DSxvSOJtXyxHUD1RUBsuevtnv4pWduPA95CBMRSZhdpHIzH0bD8AqpdUYaPbyFQGH4ZbVI21YU2oXnFz4BSYvOqc6Q/hwIuMQL+21Sed33HNjZOT+iRNzJsm48JaDT+izWl9+uKCHHWZkVRbFECXtzzZo1bPDgwYU+GuA9bh6DzVFOEiBUlMVW5i/IJCdRS4EnS/RU/VlqAt0BSpBUgly1GDaB30Pfvtg/8rHxY8p+DsHSyENZQ0xbIySVIcTKe7z9MtLgl+s9JXTPnAm2F+W15+8+9eiTrtjwxbdLL5o9d/nB6jfLD1apS0wb9IUuGbPV8Jb9puxx0503nnur68Lx+N8jSinWoCUgPHrFOVt/dPglH930+k+RS+Ipy6SEIVArwEoWTVk59/nPNt4x9tG5zmkn7vRk/aX73vvNt41vPfDiwmOWNOb2aU3naweUB9b/uCyz/Zq0XcWzirxiIwWgE6iVpMlgOjaPRiq8lg8a5QTJqq6DYBRqndrGMgNf/ur5EV+umDX+kGcOOLZNb95m5bcLFo2pGHPwmOFjFuvEOP68d45p+9aefb6Vt11gc0CBokRUfaC+1aNXTbzhoR9nLlcbpv4zyNDpvOwg8lQuvvLeE/941e13rWrNRBg1cpS4wZFDBycOPmCPK/5624V3U3qT5D6wi50JW4xONVQK8QBUnJ7G+4vI1BYgWh/1N/e/iHSzrjF/TkLMDOZQk6o8fN77psJTAbpViGYyAAB/QKmOSEmA30DWcOVvqpOoiNMY+hvcc3B8bRZxgNzRu9+85JAvtaihnyPFU9g+UmUCBpk7LHu6JqWVeOCyPXeoSsZC4WaScft6iQoqJaZNKByygVChBHzpAMsmJ/YgOFB5NOWKd96ZO/O5V2YdvmrVip1s2xoYCoaWbTl69Kxbrz/jBYXS1ewXkfFYvHHTYeJH2sxbJl866dT3Il83VZzT0dbkUtWrvMJ5N4lF1qfy7tVPzr/r42+b+zouQ80HlDO7NqCRa3MWM6qiinnSFR9e9te3ll5vgeMeaQT+eg8ynCrxE0SYbiKzUVV0msynyVPfP/P0h0tmN7d1tAYsltmCxGzCAiZR8tRu1dJbn/PRiR9e/ebFV8349Y0vPzPv0ZVfr/0C6daOoziKEtH0wcqYF5849OXT+abUtRrM5jYuuLl5cvC06//2xIsXtnXEsSDsSDQWHD92y6/+fPW00/aasNUPQt3uNV7PkA9esKI6pQAPz8qS976kTy/c2vUUGCURFrRJjlDLsTSiE9R8A3Cnp3XANR3BWMaxCQLeD4gyFxY93bvYyAuMZzKMKVME/MdK24Q7KKqrq3tKf5aNG5lCwkiniAIC1SI3wn9DIBiGAXilDX+B/LNlwVVduH6J1kHiaRYqLy/Pb1Gjrf2p3e1LAJrlu59CWDxDKIiicSqwAKOAiLdf9WAydGnIrFWnTNnxe0rJ967L7fVYRXk0/voLafIXTqH2yxUGhfbx3rbFpiuzHtx32pEXzo5/tCB4xRo4kEGk4kWNVCVg0A0WYy/+mLpyh8NfPvK4iz94eJ+dyr9ydPUbSmckGGuO7XHcR/s4qJgE+Dv3JQioEAeg+nMZtEJIwCuk4mGWVVsltm7Ffk4uiXEnPKinsoqt2ZrCmVNd6iaV9gHvrX394S/v/+TSVC4+MOcmXT1A9Wi0igwNjbvzwQOfBPiLh+j+EQIUaSLgRr/69Nvtd5l44l0L17bskc44JlGIMbi2iu47eY9bHrr78ivg3JQmwubQZtgOz7MR/5LQK3/+jMi49e6ECwfP49KpsiLRzxs0M2AEDdswDdukppYyTYCKelvUEgIgsyN5uTasJ+Qh9HL7MimPubprYh8WPCab1MaUBClQ/eFDkM6FAoBCOhtQlFI4FLkzQ6og4sT8XMKp4XdSeIOJ4gqum9M0DYgobrvIGvfFQgB+/vmKigrUC0EEIjGoj7FQQ7iHn9jxmCDbkqhj6sHafCFHNMMoCJ/emuM5MbknG7MqHk+kKQAp3r38woUBGh/remY505Wnbpp45U1HbTtl75HVswZWlSsBI6aB59ZFmFvTaZblne8asyNe+Kb5+ksf+uaDtz5adj1wDedcN3/LH9fm9nY5qSGolHnRly44CvkJbnbYyRzryEN5nnORsyhZKtNZyNXUgKvqBuotafDcwxPvWlRRTJVZruU2O2tGpdV4OBSIKYPLtvnq9yNO3f/Rg547/58RBhDaOD4SDjp/vPDm0488/fJP5i7+eY90OuuWx6LGXjtt/82N1166z5MPTr8YwgDvu5i/oHiDu9qxvb1Z/nRGhLxffmBZJ4YFQgKAKqH2842UZ/WSPGoWckQuiFMRTvTxInpX7ayngDXFHWVy/WmaBkeh5zXuZJPYpA0dOpQLyAIxUc5zakrKgpaWFigFqly7Wnl5OdRGfsOpVKomkUgMwIUBiADNEmKeUDP69OmT+eKLL5aKclR+DxO/kfb29gGapg32pTDzBnZZ0zQTsVhs8axZs2ixclaZTKavrBwN8od4PM7VJzhMNE1LDRpU/Q2uV1MeeScWVI5tT3WI2gQuIekkoVYHcPJeqK0QKnM3l2SVN+ko9L8UhKn+AdKlX0JjEApu3Uz1qKNGfhgNaR/efN/3+3y1qOWQz37aeMSylkQFkgoor0jquGnTNdOWE2xOeqw77cl42LHSLuKJXnULibP0hUrFb1ssfg5G8v1FYjFFNUbvf9AUEL8ArBkRBY1xxGlEqWRbRrd5bqchE144Y7ezX+VZecKBuLnjLrQCLrQZS9YecMj02x9/edZRbc2NRAsFyfAhA9L77LHL9ff85Y8IK2c310To3oBA83YpKR+9EKxknPYO8tWy5CB4qrjMqlodX12RzWbtZcuWrRk/fjy/rq8Sk6yyXpZKpcZgccfjcV79HOtIco7GYjHwDHSJac+fP1/B+RKJREVHR8fw9vZ2G2UPUGdVbNZGNpvNRaNRVBfr0sRaxGa9RSqVGib8gQ5UD45hljFeTuYIRhXTBIkD92wiJRN/mzx5su0/3vfZ48enVGZd8RwW/EC1AfU6jvX17/6joZYeh0hbVghqFPrZlCLlC4UneL9DJgyeUx2yTOIWIIqEOCZhKbMrQlEARLIFqrXeJxevwCRQkgJtzXednmCfv9jWMNVBJmQqa9Ozjh/7/pM3Tjzr7H0Hn1lt6JRZ4GLkAFyorzpocRgltYZK2dR9x7SU63mLWVmQHXhe6wImwSvsIBcqdWHTKoI6zZeYic8FclT8oMaFgsLSgO/x4x1GHBoMKNvV7nbLM0e9dOSZu5/zPIQBsAabG00Qt8G1AkNX3UuvvufwbXY/a847Xy05qi1ukgED+pPfT97txfdeuWfCvbf96UYIg7qZM7kw4E9UIkTOur/vhQVJp3RGYsW4FBJAOzNBOwVoIRnEUKmK1GRj/PjxxeY+n9sA/MnUZbEO+BrsZCknm/TZaaedsDYASsIaNQRxK35z4mPQFwiTfpPrirUI9mVbB+RZmCVd0p8FJVOhtrz4G8pR50VKdMkGggVhAgBXID8XIozyPfbQn98Cl5g0yKGTqod3kMR+dNdd+6yviajLCOG4BW+7AcFHZBghgUpf5WUvqK0WYj49a56iAIurqdSNRUMsFNTdWBRwzwYMHncGkf+xJph8GASdO2amcdG0XZ7bZVj5XaFQFXIKHRkXAJX6qsb4Tq1r1/c5aMqopaMHVjyt0ACv5OKZAhJkwBObfGE0HoIXFZ/l4ofhDE0NIA4ALQXFGhcSHn+Ci+puuqvVKoPn3n3Q3y7PTsxoEAR4m5sdTeDmHf8EqTVgn9+f+9T9z7739/mLm4bUlAXJ3r8a89XF5x2739sv33PokP7VC2SCXcPUqfx99vZDiszTzhLCncz/Hj+VN5Lil9cEBBuRQL6UTORxF/eVybUnwnZeBEKwmkucAfgTiw2DvxQbQH8gVfHo2Zkb8JjMuJnS63h6wBM+ol0mOiCQAl/gMeVRL3Pdx5vV42kl/wE0C8m56Isq8GOKdcxkMty28sVZFU8+EBaNRrniP3HiLFVVaLqmLPihwSuhiwQGZPlpUUIqt/S6qgZ3MHLXxuY5FRWktK5a1THizAtuu2zCAafNHLXjEc9MOfj8Z6+qf+gMkIH66Kv/5xo3hRbWWbY7XXnzod9fNLRM+54qIQTOvForbtZpSdoDL7zz04NgWz86beIVW/YPr2ZUg5OA12/zeBJsnmUqnWTc5FTEwpfVtz0yVCE/vM/8DJ4ux1Mzbd0mMa0qf+TYE87ETjV91nQXgkBkUm+OFqci7RsZQBdedtu0Udsf+sWbH39/NHL6d91+2IITD93zmA9eu3P388886r1sLi/9WWAu6ssYQ1m1/oyxQYyxIeI3fgYzxoaJ31sEdV4szltwY2R2ioPkAIEKB7RcTmYPr+EtTRmW9f6oeMkeQUVRNVGMpcdnxKYra52g0LGPWq1XLdXLePRwRXDWm0Jl2wyiYdSG0BCC59mxXf7isbm6+UTehYkjVQ9IqXDY7S20iOOgrhQYmYW04wzLPd0R4qTJZLKL0DBNPiCoNMOle+3Zk5g7m5A9tqn47KslredszLsKFByOJ7dyhJRvSUjLdwL6utlrV9V11Tnu5OkX7H3QyVes3dBWqQfDJBSKkWWfLiLvzpp3xJsffHLR9GseOK/+6tNe/5+INhRtnIuRs1ZddNPHFzR/sO791jSAdxYf5FTOYh/90FLPcmz2FoNjS4dO+esagsXBg+mSa72rP8FyYGN6KRjS9VcYdZkVJROt+D+5fe3q0aA6pmzsDUdvO3UuNIN6urm4Dq+Mnq4p5JW3Zu8xYcpxM35YsHQKskvHbz1s1ZS9dr73pj9P+xug97dcd56XlDRjhvPGp6srbpq14rV2iw1llJkqc0IKMw3cFUOSBvf7IyuXqK6qwP/u/uaON9974az9TpoxY4ZdN5ZQJLIBkMpHkoe2pSgQUkxaB3K5c02IJ3hhSAydacjehV+g57fUmdviINFIEIdj/vfq/PRoD3muiaypinQE5En0Kmg9cKLO71HCYbntxE2GfJ4ZMSAXubrCIY/olMkoWil7C79FFRgJnijcBGr9+SvQFvvJ5XKaEEZ8MPhnkdct+86s81KmLzlh9Me1EaWVMIFZxuKHlhCqJSTYzzMhRLYedwiUboqmKc6k351zwwtvf31rc1uicpftRn119B8m1f/p5D9ceNh+O90xekj1yu9/WjH8vqfeefnEE/98MCUNPBpB/gebhDnfdeVeHwyNOC/At8dcCGuiUF0nP7e5A7Y7/P4Pt97vjk++WrT+V8RF/UiwOXfmWXByKkHDDm4E5JZ5uEjPfOC/hQkBc4H7GPA9tAVbcU3TVcvN6qb7Drnvdu+dzuwl9s94pAefFQqnIRsy5bfT7jjljBs+Wbh4zZStRg9bdtFZx06f+/ETe9583bk3QxhI6vx6LzzHnvhm45bfxLXd58fzAxeknGE/JKx+33Xkq75tz1R8156t+q4jW/VtPFvxTTwd+y6Rrfkunq6d3548Eg56aFfNCyaK2Hnhnjrvr+vNdkYcxB+h94M5KW1beTgJ5TN1Wz/8t+QTEXLBn3nLM0nkccX6yuNgqvuigIhEyGheT/4SrokAJ8QJUvzQW/AiZl03wmwW0Ci1hLMPyRT/H3n/AWdXVbaNw2vtdvqZnmRmUieNJIBAEAICCUWaoKgEBCkKGhQURFQQgSSACkpRRJoIigqYiCIgRZEQegs1vU4yyWQyfeb0Xb/fde+19tlnMjP4vOX/ve/zbj1k5szZZ7e17nWX674uiW4a0doguwloiSRXkReB5iVUK9DFONK+AwMDSB7GLXRYCWppExJ+pkk9EOFz1BXWftzFL72yvtv+LET0gFn024wNxhoOZCzTSpwE0OEcyTpKXsTLrrr1C7979PmrFFUpfvO806/6ydJFdyuc5JKZrnFmWu5Nnz/r+9c/8cKaRc+/s/rhvy577rDPn3HCB0MpwNE4w/4/3hYuXOj9VzPlftMW41+YN/mXrf/46Is9tqNQezU6FbjnfdieHc9cZzw1jw1tsvLT6zyb9du/j7zpS7pjgnPJXxkhLic6CUU8LQgVRSsEdF2UaESZlZ7zkKHovew6pvAQY/GQjaPfRCS6MaAT51xw3bdnH7jwxwVLVZqbm9464bhD7r7h2q8/ZRha99JrLxIJ4GUYK9LjoNPozxcsBmkqLOaOyXGlhHUltwUt18C4+QZPAds7iiKaMYD2cOw/Zs4YP2QgHIYwjAKyHBBGiJ/9ioOPR8AxXJdIKF1DZzEQjIYVnEMbfX9/fz+pOkE5RTQBSnIUVOn00fYFrFlRFBwDwq1QdAqKox839wqFguY4DpKS6JPQsTLXYsdsNos/WJqmkeCkS0odjNfU1LjIgILSbGBgoC44E0EfLeHIxWKxFyrO8m/ygVC5QdPQIUmys/J9GUaICx9QFGUQx/eqqlwuoJh1dXVo3UTpsgGf7+0t6pMnj7Fu+cNHz7++vfi5/nyBKmg+pNRmrGYW48l3mJdvZ6o2onvGwYuIuHLuEWff2Nfbzc5a+JlbfnbjRb/86fUXkUDL/MUL2MqlSzHh9qSSsYsO/9RXGt74cOPn73n475dGDO2CpRUUaowNkfn+/3L7T3I7wSY6NPnVP/jUqkdeWr2mp8DmcG4h1wfMga8voJLTqFYKycJpLTHHLHmJuO7l8pameapKSUNFGATRCUy7ofdHTBaZ1XVcR6vykqUz5p3+6N3e7cxb4nkcJFVDNsFS5XC+HIRTya9+/cZFRy5Y9KPBfKl2woSmP3z2cyf+/pILT3oJ+Y4br1vEyuAaSgCXv0h0SeugY3JNxXNN0Gf4ovKyvVvQzdNIpKYq5kMvXA4J9kpPkCibQsUsIryR8INy7pzyKD5dBPWcYr7mTMuJa8wZHBys92+r7xKkUikP1AFYBJO+AOwun23Nl8CSIQQ+k8lkxspTSbpJG5w+GXSLKYqeTCRKmUxmlyAWCvZz067L+zi1P2OepdNpB/MJf0frM44xMDAwqCgK6AeAFSVF57ES+ZRMJreM0v6Mz86gnw3D4aUS8gaElsJ7Qv15WJooTL5cLjdryNuSe1FJJBJrR2J07uvrQw/5eCQ70+k4681m2ffOnfnqA8/t7u4vxuqxoPvcARYiVMbq92OsdTtNleHx2ItJq/GJJ147vK29c1bj2HT/H++59v4/3Xudgi5G4NlXLl1Jn1y2cJl6xvIz3NNPOnLJmk1bT123rePYYskCSWSxnDvh7K13Vx9aLBUpgWOZUFMWsSD4AET0p2nDnY1Nfy7/RRPviDHngHAGwqpMkF9qVEQ3EnHFyme7Dzts7rqR+k5G3mjlzZ/61Yee3tiZm2NCspkrlLX3SNpPxsISzyGwY06Jea4FY4ACum27ngkvDAXMoONBOAV+LR7twT6aD2PcS3pqQ3rsB/MnH/6BTxA1/DnDGOzcubPu1tuXf/uzX/juV3p7C/rYsQ0PPPSL7905Z+q47f988g72ra8JolwQ3o5Ei78Y3XGMxSOeyqkt00/y+4AnwQQhKwU0kWEYfHCF4jANCbrylyFx6JRJoQMdEPFnUZcIio3iX5uIlRRQaWXxsWJxcBakBSQgL5vNIkyncl97e3tbc3NzoBQd3gAYymaz+yJPh32zShaoPMprJjjH89iaTqdH2hd5uKk4rcHBQfCdkLcPCgPYjM7OztYw9QBcddmgBLIUJqnOQEYx5LsJS41jRDwPBHEEew7FOyPW7Nva2uBpmCHBCKo8yP2hPzcctZR0OLPZLOUVdJ257d0lVptMbp/RXP3spgH3HNvqQ06bPFdmZ5lXNZUxI8Wc4gBzFIOP4DazV9/4sN7Vkl4krryIU8QIGir+AWOAc7jssrO2//Khv3V29RXG3XLLH9J4tmcsX05Z2eVP/Ov+nbt7L3BskyHasV1oHopJQoDwEHmpJC71UXgiHpWpunK1K4SQFe6pP+oEdynjPb34tfDUcytP5Jy/hJDlP/dSZtONPuKASR+8uqmX9RZshVrIQ6tcWPG5LOEZIEDjwJs5CndByemB31Ji+6X3LK4Vzjs1BGvcUyMRNjbV+I7BNeIYYEtD4bZ49h999FH1v1Zs+8YNN/1lQT6f6dlv3xmX//j6b/wLec+/LvsJjbGFC5f5ns7HUeEv9e9msSRybNTBifvuP2JfQR4WDLMWIGNRQSGVa03JsRx5CJ1rOunciqZtlx9AWePSZ4oO7pz4m5969fUqVXb66Z9Xly1bxgYHu11FidLCB6SgaWLqkfQZ+ERGU3BGCRHcQmEUIyoS1HEMTpLR9hXIxBIvcRXFAoQm/jijxbhi3lLCL8p8RmXIwcNyQ+RyOAVa+SU+Jtv/VYQMCrodR6rrTpgwAa13ZABkBlR+F96bPn36sKq1odyB4JrzvGQS5KAsfvqRY/5ab5guQEq+qyhOCbmEMQcxz8wyVzHkm8Hd6excQ79kswVXiaR4ruR1K2AaY8S1ONLGdVV3HMdW43UpPxbxdd4Qec6KxmIMEF10+mn0AjGoznRNZ0YkwnTd8N/XNKbhffyuaRgU9DM+p8nP0OcM/316z2C6If4u9ldUlUXjiZjjesKN/M/pyQFRxg2cMjG9xoVd5zrcnLLVgmQbTXzpIYhJoBhEdopwF4AWkzs208F2DcCRD1DyS5CCll2ClhzOgA/jnu41pxt32vheCfYJ32DOvT39boMRj3TMO+oT5y9/+Iazf3LDN/8GYxACiDmSmOY/3Rza0yddEa0AokfDFr/jX1QUfcNAIQRnniFq1sgh4MAaWKaHvZ9hQdxQGzTBvnyU/fLZy2kcRyIpMpWiNC8TelKufTQFZ0xi6j2QITrN22iUBr7w8IfdF0IuIhcYsTUbeTp0AoK/IKj+hT/vS7JF8bgNr6cniO+HbXCCF+ETLTBeKvlrAqwUlGnGjRs34kNpb29H/BL0cIfvZ6j7aq8Bgn87OjrcWHV1EdUPXFg0wtWdXYPpc06Z/N6v/7ptU8cAm8m4gwqYn1yEzFvVbMbsTYwV2vc6lwULmLtyJWPzPjmzd9k/X2XRuH4CmKVALzb0s4sXLwYpLH/iiTdregdzY2qSkcGLv3pa5pILKKlHPHevvLLqdJOxcxyNmin86Jky8WH7Qq6CvIt+QEr98uItwgDvbY8gqCpgYsE7aHoD/ZznuduPP3b+YyJk+C/nMI6bkN6W0KPFfstUgTlknqOU1aSlIZBLvvx+JZpKGj6xjWbELYW7UUWxPcsWCvMgYvYIRi7plWBqLNP2ogWPq1x/Dasw7t3ysnBm8Kw/fdQBUBiSKkPKwoXL+OzZCz3GlnjwWvE8Rr6iJWzJkmBSBZvqOP5JyBxnICPlewe+5yBFfxgRu4LqWivrGohHFJoQQXmx/HffKSJ97PL7Q4IZoHENw9BKpZL0kC2BTES5ftTE9BBMD10KJvnHdQvL3akoYHCdmfQdpue6tkFzD50K5Y3KfQKA5NXVycTJ8PGdoF5SgREAZbhEL0qE4khnIySnpXUMQEji2B8PDS5RxpcyrgakrNG7xtjgp2anf7+6c+An2UIxhBIjtChzxy3grhsFqJszDKSlS+malop/zz335DdvuuP3uzZt3znhiit/cX4iHrktly9B8dlD0nH+/MWoe3uxqObc88c/f79vIBv55H6T/h2LGkShJl30I46YC6tTqer5f/BGz3axp9TMZYMHT619sNBqXtab6YACglBeQhxtUUOpn1F3mWs6an1VNatPJ54pFv0cx9ym2S/3drx9Wb/bFwXM0Ifz+jgmBT1lEQDD/Pd1T1fnJKe8dP1J333rBnYFX7ZweQU/s9xg3BCKLT/jDOLO+K9VUZYyQejt+/QihxDFlJHDk4RjRYggvAF/qPv/crLVZSb+4LzgQygW2GCD0K8cBZZ7OMo3WRhSh2pdzAHUYokfKgMbINmO/vNrg/rzQaQULeYP4YPEPXNFgnDYbfr06SgYgHfR8kw/d4F90D9RLJXcpqbKEIOEXsVFqaMCJ9asYZGJUODyr18yJcmfR8t2g4QBbEsyGBWfp/zBUNzC0I3ECmw74um6VLFxkknD2bMnq9/yvbmPPfvuM1esyXm1FHzQheBpaU6/Nk39w+v6hTGDX1Io13rleaqGruUu/MaSO7du3/nTR55ccd3dDzzZduE5xy+XKxc0/pKJCPv2d26+6q7fP/b12lSCfXrBoT978/kHsMKhUiFvBH/xxRf/Py87IvGx9Oij/8dkzZZCpRRaamd958IfPNO2fY9yZD6fBZsPJgt3QOkO7jaFE0mYoem8sSb51m9v/dLPH/rFuXzx4iX8hi9f/di3fv/jr+xIdZzmFJH2lEpOHlM1lSm6QpXuaCLOayKJV+/+4uJ7RXPRiMTY4vnAn5d5WeXZN9jE/t0DiMi8nJuzE0qMW1aJqyrq+gWGcA2UAHqqVvnM4SwruAUDJLGuugo5ONS/IdquZP6AIiVBdy9+Zn7osJcWghpUTkRkJY4QxuT7vrwgSnH8CqfGVfRuuKFQwYnHNQX6yiHAXrk/atjFeIHHeYaQjpyj2YmEcX3G2JD7iVzE0D1bW1uRCoAXQmdO2B7L0ojokXMlOmTugWQVbj8mJ/m5xBM3THKitbVVqfEbjSgxKFpUKScgBVmQV8dEwSo75ELR9ARGZXRCShVaUsCVfxfHrRgf+E9PT4/CwAGHWoxhmF6x6EYiBWfPHjMxdmyy84DJqcc2dnqLLCJXU8F4B/SMappF980N1sW3/37TX79x1rR/L2FLwr3fjvWja5S7l1xz26atbUe9+sGuk767+P4/fvbLN5yy4NCZz02d2uy++e7a9Otvbjjlnj8+e2rJ5Ozoo/ZfcvPSS1aiRBdW7xEP8H+P3uD/ts0fdCWTsbtuPP5Wz/PuEPe7XDkc2ifPufnAbWehLKsuXbqSjOod51/9e8/zlomPjCREQiv9PbREhigDRtkWL2F8ieepZyz96Pdv7+QLLbOIeEZVVPBdwF4gya4AWeazH3u269lMuelxZfvKj3YeM3/f5p3sjOU0fixHcX1koZSzFwlGynbg5fhVB5/YlVFphzvoMuThpKKQqgphNenCyhcp/yLZox2QLipe0c0rS5b4Yw/9PoDaQDlNAvCIRclHBtH8Cd+35cuXA5jlrVmzRp0wYYJY3XXwX2IOkrgKPtjf34956/driIUqNAeJ0BXznL4U/RK6jjygxDpUzD3ttttu24T4TJRBUO9Hic8xDNx1asGkLqxx48aVotEoWqWH3fr6+sbrup4+4YQTLJQ05PsioTFYW1u7Tl6g/BNAFADKfPrT/RNUVYWYHj0piXFA+3M0Gu1JJBKIK+kcl9x0E2JFNmMGKRdp114w45VX1r/11e29qsZ1OK/UZUNh7+4Bm9331NZbPM+bx/kSX/FIWmAkTpcuRZr3C5/90g9//caHWy94euV756149fXz0CpeLLis6MZZbbJq4KhPHfDTZx679WbTJMXi/wuhy8NvEmAlS74qgWnKGfShm088u5KiaRWTE2AkhRdGm+IqALgSDMY/HvaNLk00Zn3nO4WJH+12z27tzDEWNfQgWUfH8qujDE22NCqIv4H1OdFpf31tYF+23/g2Nn+FnzwGZZmP6RO5Ahkq+D+TISBjIHlVqQlhr/PSVJ/uywcZCP1J+l2cF6WOBBc9tXxzup9oa77w2xcmANyLxWI7hPJyeONPLl/OTjzxxJpcLrcvFktUIEqlkn3CCSdgMmswBul0evMwnnQwb03TbIZK9IknnkiTG/TrAhiYr6mpCas/0yZarnmhUJjgOE6VrutFFAsgQuHCIGCiCLeeOqbgeJumRb0JmNSe5wHxMFysQicFsRUgFUXiUcI0yfPAwcW+wwJp+gp9THGoPiu7JJnhurblt4PuJVLhx4pL2a3LXlO/u/CwdUfNqX3qT2+WPu+xPLpwCG5LVk/znPV96gFnX/X6lZq69PozzpgTntDSuSkaunrhrXcse/S5FW8t3L27/fBsrhCra67tG9PY+PwXjz/qwfPPP3aDT83938cYiGt3/7bsX/s9/dJHF2SyA3WOXVS5qlIbPwa3xOzicVq2BZYz0pYARB/qn/6fFUfTI+ho8GXd0JGOviksXk6RJ6JR79CD5vzpW986+1kSrB1S2h1p2zNAVUsTKHtG5XeXwzvwUami8CQrBVQFQTbQZJlCWSUJG1k6MgjhJCL2gSHAv76HADUNBYxwIHpVQKlqhO8WM1RdVzjRU1fUhQNIIJUa/bACC6+ie0z10SOKoXkREKwCbTh0LFPe5IwzPHAZCLceqkvSW3eQcxAaC3uNPbnAAc8gqQpkuzReqqpCHxBe/Uj3HMdFSZM8hmGbmyQDi68V6aspSVm3kb5UyrGB5lkYFKqRorRhI7UkJvlI1QuQRTiMYM8wJhTnFKNRL8a5DYRk+OLD+81vMewCK3Rde+HMn722adWJW7rsKIc0H2U3MEh0Je+4zosb8tfd9+i65y9YOOu1MOxYtrqalsO//c0v/ktV2L9sxwPOIAU7pXCef/wPFRRU/y023ANc+9VX33nAVT///Yrtvblqu5RjHmAmHLBvqmaIVnKqJYkJbgoqNZ/+BBBxRTUY1zB5fJEckquj5CIQw0Xa5+U33jnnW5fdfPadS698ZCjse6RN9FIp8NQFZsDXuwmq0LJkKPAbhCMYqdwiibmIB4IpEOoKDATxvPlGgQvDwVxmiZBBbiqUZ6UxkCpXMs4iBBZCBGJaZhC84LrHdCKYZ1WWq6SqYtHcx19zwCLmZyj8FmY5/0YFoEmwU4jJjAoFo+XnsCGMKRaLUFunKmKFQSAzhIySIGmQOANBzjhqnByyTHaQBQWHvJ+1lLXTYU9O7EcUUbLJg7BxvlUa8Zhz5861X1zVnvvCMVPe+Mwlz/+prav0NcsD86quEfqW7I/Jd2e5csfftv3W87zDhlFo8m/a4sWKsxTSARzc7ngJLkU04PzXegb+T99Abx+LRryX3nxv6Ya23dVMQ789YlDRw4AWkWBy+LPTjyNU312W2XWuMnRHMSTgAxyDmLAC+89Vw23d3We8u2bdjZ7nPQN6uv8EXRkjOgGRbyYunCHGgFZ5EQrQqm/7/StaOO3KmEa00JIPktSnAjCSbwTkS/7uwOmRcz3YsDLRyl/mKBVMavALRPgkeJUo16G4TOMkh4hVjsZyNBqwJO21yYVQEAIRuZB4D3wIH5ujktQDwkvQJI0BGMg+bl9oNQiymb0MgtSoZ1JgEAdC8QHl3NG+NBKJQB1Mmmy6RjRlSPDFaPsCfw2hGHEziGseDKykB/cxddZDZlVHdnX0TzBNZ9lbF644c2MnTygxkNv4I9cDXY9nOav3GPuceunKWxIR/jV+xvK9cwEhCjUpx+i7aR9fFf2/cKNntGNXew3SXq6LpV6qdJfBt/5qH3xcuMriY4EUvM846E9c/EEWjYjXlXlc9XhE8/qyBTDYoBcGZduPvalehChbfa2I4BzCxkCu7vKYLk3EJHGEhDbLZh5Vx0XpUXgAYRwChwCQD6diZYBSeOPMcm3QBkigopBvg4Hwm7ootBZmAibCR3n6mf2S57rVI3vYlUcKzb8Q2OjjyAARMrBEIgFKNPLIAzGYjyE2QpeyID4mg1LR/owPgC5NrPakzoRJalkEnNBHa6MMScFjYhPTEfaR/dmj7YscA1kzeXGix0HEUQEsc+h+S/ykSbyvr9RQVZXcc+ZRTfcn4tWK64HBViSRfPEP1fZM+7VW+8ILr1t1NijG5i9eMWyrg/8QKtGZ/x03TJ6xDbWua5swmMjK2UzR/JfvDYowEUuvgPOVYX021N3pM/R31591+B7P8X/mis0UFd9HjO6JaBQ6h8P2yQy3RVmRQaWBNgkkCk9kuP2YyLDZ8CepNdtlhl65+OgRdJRID6GcUPQTiVRRoH8VkU9QqAPScXUjHCKiJQV4X9hNvHT0jDP8rioaU1WgR310qqoa9HduaYz7zINWRGUFV6zyo7U/izyAJAqiSQ2YsVxYR9pXQJOxAGO+IFFuhhfb0eaeaNEGyxlWBmrIoYSfjKeBpJJjhhLLihIRWvP6aLRTgg9BifpGIGB54cLAfAxtlaKqatwSBgQk4CFRFx4+x/BrKQlxKJHqakPZuaePXX/p3FsPbDFWMyWmQnnQTzgJuk9VVXtzeffxVX33/n75tgNXLj3aRmcd+39wA+KvUCiyufvvf1tjLXhXrYjCdE1Ro5qqRDQFShWupwGwh5ficU3hqqYo+ExEUxRDU7hOnwPog9uuxh2r/PI8TeGGpqgxIBKMpoYUP+KQT9yhqkonqg3/qbEF7o82Gkm+d0B4AaoClH/3i2cofeBzckFcQP9N6RrGgk+ASp8VWIPAGPiNT/hdhWGg73YA6q5EPAII6/kGAXpuZAgUnalkGHQGoKrGfRg6lMJ1Q3WYomKN0xSB1xluDIfmHkIKQ8oaSM86EomgicH4uH3FPBUS5P78UWzbUFV11Hkbj8dp3iKkofAG3Y7JZNLJZDIsHo/nM5lMG1b5EH0Tre5wiWRnZDjxgRUcBCfV1dVofab2Z8ADvf5+r7q6mm4mWJSxb+h7OUsm0VBAfzdNEw3deREywMsgbIRSU6PwQYZOr3Fh5RpVhXo1/Q6J9nbItHumwjSV9/7+6U2Xbn9g+7/bOi0I1RH/h9/lRt/s7RgoJG96bPMjnuedwDnfHiDb/h/aRFKP3/PrK5+4ZvHdn3nm+ddP290z6GqGAXAJUMwK0VGJ9Bm4DHxxN9VDnO5vgmwIbENUXfLHK96Fypvj6TaCjvq6Bvbpw2Y/9avbv/8UtLDDGI6P23yFDYkdEMnMwNWX3M4ybPCxCYFBWMCYt5I8XkGgFTIGMACKbwR8AjgEmH4zo6rQiMGK64T5EAwlGlEpB+/nC6gBkioxyCCAxJrGJYIkR4k4mlGXZLFcNYRg+iye0+LxiVi5qzKZjC+bXUbu+n3nruvYtr0jKKwqippKoXzoER9JqP2Z/p5MJvGenAOoUOzE31KpFNEV0B1RVV6VTkvqAfIkUpBk9KUTyG03TTOLuZdOp2ER6A8NoC9DXICwMp1O946i/jwrnHGXlYe4FsdFVqjIhreNG71IXV3vdMkFTy5Sby9nhkElzXQ6PWr7s6rqE0XpBPhtqYFHCczq6upQjXWxcuFnpq84/eo3/7C8v3SeZRdskt302SuovZezkr2hR5k572svAFRznP+9Mhr+f2qji75x6TefZozhtdcmcykft1FBhzAJfrpHFIZog5zxR69UHPM/2sAkLopeIaMQqg6Qp+DzGCBU8MFKHjMtSxzjRX/FomXf5y4IwEgCd4CcgRL87jEo5cLAYD/pc/vAJFKtcjUlQl4FfUaRBgGeAwBT3GWW5/EIUw0j3TfOnXnDtcf86tdhSgBSf1ZVPz5GWG14jkJMRVG1qoq1c17dO1r7M0J45OjRPpDL5eBpYYHFnNhZXT3ivvAQ5oh5Z3aDpVoAmhzHidbU1GwKz1up3ESZxv7+ftlVNhxegFRkhUuDz0lCSNdTSf15uBZmuamZjE4ZUNn34GqagzY5JO5Ga3/u6+ujCxHvIQHC5HENwzCXeZ66kLh+BaKUM+WP1x/yg40XrljwdiubCJUXpDNkttqDKoCTt1dts+cff/FLD8Qj6nn5eUs070XCTPw/ZRRkGTgS0VmxaFaxvj7MbJUNUK+11+f2ejVVVR4jwpoa+n+w9THWP1jWFESm0HV6vHQi4QLLoAyqvGZqLWfV1QVwOf6n5cZgQ0Ze4g1oCxkDYSBkKFAGEmHuizYJv8gALQgQqFNvRlBNEAKtVDkVVUT8q9LIJ4wFtKsqzhX4Zx2kO4rtGwQhQo6So6pwh6tIAiqsWm967rMtF1161NSjNl7HoKvrb2hRRz5AwIv96oFnWFyPYJV2saj7ebEleA0dhzI34IJsCNwkkgAar5qamrCeyF77ypwehSDwAotFurHxeJx6fcJzD+UNWCwHJ1VbW+uMpiIrQwg6ilCdRZ81MpxNTU1BPDN0C/U9kO6D4FWQ1QfIxe+VxJPn0NXVRTdBxGAy6QJDhuMaC2Ox4FwlxhRMRw/+fdNF1/5p59M7e3Kej00IQd0URbMd035zp3PuMd98of3pX8y/6uCDm+Dl/BekXf6v3yg389hjzzT++nfPPLj/4RccYJk5xkF+SH38Pjfi8NqKouXX/1nUzH0AE80+n4jEUxXTq6+q6lp67V2XLF568cv/FaOAYJMcjoDzXP7rewVBQlHgBzy0QGCO7F1k8INq8hLKpUGcL5KWaGqm9AP3mEYlRLxXbjGVIUPRNk3NiDBmKT6XJAwM1B51hxvVhqqa8S3j41N/eNURt/6FrfF7p0EM9MJra8ce+6k52/G7wBRI6D6+F/B9JA1poRtt7smS4hA9xoANfKR9V69e7Y0fP56qd5zzGLdtBfk0kLJih7a2Nlah/gzLg9hdMPpIvMCwDynUiCSBE1wIrGDn0RJ01CYtABOyFClzEaPSTIOjYXBwULZ9IlaiQE5UPVwo0FTuwd3581doX/3c9GfPX/rWksfecZZms3mbuIf8wrGftVZVbTCfs1ZujF55wjdeaHvmnmN+zebeq7NVF/2/YhQwgowDjzz/zx9t7z3SIRBRyaeiQ54At9mnPCq77Hv1sAW0Q+LxiwWK2m0AGrXYBrttXF9fz1Ovv/7RvMMO22/9SGjVYU9QcBcEtT6xwvtuvR9C4GfqXKS8AFxPYamQU1yJnADtTGQlQcwvQhw1MAZw/D2moXTi2kxzGeCOFQQpJdcyVQCwdD91QTQ4GqDsulfFJt315X2uvWb/SZP6fshuUxln5u13/u2AxTf/7eqZM8bdglAc3yHFWoENyJpmxBA3Fryho3jX/r3wPfOAiFViFrB/d3f3iEZ2zpw5mD+yCljuwSK6dosPVX8mbUes8qMTRPvnhG5I0qNz3RJeguG1opFppAsDuAKoKLdQoIuQx/8P0FSAVQbIrZBRIgLW4aThFixY4CKf8LvrPvmzo6ZrrzIvivQi0j5lmnaqPLhaxjKdN9rZnZ+9dMU3OIwBjMJ/800wR7t/evDpGT2DpSORzVJU1eOq5nEj6nEj4nHN8H/X8FJDP+uhl/hdxc+q+BwSj4r4PeJxPWa29xfSf3n83yfiuc2fP/8/quzAl6WKWBA2yN4DYQiQNxDwY4KLKOhtg66t/IYFNEh0zU/2wSBg4lPRkKuQd0e9gF5UQMR7TGGaS6lTbvgrOMYSfVvUi3qKo3kaj9qxRJQnauJqTXTcmye0nHny7Uf/8ZL9J02iTJ6uK84ZZy/+2ptvfvDLWES588tfOPwtDNeFCxcixCWQESpyEFIRjYIK55G9FNWH2SiBKFiSSI8V81ZPp4McxQgLuaRmR4VRaq5IQ8RBbBT+sHA3TM1SlEhfX598WMOdHKoJboQBnhzxqPNQ5AOGaM97w1GowZOgk0iliPYJk1xIve2lQDvke9C7LWXhBAUbEUzI+pI7vGrRUpflWPU/bj30un3HR3YBjU7QN1owBDSXUg6m0pfLOC9vLt596iUvXKLAKMxfDLKT/5aIpPDW0dtvUOuvW1I8h9SFKKJGgAcP1a9Vg/ZI5PPxe5Dbx0t8JqiHu8E6Tt9BOB2Fo/23P+dTksly4H+++R6H7xEQwXtFcpHkEgJcQnkYzRf/5k30HQJWDLgCpQAJbahylWmKeHGNDISq6EyjUqLKAN0Mn4VllSzPdbkaVSNRPdXelJjxjdtP+tNhXz7gomd/eOeVScNQ3VWr1jUdecy3/1EqmYtmttSc+sPLT3vJN75+iQ9hNiTahnpJhsEg9irlEIebe5IygPIAAtfjctNU7VIpXl9fH2rMrtzWrFkjSYgwbxzXLdK8FcxnyAtW7IPYpQDeBsdx8uJgxurVq1HW08Mv0SptRhFi6C7CDHo5jlNyHL3Q3d2N1Vt/55136CX3w3eheSJIQPrdXMV4PJ5TVTWraVqhra1NG3I8/E7v9fX5SREcR+pEOI5axHcUfG+j4jxDL2Pz7l6dseiGJefNWDShKpIH3B2lx3L/qkAmwihkBpyVmzJ3nnDx81eoK5fajC3ZW+fvv8kmSGL4FVecvT2usz2elVc8Ow9QEUYHgnHxclzmWq7nWA7DCwSKLgQZHIgy+P/iPfzNtVz6G30G9ONgYfIcz3b0sbUp5XPHz39NMlb952fqewcEYQ6MgcwdyMYklwHoTk6fwkHOQs/Mp8llLBIjHmNKLRCKQFFAh+Z7BIrGdMVguop/dWaoERYx4ihMeeagj15/UWQnXZ0nFTXi1vNJd52S/NbcHx3+M/A7aDCot/zgpuy5Fyz92pe+csPaaMwoPb78J8cvWXJpFvMAjYM0lld7BhqcyEsWiuu6rufxwvgWCX1NzB0tPA82bdqEuYfxj3mKuVDA73mwlJdKxNOL8S5eFfOgpqZG0zQNxynqulpwHLWgaVqOvAtdhyK7Ep6vWjweb8PAj8e5l8vlmnK53JSJEycWgEsQ7gZlNPElsVhs+0iJQ9RJoRw9Y8YMSvuA4RWJk4kTJ6qqqmYSicTWkTDsuVwOCrRVuEkoqYBiCjcOYCjDUAeTyapNQ/fF77FYDPHRDMFGS7BpyGtL2PS0Or6H8yUZxpbuuvL2Dy757cu9D3b399pgraDJLqjC6KtVrgwUss7Lm7xbPnP5y2Oev2v+lZwv5bIdl/332qh3g3Pe/YWzvn9BIZ/5w2CuVAsZPMVIUJwNsKFjlwSAB7V3P0fgpxTKj9Cv5gGcCAiwykC8KvMNiLmTMXXgyMMP+tFppx35+n+50iBp0kk6bWhysQw0IhcGSUNVYRo0JUJbTNdQ8yADAIwBfvbzB8gduAxoC4QU6NLW6Gd0KuqlpGFYMvxcyVayqXX7ro7ryWPO3f+ClS67TcikMeuVV7ZOuuK6m37xr5c+Ou3QQ/b/+bIHr/5NR0fHDMMwci0ts7OQKSTE7jSuVUWqwExegcOQ4xpU6ZlMZvq0adPMwcFBWY0DRkeZMGGCFY1Gt8diMW+EbsdxpmnO0jQNWIIw9YASjUYHY7HYiLQF3d3d42fOnAl8BOXpKnKy8XgcDMe0kosvDNRnRZZztDiHqgCIS1BeIY+g4AH+FSjujrK/xBUA8UjllCHKt3ttwnMKKhUC0ITOLSKdIMxDOu1Q6LBwmXrHlZ/43Zk/fKflsbfdazOlAZtjGZBswX62E7QPatY27ZXb7B/Mu3DFuH//+qgL8AClVgD777T5+oi4b0//+98fHbLipVcObN/dWTAd6CwxRVd1nqqKxeFSFnOW6bgAF6jUAgvX2bYdl24ZZ0okYiAnreRypmlZjh2LEVjVi8UT9nELDtp8yilHbsUz+q8ZgxBkmaoCoE4XwCIFrp4f/aESQRBnJAahNUmn74cM8BJQb0MeWvOAMIRB8Cj9iWqCbOBAGAHuAhXfo6quwtRiNBoldqelVAJcyr4+95J1jDG8dFXllm27ted89apvnH7+5VenqmsSnz72Eyc9fN+1z66/8qyZ48c3FNE+XSqVYpblYmGEXoL3xBNPjHq5QCRKun7Z0IQbWSqVQKg68m3yPDeuaSY8spIgThY8JFg1R+0FwqLquiQdQNWPCoOQy+WoGcIwoCTtJMuacR8v9uqLNbsgdqCeb/xeYiWXe6ZbKIze3CT7IEI8cZSs/E/oFoXBoLwEQaejUbzMYrEYNLR7yxa6nC9U//yTg6875uLXa1/Y5F5SyPdZXFV0auMtGwVk2LWBXMZ+qTVx3ie/9lpLbmPurMSMxE7G0PvwP0hZ9n/o5hvoxcqxx+63hTGGV8UGqVLSHvqf2H6BHPv/lCam0FEQbcr0QrlQEcZA4ZihZBw0hcRl6WnCGGD0ZPOupWpRioA04Cg9/9+guiD2I6ilwl0lkdRUz8LYD5618GwIBhmPRayvXXLzF/efd+aNW7bt3mfc2DG7fnbNBceedtpRb9577xPxsZPGFR3TRHcj4XVgQDEPPC9ih1mM9rpK+Os+cI8WUHRSgR2MDPDHVyAUE3qSPqW7RBejWKBaQHh93B0WBK57dTvKsiBwBbhBMmEnGpZGLcdBQ8Y0PZZMarZp0tpB3ZKMRXgiof8nmWUyOqJt04NVFc0eo+6ESgNQlsKDgYsF6TjCdZdKpaBi4Mt8naGuuPeIby34xsoxK9bZC4vFAYtrSJPjE+iOkaA0ptl2zl61yzriE1e9/NLt92z4yhXfmPmSyxZr8+cLxEvFNlyi7MUh78v95Hvh7/mvJto+bhvuHIceewHxRpIHxRhH9n/lypXOzTc/3Lhxy5YjuetopZJlZrMFArljYOJfgpUjQeahk8GDXh4zoVkA99byvMaGau30049/5+ST5+2E4Cp2gUjrf/UKsB7KvgVUFFBB8KHHfq6A3H/6lwOyLjAFCjPxh5CHoOpJV9OypP+pMts3AIRBcKnUSMqTSPhrUUWPRpS0wt6bmE79nDHWS+Ei56CfdGJRw71m6V1HP/HcGz96YNnzx5aKA+zg/WeteO1f90LFegMM66JFp0IJCdyklDSXwCHOo5rHkO8aeZUHyr+U8ayi66K0rnN4yIoS6Jd83P0KWgL8TbKhj3zA8gavHPJxtO9QghQq54laP4UJsB5RzhUo4o3+vTEPXo5pAt/hr/iG5ylDhO2H9TRk+7NklOUm90qMaNc/zkWAyCwxKkE+uwSjEOWaaqrg2kVIEJyzhDtzvkR5+vZPnXXCZS9HV27wTi1ZBZurajl8IJgCwWU1z8o5m4tsys/+YT9/7rWrLv3Lz+beAwr3vbdh3xzh/f/0vf8d28qRjukhVn7xxRcjBx95zvKNOzoOp5VZjYDPQPAhiE5wXwdNkKD4vAn0GQwZLM8fFti6rZs+9DzvaEFt/z+M/vSBSb5hkEyfSGXACMA8AUiEVR6sA1RSpLLi0M2ixiRACX0teUIzA2zsuI6t8UhcjcYjrDpqvDqrseau204+Yhkgwfez+RpjKx2kJB767TNz7l/+1I9uv+fRM7t7LSVZnWKfnv/JO/+x7NYrAE32tUKXOp2dF/OqaBRUawS8EyfgWtmsKsB1I279/YzFFFoIYUwkAInIhqLR6KgGIc/yPMZi1IgoDYgX8ZD2/088WiQ6ma67Ngp/pOgsgA4UMoCXDSox4RomuhA1RSFOwhFPKk9FChnzI3pwPMOzI5aiZ/P5MPBir+8APNlAk5TYl6NIywgJSZ8Pn+PQLWgXBQdkiXGv6DkiP4DyOp2zyPSG8w44ny8cfemKv7+2WT/ZNPMWxpVPze0Dl3wADPJRnrs7V9T/9uHA3bPOfvrQhq5nXo0lFN1VYx6IQZCE8jwV/F5eGDxGVA5lKjnBRuZx22QuCqjUi6OXV138C+DbcCRtlOgi2TQk7yv/BiAODkW/OPQJuivgqbFtwg56BMbVFe7ZUC3knspcpWlcnfOjq776aHNzc14m+66+esm0vmxmbmag12HRJIlM+GNb8hGGHh0lZOV7iFzJw/JYfpB39Br7b9y4DYiX3v+61FxoEzPYNwbl/gHwECFfoAFdCA8B7r/CmaFpLDLE6a2KqIYyCKMAkRvkB2BdXFXVVS3OLTa2Ov7MnMZxv7z1tMOee9Zx2e0CYKepK+3H//nqpFvuWH7Jlbc9dMmurr64axbZ5CmNfScfP//Ke2/51m84vw1aEQpb4nfKYD/wieqspBpGBJwqJKwiEIqSeXkviL4c15lMRvVMU0cy3rIseLxIxsBjR1JxOIyPnLeeZVswOh64UItFz+Ulzk2LGNWJQmykW1woFODRcK9YJC3KvdSfo9FoRAxQyr0bhqHjwkrIkAzJcoY3GBNYNcCJkVj0UVkR1+IWlGtHHRQgoMwXChFVVUs+GMsGJMwDN4NhAD9QPschG4QwCZ9g2yoEjSQ4ih6A4ziBcjRYgugOCvYLHMTzvK+fdtUrv31+jX5iLt9ncwVlnjJLD5p1EK0CbDNYyHqr9uhfiQxM+Qr/8F8kG8f0JOMqIKcEsRBkHmXV40CuTRoyUgCVCGqfFpx4d0gwVTpPIaEP/0SDUGbIH2SQVTax4lgBqFh4PYS5QAWASEAt5toF1pCMsSMOO/YftMCILzQMtsUwopuZFp3FVcOP1GVOl5xGWWGQQH61zGYECgTVcD09wWuqalpnzJiCnMR/JLsx3EZJNGowQIXDRyji9lHZUBGGACKMMAqgLuMKMzSDxTmRfQRbY00srXZkwNTOmJZQo3Gd1XjFvrqY/sjpB0559EtzW14VY0LjimKnknHnsSf+ud9Nt/150aJLbzu3e6BQZeZznpGoZgfsN/nNyxad9pUvn34iEJd0N6iEu5RUIdiePXtAUhLlHPlEEyEvrb5SXUnE6UPHMf3e19cHaDNWeEiz0z4y5ABD+Uj4BPwnK7qGsYG7BLkf8fyxMBZGm3sDAwMxIjhWVfC7cw3KygISiRMpVFVVoUQRoA8HhBIz/p7L5ZolA4uEHsslN16ID/ar/WillPLTBKmsr/dbR0TrdEBNJQhUiKGpWCz219TU9HZ3d3sCZEFDHfTS8XicyipDPQuww+RyOajhtOMcoFINwkj8DQatrq4OJUkjm802SnopQVtN1ZBcjvHXX2/re/ymI75wwndfvvulNfHzi8U+B+RVfmNFEAL6lhmvUsYpxiZ5bOIZjLW/yFhmG2NGGsEKxai+Cy2aJ/2btPfkDia8iE8qVl3xYf+Ohx75Xn1m8i6Wfw8yo0M+S4cRGTj84liM2UW3X1W81avfmwJ07tq1ayG/DvGdwpcvWHpfe1f2l4NFUIZhOcU1hRLVUvcxuB7CKfvMSVbJTqbSxiFz979XIQk2Wm3/h5KJxWK/6Dpioj3ZDxPgHfjGAL0HompAZUS8xwNRXRkQOdzTbG6ohu6xcUl9Y0tt/N4LDtj30aNmNbQ/xhg7i83VkUyMGpr97HNvHXLbr/908QXfvOFLHf1WxHK4x6ycN6auih9y8Ow7n3zkJz9AqmSfdeuaZk2apObzebj4JNAI3gHMDcuyWhOJantgoJONGUN9EMFAyuVy47BwSh2TEOWZYtt20bbtVkFYIm4v59XV1R7oAwAJCABJol0an8Fx4/E4cAzb0BIt5o1whWt4TU0f5k9jCNxE4x+LcDweB8FQj2mae+QxYUFqDIi3cq5VVVXtFJyCI6k/S0l3aswQF0pueWumdduUKVOGTWkiR+h5meZQ+ZGWVExcrO79/f1dNTU1w+67Z8+eZCwWq0EpExgFCZUWArAwBOiwHXYTPeTjQGkNRhn0baCCgpJmNOrZzz03oevww7mrK+wrx33z5c5Xtjrfz+R6XEBuPVAKyURjAJ1VVCIHilYzNvMcxvrWMNa+kjFz0E/JE69e8DT94paY8OV5Lz2G8CQL27rQJA8PJ28YieHyD+GbPfTul/+lshpmWNIuuJr27pr1CzWVvSG4/LGIKn+4/7pf73/k149a29r+RcfOW1xRNWofIvlz35gJSGpg0LjKPc/M2yqzjUNmz3n77tuvuPWeqsx/zLA88iZ5B2TvgZ83IK9AhA46hQ1iOXU8lrf9uohMKjqesXPfsdUvzRkbe+DGz+6P/EDhAVaufMQiH1qLf/LHQ/7xz7e+fdai687c3duvk4etqCVVUSNTWibuXnjSkRff9rNLH+f8p9AJmTBj/HhgZuDaYwxSmd40zbhhGN01NTWdI10NcDpCJY2FEuhUXeOc76qqqoLIzF4b5tfg4OAYx3EA8pMhBlXWXJeIxnbW1NT0j7CvwVgOhh87ScUoHR4Y5t+ePXs2Vag/U4kvwpjBmdfT0/Nx6s8+rYposAhxMIJpacT25/b2djWZhKZ9MIKtkHcRUEuVFZHlaCAvg/6F6yVWeWoQEfvgArVhViHaB64UxGBEwrLsWPs9FQ60AJcswXHPUF6498gfXHjT+5m/vGJf3zkAQI7mENKV8gm+m+zPYwG8sYqM1R7IWGo8YztfYKx3HY1cDyuwC+k64curOvYB2A7MI9QWUx7s0hKIi94bwP2f5+T2+mgo9OCVNX2mxbhlu877G1vP//ktD/7q8su/GhDFiGdw1pzDzn9ke8eeL+byJWjZC1ICl9oFAUASksr0nud4aixq6AfvM/3N397x4y+h542ep5DN+x/bon4/MnkGyBf4uZQgVKAqA/ICqDLgxBzHsh29r+grG69k6Gdh7FcL534Y1fj8v9se+3GoAc/zlmkXXX7HyR+s3Xrprb996piebFFxSwU8L5PZtlGVTkb2mzHhsXvvvPGKOS2122U5EGRAtg0GOX+VkOSoSkzJF7KFIM4PuelBjgD5Odu25aJGc0HOqVQqNVoLM3kR6IiUG8ReNAS4rmt2d3c7o+xL7yGPAQ4FGBTMJXQ3C6HlyvZnrNalkmdFo5CHMoM2ZKmBWL6BnsQagJRBJvmIOQm9MU1NMH7Dxyq7du1ifgOHKTUeJMkdNnfy5MlBbBVo5ZXbn+XvJP0muRhCF7tXXCb3hViFbDkVrhLAGlTYBUd+6ME5JXux8turDrjhqttXf/CnV3c/uLnPqWVO0eHMU0nENLi18r7bjBW7iXmYTzqZeTUzGd+1gtWpfR5Tkxz5SKIGh+fAQQRD+UkhQ+5/GaJ9mla0woUF4csOgf+LvyL77K/+G/4FC4+Esv5lRHb5M/4+Uuac8gtITirgBORAwtQ9t/LD6z3POx9GUTCYyvzKmV8859qr1mzcftmunr6GUnGQWTYQ/jiKSs2jKCZHdZ2Nq6/d84l9Z/5q2QOLf+YT3fxPJBLDZUcR6pBTQ4bAV4Mg3417IDkneCSPxhE4K3Evw8bGjcDT9BeZM5SC7XtAsZjhtG/dM+XiH/3i8/sf9fXz27qL+/fn8r4EvKpaTI+oqmcaLeMbuk864dir7rvlkt/OabnLF5lZvhzjD6t8RCBjzVIJCUMiGGZmyQQ3KBGWDFkYpZEF2E6FpDsWKNlfQITCnKtABqfT6ZHan/0RUCYc5hBiEsGFiPiHb3/euHEja2xsRGke+UBK0yJhKb7XQp9RRfuzyAVQNhStxuKDw3ZOiRVdSkDBTSLeeBG3Bd2OQ0+qqamJRGBcCG4xBr0GMovipkCyerT2ZyQOpUWW4Uag7jTaoJL5jVCoIlFc8vpCFnWpW7KXqUu/ve8Tv7rv/cN+/XLv8vV7vP09s98G4yNS+X4sLWJxehp+i69XLDEWn+ypLaezw+fEu046oPonB7Tk33hvm50aKDluRHMt1dM1VeERzlwy89z1sMw4lut5GoHsuIYror/RE1GApSGdQboW/E2ubzClMnPkc99UXjj9fZgboqpQZnYcyynQFSkKuhfWQ0xa0ASLJIHPyM85+7E74N33revu/Fzn7l1Htvd2Nesxox55+nzJ7KyNJbZPaB638p5f//A5hfNu/uASAeL5XwP1lng4v7JAqEIPhgBP1TViSjQeVeKKxeri2keNae1vc8eNe+LSz+z34YfLlqnLz+DQryRjn4hH2PU/fXjeMy+8dt4Bx33tjM7BTF3BQYinuoqq+/USy9brq5PsoJktjzxwx7VXThgfb0OAtXjxYiQOgxssPAKiDEgkECpQ7g0Jcbj9o5LIEjQYy3poPApeEXJ9R9tXcIgQeEgaB+wIbvi6uroR58H06dMp6YiDIlwWrMyeW3RNZrC92p8Rm+wjI9VSqdTW0NCQGW5S4/sGBgYmSj4DQCtl/4CZMfV0Q3oX57xCjKI8H7344OAgScQBVsxKJQXdleL+2ul0Gj0S5nD7IocQjUYbIUGJDGp4EuMmIQk6nBqOxIdzzsegLwItp5JcRXg6xXQ6vSW8L8bf17/+jn7ffQfD5Z10yvdW3L9idd9x+SL4GRFaCo4wUZcnrSIZ4oiqEnNUXpuIbT9k5pjfPPOLQ+4R4qP/l25lhGE0QkIehPvHhr4m2wHuxAl/dlTh3v9kk70jfbsLk4+6/b2Na3uyajqqORAi4sxTDF1l6ZjKxqSMLZPGJJ78RGPsiR98ZtrLId0QGh8wAtlccfxF37r15Nff27hwz2D+uJ7BInOQrIQAk2/g4auSyPHUiWM3nbfw6Kuu/PaXHqcxMX++xla+BCqUivPr7++fClyO57pWKco1xR+T6LiL2bbd0dDQ0D7C/EFJf6plWeQhY+BHSTUk8Hr3pNPp7hH21QYGBiZHo1Hq8wlVLfBhPZlMtnHO+4fbd9u2bdH6+vqpcjEX32frjl4wFRO0iLtQepafl8KtBO+WHsJIJY5oFC2TfpZSsryUSiXuqE6YxSXY5Pe0tbV5VVVVcK2AiiIPw/dyaRYNa93kvqIEqkAjFtMxxFtPHsISCD2OsC82wRUZLtvIuM/Z25AwBmMglaieuuXoay67/Y33lr3ad3nHINeYV7Q5AEvETi5Yg4LWXCQQVTxe1uvwSS9uKd4446x/f+9LV71y76+uOPjBlkmJDZm8w2YvXGZcdlyLV1Oz9f9vDVO+HChjC33NxVF6VJb7MlhnnKEUly+Xca/8LOXxFi5cCHowfMf/0l6PEi8pUFrw9LRiJOJKjBVYbcTcOXVM9NV9JyT+fN3nZv1b4XzwYcbYlaExlIhHvBtvXn7Y8y+/ed7UT573uc6e/sYshgAlCpnNjRh4ULln5giP0DR2jD3/8APufPi+a36Fw77xxkcNyGWzlSuHBfVgzJZQjo9GFQ40ImgYOYjDdFliHHb+4D2gGJEYRMiAmLNkmtzTdTPKo3rRK47Y77N8+XLvmGOOMRgrOSIpHk4tj0o9gHA8l8uRVy1WZ0osmooJj4ZEmCvOs7+/f5osRaTT6c41a9agtZJozeQ9EAfWisXimEKh4En9RlkC4VGuugW3e2BgoDR58uQgYSL37+zsBDahvtowrH7F1L2C50ktBpxULBZDZtbdtGkTXJyKQdff3w8QQz2LMBfAI0E9BUOEVcrq6enpDJ1roFSNxo4xY8bEPc+rdgzD4kVqySCrCjRYoVDgvb113YxtCvYLfQ/v7u6u683yqhmT67bc+/iGI+96ov3O9bvNfcxCjwOFehd8reHyXyB9RhJnPk2wp2hRQ2HNMbPnoJYxf7j2wjn3Hjwnvd5fVEEBv8bzPMpq/k+tqv+dNgmS8jxv4tHXv7O21+Rd+zRFntp/QvLpH506+fXqhNE/mLckLxPdt1hEY8sfWznpkafe+Py7a1q/2JcpHNHZP8hcMwunzaEkKADPFMszl9m2WpfU2L4zJ7x81eXn/+LEYz75Sltbjw71wdraRCaVSoHsRMF4DLvebW1tWDQb4G0GnXUhyQFFUQY7OjoG8Vm5L8YUQmIk88aOHTtWaCzIRQlcH+j/Ad/HYDqdDkKONWvWENuR+G5lYGCgEd71UKk3nEtVVRUSbUVwkw6dCx0dHShLUslSNgCGad7T6XSH4JQVMrWhDZgEsDDLGmm4BgbehHg8TlRQw20Co1ArqgDYZEsytU6jzXqktHlvb+8kwLmRj8ALZUIhVolu9t5EIjGsGyZdOCR5YMSAyQ5frG3bu0Ypx6BFm/YNvUdZW1VVQWSxs7a2doAx4PGX4nyqP3fFa794bUPP+d0DICNVHU60QKzSGIQwBEQzaBdd5tiqbsTZmESkb8a49ANXfKHlj6d/Zvz7xWANWqguXrzM+2/YZv0/vGHyPv7itimfPWpSt8I5SRnTtthT9J9obiwaYXvactMu+9Gdh3y4sfWMnV3ZY/pLVioLj9otgdLF5Z6leFAmhIfu2i5zXDUWj7F9Jk/YeNzhB1xz7x2XLR/MUJQb4CW2bdsWraurmxwaq9QjIHJXwNisGRreyg2hcSaTmSzjZInKkpWI2267bfPQjs9QWF2dyWTGQbZdCK/aIbYjhNWbRlo4sC/KkqHPE3ER8hOgHojH48TpOFL7cyQSqdY0rQi+hopEoAAwQD+B8KqIVyKRCLljYlKPZhCaHMepxskIxSYANQqImYBtqK2t3TGSQSgUCi2O46RlfIPjI+63bTtVKpV66uvrd40Ulw0ODk6XhK9SPVqAnjTXdTtBTz3CvmQQIBATeg94BRNW17Ks3fX19b4O5BnLFag9xQ3Grr1z7Xl/en7TT9Z1FJodx3S5HkGo4BP1kVJByLESNtVnM7dddM5q0VrWGHessTHt8RP2r33sxis+8ayh8QGfuXyxwhbO4Wz5QmIG/H/RcwiV7aQ7HGzpVJx9tHrb1Lvvf/rTb72/4bgd7e1Hdw2atYOWIITlzBHCsJCCJrU5Incx8xr6HVomjOk+5fgj7rtl6SJUQ+AFcMYWcySUwzmrRCIxTcwB2QkouQxRHtw8Wr4rHo+3SNUx4lXxdLvklWxX162qaLR1JD4EyA24rjseCxvyXIIMiKoRGM/DGQS57+DgYL1YyH3EimgB+E8MQm9v70RN05LSa9mrF0RaF9FKDCglrWNCmWnEDcgpw+/HDmIbwY0gtehH3JCHQKVCClSK/YkU8uOyr6HqBjwTqVaNazD+w/1ksoVCoGKxSF6G7O6jbfkZQlJ6ufLDRbMf8jzv31/6wYrbVq4bOKMjj+YfE+VJJN+D2iElHP2EvS8U7OkquAZtO+O09av6roy3cNO/Whc+9mrr1hO+/vyfzz5x6vKzPjf5/TK0Gv9drCxevIQtXSJ6gP/bbh5fvBh8CWfQI5e5CjwBx/X4qlUbZt5+3wuH7ujYddpRn//R/J6BTE0OySykl6C9RvBxFW48RN0AOSdFF+AjIpGE1jympv+oQ+c8cufPL308EdXa/vnBB2hFVkXuo2J8jRkzhnoDZBgtvAdP11EYo3B11OcA3Iuu65ZoNHJhDGiMmaa2/Mkn2ce1P6Osj+StKG0iJAC3x8c9e3jDyGPQL47jQAlNLuqjYscxN23bJmNANIVDPQTXdWukdZElFpICVNWBRCKxa6QvBoJL1/WE/zzoYhxQraHvRdO0LNz+kfYtFAqTLcsCGYcpvAvZJIXz60qn012jrPLIgQREKSFsguG6LqG/RtgXtd/pZaJLOiZRW8FA2La9GyHDXvsuXKYuXMjYsoULU/f8eePhf1jRft3qXfahgzkRr5IqOGqEUv+FnpcQR/VPjxDByME4JlibFF1PsjEJ7tSnIq/Ompj4x5ePnPDcKSeMXxszuFUMms6RxZ/tMc8HjInaoPd/4+oP4JvfW7KQM0bJSpqUwgBgEtbfcc+zh61ctXHu9h19C3ozAwf1ZN1Epphjnp3H6uP4lGrUBwmH22diJtebucz1VCOSYOMbEoNzZ0968JfXX3BvY2NdRzabberuHsA6tSGcWR9yjvFsNtsiPQPxHmmA4PdoNLpjqBK6HCPd3d1pwzCapZcsQxEZfj/77LObzjijkmhH7tvb21vlOA5CBpp7YlzSBAQFWjQa3TKKh0ChPpKaIFlBIl3TNEJQ2radq60lYNWwW6FQmAJPPCAaCn8xegZc160OUadJjwFY7cERDAK5Kfl8fpJw+yXJAyYX2Jy5rarZVCoV4KWHbvl8fqJt2wmJLUC4IYhY0fCxZ5RJjbzBNMRcSJYIYxJ4PZB5GyVkUBEyhBJCUpgWbhr09HaMtC846CxrfO28eWNRUqxeeveHX33sze7LNnY7zaXcAPKJKGsJamd4CCB1DYGNMC+CcALrHPjHVI1pSRZXPZb08uaE2uhH4+tT/zr6oKbnLztn6vuGwnuAfwwtZ5yxZQqM0+zZazwgS/8PMxA0OzH5X3yRKStXUuwcnL5IyTLdUFmhaE/62a/+PfH9tZvndXT3HL1zV88BvdlCY85VGeo6RA/PAUkGfNpWmGtzUnT24yrSjfds02N2UY1E4mxy49j+g/ef/NANV57/wJQpY6DsFVu/vjU9blw1VkS1o6Nj1zAwe44HItz+ybSgidA5zBy2adOmbQcf7FeihjMInPOJ6BYWeBcqSRaLRfAnIg+wZbSQQVGUsQbGvU+sQmE3PgPu0Y8zCKZp1ieTSRiAmPibDBv6EwkQ/Iy8GCPUF8/LDqvPusIg1EhZ9hBNOherPBmE4SDKfX19Ew3DQBwE1irCaYuKANyZHAzCkMkl4x1eKBQmhAwC3YzQOXSlUil4CMPCQQcGBlrweYQrpriBfrjiuZFIGtWLgZHgoDAmSKQIj0KVPRbosqyurm4dZV/0XzSDjKboqbHJjSkMnKpv3PL2V55/u/287d12lW2bgNQ5XNHBJeLDsv2bF36kojsS1wZvlwgxPOZYGmyEpkVYXVRhNXF9S2Nt/I1ZE1NvzZpQ9cq3zpkGKq+SpnACSZc3mYNYQ+8uXkyrsV+S+l/kVZRj/FAv1pIlHB1/6PljbBlWfSpZhvczdMZKJvFTRJ97Yk3jG+u3zv1gQ/tBbbu6Ptmfy83OFO2xA0WHFUs56pwk5lRFAR8aEJ8ADPgwTcGDCX1GVP89B6TKXKmKGmzy2GT3oQfs++hPlyy6t7bW6C0Wi7FdXV2xhqoqEIoSFg5Z/r6+vp2TJ0+WMHyJLOTiFqX6+vqmxBQly0RWH/vYih0xmAHtxM1hNiWxBftmMhngbWyJQBQhLGF20uk0qOTC5UUJG8bP1fBgwvLtIjfGBGamXPIYctxcLjfW87wxwpsPaNdFw2KmpqYmyP0NaQ8AnBoGIS25TyomtuioahCqLpacJCKmHtXSCA8hhRuPuAUuVpEVucIV3cpa+fr6+hH37c33TtQdvWoI/zwIVwH4yCWTSZRGht0ArIJGnZTEFqgvnLsGD2GkKgO2TCazv/g8hQkSOQmDgOrlSPtCluvUU0+dhLyK67qxoqNEYzHNqk4k+O7d3XOuuWfzSS+u3nPazrwRN80inoLDFZ8RUHgFIgkmZ5UfHVGqQYwSD1MdvQ+oVahxruo6i2kKi3muWxu12pqqI6vH1qTemFAfWX3SUZN3Hz2vbkvM4N0ly28BILT98JtIpK3lbP5sPj9gUnqRjRkzJ9irs3MNB0UyyJVW0n/WemH3frgNUGYkRwkD7XmRTHdpyp//+VH63dXb9tvV3jd9IJvft6cvs09v3m3MW2a8YJqsRDVYLLhk2ogokSChKBKQF1CWciP2ZR9HS6VDGNL6hMqmNDWsO+SA6ffe+bNvPGzoSpdpuWrb7t0HpWIxCBA5zDQ9R1WJsRuLhwDCDeHu8TcRMkzTNCfPWMQuFovwKuCJUii8atWqzUcfPTyVHjyEWCw2PRRuUGIRi6QBbpBIGgnJYZ8MEoOqqk5BIr58KrSCamBHjkajIQ3Tyg1Nh5xzhCrSc5G4BFAXDNTU1IxIslooFKY5jpOQPUKotzdjX7RCxmKxbHd3tw24MJ4QRBzGKQrvVhReX18PgANVAlDDxwaeAxlSmKZZqKmpoTKjuBlOa2srnzx5MiaZVigUqgH2AQZACLxISwkPoThmzBiydp2d1DbqAsw0YcIEuFKRSCQC7wGuFxKdcOX8RE06ze2+PvIkqqur6SF1dHTQzUBSEFLauq7HQmIytEkVm2KxOAjOBHG98qHiWtGQFamuro6GmKMkmIlcx1QqlZMPvaenBzkWZWuHHZ/eHLFSqZS1cVt+5jUPrL7o/c09Z7Xl9CgGFqYLJcJ9vTNRqgxXI0SeQfa+iEYHkq2mhgTqP1J9LWaPuAQTisPShutUpZKtMV3b0FwX29lQra3Zb3rVQNLQ3jzhmJmFxhqGOjXlgqqSummZNkmcjWI09tqQpdd0zvJF6GSSVBlOuuH99R3GmlVbqvsd7aBVH24DL8X+61o7kx5zP9k3MNCYLxRr8jZnuRKqfhivsp3acWGKKeXqaxdS70cwfIhnwTcIfieA5zLbQvigRqJJNrY26c4YX/fCKUcf/PvLLv7M63BSe3p6Yr29uWhdHTzogV6MAUWM3fb2dg8QevHM4j09PRHwe4TLg8xvNgJRbwHJQbPJdCezyaQrgrGIv+/ZsycNGnG07gssDY1FwJexCGIeuDU1ntfdTXdXjC2PdTMlH8/T/JEVCxJsAXLXRzqWbNs2bdsmzwrn3NHRQeMYbdT5fL4qJNsmiY9h4NBsVcAqLIF8SIyKeQQPGl2ZNTjFQgGkqhUdMiBZRZu63d3dTfuibDdbnmAymQQMuMy2ENqQBB0cHJyCpAdidjmpZOdWR0fH1nAb5ZB9Y7lcbobko4dRkA/CsqxoKpXaOkptF27YhFCXI0lfBXjuSGTE+iwsp23bTSBe8fcpeT7jmn/uqVRqzUgTQFBbj1FVpxRhhscicDf9gYMsciwW24uUVG4fbeqZsO80hHKxro827Jnxy7+1nfvGhvxZO3oLzZlBRCGWy7SoB8eLomDiNZLsQ2Xvld4PGpokk5OfQcM4JGYSECIhSQH6PC3KIhq4gSwWUz3mFgfN+uq4rbpWh8t45/TxaYXb5rZdfbnWsTVVXnNt1KlN6Z6huHY+Y+7Km1Ze1xVV19SIqmmGpiox13WrunoHlLaujFooFKKNddG5nb25yK6OfpZMRKf2F9xYNleKsEhSLRSLiL2YCdkGU3BEYPXHogERRb+Y6HNLkAcglZ3LcHBpGH2hdNtlVgGdpSrX46whHWfN9Yld+04f//jZZx7/pxPnz/oAN2bHjj1jGfNiqVSExiXkzauqqjaMVns3DAPxvpwcNK5M04Q3MKon3N/f3wJ8TCjhSIsEVuMi5131oyTeBwcHZ8neBQkwwoMUFYKOVCo1bOu0KK9PkzkFkfCnihx+Bo9CdXV13wj7Grlcbl+R3ET4Ax5SSdaiJhKJjViT5ecpAQfhB3GjZD6homMLiSE/YwrIJEmygQgiChyDrvtPcfr06WgEAZQ4mJyyjRqeRiKRIGEXIVRBOIESK6nc4x6QhUJeLNhC3wPrmcd+MgYEK5M8SeAkwlTtQzAhyH1QtyOg1o7jFom8FTe2VMIHBRFsxXFl67ftC2LEMPNUswTC1yCnEm522Suf4qMgC+ls1tT3mzl2x/1Xjb2OMXb/1Xd9dObLHyhnbeks7NOV95ARxin6cuUoPeC7Qt8mpoaYMuK/YjL5F+uD5DgmGyjc7DwVuUquy3NEC8CN3l0Zg3HWwhWlZUNHB+NcOwT5Tt6eYTrPMM2zRTxO6D//+0S5lOJNaoq0mGlb1NrtbOz2+zjAhNxTFO3gnLHCoOPfHQjhABfooIGTe0zH+QIuHMqqEmKwAs9Kz8tf+FxoPDAHkY+rpg2DNY2tyk+a0PzCkYfu+6erv3vKvzSV9zx0l5fYurW9Xtc1o7qaqMYxQQBHVeRYkY6hxAcNRywKsWOR97JF6Zta5UdaZMTYp45fV3dtzbd7pgAA7R5JREFUbko3jymxvfMR4TGpCK0TUnIWsmo0STUNSoO+syLHnzy++D66JgIO+X04pZB0gVVVVTV0XzlW6PAyDBEMaBK+TPcoPIbJWhcHBqa7ul7kphmJVlV1omNr8eLFqkBUBZMFdduf3HrrFM3R7WiU8Zxta5rjIHliQhrm3nvv3bV0KVh8994WLVoEBZumkqqW3GwWZRScGMqCqPuzqVOn7hwxNl24UH3v6qtTyHskEmZC15OJmhpwURhkyERz04jZVzQ3uW7RZhHDA9+igIsimeJVV1ejhOSOtK+u62mG3nBFiZiM8aTnWX6+2ymNVspBGKbrOgBPCrpXukqeG0+45rjqasQN+i8eWj/7jbX9X/twW+8JHVk31QfIopMTNXW0SlNjhKBklIQqFIX5dGViwvon7LMQlyeZSLpJVhX6nZLOHmEj/H2QnxAt0f7nwS1UZnTyJa2Y37sm3pNsjYBYOKLdGvEMrSi04Pj7wavxDyIavOX3yRQBibr4zWF0NrAgLkP1wKa+E5aOqKw+HctOba5/8+DZk5787te+9OyUmdEN2Zx0QBeqg4O/rXV1NxVhEfSrGCJnVPS8iBOJeLZAxo6GzsMsolVSVhOYTzDcP1pFDB4CcgmUpCyVuCXUmHH8YrHYM1pzU6lUmo6FVIShkqyEKAFEeX2k5iawik0W5KsBKYLMlbmuuxsewnD7bty4MTJ27NjJISlELhdZXG82m20Ll2C5VyhMgZYbEoKdnbmuqVPH0c2IRXWWL5j4AmALKA8AiroQ1xtGCy7IFH/r9dk29+L1ophNIKnwd4xyySgr98e+gcWS1yv218U5yPfwc7ytbY9dV5fqi8fj7SNNahAqYWIbuHHkkZBwLEIjCl3i8fiukYwJ2JY8z6tFDVixrCjiFR0ZYL8PwhXIy2E34DlAgY1GLh8xqduRCLP7+nLxfN7qnjatsTMR4ax9pzvz54+sO/mt9e2f29qV+dSegqblTIshsvalkKBHguy6i1WlskJR0Z5ensi+GZCT3H/RSk8LtzAIwqEIMGjUPxZCqpeNhDAE0sjICS4eFL4jyPyXJ7lMd5TPS/BJ+OeDyQ/n1WOOozFFZ7oRY8mozsYkvL7JY6vfnndAy8tfOePw52fv0/xGoVTupkQjFUqsWKz6+/trI5EIxeSmaYKjgLoIhXhxYTTci0DVpkPPnHIB3A83BkYzCAPF4nQYApQHo5EIGxS9CSiv9+Xz/Y2NjZ2jdDu2wFMNdStKNCae72jldVAdtgxRTCcPAUjg+vr63UBejmQQGhsbp8kEpZgYdJ9wvalUCujJcrdjQZRVBgcHYy0tE+DGVF2z+N4T129uPfpTxy+aqWvqdJWDU4nsuXBC/PVIOJbAjYpxiZIZ+ZN+qttTPCIupQRxJRmo3+bo028Jv0U2AsjFzvYPiAFLqxcvFYq7xk8Yu+6bX/3ik0cfPe/Z19esUQ+bM2dYDgZ8C1hsibMRsY7hOdyknIPKTRPClNrP4vGwa1fB6swLBYXFYiCTiEYEi3TBdZ2YT3o5amcfBifCmkgkUmJF5pR4CfmHSDJpuKlUhC+kfn3GqhrA5882eJ73yzfezcz8ywubTlzXPnjK5k7zE71Zs66voDHHRhmuWA4t/Aq+L8oqbpZkQyJeQ3F/pcMVjkF8OjdJxCqUlWldlNiIoUC8EGVb+Uv8B0R/lrdBKGqTSkqZqMX/Rg6ePJKGY7ZJWqs4YJWhsnQiyupraje2TKh97bD9Wl7+9oXHvWMYDHG0xlhRK5QchS1cyBfPnk2cBKB6w7MGqAlJPVReWclHcYpyXYBjqbyIyrJxsVhE7IxJURThp2yL52APH+3ZesWij2CMRPiA78K7sgMymUzy/0AyIDACYsUmr9VLp0dL8dKDkaGx1C4ROTQOgtaRNsGQBMQjFkEKqWR/Bs6jvb29YvFGhjLZ2dnPpk+f1PaLux4+5cE/PHFle1fv7P4CztRlHqjCfEzIkKSXOM+KVUv8R34u+JMPICkThgrHIVzMDr5D7Eun6686Qciu6mP5+raDXn1v6xnHzdv/2t//ZvGtwzEyh9iXgDiMqqrqGZzUcOhbTcvSuaa5Sys9iwqlKC8WowQMwWE1DfEmk9x5wmsZdUObKx5CJBpxucl1HBuNKyAQWF6JVpNwXZIK8zzv94XBwj5/em7nJ1Z+2HPkxnZ28J4+bXpvgan5Qh79OcIeeJAh8rHSMlEX3LnyCl2e4+HVX1Ko733rK5ia/LsSjJfyoxKeiF8i9BcCckFo8UEagNBYYD4Ef7zOHRaNGKw2FWF1ab11XG3N+1Mn1r96yokHrjl+/kzU1wHwSpRKg7Ft2zJVMJyCRNRly5f7EIchz9bPrvOYqTJPQ5uij0NBTiAGtuPwFQx9NqiOCWg8uBCI6EdokbBEIjGqQRAU6ajv6xLejnGD98CGFD7HynvL3cHBQQl+k+xfQf7L6+8fTe5Q4nMigh80oHVH+FJTU+ONtC+6fsGWJP5G1b0wxqipqSncxOjxbdt2Hjh5cnP+su/fes6yp1+5pqOnnyIMtKJLbqYA0RiUxuRVlue6HEXBqhX6iHy/4mcx+MrvyeyyCH2HAnj8VcjzFM1ljqPWpxP8a184+ts//fG373/++VWR446bK+NADxyOULWNj4nHnQEnJbT16GHKMile8XiccAawkU0dHbxTUQCicnGTMplM0rbtuBuJOBEfX05QUjyYQqHAamtru4fe/Pb2dt7U1ITSZ7KqqioqE0clpaSzAkOZlspLqVQKLhpnQFeMo/2QV1Ha2mxXSfPEzGbDiMfrghT8k8/vnvjSB21H7tyTP6q1qzCts9+dlLOc5GDJYgUTHI42MTcJtx+0YqAS8LnGSJ0IxMXSvUf3H0m5i9ieDdeMJcIJCLkTx7r4iL/5eQxa9pFPUBio1ehBqiwa0VhSc1hMV1g6neqrT0XXTmlKb5k1rXHT3NnjVx47fwYAX13JZLSYzRarW1s7qvP5gldbW6MYhgtRElpx83nGGhoSeD7h9COX9zmVSqF2jmdL+H/GCh5jMUKcinCwH6sfqP2oCkZ82O1eEyP2LpQOSdRETCRqVS0Wi0g2FseNGxe0Ibe1tcGjQN6J8hu1tbVAFBI2R94S8R2omAGRm6Ny/bhxdDNFyZKjjF5XV1cHbzzM9oXSJcYmGgAHBwcLgsGoojSIyVsoFGrRJo3rBSYCBlOqNSFcyOVyJVFWDW+AFcBzqYMYkqjuOUh6Q9NRwA3Qsu2MGwdic7IwjC298d7z7/z947/b01tweMSAPVCRefaZ+0S2ObzKVDjaktevnOTy28yG8RwE1FQuNWW+v3BMLAekNC3SVBATBbmmiqI5bqmgzJnY0L767WXn9/b2gsMhg3gfsSBqwsgeCyaZkcoxSKjMkVaaW1yxFRLbJFrtYjK5rWEESizEdNlsdp+QxJZcGcj69/T0bB2FgRrAF4JM+0QznqtpBAID4QwmxIaR3F3P85A/mbpnT7Hh8Re2N25q65m7syM3ta03NzGbL07pz7txy9W1vKewEhTcKWeA8r3I8ON0HUEXHzysCifJNy6OQCjo0DjAPj7ORjUiVD1UkN6wHRbRGYsrJmtIR8yormyprqnuqk0l3poxvmbzJz/R1HrKifu+h2q28KiIFGTXrl0Rx1F1w9ATeW5vnzrOz1kN93xQajM8w2FVzOGD8LIgHeC72slkcusoJXI8n5lYGUXDD+UIoAYWYRGlo6Njy0gl8t27dydSqRTARWSQRfJPTmBl1apVG0cCJmHfmprERFbyvVG5iovzsKurqzcOJzYkTrAqm82Ol7E+ATVEnI/vGa2MOjAwUKcoyjgdXlKo7QDnjB6kWCwGIzxiLwNoC+R1ao7jjd1v3pd+uqd70FMiEe46Nrr4A2MgTrY88eV7QcAq+czEXB9S4fG9WeEKVIQIIaMhvILg5gTHFsQjOFRgLxwyWKjl7+zsav7jw09/4pyzT17Z0dHrVVdrcM9V3dBL6LIcGBigmFM00wytYlC7tHT/Xc21NAWE3jTgWL1gJRwhP6EEyC6/78IKkc+CpWa4B07f09PTAysv37NU1YVOJ+2HzjaEKCIu9V1Az+PLljPljOXL2Xvv7Y43NEd5Q7W686KzZiLEeBkeMBFnFqyax//Vs6uvt2/25p7s2GKhOLu/aE/c1tbLEonoJ7r6C3pv/yDTjUidHomkLVdhJbPIbLvEPGBhQBaLye6pLGbEeUTXWKHgdDLm5uJxgzek4p5VLK2pbjByU8dWWdw0V39i36b4pIaajw46aPpH6TSBn/qHJnh37Ng9Bg9eVVWu64aSTNY5Cgj0FKUUd8jSjHiPhQKy6xWoa9BTFdKal0aSns8wEvN8zZo1yqRJk1AilOU8WhkjnErUQLTuFRbIc1B8LxHYo0BXUaB16bsWLFgwzL7+2uYLs7omwPCKvw8tEHDxWSQiVZYrwkX5nKG9gGqdqEBYoW5LeLZBe/Zw5xwon/uQaVn1kGNKJi/38h78lEhR4hIonNYuvOSnF7d25RuZYcAigceyIgcQ1DTlSk12IDSRA3depJN8dzKY/GHDEj6X4AyDH8J5hvLnAjshrI5/SEqIeWYpz97/aK1+DjvZ5hxjNGswbrgGN8iig8VWYqekctPQExlCRBEIuYxE7SbZbARyzf8SzpGYqojFRtpXoNqInFaEFL6wLecMLaJ75TM49xZipTnjDG/abwcBnjGLRcUeHBxMZEy4uSUHFPvVMXX7GZ9rRKmNYKqGhv4BlVmWw4qWB6M3iTEzkc/beu+gm+rP5o2evrwymC8x23JVx2JmRNWYEVG9cfVpJ5VSrZaJNcDtY6IDJYfvgGYHBi0GHbL00L1ApUZWXPj8+Ys1QKCXz17jeUuWwA+msIky8REDrFcq03Ug8gKNjZFule+q4xZEJQiO7nE4aYi6+5BnC+CaJDQJckly0uA7BKtX5UAI54+8IJQQ+BtyoyhkXLNmzV4GQbapINwslYB3IpFiNbRYWFByhlEYeo1DJkcg2CppB4QHO9p9oo+j2oCJLa9beBpAKcrjDHdsGAMZa9JxtU1btp6TzeU8rqkKaQrI5ZgmdcWCLr5lSBZ6qPsfqiQE+4UzVcLTKEcgoerkcHmDikuQZgl/05gaMdhAfwZuYyeasjiPY9CgtIhWGpex1Kg3Efx0qGGHOBQoxwAr29PTM6JQ5pw5czBYYYG4FUpqYrAhTxGm3hpu0zTNHyTlFm+UsZDcQnvtiA8fwTJEMD3dc5KGYSapXKXJa9BAULp2znKO3qLly9d4pg0q/cUYmNb2rq58TSSiqSm1a3w8WRrPkpKbknDEnHOf736Ybe3anUxLeRovek1Y2xOaquMxGkYkp6oRWwDSKHHmK0qLbckSWSbzkJWne1QVtXmxSO9XVVWNKkbqE/KaGlCiAiEbJlIdcQONWDabJWMr3wtR+I9EBOvJSR2LxdD3ECqW+DX/kY4nPQTkgaJRU5WIVrk/sBJeJAB1DLtVVVWhWgB2MWuoYSoUCqTtOPLxiVFZnqusRNC/yIuw0TdADgqi7KpoHZ3dLQzAJwi8lK9/GB8jlEEML9shD6Li00PfD1G7h21MSLMi+MrQpYaOOyS0QNJMM+AzAIUFemnJqkw30uIqeOhGHThYgeDKhhJXdCNdt2iFJBJH2oZa8jBJ+t4fFu5efX29M1AsmgA8RYCY82NMKn0N00W31xZhEe5xA7o6yHWIY0UY56Vh6c89b4nD2BIOpSvHcWy3qEeLrBgtlYDJyCieYbiqbRdCQRlbvBjI1ICYhdXV6VD6iRn1Bpp9nJiiQOlLIaCXbhXhsgs0697nG4mwjGmqyYJX8qJFZg6UiPXasiwOYaDRNuA4QG9ucBOgMiQ2MGiFUs6oz0ckGv3PCUNAg0m41CN6f01NTYC8S5lCB8nH8HPGYjB0n/D0R0eBZ+goQ9CfyPAbZOxxD0ccjwhvBUEKVSCk14pDj2YM/OP73I5CyMg/sNg3BM8edsvl4KF6NsYji0a51tE96DEVykJyRReJgMpZW07yBiu/P48ChyBcs5L1bvqDCBBCjoD8SrkPAfPCGABKIA41R6HwA/s5DosaUXb00Z+CQdCg9SjZo4nQghk8WZUcFkoavpHiQWMfG8q5pomVJMLr61MBHHRo6/UmtkkZa44FQsziUQ40IoUYyD4jHADJ5tAtFAOCrRc9DHRcVXWi4GYUHxvOIND54xwGBgb8AW6aWEYUIPN8TETQJDNSPE61dzR7uaWSCQOC/xvMD63ILwzngDk8C1/EFFsymSSBERivSCRi0XOzrAjIddMpv18nPFrC5w50XRqBZAw7xSAiTpwVOF9oEY6g9kXXHIvFgOfgjEW4pZikYSb7WTo7O4cqitPgkEl/KtWZXMVkFO6zuD2BrmJFiVD+3tnZCc1DzUCtPxIBdsW2FKIBJGJhebwQLDrQAUJDEZ5/AspIgo8A5VDg8QCsDZ1rMCxC3xf2XMJQfHRMoXcmDLir8F7E9QWrtPQ80bUbvr7hsDqKUlQ5NxRH10tR3fWUoBod6AvI9EHl734oIfDOZQGZgAWIAKjDPdegWCTyEHstvaHkoqxmBWVMYXQqhhgRmtrMKaqTxiY3nnPGp9/d3bq7Cq6W1ImQhgGJGpF02WvlxHvh+MpXdVL0SARyr9SAQrVbrH5DbqT3Pns/8ArgDoZWHNnsNdzxggeGVQpVEB/2asgJL+vDQx8anT/2B5W9bKmVxwdwCqhIeBlDjhMcF69EIgGIbpIacyJ+eUG40bheulb5Gi4BRfqDfh9JBAYItXsMVoC/xDnutR8SaOIe4/OG8ISIrBQNPdImDzNQ8TtwILRSDsEVUH4AnXzDnG/wO5iD/EfjV3PkmJC12aHHlb/btg30HyC1ABIFnoTM3C9fvnyv/WnIco4uQ6INFOdN5+OXMiLwHHAeQck7NO6kF0P3JMT8FST6AD6qPF7wrOhchOKTjx8X3gXxOeo6vIsRsRHYV1VV8tj8ihdXeKL5aC/nBDkFkQsQWJMKtWL5EZk88zt45XMIogP6qQw69P9M+Huf8Uo+bvoqwUEo3xmaNgiWLVfsSySmCrNtpbk2Wrrhh5ec/dXzPvv4iy++aCxYsIBWX9SoEcuhZXRwEEpjBVCx0w2Tk04o34LivbepqUnSWvv16o4OSaaZBg4B+yHbLTskc7mchTbXVCoVFsUNLDbq1jU1NSlN0xKI3/BZv1076alq3ojH42Z7e3sOGWlIaonkpByoQMrVoN4sv9i2bSuRSKAOjevKVVVVoZyJPAUaykg3E59ramoCPgL0d4ifgeX3b14+z3kigWOhRRa5C2/cuHHkFAjabmD7MVFRmyejgvOFJyDcVyQGs6LGTd/Z2toqM9MA+ehOJJJwMhlHIvWIm1JRtIRPsQ+8Bu0LGnOEaI7jgD8QtfpIT09PHKtfLpdzIK8n0YKxhgYn6RPUKK2trZ6s3Mi6PmZZPp+PoyUeQCOEF0J8lfQ/+/qSmWJxkyuRejh/jA0cW1WrE9XVaA8ll9mJ+zS7HoxdTU2NJaD0NBpxvPAMApYgBCEO/oaQoLpazbS29mSHJC2peW/y5Mlok6+Nx+NqoYBKQh65CrQ1AwyF+wT8QkngFwj/gA33C+NZVdVqjNl8Pu/GUFgiGAFzMS7q6uqQRyvhuQCIJD3UtjamTJjAdAHzlqEIbcAlIF8VjRb6GKsxWxnzcIc1Hwko3HACrwgPIRS6+z/7wBXFc7CcMEWLQ1aPIKmuY1ODjYRn+x8PhOcEUZgvfVZuxBE2gNR9JS+AdMMApcH5uAyVVeru1YGeVllUV9jEsdUffXvR6bd99dzPrluyZAnKTiOqTmezDnFEgma6yIpOtAT8tqEUi5Y7Eq+eP1By1Z5pptVotOS6btBCidXUdd3caEm4fD7fAJ46TdNAaWVBXDbCisw0bQ1xc3Nz84j7FgqFJl3Xo5bFFE0jI4ZBDGCMkUgkCuFW1WEaYFp03QOWAfkC371G/Ok4rMqqyvCxw9ftt23bhkEFXj0yRAiF0JcfqYo45qBpNDU19Y/Snp4cHCwl0WhP6U6U2VzXTrm6kjEz6HJtGylJPjAwEAUFuGrblpFKOZo/KTWAgJyBAZNXV4+opJzNZtGcVE1NboUCFi88F/ImYMMmTuSj7ou2bv8314KUrWcYmDApcGQgqTjSvoVCAe3GIFKgkABw5Gg0ivum9/SY5pQpU/YCrMltcHBwUrFYNHQd3ixY+qhSjedGnAaNjY25Ee4xjHIzPo88gxqxObxZkRh3u7u7BxsaGkaaAxhDKQfMUsJdwvNVHQfXokejNXsq2p+DjgSRB6hAGMuuOj/Od7ldUg6c1fLGvIP2fzdXMG1VVWwQWjKF6VbJMlGeB9IJk98P2lzUUUjQiOqyvtNB+HYFiR76E1YMhFzUGO4j3RUhpYhRrjKAhTTXY3m4nzOnTV79ja+f9i4eRltbW2qYspOMlyhRI5hxyEWOeBHb1T2TuyAg8kat7dJNi0YLMoMrqNZQLkTsPmqiBlZcrFbkIqIRhnwkz8PEIy9ypGMXCgUikInHSeUKmWufHsTz3N27dyuj1e0VRclpWgS6FjykGUjeTTfrDre5VuyLVUh6UKIM5aGBK9+bx/H14ZJw4XNAmU30LlP8CiaEiOGBgpwPU3sP9kWS0zCMoldiYL+hjg3oTYYaf/YKoUK1dxhK03ANm2sudOZokYF2Il6jPR94PrKZKZUynCKebakErATaoK3RcjF4PiAyEfSA1Bsj6ccE+m+kZwugFQAODnKHyDsBN5BQVQCnw9ogI2EySB4A3iooGlXQXoiSqCT4GSlPIK4L4wOMLr4EI3gLfB2+irGslWkTg1OS3yN+k948d1WFK/W1qZV33n7lL2RdVtSoEVNviccjgDfSrtAAdEGKwzyWz5cwqIDJDNc8MWCla7w1Ho+CSaPiFPCAC0UT3w9WJ1h8XEBVV1dXrWBJApJwpIdfFsowPJsNXwcecRPuFWI6rFp0E0NAkVH3hSsI9njJj+crbJNBkcK1Ix4WDz0a5YRRMFyXlwTBJ1a+ZDI57L5yokcCFh6stFL/nXIW6Ijbi+dCbnDBEd54RKXgwxgjNTVeqa9PxqMjGkBwaIC4VJwHQgEMPjwn1MVH1XpE+IIVFvczyoh9CIuCNCxStXvYY0vlMMhEZ3Ow3VGZ9APj0Si3mMam7hC4jVZ4JEpJCoABvvwx14ubiEmJMERQ7dG1CxDUx0oG+Hkfron0GlYrS/kYuYEXX3yRz5s3TzZg4Tmju4lCAJGLwbWPtkjJ++mInBF5DsKoVTwfMSEFhiBslENIRJlUBFCskKWBnXt/zZqpDdXVMcwZ1y0NPvPM+5FCwRzWnb3vySf1o/fZP4rVBvQXqmprjqPiIcK71N59961ooVCSPHTyqGSpVm3YoDfV1SUy2WyD5upOLKZoqVQKgw7x4XDwU6oM4AckWzI9PYphJDzXcC2KTvxeBEwSSFwFRZLhsvI4P2GEKIEjoKReuBpQycDlbzAEgFED7YZEHJJp4iGIxOGID0+cBx5U3i0yRVWinAseByWVStE9WrLkRXX+/MVMMBnLZJONDPnAgFlQFMK6K9FoFU+lDNkJONwWjoGxkhCEmgxfoSBbdMOf3es+gU1MogEpiearEhMCDq7/8uXLwxOM9hfPJ0gEismooHpqlSzyFAcHTV5TUxPU1UObImnrqqurMX6R71BKWG113RZiqKN6CDDQaD4DlqQ3l0OZR0lH0tzRVNW2yWDLxDR4QNjy5ctEa6i/UIhqkgQuyesaVbdE7Cvp/mVFiFZo8Z176XUD3wHG6g8//FBtadkn6nklL1Fby5MRZupMR1inwzCtX7/emTVrFs4B44KPGbPWk8nP8HMOHQ8GhKpcQ40fjzXO9woyCSiz/JJgw1dkFbsotqG42uH7tty04tnfXrdt266WMWOqkXSjpFk6nQZabVjT3NnZmYCCMyYwTgarmMi0S3ksEEhWGJPly5crUCM9qOOgurpoeryqRvNyVYA7LgdRVVUVcNqwfHsNgrVr19al0+l64MhN0yQ2aMMwIq6rRQCFndLcDHy4F4bAStYnMEG7rlsvWk0J5CLcSYItj6ZE1dra2aiqiupEXEe3rCpP12Pc4gXg+6ycNzB58pjdw7nvFDL09U3ksRjINohVCvcLXZLZrMm3bNnZeeihs4PifSoZY4OZPO4han9IkkHBB+dVkqLw/cX+MZmuTEM2q2ydNau5d7iQ4Ze/fDpy5pkHNbuIo/GEOD0fDQK7cOsffHB929Kllfh9ef59fX1pW1XHoC0Yk8y2VTMeV4u5nBPBcJg0adKIVHPQIlAUBSQ26kChYLhF15k0qRF5DtLKqK5Kbe4fyABdBl8c54w8xiAiHEDuwWvQ05MxEwkSWsUCUWRFxl2DmpsCPgR5vSDHPeOMM9jOnTunNjc3w/ME3VlGOFdTxHHhiW5PpxN2NpsPl2MpNEB/BToqxZiQHiByAIbjOF3gJhhhXKAdf4YgTYX7j3O2Q+Ed2MV7MO7POKOStVrT0E3pThSeMi8NDnYX+vKWrTlGVU1NUo/H8Vx3plJxK5utmEbqO++8E5kxY8Yk2Y8RqmaR2yAIUhAa+ycZa5rvFagPQWAByA6Emo6Cy1FsQ2XawftM+umr/3rg1tbWXRMaGqqxGtIqnUqlwFw0UmIjmslkpuiubpVYiTrSwquGkMeyRkoMMsamdfbv0hU1UkzqSTqGZVkQgUUnF7T2ZAwlwxj6VUyUenEsKWiJnDKahHDnaLAmEzEzly/uxSItKOmJkVl8P1pPLZhqHo+3GYbGSiUL3yetsVzRakNeBOC9ETHw8HdcJwgtAoHN8JbJZPYTZUhq6S24bnJgIBedNKkZE2U1zv3hPz9/6Cuvvj+3u6/3iM49e6oUPTLFdOwYZ9yyS+bm6nQ60zymbuOc2S3vfeMrX9wSqyGKsQ/QFScmBQRl0Mgo75cmCGww+SRsWxP3CAZoQJYCh5Zw0dAzblzNRMYikvQmL16Ey0eye8gl+tKAqmru6exEvQwsQqyxsRHH6t+wekPz355+5ZgtW3cdsnXHjrTH9MkOZ7WuU2K65w1GY5HN48aNW3f4Qfuv+dpXT32X6TomYHLTph1Wc3N9UbGUKJqYqqurtw5zXFeIwcx98cW3Pvn4U6+M296+Z/+uPbvrmKpMt1xPixhxxzXN7TXVqb4JjfUfHLDfrDcWLfriy4rCezA9tm1rn1Vfn5J5A3LbqSTrh4V7RiPOEfylkgOBJOFl6be7O985ZcrYgF08nUqwFS++M/fJZ187YGvrrk907tm9TyabndafLyZz2aLnMg0T1tUNxautSalxjW+vrW1or6+qfuWIwz/x7nnnfeZDQ1c7AS/avn3HTHxnVVUVzhFVs8ArGNqIx2NNC7wCtbJXZv8rtpCHcPCsyTe9+q8Hf7Fhw/Z4c3NdHiUi9JCbppmvqalBTEkIW1HeIjcGO+fz+UQmk3ESiYTH83nFjcUAEaYYNZdLZMaMESWt/vf1ydUH0KC5b90th2W93nH7xQ/Kf3ry6S+xQVbM8AzhvbGCyZLLjh07lIkTJ+a/9o2bH/1w855PMMV0XOYqjml5Dlpzfe5+Uf2w0S+EpDTTdIOKGYl4rP9Hl5+58Pjjj9re2dkJfjatuloz8nkFjfnkRuN4+TySbo7Z3p7xZs5s6r3kspsue/ujTd92rKLjUTbVv1mOp/j9SqAaA3AMTGMqJVs9q1Bik8eP4bfddvnCxvr6rYClot0a5U5d1zE4qlBGc2Kx0raPtrL99mshpqpXXlnVcu/vn5q/blPrCZ09PVP7MiWWLSKPqwgGc7HpGlWLDMVl1VGV1aST/Z+YNf39ExYc+uvvXnHeXwYyOWLR+esTry946vm378qbJOOlWMWch0ulShJTfVCJojKu6UxTVJe7tjp9wtjXH/3Tzee++OKLkQULFrD167v1ffapN3/2i9+d9JfHX7w1X8i7TDFwS/3VxPOYqvuldY7ucZ9h1VE9VZ03d9o/77j1+99mjGHlS993358P/8tTr39hc1v7YX39/fGBQkkA10C8gu9CisnBcskiusHSsQirr67qPGDO9Oe+t+jMhw86bM6KZ555hs2YcWBVY5WmWbpOCeGuri73/ffN4vnnHWi++yYk4R7+ykdrN53e2tE9baDIqX2cOwXmAlmo6KAIYZhrzLFYTPNYTSrBGsc07th/n8nLfnTZV++bOnt8x5o1rYkxYxJoHUZopKGWLyQDCslkMmjD7wOFTHc3GVAk/np6elA6lA1XlMRFibmvry82bdq0nuqqZN/jy1844JEnXvjC+g3bTt7dveeAgbyl9hdcZoJsFmMZD0Yz/POk++Mga0JVOa7FWDKeYtVxhdWlE+0zJjX++8RjP/XYV7/62TdQXFm9enNtY2OTi4IQEquYf5i3yJHJ0E+r6M4I6owhQ1CRb6ROQ1fX1T02tEuHtFEODOQbotFoES4RtOwzmQyBHoQM3IgcdxDGyGa1ZKbU40yuO8DctmNHy20ffvOWfn3rfjY32dvmcvbX9fc++cMT/3L+JF5dgXcdLBZnR6MpeBxbprY0/uapl9//x57Obv+h0mNBTUmgWMMoyzK9HAJhxm9/6P7jjz/qa3nTGxuLqSpjyTWJxMilq1deWXXY3154+ye7d+6OEiSDBrwIyeS/EnwGB4M4oPLMiOhs//2m3Tl7+vTe3bt3T08mk4OFQqGEej6MTjqdXh8+ztNPv3rgnfc9cu3763ac1J034WuiOIeEDjQNQTkF1ilhzcHeDFCJB+ZjrzNTUjoH8tUbdnYueOnd9QsOPW7Rv7969meWTp8+fZVZXDnh3Q3bW/JI3YCLhPpYJHFKKKzEz7g2CMconEpq6fS42v7+/kgyaUNxeNM/V6zKrm3taMlls/6kCiOLJRcH7gGVkzHZXDZpQu3McWNr7RtuvK/5T4/945aPNrYf0lcUJCxQvlIxOQOMKzoFaATiEyWHeV2ZgtI1UBizceeec19ZtXbhYQfse/efH1pyDee8c8WK1ckDDxx/eD5fQk1++/TprP/5l2+46guLrl3U2r4nUSjkGQMiWdVsYL485PiiOvFRCfSs66kKK9iOV+jNKe192yeu2d7+vddXffSNo+bO/vHDf/rZTdnc3umyrq4ulPf2QVafVn7LAtqSPInBwUGzvr5+2BZmXePsoUeePf6YUy7+/oVX3b5gz0BWy+Vy/jNRFIepmscUjSsRQhQLwBDdDrqfPBKnFkm8kcnnvEzOVdrau5vWbt5y7ivvfHTug4/+/ZXPfnrBz6695oIni5SqqzjnplgsViX7PERvsTjCEOegsuGZilg0ZlRFYfPnz6fMJlYXQSApYcCweo7nEbuyVXALppfz6+kiP8GHvqI8qpTcbLwxNUnJD7K6a9/5wi2t9pv7DQ702rnerDPYPeDszn106g1PfP5vCD8WL/YZinFs3fMK6YiSXbVuXeqHPzjvuYUnH3BzMmozxS2YimK5XPVcriku1zWX67qrGDr9rOh4j7mKrroAqL350bZjrl5y11cnjx/L+7NZqasXXJ98iVgy8p3r7v5ZR08uqiRSFjei4vvxUl1FQ2HTpe+n9zTFVVRmKSpnRx486+VlD930vdbWVqgJIQwI0FeaphUWL14hJfT4V79x/bWXXPXz15957YPPtvcP6KZjOxznrUU5U+AKMABBVPRjoQjigWrOI49MpSUZq7UR8bgRczoG8u4/31537JU33vfCZZf/5DstExq7IrqO+2MDWeHfD5XuEdfwr+rfN5W5GNqKobm6odEsSKdTaiymRg1YN8Y0VFV0w3C5EbHpe3C9+I7ghd/xPZ5/HxTXrUlF+y/97s+uuvG2+1asfHv1Ib3ZIvQaHSLkA8DFF61QPddVcF0IfWm84vo404Cd5JoCn91p6+qJPvnSW5cfefw3Xl7/zof7LFgwJ9/R0a02NY0d+PDDDbMOP/6Cfzz095WXr9u2O1FwmM0jcZdTp7GHMYzyt0rf74Lwhejh/XuoaBrTDIVr3C2WSvb6ba3JPz/70k+P+vTFzwIDgXuBnJMcIwhOUfoVLM4gHyaOAd3/3VmxYoWcM5zNny+f89gvfPmaB76/5M7n/vbCO8dt3d2p5YolG4xePBL1mAb5G4RfLnAWvqAwAft9SD6+C3IKxFfIwCLuaoj/uaF7JtOd9p5+9+X31h1x2/2PPvHZ03/waLaTFN7Z3EWL4LqxeDxOyUZBNOv6IUMQdVe2PvvWIDAYtqHY2kHTJ/343dceuca06KbJie4JooU6KC352gvU6OGTgGhmtiHRMCLp5e6+3ZNrtdqUwoyx33zh6F91mB/tozkGfFmNCHsg9+3ZVrI+pR874evnLzrg6ofQ2Qe29EJf3wRTUZB4s9ra2rQZM2akTvrcN+969uV3DudGzCWpdh/+HgZYBKugAGN6nplnzQ1Vg6889eAFk6fWbVy1qnvrwQfvBVxSdE1xF579g8V/feGdJUUHop+uRh4IfW35e4OVFoQumAn5fm9Gy/jMayv+enhdirXu7OmpTSPZ6LPmgp7e3LMnYyGO9Dyv8ZiTL/nNqnVbPjM4OMi4EQWiC4M2ZJ2HrOISWh40qovTEChYAfZyPMtU4xpjh+43e+tb67dNzgGCLEq8ATBMPPdya4oKumf10BnjX359xUMLd3V2NlbF42h0itTX12/74nlXfvLfr65+diCfdyHVJfkz5YLii8/I70IfhMnmtDTm2joHEgMDGcYNnQgZfGRs6Poqfq/Macm6mH/5sBWOzaySPnva5Lbf3rX0hHkHzdh28y//cNZ9Dz3+q63bdyaYnrC5qsGYD1n5fGh8mW/Dq3ziskdH9LN4jmWrSkSfM3Xi6x+88sCpnCO3UJaDj0ajUD9DQhgeH7pQic4f3Z4hOXgCbP7iF3/67G8feeqXa3d0TTZJ00BzEadVPD+Z1A+S+0O7gocpWNEFUV+aRA074AvXtKi675TJOy4+53NfXXTx514AcW0u97txkFGkwY1EfwCfp2MIN1uWHeUJCPeE/gGGJXRDZUIP3gConfEvjAFw0cikIZMadaMjk5J6jI+rHtdtJA3n0tdOvK298NE+Xl5zHNc3BuTxg50vonA3ajq9+T3zOFNl341XFPwFwPQ3NDQkoVTz5GO/un7W1EmdHuHtQ3DpoEdCviFYypF41TR3V1d/1XnfWvJNxozt11xzU8U5g4Yel3nX3X855IXX37m6WMhAhQkWeSh0o0xNJqo2rmU5DXVpZeGpx35/XK26duvWrfXAv2IPXfdp1rq6upzJk8fYG7dvn3rY0V9f8cr76z8zmC9ZPBLDtw0xBqHBKgeI6BoNP5vgJ7mmACSkaV7eZu6KVR+15EoFeuDljlMhNkdf6U9ImtTErVlgKgOYkDXYBVvRNKdoWRYMpqljqRaNl0Fvi+CyKN/8MiM00Gart7QnBgqmR9fnkbDlkFBj6DgZAocP8xcQTTXTmR6x1m7vnnDx93529y23/O60W+567L6t27sSXAM+xtHQA1vO44pjBZwdvpGp7NMK/SaZKlVDd7hrrd687bD5x1/0hOd5geAL+ivEQih5DCh5GImgBVo2N83XFIU5Z5139TdvvuuRv7+/oXWy6bg218H6CU+lLF5TbiqWAMGyLZOmVgyyUEOhZMwWejj+81W5HlNtzu33t26f+NP7lz196eW3noxKxubNu4jRSSbOEdz5wYEM1vYyOJVEqpQUHcYq+Rl4oZaLJgnoyvn1WmKr2GsHWbryn0LTd/516iM7i+/up5iGTUuqJGcRjEkOpA5drkZiyW1eqMNYZnlhkdHg1NeXz2matv6Hl3/l+2PrqlzPsYnUPLh/e12TP6E9DuxXxH5/487jvvW9W778r3/+qkTiqeI8BeOv9quH/n5XR3/J4HoEsG1/1MsbR89HPBCfmRrxsg2s6tFHznv4tpuv+M2DDz6Vbm5ujrBYTIHeH74/l7MjEydOjGUymX1O/9LVz72xesNMCzGAwnUx2MtGuzxvQ4PCHyLiSYYGU2VpnIYMvg/UzZrqzwwhC+d/oNy0VibUlceGk0Jv2BZqGZAodqj3QfEs3GO+98SteOpi1ZJeCAhnaTXBrmUR1yEDxH8RZidE/16xSMqbQXkKHRxva7bsnP+T3/ztke7BnKZAjwMhSIXnWz6XcjhM46scPpfHaGVazS+y6K5nW++u23j42V/+wVLNv5f4mINSbSyGSNbHZqBSRGkfz/MOPvgihbGV9mlfuPSKp198866de7ocxUCXKoUu/iUED484a0PzsdIr8C9DPvGy8ZL3d6gnCUOMdCRXVWdbe7ux/JmVf/vB1XedtP/+0+2BfF4zDL/BDtLaweHKLcfBMkH9RxVPeAQ8FSZmLpcH7x0IsiKhDrPhQTg0HCinEL3obyfcvz339v68oNpI/lNHpJxXKikTuQ4ztVixse/I6Wf/kc4WvfriuPJe4JjjxtVYq1dv0s49+5R/fO7T8++KGlHEoVAWqXQ7Zet1cM+QnzPUTMFyn1n5/k0ffLhjGmME/KHvNXTN+cLZP7hm886euUw1bM/nHCvflzCnpDQ4KCuYWW2/6c2b//zAjZcUCiX+mc/M4yVGXXRuwXNtDBZNczCx4id+/tI7P9zcOpVruo2cccXKvdftk25+eeIEa500pCJcoI0o0mWSk/pL/KoIVQD8QRUmy60YhmKQuR65u72a5jpRRTHUGDUIeSXXs2FPfZRWaPEIvcRoDd7YC1k3NBSqGP98mPPZq8wv/vX7Gfr6e13ObM/FLCdLhtkmPuPu7VEFBo3LwVn+Tn8KlA0e3T1V17P5vP3ah5svfeSRFYfiL2vWdAYrLcJA4CKA/4xGFWPHjj2xDz/8rXXBoqXfePH9rbf0Zk2HGzGIdhJloeADCD3csj0cavz9ZyWeW9Ad7I8H+XPloh362XNUrire7p5u4/Fnnv/Lmg82H9AyYUKpt5daVSBU7t9gub5UNDHTwUKU3oLWGcqjoSPwhcuYqlhKpC7VaDREGnimlIEiEzD5ZhjyG+whyJujasK74InjHtxZeO9IJ8NtpP/oYSi+u0MLgwt4qqUkqxoLRzefe9bByTm7F3uL8YTpJAIeAJ9cAgAkdcKEMby1tbXh3ruuXnLgjKaViudoQkk0dBJDxp3P+ERwjG1tO6suuvTqX2LQzp27SFMU7tx9z7L5r76z+pp8LudwCK7KyRa4a5Uv8hhLOW/8mGr38m+cs4hzTgzPEMQVoBZHd2pKXV3Z0tixYwfOOv+qH7310eZPMEW1wXUZtMcPZ05Dca0/yeQEkbOtfF6+mGL5xocfghz8YbtS4cwFDWii0uAfCBRoUdNgvOgnokAiEgzDoYco3+hyCOGHIRXOTdh0DDuOK04yKGDI6x16ziSbiXqlGCUkhBU06FW45JjgQZrMq6ABCCTtQseQhphW3Eic7+zNKg8+8sQ1GCtdXe2oqlGZBeA5qRWRz9vazJmT2T33LD/pX2+subs3U3A4AAQVS3sosJKAwGCo0sewRNoMsuCe43qO57+occi1AdXyi4UyXKs0mL6XI/7nugrXo87G7Tvji75z/e1IWmazvYCtm1BS2St6Cj+MoM9EIBijMSNqWc7kTZt2TIEg6vpdG+qWLfQaslap7m9dj8xc2fdcc2Nd4/r6+vqN6fSmbel0ems8Hh8oFoszCoVCS3+hv2V339aJnudNuOjpE//Ua639klpSbaoDBiQt4uGgkdItcU2tzp3WcsXXvzr3W++1de6Y/p3+qyb19vZOBJowGo3uTCaTG1KRyKZ0Or0FL6AXe3p6ADoa+Pujt18+tblh0LNs2FOflFEOrKG3jDwTR/Vcy3lvw/aTv/3dn333nXfujTmOO/0Xv/nTXZ1dXSpSEsg5VHoEoR/EREShN56Iq+csPPXeL595bDcALZABy+7Zk6yKVrXG4/HOhx++r3fGjPFtt/7ijwe88MZHZ9pYaWnmyUkenhV7T1SZbIIsJFdUlyH5xxWbcW5zRSEQkfQd9nLJh67IQVQlMCkhjnzhP8gdet5+bceGpJHc2N/ZCbHc4owZk8bapN3hx03B1AkWEukCl3ML5bwUxpdfkvSbloiiw6FmOlxXwC1ZcbMrbsPehqO87offg5RdsMpCZxelW3qFjsPk4lg2aOXzFY9dLgKgt3E9b/Wmbce/v2rdyQsWHFBHYzGV2hqNRtuAlnzvvfd2vPHGG0Di5u68/8/3te3a5aFdmMZQxbkPXaHE8RTcDxBoe4rqWpruFtV0PK401NUp9bV1SlTTFN0taJ5VVD3qnlSkiEowbyWlQTlUpjEMII79/sbWfb564dJz9913zpZ//vOfbaKXQSQh2CgbDXJQb6NAzJKG4XldudbszKY5Az954ZtXbsy9892BwmCNrqneo4nmJ346819fqZlyMFZF7u3yPLvKBn2RnbF62Pialux1KxbdssN643Qn79pMU3332JdzomMpCgPul6s8bX123LcuOXvOhS9t7t3cMN5oKJocEnOqBoRbLBbDXRg2kHn55ZdrjjjiiLbzzjr1mlvufvSOgWze4apWJnIQLnM5thQDRtV53nTcv/3zlWt+xX6w/HNnff8r67f3zmZGFEkiP1SQgyTw6st0cngoiqaon/rknOd+ev23f9naurOmvr4aJcaEkkpBPYdYfC666CKiZT/oyK/8pLMvAxCQ0HMc6hZUjHhxXH/dIuEEq4gyI2d6SrSJg1MO2Vj0HcJGKFDTUmgQhkMmOejk4BcKTr6qU+jIYhIoHvd0TTVt4s+Qcw0cG4iBK94Tg1FOKtFFGzgxZffYN0Ai/0aM3w5XjCiI/uR5YuVD5UIs5WGPbMgyJs9BGh6ZW5ALjf8Zh1EpUwEoHFUHX1UK36wQoAz3KRRyhT2GymdB2jQqd3f3DBgP//W5zx8wd9ZPhmkT9wxdtU/94uXXrd66czzXFNt1Ha1sOsUxhtpqWoAdx7OZmk6lWH0qurVlYvNrs2eOb9131rTBsbU1BdO0vDWbt8Y3btoxZfXaLft29A0e2ZVzVOZasDa48CA5TJuvpRMYB1RdACZ768M1l/Z3mPfEa/l2fyJK1zM0+CR8OcxmVParmJO38ol9Js7JXvnPc7693Xv1hkymn3lFzTVVj9nR7s99b92R/1q3ffvnZk2a1L7OWxebyCZ6A4W24viGWYNLnr/4tg/y/zjdyVkoBWnk9fjlVcrFgs2aR2An0tYxTRdc9NV5l720rXNj7dhYY8n0a77EDCO6tYZ1qnG2RxxxhLN27Zb6a77/lT9++OHqeY8///bZNgYEyhQy+RXMhpArjgKdZrjt3bnqA448/9kdHd3jLYTgig64YXlilmtV5ZWHK3DllBnTmnf/8293XJzN7sml03Hc5wqGngULliDWtC+/6lcXbd3dNwcsUJ4PPyt/5dAyXGgCoEPbs0qKpjK1vqbKbqirXt08YdKOXd0DLxqqYtTXpD61Y2fbxM6+wr69/QOqZxchhe16uLu+luzQPFXoWsrHK3v9aGGlVZzuT+h6vHzeEdLR0qKHvio83oPbVXFdrmfbyNHwupr0QENtYmfPQGZTQ3X15IGMObmju7u6UMgCnYeSnJ/WpO+V8Ud4JlV6O4FFkqukbXmKW1LjsThrGtPQrmp8E2Ihy3FbdnV0NQL5yvSIizy1XCIDhcGK2pqQBxDHsMwsW/Xe+4cGfxVlSNkf89e/vTj34itv/opluy4HFTaFHKHka0A+EpI2cExH8Vx1clPj1iM/degNv/v19/8BDB9Y/gUUnjQeT2fHAg9xXyIeG/jpTX+a9/u/P3fV2k3bP0dVMCNGLOphT0pKKviG3+PMiDlb9wzEL77q+gs1lV+n+QNcTMbQ/axgUPavEr13zHWhBcF6U9FUpKO7/+DW7AeLizznamqMeTGipmde3rO6jI0H3/Leuc/lct5J8TjLb23fkGhpmrX+B89++YoN+RfPsgumrXBV87C4i3uLKF8hcXLLjUdr1X1TJ335O/N+9Pim3WtmjE1MQIkLtV2ithbJm3KNpuJUA0JXPn78eLZr164xy/5wy/UHHnnuQe+v374P6r0eBdrhByxXffE2hzAl8z7Y2DqL0HWEspY6EfKQlWadbrGZ9xrq0vzbF3zhMogM7diRH9vS0gh4KIwCEq209K1cuRTeRvqgI8+9aGBg0OOGIur3Ie8jmJihYwn4r2fb6ri66uKRhxz44FfPOu03J530ya2iH6GtKp0ovTOYw6Cpe+aZNw741W+Xf+nt9z44s3tg0GA6sBkinx0MdDkKRXUkGJflFdJ3EvZqyMOmRyIs4iuvhf340PipyEuU7zdR8riOMrGpvvXMzyz47aXfOPf58eOpSQ7eVLKzNeded8s9Zz7/yis/2rq7p4EpKu4Qhkg5pJGJUXHfgq7dkBvIVVRtXVaVSPH9p+/z129/beE/Fy48HohQNMbhgU946I9PH3z3A3/50rub2j9pOo7LFU/xp4afWPcX2gCUHboeRwGisKOrF0hdkPEE29KlawnY9vkzvntZR29BYzqUkwAyGlpCFf+R9M2u40RUVT384AP/9sITd3yLc95+3RVfHNvRkamtbkxbNdGoEY1GI8CWOsViqb29Xc8Xivyyy774RnU6edqXz/7RrY8//9J3ewZzKGdqga5qYJVFEg9hmqIoRcv2PtzQeo5luz/1R3/YA6tcIEIP158MKsV4LNOcas6/tfu5uaabTzs5au+mNA5WCddhutuv2O3Zdfte8uxRT6JE29I0c+2VL5x7yfr8izfYBdNhLilhlE/ST4x7llN0o8la9dBxZ37rhqPv/Mszm/4Yn1AzlToVQ+2bAHxgz9H7z/v7OWrA4koKv7jxu5c2j6speLZFOJ/ywAkJnYrffQuOqQOLJeOZ8EMMuWIiqw90jK4p6hdOXvDnb174+XfXrNk8vqWphvxhtENDldhxCCNM8fJtv3z4hB27ds9kHDlGaRaHAHKCpKHMFyBTpqj7tEzY8LMbrzr2qb/ccvHJJx/y3qV33FHcsXt31datbTPeXbth+vsbN9Zv271bP+mkeduf/sut9/7+V9d/cZ8pUz5iXFcY15BfCD3n0HUMXc2D8TEk4VzeIhifcLwq94KHH0oUhgWh/UtEAUD55H7TP9z+3vKFP7vxkuVjxijZzbt2qetbW2t37tyZGDsl2XH/Pd//5avPPXDcfjOmbvYAnsb4C4N0JJO3pOiX4URQlyfNJi8dUZzPnzz/K+++8ocvLlx4/L97ezvtXZ27mnfs2NHc1dXFzjvn5H+9/sIDXzzhyEMeiBgJephB5C0sZDlRJzOA9CZnqs76MqXEk8/8u1q2d8M7QJ2/d1dv84bW9i+aNrhpibtS3p0ht1HkLhDcK4p68AH7PPrCE3ecA2Pwy6efjoDR6qCDZsabGhvRQwEUZCmmaaWampr+MWPGgGnJQ+Na/2BWefIvP7vi2CMOuj0SSWgeFHmDuVxZIBD3iZT+2rv6p9x991+OFSg+aTWGjEfpyYS+iwM/zxgg9V61XpvDesc1ZDyouuPfI39saLyoOR3mhgO+8uiRD12/4vJvrO1f+XMwnDNHRThKZlcmfQE9cDzLTdTUqUeO+/L11x16y6/dxbZy0vRzSGQQSEQQnYAjXWbpUcmgKTMECi3fI7qkUgmZ/dymTTv0+fMPfPtLnz16aSqVUjyHguu95ngwesuRVqheFRrU4ZqQv3o6ClfUg/bbZ9U9t//wjl27do2dMGGMAmFZEIwqJoH8wdFA6M54LMpeef3thX3AxKsGMsWhXFAohJGn45fCXCRH95s2YcO6Nx4+5byF818rFEwNg++OSy91q+Lx/oaGantMVZU2tbExVZ9MJjo7+/VtbW3qyScfvnLd238+ct9pE99WFVVFjrqc9h+6qg478SlGNy2bL1y4sHyPUWUon2QQv4fLYeVF25ekJ8/Atvn0iY3tb/3rvrMYY5s/2rjDGiy6WP2S4+vrE+l0WgEa1XEO0seNq/3wxh9987zxddWEuCvzkA5zinQOoYnrmm5UdZUTjpn380d+86Pf5/IldUd3dz4SSWQTRqIIjkpN0/p37NiTe2vduvwTD19/WcvY5GtI2ovid/mZB0FUeOWEj6HZpsNZX9fgHJza2rVr+dq1c+j+3HjH7764qzcfZxSu+ckUH1oTqiuWyy8uwFPTxjdsfuXpey4FtRme7afGjKEQOZPJ6KxYpNGH9nhRxfPrhZ6H54L2fdbZ3af++cEbr5oyvv5VkJJR3iRY0MS9w+otQwjFcwezGfbme+tOFQah8vrKZV6RkRRf4kNQKQET3ZXZWHPUPp/+MK2NeV2rsVUFdGr0d1ytn5VWAYMoKm6XvemT7/Y+9mO7WPDbJ8JJXfpeBeKKbqSmSj2w7pTrf3jYTYuLXp6xpQxdgCh1RoktxuQqi0TAYIT3qG0ULniYMViy0eLfjKIQmQVe48c32O+uXVt1y4+/c9cRB+7zhKoqGpRKwxdfWSYLHlJFfC1LVfJ3Arf4+R8+ZcLYwau/ff5CDHBNi1O7M49GwaxjmIpJepODg4PEcJzN5ePrt+45HAx0NETEPQ/bg3CegmCzVpE31SWzNy+54kzOOVSV8EdbSJkF5V2pGgTAVjyuuc01NT3PP78K4crA73/6g9MaqyLtzKaB5QYTWEykMsBI5kR8rkw6PddlqsJBvkEVDEnMgtEmSonBwlOh2BXO09D32F5M5/yszx3+AJDrH23cWDt94hgvaRggsPVpxCPM9XUmVsGj4l889bDX958x9gkFJLvCHQnfs+Bcy8OK8CvMsdVJTfXblv32+ptLxBjGnHq/4xC9I1QWB3/DmKqo0pxOg9VLP2beJ26pjmkcmpaogoQyNyFvUpoHnIDGSrbNWtvbU8CBg5xk+fIznEQ86m3ctP0z2VwGqZtgZQ6UAUJeJjlnruUlYjo/cf6hNwjOTmL9fvLJJ1HCjGmaFgdrNuFiUgZa8YuZjGsVixPoWUiWcMG+bJ5y9ME/TxuexwihWX4UZesmHQGP5wsF1rpj12E+KCJk/XzYgbQIIRMvvIWSSb22WsYq4ITXXzj7hq9Ue1N2WkZBc0vcQcKW7KpovwC2HeU+s1ByuKty8gykIfBvCBmDWCqtHtGw8KfXHnX7zW9u/teYwcHBhsHBwfpYLAakY69lWQP4l5tmn67rnYZhdOWVfD+aTDzPQ/cdXjWe59XgZwhgikHSA1XeWCw2MGns7L521u48/Zfbvjlz4tg2xOFYdctGQYYNYes4NHQaEt/DGlglN50wlNNP/vQVn/3sUW2vvPKRHYup7QLn3hWJRHoURemzLKs/l8shhEj+5jd/ndczmG8E/oiQg35HQTg9GcrrUKnONQyNH7/gkAdOOGbf7W++ubYOOQi8QDQyMDAAdxVEGV3i1el5HrgC9piMdR133Fz3vfe2VR/0qX0KZ592/P2xKHEkiKRBOaYNzkC8J8FPvmHg3HG99Pr16+txj3dnMuh2VPp6Mv0qyS6GDdmQPGVwKYrHbFNtrk8VFn3lS09ns1ljcmNjxjTNwVKp1JdKpTrcVGoPN3lfd3d3Gq+1a9fWmpYz5lPzDvh3OpkI6XNUOjY+PqA8bkH7p0XjbO7+s58Fin7dunVNOG9Arj3P63Bdd7eqqjsjEa/LUpT+eDze29HRX33nL7+/rTqV2APyOQn9q/hPgGkIr54ey+YLju14+rIVKxJ9fX3VuXxx6vbdnfu6aI6HVQ2HZ+HNvzGItdUp4xt2/uLn330e49jzPBD0jF2yZMn46upqBYzQqVTKTafTdlW0CiQw5pgxidL06QwVvDGe52FM4Gfw0VX//MeXbm2oTuxEP75vvENhX2U6XmG2yQYy/bP9clGQSS0PwIrMe8h9KBZNp7oq1T8wmCWgDQbhvz/694Lfbrl8RY+3fYJmJxyueaofPvgrHDkolCwpg2DEaeCZuXoyqc5OHPfT78+76Zebdm+aPHXs/kTSAf0/TdN2pVKpEduQBwcH9wHjjKZpCA/sbDaLzCZouOA97Ein00NZcHFB+Z/f/shFN9/1pye7+/o8rqPcJ8EYQx5YaFTLQp8PHpHhBoALlnrsIbOeuWnJ1/7V2to67Ygj9gNZzLDqwh9++GFNY2Nj06btO47K5TMKU6hBjLzuoF5c1riVkxFYFLWpqWHwthsu/Wsul5s8e/YEp7e3F4SiFmi6sVrEYjHSdRxuy2azjZMnV4/ds6e77uYbvrny8X++vHnjjp5pUDxEjZtcWZpQlSu64NaRbiZ+SMdi6bEEc83n4yyZ/Ghj684uDc2XJcFxI8ZReXKG3WtI1WtqfV3V1ubm2g0w8sOdL1zggYGBKfB0GhsbQUITOeawubvv+cMTXn/epIbvvaAk5Z3lv0okHmX7zpr+OigJGhoawIRtJhKJdSM9nxXvvVc9btwBzY2NY7e17hkY68sOeKq0AxTm0twKaIREI5vKjIhPTT9r/Pia6upq48131h3Sn83V+8DlQNQg9JxlfwctoJSLaO/K1k75xMKXPWZDvQcLBfltpknhgaxHBM+G2j+BPVRRuqcMvYsHaVtFkCdqXZliHTgkwihtf2fxXGXVQddZ/0CfGvQylK3q0Bs7pEaqcN4/kOFof8YtWbxivnbsfsduWTz30W82x2cNmiyHmrjLRE6BAeYtOueo2YI6rv1uG84cR0tE1GnxT91w/dF337Gjd3P9mPiYkuNoILewQDktB8fQNuTQe0Q0IVV2EVfhFYvFqFMRn/MTPOWrwu/XXnXuM6edcMjDEUPVILwdoA73umllixqUbKSPo2iu53B1xrQpbX995PabBgYG6tPptAb9gRHOlQ+Ae4ix2MZNOyYjt1r+vlCJT8bb4hn4XRycTZ3QuLKmoWbj7t09ROuG65V9IuAwHK61XB4XeQxMrny+iAz+utnTp/wdjE8+vx+h10LJxHDwj//6iR4xDHKq6pKCsuNQs1M0qqKeP+TeBeFW2bsUHiEaGVhjY1NhaIu5H4EFeSCCvsNFhqR7f3+/Oe+g6ds1VeuqkBiX+ZUK100qbbkKOprG1dWhopBVFIU0DKFlMcK94gPbc4grdnKFb0VVrTLHDjbxctes71xLT0VhPkE103iJyoL6m2++reVzWY1YePzbLESRh3QzykNwhfVm8vHWPT0t2zv6J27d3T9+S3tf89bd/c07+zJTdvXnWnb2Z1t29WdadvUNtuzsGWjZ0TPYsq1rcNKW3f3NWzsGJmzd0zdp6+6eyTv29E7auqurOZMvRqlCFoSe5euRLzoHPcL6izKzXZFELGszyIEqB6Z/7QCnKN7KlSvpk0sWvOiAn2DG+Flbr533x+80V03vZrGCoqoA55BrEHgf4UEDehklGtUmKAfeeNuxf/jptu71VdXaGGiJRhTFishOsUKhQDE3fcOQXAHeG6K+RHEzfoaykKyVgyMxPFYRZz30k4djv/nVtb+bu+/0NZ7t+KHDcHmqAGQj8ygS9YYSY4GNqY07V3/3QpQYd/b15fPI/kJ8ZaS8xtq1r2NlyppFexz1GYXja5ngE3XiYFK5thdVVXbgfjPeQwys6wGMiGYUCFJF78iw90jiH9Cfz7kCk1SY2lzzRkxH85UNDfPQAAgtvX47hn8eLij16d2CZSlEbQ9CF5yHQz3qMtApZ6KD3KgciWXQPatOxn0IWvgchQJSSKGIYMAElwdMOh4vKAxlb3+hKtuu8nOpWNFch3Q8mmtqKohXhWbFsK/33tsEN8esSsV0v61dIB8rSIcrxwQhLRHT+X9Nu5oDj8azTHs/JFjQESW8gABlHuiWBDSnPppUFKxccDD4L8FJgchavrBA4DMKFy98Bp9lrqKKz4CHAtwfIhQNhZ7lcR1e6BlnAJyRhxA8xFD8FYSSooUyeLS4cWJklCcYc3d1bXUmN0996apDfnfm2Mi0Li9iqooqy1vlzLPfp+DYSsLQWmJH3v+bz/7z2vUd749NGk1SnkwxPAMrUAlEK76ozsgbehdMM6CZRp9AwJjLfdbeYbc1gw1WoVAoPHz3DUsmNDf2ezYAcaJQFm6CCYNdQt+GWnE8wpQvnnToz89ZePybW7duTdbVJT+Wfffhhx9GVrifG5FmKqmXo2vxxeUJGTSbea6qsRIbO6Z2HVxnHo0WRbWF+iI+ju4cm1BiKqVSpMKc+NS8gxTD0H1WnpDxpy181wLqOVSvCG7LTTMD3HvJdVWoVxVUsvrBXSqXAoPJGqxFQdxNSMrRN+KVBJdATFHAYA2Pws/Ry2aYcN9BOAEsY3txHgVu4wcsNuSVTJiA3OHwW1MTNCaY6aElOXxTgnAy3EtR/tFzHDC3YDFKEXsnY/192ULCEXZbOoJDH1KFTKMIJXy0pC9VRt6TL5MGBKX4nTj6yj8DeYm/uS4IgP33/PhOKecMyonQcqYo5JXC6DloNB6aRAvusAB7BCfr/zSUPl5a81Skvqejv9WZOX7Oqu/Ovv/saj6h01Jyqmczx+f8A78eZ0rEtYz6mDa96oj7f3nSnxaZ1xWUxtgUJeZ3SIIVBkSsUqOOMSlJNsLms+0GGAV5ctDn0/go6s9Llx4N6bCdH6xeX2uaOaPiRgX5E0msIro+h9TS4VJkBiyAUSJM1yORSCRgkh7puDNnnk1/yxdNKnv6RxwKpwjlcgTASovorLGpESeky6YuQedN+RYhsDLiBvwD7pHtr+q1TeMbkn73mF+HrkichoPUYBwRvRpOJqnG48QNAQ5pMDxT04HctyLEDE2BipUJKyHhia2Pe7aG65oQNEErPUG/aYUSacWh/u9Qg0bwEc9LRUkrnajH4cGNZjzHjp1D3bPwnfaW2qzcLZSPJ/izoVKokBQaC8V8AfBQPxQNOlPD3yNCcj+SkoumDF1Dr4rEduiCfY3hEPxefki+LxqBw/e+8oaVf4fBd208yMqYJgyCke2Uwb2GWzSCahE0EKqj4/iu7m3jZrfst+PSfe68OGGP6zIZmi64paDpgltWpFrXp8XmPXbPSU98neLfpSRKidUOD4AIVKHCjLZg9GiDlHWkhydEVKQMfBh8LkIjgaQfRs1G/Ktd9eN7rtvT2Rv30WxDEy+y9Oh7SYQDD2yFqhYt5v5j5aqL7rhz2cktEybkOzv7aX9oCI50znPmTMB5plVKVImEbvgZBc9Kxnx4ACA9NdjAYBZ5EUd3NYi7Rol4w++qKxUKw0piVGwwVqqjRBmzqnOZLMgifIIesXwFXYjDOgvk8vsGwabeE5dzm5a/oqjnBQmDIUC88hIuf/ZLzSOoXQcbni2LcpVo+y2L5NICfVBxj/wYvpyMrVzc/OHADR3GBJRhJFILDcSRj2nhC1QV4bSM+8M3IiRZUHF9nssKRShrs5KmUcklkohGdNElGnZdZMWmHIaKTlD6JhmvB/mW8OHlL0MndDjsDJczxbMIA+iEAfHfk4YD3oHJEgbKXSKWKUeAoZVQ/Cz7wTFjkAYzTUeZP39+OAFEPN1Q/6lLNtq7+lsTB02ft/biObd+vSE6aWdsrKbr1UyN1ST1Kca8v996zLJzr3r0FOOdd95BcodifbiGtm2DccmGUAnIQ6B/l8lkiBVYuJBDk3WBjBZevhvrltDjgISkQDfiMxVJJDzwaFT3jjvtkls3tXU2c192zfffw85v2LJWomH9e6RqvC+T8+783WNLC5lMdX19OgsPByzKvsEOzlcRfHrq7Nk6OASq4hHdJ/0nl3xvIhP5AIVF82ybsa7OflBd0Vu4T6WS7yX4kvDB8cIvenL4Vyg866VSUcO8Wruh1SvmAYoqw5HLScXQO8HAQqLdj4CB4fAPB1VEpsXjUUpjV+4tXdVhBq6/ClJrcPntynMX5+0iQkLp2fINfs6lQFxSgQqnN1hpw5NVeHycg5iReAkRXhqG4YyWqC4WSWm6QMQOkrEkZKCHvSZy8hVm2nhKoNunBqdEOhbdo9GMEajnkPc+4mJN1+M4FS+iB8K/lNABiM0lplr6lxhynfK/dmhf/2ewuPm/y/3Qo+mInwlG4sKFrK9vYH4G1BHljyGEMf4CKTKklLR1WCKqw/1qvO++h6C+SwxJ+Xxe6W5r654wYQKaL3hzdDJbsmaJdv2+N6xb07r6rUfabvqyza1ZVdqE16489OanejI9k7930gNw9M1MJgNNh+3DiLxgINCKlM1m5wgSSDFPyBeSKrlb58yZE5bZpluMZOPAwEBVqVSaCkQjfsd+u3Z1gbMre/X1d53zqwef/LwNGwd2D4GPr+T1Cz+kUC5BQj9d9MFwZ8vO3XUnn/n9JSuevufEyy57SPn5z7+ULuQL4/L5PPgZaImYN28eqUHVjx+PQfPRwED2PS1iHGK7oK8ewhcRHFM8DK64JcdWtm3bfhBj7JEnn3xMXbRoUThvANlvJF2mieP5SThxr5BoTSaTu1atWrV669at2sKFC60Vr6/5lAsvHGR1oNgSMfnQrm5ZWRFtvzjm9mee+buG48MrxESbNWPSuOdeemdIlUaEWWGtULnSKJyZxRLyD96GDdtbBgYGiOYen5D5EOQPqqqqMC4gdKtAMAKeFbhV/esaGumGuRjEMAANvMLV1i3tO486hm1bs2YNX7BggVtXVwecSjPCLtyvSCQC8SAX3ul+n5xYNHS19bjPX+aSDQp7TOXCYfB8xLOiGzeYKSAc2fW97/2q65ln7rBVQ1sbi0XYQMHC0ls2AAGUIbSqi4Oo3GNRPRJAnMW3+9RzQSlYhAlBLmOIAxzkOuSUUCpwGr4r7BtT1L0VTWeuHmd1tY07tIoQL3StFX5RyOUQbrjJogzIO1L7BWIQJBkyYyxOCwOSz548e7fK9FtszySp8p58W0Jz0ybSHiKL7N+SvcUx8Xe3r68PcthSBouUjMTfiYBlzpw5Q9ufCQaJ/XcNDiqKmVWhMwvZK6jdtrSM5x+8u/aTf1j+/A+zuSKysX67q/Rzgx9F3Cfdr3LvbmgI+mPOY579ztrtx33t4p9f9bvfXvnjRYuOS0yfXg/InaIolmoYJMFG9ybixUDR7Vx4yU+6Vm1qZbYv+ViOAfeyRz4QwSwV2fpNW4/COxdddJF90UUXDb1fyJuQCIJU6RECsNDyi0L49uCDD6Z+mGhEd2cccs4JRVPCzSW2YoiXIC9ZDlaFM0NXwZhmL1q0yMdH+N5dRQWpXKuX2I1ghInrJFcc55fQNCUJHVB53qVSyfQMw0UTN/e9CGyy/Own+sQ4DEezoVFcGTcDKNMzgDKMFDb1MhnS9kAoQp4TUK+qqiKHZVTzKGwkXOdA+ZjwGeEwIWwX6LDAL9nMMW1P3h/Of+W8996GXXf+4R8l5uVQ8vXbpYUhCE5PvMdEl+zMlnHbjp53yPe7+wZQDreyuVJRNVRtQnNDna5jvqC72ba4pnAVTQ+qZvT2DQ72D+TycEZUjACNK/jZs13LtCyvvqG6JmLoWjoZiRi6HtENTcfkwT3c3rqnvbNrYCASibNJE8bs1mTjW9ni+PUib6+sqj85CiUTEz3PSwwouWG7DSvHqcdf7X41ubNz20RVjZrV8focMxAemjTh5cow0gZZcjRkYbBAECPkCbiQ8H6RvUh0ZOF9IAuO8dWUStnFgQGLRTzPskhWHX/WvnbFLdfu3NMbBcOpz1kYCuDE5C8vNgJ1R9UR+Wb4M3AXNTWbzTjPrVi55J47H3lx7tyWdVA0Qv7D8wwZ0lAraLVm0CCfMaVpLYRAcnkTrdYhqG+ov6AMxUVI67bu7j3o3vv/ejxj7NmFC5epgMeGLtvnxPOZfikpK3MogmZbql271934wCdvu3/ZkZ5rYhX00arh8CA4B5nHkLMgSDSHp4dKSt5+35l/vmRkhjFwoeQfpr/fvUrengnJOPlsuWkqxUo9Snk8+N+hZx2yXBWLmThpYZRUzWd0Wrx4MV+6dGkwlf2KFqJnX9WJo14HZWZqtpAoLbSKh76/Io4fmq+o0LTgBxww44OqRHwP4+pEMfMF+K1yvMneQ7zTO1gYd+kFX1g7c78pqCgFh3JdD52sUh1LJkY98R6UwIYt2wilKlwfQtUmoc4lQht4lsb2kgSU0ZfLTHrFLS4jqfyHKBCHRCnuF3iUKI9ZlmVKOaumpqZhDYNYRdxsNkuUUiXOVQPKxgbCQg6X7WMVguVQChkg34SZXFkQXTDivnBF8dB1Ty9u376T7bvv9I4zzvvhzz/YtGsOVzWHWGPChx7SeiyXn1AWK/Qwg//4rM26xnd296p3PPT47772jTNPIZ/XsrREIoF7hJUa4CA9OS5J+5z++eNW3f27x7Ld/Zkk47pYoocMsBBrMddUr6tvUP39o/9c7HneCkjfDZH/JlJP0WKN8hPKbOgBIcRnPp/XamtreTwW8R7569M39wzkVIiA+G3g4YlVTmaWQxdhDIBj23tTOZHi+iidCqeC752Q8+0M6PFIJ5N0C0QeBLR7pEhsWRZ5n8M+1FAAHmYoCCarcNvC93AoJBEgLSRAAOwCVyxD3RrNc56n2LbmG6KAVEV6yfLxh3A6YWunkLMdFnHmsahR+vRnv/3BR1u2T/C7izxfEisYwT4TYHAERXG6egdjly/95c9iMePUQmGasXDhHOfkky/W2/bsGcdMdHFjYQQHZ4RBNNmyuLJrV09+4cKFA7Nnz+Zr166lr+vs7ORjxozxFi68ONba2j6eqJQcO5FOx7yCbVueCefIUp999r3kXXctpYz07NmzpXKTcJHlvawgW63wjSDSQrWMzs4+zi0eRclPsZTSo4+upGTdkHtPj+eZZ95U99uvGSL1EdVRVUdzNW6aAC3TYLj//ucEj1ZoR2FpX1u7VvvEpGkxw4gosD+MKaVEQrMRMwPA5HNFDL/Bo7CYFe3dk1VgDH5+2/2nvPDGB1+z7BKos4ioYpg0ctlh8N8mg1fmDBD/GbJAUPuypjvrtndO++LZP7z+sYd/+r1167bGZk2KlVAasKxsXFFiGOQ0q1pamrdNahqzpnV31yEENyWRjSGDL7RYU2jCmfP+uk3zLvj6jTfHY8Z38JAXLlymCE+BhpjmaIUSI8VWQeDAFMdxIu++u4nHorr72bOvvPXx5944Gnll34sIj2s/4xx2ECoipeHSK+HnFUzUcsw9tMlJZtd9yR8K70jgx3+RLDvc9yD5N2QDAr78ZRUMSWEJkbIlot+HyMPDK9V1PSIFW3VNs0oAxVqM2brjtwsH/BeSalFcm1w09vJ+Ko0lOkL/+tfH2Jx9pj770jtrTx1AAld+Z2iTnM7i61XHtt1X33r/pOuv/83511zztd8vX75WXbZssWuaNRCahRQihpPHuWmYZoTk2FKpiIOGs/AC4f+8hC9btgA3qY7ETRizdnb0eGNqaoDmtSA4u2DBePfoo8vCshUQo4om33B0JkIKLNZV1QS+MQ+bu//ulpbxfRMmTDCbW5rdK39wdl5T0dgsXv7PAGQ7Xzrz2MHx48cPTJ0woX/y5Kbc1PHj8y0tLdlp0ybkpk+fmLvu2nMLwX4KcwxdcaIRzTF01TltwYLMlCnj883NDd7kyc2lyZMb2xsaGvZ0dXVRDkF4DhUv9KOLnz0rb/HGxqTLTNbym4efv62nPw/+QT5seaocjYiZ6DAdrcIEs5YPPDwQQjPFFzRRbde1V7z1wRk33nj/KbNmtexc27pH5bykInSQiSuR/bf3nzPjb8lUje88V3RZl783nCcDkLdglZynX3rjsgsX3XhLVSpOXXV4jnc88wwmfixrZcE2iwnmZDIld3d3N6+tTQ6eeuqCwmlnXX3dE8+/8d2iZTo+Ceney3kAxQ0ahUKLsmhyGtr+zKG45Gf4y/XrcEoolJiQc0nh5AHojFEnHsqBWjQqoIueB+EfXAMdA65+WTULQL2hhtz/vYIgfIgVw/5z5vgtydFolHYWnZUgmvMiUZCXgFitxG3bZxqVMELp1VTifodYR6L98xmlUIFbtmwZVA/5T3948VNNtckcs0zELRXOZmUeXJgHVeMDeVO59f7lt/3uj/9coKncOemke3lX14ABI1mWkjcAXSdFqNCYkvdLMJ5T52PyiBO+dv3suWfftXN7d/P4ceOctrY9lPuL4oYPmT/UahecXgULjXQbZbcbI7qaLds7Tz3uMxe3WA7UvlRPVYhPj+9/6FkmeHz8Nk8AT0CP5wL9TQmkBSct0tHo65Oa++gteaxDPnWe5YLIgX71kZng/MbVHHX8IkXTVB3A0Ggs6dVWJTeddMK8J88586R1u3btGpdOp2UbbnhDGyjw706pBHc0nT/mlG/es3lnZwPXVNFMNCRhKhGZtCLg3IteKqqx+YfP+9OKtz48J1fMUctacIsq7El5sGPS9uVK7m8f+/dPPnfyEW/vd9A+O3fv7qlJJikxSt6VPN+urq4/PLXi3e9n9/TX+jSL0lfzk3Ll0ePHwpSU0gx1T9+A88Dfnr/i4CMWTbl58cXXHnXUrLWXnnSSH/CSg0xM1G5VFda8ur733ls/5XvX/erRNz/adkq+aLpohQk4EGUZUDz0yoYkYSgDVn4XOR209wYEKwIZGroB/ncGY4fyDn5EFNxquFMKtEuYxrnBkc0zGYPwqMST0E0Kh0N4nuRFhA1yhQGVP+3tuhVxzf530Rjr7+8nYV0B+YY3YqPGxakiEycI+KlnXB76Tnld8gLCuRCRLyEaQOSpXLZy5UokuumjaoLvOPlzF/9m6+6e75RAKyWS4mWHqny+wlfgXNe9ju7e2u9d98u/f/Oyn1/x+99cd39z8x2FHTs63d5iv1cTieiKEkkoUZUweXuKRV5TUxO2kl4iHnF/fsvvT9jvkIW3rGnr2xfAqaNPu+yeJd8791tf/vKpz772wQeR2ZMmgXkJ5xrcMJ9CLVxVCxmv4OZTVhfzRGdbOvr229IxsB/x8ONjyAXhO2Ad/Whi75SAxMSH727IH66wvnSXygzBkhVazDZm6Ap7/vX3rlj58qp7f3PnNT9Zt3372IEByiLTqBZBHOHs16xpVebN22/Lty67+YdvfbgJbjLpHYTDgkpoclAstlVN1w4+cN8/vfH8nefuN/dcfdX6LWfaVpFkuCptQSU5rQeYvMLd7R1d1d+++pe/fvHZu4/csqW9tN9+E1ltbS3k3rXu7u7mvr6iUl9fnzv4EzP/0fbvVec5bgn4Wm2oZasYOLLpSdPVnGk6L3+47gunX/CDEw47+mtPXXf9fa8ddegBW6O11WsjUabn9gxOfHv1+pnPv/j2cR9u2PHpzsFslQtUGYxBmHRUdu35DFF7hQWy25Esoetx2/GqNm/eDLEZpbs7g1W+a2dbVxdhcawyv2j5CGXDEqCJoU4dNWLwNAcHe7ZVVcUh305hT39/v1JdXe329PRABLhaKn2XSiUcKxXKG+wdwIhrCPKhxCjjubMPakFrcH5LRwcfm0yi5I5W9F6BU4E34owZM8bN5XJaz64M8CLVnz3rckqUQkayfBVhEL+fBizbT5dVVyUjrvv/K+w74Owq6rZn5swpt27fZDe9QxIgEDoIAUQBGwoBRVGx0HxFFFHUV5Poa8WGoiCIBVQwUZpKkZbQS+jpPbvZ3m6/p858v2fOOXfvpvAdXTa7e0+b+i/P/3lk45YtXcnmZktF8UslLJiVP777gus+t7s3l6Q6RK6Eyr8c0FLjsUa1KPQMj2T//M/Hbz/m1M9c8pvfrL79C1+4aN0R2Vm91apNXC9IRAHCansoqYhSdMQ/2J1/e/iwu1Y//PkVN97+yaF8kRArEzBukR09A61f/r8//fHFV3d99fe3XP/bB158JbWoLNrL5bKEGC1i4HzixKx7vHoHNjaRQlsPYXkR6rqGcRE1acOivImBnnoTFAyF9SM8HIXhIK/x4Ue7Sb0VFUaKYZCpBR7q5YMjRWPNv5+8JqFphV/d9I0/7d7dk21uTiusAfwiBCxzuVzmxBOP6Lr19/84/oHHn/tmuWIH1EBVHvIx8Xiqi4RH74torRQBXzBn2sYnH7zlakpvZUNrvn/17GWXnbp7X6ET5KuK1y96u+ip69otlM5C57yyccfxn/r8iu/dc+cPvu66Ck8PefJMIpGY2tqqIt+Z795w2Z9efn3jhXv7ChbVDTUzQ16KA/n5xwONCgijCeEG+4Yqqb7hwYvXb9px8W/vfAD063mlSOC5aTuQZtlxwwwfotDA6MT2f53ZW6v4r82kmN04gqDE6fjw4wkhtGaY+LouMCDHegaHigoeEJUBj0djxyfo+BwOrxsFp50jjzwS0e4D57aUzLbtGYQ4rOy5pq7rmAB6LdhXs7brF/T6CRr3D5XZZAKArmyjnkJEXm9oaNgONqKD3ffOOx9NnXTSIiXbGl41ulFYjFmbDxOstyjuZlmqTsQy0qwVIqr4o+/n+ZQps8fec+rS3//5gaevtX0ftO+q8w9aHj6e5YK8oCw4jlz7xtYz3ti6+4yb//zAliNP+eyLUzrbN97wzZt7C9XqMDOYnDypeXax5B6+edPOKb3DYwt7+nMLh8byJPBcQs0U1kQN0pdUN8RQvpS484HHfnPGOf9z2HuOOuo61y235nKlZCqVGjMA7a639WJy0igqMQ6drA0g9QdVPFEL4KiBEHfJ/vtb9LsJOepaj4//UGMeD0mr9gvvRx07ztvAdFPki0X58NpX/scu+U+1tTUMBKHSLhYFqhRbfL+fENJ425/vu2PfwAijuhmmGOvMvYmLtJpkCtQ1qbUx+NJnP/5ZSmlh6dLLddrYOPrVr//ya7ffff9f8xUnoNwE51Nt9B3w1mHxjVaplvwnnn3pa7ffev/zn/rM+x+ASTk4CN4S6ScSiWDz5l3+4YfPfnHZCUfdes9DT3/FCQKfMmh/19bKuok0cdBEE1sD5DogTJRdj5TdAECzljAMoijZhNIcUIg8ZFTG4bATAob199nfZajdWi10oX1HIc7FFA8D/mBxYz+4Y91j14z4OpZrLO0RK2vs+x5oVhIK5GsiAQggIqCBU9sx4olzkNE2vpFEfarqOJhSNgY0RghdpezqlaOisAXFvx9++GEV+YMY9MQ6tXHSl/Emqs9qqHGgNBSELYG4VbiZhmQDHRoaarr11//7+9c2fO7dr2zasZhyaCcQxekQx2pqxmm88oZl75E8kwxy1SrLdVUOo32jh6W37SDMLxMfxUi6SSgziONL4jhxbBBSAgxiMEo9O3xwxWyuaEkK5aJ88sX1Xzzp3Z9b8tg/frmyvb19y+DgYELPmpKhPDRkEosDSHWBxFo0tc6cj/n0J+wuEwM44y5B/Yyr25H3P+r7tjY8xkUm9v+cgpAZaTJUchpvvv3v89PptBtUKolKpZLyPC/Z29ubmDZtmveBi6//2du7R2YT3QxqvVtjz45Zb8YXJki+mTrVzll23C+uu2b5S3iI9et/50Pj8dabv/G344447C9cU2wTUUFHnUkVhb/jBVYi1qYbrGdgUP78tjtvllJOqfO7gWAUs2a1BFu39qb+dNu3bzxyfucGKiHlHWNdQ0RZfZqtHv8+vgirD2sqa4ICN1WNypB8D/G04G+p8YXGo65u9NXF5lRcU5Wr19hAor9NaCcTpUJ4h2gpj+p4w+tOcOv3J/uNhwJiCKEtzt6htF2pEVcqzLIsywBwKKR/jyfq+B49HlKgB1oPYBmlqruQqY0xKwfcK14YdL0dn03hTzWRkPEP7OdaR2BuxL0UqljRSwI8pwIkqk5bV+Y87pm77srln5g3fZKLUi3lONWMsnpEWNTHqh/i8ngFvUaEQwCVVKxU/bxLRFnopOgQWah6wvH8gGokQAl06Dyj35VFP0FjAiJPlEjmBK7/0obt71p69mdv++kv/nxye3u7XRx2NdbR3igJWIjr7MLa0K4zK8dHzf5Bp6hhah/ff6DtzyJc5yzFnVffmTXsfA2RFsXUwgKS2rW5oYz/5195DSZnqVRyUb8Q5HK51Ny5c1M//tkfvvDsK29d5HuOSq9NIJk8YDFQPjIiZvyI+dNf+NNvvnXjz+7+ezL0VfDCq2SpXGWP/fuW62dPn7ID6k41gEzNmqmrlIyeWXWHkQo27uye+v4PX3OzrnPwSChATJgJMLy2NgumsLzt5/973awpk8ZkIAH5VBHGOJh4gLccmai1eExtroakE6EqkPra78QDmr2WDQiLCCdSkE08akWG0vPCTooGuhBBBPas0/GI7YKwKKx+ktZ6ocYEcLAjEtcNdN1FXAil21iAQt7OGiV6HeVcvHnV7T2qGE8GwkIZGCE8CDTgX96xoCqVcjH5EoqGvdbB9RviwZ9aCI8Y3EDYPqPrBkBzPhfC4TxwLMuqdnV1WR/96Dlvvu+sUy6c1Nwo1MQMGczHg/nx7epW0Np+HMaPACLjFGA+xlnonkUzXqXGUbahSqHHrxWv9LVeD+H5VEtyybi9bVfPnIcee+HbsFhKnsvYlCmTn6OajsgHzMsacuoAA67eTNpv0MRzPO6TCS1WW2Ria/UgKc36K03o4WjXiQcaXhTWzPh31L7h7QtC2GUUNk2bNi3Xvaf78NvuvPebuVwO5BKoE6/zN+tWzDh/jbEdOKRjUlPu1l+u+goa54Tpx9U/GSiusfD3X3z+mZ9pb2mEqvT47hK3SaxSXPc+kmrcl8J/8e0d53/hCz/51PLly8sjhQIGqARyNpFIkN09PY1LlszbfN1VH/vUlLbGKgRYIj7qiW5T3W4+XrJct5uPr28TAQT0wEExwdqINTyjf8fjqPZ5dQslmKr8fstShVIMGRy1FwU1x14eoI0RVfLVCFCjZ5PynSsdly9frqL+UiqsSuBrGvLvxZrctipyqrfK9ssk11xTEMOYKv6g6zrAUHW+6oGH71tK7wPbeq2t6yoH93dV6lOzXFdRdUuIimETG7Es3yYyABqyra3Nv/zy3+k3/eRL/7rgvadcOqkxA8UtmPDKIpzgIkYLXry51t4rLuiKNsbau9em1Ljy+ATrtVZEF7WUopTwfS4864Ql83b+5lffvhIDKpXSA/bud51+06SWJkqQfFUZEUWvV//G0UPW045PnPDq13W77QFHvJKDU3G/EFP8r/rg2bg5HjV27ZfxSi0IQam5dMn8w2Zn8at8Pqg89th6FFcNfeLKlV/etW8kQw0LOLgIkF43OerjCPiv74mEpbNLLzjzJ0uPnLlr166ezFFHTZ7wChddFOb7f/idK5456egjbzJ0C6AeJS1Ta4h6QFe9pWkktZGSI/799Ks379nWc9SMzs5RMJ5ZlmI+Zp2ZDNm1q7vx6s+d/9KXP/vhizqaUnlpF5D7rKWvJtjGtX4Zdx/GIegTrTB4DcpAqllFqmBOSA9zQ6sLAsbXjUG08SRWpcrqc1FmyQ8CxoF54JH6sys9pPXqFqjIvK6XXYuzVnHwL4Tahl17YNoYh4Jfww9HHUpU/oxCtPBKkYUUDrv6tp+4GKh2UahoAug6FoYD0UF1zzA6WgGKaVB4FKpi4VVq4f/9rNvwRPV+AFa6nkB8ItD1pAnafeq6OnWoBjV0oC9/97vLIdTCbvnV1+++6P1nXjB3altZujYAcggG1wJ1YR3NxOzSuFU47vLVaEDjJqhZ2vWsV/VQAmQCuZCuI5Ka5EsXzvnPi0/88eLD507p3zc0pEMigH31y8v/dcSCOQ8yntIkZYATj0/A8fV+wuCb8KvogSeYV7VFZILjWBeLqBvT9Rfbz12oXVL9Z9zCVNZR4NGMaQQXnP+eHlh6lYrtf+ayD+Yvu+oHX31t69DxxEj7EkJnEf31AYMl7k/hQ3pHO37Josd+vOrav77xxpbk7NmNxZGREfhgQLOZGzZsAILF2LZtm37vfQ8k7//b979z+OxJbxEBf1GhDCdeNPbbsGEwlA0zxBPkjv7R9GXX/uA2uDibN+8o5PMOgk9OyePljo7W/ObNuxqv+9LHn/z1979+6VGHzR2gnstl4AsIuaosS30evNY28a9if7Z+ENRVTCq4AewpLyBemc2Y1BQARzLetnULWzSxYzJZpJPD8RhzHfpgwM4TSkGQ4pkmqjLwCVTR1a1Kdd0fezTjvqwiW9F2795txe2MNo6+Y+ImURLPOS8yxsoRY5KnUttMCQPVFoIwnrX/2KkFMVk+Z4PItRgEAYA8WAlxT6urqysRf+Hnnp6eJG9R6N0KakDDNp+4eYxbyfXjXRmCpOr4WExG0KV6ENhl33dwT13X0c9udF/9z39+NHX7zTc88Nvvf/3dJx61+K2kwbn0AgR/1MKgTPp4Rh0wD+qfJ7ZOwrafsF5F/V/jYFIZWBKggKqzvZV99ENn/ebFJ/94SalU6u7uHtQSjFVLpZLDC4XC9P/e/4svHnPGFVPf3L7vGBE44FFS6d+JDlkUAqrr33gjid2dcXuz3sxSkY0QvVGTq47y9zVToDY9wxq4mCI23KGAE4q2l1DMRArhUW4aRy5a8O+lR817KALGBPf85d9nXH/jX79ecoTPuK6FOJaIgKr+TrEbQ5kUTpXOmjFlYO2/bvk0VHLi1xoeHp46NjbWgoE4depUwEbZ5MmTeUNDK/Lh/Vdftvyy79x41wsDw4MaeO9C5E1dWer+CxEyM4z4L2/cedyln/nO/6752w+vdRT0oG7ZlVLr6uo9+YILTi+cffaJn7n0c9+66qWNO94/MFbBpcCdp9ykELw14doTF6S6dwz/rzCbgfR8nsk0aCceNePxw2fPvu+Ofzxys69g0ywCYo4v5iEqqNZwkkpXisCVCcsYsh1vqO6Z9SMXzupY98J6hQCtKWLV0hhRBL02LMKAp+cotqUUY2yW4ziQugP6TlAKRmVVG+E3NTXVCnyie6XxpOHlkOhBKUbs5sS0/rXxA3Q2oPZ045ub+z74kWUopVYHVMiFEE2gNY+BUKOlkpdMJs3ZLZ0ljdFtZ3zomlpipGbWsv0Xu1ocCreWrucFmkb7xbjG+IQjn8/PQ33JxRcv0846azfr7OwcOvvcEz585bU//uzaF9780t7+fMq2K5gusIRCA20Coiv+T727W8vuxuZWbX4RpqTyQHsPWToG0diF86a/+oXPXfjdTyw/e31f3/ACSls3TZ8+zmqOXC+uU3rtqd+9/8xj5/6nrSHJiVfVpFOk0nOo9D0KyxUreiicq7gW8G8qBL77ER8DJOp9mN/hl+cS6TsqAhVdA79XMva4FioScY3w97iGOi/8rK+upb7U/ZRrhFK0gEq7SDXhGicfPX/HA//85Q3RCiplSXb84Fd//EN31y5OZJUJ342eOf5S94qu7SnadeF5rKG5lX34Q+/7lsZo7/Lly8PILKJKiQScTmXWokIoLr1uaEjaGzbsMK7+3AdfO+3Ehd9PmEyTbjm8pvqC4g/yWEqOdfwrcKn0yrxSHBaPv/Dml7757T9/ANfGPWvLIaVBU1NmqH9kpKpplZ4HVv/00msvv+jCk5fMf7Mlm2KILaBfSACGIolwsSLZZBpTq00oxqpUcPB7oOIC8GMIp0gt6vPDZ04ufvRDZ6367/23ntPXN7wPOCXpVpFrRWZStYl6ft+lAn2If6t4q69J2NBCguJOjZu6tF2zZWrNtlNBMBMc4CFypK7dVV94VSp9m4rAgypPpC+GLcZA+TFHTF5KG667JwAtDFWJos6t3cuSgaNLp0SFwDOF4zAcT3jW6J7gIxe4j08FlRqxQnr0mH0bal6hwVMj1bG1IICCk6PrQmDddasV6lVyRLpFovo3Hk/oV9wniMeqq4h6sIsqMs8QvjAx0BC+gyICglQAfsxms+7wcIH39vb6t/7y69978K4fHHfeu477zcIZHYWMwTXplZl0S7gPasvR1wHTIkJVpvpb9XmYKVGmGao0MV59ZMEgbIu+ReRxUmOanrhkwcZvXvOJb77w6O8+8YnlZ7+4ezfUDSyHEGTnx9uYg+RibGwMYBn+2P03fXLF92876/XXt3y2a1/vsaPFYpoQboZcjRRllGiRMCag8jKRDRhR8oRKr5EYULSwxmgwEajP1GpSwziH4mVXp6pgaRSShuMfC/WhnEogh0eoYNT3mxomjx13zKKnbv3lt//ICcmhk7/3ve/5X/z2j79bLjszp01qCCgzNUE4Ql0+wAeqZCdcydGbYZCEUp/rFj1q0Yw///qHX7xDyNP5mjVrlDRc1IdxekrzQKgQ/Q71CLNnd6qS+dV/WPmT09931Wk7d/WcGuWeOHxtQTQVW4iC8JHtJCkRLmVaRkCt6dkXXv6JlIUXKc0O71eUIjKW5YiqcDbu2ZW94UuX/PuGL13y+M9+c/cFTz3x4qe2bN954mixaABw5CJ7pVmhyYggfMSJQhhHWQVJmQbJJEzS0Tal67glCx/81levvnXmjJaNt990PfnkZ75lT2/vcIp2SVJqEKDYVTY18BijSBvg5VGcx4SGxUVmSXtre2kiUCrs5paWNJ3a2uL5gglJ9dD6A+etCKAVD5UMRmgQmk/M9BjjmqVrEJEZghhJeBkU8AkpHVCeGei2g0HSRWd7U1Ct2lJSHRUHEMgFA1DoHKnPqwpCZUxBra+ztSmYPLk5CruvBJtmPAEQn0FcQTUa8hC+T5RYAzbl1oZUMGtSEwHXKkrPheQ+1Qy8UgDufcxPIpDOxJQzvFQqydPJpA27HLUeyJIc5FCZMKfouI50QB3op1JNdOfOfdPmz5+5576/rvqfwSH71z/9+e8//Mobb36wu3/o6LGya1Ucl9ieAPl26IZikwxjQSHPpUIXhFYLWNsSOieZZAPpaE13zZwxbe2Zy0598PJPn/cW9rn+/v4sIQZvbk5XbdvGxjwhjkOruepshzhY1pq7u4f6jjhiXne0WrQ8fN9zDZohG4lHAz2pmclUcrIQ0iHwdQiFzj300JDH4XbFHQ1s6frMp5xrmu9LDjAdJW7FoDo30mYL8YNqQFmkJSA55K6Jxiy3HOR9Ihys1cxERs6VWqABXEQNS0tbmWSGE1YlmlY47uSFiDbzsbGxZt+npba2xjexI/78t3+dvXjuYSaXMouyNcdmJV0XWjqbbDA5BAiItIHkIDCSKXMdz2Vaou+Mc5bsdT0U2ERecjQ5i8UixEiaI+sZloNnKIkxU7ERJZPJvbHJ/Mgjb0yx+3Mk0ZKxUEvAOTMSekKTOrNSFs1QDgkLzFHeWK1U+rp3juzc3rOdHHnk1KHzzjvP2d+sBHEM7mGaREDHwfdL3pQpUxRL1Fuv7mr413+fPuPNt7YcOZArHTs4WkgmLKOD0CAVeG4gfSen6WZ/Z3tHz4wZHW+dsHThC5/6xLlQhoY4iXv66Sv42nUrgzWrN+rJZGVqUPGZ5gWiqTWTNCyeJAZrYBpNSCksSoWPbW1kpLI1P8LHUm2uc845JythlbidoITkOM7h3bsHjcBhgWmxRMXxS9BBwwd69w0PZRtSJmWG7sFqFNTPGNT3rPLAsmXLvEKh0Ikio2jXxk01EKRUq1Wvubm5a79dlj777BsL7aJnWJyZRLc4ChaTpg6hHoNwqTPJuOMFpaE9o7sc4gXpFBdj9qk9F100zhcwOjo6Q0qZ4Vx4UqoSaNSJa1EAs9jU1LTn5ZffnuaUgxnVqu2LgFQZkZWgRLymSakWq8Foo5Q2AQ9aLLgDftnLmc1p2mKyvvnHzMciV1+HMUF4JpYJELpwUXYS17eUKR3rzH52FAKxEf+BtmNHfta/HvrvSZu27T62r3dgRr7qzB3oHyRWwpqk6VoDUSnGwHGrlX4rla22NzXmUgnzzSMWz+s+8siFL3/kgye/2tiQyefyxba9e/vTlkUz6XRapYuDIChjnGWz2W5KqdIwUQtCLpeba1mU+r6WFEL0fPaznx1bs2ZNGKXYz+QBR2gk1R2TWeBlYpbhnfsxF+1/7vSIZRfnhemM8ODRuQdNBSHIRAiZE92zPDQ0JAqOk25BuZDkQTqd3ngocghg0gkh9ekC3L8GWY1iBvufEy8IwIa3RRwM6kAuPCqddSKVpIP6itGu11j3njEgBp8fjHT7Dva8GDRzLGpRDJgwteeoyrWK5zWVxsZ6582b143PgpoA8NhCsQyChWn54XyacX1qpjGJd4RaFXZgqGsZfX19yTfe6Bo799wTigdr56h/ZkT9iWfErqkCe9GzvxXWHx14DAwMpLPZbDty7REJB56nGvUXmIK2kUMcqFWwLAuSe+BAUKWHYFl2QFxiGH686B6irdC+DdGPRvScyWg8Iui59VDnlsvlTt/3IXkmMPYdR0XHGBYEz/MKzc3N73TfeZ5XbdX1BNoZcRBQwdnRMyAmMcHiq+9b27ZnuqWSYaTTHiEOJQ6RLqV6wAKDBaynsbER9RUqghcXYuHAfuJ5AvMHcwGTt7k4VmyWUiaTyaTHLb4JXZFKWkGlCjHc8eDiz362OvGxj505symZchzigBJAaUZwzm0wbIG+sB7GDXotnRDDF8IF+EOgkm3Fiqd4X19GexW921GUy8gy0dvbayUSDUwIG5FS5FgVMzJjHsRRQEeVWL16tb1x43K5aNEaZWHDalq+nJA333zTmjp1atIgBvFAwKxWRYMIYUPW2t20qWCtXr3axbnhY60kKFXFzz09PelUKqUILTRNk5mM7mcyGY9St+R5qhPUvcbVmXAuoRs3rpTDw5Uk566lzELDIFI6nAtQ6MmgUjEDyGdDMfcQrMwU1Fqc85jjT4RmZcgCrDo57GlVbl1/cj6fV5qLkR9co4eHjwppR6gao42ie09Q78HDoQgLFk28GGCgZizLTU+e7OBcQlayVavWkmJpXUDpRVVC1mwGmnJs7Fo/l3NptVqlnkdaGdMaeZL7qRT3zj33hGD/AbpmzRq2ceNGBFBTuq634Hl934cvrYSzgwBV0qbe27vXWLFihQ/Bm6itagshioKgRGzbObNU8rEAwAJE2EUFB9evX69DrDTum7B/VtLly1WbaKVSKeH7PnarMMBnWYw4DmT56ts0DBWqZyZsDVlDRkZGUpqmNUHKDvUrmERg7wbDlu/7UEGu4Y7rnnv8gtHODBXzaNeWrovustSzYmygbeJnjq6DRSxJKbWDoAJtjMAwEGKyEBwFqcvBZOlqST9wPYBSPLovZrea/CmSIkYD6mzChG/EdKXGFZSkTzxxubFnT0+LrpMy59w0DKNAKS1aCYtX3Art3mkXFi+epODQBOTH69rlihWr5cqVYeoHvKWuU9JJqGyGdkalJ2JiBzCE02KxqAhM8SHf9/uam5sV+eX+L4U1oFgszo4fODRpFUmlUlceGRnZPWvWrP2JUuPBlyiXywuQRor8cjXQ8XAqwLI9u4ceSw/K0Q+xT0Si0fGgZ5dSEWcEGDCYKNlsdsc7WCat0DSMRE3UDh01iKIvTyaTarc92IHzKKWdOBeWAVigy2WGUgkuhKiiQOZQ5+ZyuTmc80wQBLEpFpf4YpAOZbPZ/S2EWv7j9a3DUxoNTWdJqsnANxuTPAHycZ3o0LHs6ujoOEDnEgUpnpCs4BTmVcYMJ5FwOOc8hfaJpe1s2+6ZNGkSrLsDDqTbGhsbMQ6qeOaYnIRSl9u2zNi2vb2jo6N8iDZOIyOTTutgy4ZVgHGESaPo7lKp1JuHsv7Qxr7vT2XMK8I2UOPOJIxi29QDJ5FoeketSqxH+1lwqt4B7wuz/1DnwmUgup7VhMAgUiCreBNwHGe0tbUVqeyDHpgv8fiPJAOgGQxFZm7b9lBra+u+Q507NjYWuwwhKoAx9A8IHZH27M9kMgOHaGMsnAsxFrG5YQP2GbNMj1LBhZfNZiEsfNC+hUtbKpUO2y+bFXI/MGYWCoWdU6ZMGXcZYHZEHxbFIkFmIn3HvVuOWL/VOdpxA/+MY1IbL/3AgnUvvNBtG8agv3TpUrpnzx42c+ZMdc727dvZvHnz0EGNvu8nsMNFN1UR+pC6KXDS6XS5u5uIadOUKRR3QCw00lCtVpVvhc8XgsBOIIysaWYymbR7e0m1s5PI/v5+pP7qeRy1anW0BWK31WpVIjOA86FRkEwmUTILtWVH0aJPdIGUGT82NtYWUZhL27YBmkeUGzszLKciWKUh+wWlnz179ki8M3QBOzs7QUkGC0AtYtGgUgMKC2QQBKOIJkf3rWWyo3vqiUSiBQsaFtQoHpGLPqc3pPQ8Pp0vebH1g3gJBg5k6xIgE4IEQ3NzAiCs1Mo7dp6bK5ZP0A2++6sXtPyjY/LkXVu3SXPePML6CWGTo4KbMUISRrmcoNUqo6kU2kYzTUUKIsvlMnb1Umtrq1ps4/ZFP6fTaV1LpZoTkawa2hFs0jgX7x0EgZ3JZIrQs5w2bVrUxr0aIZ10eHiYaZqGOIyqQK1UKhDaoLRaVZNB07QSyDXxbu3t7cqS2rNnD505cybr7++n2Wy2KeKI9PCMyWSSwk3K5/NmQ0PDaH9/vzd58mT1rL29vSRu7/7+ft7Y2AgOwjhLoQhj4jFfKpXysGz27NkjZnqeJPPm0f7+fm3y5MmCkDGjWrVg4SGArEVkKorABSc3NDQM7Nmzx585c6Yax7iXiokzBguFNzQ0JEulElTB1JjAkUwmMV6ClpaWETwO2krXUUQUHniHcnkwq2lZK4Joj6N8wvgV3h3nxuNYjYvodJHL5SzDMNIY+/G8i64Tv28pfLfanIvnDysWi2nOeaI2lvdbTRrOv/7Z37+0w76w7IH2PSAZU5LTFmX/cs//nXrVoVYhHJVKZZrjOOgERE/rVyFYEGPJZHJCgKj+QByDc56GyQazWujCNx0iPItZGTMzTCk9pPpztVqdFXcYBk+kGA24K6Sz98G0Oth5MAnPPffchbgnor3xgIlNXcdxultbW2EtHXDAfSkUCnNigZhIX1FFkJFnHhsb2z59+vSDltfKIZmpJCuHe54HwV2RrxD6m3u6PzHkBAni+5oviANOUE9ImtaY/95TJv/0/ad0Kp+2UCi0OULMaG1I2Ou3VFqvvGnXzXvG/EXVal5RaU9vaey9bFnzldd/fMa/DrhvsdhelLLV8zxqKCEjFS9Qg9r3/VJT06F340qlMgN5ezwz9CHRTKBLB6EJY2wklUodEIuJD8dxFgFjoCtgjkkdosx7pRGRTqe7FMDp4G2sjY6OHg6KdPyIhSfezRlj4PPc0djYiEXxgANuyoIFCw6LaeRsW7mmyjKE39zS0rLzUOrPAwMD6WQyCWwEaPoAkoonlppE6XR6yzvErFIjIyMzUpzbMHIiXki1oMDyamho2PUOsbKWXC43XdO0uBpTQNLQYyBG0r1EIrHjUG0cxbsmxWilaHIrlzeZTJbeae5Vq9WZQRA0xtZVuMosX63RNRcFF3752ev/u51fWLUrPpGKnYyUbIP85232iYtueO4tSsiNp614iq9bdUYNh14fQKlTa4448pRpBVLL+HO1hq0/8AL4KCYUJMcQOVHhCZcIYpZquej9GxN+IiYy3E3U9UELEqMFmw4G7sjIyCHPRWqoVCqB3BMR/QCNT0y1eqISEZJshzw3zmHDYox9RCwq0efdg2kHxtcZJsMkIVVAiic593blnMl/f9G7ZteYQzRqq4ploFF83yOdmSQ5YWHbU1JK1aElUqIqXEiMhmt+9txtr/bo8witADDAiE/Epn6v83ePDty5bn1p4enHpvpXSEJXRpV9Fca4W61iQAdRfyhfEn2C3fsQ76mee2RkJDAMw2bI8gkho4Ee8u6PMztPPD8sS4QJjly/6k+AjogkGOHKMIjPUw16kDZGmyI+gUUhZucOh5TCDtTO3//cpUuXYudTSmCRKrjanJDNwN9hhRzQQeOBXxJZNKqNxinLagtDSMxRd9SXUqdSKRsZkwhGrABW6ssMZfdWrFgRPmsYliAryUrFPp7P57EZ+rquw0IVxCTSdYB9kK6s1shz9w/0x12A36sdPsryK2kE/BzHuw51OMQB71kc7IfZISlZQ4HUbFx88ROfr1ZKgqEgSKpaWgCmglIpL17Z7F8qpPxZBH6oHXXlyTTeQeLYACaZS1yYuhNomvY/4BPhO8wcvAiI3qllSup5eqUS/u1gB4I8o6OjslKRnmVJpJ7Uii6lNWbbJbDgxCb9we6tGs7zPJZMJn2Yhwohbyu699gFeMcDMQ3fZwikwW9GTCUuCT4YVl59F62tkpTLihOBGLzi+1XH9iuO8CqaUCTISIEq8RzpKtxCAD9XXW/fvnwwZUpD9x33bf/w9kFvHvHzPuMJXYQ1CRqVRa+7mG287cGdn9DYUTeuumg1W7k6DNRCJAYumOu6iIkogZ0oiwFX64Cg2/hjU5C6jOsiUAVYQFagvlrqYGcqlChEcuCrRik9WG+I5MdaGiqgWz+h9r9KZM3gQeA6wKVTcaqIMuxQh/pbnViNDotGTdQImfgO58i6IKZiCo//fqhz6zAkNVO8Tl6Ace4VnJIIaKLOsgjhEGRV/I+ofdVigH/aIZV+hEn+/4r5RouCsr4i8Rkshlg331FV1yQSu3GoUQ3rK36F118a6SxXKg1EOgw0YKqCVuGGXEpIhQ0WCqGy7YQqkvGGxIG+joJ2UcOFVhkomt7poaIXwuRVL6WKVSdGmQ95WnOz6U2d2tLc2JjuSOj65KYmK93SkrTa29vB5oPUVAibPsgRBXQQcJNApjCPKf0/tExE6XXQAwpAoVtRs25gmagUI2IDe8ihtQNh24UDG6a3aQeB8BkATdLjIbCJh2I0UnCNcQ72w/jczf3DGJz5p9cP84qrKQSYKqoP8dxESpf6bk4M5XJzkxYjZM1FE3aUVCqlYjeImyDvHpv+79S++A98d5QZRPGhWF0bVtH/r49AaFrT3oQFGSlJKTmR2II75M3DBSDO0qh7x8+w/y5df2zfvr12XV0XqF/wkdbEwgC3EjGKQ73r5MkqSx23SczvqJ75nd4z/l4qeeBuQJZF6V+iFiNBm/XW1tZWKeWM6AsYl5boS2144LcDzgVpV+VqQ6w4dI8MmoyEjw9183HXVT23aZrq81gUEFd7p0NKM2LfVouHz+OU2dEntIwkLNMhJVcPSTEV2amCwuKj7c0JDEYnMgXjB6mt7hFxJSaFMs/UolCTRJpoVu3nOsRmJypbABLRQ4xK+KIRb3+c2qut3itWqCKU4HPf2HTu5lH2y1xp2AuAkjctwi0uNJngcxsFWIo+owbPQRYFBPUiKCnSLZyYJnpAaUlHg3X/cxTicePGjer544BPtCDo2DXBvDuTzDxIw4cu08jICHLsiLMKQtKBEMi2uGHp9H58HLW69uj8rVt7MSizRx/eyB/YUGbE1sKTYni7FAAgs4625IjrAjG3OqwMkRJxAMQ51AUR+IzVtetN/v36p5YOVcFA6mqep9wEJTyjoqGuCybgetaj+n5VB/QUQ+IRBQOHCasWUPy7pQXzocYWfEhLLtpkFGN1dD0G3YGDQJtrr1Aul3XOeRXBZpc6IsmSWAxQdUgRgIw/N549Dq+B64LNKjK7FWV7BNVWfIUIou/v5oAVGtYq/p5IQEwXDrdLmljTWI/d0/rzV659dKS8z6KepgaNntQksyhhWoIsa7rovVLKl7ryXayRNIJ9FkkjZmQyeNcJ73QwKwHPUkYuO0pZRnoc8eJ1sLlX59YRSqq1OaUkEdiqVYqqGR8cmDMpfadhNkFTCnBQnxDPl64nU+kGduJR7fcg/xqxGNVMzPjflmWpfHnUgGohAEYVAw+R4f0YaupfDOfi94avaYlwTqodBC4Is23l54J1t+bf4T+rVlHlhhSqwYyXdoxkN/XkW7YOeI0be5zGN3dXml/bk8/uHhppjFl5DsLIA+mzhGVZXO1ilhWKYIXvYLS0tIR1IxMHqfoZWpKRjxZhKmp/O+ROEpnIsVma0DRLJ8Q2NSETwoficNQssUGgKgxFVLkbtvO0aY3W6Ojooms/MWvzvElsH+GNnBDmEer7RLi+9E19UpKWP/f+2X8FpRYo2qOJpFYj3/fV80YmvHpuBJ6QYTlI/9R+jlLLRtQ+sRWvgmWxW1H/fvXvHJ0D3UnELrCIhKBNTUuMjY3F9z0YFgTuHDIRiWg8qOeN2l2DCMnB+jVmV46yRbiRmWAJnK9cO/jodaJC4wCQ6Nz29nYEwpNxG0XjQS0m+D5v3jw17uqfF2MTP2MTicBNmi51QZIkkFw2j9i72/sKG7KD1e3ZYXdrZsTenB2svp0dcbdnhawiRQsuvHjc436m7/tmFCRXls3BUut1z4FUazJuq6hPcF4MvT2AHSpaLuIAL2JmyrKBWZPu7ydyYKCfHHXU5G+c8Okn2neNWBeBow9rQjalkSWz6J/v/s6pv/jF5X0pPpkzMRjin9F40aWRVkEUWHHtwUrAKg6/NZlMIortAU8Q70Z+qy9QUwHzDCsyIsie56lFxHWBHlOpHvjyxXw+gAmVGR4eJr7fKhgLFWnefrvLAJPxNateatKpI4VhophHC0l1JKp2oCmDncha39PDZ5omavjV87a3tyNQhp0g5/sAiCunT2UFgO9OJBIqLYmoMX7X39+vUkNIhU0OeQQx0ArYdRH/wMoKy8y2YbolmOuOIZYRB6lUCm3yZM4GB4Vs4S3cNu3RXA6ucNr2g7xXNyzrbJK4XtWH75wcGRlBqpD5Pu0jxBr+9HumXug8mrtr75g2z3VKUPwj7Q360PtPmvT5U5c079q5c7ShoSEUwG5tbUUAFds5YLVqp1WzITJb0D+jo6MNTU1NYpAQ0PfWfGG0E3ZaBBQxwcJzDBcyN5xzlPVi8mXwuZaWFmXeA+zCOjvpZEJopVIZi0xatchYliUc5nDmMMQTjOi+mEywGILu7m4xbdo0PB/M3UKUHq2Bu2JfF6lHmNvReWJ4GMmo8L2gB4p3ZSgECiXgMaGqcDdUYHd4mNel29G/yjXgnCMYiVxwDoHWuCdUfClSl4LnhbEYtWsskAPLT2tp0fV83hszTdOtyIpIkzSxXVv3K0RowqI63FIlykcB1Jck0FBzgXGWHBjYxbWsFmfx1FyAdRP3D7If8XyLU6y9vb0MzxzFOfKapiG1C70FNf/QhpFLmyXDw5IgfhWmZTHe1TwqlUoVJI4qFRHAveClUmlGMincOXPSRj6f793w9/de/N1b3rr7pa3DJ7meFOecOOX1Ky+cu2bt2rXascceu4CUiEbBOCclIqPVqEDEampqQkrlUOAVK5fLHYZ0ID7vlBxfT+mgfsZDsFQqtRO8iAc7FxqJpRKdjYFky+HAkBbOE41tKYjHWo7re4LqVBI/onoCeQrKigUlgYHOasr6KW5mtWTsgmBnwuaYSqUA+Tzogf4uFovw9wChBZRZS6fTRoUxPSgWK9lsdkJpbv0BmvVCoQAIL95XT6VSwHjQZJKyCqvkUolULcX3n7Vdrdi0w4WszvKOFnPL5DBBW4WmtRqa1t3cnI1TSD1SyuNv+O2WS7tHs9OmtbPS5ec0PDV7eudbg4P5Ke3tiqVYWTmlUommbXs3bWs7KPAFTNCapi0CZiMhhFMk4zuHTinAK29OmTLloOdioaeUzsZkQ9YGi0cqlUKC2ywSAoXoLe/QTlMTicQs3BdFUMVi0W5oaAhAVY9xkslkNr5TG5fLZYUjwftBdCSyHoBULKbTadRvHBLUVCqVWmLXJdTglAFgvUBNNjY2HjLFF0L9LZyA51SZNPweu3OlwocaG1M76z//Stcr0wQqHRXHjFDELoqBwya0SiskwTJTCCFHOhrvOhRUW8oVLAi+Mq9YLKp5h3oEm9guYkKwJqL0+o5Dzb2xsbEFumVJGT4vMA3qmUulkjE2NrZz+vTptbS+KnEDbrxarSbznmdW7Lu16y9bdD8j5H6c9QAh5CqyWuvqWq5LWcRAQeQ11hFUK6RtC3/LlmEGyDMuumnTkLrhwoXLFXyyu7sb4AykilipRIgtEDHzuQZZei78Rx7JaTh35UolOzXBrAG4JYoWi7QEwJ+oYF6lgo7whOdzn/IEIaIcVX9FFGtMI36o9aI7tpMqlXhATSqkozow0JIqFaV87IikLb6lmkSDg4McgBLULSD1hRW17HmKP7BYVlYFJWQFBd8ivp++Yhn7wqIhlDMDuIO8ucLGx8WeihGaEGal024MwV2+nIqHnzBCsqnarce9RVC/UVTz4Z+O9IWldmO1alxxxauYsDlKyK91jYB1F1bAog3bB1unT7JiebR4VxWktVXFYvYbL+pOiP9gN/E8TxZtmwVI1XGuC9/PMEozg+Vy6vSnVrgryTKybNkyNYkQ0wEk+I09e9g0eB7KDPC5cBXIxTdM1O+oNg4j5Qe57+joKNxCuDVIacKVCqhFGXxi3/VTK1avMFZdtCoE6iwn7PSrT6dI052x7IwgWU0CFFXDuqCRYbHANcGiVP+uKv60DMxUnbS32CvH7DHN1ExYDSbqYSJ8jYzwJGzFU6fzVUPrJLlo/LlPX0G0ZctWIO7kl8tlFzGRyI2pBeXieNcasoY9/urj7N1L3y0qeytc6iGzBDoatjXi/vjyFZ8DS/pFv5MKd3jFUyv4smXLyG/X/BbjiCwny9X9X331Nm3uXOlRj5rUCLUonbzUkCKQbrio4b5ryVpt7dqVZNXadSIu6nz4pb8YSw47iztu0U5olmOaWbhuRpwJREJtQoyhUCgchvxsteqnGhtT294JApnL5RYzxgBtjXEGWG1gFWD0Fw6VSni+SyY6+MjMmR0plxArBpMorvzo35vG/RtgC8ax57AQGhsbO8JMl2dYFrMMwzSI72QIT3ctv+alDz60y/9FxRvzSRDwkESFBUxPaksme/e8+ofzriXEbSfEAMgIbgF2buyewIMfAF2Oc7sAdyHo7Ejp5/JVt6EpVWlJJuMiFjz7jqRJpaHrxPU9Uh2HudDdu/tmNDZajchgRAuuihwjINbQ0IDCmxHUJKxaRcVDj+07+rO/fnt9X9FlSsRD2b0UqleyNcHor65ecMnHzpnz4uDgWNY0WdehwDjRs89V7+a6rOg4yPjYhiECxjSjWHT729raivvr/+H72jf2NMzuTC6e3t6ONgKs2o3aqTkqIHq5viIO+I84pgMaQTfws1G7TI6+Y6JgIa1YlnVA4VpdGwMePgWgnbJXpq7mVqc3TAeQDH5wU4OV3egHiuREYVkc9VgKQ6LddMtP5metNguYhqgwChsV+gcuTvGdwFJSyqnEU8V2lOikms/nCxBrIYaXmpyYlqdJ2p1gKUU/hwPVxSUP00KQ/v69s6VuJNNGOkbjqpgIaN50Xe9Pp9MoKqsde3MDc7797Pt2jJR3Ey4sEig+0FBD02ea/OpRv77krMM/NBTNgb5DPK/mOIXZnqdQm3HZgIouudQ1+0b6Rg6bcphCqKg+IRZxZTWe5BjriKWh/5I9wz0ZjWk8yZNq/hYaGnqm1xc3IdWBvGVjI5P3/GPDyd/+8SaHGSA/0aTt28T3K6w53Wje/8gO+31ntVUQ5Isq2UgyGTi3/Gnzgp4xnv3aja+7MByIX2FpU+OWldLTabLvC58+8u2TpqEarMV9a9PwzDsf2fGRnpHi0tGKPb2cL+vNjQmtNZXcsOK3m56+4ZLDn2xoobsoXUViuXP47shlozfggu3uLk597e2+BXYgWDpDm8qSzBaBh+L3OPKLOn7FASiZnPTHv+9Y0pe3m3SNeK5P3EyKmumklmGMFNZL2X8sPXgNBS6VKzmJKZNMuzndBtio+Zs1e5dt2DF6cr5UPqF/zNWO+Ngjor0tTRsMfXNne/bN846dvG7ZKY2oO891D4xlkhxYBURibWISE8AYoMomFpNQ5d5EPR//NuJXhcIrUxTcROpEFiITYvduaf3+H28dL3Rqcq6hLC1whU9vunvj9KZEInXCEdktC6akd1cU3p7ptu3A3z3gBS9as4atueiiQKc05wox/OSmTUf+d9OWj+4eHex0SDB/1CuT5mRSS2lky7UP3vn2sjlznvjQwpNVuTnmyc1PrU6OVMsnfP2R27mpG6yNJ9saTCvti8BvNFPirMNPeV2pBx6qgSllJbtkgORlcvNkxVb1r83/PXp9z6un95UHZ134t4/PcIXDDE0H4cHW9mzb7mM7l7xw4cILX4Hb3j862kbdEmsymoRL3MgKM6lQjNPjC8/v1v+8g2S9+aVSiWAo/e3138/LaNmMRwO/MTEpf/zko9+YPHkyFiL32Z4nlqx49n8urbiFJZ5bneb6FZLgGT9hZdYvaF/88KRJ02HW2/vG9iUTIoG4gspAaJomKqSimJv++NZtSwNSnr4vt3f0ib1/nQlDVVFNhOJXIR2HgUw3p1vKr3WUt44FY87YKT9+/htDIhA0xTMJYQTujObpr3541mW5tWvX0mOOOQbaNhhDCueDseHQPLVko7ugc0EOcYj7dtx51M6xt5aNlAeXfO2xi2fYQYWkzIxjJvS3OzNzd5885bzXFrcuxeJj9492W7qe1PaH0HGPMSNX9mgm05T/3YM9d24qphczHpYbwOIUQZkQL08WvUX+cf45c1d19fezKW2NonekZMyZ2lb40793/mi33Xg6VaS4Ifkp5DMpSZCTZ1mPtrcY59x178Zz7n2m9xuv7SydOBZkjarvEt93CRwp4o8R3UgcnU6MXfqPtdvyH/3G86v//I2TVtIUSpNXMODNER2zbYt1dGRHn3y2/5Ifrun5Ti7IEz2RJl6gE9tTgrhKrzHivdJE4JAtA+YZ16/uOwM1VdD5oEiAgMfM98j0lCi8fP5hR2IRD+ndxncxDKTXt/f5R8/r6N++q2/pD+566yuv7rbPGCizKfmqR5wK0LZCkY7KPo/oPHli0nTIfc/vETPuST973vFtd37l0gXrUBXbM+xqaUE1l7jAxWOATiwA02OC1HqxV/UQEPiBKKoKqOmo249S73t6izMferu4bm+pSsD8TTkqR33i2jmSoia59kPTPnX9pa2vFXJ2U6MF91ZpGkx0F5Yv17AYICB4/d/u+/zJP/nVpd1jg0cWfZf5TJAAz2UwIkccYunk+ESvJI+++aK87elHnrvx0b/+5Bvv/fi/nt6wcf4G0fd4NbAVEVvgQz7ZJ65wyZRkU+kjR5697J0AaXsG9vCFsxYqkNS3HvnRpa/3v/3xocLQUQ6vEsFi1vXQDdQkO4X1M7Jux1py1/o7N5w4/aS/fWPZN8BC4u4Z3pNsNhQJCqPM4w4NJd2XrV2mNq/Xhh66qOwO/NIZFQRyL8+VHSJ8QQIjIJ3WgrfOnPvop//z9t3vf7jr7k8MVvYeHyTKxAtsAjbBMCMsiZkwT9069ty1a3feu+3E9vf98XPHX39fqVTijnQCLlLgriD7Rja7nBjk+d33/mBY7nwPdRjZlAuIbZcJlQYIhkJ6NxVupkR4Nnlm4K8/J/1M4dEUlQ+BGkBAaMIgJ7vnX0AIu3dbZhs9hhyjQEYROEwbKuxj09rn9hObTPrlizf87+ahVz5aCgYX+LxMHLtKfA8iLpQwTyOGa5y2efQ58uSm1V6rOfWtkzrPWf3xY7/0IHFIsGX4OT5hQQjz6KrQIhi1tdJgoSwIdcFbpkX03uA34yO2isKXAVxFFNMkqtgjMVouekPFqiBMOe+h/gEYdhjRbKENXPn1J779rd+9/t3ufKj/SLgLKfaITh7SeZx4gSvHig4Zy4mG7SPi8299/tEz/3DXmxd95tKjXnul77P0fe1mMFYd02D9VF1ZKQhXFr0gIF6VEerSMGwbT6RwMgF6UHWrEl+heKynOPaIpqtlekQTdmnIiWBw4w2yciWhMOUhy/eFH732tYfX913VnfNSYS4CwIwoSg/sEMiXhCCeV5R5R9K8ENqOMXHa67vHTnv4ud5nbrn2iBVzD2vd0tU12jxpko40Z1zrXiPv0akWsfvWam9qvICMc4iiwoIpBYEwkMTFJ/K2lMPlvDNadHWiY4RVItb2QHgmYcPFAPat57mO5jKVz0JWZcJiQNasCW6696HT3/X9m27dMDZ6WE76RPNtArZPkCtzMAGpUExApefLqufIqu/wQa906i43d+old/3oN4tbO/76xPNvAlqpcaypGDGMSFeTtEpEkQhy0DqS+EUXzloY3PfmQ2ff/Nqdnxnwh+a6pTKhtkD9f6BKqWLpWwFyIiEDGRBbVLW8O7a4Twz84KV7XvjEx+Z99CsXL73kqb39W6dMbpppIVWZ2M/qK7ulYHRoSAR54gsfY15llSAWx3Tul3+y7jufX9d991VOkCdMMMJtpLQxQLla6IAv9Eu+LJAhlucD8/uLu3+4eei10372vruv83Mlm1CHwpwXegrLKBGsVHGcnCAV7lMfADONxeMylMiUBE8IUcSSVwplhmFEeKESAkINIuCMdLiK/4kQEBF8NIYu86F8b2LapLmDt7164wdfGn7wezl7z+xqoQLHXXKdBZhcGohoQehkU+KWPem5Dg28nD7GepZ2VTYsfX7g4Q9ee/QvvnvktLO2kxWEkVVRib5K+fkKOwAOLqakmSWCHUFENaYxxlNMMB8jtuC4PImgmU0VwyzRmeRE+kpvSiJaTnmo90gk7Rr1zn9x21A2V3bh2IVCqeC4w9ujoWPyTYVOVroT0vVs/419Ys6PH+x99Pa7t5720VOnbx4ZGTOSVhZpooxpECtMiSdwS+iURSyTaElQ3Mf6DWF+qqZNRzmQIHgfSbhBNc2hAmMskg+IVpN4MWg753+e/PMzO8i5FeCWmB9A7QbDCOv7OMtvBN6JdSlxfd8ORm2HPL6ZvOt933r979dd0HrJ5Zcc/VZf30hrOm0gTVnjbVCzgvIAOYYJsnl1JP1IS4cFSIGb0qkKKj62bkggb00ZuAJhiKq4KJHUIkyzGPZJxDl44OuEhLoGccpVJcoJCa7/+4Nf+MWLb/1yTyWPkePryLRzjVEpNEnBnSnAnqYAU4qhkzGiJSz488FwsUgfr27+wuZ9u87koDKEIQMQjxIAoURP6FQzFBwjkn0+AJ9ATWaIr93/3Ssf2fvMt0qySDg4sgJNZeGlcpNiwRs8Q8z+LFAvo1JJdqkqdoldC2/e9Jv795X7vnjdadc9vGdg36S2dJMdVyau++06dd+EaEkKu5sJjzD8TxE3Y8QSkwxVe47dlfvjSb5rE10zBf4eCMnVUIGpH05Q1T2aWqWIANJ8I3nq3Gse/LDxqw/e96mhXE+7IdJ6lpg+V1NJ03SDM0VrLRlTC0DIHFjXxRi+yjpU/8BA5RxuLl5VEt20mOcKG/rp84sdQGgCM2KOlQataZNmuz9+5rqVLxcevNINckRWmI8WpRpaSPAw4RgtQNimsTBQRkAqpQpBKlWxg758yndfv/QvNz35rfdevew7r5/+1Ol83RnrfCjtdt214zWkwfrBoCdR9BRWUEYrGrQpKbFtCQth9PqHXhpMJpN9rz//BHy+IZVLwbZQy6GHPG+UeGTTQCWb85mkZgKvruHqqoRSVYsxPCacvJqQiPL4mK5TTvytw07rrQ9t/08QyLkvvDDU+9Iz+5Bu2+14dAz4Asmia4SZhHHa85jVOazURVco1Fn4hXtjtWMB1TQ/ZfiparUyfWhoeMFQoXDYvqHh+VLKI9575aOPPLnVPrcSlHzKQYaoUEJhjCLcM3wS3zc0iQKgBtBukmLcmBrhwt+W89u+/4+Be37xl81TOzpatnR1de1wXbcCvoQv5788U8rylMZmOjkceeOU2fFlFXWRowKhox0dTfuqVU133dIRLRnSEWAlAIVl2OM1xWyc0z1Yxs48umdPZU82m92dyWR2tLcbfG9/bjbSqd/627+u/fMb22/eU61qmmEGXCFMsdDV1LaFhFEtAhG4jhSep5K4RNOgWwx0B/MKTrC7OHo4YcxQC0E80tEMWERgK5ikVK6UOwZyA3Oq1ar62jm4D22c/vzfb/j2g3vWfqtYzAVmwAW4mMFXHxpi+CZ9BcVQmjxhVkv5D2F0AOShnDtakCuNmf/ee//tv3ruV+fOnDT1zR3FHduz2WwXqm/33LZlOnLw7YmOGSFX6bgZifVRuoyU3DGdSU9ahoEhBCh/QCTGlRJQCc22aP1Xj0UlYEpcOsTbG7x+1nceu+ILbY1TNlUbqgrPXgmKJ8K+kCBe5DKgLJa6C4dnjZg73kcEFSQAGS6+CMZzAIJcxmnAfJW4sA4//JMtfX19e5/f+fzujtbpr31v3RevfLnwjyurxdGAVXXwyKImSbUdxiejxKcaVnGYGkJIn/jYO9W+qcHw5Fz3TH/U39v2fPkfjz615cFT1y5b2z40NLQAK3uw6owz/JTFAglYvQ95dyWLM0GPFeyS8Acf+VIJPkwQCZeECqyxSEjc3PG8Z0FErBrCrGUQwPrArAWXoibARowRHY6CcYVdVXTle5v6xazPffuxa87/4Hzn/PO/U0F7lyvE1vQkfA4D2UMpXFxzHDgXcV0oXmwCkV3k0wVsD40StfNwIh2N+ajbQPWbDSfbKBft5JTWVnLRdU/c8OQm+xgvqIBSl8Nli10R6XuK195KNPJMOskyCY0kUNeOCKzvwmxX0axQ2IRxymTQVWQtdz7Ui3LkSYsXL3ZHQxHyjBA8Q4jWoGm6HmUWI6hwnXUwDh5CaivIZBRMkIIXOVz6/ZqVFRKPY5kA27Vya/yTT74DhVEg8vC291X16ZMa+JNvbDr99y++9cPeXE5wbGBSahgpajApOUMWeDAX0imeSiZYg26gqohouq75SV3zIVEagptQAaYYnxSHe0yrj8eAwDEmeEiNljSkAZKVhjGn1Dy7bYp+x/OrL3yq74VVrl3xTW5iB8UAUbsjCeDceZQmOccuqUldiSVrSc6JJRnlBJMstKsDoTFHk0MDg2LN2/f87IntT8xa0rGkfNurt/kQSWGCIxtkGZqZxa4b8rXHYkE4HyAhHbKYYJAX2Kck97VAczRheBpLBBrliOUogFNNxRCbukY5t0vVYEvuxS+/sHHttA7aUR51yli8QRCb4CmpMVMYAbdZENVHhfte2L+wE9VXOmAs7WvE9DVpBpo0PC0gri6IqwVSxX4Mw2gyX3nla957l7y3/O1nrr7ktdK/rrJLJU+jBthfsZOqPpCBFFLzmN6scTNlMq6bsKgYMwXUWdDJoMBSNc6US848IyiIgba/bb/xLiDJkb4Nqx0JRT4yopUaR3UqhzBynWvH6W2UrKuNV0xq2HK1c2rfoBEdmouCUIu1pHQyKUnemjMtPZJO6mRbX5H3jNEThqvS8D1HmWqh2R/JWjHCq3ZVvLzd+5jjyRspJfvwSMYt2986rlN7quhZilp3YETM3DYkZgtVBKTFOiES3vn0LOnvyLCtVVsKXTOIoSsrRDpBwKY0ZMeQOfJ9M6gGgTujs2nktvt2nPzEFucSn4mAMujbRdH+cJGRKdPS5rbyrYvnJf/a0my+nkxIN23wo159e+i4V7aMnLuvzNMqqIB3UWpVWPwCb+MA7/zYdWt/KqW8dO3aPXTy5CbuVRCV0H3heoGCQ4SRpvF2DOX6CNVRrxAewHCk0/gQsqshm3PorUQrCBYEESAso47lyxfR1aujElnfV1Lo37j/6W/0+NLSOFiphZIlV14btkYqpWWa2mzL3HbElEkPTMqknms2jZHp2ezsXaOD018c2LF0e2HkfQVR0g0MCsUDHjU3JowyVaMugCcTsgKhvFyTXErTYQbxSfbPG1Z/t0qq1AwZddUUhbepGJQNwiaZk/sOa59794Km2W+28man4pdTW3Pb5mwa2XjBKB9eAPYy6hEGmwFPoQnuF7xc9sa1P7lRSvmBNRvXhNTdrmqdRktLtCAwCBH2WI28ppmonEAqeJqwhJcl3Gt4pT0zpWpxUw5W9yaL+vBxNikyWtWwxIXaLeGkQHJbVviQdV/vHdcQQq4pjG0ok6lH7p7eePhzjSzLytzBy2X3Vbcf54oSgV8fEpiH5OOUc9JuzH9F03hR8AC+qnK7NMmppWfosZ1n4Ak7Ah7Qiy5aIx7b+szsP269+lbPKQpNmFxtqMpBQbPLgCeIlgim7Dosc9Q9LXrHC02sQ/OlmLw9t37xruFNF1X1vvYAmwUUEJUOgdRkkXoDie0zv/f4FV/49rt/t2pChDGMYIyXe4dc6VELjDNVTTgiKftxIdbaH1QFrGA8yRZ2Whs+d/aca675+Mxn4ounLSaff6NwxBU/evYH63dU3h8Afg4oN+aSsk5Uy8jBCm+98Q/bjiVkQfeKlZLdsGr+i9kUPzNf8nh7k+lf+uVnfr776ZEvO8TGdbkanBKvSrS5ney5J351DnTrco1pE6ot6tHyJVRwEh2UcLquG1W7HBCSYb//1+7rRx3EC+HxxaLGIG7XZYZr9F3zU//7n1+d9Sud0aInpB7l6ZE7Xrt159AfL1r55nff7vWOJbIqZCR0Q4nUPa8q1u+hH/3P2sEb37ds5lvQnJQW9+AKC7Xb18vLxZMbhWVogjCcqbqHVbCngA3ZCMdCvGXVEfcxShw3OKCz5k+bRm769zMf2jhaPBpqbJiCdYK3CHLLjK6zs6ZP/9G9l1/+fZPSUgwdRa1XFNm6+Z8vPXf0D57752+2+72Hp7DdIaAbTS71pVDt46KvIPdAAGPMHRNTm6cOX/fv73+5Lxidwn0WCIYeDiPuKqxmUjbbmnH3w5/92w1JLdFVCapoX2Qu3YSWGK4G1R+9/88f+cletvNK4geChfLmeBPu2yIY1gfPu+Hh607+8Xk/f7avrw8TEH0ECCPc1ZpHq9T9IjsBOUqZDFiKT31h+WFfuP6CBZ/Gex6NOifwpTy27Z+n/fWtn/yin+yez1wdas5RZZ6S8mVu1ZXDrOt8KeV31q5dWSwW+5q+feYtvyCE3A6cTTVPjvjUv5f8rerlGOILNb1iWMyMkovmX/+tixdf+ljOUXU3iJOh2fHehltSDFumF1SAP2BXPfqhVXnSk6KejtQjx3so5QMeBMkGU+sUR9318/c+iOAUsCo0woW8wan1u+29fT+98fWLbh2w3jonKBNkZkFhDyOTuwVXbnRe+HjPaM+dMB7ighsVTwkVh8KBWUPwSWST9gebAaVHPCLBRKM2xAmHiohSTo+cZm5/7Y6zzqeU7vzSJyakvuhRh2fellJeePwl9z/zSpdzHEDQyE6EC4xyukQhSLDXdo8uIYTct3btWrU6FcrwWfHeLu0d9RwGlC6wCLF6sYocuKRcUeF7EHjWMOfxJL/88t/Jm276pFsoOHTO1Dbnz/dvO3dnv32cykdp4QuFw0YTsAxOnKFd99Cvz/w5/TUhK363PtnXN9YmdIgiCF51BV8wJ7vnzbtOv2ruBU/8dWcuNY8yRwBEp9qFa6K7zPidj+645H3L2t/MlYloQNcrZS/BI9arSCw2VggaFwwaf+4UQjOI66k4JPoKC8D450IFeEQkDrJ2Zx94dcPF5XJJMguV1pF+BlPRCplKpNi5s6Zdff8Vn76FXnGFykQALbdw40Zg322f+pk8la0XnHDKhgtOOOXjR//o6rt3V4YWJCUT0CNW944VoJUQb+hAgINASk5MzUQQWntx7xvnB54vNcEQAIv3kiAgQptrzHjq4c/+7ePon7tevCu7b2Rf0mVucUdX0V90yyIgC6EL8aVzbn/fvC628yzq8YByeMWMcI3JIBGQTcObPqUz/dnhYDho4EkVBQTlZAwymnAAOKgFNMumbL/nfa+8F+xaqzek0ie1HVuBvWRrsvHs+RdsO6blXd+86pEzbxkrD7Ry3wgtBVUJqXRDaDUodL7S8/ysZctWvjFcGG52R/qYI9xE2kwb+6q9klmI943jYTHj4EUgDLW1+LKWc0foy3vWWpObZluO7zAzIIFgut+kJ31NY1pHYnrlrYFdC4fKOy4MbCH0EAQdWmSaDIwGpjXL2at//t4Hrxwu9EwrO9VZppZ2LVMzCqWqHaxw+OzO5r1Syo9c9p+TXxlLbF1EyybiRGHgzuPCTo6l79n803NV2SS+StUA4r7hQIry4mE0PVSbjS3TpQsyIdRx+SbamDGlBrLYmhZ8eJ5aIAJBGlMaveKcTkhN7wPQKLpXTXb89NMV1Nn9zHmzf9KcTiCNFzqisXanCKjr+6RUtFFTQNe1ozYntFwuv/wD6loqxY95X1N0jsxvFcaMYlTh/ervTX/3u8ullFV/tKJmo//g0z3vHStD8BYfwBQJpyhhujYl6b7x/B3v/nk8jlZevtRu6khoU1p5MW3QfEcjrezqs9OE6N0fPqX1+ymLR5p0OBBDRYFNmeztq7wPJixj8MThybkJTpgBdaUQdh2O0tgdiGoO43ZTT+YRSLIFbrg11A/yKNOB2HlSFVzRhQvbsGyqZMsbW3ZM3j5cWBhux0jDhhkEhTK1EmxxNnPvv774+VvmrFgBhmUmV68Wq5cvF4Anp1KpUtpMF2ZnWp0d/V2gdEj87D2X/iYjNQ/aSHE1eygYI8PERzgiVJq1aI/6bdk2+8Ynbj9zyM1NAhNWaO7Cy0NSW7AUTQSfP/piKHGpF/nECZ9wmxPNcnbT7MJ7lywpv3rFq97Y2BhQky1XHv/5G5MsiZBn2AAwmXUG94MUg8J73MBNTGmeoor4lDI49QtwicclR+PUkpRGMkmPnXLO/2mUF09fcTpfvmh5kMk2lzNGazA9PT21p3/HtJaW9r1HTzrzbzylQdNODY7xYl8aaBmpbR/bsECt0b6vW1bWSLO0zkmCGOBf5ZHXWWsVuAuhzm6SGiqGNqd9jtasNSfaebvZYLbLyekm1zSzIgf2Mov0/3Pjj5e5JG8xGAbqchTCopKZAdPdhsK1h/8CquU+F0aqmbckM2ZGM1l6aPoka5CsUkBC9EV1bubkrxs8QwQFLUoYO9KYRjzHlt2FrefWyl4tXcX3Irsq1vSLX0ISFaSVhLx627Eh+9GaNbAEFaw02tbC4ag05aDoxdi0BtF9xfKFG7Z0DbfsV4qrvtatO0Nd68pPHLlhctrrJsxikkVZi2hFDZwSSZgSACJOwmuoc2+77VjA52XaoimJDKgqV4+taOSIdOye+El9bv97I+gGazJBPcUlt72/vBgWaJwMjSreJTagkw5vfqBUDSb17+yfVB4uTx0ZcRaUhv2ZuRxv55I3Cpc3tJqUDw0NNfzgyvnb2tP+GBEMKTw1i3E7RN4H8s7MffsK8ztakygHxWJogjs1ZAOLy58n0jAESvgsfGbGqtjPkxI9GI7K8eUgzN+qfwMxGpaMn+GvCvkLyF2vbOms6AYHFbfavaPFEVqyzSIQFx+96KcV1yObVq1y43Leui9gTxSPQkemxdjZt6/hzGNO2rxk6py1gQW9W4CnarHXaMioF1GEORURFnFv6Nt+uu0CxBR5ZGFaQ0pO6YyGqW8uP/bDI/vy+xRxyFh1rF1y2VqpVDoHZaljVI5O10xtUW++94jzjzq/1JGc0iUNlbWLJGgJoz6VVel03PrSrfObEk0ZRoNIuNV1kDZV0Q01PqMSflNqSa157IYTvn+fIAFZt2qdGkSar6lSYvAvNiWayl7RS5w198I3LC3lURNFuApzEQvsKnxET2kfLGXFS6fagXOhcx0wYyZcZDQiNWyFgWOEGlTqloFCJzU+hbQgngJq+RQYu6GTgjQjJcq91Ybtfe+W0kUKK1I8VV8CWf5mbdqz86ctcbeP7GvnvMF3iNYcENIZaIlGQtob0Z7g05RSNn3rtJ/usqrNfYEQKCUKtwSYAgGh5WDsKA7lnVKJ0ilTssHSjz4cObPwKGt2qkojpROKwiq19tXexBEzTXfPnjxdunS2s+j8+8bln2sUGbCkDNKQslCBNViuujIqc0XZqiI8iEgxBApZYNY3JpODZMSfRlXSIfaIQ+XuiqP8lezO0VEfJY5BUyD2vjpiLl26wP3MDS+ACC1aEFRgNEwy6Kj8UO+rP7t5c2JxR4dioImX9UwmA9bdwsyZMysbXx2bXHHZVIL4HYLJ8XKuWIXL5L+v29fMu/iRq0JGlEAngcdRV6NsPqapIDlODCTzIeMwWEatQ2g2h++iIm3CoUbyhW2jc5ZPnflorjSabG1odsuunwFpbxQAGN9FkGwPM0co8832FYsmcRyi64muSrXaFL5LxPuiPh/2AcSppndYTRAy6e0t8i57SDtx9uzyx2/551wb7ckRKw1xOxQBUEJYk6Ftv/acd++8qFBoBSVWTLsVU9ulg8DNBwHYkQWsBTEygpqH8rzOGfe/afed7QFQFIpAht2vXBmFsEz5vj/kj5V10km8vYXeTgVzQUJRrVxhSoZ7lPSQvvmn3PLBFwTzoU2Gac4kQfwELYisLmIeAlMHmT9R9otJrinlvHGFYciEMqEXgvIJhJB/DxbGyp3NM0oVrzyiAENRuEtNKNyBI3ufBr1V2+a9mzNTm6f6gDczxoYBQ7aE5XNXaEEmMCcZM8qMmEVNrzQTLwqXh7A96tg20bI0iyr4vr6+fhXTd129oaHBL+fLFoN7BKspCiaqLuOqQpKYNAUrnedye6qunuxWnmS0F5TLZUBu0UpOzhucoekg1oo2rJBbTRNFRnrFznddvGbJW1DSU2NJED2Aa4nCK4C8MN3gz8F1FpSU7eFGoZJ94GkMn0U3Ocn5gxxkDJMidjwP6oshYqheuT3MK2csHVHq7PTJeiaVSnmNkxQ7UsgnII39d4fQHvJpOZ1gQz9f/Upyfuf86U6xqFJ8EfMN6s/dgIEohPQDk6FKwJDmjxaiaElHuBG/aaGVgAYZsAikvUmTDLxR0XEhgTMeXIv9aTVAw/fiDTzT6DEvZVFLpTwxmYrFovunP83cAiDS137xamqsyqxxdd/x3RZTuq8km0hVEdkSeLtI7YUgKJVRVfLk6kv1FWayGw2/cOuKMRq+55Dt28ZcetYsBOkUdPChp7snAc4VOuAx10q4IGJeakxgjLQarkC5Yjelid33PjTAVCU9lvU6ZWW1mBFJUpaGwqSMpvmN7YkEgnL9gSdmheNYgT/iDJIEGrI9rSYFDSp+B0sbQDYqFmnOOWrxeXehsGk/FmlVLnvJn374kl9ykGbUarQqKtvAlA4ocUkCBVWhBSOTjJNOFQvFfK9F+zHBKSl61XSJldMqLSfCxUUB8xCKwn4Q4qqimwiiMZjMoSy46vcQhAM0JenP9yc1qvUKVUVExOX3ndundmWMJHhjkR4r4rVMcFxwfiadKAjGAhoE+VQqdUDJdm+ht8yZyao4J5YjD1064tgOybAGxXzc2dk5gT/v6T1PDzGkAcY3y5gAjqq0p/J3iWHbltfRlAYd/wEHdnjO5HRgKTjRAXAMKSsVOJURlxQzrpfPIIuisiCIDgBnEZG2h5xP4b0xxagP8wZAQEVOrKYqwJRVzwmpoTxPlRcr2ySOPMf+ds1KCGea7wfCd4hDrfEEX7xFjafQFSFTAPeQl6qCJnmWmybVQG4XCZ+4mqZVdB0oSaDLEUyP/NqaQGrsU9fa0g+ECHQZshu7ruJcDOEGqpipbmaoMaMKntTMhs0MwteoIk6p/aJIpPMDr4Z8dg1Wa4DBE1bTxsim6MaAPAAQCYlhUIgHknIuqW5IynUgD5Vusvq9qrAKhNJRjiL+KsgXqvQqEJHtoBUkPX3FU4CN0yTqDEKp0nhOT3AZdF25b4KZGkRUFFLR1PHCIRBpApZJafACRCZQ3FHxfd/xHSW95znlipROFQUH4YIdZ5AgEFoLAqhngxKSn0gkyuAVYIyVp02bVrP/cP/lq8N4kC6pLqoK9xBr2o5vJuEzGaC7j54uqxvGZAWYDMELarGNTwAmCHqe1GOCeUxBgzTBJMjwuBJtNNSXwQ1p6IbUYB2oyYWJDmgsVTAifDelkQDSZ+GKhUqdCIoJSDuGCbFwnKgIewgjxiApBYK6mpSeroeq05Fac+07lCSQkEDMOQRQ1uRAo/2vRjoTnRO2kWBCU85hSEtY8/Rg0ApPEIsngDVoGHKHjPFYUe1L3XvHaE9S+DIpXGwSdeRauCyMPsYlJ4agARfU1YUGDnHNkga3pC5NyTxDMteQ+BvzDEEJ0MTwscJKTjU2wY+G+gcEt3TuYZdGcjuqvKuj8gqTyQS8q8qKEJoBgk4QI4evGQ76WjImCvqFO5zws0kupyZ05nkh7rPuoK6LUFBI2Dk+JdQcHpd+jFaiqO6cmtJESZQUHOU3JKVwMxOKg1RHKel6GtpF2XTSjJWElKoPFgbbtoPe4tIwjBdIhwa2VIuIevR4QQoHrAQ+z69C0pxItxLJ2uN3GCC+FL4Lzmz8jUrPZgKiplhcgLDDHPZRRO1JqWmksS2l4jbrNg3F/nkkhR3tWnWMSZFCtqI8t8NJqmIumqqhDbeZcYhvtBKKgJg6xSDTiWFCyx7R/XzC5IIq16pGPIQUIRWuQ/ryOehpzNB1PRmVD3tIDkQDvC7AEWZoYmXjhKYdH0FZ/BApqlyFyEqgzLMrZp0GsvQrvivd+Fkjayo6AhFQxweeCF8ecaVPXOoRl3hgn5aOjzIiaCa7kAMkHvWlR3zonks/8KUPBjllrQbSEU4RiKJNZJMai5xwUwRhEKwea4cdVUgfG1IpJoyV0lBxqnCzGf+OkV3b5KO2DtHIwEKHaKSYRTq8xUZ1XlJPIjkYt0ltp1GCzSIglmYCPNXe3twOy7nWxvX3bhRNAQ2orxbeCM0dglPDIDKAskHgUBTt+YFLA+pRn7rEZyDRcAF5Jp6rCgpp4LvUFyBRcKVgvhSaSzzhStd1BNxgVdxEXDUp9fHCoDqwi3p3rBNhahu8ZK6L/K+ii1b4/biB4kVB/YLpMBtV7j+ft4NY+io2XmL4Etc0JdQ5vi3H14pX4LrAmS2l0+pw5lJoqqmkiwbsWlxLEPtotRsom7KZKj5Wgu0P9FKhzalm09ronfRhCzBwANBi0F4oPaieqDUF1VKdUmYSTIBwJaSE64bK5SKopOZNuDyGYX1MaLUrQQfZw+qrTW1NksltGcUqdfrCNgp8F1an2oIWx07idoSboXJztjSIQQsRNiJS8ogWy8hFCjUisVCRjMkwyBrhT2p26IMZnL9t6joKyaJQCgZ0ABVROVb0Dt/R0zNrelvb1mq1Cg4Fj1KHuy5FhSYWpANydmluyuN/cs05WDcMRNmieozQsowHEGVIW0anjJWr1T1M50epYY2sZ7TLQrY9oRskaSShkB1eSpeE65RomqQClJGoAISboKksI5Gqu0J31vVQOStQzMMTDQmSTaSHESRcuGghwaKAkp+wdiZ2K0NjFKEsn3jwBSvE8xtqWkiHOGqxitpQDTHIdfHciRWl6jBCSwZuUIi7q1mBCogEuR6PpFnVZoe6X2trcjCgwT7NZAuQ6FddF7krGB6WkSaJbEYxM6qpBx80VM2gvu0SpxIiSkDfijgE3EQVf2GSeKg6jlYqQU2i+NuZ5YdopLoilPhh1JCLIuUh+QVyRiD+UA6vpkqqVMi1bjIT1P8wAEDUKjLAyjB1axVoMZFpFDrDryLK4XgB2C9IGaU+WIIyrQL+Smh7qZcsu4HmqKKiKAceriVR+jE8NK4jPiJBcw22GEXYiWKflcuWCRDL3PCZ+cU/PbJzpN+h7fFyFS3QqPLRPnJ803+//+nFq7f0OUnDtIaELwZKOXvEBY2mcuO40DgMAF1Ay8RA2JPJppRONQ2VGYFo8BVZjug7bHbqeZyzbOUysW4VIV61CnbBsAlq5RHhmqVcDgIAk2UT4mElk+Ns+gphEyPtomYKFxaApglxk1wEsiUNPiXSdsLcKVv/vbeflJwQeRDStiuetGDIC8yV9z581l+++LkH/vnaa40fOvJIzphvRCzFrLe3N7YWKL3iCk5uu837v0fvP+GmNx8+gym8CSrGxmsxQgOLUsKlSluuWrlKOZEtja09+4qjR6E3YPqqEAyhwiWCLWyY9vZvPvC9724f3eulk2nJNN0s2IV9qN+teK6EbaCcbS0tuNSkEE7Clz7PGtlUwtQnaVTPwubRDF6YlGx9aaVcSY+97diwvZR7UDegarmkyOkEoI0qzISpaSiHPfhRS6bGUPFakBJlw6oI7SDrgRuWvdQb3VFgEatiWC5ApEhYB9Lho0FXAGBp+J++/5SeEutZAHS1shzDLkRins3OvOvFr7/nTz95o/+pHGKqDNHFQIJmkxsGbUmQ9ByQQQXEd3zmDksZlH3ikxRvTDo+GeUJK0t0yYvVoV4wxHLfD8OKauuISqVqz67eHoNHRVI0Pwh8qCfLkvo5LGuJ2zrezWuwtZAGTBMpTECw28Sbt3IT1P28g62q8aViEyvEs+mCm9REtJl4ilNHMScDdKLX+VXRosBAK4boH6l6nqubjQli4/cuAiFgtVXpPNUJKZMNHXHpU6+TnPkeQqsYPmqFgQEPi/LpjSOHtUxreXNpW27YshpBhgCz8qAyb9HEQb48QYjbWvXsDFciRJpNiFYoFsNKvFXR+OI8oYBdNauglnoEjBVMQSrG4gbM17VSiEfyPIVLr/WPGiCqiRhhukFKtsgTYowJUZykabo5PDzcctV5p2768ZOv7xjV9DkEdbZShJW1nDOPaOKR7v5Lv3vvQ//45SeXP3bC1m3zOWNZwJsNk5T/9MorgqxYAX8cPph391P/nvzN9Y/eUaTC0pGaiWqAQkcm3LWkECjB0w2srMuJRtdQ/9Orv/VEQqbO8wso+Y2sOV/V28hd1d7ZtmH3v/vI00B+0kEqXhtNGa+8QxsjWIpNZk65UPaw50HExUpYI9WgGqbGQ4VpPA8KhhQyUM3/WsxFLWKRf6qjk8AkG7M718x3dY1iUSGro9WgbqCGgz/iCghHbeQF4kCmAPVLCODFbl3UVfGmh0ihQClZ/f1qx6pwNHQm5rw17Gw+kwKfGONQ4ce7RO7MvbWQ8eSr7572/gMk25DBIC45HMLrpEIaXenaRsrYbdFk3paVuWSEDNLW8bGsgooJDbsQcjux4ROluGO/NrZP8HdHc4QwXKbpaotWaQ41/+Nqx8hXUqWwkpZsyWdMNVBhpcgdkJeOA4sovZY8pqEKx0i8CESAiGjHCZ+GUh/qt2XEJligSDAc162UYCLVovmxxa1AjyD+JSODw7JQLut2sey7oxXhlf2gkskY7rZt28ybHtpmlh1hHLMg82RSjd5xKwVZPxAC7Bxl0y+4Yd3VltW4d1v3SHLrroGZjz7al3r++a7EQ9uk+cendls/W92VuHn1hnSo+zi66GM3PPvTp1/tPTahW8W+EW9kMO+N9o+WQVaC4JG+PuSiYMyCCR9HqKL3jwYZ/HEt1NwITMbKWoMCHFnpdJhsq1V3xjsgcA8IGJRVUDEnOCmA8q7iqniLf2R744MK74raEUDlkGnQdPhfdITqid+uf33Np35z++dmT5mSmt7REUxpayu1Zdv6vnfRRS5ZtQoRV37N6j9c8p0XH14/WCks0qou7gj8sgru1WKcUeglHI+Sr//6eowz7eLF71mbcrjrB6ihCPEARKdUT3DhWn7qiw987zvwRoaH8207R3vaXt/9euND2x4yV29YbeDrpoduMlc8dTPaGLRgC77y4NduennHC0tSLOUM28OjeVEdKDpFiMmjbfkfr14BrkWNaFSoSnG0marGjCNyanJiImJcOQg4S12HNYl2VteATqSaVEHCVAC90KwZd02VaUSJE7oeXnzOq2Q9l+vX661mpxZ4JKxCi2EuYc2HsFIaGfH2GhD2zpfzqfVyvb56wwrjpm1fRKrZwPiUcoPuSzf57mkffd10GqVgkACJ6kY0OJu6sLWx7FfuX3ZjijeoBWD1hpvTj77xaGp9z/rkluEtib5yX6F7cIe/t7SDGZrRUi16S1asvfSH/f17p5MWoj+1+4/WBrnB2LbtIZP/7Gc/2wazLlrha2ZL7Jviv1xjZKzkIlrd+4lf9PBXb+vwYMqA+jBAuCOGLkfBMeVhhJBa/LZjakuzBpbiSNBiwgqYDVkV52g6tnkEpuJgk7oWIpPxOB+cNq0d5JnqhA/+fK0Gjcnv3Py6v25fgZSRdBgvwlKlVV0F87j7Hus+/CPvmR6WY0WHqROytUd2tLW1HX5pBw8qlUrq51fNeO3hl14a2uPrrRSePe6pHDKquQELHt/sXvaea9cFv7/yuBXTD0t2LZjtLSYkAB1VgcybicUJfvucex/bduq37tjy6a39fN7G7p2nPP7D5LtnzJi0c/lyqa1ZQwOwFAdBMG8uxDUbGooJgzcKbD2Rgk+dVaQQ1MJXQcEcuBSHh8tTPa14XHODxwRYgZhie4mqghFP8wkLGBkbc8sQ2L3wwr8XVq9eDpbetx3HmXfbJ9/30nErflXcV6mmGMZ3WImqwP3MdeSA8Br+8Nrbtz++ZcXbR0/reFrPWM9mDaP8pXv+1NpTKS456sZvntVXGl1UtUvEYlQIHf4JzD30P3AXyqZSTFKIdCUSiV1j1WrnvMWLE3v69rBzDj9t7Ldr//LMUDB6lqrg1CJXA7UGrhS7nX1nv/uWj97y+FX3fLOVNIDVetYSsgS8l4i7SLJIcQM2rt/11pIfPP2DL22ovHXqy93Pn/TsF579wNTM1N6Y6KOrqyuhN+tHnLHoMjRkd2O6uZkVw0RwSFpEx4PlAL4QsvvL675MVi9fLfr7+xGAnhPL1s2fP5+7JZc1pFKNKNBFIBLrt3puFXeWEj7pjvxOUKv194/2HzZv3jyXFAgdme9pizMtAXVlD6NsBpIOkkqmllAQhVs+ebX/iXOvpN/9gyMrbaRKTly6aCnmmZMbLnW0tk72CwVpl+wu/fTD3/2Pu7dM/+QAz52lyqWj8cLNUAagaOy66JMPnChs2//T8kVf4KREXGKQPmIQcFoqNmcp5c5HtjzygX9u+cn3+9gbi/e+suXkWz+w7lMnTj4/Wyr4hjl5LuVApa1atYqkLEbmfeChkLIlzmnESwJS3iyKfJ7+VPiXVSvA8SEkYNXYqWIzJjSH1GdAcoPVt1yuYEWH2Aq+6s2ZWEvC4CG7RmQF1uVV1PSuzZQwzBiJceCXSxc1bjafHIZFwBEXDD3ZUKVusEzN79zVv+ayG1750ruOa9zaO1Rl2/aVD3tjr7z8syueKT3+23d9s6svb+YCrzprasv2E+aaN/dtoKtcF2K3iERGwU3GtELFFU9tDj535g1PvP+4eU0Pnn1q21tHzWnantarlWc3j5qbd+SWvLI9d+HrO0ZOLCJOY2re2wP6zHd/7e2HB7bnzps0L6TJhlpRqVRiPmMmIQ5atjEIvBghWud2wdBFViAUqgl/WYWJ74CuRSXiYTtOcGwDgqSwZRBYISSO8OP8zXuHSodNb918yfFH/P43z7/55bKUPtM0rloshDFTLnxZChy5MXCP2L5r9xGWQb9gWZT4miQ+J8RzK/AZhc5AEiEY9RFJkmFePA5s4l+onA1h46rgzPc8PWEprYvGL574iXu++MiPznQ0hxpMI2qIYAdxCRNeIPY43ctPuOncM+Z2zL77pGnHvnhU+xHZNqtleOvYts6tAzsXvDn85nu2jG48pyIKCb3KvREtd8zZt591v5QStQiKgHaaPU2UZCmwmIYYD2hDlApTVEkQBQGpohiTGhNpoyEof7PA6EVKx1JRm0fiqrpPfd0jSjMkTCGr5Hs8gqMUJkXEuuwntGSwd2Q7YlQI9POAVjGgKw36pJ6c0z1dpYVr9aVEc0tS9lq7L7zskTP+dNfbN/91lrlo9r7S7hmv5h5aPGgPLFp57J1XTzGmFB1DKVdX/rD+56se6dp7VtkZlSCbCHFxYOugzKu6opu9+tFPPXDce4/uWLZuUfMJz81pXMiZlrL+uOE2f1vf+oVfePhDH+sZ3fQ+WwxTxrg32LTttOsefd/dP3vvf64VyOwEZWdiXFX5sjUhopqTg9/rqnDsIAcW+RggEk5FFazCZq2BHSlk0J2wABzkKuGiosBt4Y9hs8XXjX458VB/+OCymRu/c8e2vb0jcnqEZlEOeVhgUZUb+oK23Tnzb0/uHCGu45Cca5Gqr5OpqfwA8arZpjQbqgji9PeXtHt+dNbvjvrEvz7yVh85imnIG2qohleYfywLvmsHO2138q7XxOX/2ZQnrWlY7oFfdQgfKUvFDKx8HwPWpqZTzQ7eHmTzzvvWC//asXHwg3MXtW+HJgRIfxhFsI/qlklTgMdMyOpEESiB7ITvTghW6XrCA23cOGdc+PlYNQ9GTUAiqNnCUOgVx2HNVrC3d0j/8WcvuGNDV1/LU7nyJ6uEIKKF4FMINEGu0Epgg1c0SbbnMcd1FD0M40Lq4BfTFIxK+k5VABfIITyG0mVl3EWcAeFd0ZkBhfow5yxtpGnPaA85++jTXjtl06O3Ppl77arAEx7zKcqjlSmPmQguh8HyWOvI4FtffGN48xfTMiFF4EECPFMKiiTgrmJS4ZQLArIEh/h9Wt9x773jrIdWr199zkXHXlTY6G6U08Q0cOchzpCxaLrBV2GT0LmPKiQj0z9q703hv0BGm8/nDQigqPYOGac8TWi+coVrEJkIuFa7BGUw9ITQfBUhA1tHoDJowfy2o1/cN/D2ySgfCGPd8NhwLUYr5bLszr7xqft3b/uUdKh0hU0DrUySmRR5pv+fkz955FffJgWRAjX8Z469/pmv/OtDv91qPnu1KAeuZmhGDOQCVF54NBjzupue2PuX85/vvff8jJUNNGoKX7p61QN/XEU9rGlYsH/1IEf87frLJ371P+ff8qNl//imE8i+/SKjkQMep6Wi34WFRgfO4oa0EcGl4tRX6MerHJQA6EIRVDSnxgseDnlESJUw9lAPztm/jHLCsUJd9+wjk/dawKkoyG3deQJxw4osVwvB3pGq7Cu6supWAiLG/LLH2n/5z+7WTCbTH5SFTyw/SYiT/cu3jl81p90soj6ZUA38aVEcQ6FeNICRpPSDou3K3UMO2TXo8b58VbqBHVA1kUyqqOTUlEAMrmi/2l057Iu/ePsv8O8ef7zkSsm9wLax/TtCMEBLIyreOiodVQ438d2hQ4Dkq7B9BiyE+kwMUFKNiJRSglSruqK3W75pTW2VgcBMSzbh7cnlyv9Z9T+fO3lS060Z0+DC9bCm+UDTK6AqXgKTnlGO9QEiMqqYHKJkIQ9X4BicmskEe8/shf2cIIkSBxVBggM4Byg5w5iTUkCxbdeRMmjlrSPdg92FOz7+/a8enZ7zELW47oMNxQ+DcAj6EZ0xruuSOtL3Ko4crozQUSefGS2PEtBLazaFk4oshQIMMoNBBcreJ3pO/NuGex7Czv6c/Vyk7K2ox01NGBkf61dE6xlSqIU7vS40veQWNLIwHMRQoIp1NBDnErYSiw0RKjFcuV7iO8IXAIYcT6JYDyOpJ0VxtNhw2VFfX9vMp+ZIUmhaKFYUZWLAqs0oKYvAqRalRwuUEl/wwHLwBBtHXl6q0rWi5K1ctlaQFQH72fv/ef1UdtQTPEsMqklP+R9w1ZDsI0wzuClNinidLfNenzbi7NFzpT5pV5xAenqAALyiogkzQZyUuLPVf/borz/zqaubmppG4PyplyxVA4NT31RlxDHktjZIkX2r2UnRsUrmSy4n1OcSQEe1GIT052EcwCcmV7XsmYB7/58ML6mqQqR4+NYqF0OTPQKBHZCNkHIl6esbm33j9cc9cfQUbSu6gFLmhywWMQuRmmwa1Til3EAZB1xeUiJJsva1QRRNlQsFB2J+Rs9gNXXE4c3dv/3i/Mvnt1v7ND2DAnYA6UOSRmVuKhtNscQq6AIL4IqrCJ9ibIktJWDKAj9gLGMdNi3R9a4lk1ZgTh9xRkLpXUqlumx6VJBqTR6yRoUW3ovrFrF0vZauNU3TIcSsKBJmQKgF8lUhsjP0OHSi6UkALg9QwapWq2DDDiaZZrC5pyf7+Dcv/+aVS+atmpkySkSgol4o3jTJEKxliAupL4V/V4pyQeC5sD2Y1pZOO2fPOvzL5x153I95xsJ4jDhoo7XJE0hkKJcBWSzofmhBYFdF1UtqyWoulzPu/dTNVy/Wpt8vmdQVjgDWBCgruXJPFYutwhgCrEs0aSjCWVgnsFeihCVQvBbxtQbdaqMdg0d1LrkJjXHK0lNCjClXRVWCBBoLvGibQQxApXNRXKIke2DBGqeT09Ui3tLS4gNSHwvsuAaezQA/ea3sTJ0fB9ARRtEYMaI6GqpwwWFvcslFlRRkU1O254xpH/lepiFDiQHBXU0gyROW+FPCOK6gUepjV0OYnikU60il5zjUyDMnGW4wK9UZ1Vsv+O9HOs1F95itXPcpJBdpACJaRa6CbYyDCFWjzDcI8w3JsHYoAAcmKMamQs7CXgqsJs2cTA/f9cHDPv0n4pCsqsqJCrRgKaK8EKBKoEB8gsGhGcDG+Ci8xVJ0eiTUUeO/gNkGkLX6PPUJw3dU3aCOTlUc2RpXCj4hnieCZcZtG/0bgHacE/Guab5kTInBKOECTRG0koNdg3MfCs7izm8fd+Xs1sQmxhu49Bx1LTjZIeorGqrhthWgnNwLGO3qrZyvILVpTWET2pLc3zdQct9z8px1T31z4SnHd5J7G1MJTfo+IwEgdqiQ1KJr4gEU0756NCxoVPk8wideFXsEa81ktdMXttz3j18se9c3/2fxI2jnw1pb1WrLA6gEO4xpusO4DgwISn4gDR9QCty34Wtmyrd0CCnVoKySkIrncL2iSGspAl+qiNWXaH9NC6jGfTAm1avxKAitovDWFYX3rJaW7O59g9N/csn7//X4/1zy+Q/Pn3X3nGx2pC2TYYaV4AGjGia+n0xoMpPRAgP1DZrWpifsd0+e+fAtZ1/wkdWXXffLl7p3ejRhQv6nLgxNQ4xWGAxScvMwuU1p+qiYhNanNM3syMhI6r7P33LFyU0Lf9CYzNIgxTTPBCMlwuhoY8VdGaP1QspHUJkhcCIAUYR+esASzSk+q2nuo18+/uunrXrPd1bjs4vIIhTQmb6udEZHPGGPoSyDceZzHV+gcwW8j/g8LGwQ61ati5XKa0lJ/J4L7uq6dCyuecoU8migaSzgBq6Da6LPpG/oTNgBlA1BMakYiRQPZMpsLPflu9hlx19/e3sw90vJZCMHPlt6IIujEDgSGoWxo8AJQgbgYYKh5ZEyzZ/4wrZnl06fNJ2uXbsSs1iuWLkS87Vwxwee/9hMedJ3MkarFJqn+b6LfA84QJGGEBRUHCqKqtYc1JshIKqo+ARok7jLEk1pbVrq6Pt+c+YzHzx95tlvD1YGdexWoXlj0urhH/g3CNg4hRaDZhGKL0Y48V2iM7fZ1Kh01xF/3Too3qop5lFJUvCHmYYlSY/4FB1CDZN4roJE57k/Dget37XqFIR06TpNRHgaA9xYswhEbCnKuwNBhOc2ToSF1o5gbGyslC/7wdy5k/e8cHPig5/84cZfvbHDOm+46hHY8KFcgx8WJAkgbDXNspJaRxMZXDi74d6wlMFoRCVfRcogm9CqAwP5ROeizj0v3zHlgmt+8NKn175tX9M1xo4uyjSHTrMK9nllEvJ+YkcLYy9YoMGZ0pzUyLS2xDPLlkz/2a+vO+yBxb+O3ZtVIpfLQQ7eREgR9URMlwnLSFt4VE2WSOCDy1QnIJoAqswTShJOvfO+fXmtsZFZ0jE0nZpJrKEhhyeCeCbAEWqXcuzaOUjYqnNRxYeUJVSFET1vyVrWlj37rMNmznxzzVc++wDpL2VuWPfcuzYO9B2TL48enXMrSVdjXiKh21Pb2qqLW1p2ffL4k7sPnzFjE7QUXSlnfej2H84qg7IN8GVwcUZl3LhpUyqr5Pcgy97Q0KC7AjI5YVA+YZoksCxj18Au+rcrfvOtWx7/6wMPdD153aA3+IEScROe6xAK1nzhK4oOhYpW9WOhrhWoMXXNJK2Jlg2L2w7/5a1n//KO++Q/YitSvTdjzKIVn5MGMjzqDxX0dsqlQzn2yMDziXSIRsGrI0UyYzaM626FB40Xs6JbVJlhGchUxXaaHGJrhkEJT2sqoBc4ktMMJRXpaiazZPdwV0pQSFzyWEqPZXmL3zu8Z8pN7334Vyse+XzXZvLKz0rG0GzN9IliBgU9jIP3hYdGiZ4ySZK1kUnW/IcmZ6Z45bKbPuMMJWmnlKbx3ZVV8pOz//69B9++55Hn+v75f125zWd7yTwn3Ce+I4hXADll+CrUIoQnwPIZpuc1N0uyevvLs7Mn/Gzlqb9YVywWW8bGKlZKb5K8WCy2l8uETprkO5/82ubfN7aIRbqWMH1BfFdoFfQKN6zUnI7sq7ffL+eOjFRsw0ApgK2Y8086YvL900fFFtu1iwjZUp2ZmpbghHJ3ckIg3dfDGLWKxeIk9XhhWkGJTuCrWFT+wfCczswf38utIyWWAEE9T2ol6eseGnfB9GQPdvLe3t5MBkyjSGralCdbkhh0ENf0X3ihW5500rT+R35x8gfv+Mvusx5+tf9juwaLi0pSm+37WgIDS2d8pL3B2n7Y7Jb//uDzsx6c2tG6+RsXD6SnTjUhlRxHxUU22yt7e1nzht19mQUzJj9ECHnxrn/tOOz+p/uP6x7zTspX2Xy/Ktpsx0bQMTCthG/qrNCS1HZMmZR+5tx3zXrssvMmvw1L/fMfHp6S1U0vkQh8zq+Hph7QkjvCkZ30NZ2XjpxMf9WZ1ZKWlrRMPdFQdQjA3uXONp83tetjuZxs4rxsjip1m9xAyqDlo6dnbpvtWWkQ9AYikL5PqpouaCYt5YfOaPfu+j/Z3NdX1FIpGSsgV5qamlCJp/yxXbt2GcctOnygWHW0D9x2m/mvyy8f/dHF74WUJ2juUNugR/RwY0lNc/8VKCxD88beXaZTEcYxc+cGfYX8B1H0YKlS5ZC1X1IqWJrTFiuD99e7crnhk6ZMiYIdBGk9qBor9eHWRKv59r7t0xZPmbvhSnLJZx/bsvawf297+j17C/uWjZXz04kIppe9ErDx1EyagaHxEUMzulqtlrdPnX3Cs1cd98nnkFb80sgVDW28DVYeq1arUCJHNe0uljAVHwI3Mg92yMWGUKqVzGwwGpuZQWWBjo02mNO2FJ2nmvpL/XpSJIN8Pg9i2L2toVJyAP1Js7ONlkYqlVktS3+vCWe6L0Wp4hdHUEiBXJoeNFgnd55b/an4Q3tPoWeLKU2QyoiRkRGBrBICyQmtgfSO7pzR0TRrg58nV922+f9O3pp7YVFeDM6iXM4OCDUDlw5aqXRPa7rj+UXNZ66+8uhr1//Cn5Xq6emxIHuH+RPFNwJapmzIHErMaprV84HFF3/8iV0Pdj6++x9nVsTwmQU2NCcw5awqdXQU1BgW90zDHDFIem+LNf31Yw57z5MXzv34s1hXdva/mWixZvRqwaTAESOUDg8PHw7hSkJMmc2aO5MmVTElpdpUF0isOFLL5ytLNE3YseItBnWsZdeY0Uejain1eUBTIwg1XsIsFotz8HkAk1zXxXlYvwQhDt+4ccfeY4891ksnYqbnSBQTXrapkXzJbx8eHm4GTFHqyg1hAQuMpJ50LcvaSg5yGJwQR7l99uyBfW6bbsjG5vYGlO2CPt7P5XLp889/Yx9IWg52/sjIyDTTNFtzELkIAmdmR8uowhyEKMtJu7c4fHP3aIYyUpnS1lA48shkMZ3Q8g5WZyEb9/SNdJiE0HQaa0AorAmYuOd5+Ww2C/xC7YBMjqkqyCkpVsXsKO+O1k8S182UXFci/aUJrSc+F4s93jE+NK6RQtmfGi5qTmOp5CmRXD2bdb1CwbJtu2fSpEkTdDuvvHn1DRcec9RdHzx9Yc9bO/ctrgTCT5pGWtdc1PCDTMWuum7h12+8sfPX5503YRe95o5bL7t796t/8Eg1gD8RxzEEDYJENqF9fv6p1654/2V/oytXjpBoV9v/yOfzYPrtqFarTpn5fGZLB4KhqvoVr9g9MuK+PrY+QQPCm9qnVk5tWjRmUKMYEB87e2P/aP8U2L8t6RZUaGJ4aYSYVErba2xsnKD+rBNTxQ10TSclv3BM9GuUObNCodBhmqaGcVllbGxSOt1/sOfVYCxL74iobzDuK9GzZkiFzBiuDOfb2toOqXZdLBYXV71qEyecNjU1QU17X3St+YPlMm1PpbYZNFlRplSYVQ+BqFJCg3U+XtAwUFZOmcpnIWDm+8ae3Bu7j551Rk4VTRGdBNJhvdWxqXsLbzUETiDmtM5LdiSn4VmVqC3u2TPck+IaN92Su72+tB0BnwiJ6NBn3x5oqLor8tCdJGThBPP+Vw9v5586uX1MQHs9qo7hnDs9PdVUuVzQv/+HN9z/uWhNJAgKrAFiaHjqVaK7uxv6DwLluKZpYoXTwbWnxCco40uXLtUJWRGUqjFGITwXz3H7f35l9fQMNWYyJvxsjUiTGQZKtj0W8vWF6sL1AqYXXbSGrVmzUT780vb08UdMcRrbzDFKvWJPz7BfLgdJnmI8nTaCtWuXKSu7hjONDsUgZVmoNnZSBqlImZQDA/lUyaNZzfe4ZWn27MObDrYQUUJWs/7+XFNLOtTlDNNcVKlIR+KgqoPr/VVK1zDPDouANmzvoimebHSBymBMa0zqPmTCOOFwLdVNwhz5RdR3x/toxe+WGVv3DKe5xnhTyiC6Hqp6yyqwXFLB23DfK267jd92xRXeZ26+79J/7h364UM7/3XVNbf//ftzpk0GEAgWxMhwodDs+H6zBsYPzr23XnqpRtqaZBr50cP3nvmTZx+9EcWLyqkEsYxiEWIg2Wct1CjecNolD+XzeR/y3wTqyxNdRPUzdmMlPGBZGmCAhUKBF52i7hCHcMH9GZNnwDKceCwPSSP6in16iqdiBXK4SBqQhpaFeIlq6/F7riDM26TK5clJV9u0a3SHozFulAvFyZlMa9CUaFL6phQGe6UW15oAI1Yp85WU7hvaaRNuJoQUHSwAhpxSwXUzKZKAPqvS6f1iEbWxkc/nbYtbI9TzaM9Qj/SCYhtYu5qTzZV2mh1TgroAVy0idMXyFXIVXRUrdqvFPYxNmKDRUw4nNhrLos6Smcu8FXIFW7tyLVtH1sWQfAVjxvmOQ+aOjo6aAbNbXNcFLWA1rac9lebVwxLr2kMWCoXDo8CN5nnevubm5vzB8AKYxLZtT3EcB4VJClsQD3LmMcPMmtCoryfRqD/XKhQK0/BZLAJg/nGifB6CP7lcbvd+BBy1o6enJ9nc3AzyCeZRz6AeVXm/KLgYZLPZ7XEcZELnUSqxAwUB64RUBb50XXimaSKegXeRjY2New6lTDw8PDzVMIwM57zsOA52Dy2RSMSFUXZDQ8NOyIyvWrlyvGYlula5XJ7i+35auTahJDwi1yB7hYVQymazw4fCZIyOjk5PJBLKXkCEPhrofiR5PgAL4WDnAi47b968mXi32BLDwgnLzyuXLZ+xXvQt2vG1p7e1XPjAE2/tsiutxC1Dl4vMa0m9/u5FC/75kYULHjtpyXwQdSBug+eARQIhn+CBV56b+7c3X7vqmd1brh7zq7rFUPyL2C0WBFV55GsZk5/ZNPW393zqf7+yo6enfe6UKfsO1caxhRCbwYZheMo+MImoVquyqakJKEUZS8+rNg6ridS5xDSbieMggKgqZgGJx7V0Xa+mUqkDF5PoGBsbm8U5T3KIfpomHkaHgjQJd92hdDrdf6j+qVQq0yGnrhYP0AC4LoQTsNlhko41Nzd3HWL+oHhqHvQgeYqrCRBnzqI6n4FMJqMUnA9yrpbP52di7Mb3hjWDPm5ra4P1ORyptocLUcjsRtBuy5YtM4455phpYFKlVOHgFQUAvhIJpg8O5rsnWAhoQMMwHFQBNjU1HdS0Gx/vqs4evr9nmmHWxqSUgTBlpHfkoEVK+x1qV6+/Jl4gIuA46NHZ2YmdRCMOoVoSNYWhDD2+G4YxgU35YA+sacIF3AW7ZRS/CLljIjq1Qx0yAQHeQHH2QHrbohaq8FXkGD/XApyrVh3wALX64tAawLOGfjx19ThhdagjLs/GAhQNdHUNgHeihfCgx9KlS2WxWNSk4wTSMPB+HDgpKR3BDIOQTEZ9LmPq4sv/efzWroo9mRIPClayTJh8Y7R09Nbn1h991ytv/m/rP40uUwv2tiZMFgiXlO1qUGJkyijx5455FTMoO8TiiMoAiBuS0yC3UPV8NltmB++5+Ks/r1QqrW3ptPGObRyl9uLdzwGICSa0Q1B2rfgN43aOEa5qssE0rlQ0HtZpq8Uganc1SR3HOSDtWn9gvEO+XZWRE1ta0lLjiaj8d5RQPIRIbRQoRH8qhw0p1ehPmEeHUhKHcB8FCyrcc13qHmlISDefNyOJNl+IjAq81+/WtWPtWmoeDWZ4gli0lLpeDXRdJfkxN7B4Tvi8otUIIRPLly8XSuGd6AJ1l4goh/NCMzE+FEFk3QGXQUVDo8n1Toei1MZhKr/HFFgdonJi3tnZechJvX37djmtvV2qBcQ0FbjLdV0WlSSD2/BgDR+vsur5DdMIXIWiCQlV0JBovZis41DP7HkeCopiLgbcVzPhTmLnObhpp45EWcqSRSRXnzex6CGgj4EWylD9/w+MdCweEu0GK0P1RKi7eOiTMDLDTEDsP6oFCO31/7kfIH8+0XVUjkF4B5wRHFRo+OPePTl1wQ98985LH9qy6yM+DXyq64o/TKWkiRRV35fVwLUGAns+0+l8WciHqtmoArCgoCSJLkWgQwJd5TRCHgiNMVEhvmxJpLTLF5/xQ2JZXmF0tKkpldo/ej/hQGAagkxGMol+pdFYMrD7YTd8x5fVNM33feBcVFqzDjSoVpl3Ohf9Ei3aGqxbV4ZS8pRSljzEeKg/TNPEs4aiv9EiAVKZQ40lHBANQGPYijYL62cFUvJBsejqhiE9AQKHQx3Llkknn5fUddUOH62ySl/RtCzS0NCgxkbknk44ddGiRbCm8EdEhRGai3ESav6BhKv+84g+w5VAOkrVG6g7HQpeHE7EuOGp6xZVzXy0Wh6yMebNmwdAC8xfTFAdk5K4ao5Cjl6fOXPmAefGzzA8PKz8Q2op8zuawC6em2OSYQU81LnKpDFNuDj4rPo8TO9SSQGlJBiJD3luqorVFBKoFu7lgh4lfHdM7v/fpBYw0fCG0WMwTdMscAzEi8mhdiAKDfGovWNLCM/teYqO/R37xwtVcGMyWbR16BdTykvFbnVSd650rpfIAsGh0t1K6hLWB5YqzoGkC9V4BTjOaaAxTX03PF/oipEIyuwaFcDzAAVDkOUK2KRsVrtiybtuuubs81/uHRycngmDdHV1GAe2cSqVktQ00S8qb68068Pn12zbPqjJXvtdMqlM/bjJo3ZGX78jpwEOjIfwG8rgfTNebD1PjYt3tOAicx3PGLetwh3gOrZt80P3j9r9FeuXy1xQiOG+HD9ZFmVNTaoM41B9iwUeRVcqs+e6rplwXcw9o1qtJnO53ISS7frj1VdfVRarZVlR0LXWxnh/fX/rnCMQYts2dkFViAKfHatGvHLMnDmTdHd3q4EN30xpIWA1DAKi62nsXIgLIMVi7N69G40qQUeF8+JFYnBwUIeZVr+zWqkUAEVwAbCy6l1dXcy2bYHFo/4AlBT+N1JJEeBDOawW0UlFVILu7m6oneAZSd3LUfwc2aOgx1L3xg4EdwfADeSJURW3f4PgXadNm4Y4gOBcqwD0BOvJq1R8AZ3yIADARqCdYBVF941BLereGFvVKq1o1aqQuq4Ui23bdhOJJAYORoYV3af+3qptKpWRgHNeqoCVFIjGaEFgTPiNjU2qkg/nwaqK2rg2gH3fL2PEhe6R47tuyJoMC+eI6dOJOre5+brP3PlY32M7u7+wr5A3QfdGTB5AihLEBELR+obatSixiNndIgZoRbaPWJTgVDiUagnL5HM4H7tw/tHX/++5H39sV3d3Q2dbG3xSpO9EX19fCt/xHHjuuH3x90KhAFeo4Om6A9LDiuch1hIkk0m0PwJ05vbt27Gh0Oh98d5Y6NCvga7r4HxE+1CAnrCwKBRoGJOI27g2DmM3rlQqKW5N/CKRSOBaSFGweNPbvXu3BVMa4xGficcznkFlblCMQh1m28rFq1kJUKlCG2PuzJw5U40JPG/UT0TKqjpXUQBQGkIrOAdbJ8vnpRb1bf3zqqO/vx/xKzuyFgOkbbllCYMxWFMym80SPHM0X+vflQ0ODqo5DsX2RCIBdwobm2Pj4cMDsSm5HRv3RJ40FQyb6nneJE3TVGoxNrNVQM/z8s3NzaqM8mBHuVzuDIIAxCDKf42DftHqBiZbpPsOeoyNjc3Udb1B2LYrdT3ATgEN6mglG81kMkjRHPRAUJQxButGkajGkxLn+r6/t7GxEenCAw7EEmzbnh2ZnRP+FMU6kIs+KAnKhg0bjJkzZx4e80lEgwtLrtICHxsbm5DK2e++yVKpNKuuzlmt3GC0URyUBwmSxgcwIywIZsoQWou5H1EtqYCak81mN71T/6AU3XEcD8G6pzfsOP6WJ1+68pXdXWcOC5rNo6u9ipJchysAxLJ6DDA7S19VBqIeS0HYmSTJFKLjVumYSZP/+bUlp333+COPRJqvXqBSuX2lUukI3/eRZcF4UMxV0biCyd/V1NR0KKZhMGPPjywdTJw4fqIC0e/UP6ieLZVK8xRjrOcBCIaNCl4GyHWRQep5hwB4olQqzY0sNPUVWxT4XTqd3hIpgR1woLTdsqzZWDTiOFUUVFYuSSaT2XqovpVSNpfL5engP+Tg8o+x8uHfvGw2+07pzHaoescuZpx1QbtDJyfCnxz0qFarmAPpaP6EZIP1UVFN01TkLPq9CobhJtgVyTscyWQywIpPTGhLuDHUWK24lUqUjTzEgZSTDzZkdLqUwg5fSFkD9amqg0Vui8WY5TtCMEQR6yAIkOs95LmR6YVJLPfzKxUgpa4NDjh30aJFWIjwUspqiVf0aEF6x3YCMCeVSqlUWZzdiYJr9QvwIeMpEFrROYdlptwn+K6IdMc73qHOxfvBhcWOtGvfwLTTFs/dd9riuf+7u7uv9abHXj5149DQ+/rH2Al7x0ZYBeMZTAWIWqnTQI+vwZwkUxua5eRUYtv89rbHrj7ltP8cNXPu00iXnb5iBV+3SqHp6uMyyteOTHg1FqJFAX9TbsI79Q/6Ei5fNDnJfoHadzpX9QXO1zQkaQSsTLVYY0eNd/VDHLSuQCkOPse/O2gQu/4Z8L7hrWrjX7l9/79zC4UCFixcA1ZdLGeu3gN9F5VjHzSbFreFAWsSv0c80HWV8nN0zXc6FEhGD9vb/38K4V3QxfsGdQAAAABJRU5ErkJggg=="

def _b64_arquivo(nome_arquivo):
    caminho = os.path.join(os.path.dirname(os.path.abspath(__file__)), nome_arquivo)
    try:
        with open(caminho, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except OSError:
        return None

LOGO_ABA_B64 = _b64_arquivo("logo_aba_padrao.png") or _LOGO_ABA_B64_PADRAO
LOGO_ABA_B64_ROXO = _b64_arquivo("logo_texto.png") or _LOGO_ABA_B64_PADRAO


# =====================================================================
# 2. TEMA / CSS (replica a aparencia do CustomTkinter)
# =====================================================================
def cor_texto_legivel(hex_cor):
    try:
        h = hex_cor.lstrip("#")
        r, g, b = (int(h[i:i + 2], 16) for i in (0, 2, 4))
        luminancia = (0.299 * r + 0.587 * g + 0.114 * b) / 255
        return "#0d0d0d" if luminancia > 0.62 else "#ffffff"
    except Exception:
        return "#ffffff"

def ajustar_cor(hex_cor, fator):
    try:
        h = hex_cor.lstrip("#")
        r, g, b = (int(h[i:i + 2], 16) for i in (0, 2, 4))
        r = max(0, min(255, int(r * fator)))
        g = max(0, min(255, int(g * fator)))
        b = max(0, min(255, int(b * fator)))
        return f"#{r:02x}{g:02x}{b:02x}"
    except Exception:
        return hex_cor

def cor_rgba(hex_cor, alpha):
    try:
        h = hex_cor.lstrip("#")
        r, g, b = (int(h[i:i + 2], 16) for i in (0, 2, 4))
        return f"rgba({r},{g},{b},{alpha})"
    except Exception:
        return hex_cor

CSS_TEMPLATE = """
<style>
:root {
    color-scheme: @@SCHEME@@;
    --cor-p: @@CORP@@;
    --cor-s: @@CORS@@;
    --cor-ph: @@CORP_HOVER@@;
    --cor-pd: @@CORP_DEEP@@;
    --cor-texto: @@TEXTO@@;
    --cor-cinza: @@CINZA@@;
    --card-bg: @@CARDBG@@;
    --borda: @@BORDA@@;
    --fundo: @@FUNDO@@;
    --btn-fg: @@BTN@@;
}
html, body, .stApp {
    background: var(--fundo) !important;
    color: var(--cor-texto) !important;
    font-family: 'Segoe UI', system-ui, -apple-system, 'Roboto', Arial, sans-serif;
}
#MainMenu, footer {visibility: hidden;}
header[data-testid="stHeader"] {
    background: transparent !important;
    border-bottom: 1px solid transparent;
}
[data-testid="stDecoration"] {display: none !important;}
.block-container {padding: 0.8rem 1.5rem 2.5rem; max-width: 100%;}
h1, h2, h3, h4, h5 {color: var(--cor-texto) !important; letter-spacing: -0.01em;}
p, label, .stMarkdown, .stTextInput input, .stTextArea textarea,
.stSelectbox div[data-baseweb="select"] > div,
.stMultiSelect div[data-baseweb="select"] > div,
.stNumberInput input, .stDateInput input, .stTimeInput input {color: var(--cor-texto) !important;}
.stApp a {color: #17a2b8; text-decoration: none;}
.stApp a:hover {text-decoration: underline;}

/* ---------------- Scrollbar ---------------- */
::-webkit-scrollbar {width: 10px; height: 10px;}
::-webkit-scrollbar-thumb {background: @@SCROLL@@; border-radius: 8px;}
::-webkit-scrollbar-track {background: transparent;}

/* ---------------- Sidebar ---------------- */
section[data-testid="stSidebar"] {
    background: linear-gradient(185deg, var(--cor-s) 0%, @@CORS_GRAD@@ 100%) !important;
    border-right: 1px solid rgba(255,255,255,0.10);
}
@media (min-width: 769px) {
    section[data-testid="stSidebar"] {width: 14.375rem !important; min-width: 14.375rem !important;}
}
/* Esconde a barra de rolagem da sidebar (conteudo fica todo visivel) */
section[data-testid="stSidebar"] ::-webkit-scrollbar {display: none !important; width: 0 !important; height: 0 !important;}
section[data-testid="stSidebar"] {scrollbar-width: none !important; -ms-overflow-style: none !important;}
section[data-testid="stSidebar"] [data-testid="stSidebarContent"] {scrollbar-width: none !important; -ms-overflow-style: none !important;}
/* Remove o controle de redimensionar (alca na borda direita) */
section[data-testid="stSidebar"] > div[style*="cursor: col-resize"] {display: none !important;}
section[data-testid="stSidebar"] {color: #ececec !important;}
section[data-testid="stSidebar"] [data-testid="stSidebarContent"] {
    padding-top: 0 !important;
    padding-bottom: 6px !important;
}
@media (min-width: 769px) {
    section[data-testid="stSidebar"] [data-testid="stSidebarHeader"] {
        display: none !important;
    }
}
section[data-testid="stSidebar"] p, section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stMarkdown, section[data-testid="stSidebar"] [data-testid="stCaptionContainer"],
section[data-testid="stSidebar"] .stCaption {
    color: #ececec !important;
}
.logo-web {
    font-size: 1.85rem; font-weight: 800; line-height: 1.1; text-align: center;
    padding: 2px 8px 4px 8px; letter-spacing: -0.02em;
    background: linear-gradient(90deg, #ffffff, #e6c8ff);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;
}
.logo-web small {font-size: .5em; letter-spacing: 4px; font-weight: 700;
    -webkit-text-fill-color: rgba(255,255,255,.70); color: rgba(255,255,255,.70);}
.logo-aba {
    text-align: center; padding: 0 0 4px 0; line-height: 0;
}
.logo-aba img {
    width: 124px; height: auto; display: inline-block;
    filter: drop-shadow(0 0 5px rgba(255,255,255,.20));
}
.nav-secao {
    font-size: .66rem; font-weight: 800; letter-spacing: 2.5px; color: rgba(255,255,255,.55) !important;
    padding: 10px 12px 4px 12px; text-transform: uppercase;
    border-top: 1px solid rgba(255,255,255,.08); margin-top: 2px;
}
section[data-testid="stSidebar"] .stButton > button {
    width: 100%; text-align: left; justify-content: flex-start; gap: 10px;
    background: transparent; border: 1px solid transparent; color: rgba(255,255,255,.85) !important;
    border-radius: 9px; padding: .40rem .68rem; font-size: .87rem; font-weight: 600;
    margin-bottom: 2px; transition: all .18s ease; box-shadow: none;
}
section[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(255,255,255,0.12) !important; color: #fff !important;
    transform: translateX(3px); box-shadow: none;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"] {
    background: var(--cor-p) !important; color: var(--btn-fg) !important; font-weight: 700;
    border: 1px solid rgba(255,255,255,.28);
    box-shadow: 0 3px 12px rgba(0,0,0,.28);
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"] * {
    color: var(--btn-fg) !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"]:hover {
    background: var(--cor-ph) !important; transform: none;
}
section[data-testid="stSidebar"] .stButton > button[kind="secondary"] * {
    color: rgba(255,255,255,.85) !important;
}
/* Rodape da sidebar (login + Sair) sempre no fim, adaptando-se a tela */
section[data-testid="stSidebar"] div[class*="st-key-sb_footer"] {
    padding-top: 10px !important;
}

/* ---------------- Botoes ---------------- */
.stButton > button, .stDownloadButton > button, .stFormSubmitButton > button {
    border-radius: 10px; font-weight: 600; letter-spacing: .01em;
    transition: all .18s ease; box-shadow: 0 1px 3px rgba(0,0,0,.08);
}
.stButton > button:hover, .stDownloadButton > button:hover {
    box-shadow: 0 4px 14px rgba(0,0,0,.16); transform: translateY(-1px);
}
.stButton > button[kind="primary"], .stDownloadButton > button[kind="primary"],
.stFormSubmitButton > button[kind="primary"],
.stFormSubmitButton > button[kind="primaryFormSubmit"] {
    background: linear-gradient(180deg, var(--cor-p) 0%, var(--cor-pd) 100%);
    color: var(--btn-fg); border: 1px solid @@CORP_EDGE@@;
}
.stButton > button[kind="primary"]:hover, .stDownloadButton > button[kind="primary"]:hover,
.stFormSubmitButton > button[kind="primary"]:hover,
.stFormSubmitButton > button[kind="primaryFormSubmit"]:hover {
    background: linear-gradient(180deg, var(--cor-ph) 0%, var(--cor-pd) 100%);
    box-shadow: 0 6px 18px @@CORP_SHADOW@@;
}
.stButton > button[kind="secondary"],
.stFormSubmitButton > button[kind="secondaryFormSubmit"] {
    background: var(--card-bg); color: var(--cor-texto); border: 1px solid var(--borda);
}
.stButton > button[kind="secondary"]:hover,
.stFormSubmitButton > button[kind="secondaryFormSubmit"]:hover {background: @@CORS_SOFT@@;}
.stButton > button:focus-visible, .stDownloadButton > button:focus-visible,
.stFormSubmitButton > button:focus-visible {outline: 3px solid @@CORP_SOFT@@; outline-offset: 1px;}

/* Botao "Gerar com IA" (Central de Planos): com help, o Streamlit aninha o
   button dentro de spans (tooltip), entao o seletor ".stButton > button"
   nao casa. Fix local por key: garante fundo claro/texto escuro legiveis
   em qualquer tema, sem afetar os demais botoes (ex.: barra lateral). */
div[class*="st-key-plano_ia_gerar"] button[kind="secondary"] {
    background: var(--card-bg) !important; color: var(--cor-texto) !important;
    border: 1px solid var(--borda) !important;
}
div[class*="st-key-plano_ia_gerar"] button[kind="secondary"]:hover {
    background: @@CORS_SOFT@@ !important; color: var(--cor-texto) !important;
}
/* Tooltip do botao "Gerar com IA" (e demais dicas): fundo claro + texto
   escuro SEMPRE, mesmo com navegador em modo escuro forcado (color-scheme).
   O wrapper usa data-baseweb="tooltip" e o texto fica em stTooltipContent. */
[data-baseweb="tooltip"] {
    background: #ffffff !important; color: #1f1f1f !important;
    color-scheme: light !important; border: 1px solid #d3d8e0 !important;
    box-shadow: 0 6px 20px rgba(0,0,0,.15) !important;
}
[data-baseweb="tooltip"] [data-testid="stTooltipContent"],
[data-testid="stTooltipContent"],
[data-baseweb="tooltip"] * {
    background-color: transparent !important; color: #1f1f1f !important;
    -webkit-text-fill-color: #1f1f1f !important;
}

/* ---------------- Inputs ---------------- */
.stTextInput input, .stTextArea textarea, .stNumberInput input,
.stDateInput input, .stTimeInput input,
div[data-baseweb="input"], div[data-baseweb="input"] > div,
div[data-baseweb="select"] > div,
div[data-baseweb="select"] input,
.stMultiSelect div[data-baseweb="select"] > div,
.stMultiSelect div[data-baseweb="select"] input,
[data-testid="stColorPicker"] input {
    background: var(--card-bg) !important;
    color: var(--cor-texto) !important;
    border-radius: 10px !important; border-color: var(--borda) !important;
    transition: border-color .15s ease, box-shadow .15s ease;
}
.stTextInput input::placeholder, .stTextArea textarea::placeholder {color: var(--cor-cinza) !important; opacity: 1;}
.stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus,
div[data-baseweb="input"]:focus-within, div[data-baseweb="select"]:focus-within {
    border-color: @@CORS@@ !important; box-shadow: 0 0 0 3px @@CORP_SOFT@@ !important;
}
/* Campos do plano de aula (Central de Planos): comecam com uma linha e
   crescem conforme o texto (field-sizing: content, Chrome/Edge 123+,
   Firefox 130+). Navegadores antigos ficam com altura fixa + scroll. */
div[class*="st-key-plano_"] textarea {
    field-sizing: content;
    height: auto !important;
    resize: none !important;
    min-height: 44px !important;
    max-height: 300px !important;
    overflow-y: auto;
}
/* dropdown / menu */
[data-baseweb="popover"] [role="listbox"], [data-baseweb="menu"], [data-baseweb="popover"] ul {
    background: var(--card-bg) !important; color: var(--cor-texto) !important;
}
[data-baseweb="popover"] li {color: var(--cor-texto) !important;}
[data-baseweb="popover"] li:hover, [data-baseweb="popover"] li[aria-selected="true"] {
    background: @@CORS_SOFT@@ !important; color: var(--cor-texto) !important;
}
[data-baseweb="select"] span, [data-baseweb="select"] div {color: var(--cor-texto) !important;}
/* labels */
[data-testid="stCheckbox"] label p, [data-testid="stRadio"] label p,
[data-testid="stSelectbox"] label, [data-testid="stMultiSelect"] label,
[data-testid="stNumberInput"] label, [data-testid="stTextInput"] label,
[data-testid="stTextArea"] label, [data-testid="stDateInput"] label,
[data-testid="stTimeInput"] label, [data-testid="stSlider"] label,
[data-testid="stFileUploader"] label {
    color: var(--cor-texto) !important;
}
/* multiselect tags / slider / file uploader / expander / tabela */
[data-testid="stMultiSelect"] span[data-baseweb="tag"] {
    background: @@CORS_SOFT@@ !important; color: var(--cor-texto) !important;
}
[data-baseweb="slider"] div[role="slider"] {background: var(--cor-p) !important; border-color: var(--cor-p) !important;}
[data-testid="stFileUploaderDropzone"] {
    background: var(--card-bg) !important; border: 1px dashed var(--borda) !important;
    border-radius: 10px; color: var(--cor-texto) !important;
}
[data-testid="stFileUploaderDropzone"] span, [data-testid="stFileUploaderDropzone"] small {
    color: var(--cor-texto) !important;
}
[data-testid="stFileUploaderDropzone"] button {
    background: var(--cor-p) !important;
    color: var(--btn-fg) !important;
    border: 1px solid @@CORP_EDGE@@ !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    box-shadow: 0 1px 3px rgba(0,0,0,.08);
}
[data-testid="stFileUploaderDropzone"] button:hover {
    background: var(--cor-ph) !important;
}
[data-testid="stExpander"] {border: 1px solid var(--borda); border-radius: 12px;}
[data-testid="stExpander"] summary {color: var(--cor-texto) !important;}
[data-testid="stTable"] {background: var(--card-bg); border-radius: 10px; overflow: hidden;}

/* ---------------- Cards ---------------- */
[data-testid="stVerticalBlockBorderWrapper"] {
    background: var(--card-bg); border: 1px solid var(--borda) !important; border-radius: 14px;
    padding: 4px 12px; box-shadow: 0 1px 3px rgba(0,0,0,.05);
    transition: box-shadow .2s ease, border-color .2s ease;
}
[data-testid="stVerticalBlockBorderWrapper"]:hover {
    box-shadow: 0 6px 20px rgba(0,0,0,.10); border-color: @@CORP_SOFT@@ !important;
}
.card {
    background: var(--card-bg); border: 1px solid var(--borda); border-radius: 14px;
    padding: 14px 16px; margin-bottom: 12px;
    border-top: 3px solid @@COR_BORDA@@;
    box-shadow: 0 1px 3px rgba(0,0,0,.05); transition: box-shadow .2s ease, transform .2s ease;
}
.card:hover {box-shadow: 0 6px 20px rgba(0,0,0,.10); transform: translateY(-1px);}
/* Cards via st.container(key="card_...") -> classe st-key-card_* */
div[class*="st-key-card_"] {
    background: var(--card-bg); border: 1px solid var(--borda); border-radius: 14px;
    padding: 14px 16px; margin-bottom: 12px;
    border-top: 3px solid @@COR_BORDA@@;
    box-shadow: 0 1px 3px rgba(0,0,0,.05); transition: box-shadow .2s ease;
}
div[class*="st-key-card_"]:hover {box-shadow: 0 6px 20px rgba(0,0,0,.10);}

/* Lista de alunos (Turmas e Alunos): contorno ao redor da lista */
div[class*="st-key-lista_alunos"] {
    border: 1.5px solid @@COR_BORDA@@;
    border-radius: 14px;
    padding: 10px 14px 4px 14px;
    background: var(--card-bg);
    box-shadow: 0 1px 3px rgba(0,0,0,.05);
}
div[class*="st-key-lista_alunos"] div[data-testid="stMarkdownContainer"] p {
    font-weight: 400; color: var(--cor-texto); line-height: 1.4;
}
.aluno-num {
    display: inline-block; min-width: 30px; margin-right: 4px;
    font-size: .72rem; font-weight: 700; color: var(--cor-cinza);
    letter-spacing: .03em;
}
div[class*="st-key-lista_alunos"] [data-testid="stColumn"]:last-child,
div[class*="st-key-lista_alunos"] [data-testid="stColumn"]:nth-last-child(2) {
    display: flex; align-items: center; justify-content: flex-end;
    padding: 0 2px;
}
div[class*="st-key-lista_alunos"] [data-testid="stColumn"]:last-child {
    padding-right: 0;
}
div[class*="st-key-lista_alunos"] [data-testid="stColumn"]:nth-last-child(2) {
    padding-left: 0;
}
div[class*="st-key-lista_alunos"] [data-testid="stColumn"]:nth-last-child(2) div[data-testid="stLayoutWrapper"] {
    margin-left: auto;
}
div[class*="st-key-lista_alunos"] [data-testid="stColumn"]:nth-last-child(2),
div[class*="st-key-lista_alunos"] [data-testid="stColumn"]:last-child {
    flex: 0 0 auto; width: auto;
}
div[class*="st-key-aluno_"] button {
    min-width: 34px; height: 34px; padding: 0 !important;
    border-radius: 9px; border: 1px solid var(--borda);
    background: transparent; color: var(--cor-cinza); font-weight: 700;
    transition: all .15s ease;
}
div[class*="st-key-aluno_"] button:hover {
    background: #fee2e2; color: #b91c1c; border-color: #fca5a5;
}
div[class*="st-key-mover_"] [data-testid="stPopoverButton"] {
    min-width: 34px; height: 34px; padding: 0 6px !important;
    border-radius: 9px; border: 1px solid var(--borda) !important;
    background: transparent; color: var(--cor-p); font-weight: 600;
    transition: all .15s ease;
}
div[class*="st-key-mover_"] [data-testid="stPopoverButton"]:hover {
    background: @@CORP_SOFT@@; color: var(--cor-p); border-color: var(--cor-p) !important;
}

/* Topo de Turmas e Alunos: botoes popover da 1a linha (Turma e Adicionar Alunos) */
div[class*="st-key-pop_turmas"] [data-testid="stPopoverButton"],
div[class*="st-key-pop_adicionar"] [data-testid="stPopoverButton"] {
    width: 100%; height: 46px;
    border-radius: 12px; padding: 0 14px !important;
    font-weight: 700; font-size: 1rem;
    border: 1.5px solid var(--borda) !important;
    background: var(--card-bg) !important; color: var(--cor-texto) !important;
    display: flex; align-items: center; justify-content: center; gap: 8px;
    box-shadow: 0 1px 3px rgba(0,0,0,.05);
    transition: all .15s ease;
}
div[class*="st-key-pop_turmas"] [data-testid="stPopoverButton"]:hover {
    border-color: var(--cor-p) !important; color: var(--cor-p) !important;
    background: @@CORP_SOFT@@ !important;
}
div[class*="st-key-pop_adicionar"] [data-testid="stPopoverButton"] {
    background: linear-gradient(135deg, @@CORP@@, @@CORP_DEEP@@) !important;
    border: none !important; color: @@BTN@@ !important;
}
div[class*="st-key-pop_adicionar"] [data-testid="stPopoverButton"]:hover {
    background: linear-gradient(135deg, @@CORP@@, @@CORP_DEEP@@) !important;
    color: @@BTN@@ !important; filter: brightness(1.06);
}

/* Botao de exportar (impressora) ao lado do titulo "Turmas e Alunos" */
div[class*="st-key-export_alunos_pop"] { margin-left: auto; width: fit-content; }
div[class*="st-key-export_alunos_pop"] [data-testid="stPopoverButton"] {
    width: 42px; height: 42px; min-width: 42px;
    border-radius: 10px; padding: 0 !important;
    border: 1.5px solid var(--borda) !important;
    background: var(--card-bg) !important; color: var(--cor-p) !important;
    display: flex; align-items: center; justify-content: center;
    box-shadow: 0 1px 3px rgba(0,0,0,.05);
    transition: all .15s ease;
}
div[class*="st-key-export_alunos_pop"] [data-testid="stPopoverButton"]:hover {
    background: @@CORP_SOFT@@ !important; color: var(--cor-p) !important;
    border-color: var(--cor-p) !important;
}
div[class*="st-key-export_alunos_pop"] [data-testid="stPopoverButton"] [aria-hidden="true"] {
    display: none;
}
div[class*="st-key-export_alunos_pop"] [data-testid="stPopoverButton"] [data-testid="stIconMaterial"] {
    font-size: 1.25rem;
}

/* Dashboard (somente telas >= 769px): grade que preenche o espaco visivel,
   com borda superior colorida em cada campo e fundo branco puro. */
@media (min-width: 769px) {
    div[class*="st-key-dash_wrap"] {
        display: grid !important;
        grid-template-columns: minmax(0, 1fr) minmax(0, 1.4fr);
        grid-template-rows: auto auto;
        grid-template-areas:
            "mural  cal"
            "grade  grade";
        gap: 12px;
        align-items: stretch;
        min-height: calc(100vh - 340px);
    }
    div[class*="st-key-dash_wrap"] > div { min-width: 0; }
    div[class*="st-key-dash_wrap"] > div:has(> div.stVerticalBlock.st-key-card_dash_mural)  { grid-area: mural; }
    div[class*="st-key-dash_wrap"] > div:has(> div.stVerticalBlock.st-key-card_dash_grade)  { grid-area: grade; }
    div[class*="st-key-dash_wrap"] > div:has(> div.stVerticalBlock.st-key-card_dash_cal)    { grid-area: cal; }
    div[class*="st-key-dash_wrap"] > div > div.stVerticalBlock { height: 100%; }
    div[class*="st-key-dash_wrap"] div[class*="st-key-card_"] { margin-bottom: 0 !important; }
}
.card-titulo {font-weight: 800; font-size: .95rem; margin-bottom: 4px; color: var(--cor-texto);}
.card-sub {font-size: .82rem; color: var(--cor-cinza);}
.tag {
    display: inline-block; background: @@CORS_SOFT@@; color: @@TAGS@@;
    border-radius: 20px; padding: 3px 12px; font-size: .72rem; font-weight: 800; letter-spacing: .02em;
}
.postit {
    border-radius: 10px; padding: 12px 14px; margin-bottom: 10px;
    border: 1px solid rgba(0,0,0,0.12); box-shadow: 0 2px 6px rgba(0,0,0,0.10);
    transition: transform .2s ease, box-shadow .2s ease;
}
.postit:hover {transform: translateY(-2px) rotate(-.4deg); box-shadow: 0 8px 20px rgba(0,0,0,.18);}
.postit .pt-titulo {font-weight: 800; font-size: .95rem; color: inherit; margin-bottom: 4px;}
.postit .pt-tag {font-size: .72rem; color: inherit; opacity: .78; margin-bottom: 4px;}
.postit .pt-conteudo {font-size: .86rem; color: inherit; white-space: pre-wrap;}

/* Post-it do mural (Dashboard): conteudo truncado, revela o texto completo
   ao passar o mouse (sem JS, sem clipar dentro da lista de rolagem). */
.ei-pt .pt-conteudo {
    display: -webkit-box;
    -webkit-line-clamp: 3;
    -webkit-box-orient: vertical;
    overflow: hidden;
}
.ei-pt:hover .pt-conteudo {
    -webkit-line-clamp: unset;
    display: block;
}
.ei-pt:hover {transform: none;}

/* Lista de post-its do mural: altura fixa com barra de rolagem interna,
   para a pagina nao crescer infinitamente. */
div[class*="st-key-mural_scroll"] {
    overflow-y: auto;
    padding-right: 6px;
    scrollbar-width: thin;
    scrollbar-color: @@SCROLL@@ transparent;
}
div[class*="st-key-mural_scroll"]::-webkit-scrollbar {width: 8px;}
div[class*="st-key-mural_scroll"]::-webkit-scrollbar-track {background: transparent;}
div[class*="st-key-mural_scroll"]::-webkit-scrollbar-thumb {
    background: @@SCROLL@@; border-radius: 8px;
}

/* ---------------- Calendario ---------------- */
.cal-table {border-collapse: separate; border-spacing: 4px; width: 100%;}
.cal-table th {font-size: .7rem; text-transform: uppercase; letter-spacing: .08em; color: var(--cor-cinza); padding: 4px; text-align: center;}
.cal-cell {border: 1px solid var(--borda); text-align: center; font-size: .8rem; font-weight: 700;
    border-radius: 10px; width: 13%; padding: 8px 0; color: var(--cor-texto); background: var(--card-bg);
    transition: transform .15s ease, box-shadow .15s ease;}
.cal-link {text-decoration: none; display: block; padding: 8px 0; color: inherit !important; border-radius: 10px;}
.cal-hoje {background: @@CORP@@; color: @@BTN@@ !important; box-shadow: 0 3px 10px @@CORP_SHADOW@@;}
.cal-aula {background: #28a745; color: #fff !important;}
.cal-link:hover {background: #17a2b8; color: #fff !important; transform: scale(1.05);
    box-shadow: 0 4px 12px rgba(0,0,0,.20);}
.cal-scroll {overflow-x: auto; -webkit-overflow-scrolling: touch;}

/* Calendario com widgets (Dashboard): escopo nos botoes de dia.
   Nunca quebrar um dia de dois digitos em duas linhas: bloqueia quebra
   de palavra (overflow-wrap/word-break) e, no mobile, reduz a fonte. */
div[class*="st-key-cal_d_"] button {
    padding: 0.2rem 0.1rem !important;
    min-height: 2.05rem;
    font-weight: 700;
    border-radius: 10px;
    white-space: nowrap !important;
    overflow-wrap: normal !important;
    word-break: normal !important;
    overflow: hidden;
    position: relative;
}
div[class*="st-key-cal_d_"] button [data-testid="stMarkdownContainer"],
div[class*="st-key-cal_d_"] button [data-testid="stMarkdownContainer"] p {
    white-space: nowrap !important;
    overflow-wrap: normal !important;
    word-break: normal !important;
    margin: 0 !important;
}

/* ---------------- Aulas do Dia (Dashboard) ---------------- */
.adias {display: grid; grid-template-columns: repeat(auto-fill, minmax(180px, 1fr)); gap: 6px;}
.adias-cell {margin-bottom: 0;}

/* ---------------- Grade Semanal / Agenda ---------------- */
.gcell, .dash-item {
    background: var(--card-bg); border: 1px solid var(--borda); border-radius: 10px;
    padding: 8px 10px; margin-bottom: 6px;
    transition: border-color .15s ease, box-shadow .15s ease;
}
.gcell {position: relative;}
.gcell:hover, .dash-item:hover {border-color: @@CORS@@; box-shadow: 0 3px 10px rgba(0,0,0,.08);}
.gcell-aula {font-size: .72rem; font-weight: 700; color: @@CORS@@;}
.gcell-turma {font-weight: 700; font-size: .9rem; color: var(--cor-texto);}
.gcell-disc {font-size: .75rem; color: var(--cor-cinza);}
.gcell-del, .gcell-paint {display: none;}
.gcell-colorido { color: #1f2937 !important; }
.gcell-colorido .gcell-aula { color: #1f2937 !important; }
.gcell-colorido .gcell-turma { color: #1f2937 !important; }
.gcell-colorido .gcell-disc { color: #374151 !important; }
.gcell-tempo {font-weight: 400; font-size: .72rem; color: var(--cor-cinza); margin-left: 5px; white-space: nowrap;}
.gcell-colorido .gcell-tempo {color: #374151;}
/* Grade: botoes nativos (excluir sup dir / paleta inf dir) sobre o card */
div[class*="st-key-cell_"] {position: relative !important;}
div[class*="st-key-cell_"] > div[class*="st-key-del_"],
div[class*="st-key-cell_"] > div[class*="st-key-pal_"] {
    display: none; position: absolute; z-index: 6; width: auto !important;
}
div[class*="st-key-cell_"]:hover > div[class*="st-key-del_"] {display: block; top: 3px; right: 5px;}
div[class*="st-key-cell_"]:hover > div[class*="st-key-pal_"] {display: block; bottom: 3px; right: 5px;}

/* Grade Semanal: visao planilha (coluna de horarios + dias alinhados) */
.gh-head {
    font-weight: 800; font-size: .85rem; color: var(--cor-texto);
    padding: 2px 4px; margin-bottom: 6px; white-space: nowrap;
}
.gh-time {
    background: @@CORP_SOFT@@;
    border: 1px solid @@CORS@@; border-radius: 10px;
    padding: 8px 10px; margin-bottom: 6px;
    height: calc(100% - 6px); min-height: 58px;
    display: flex; flex-direction: column; justify-content: center;
    box-sizing: border-box;
}
.gh-time-aula {font-weight: 800; font-size: .82rem; color: var(--cor-p); white-space: nowrap;}
.gh-time-range {font-size: .8rem; font-weight: 600; color: var(--cor-texto); white-space: nowrap;}
.gh-vazio {
    border: 1px dashed var(--borda); border-radius: 10px;
    height: calc(100% - 6px); min-height: 58px; margin-bottom: 6px;
    background: transparent; box-sizing: border-box;
}

/* Grade Semanal: coluna de horarios estreita (so indica o horario da linha) */
div[class*="st-key-grade_dias"] [data-testid="stHorizontalBlock"] > div[data-testid="stColumn"]:first-child {
    flex: 0 0 130px !important;
}
/* mantem o botao da paleta ancorado enquanto o popover esta aberto */
div[class*="st-key-cell_"]:has(div[class*="st-key-pal_"] [aria-expanded="true"]) > div[class*="st-key-pal_"] {
    display: block; bottom: 3px; right: 5px;
}
div[class*="st-key-cell_"] > div[class*="st-key-del_"] button,
div[class*="st-key-cell_"] > div[class*="st-key-pal_"] button {
    min-width: 0 !important; min-height: 22px !important; height: 22px !important;
    padding: 0 6px !important; line-height: 1; font-size: .8rem; font-weight: 700;
    border: none !important; border-radius: 6px; background: rgba(255,255,255,0.85);
}
div[class*="st-key-cell_"] > div[class*="st-key-del_"] button {color: #e74c3c !important;}
div[class*="st-key-cell_"] > div[class*="st-key-del_"] button:hover {background: #e74c3c; color: #fff !important;}
div[class*="st-key-cell_"] > div[class*="st-key-pal_"] button {color: #4f46e5 !important;}
div[class*="st-key-cell_"] > div[class*="st-key-pal_"] button:hover {background: #4f46e5; color: #fff !important;}
/* Paleta de cores dentro do popover: grade compacta de quadradinhos */
div[class*="st-key-gc_"] button {
    min-width: 0 !important; width: 30px !important; height: 30px !important;
    min-height: 30px !important; padding: 0 !important;
    border: 1px solid rgba(0,0,0,.15) !important; border-radius: 6px !important;
}
div[class*="st-key-gc_"][class*="_Azul"] button {background: #dbeafe; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Verde"] button {background: #dcfce7; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Amarelo"] button {background: #fef9c3; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Laranja"] button {background: #ffedd5; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Rosa"] button {background: #fce7f3; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Lilas"] button {background: #ede9fe; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Ciano"] button {background: #cffafe; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Cinza"] button {background: #f1f5f9; color: #1f2937;}
div[class*="st-key-gc_"][class*="_Azul"] button:hover {box-shadow: 0 0 0 2px rgba(37,99,235,.6);}
div[class*="st-key-gc_"][class*="_Verde"] button:hover {box-shadow: 0 0 0 2px rgba(22,163,74,.6);}
div[class*="st-key-gc_"][class*="_Amarelo"] button:hover {box-shadow: 0 0 0 2px rgba(202,138,4,.6);}
div[class*="st-key-gc_"][class*="_Laranja"] button:hover {box-shadow: 0 0 0 2px rgba(234,88,12,.6);}
div[class*="st-key-gc_"][class*="_Rosa"] button:hover {box-shadow: 0 0 0 2px rgba(219,39,119,.6);}
div[class*="st-key-gc_"][class*="_Lilas"] button:hover {box-shadow: 0 0 0 2px rgba(124,58,237,.6);}
div[class*="st-key-gc_"][class*="_Ciano"] button:hover {box-shadow: 0 0 0 2px rgba(6,182,212,.6);}
div[class*="st-key-gc_"][class*="_Cinza"] button:hover {box-shadow: 0 0 0 2px rgba(100,116,139,.6);}
[data-testid="stPopoverBody"] {min-width: 0 !important; padding: 10px !important;}
.dash-item-titulo {font-weight: 700; font-size: .85rem; color: var(--cor-texto);}
.dash-item-sub {font-size: .8rem; color: var(--cor-cinza);}
.pend-item {
    background: rgba(220,53,69,0.10); border: 1px solid rgba(220,53,69,0.45);
    border-radius: 10px; padding: 8px 12px; margin-bottom: 6px;
    color: var(--cor-texto); font-size: .85rem;
}

/* ---------------- Metricas ---------------- */
div[data-testid="stMetric"] {
    background: var(--card-bg); border: 1px solid var(--borda); border-radius: 14px; padding: 12px 16px;
    box-shadow: 0 1px 3px rgba(0,0,0,.05);
}
div[data-testid="stMetricLabel"] {font-weight: 700 !important; color: var(--cor-cinza) !important;}
div[data-testid="stMetricValue"] {color: @@CORS@@ !important; font-size: 1.7rem !important; font-weight: 800 !important;}

/* ---------------- Abas ---------------- */
.stTabs [data-baseweb="tab-list"] {gap: 4px; border-bottom: 1px solid var(--borda);}
.stTabs [data-baseweb="tab"] {border-radius: 10px 10px 0 0; padding: 9px 18px; font-weight: 600;
    color: var(--cor-cinza) !important; transition: all .15s ease;}
.stTabs [data-baseweb="tab"]:hover {color: var(--cor-texto) !important; background: @@CORS_SOFT@@;}
.stTabs [data-baseweb="tab"][aria-selected="true"] {color: @@CORS@@ !important; font-weight: 800;
    border-bottom: 3px solid @@CORS@@;}

/* ---------------- Dialogs / popovers SEMPRE claros ----------------
   Regra fixa: toda janela flutuante (dialog, popover/tutorial) usa fundo
   claro (branco) e texto escuro, NAO importa o tema selecionado. Isso
   evita fundo preto com letra preta (ilegivel) em navegadores com modo
   escuro forcado ou quando o usuario escolhe o tema Dark. */
[data-testid="stDialog"] {
    border-radius: 16px; box-shadow: 0 24px 60px rgba(0,0,0,.35);
    background: #ffffff !important; color: #1f1f1f !important;
    color-scheme: light !important;
}
[data-testid="stDialog"] p, [data-testid="stDialog"] label,
[data-testid="stDialog"] .stMarkdown, [data-testid="stDialog"] h1,
[data-testid="stDialog"] h2, [data-testid="stDialog"] h3,
[data-testid="stDialog"] h4, [data-testid="stDialog"] h5,
[data-testid="stDialog"] li,
[data-testid="stDialog"] small, [data-testid="stDialog"] span {
    color: #1f1f1f !important;
}
[data-testid="stDialog"] a {color: #0b6bcb !important;}
/* Controles dentro de dialogos: fundo claro garantido (evita fundo ==
   cor da letra em navegadores com modo escuro forcado). */
[data-testid="stDialog"] input, [data-testid="stDialog"] textarea,
[data-testid="stDialog"] div[data-baseweb="input"],
[data-testid="stDialog"] div[data-baseweb="input"] > div,
[data-testid="stDialog"] div[data-baseweb="select"] > div,
[data-testid="stDialog"] div[data-baseweb="select"] input,
[data-testid="stDialog"] [data-baseweb="select"] span,
[data-testid="stDialog"] [data-baseweb="select"] div {
    background: #ffffff !important; color: #1f1f1f !important;
}
[data-testid="stDialog"] div[data-baseweb="input"] input,
[data-testid="stDialog"] div[data-baseweb="select"] input {
    color: #1f1f1f !important; -webkit-text-fill-color: #1f1f1f !important;
}
[data-testid="stDialog"] textarea {
    color: #1f1f1f !important; -webkit-text-fill-color: #1f1f1f !important;
}
[data-testid="stDialog"] label, [data-testid="stDialog"] p,
[data-testid="stDialog"] .stMarkdown, [data-testid="stDialog"] .stCaption,
[data-testid="stDialog"] small {color: #1f1f1f !important;}
[data-testid="stDialog"] div[data-testid="stCaptionContainer"] {
    color: #1f1f1f !important;
}
/* slider de cores (select_slider): o balao do valor selecionado fica sobre
   o fundo primario -> texto branco; o thumb e a trilha usam a cor primaria */
[data-testid="stDialog"] [data-baseweb="slider"] {
    color: #1f1f1f !important;
}
[data-testid="stDialog"] [data-baseweb="slider"] [role="slider"],
[data-testid="stDialog"] [data-baseweb="slider"] [role="slider"] * {
    background-color: var(--cor-p) !important;
    color: #ffffff !important;
}
[data-testid="stDialog"] [data-baseweb="slider"] [data-testid="stSliderThumbValue"],
[data-testid="stDialog"] [data-baseweb="slider"] [data-testid="stSliderThumbValue"] * {
    background-color: transparent !important;
    color: #ffffff !important;
}
[data-testid="stDialog"] [data-testid="stSliderThumbValue"] [data-testid="stMarkdownContainer"] {
    background-color: transparent !important;
}
[data-testid="stDialog"] [data-testid="stSliderTickBar"] [data-testid="stMarkdownContainer"] {
    color: #1f1f1f !important;
}
[data-testid="stDialog"] [data-baseweb="popover"] [role="listbox"],
[data-testid="stDialog"] [data-baseweb="menu"],
[data-testid="stDialog"] [data-baseweb="popover"] ul {
    background: #ffffff !important; color: #1f1f1f !important;
}
[data-testid="stDialog"] [data-baseweb="popover"] li {
    color: #1f1f1f !important;
}
[data-testid="stDialog"] [data-baseweb="popover"] li:hover,
[data-testid="stDialog"] [data-baseweb="popover"] li[aria-selected="true"] {
    background: @@CORS_SOFT@@ !important; color: #1f1f1f !important;
}
[data-testid="stDialog"] .stButton > button[kind="secondary"] {
    background: #ffffff !important; color: #1f1f1f !important;
    border: 1px solid #d3d8e0 !important;
}
[data-testid="stDialog"] .stButton > button[kind="secondary"]:hover {
    background: #eef2f7 !important; color: #1f1f1f !important;
}
[data-testid="stDialog"] .stButton > button[kind="primary"] {
    color: #ffffff !important;
}
/* ---------------- Popovers (janelas flutuantes) SEMPRE claros ----------------
   O Streamlit 1.58 usa stPopover / stPopoverBody / stPopoverButton.
   Fundo claro fixo (nunca preto) e texto escuro em todas as partes,
   inclusive no botao que abre o tutorial e nas listas internas.
   color-scheme: light impede o navegador de aplicar preto nativo
   (modo escuro forcado / tema Dark) nas camadas do popover. */
[data-testid="stPopover"] {background: #f2f4f7 !important; color-scheme: light !important;}
[data-testid="stPopoverButton"] {
    background: #ffffff !important; color: #1f1f1f !important;
    border: 1px solid #d3d8e0 !important;
}
[data-testid="stPopoverButton"]:hover {
    background: #eef2f7 !important; color: #1f1f1f !important;
}
[data-baseweb="popover"] {background: #f2f4f7 !important; color-scheme: light !important;}
[data-testid="stPopoverBody"] {
    background: #f2f4f7 !important; color: #1f1f1f !important;
    border: 1px solid #d3d8e0 !important; border-radius: 12px;
    color-scheme: light !important;
}
[data-testid="stPopoverBody"] p, [data-testid="stPopoverBody"] label,
[data-testid="stPopoverBody"] .stMarkdown, [data-testid="stPopoverBody"] h1,
[data-testid="stPopoverBody"] h2, [data-testid="stPopoverBody"] h3,
[data-testid="stPopoverBody"] h4, [data-testid="stPopoverBody"] h5,
[data-testid="stPopoverBody"] li,
[data-testid="stPopoverBody"] small, [data-testid="stPopoverBody"] span {
    color: #1f1f1f !important;
}
[data-testid="stPopoverBody"] a {color: #0b6bcb !important;}
[data-testid="stPopoverBody"] input, [data-testid="stPopoverBody"] textarea,
[data-testid="stPopoverBody"] div[data-baseweb="input"],
[data-testid="stPopoverBody"] div[data-baseweb="input"] > div,
[data-testid="stPopoverBody"] div[data-baseweb="select"] > div,
[data-testid="stPopoverBody"] div[data-baseweb="select"] input,
[data-testid="stPopoverBody"] [data-baseweb="select"] span {
    background: #ffffff !important; color: #1f1f1f !important;
}
[data-testid="stPopoverBody"] div[data-baseweb="input"] input,
[data-testid="stPopoverBody"] div[data-baseweb="select"] input {
    color: #1f1f1f !important; -webkit-text-fill-color: #1f1f1f !important;
}
[data-testid="stPopoverBody"] .stButton > button[kind="secondary"] {
    background: #ffffff !important; color: #1f1f1f !important;
    border: 1px solid #d3d8e0 !important;
}
[data-testid="stPopoverBody"] .stButton > button[kind="secondary"]:hover {
    background: #eef2f7 !important; color: #1f1f1f !important;
}
[data-testid="stPopoverBody"] [data-baseweb="popover"] [role="listbox"],
[data-testid="stPopoverBody"] [data-baseweb="menu"] {
    background: #ffffff !important; color: #1f1f1f !important;
}
[data-testid="stPopoverBody"] [data-baseweb="popover"] li {
    color: #1f1f1f !important;
}
/* Wrappers internos do popover herdavam fundo escuro do tema Dark:
   o divisor entre o body (cinza claro) e o conteudo e o envoltorio do
   textarea (stTextAreaRootElement) continuavam com cor de fundo
   escura. Forcamos cinza claro nesses containers. */
[data-testid="stPopoverBody"] > div {
    background: #f2f4f7 !important; color: #1f1f1f !important;
}
[data-testid="stPopoverBody"] [data-testid="stTextAreaRootElement"],
[data-testid="stPopoverBody"] [data-testid="stTextAreaRootElement"] > div,
[data-testid="stPopoverBody"] div[data-baseweb="textarea"],
[data-testid="stPopoverBody"] div[data-baseweb="textarea"] > div {
    background: #ffffff !important; color: #1f1f1f !important;
}
[data-testid="stPopoverBody"] [data-testid="stTextAreaRootElement"] textarea {
    color: #1f1f1f !important; -webkit-text-fill-color: #1f1f1f !important;
}
.stAlert {border-radius: 10px; border: 1px solid var(--borda);}
[data-testid="stDataFrame"] {border: 1px solid var(--borda); border-radius: 12px; overflow: hidden;}

/* ---------------- Divisores ---------------- */
hr {margin: .7rem 0; border-color: var(--borda);}

/* =====================================================================
   RESPONSIVO
   Estrategia: em telas <= 720px as colunas que nao cabem quebram linha
   (flex-wrap), empilhando formularios e cartoes; o calendario do
   dashboard permanece com as 7 colunas de dias; a grade semanal usa 2
   colunas por linha; a tabela do calendario de planos rola na horizontal.
   ===================================================================== */
@media (max-width: 720px) {
    .block-container {padding: .6rem .9rem 2rem;}
    h1 {font-size: 1.45rem !important;}
    h2 {font-size: 1.3rem !important;}
    h3 {font-size: 1.15rem !important;}
    .card-titulo {font-size: .9rem;}
    div[data-testid="stMetricValue"] {font-size: 1.35rem !important;}
    div[data-testid="stMetric"] {padding: 10px 12px;}
    .stTabs [data-baseweb="tab"] {padding: 8px 10px; font-size: .85rem;}
    div[class*="st-key-card_"] {padding: 12px;}
    .gcell {padding: 6px 8px;}
    .postit {padding: 10px 12px;}

    div[data-testid="stHorizontalBlock"] {flex-wrap: wrap !important;}
    div[data-testid="stHorizontalBlock"] > div[data-testid="stColumn"] {
        flex: 1 1 260px !important; min-width: 0 !important; width: auto !important;
    }
    /* Calendario do Dashboard: mantem 7 colunas compactas */
    div[class*="st-key-card_dash_cal"] div[data-testid="stColumn"] {
        flex: 1 1 0 !important; min-width: 0 !important; width: 100% !important;
    }
    div[class*="st-key-cal_d_"] button {
        padding: .1rem 0 !important; min-height: 1.6rem; font-size: .62rem; border-radius: 8px;
        white-space: nowrap !important; overflow: hidden; font-variant-numeric: tabular-nums;
        letter-spacing: 0; overflow-wrap: normal !important; word-break: normal !important;
        text-overflow: clip;
    }
    div[class*="st-key-cal_d_"] button [data-testid="stMarkdownContainer"],
    div[class*="st-key-cal_d_"] button [data-testid="stMarkdownContainer"] p {
        white-space: nowrap !important;
        overflow-wrap: normal !important;
        word-break: normal !important;
        font-size: .62rem !important;
        margin: 0 !important;
    }
    /* Grade Semanal: 2 dias por linha em telas pequenas */
    div[class*="st-key-grade_dias"] div[data-testid="stColumn"] {
        flex: 1 1 45% !important; min-width: 45% !important; width: auto !important;
    }
    /* Abas Entrar/Criar do login lado a lado */
    div[class*="st-key-login_card"] div[data-testid="stColumn"] {
        flex: 1 1 0 !important; min-width: 0 !important; width: 100% !important;
    }
    .cal-table {border-spacing: 2px;}
    .cal-table th {font-size: .6rem; padding: 2px;}
    .cal-table td {padding: 5px 0 !important; font-size: .75rem; border-radius: 6px;}
    [data-testid="stDialog"] {max-width: 94vw !important;}
}

/* Botao flutuante do menu (criado pelo JS injetado).
   Visibilidade controlada pelo JS: no celular fica sempre visivel;
   no desktop aparece so quando a sidebar esta oculta (para poder reabri-la). */
#ei-mobile-menu {
    position: fixed; right: 14px; bottom: 14px; z-index: 99999;
    width: 52px; height: 52px; border-radius: 16px; border: none; cursor: pointer;
    background: linear-gradient(135deg, @@CORP@@, @@CORP_DEEP@@) !important;
    color: @@BTN@@ !important; font-size: 24px; line-height: 1;
    align-items: center; justify-content: center;
    box-shadow: 0 8px 22px @@CORP_SHADOW@@;
    -webkit-tap-highlight-color: transparent;
}
#ei-mobile-menu:active { transform: scale(.94); }

@media (max-width: 480px) {
    div[class*="st-key-login_card"] {padding: 1.3rem 1.1rem !important;}
    [data-testid="stMain"] .block-container {padding-top: 2vh !important;}
    body::before, body::after {opacity: .55;}
    .cal-scroll .cal-table {min-width: 460px;}
}

/* ---------------- Toggle Criar Questao x Criar Com IA ---------------- */
div[class*="st-key-cad_modo_"] button {
    border-radius: 12px !important; font-weight: 800 !important;
    height: 44px; transition: all .18s ease;
}
div[class*="st-key-cad_modo_ia"] button[kind="primary"] {
    background: linear-gradient(135deg, var(--cor-p) 0%, var(--cor-pd) 100%) !important;
    border: 1px solid var(--cor-pd) !important; color: var(--btn-fg) !important;
    box-shadow: 0 8px 18px rgba(0,0,0,.18) !important;
}

/* =====================================================================
   CATALOGO DE QUESTOES: cards em grade (modelo "retangulos")
   Adaptado do mockup Central de Questoes (Tailwind -> Streamlit).
   Usa as variaveis do tema, entao funciona em todos os temas.
   ===================================================================== */
div[class*="st-key-cat_card_"] {
    background: var(--card-bg) !important;
    border: 2px solid var(--borda) !important;
    border-radius: 14px !important;
    box-shadow: 0 1px 2px rgba(0,0,0,.05), 0 1px 3px rgba(0,0,0,.06) !important;
    padding: 1.05rem 1.15rem !important;
    height: 100%;
    display: flex; flex-direction: column;
    transition: border-color .18s ease, box-shadow .18s ease;
}
div[class*="st-key-cat_card_"]:hover {
    border-color: var(--cor-p) !important;
    box-shadow: 0 6px 14px rgba(0,0,0,.08) !important;
}
/* O card e' o proprio stVerticalBlock; o stLayoutWrapper que o envolve precisa
   esticar para a altura da coluna (para cards da mesma linha ficarem iguais) */
div[data-testid="stLayoutWrapper"]:has(> div[class*="st-key-cat_card_"]) {
    height: 100% !important;
}
/* O "spacer" dentro do card cresce para preencher o espaco livre,
   empurrando a barra de acoes (Ver / Editar / Excluir) para o rodape */
div[class*="st-key-cat_card_"] > div[data-testid="stElementContainer"]:has(div.cat-q-spacer) {
    flex: 1 1 auto !important;
}
div[class*="st-key-cat_card_"] > div[data-testid="stLayoutWrapper"]:last-child {
    margin-top: auto;
}
.cat-q-titulo {
    font-weight: 800; font-size: 1.0rem; line-height: 1.3;
    color: var(--cor-texto) !important;
}
.cat-q-badge {
    display: inline-block; padding: 3px 10px; border-radius: 999px;
    font-size: .72rem; font-weight: 700; white-space: nowrap;
    color: #065f46 !important; background: #d1fae5 !important;
    border: 1px solid #a7f3d0;
}
.cat-q-texto {
    margin: .55rem 0 .7rem 0; color: var(--cor-texto) !important; opacity: .95;
    font-size: .86rem; line-height: 1.5;
    display: -webkit-box; -webkit-line-clamp: 3; -webkit-box-orient: vertical;
    overflow: hidden; min-height: 3.9em;
}
.cat-q-tags { display: flex; flex-wrap: wrap; gap: 6px; margin: 0 0 .85rem 0; }
.cat-q-tag {
    padding: 3px 10px; border-radius: 6px; font-size: .7rem; font-weight: 600;
    max-width: 100%; overflow: hidden; text-overflow: ellipsis; white-space: nowrap;
    background: var(--fundo) !important; color: var(--cor-cinza) !important;
    border: 1px solid var(--borda) !important;
}
div[class*="st-key-cat_ver_"] [data-testid="stPopoverButton"],
div[class*="st-key-cat_ed_"] [data-testid="stPopoverButton"],
div[class*="st-key-cat_del_"] [data-testid="stPopoverButton"] {
    width: 100% !important; height: 34px !important; border-radius: 999px !important;
    font-size: .78rem !important; font-weight: 600 !important;
    border: 1.5px solid var(--borda) !important;
    background: transparent !important; color: var(--cor-texto) !important;
    box-shadow: none !important; transition: all .15s ease;
    display: flex; align-items: center; justify-content: center; gap: 6px;
}
div[class*="st-key-cat_ver_"] [data-testid="stPopoverButton"]:hover,
div[class*="st-key-cat_ed_"] [data-testid="stPopoverButton"]:hover {
    border-color: var(--cor-p) !important; color: var(--cor-p) !important;
    background: rgba(0,0,0,.04) !important;
}
div[class*="st-key-cat_del_"] [data-testid="stPopoverButton"]:hover {
    border-color: #dc2626 !important; color: #dc2626 !important;
    background: rgba(220,38,38,.06) !important;
}
</style>
"""

def injetar_css(config):
    aparencia = config.get("aparencia", "System")
    dark = aparencia == "Dark"
    tema_visual = config.get("tema_visual", "")

    if tema_visual == "roxo":
        # Tema de TESTE: Dark com tons de roxo
        # (primaryColor #7d3fe0, backgroundColor #090710,
        #  secondaryBackgroundColor #151221, textColor #e4e1eb)
        dark = True
        cor_p = "#7d3fe0"
        cor_s = "#5b2ea6"
        cor_fundo = "#090710"
        card_bg = "#151221"
        borda = "#2a2540"
        texto = "#e4e1eb"
        texto_cinza = "#a89fc0"
    elif tema_visual == "branco_novo":
        # Tema novo: "Branco Novo" - Clean SaaS claro com acentos roxos
        # (primaryColor #7d3fe0, backgroundColor #f4f5f7,
        #  secondaryBackgroundColor #ffffff, textColor #2c2e3e)
        dark = False
        cor_p = "#7d3fe0"
        cor_s = "#5b2ea6"
        cor_fundo = "#f4f5f7"
        card_bg = "#ffffff"
        borda = "#eef0f4"
        texto = "#2c2e3e"
        texto_cinza = "#838896"
    elif tema_visual == "teste":
        # Tema de TESTE: inspirado no mockup Stitch "Teacher Dashboard"
        # (sidebar roxo-escuro #1a0b2e, primario #8b5cf6, fundo #f3f4f6,
        #  cards brancos com sombra, fonte Inter)
        dark = False
        cor_p = "#8b5cf6"
        cor_s = "#4c1d95"
        cor_fundo = "#f3f4f6"
        card_bg = "#ffffff"
        borda = "#e5e7eb"
        texto = "#1f2937"
        texto_cinza = "#6b7280"
    else:
        cor_tema = config.get("cor_tema", "blue")
        presets = {
            "blue": ("#1f538d", "#14375e"),
            "green": ("#1d7a46", "#104f2c"),
            "dark-blue": ("#1a3a6b", "#0e2340"),
        }
        if cor_tema in presets:
            cor_p, cor_s = presets[cor_tema]
        else:
            cor_p = config.get("cor_principal", "#1f538d")
            cor_s = config.get("cor_secundaria", "#14375e")
        cor_fundo = config.get("cor_fundo", "#2b2b2b") if dark else "#e4e7ec"
        card_bg = "#2b2b2b" if dark else "#ffffff"
        borda = "rgba(255,255,255,0.10)" if dark else "#d3d8e0"
        texto = "#e8e8e8" if dark else "#1f1f1f"
        texto_cinza = "#9a9a9a" if dark else "#6c757d"

    txt_btn = cor_texto_legivel(cor_p)

    cor_p_hover = ajustar_cor(cor_p, 1.12)
    cor_p_deep = ajustar_cor(cor_p, 0.82)
    cor_p_edge = ajustar_cor(cor_p, 0.66)
    cor_s_grad = ajustar_cor(cor_s, 0.80)
    tag_txt = ajustar_cor(cor_s, 0.55)
    cor_borda = (config.get("cor_borda_card") or "").strip() or cor_p

    css = CSS_TEMPLATE
    css = css.replace("@@SCHEME@@", "dark" if dark else "light")
    css = css.replace("@@CORP@@", cor_p)
    css = css.replace("@@CORS@@", cor_s)
    css = css.replace("@@CORP_HOVER@@", cor_p_hover)
    css = css.replace("@@CORP_DEEP@@", cor_p_deep)
    css = css.replace("@@CORP_EDGE@@", cor_p_edge)
    css = css.replace("@@CORP_SOFT@@", cor_rgba(cor_p, 0.18))
    css = css.replace("@@CORP_SHADOW@@", cor_rgba(cor_p, 0.35))
    css = css.replace("@@CORS_GRAD@@", cor_s_grad)
    css = css.replace("@@CORS_SOFT@@", cor_rgba(cor_s, 0.12))
    css = css.replace("@@COR_BORDA@@", cor_borda)
    css = css.replace("@@TAGS@@", tag_txt)
    css = css.replace("@@SCROLL@@", cor_rgba(cor_p, 0.45))
    css = css.replace("@@FUNDO@@", cor_fundo)
    css = css.replace("@@BTN@@", txt_btn)
    css = css.replace("@@CARDBG@@", card_bg)
    css = css.replace("@@BORDA@@", borda)
    css = css.replace("@@TEXTO@@", texto)
    css = css.replace("@@CINZA@@", texto_cinza)
    if tema_visual == "roxo":
        # Degrade radial roxo/vermelho no fundo + sidebar Glassmorphism (tema de teste)
        css = css.replace(
            "</style>",
            """[data-testid="stAppViewContainer"] {
    background-color: #090710;
    background-image: radial-gradient(circle at 0% 0%, #1a0b2e 0%, #090710 40%, #090710 100%) !important;
}

/* ---- Sidebar Glassmorphism (tema Roxo) ---- */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #13101d 0%, #0a0812 100%) !important;
    border-right: 1px solid rgba(255, 255, 255, 0.05) !important;
}
section[data-testid="stSidebar"] .nav-secao {
    color: #8f85a3 !important;
    font-size: 0.72rem !important;
    letter-spacing: 1.5px !important;
    text-transform: uppercase !important;
}
section[data-testid="stSidebar"] .stButton > button {
    background-color: rgba(255, 255, 255, 0.02) !important;
    color: #b9b3c7 !important;
    border: 1px solid rgba(255, 255, 255, 0.05) !important;
    border-radius: 10px !important;
    transition: all 0.25s ease-in-out !important;
    text-align: left !important;
    padding: 8px 14px !important;
}
section[data-testid="stSidebar"] .stButton > button:hover {
    background-color: rgba(125, 63, 224, 0.15) !important;
    border-color: rgba(125, 63, 224, 0.4) !important;
    color: #ffffff !important;
    transform: translateX(3px);
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"] {
    background: linear-gradient(90deg, #7d3fe0 0%, #6323c2 100%) !important;
    color: #ffffff !important;
    border: none !important;
    box-shadow: 0 4px 12px rgba(125, 63, 224, 0.35) !important;
    font-weight: 600 !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"] * {
    color: #ffffff !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"]:hover {
    background: linear-gradient(90deg, #7d3fe0 0%, #6323c2 100%) !important;
    transform: none !important;
}

/* ---- Logo da sidebar cabendo sem aumentar a barra (so no tema Roxo) ---- */
section[data-testid="stSidebar"] [data-testid="stSidebarContent"] {
    padding-top: 2px !important;
    margin-top: 0 !important;
}
.logo-aba img {
    width: 80% !important;
    max-width: 80% !important;
    margin-top: 0 !important;
}
.logo-aba {
    padding: 10px 0 4px 0 !important;
}

/* ---- Botoes de criar post-it (Inicio + Lembretes) no estilo "Criar Questao" ---- */
div[class*="st-key-dash_novo_pop"] [data-testid="stPopoverButton"],
div[class*="st-key-anot_novo_pop"] [data-testid="stPopoverButton"] {
    border-radius: 12px !important;
    font-weight: 800 !important;
    height: 44px !important;
    background: linear-gradient(135deg, var(--cor-p) 0%, var(--cor-pd) 100%) !important;
    border: 1px solid rgba(125, 63, 224, 0.6) !important;
    color: var(--btn-fg) !important;
    box-shadow: 0 8px 18px rgba(125, 63, 224, 0.35) !important;
}
div[class*="st-key-dash_novo_pop"] [data-testid="stPopoverButton"]:hover,
div[class*="st-key-anot_novo_pop"] [data-testid="stPopoverButton"]:hover {
    background: linear-gradient(135deg, var(--cor-ph) 0%, var(--cor-pd) 100%) !important;
}
</style>""")
    if tema_visual == "branco_novo":
        # Tema novo "Branco Novo": fundo off-white, sidebar branca e detalhes roxos
        # (so este tema; nenhum outro e afetado)
        css = css.replace(
            "</style>",
            """/* ---- Tema Branco Novo: Clean SaaS (fundo off-white, sidebar branca) ---- */
[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #fdfdfd 0%, #f4f5f7 100%) !important;
}

/* Barra lateral branca com sombra suave projetada para a direita */
section[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 1px solid #eef0f4 !important;
    box-shadow: 2px 0 15px rgba(0, 0, 0, 0.03) !important;
}
section[data-testid="stSidebar"] {color: #2c2e3e !important;}
section[data-testid="stSidebar"] p, section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stMarkdown, section[data-testid="stSidebar"] [data-testid="stCaptionContainer"],
section[data-testid="stSidebar"] .stCaption {
    color: #2c2e3e !important;
}
section[data-testid="stSidebar"] .logo-aba img {
    filter: drop-shadow(0 2px 6px rgba(0, 0, 0, 0.10));
}

/* Titulos das secoes da sidebar (cinza claro, versao limpa) */
section[data-testid="stSidebar"] .nav-secao {
    color: #838896 !important;
    border-top: 1px solid #eef0f4;
}

/* Botoes da sidebar - estado padrao (inativo) */
section[data-testid="stSidebar"] .stButton > button {
    background-color: transparent !important;
    color: #4a4d5e !important;
    border: 1px solid transparent !important;
    border-radius: 10px !important;
    transition: all 0.25s ease-in-out !important;
    text-align: left !important;
    padding: 8px 14px !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="secondary"] * {
    color: #4a4d5e !important;
}

/* Hover: acende com o roxo da marca */
section[data-testid="stSidebar"] .stButton > button:hover {
    background-color: rgba(125, 63, 224, 0.08) !important;
    color: #7d3fe0 !important;
    transform: translateX(3px);
}

/* Botao ativo: solido roxo com sombra iluminada */
section[data-testid="stSidebar"] .stButton > button[kind="primary"] {
    background: linear-gradient(90deg, #7d3fe0 0%, #6323c2 100%) !important;
    color: #ffffff !important;
    border: none !important;
    box-shadow: 0 4px 10px rgba(125, 63, 224, 0.25) !important;
    font-weight: 600 !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"] * {
    color: #ffffff !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"]:hover {
    background: linear-gradient(90deg, #7d3fe0 0%, #6323c2 100%) !important;
    transform: none !important;
}

/* Cards do conteudo principal: brancos com sombra sutil */
[data-testid="stForm"] {
    background-color: #ffffff !important;
    border: 1px solid #eef0f4 !important;
    border-radius: 12px !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.04) !important;
}
</style>""")
    if tema_visual == "teste":
        # Tema de TESTE "Stitch": sidebar roxo-escuro + conteudo claro em cards
        # brancos com sombra e fonte Inter. Isolado: so este tema e alterado.
        css = css.replace(
            "</style>",
            """/* ---- Tema Teste (Stitch): Teacher Dashboard ---- */
html, body, .stApp {
    font-family: 'Inter', system-ui, -apple-system, 'Segoe UI', Roboto, sans-serif !important;
}
[data-testid="stAppViewContainer"] {
    background: #f3f4f6 !important;
}

/* Sidebar roxo-escuro (#1a0b2e) */
section[data-testid="stSidebar"] {
    background: #1a0b2e !important;
    border-right: 1px solid rgba(255,255,255,0.08) !important;
}
section[data-testid="stSidebar"] { color: #d1d5db !important; }
section[data-testid="stSidebar"] p, section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stMarkdown,
section[data-testid="stSidebar"] [data-testid="stCaptionContainer"],
section[data-testid="stSidebar"] .stCaption {
    color: #d1d5db !important;
}
section[data-testid="stSidebar"] .logo-aba img {
    filter: drop-shadow(0 2px 8px rgba(139,92,246,0.35));
}
section[data-testid="stSidebar"] .nav-secao {
    color: #9ca3af !important;
    border-top: 1px solid rgba(255,255,255,0.06);
}

/* Botoes da sidebar - inativo (texto cinza claro, hover roxo translucido) */
section[data-testid="stSidebar"] .stButton > button {
    background-color: transparent !important;
    color: #d1d5db !important;
    border: 1px solid transparent !important;
    border-radius: 12px !important;
    text-align: left !important;
    padding: 9px 14px !important;
    transition: all 0.2s ease !important;
    box-shadow: none !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="secondary"] * {
    color: #d1d5db !important;
}
section[data-testid="stSidebar"] .stButton > button:hover {
    background-color: rgba(139, 92, 246, 0.10) !important;
    color: #ffffff !important;
    transform: translateX(3px);
}

/* Botao ativo: roxo solido #8b5cf6 (como o item ativo do mockup) */
section[data-testid="stSidebar"] .stButton > button[kind="primary"] {
    background-color: #8b5cf6 !important;
    color: #ffffff !important;
    border: none !important;
    box-shadow: 0 3px 10px rgba(139, 92, 246, 0.35) !important;
    font-weight: 600 !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"] * {
    color: #ffffff !important;
}
section[data-testid="stSidebar"] .stButton > button[kind="primary"]:hover {
    background-color: #7c3aed !important;
    transform: none !important;
}

/* Rodape (Logado / Sair): cartao de perfil como no mockup */
section[data-testid="stSidebar"] div[class*="st-key-sb_footer"] {
    background: rgba(255,255,255,0.05) !important;
    border: 1px solid rgba(255,255,255,0.10) !important;
    border-radius: 12px !important;
    padding: 10px 12px !important;
    margin: 10px 10px 4px 10px !important;
}
section[data-testid="stSidebar"] div[class*="st-key-sb_footer"] hr {
    display: none !important;
}
section[data-testid="stSidebar"] div[class*="st-key-sb_footer"] [data-testid="stCaption"] {
    color: #9ca3af !important; font-size: .78rem !important;
}
section[data-testid="stSidebar"] div[class*="st-key-sb_footer"] button {
    background-color: transparent !important;
    color: #e5e7eb !important;
    border: 1px solid rgba(255,255,255,0.20) !important;
    border-radius: 8px !important;
    font-size: .78rem !important;
    padding: 4px 10px !important;
    height: auto !important;
    box-shadow: none !important;
}
section[data-testid="stSidebar"] div[class*="st-key-sb_footer"] button:hover {
    background-color: rgba(255,255,255,0.10) !important;
}

/* Cards do conteudo: brancos, cantos grandes (rounded-2xl) e card-shadow */
[data-testid="stVerticalBlockBorderWrapper"], .card, div[class*="st-key-card_"],
[data-testid="stForm"] {
    background-color: #ffffff !important;
    border: 1px solid #e5e7eb !important;
    border-radius: 16px !important;
    box-shadow: 0 4px 6px -1px rgba(0,0,0,0.10), 0 2px 4px -1px rgba(0,0,0,0.06) !important;
}
h1, h2, h3, h4, h5 { color: #111827 !important; }

/* Botoes "Novo" (popover) no estilo do mockup: contorno roxo claro */
div[class*="st-key-dash_novo_pop"] [data-testid="stPopoverButton"],
div[class*="st-key-anot_novo_pop"] [data-testid="stPopoverButton"] {
    border-radius: 10px !important;
    font-weight: 600 !important;
    background: transparent !important;
    color: #7c3aed !important;
    border: 1px solid #c4b5fd !important;
    box-shadow: none !important;
}
div[class*="st-key-dash_novo_pop"] [data-testid="stPopoverButton"]:hover,
div[class*="st-key-anot_novo_pop"] [data-testid="stPopoverButton"]:hover {
    background: #f5f3ff !important;
}
</style>""")
    css += """
<style>
/* Em telas baixas, compacta a sidebar para caber sem rolagem (todos os temas) */
@media (max-height: 900px) {
    section[data-testid="stSidebar"] .logo-aba img { width: 96px !important; }
    section[data-testid="stSidebar"] .stButton > button { padding: .26rem .6rem !important; margin-bottom: 1px !important; font-size: .84rem !important; }
    section[data-testid="stSidebar"] .nav-secao { padding: 5px 12px 3px 12px !important; }
    section[data-testid="stSidebar"] div[class*="st-key-sb_footer"] { padding-top: 4px !important; }
}
@media (max-height: 780px) {
    section[data-testid="stSidebar"] .logo-aba img { width: 84px !important; }
    section[data-testid="stSidebar"] .stButton > button { padding: .22rem .5rem !important; margin-bottom: 0 !important; font-size: .82rem !important; }
    section[data-testid="stSidebar"] .nav-secao { padding: 3px 12px 2px 12px !important; }
}
</style>
"""
    st.markdown(css, unsafe_allow_html=True)

CSS_LOGIN = """
<style>
/* ---------- Tela de login: fundo azul com simbolos ---------- */
html, body, .stApp, section[data-testid="stMain"] {
    background:
        radial-gradient(1100px 700px at 88% 12%, rgba(10,123,184,.5), transparent 60%),
        radial-gradient(900px 640px at 6% 92%, rgba(30,64,175,.4), transparent 60%),
        radial-gradient(700px 500px at 70% 100%, rgba(14,165,233,.25), transparent 60%),
        linear-gradient(158deg, #011d44 0%, #023a75 38%, #055c91 68%, #0a6fa8 100%) !important;
}
[data-testid="stSidebar"] { display: none !important; }
[data-testid="stMain"] .block-container {
    max-width: 470px !important;
    padding-top: 5vh !important;
    padding-bottom: 4vh !important;
}
/* simbolos flutuantes (livros, matematica) no fundo */
#log-symb { position: fixed; inset: 0; overflow: hidden; pointer-events: none; z-index: 0; }
#log-symb span {
    position: absolute; display: block; color: rgba(255,255,255,.10);
    font-size: 40px; user-select: none; line-height: 1;
}
/* ---------- Card central ---------- */
div[class*="st-key-login_card"] {
    background: rgba(255,255,255,.94) !important;
    border: 1px solid rgba(255,255,255,.75) !important;
    border-radius: 24px !important;
    padding: 2.1rem 2.2rem 1.9rem !important;
    box-shadow: 0 26px 60px rgba(0,20,70,.35), 0 4px 14px rgba(0,0,0,.15) !important;
    position: relative; z-index: 1;
}
/* ---------- Marca (legivel em fundo claro) ---------- */
.login-logo { text-align: center; margin: 0 auto 10px auto; line-height: 0; }
.login-img {
    width: 300px; max-width: 100%; height: auto; display: inline-block;
    filter: drop-shadow(0 10px 24px rgba(31,83,141,.25));
}
.login-titulo {
    text-align: center; font-size: 1.45rem; font-weight: 800; color: #111827 !important;
    line-height: 1.15; letter-spacing: -.02em; margin-bottom: 4px;
}
.login-titulo small {
    display: block; margin-top: 4px; font-size: .6rem; font-weight: 700;
    letter-spacing: 4px; text-transform: uppercase; color: #4f46e5 !important;
}
.login-sub { color: #6b7280 !important; font-size: .9rem; margin: 4px 0 1.5rem 0; }
.login-hint {
    text-align: center; color: #9ca3af !important; font-size: .85rem;
    margin: 1.2rem 0 0 0; font-style: italic;
}
div[class*="st-key-login_card"] {
    overflow: hidden;
}
div[class*="st-key-login_card"]:has([data-testid="stTextInput"]) {
    animation: eiExpand .55s cubic-bezier(.22,.85,.25,1);
}
@keyframes eiExpand {
    0% { max-height: 330px; opacity: .35; }
    60% { opacity: .85; }
    100% { max-height: 1400px; opacity: 1; }
}
/* ---------- Botao Recolher ---------- */
div[class*="st-key-login_recolher"] button {
    background: transparent !important; border: none !important;
    color: #9ca3af !important; font-size: .8rem !important;
    box-shadow: none !important; padding: 2px !important; height: 32px !important;
    margin-top: .6rem !important;
}
div[class*="st-key-login_recolher"] button:hover {
    color: #4f46e5 !important; background: transparent !important;
}
/* ---------- Abas Entrar / Criar Conta ---------- */
div[class*="st-key-tab_entrar"] button, div[class*="st-key-tab_criar"] button {
    border-radius: 12px !important; font-weight: 700 !important;
    border: 1.5px solid #e2e8f0 !important; background: #ffffff !important;
    color: #64748b !important; box-shadow: none !important; height: 42px;
    transition: all .18s ease;
}
div[class*="st-key-tab_entrar"] button[kind="primary"],
div[class*="st-key-tab_criar"] button[kind="primary"] {
    background: linear-gradient(135deg, #1f538d, #4f46e5) !important;
    border: none !important; color: #ffffff !important;
    box-shadow: 0 8px 18px rgba(31,83,141,.28) !important;
}
/* ---------- Abas da Central de Questões ---------- */
div[class*="st-key-cq_tab_"] button {
    border-radius: 12px !important; font-weight: 800 !important;
    height: 44px; transition: all .18s ease;
}
div[class*="st-key-cq_tab_"] button[kind="primary"] {
    background: linear-gradient(135deg, @@CORP@@, @@CORP_DEEP@@) !important;
    border: 1px solid @@CORP_EDGE@@ !important; color: @@BTN@@ !important;
    box-shadow: 0 8px 18px @@CORP_SHADOW@@ !important;
}
/* ---------- Inputs ---------- */
[data-testid="stTextInput"] label p {
    color: #374151 !important; font-weight: 600; font-size: .83rem;
}
[data-testid="stTextInput"] div[data-baseweb="input"],
[data-testid="stTextInput"] div[data-baseweb="input"] > div {
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
}
[data-testid="stTextInput"] input {
    background: #f8fafc !important; border: 1.5px solid #e2e8f0 !important;
    border-radius: 12px !important; padding: .7rem .9rem !important;
    color: #111827 !important; font-size: .95rem !important;
    width: 100% !important;
}
[data-testid="stTextInput"] input:focus {
    border-color: #4f46e5 !important; background: #ffffff !important;
    box-shadow: 0 0 0 4px rgba(79,70,229,.12) !important;
    outline: none !important;
}
[data-testid="stTextInput"] input::placeholder { color: #9ca3af !important; opacity: 1; }
/* ---------- Botao CTA ---------- */
.stFormSubmitButton > button[kind="primary"] {
    background: linear-gradient(135deg, #1f538d, #4f46e5) !important;
    border: none !important; border-radius: 12px !important;
    height: 48px !important; font-weight: 700 !important; font-size: .95rem !important;
    box-shadow: 0 12px 24px rgba(31,83,141,.30) !important;
    transition: transform .18s ease, box-shadow .18s ease;
}
.stFormSubmitButton > button[kind="primary"]:hover {
    transform: translateY(-1px); box-shadow: 0 16px 30px rgba(31,83,141,.38) !important;
}
/* ---------- Alertas ---------- */
[data-testid="stAlert"] { border-radius: 12px !important; font-size: .85rem; }
</style>
<div id="log-symb" aria-hidden="true">
  <span style="left:6%;top:12%;font-size:52px;transform:rotate(-14deg)">&#128218;</span>
  <span style="left:16%;bottom:14%;font-size:38px;transform:rotate(10deg)">&#8747;</span>
  <span style="right:5%;top:18%;font-size:46px;transform:rotate(8deg)">&#128214;</span>
  <span style="right:14%;bottom:10%;font-size:40px;transform:rotate(-8deg)">&#120493;</span>
  <span style="left:12%;top:42%;font-size:30px;transform:rotate(-6deg)">&#8730;</span>
  <span style="right:22%;top:44%;font-size:32px;transform:rotate(12deg)">&#10142;</span>
  <span style="left:32%;bottom:8%;font-size:34px;transform:rotate(-10deg)">&#10003;</span>
  <span style="left:3%;bottom:30%;font-size:28px;transform:rotate(6deg)">&#8764;</span>
  <span style="right:8%;bottom:38%;font-size:30px;transform:rotate(-12deg)">&#9633;</span>
  <span style="left:26%;top:8%;font-size:26px;transform:rotate(12deg)">&#10010;</span>
  <span style="right:34%;top:8%;font-size:28px;transform:rotate(-8deg)">&#10052;</span>
  <span style="right:28%;bottom:26%;font-size:36px;transform:rotate(10deg)">&#128213;</span>
  <span style="left:44%;top:60%;font-size:30px;transform:rotate(-14deg)">&#8612;</span>
  <span style="left:8%;top:70%;font-size:26px;transform:rotate(8deg)">&#9650;</span>
</div>
"""

def injetar_css_login():
    st.markdown(CSS_LOGIN, unsafe_allow_html=True)

def html_card(titulo, sub="", tag="", extra_html=""):
    partes = []
    partes.append('<div class="card">')
    if tag:
        partes.append(f'<span class="tag">{tag}</span>')
    partes.append(f'<div class="card-titulo">{titulo}</div>')
    if sub:
        partes.append(f'<div class="card-sub">{sub}</div>')
    if extra_html:
        partes.append(extra_html)
    partes.append('</div>')
    return "".join(partes)

# =====================================================================
# 3. MOTOR DE GERACAO DO WORD (em memoria, para download)
# =====================================================================
def gerar_documento_word(questoes_selecionadas, config, incluir_gabarito,
                         modo_acessibilidade, mostrar_descritor, titulo_prova):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = config.get("fonte", "Arial")
    style.paragraph_format.space_after = Pt(config.get("espacamento_pt", 0))

    margem_inches = config.get("margem_cm", 1.5) / 2.54
    for section in doc.sections:
        section.top_margin = Inches(margem_inches)
        section.bottom_margin = Inches(margem_inches)
        section.left_margin = Inches(margem_inches)
        section.right_margin = Inches(margem_inches)

    if modo_acessibilidade:
        font.size = Pt(16)
    else:
        font.size = Pt(config.get("tamanho_fonte", 11))

        tabela_cab = doc.add_table(rows=3, cols=3)
        tabela_cab.style = 'Table Grid'

        celula_escola = tabela_cab.cell(0, 0)
        celula_escola.merge(tabela_cab.cell(0, 2))
        p_escola = celula_escola.paragraphs[0]
        p_escola.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_escola = p_escola.add_run(f"{config.get('escola', 'ESCOLA PADRAO')}")
        run_escola.bold = True
        run_escola.font.size = Pt(config.get("tamanho_fonte", 11) + 2)

        if config.get("mostrar_aluno", True):
            celula_aluno = tabela_cab.cell(1, 0)
            celula_aluno.merge(tabela_cab.cell(1, 2))
            p_aluno = celula_aluno.paragraphs[0]
            p_aluno.add_run(f" ALUNO(A): {'_' * 75}")

        celula_turma = tabela_cab.cell(2, 0)
        if config.get("mostrar_turma", True):
            celula_turma.paragraphs[0].add_run(" TURMA: ____________")

        celula_data = tabela_cab.cell(2, 1)
        if config.get("mostrar_data", True):
            celula_data.paragraphs[0].add_run(" DATA: ___/___/20___")

        celula_prof = tabela_cab.cell(2, 2)
        prof_nome = config.get('professor', '').strip()
        if prof_nome:
            celula_prof.paragraphs[0].add_run(f" PROF.: {prof_nome}")

        doc.add_paragraph()
        if titulo_prova.strip():
            p_titulo_prova = doc.add_paragraph()
            p_titulo_prova.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_tit = p_titulo_prova.add_run(titulo_prova.upper())
            run_tit.bold = True
            run_tit.font.size = Pt(config.get("tamanho_fonte", 11) + 2)

        doc.add_paragraph()

        if config.get("usar_duas_colunas", True):
            new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            sectPr = new_section._sectPr
            cols = parse_xml(r'<w:cols w:num="2" w:space="708" w:sep="1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            sectPr.append(cols)

    for i, q in enumerate(questoes_selecionadas):
        p_q = doc.add_paragraph()
        p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        tag_descritor = ""
        if mostrar_descritor and str(q.get('tema', '')).startswith("D") and " - " in q.get('tema', ''):
            codigo = q['tema'].split(" - ")[0]
            tag_descritor = f" ({codigo})"

        num_formatado = str(i + 1).zfill(2)
        p_q.add_run(f"{num_formatado}.{tag_descritor} ").bold = True

        if q.get('enunciado', '').strip():
            linhas_enunciado = q['enunciado'].strip().split('\n')
            p_q.add_run(linhas_enunciado[0])
            for linha in linhas_enunciado[1:]:
                p_extra = doc.add_paragraph(linha)
                p_extra.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        caminho_img = q.get('imagem', '')
        if caminho_img and imagem_existe(caminho_img):
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()

                dados_img = imagem_bytes(caminho_img)
                img_temp = Image.open(io.BytesIO(dados_img))
                largura_px, altura_px = img_temp.size
                img_temp.close()

                limite_largura = 7.0 if config.get("usar_duas_colunas", True) else 10.0
                limite_altura = 6.0
                proporcao = largura_px / altura_px if altura_px else 1

                if proporcao >= 1.2:
                    run_img.add_picture(io.BytesIO(dados_img), width=Cm(limite_largura))
                else:
                    run_img.add_picture(io.BytesIO(dados_img), height=Cm(limite_altura))
            except Exception:
                p_erro = doc.add_paragraph("[Aviso: Erro ao renderizar a imagem.]")
                p_erro.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if q.get('pergunta_direta', '').strip():
            p_pergunta = doc.add_paragraph()
            p_pergunta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_pergunta.add_run(q['pergunta_direta']).bold = True

        alternativas = q.get('alternativas', [])
        if len(alternativas) > 0:
            if isinstance(alternativas, dict):
                alts_keys = list(alternativas.keys())
            else:
                alts_keys = [chr(65 + i) for i in range(len(alternativas))]
            if modo_acessibilidade:
                for letra in alts_keys:
                    texto_alt = alternativas[letra] if isinstance(alternativas, dict) else alternativas[ord(letra) - 65]
                    p_alt = doc.add_paragraph(f"{letra}) {texto_alt}")
                    p_alt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                for letra in alts_keys:
                    texto_alt = alternativas[letra] if isinstance(alternativas, dict) else alternativas[ord(letra) - 65]
                    p_alt = doc.add_paragraph(f"{letra}) {texto_alt}")
                    p_alt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_alt.paragraph_format.left_indent = Inches(0.2)
            doc.add_paragraph()
        else:
            p_linhas = doc.add_paragraph("\n________________________________________________________")
            p_linhas.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_linhas2 = doc.add_paragraph("________________________________________________________\n")
            p_linhas2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if incluir_gabarito:
        doc.add_page_break()
        gab_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
        gab_sectPr = gab_section._sectPr
        gab_cols = parse_xml(r'<w:cols w:num="1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        gab_sectPr.append(gab_cols)

        p_gab = doc.add_paragraph()
        p_gab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_gab.add_run("GABARITO OFICIAL\n").bold = True

        tabela_gab = doc.add_table(rows=1, cols=2)
        tabela_gab.style = 'Table Grid'
        hdr_cells = tabela_gab.rows[0].cells
        hdr_cells[0].text = 'Questão'
        hdr_cells[1].text = 'Resposta'

        for i, q in enumerate(questoes_selecionadas):
            num_formatado = str(i + 1).zfill(2)
            row_cells = tabela_gab.add_row().cells
            row_cells[0].text = f"Questão {num_formatado}"
            row_cells[1].text = q.get('gabarito', '')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =====================================================================
# 3.1 JS INJETADO (componente): menu movel + swipe + cookie de login
#     Roda dentro de um iframe do mesmo dominio e controla o pai.
# =====================================================================
def _html_js_movel(usuario=None, token=None, limpar=False):
    import json as _json

    payload = None
    if limpar:
        payload = {"limpar": True}
    elif usuario and token:
        payload = {"usuario": usuario, "token": token}

    _BODY = r"""(function(){
  var d = document;
  var W = window;
  var COOKIE = "ei_usuario";

  function isMobile() { return matchMedia("(max-width: 768px)").matches; }
  function el(sel) { return d.querySelector(sel); }
  function side() { return el('[data-testid="stSidebar"]'); }

  // Aberta = transform 'none' (o botao de recolher permanece no DOM mesmo recolhido).
  function isOpen() {
    var s = side();
    if (!s) return false;
    var t = getComputedStyle(s).transform || '';
    if (t === 'none' || t === '') return true;
    var m = t.match(/matrix\(([-\d.,\s]+)\)/);
    if (!m) return true;
    var tx = parseFloat(m[1].split(',')[4]) || 0;
    return tx > -1;
  }
  // Streamlit expoe um <div> wrapper; o clique efetivo e no <button> interno.
  function sbBtn(t) { var w = el('[data-testid="' + t + '"]'); if (!w) return null; return w.querySelector('button') || w; }
  function openSb() { var b = sbBtn('stExpandSidebarButton'); if (b) b.click(); }
  function closeSb() { var b = sbBtn('stSidebarCollapseButton'); if (b) b.click(); }
  function toggleSb() { isOpen() ? closeSb() : openSb(); }

  // --- menu flutuante (sempre que a sidebar estiver oculta; no celular, sempre) ---
  function ensureFab() {
    var side = el('[data-testid="stSidebar"]');
    if (!side || !side.querySelector('button')) return;
    if (el('#ei-mobile-menu')) return;
    var fab = d.createElement('button');
    fab.id = 'ei-mobile-menu';
    fab.innerHTML = '&#9776;';
    fab.setAttribute('aria-label', 'Abrir menu');
    fab.addEventListener('click', function(ev) { ev.stopPropagation(); ev.preventDefault(); toggleSb(); });
    d.body.appendChild(fab);
  }
  function syncFab() {
    var fab = el('#ei-mobile-menu');
    if (!fab) return;
    var show = isMobile() || !isOpen();
    fab.style.display = show ? 'flex' : 'none';
  }

  // --- gesto: arrastar a partir de qualquer lugar abre; empurrar p/ esq fecha ---
  var drag = null;

  // Bloqueia a navegacao de voltar por gesto (overscroll + touch-action),
  // no desktop impede recolher a barra e desenha a "alca" de arrastar no celular.
  function injectStyles() {
    if (d.getElementById('ei-styles')) return;
    var s = d.createElement('style');
    s.id = 'ei-styles';
    s.textContent =
      'html, body { overscroll-behavior-x: contain !important; }' +
      '@media (min-width: 769px) {' +
      ' [data-testid="stSidebarCollapseButton"], [data-testid="stExpandSidebarButton"] { display: none !important; }' +
      ' #ei-edge, #ei-mobile-menu { display: none !important; }' +
      ' section[data-testid="stSidebar"] { transform: none !important;' +
      ' width: 300px !important; min-width: 300px !important; max-width: 300px !important; }' +
      '}' +
      '#ei-edge { position: fixed; left: 0; top: 0; bottom: 0; width: 30px;' +
      ' z-index: 2147483646; pointer-events: none;' +
      ' display: flex; align-items: center; justify-content: center;' +
      ' -webkit-tap-highlight-color: transparent; }' +
      '#ei-edge::before { content: ""; width: 6px; height: 84px; border-radius: 6px;' +
      ' background: rgba(120,120,120,0.45); box-shadow: 0 0 0 5px rgba(120,120,120,0.10);' +
      ' animation: ei-edge-pulse 1.9s ease-in-out infinite; }' +
      '@keyframes ei-edge-pulse { 0%, 100% { opacity: .5; transform: translateX(0); }' +
      ' 50% { opacity: 1; transform: translateX(2px); } }';
    (d.head || d.documentElement).appendChild(s);
  }

  function ensureOverlay() {
    if (!isMobile() || el('#ei-edge')) return;
    if (!side()) return;
    var o = d.createElement('div');
    o.id = 'ei-edge';
    d.body.appendChild(o);
    syncOverlay();
  }
  function syncOverlay() {
    var o = el('#ei-edge');
    if (o) o.style.display = (isMobile() && !isOpen()) ? 'flex' : 'none';
    syncFab();
  }

  // largura real da sidebar (le o translateX do estado recolhido do Streamlit)
  function sbWidth() {
    var s = side();
    if (!s) return 300;
    var t = getComputedStyle(s).transform || '';
    var m = t.match(/matrix\(([-\d.,\s]+)\)/);
    if (m) {
      var tx = Math.abs(parseFloat(m[1].split(',')[4]) || 0);
      if (tx > 5) return tx;
    }
    return s.offsetWidth || 300;
  }

  // Prende o rodape (login + Sair) ao fim da sidebar, sem folga vertical.
  function pinFooter() {
    var sc = el('section[data-testid="stSidebar"] [data-testid="stSidebarContent"]');
    var vb = sc ? sc.querySelector('[data-testid="stVerticalBlock"]') : null;
    var f = el('div[class*="st-key-sb_footer"]');
    if (!sc || !vb || !f) return;
    if (vb.style.minHeight) vb.style.minHeight = '';
    var fw = f.parentElement;
    if (!fw || fw.parentElement !== vb) return;
    fw.style.marginTop = 'auto';
    vb.style.display = 'flex';
    vb.style.flexDirection = 'column';
    var sb = sc.getBoundingClientRect();
    var fb = f.getBoundingClientRect();
    var spare = (sb.top + sc.clientHeight - 8) - fb.bottom;
    if (spare > 0) vb.style.minHeight = (vb.getBoundingClientRect().height + spare) + 'px';
  }

  function inXScroll(e) {
    if (!e.target || !e.target.closest) return false;
    return !!e.target.closest('.cal-scroll, [data-testid="stDataFrame"], [data-testid="stTable"]');
  }

  // Abrir do gesto em qualquer lugar, exceto dentro de areas de rolagem
  // horizontal que ainda tenham conteudo para tras (o gesto rola o conteudo).
  function canOpenHere(e) {
    if (!e.target || !e.target.closest) return true;
    var sc = e.target.closest('.cal-scroll, [data-testid="stDataFrame"], [data-testid="stTable"]');
    if (!sc) return true;
    return sc.scrollLeft <= 2;
  }

  function onStart(e) {
    if (!isMobile() || e.touches.length !== 1) return;
    var t = e.touches[0];
    drag = { sx: t.clientX, sy: t.clientY, dx: 0, dy: 0, mode: null, w: sbWidth() };
  }
  function onMove(e) {
    if (!drag) return;
    var t = e.touches[0];
    var dx = t.clientX - drag.sx;
    var dy = t.clientY - drag.sy;
    drag.dx = dx; drag.dy = dy;
    var open = isOpen();
    if (!drag.mode) {
      var horiz = Math.abs(dx) > 10 && Math.abs(dx) > Math.abs(dy) * 1.15;
      if (!horiz) return;
      if (!open && canOpenHere(e)) drag.mode = 'open';
      else if (open && !inXScroll(e)) drag.mode = 'close';
      else { drag = null; return; }
    }
    e.preventDefault();
    var s = side(); if (!s) return;
    var w = drag.w;
    s.style.transition = 'none';
    if (drag.mode === 'open') {
      // largura fixa (300px): o conteudo nao se adapta, apenas a tela vai
      // revelando a barra conforme o dedo puxa para a direita
      s.style.width = w + 'px'; s.style.minWidth = w + 'px'; s.style.maxWidth = w + 'px';
      s.style.transform = 'translateX(' + Math.max(-w, -w + drag.dx) + 'px)';
    } else {
      s.style.transform = 'translateX(' + Math.max(-w, drag.dx) + 'px)';
    }
  }
  function onEnd(e) {
    if (!drag) return;
    var was = drag; drag = null;
    var s = side(); if (!s) return;
    var w = was.w;
    var lim = Math.max(60, w * 0.25);
    if (was.mode === 'open') {
      if (was.dx > lim) commitOpen(s, w); else reset(s);
    } else if (was.mode === 'close') {
      if (was.dx < -lim) commitClose(s, w); else reset(s);
    }
    syncOverlay();
  }
  function commitOpen(s, w) {
    s.style.transition = 'none';
    s.style.width = w + 'px'; s.style.minWidth = w + 'px'; s.style.maxWidth = w + 'px';
    s.style.transform = 'translateX(0)';
    openSb();
    setTimeout(function() { reset(s); syncOverlay(); }, 400);
  }
  function commitClose(s, w) {
    s.style.transition = 'none';
    s.style.transform = 'translateX(-' + w + 'px)';
    closeSb();
    setTimeout(function() { reset(s); syncOverlay(); }, 400);
  }
  function reset(s) {
    s.style.transition = ''; s.style.width = ''; s.style.minWidth = ''; s.style.maxWidth = ''; s.style.transform = '';
  }

  d.addEventListener('touchstart', onStart, {passive: true});
  d.addEventListener('touchmove', onMove, {passive: false});
  d.addEventListener('touchend', onEnd, {passive: true});
  d.addEventListener('touchcancel', onEnd, {passive: true});

  // --- auto-ocultar apos navegar pelo menu (so no celular) ---
  function watchNav() {
    if (!isMobile()) return;
    var side = el('[data-testid="stSidebar"]');
    if (!side || side.__eiWatch) return;
    side.__eiWatch = true;
    side.addEventListener('click', function(ev) {
      var b = ev.target.closest('button');
      if (!b) return;
      if (b.getAttribute('data-testid') === 'stSidebarCollapseButton') return;
      setTimeout(function() { if (isMobile() && isOpen()) closeSb(); }, 450);
    });
  }

  // --- cookie de login (lembrar da conta) ---
  function setCookie(name, value) {
    d.cookie = name + "=" + encodeURIComponent(value) + "; path=/; max-age=2592000; SameSite=Lax";
  }
  function clearCookie(name) {
    d.cookie = name + "=; path=/; max-age=0";
  }
  function getCookie(name) {
    var m = d.cookie.match(new RegExp("(?:^|; )" + name + "=([^;]*)"));
    return m ? decodeURIComponent(m[1]) : "";
  }
  function syncCookie() {
    var p = window.__EI_PAYLOAD;
    if (p && p.limpar) { clearCookie(COOKIE); return; }
    if (p && p.usuario) { setCookie(COOKIE, p.usuario + "|" + p.token); return; }
    // Sem payload: ponte de autologin. O cookie so e legivel no navegador,
    // entao passamos o valor pela URL (ei_auth) e o servidor valida.
    var val = getCookie(COOKIE);
    if (!val) return;
    if (W.__eiRedirected) return;
    W.__eiRedirected = true;
    var search = W.location.search || '';
    if (search.indexOf('ei_auth=') !== -1) {
      // Tentativa anterior falhou: limpa o cookie para evitar loop.
      clearCookie(COOKIE);
      return;
    }
    var sep = search.indexOf('?') === -1 ? '?' : '&';
    W.location.replace(W.location.pathname + search + sep + 'ei_auth=' + encodeURIComponent(val));
  }
  W.__eiSyncCookie = syncCookie;

  injectStyles();
  ensureFab();
  syncFab();
  watchNav();
  syncCookie();

  // No celular a barra abre sozinha assim que o app carrega (somente 1a vez).
  var didInitOpen = false;
  function forceInitialOpen() {
    if (didInitOpen || !isMobile()) return;
    var s = side();
    if (!s) return;
    if (isOpen()) { didInitOpen = true; return; }
    var b = sbBtn('stExpandSidebarButton');
    if (b) { b.click(); didInitOpen = true; }
  }
  // No desktop a barra e sempre visivel: se por qualquer motivo fechou, reabre.
  function enforceDesktop() {
    if (isMobile()) return;
    if (!side()) return;
    if (!isOpen()) { var b = sbBtn('stExpandSidebarButton'); if (b) b.click(); }
  }
  W.addEventListener('resize', function() { forceInitialOpen(); enforceDesktop(); watchNav(); syncOverlay(); pinFooter(); });

  W.__ei = {
    isOpen: isOpen, openSb: openSb, closeSb: closeSb, toggleSb: toggleSb,
    sbWidth: sbWidth, isMobile: isMobile, syncCookie: syncCookie
  };

  var obs = new MutationObserver(function() {
    ensureFab(); watchNav(); ensureOverlay(); forceInitialOpen(); enforceDesktop(); syncOverlay(); pinFooter();
  });
  if (d.body) obs.observe(d.body, { childList: true, subtree: true });
  setTimeout(function() { ensureOverlay(); forceInitialOpen(); enforceDesktop(); syncOverlay(); pinFooter(); }, 400);
  setInterval(function() { ensureOverlay(); forceInitialOpen(); enforceDesktop(); syncOverlay(); pinFooter(); }, 500);
})();
"""

    _payload_js = _json.dumps(payload)
    _payload_lit = _json.dumps(_payload_js)
    _body_lit = _json.dumps(_BODY)

    return f"""<script>
(function(){{
  var W = window.parent;
  var d = W.document;
  if (!d || !d.documentElement) return;
  try {{ W.__EI_PAYLOAD = {_payload_js}; if (W.__eiSyncCookie) W.__eiSyncCookie(); }} catch (e) {{}}
  if (d.getElementById('ei-inj')) return;
  var s = d.createElement('script');
  s.id = 'ei-inj';
  s.textContent = 'window.__EI_PAYLOAD = (' + {_payload_lit} + ');\\n' + {_body_lit};
  (d.head || d.documentElement).appendChild(s);
}})();
</script>"""


def injetar_js_movel(usuario=None, token=None, limpar=False):
    st.iframe(_html_js_movel(usuario, token, limpar), width=1, height=1)

# =====================================================================
# 3.2 PONTE DE AUTOLOGIN POR COMPONENTE (sem redirect e sem piscar a
#     tela de login): um iframe 1x1 do mesmo dominio le o cookie
#     `ei_usuario` e entrega {usuario, token} direto ao servidor via
#     `streamlit:setComponentValue` (widget json).
# =====================================================================
_EI_BRIDGE_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "componentes", "ei_bridge")
_ei_bridge = components.declare_component("ei_bridge", path=_EI_BRIDGE_DIR)

# Componente de Mapeamento de Sala (arrastar e soltar). O HTML envia a
# nova disposicao das carteiras de volta via streamlit:setComponentValue.
_SALA_BRIDGE_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "componentes", "sala_bridge")
_sala_bridge = components.declare_component("sala_bridge", path=_SALA_BRIDGE_DIR)


def cores_marca():
    """Retorna (cor_principal, cor_secundaria) do tema visual ativo."""
    cfg = carregar_config()
    tema = cfg.get("tema_visual", "")
    if tema in ("roxo", "branco_novo"):
        return "#7d3fe0", "#5b2ea6"
    if tema == "teste":
        return "#8b5cf6", "#4c1d95"
    presets = {
        "blue": ("#1f538d", "#14375e"),
        "green": ("#1d7a46", "#104f2c"),
        "dark-blue": ("#1a3a6b", "#0e2340"),
    }
    cor_tema = cfg.get("cor_tema", "blue")
    if cor_tema in presets:
        return presets[cor_tema]
    return (cfg.get("cor_principal", "#1f538d"),
            cfg.get("cor_secundaria", "#14375e"))


def _html_fallback_manual(tempo_ms=4000):
    """Redireciona o pai para ?ei_manual=1 se a ponte nao entregar nada."""
    return f"""<script>
(function(){{
  var W = window.parent;
  if (!W || !W.location) return;
  setTimeout(function() {{
    var search = W.location.search || '';
    var sep = search.indexOf('?') === -1 ? '?' : '&';
    W.location.replace(W.location.pathname + search + sep + 'ei_manual=1');
  }}, {tempo_ms});
}})();
</script>"""


def tela_verificando():
    injetar_css_login()
    aviso_banco()
    with st.container(key="login_card"):
        st.markdown(
            '<div class="login-logo">'
            f'<img class="login-img" src="data:image/png;base64,{LOGO_INICIO_B64}" '
            'alt="Exame Inteligente" /></div>'
            '<div class="login-sub">Verificando a sessao... Aguarde um instante.</div>',
            unsafe_allow_html=True)
        st.markdown(
            '<a href="?ei_manual=1" '
            'style="display:inline-block;margin-top:.25rem;color:#6b7280;'
            'font-size:.85rem;text-decoration:none;">Entrar manualmente</a>',
            unsafe_allow_html=True)
    st.iframe(_html_fallback_manual(), width=1, height=1)

# =====================================================================
# 4. MOTOR EXCEL (relatorio de avaliacao, em memoria)
# =====================================================================
def gerar_excel_avaliacao(av):
    gabarito = av['gabarito']
    qtd_q = len(gabarito)
    notas = av['notas_alunos']

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Relatorio de Notas"

    font_bold = Font(bold=True)
    align_center = Alignment(horizontal="center", vertical="center")
    align_left = Alignment(horizontal="left", vertical="center")
    border_thin = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))

    fill_adequado = PatternFill(start_color="00B050", end_color="00B050", fill_type="solid")
    fill_intermediario = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
    fill_critico = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    fill_muito_critico = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
    fill_vazio = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")

    col_fim = 5 + qtd_q

    def apply_border(cell):
        cell.border = border_thin

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_fim)
    c_titulo = ws.cell(row=1, column=1, value=f"ESCOLA DE ENSINO FUNDAMENTAL - {av['titulo'].upper()}")
    c_titulo.font = Font(bold=True, size=14)
    c_titulo.alignment = align_center

    ws.merge_cells("C3:E3")
    c = ws.cell(row=3, column=3, value="TURMA:")
    c.font = font_bold; c.alignment = align_center; apply_border(c)
    ws.merge_cells("F3:I3")
    c = ws.cell(row=3, column=6, value=av['turma'])
    c.alignment = align_center; apply_border(c)

    ws.merge_cells("K3:M3")
    c = ws.cell(row=3, column=11, value="DISCIPLINA")
    c.font = font_bold; c.alignment = align_center; apply_border(c)
    ws.merge_cells("N3:P3")
    c = ws.cell(row=3, column=14, value="MATEMATICA")
    c.alignment = align_center; apply_border(c)

    ws.merge_cells("S3:U3")
    c = ws.cell(row=3, column=19, value="DATA:")
    c.font = font_bold; c.alignment = align_center; apply_border(c)
    ws.merge_cells("V3:X3")
    c = ws.cell(row=3, column=22, value=av.get('data', ''))
    c.alignment = align_center; apply_border(c)

    for range_str in ["C3:E3", "F3:I3", "K3:M3", "N3:P3", "S3:U3", "V3:X3"]:
        for row in ws[range_str]:
            for cell in row:
                apply_border(cell)

    ws.merge_cells("F5:G5")
    c = ws.cell(row=5, column=6, value="SITUACAO")
    c.font = font_bold; c.alignment = align_center; apply_border(c)
    ws.cell(row=5, column=7).border = border_thin

    niveis_resumo = [("ADEQUADO", fill_adequado), ("INTERMEDIARIO", fill_intermediario),
                     ("CRITICO", fill_critico), ("MUITO CRITICO", fill_muito_critico)]
    for i, (nome_n, cor_n) in enumerate(niveis_resumo):
        r = 6 + i
        ws.merge_cells(start_row=r, start_column=6, end_row=r, end_column=7)
        c_n = ws.cell(row=r, column=6, value=nome_n)
        c_n.font = font_bold; c_n.fill = cor_n; c_n.alignment = align_center
        ws.cell(row=r, column=7).border = border_thin
        ws.cell(row=r, column=6).border = border_thin
        c_val = ws.cell(row=r, column=8, value=0)
        c_val.font = font_bold; c_val.alignment = align_center; apply_border(c_val)

    ws.merge_cells("K5:O5")
    c = ws.cell(row=5, column=11, value="Classifique as categorias")
    c.fill = fill_intermediario; c.alignment = align_center
    for cell in ws["K5:O5"][0]:
        apply_border(cell)

    legendas = [("ADEQUADO", fill_adequado, 90, 100), ("INTERMEDIARIO", fill_intermediario, 70, 90),
                ("CRITICO", fill_critico, 50, 70), ("MUITO CRITICO", fill_muito_critico, 0, 50)]
    for i, (nome_l, cor_l, v_min, v_max) in enumerate(legendas):
        r = 6 + i
        ws.merge_cells(start_row=r, start_column=11, end_row=r, end_column=12)
        c_nome = ws.cell(row=r, column=11, value=nome_l)
        c_nome.font = font_bold; c_nome.alignment = align_center
        ws.cell(row=r, column=11).border = border_thin
        ws.cell(row=r, column=12).border = border_thin
        c_cor = ws.cell(row=r, column=13)
        c_cor.fill = cor_l; apply_border(c_cor)
        c_min = ws.cell(row=r, column=14, value=v_min)
        c_min.alignment = align_center; apply_border(c_min)
        c_a = ws.cell(row=r, column=15, value="A")
        c_a.alignment = align_center; apply_border(c_a)
        c_max = ws.cell(row=r, column=16, value=v_max)
        c_max.alignment = align_center; apply_border(c_max)

    r_cfg = 11
    ws.merge_cells(start_row=r_cfg, start_column=1, end_row=r_cfg, end_column=2)
    c = ws.cell(row=r_cfg, column=1, value="NUMERO DE QUESTOES")
    c.font = font_bold; apply_border(c)
    ws.cell(row=r_cfg, column=2).border = border_thin
    c_q = ws.cell(row=r_cfg, column=3, value=qtd_q)
    c_q.font = font_bold; c_q.alignment = align_center; apply_border(c_q)

    ws.merge_cells(start_row=r_cfg + 1, start_column=1, end_row=r_cfg + 1, end_column=2)
    c = ws.cell(row=r_cfg + 1, column=1, value="GABARITO")
    apply_border(c)
    ws.cell(row=r_cfg + 1, column=2).border = border_thin
    for i, q in enumerate(gabarito):
        c_gab = ws.cell(row=r_cfg + 1, column=3 + i, value=q['resposta_correta'].upper())
        c_gab.font = font_bold; c_gab.alignment = align_center; apply_border(c_gab)

    ws.merge_cells(start_row=r_cfg + 2, start_column=1, end_row=r_cfg + 2, end_column=2)
    c = ws.cell(row=r_cfg + 2, column=1, value="DESCRITOR")
    apply_border(c)
    ws.cell(row=r_cfg + 2, column=2).border = border_thin
    for i, q in enumerate(gabarito):
        c_desc = ws.cell(row=r_cfg + 2, column=3 + i, value=q['descritor'])
        c_desc.font = Font(size=9, bold=True)
        c_desc.alignment = align_center; apply_border(c_desc)

    r_hdr = 14
    headers = ["N", "NOME DO ALUNO"] + [str(i + 1) for i in range(qtd_q)] + ["ACT", "PORC", "SITUACAO"]
    for i, h in enumerate(headers):
        c_h = ws.cell(row=r_hdr, column=i + 1, value=h)
        c_h.font = font_bold
        c_h.alignment = align_center if i != 1 else align_left
        apply_border(c_h)

    alunos_ordenados = sorted(notas.items(), key=lambda x: x[0])
    distribuicao = {"ADEQUADO": 0, "INTERMEDIARIO": 0, "CRITICO": 0, "MUITO CRITICO": 0}

    r_aluno = 15
    for idx_a, (aluno, dados) in enumerate(alunos_ordenados, 1):
        c_n = ws.cell(row=r_aluno, column=1, value=idx_a)
        c_n.alignment = align_center; apply_border(c_n)
        c_nome = ws.cell(row=r_aluno, column=2, value=aluno)
        apply_border(c_nome)

        respostas_dadas = dados.get('respostas_dadas', '')
        resp_padded = respostas_dadas.ljust(qtd_q, ' ')
        acertos = 0

        for i, q in enumerate(gabarito):
            letra_marcada = resp_padded[i].upper()
            correta = q['resposta_correta'].upper()

            c_resp = ws.cell(row=r_aluno, column=3 + i)
            c_resp.alignment = align_center; apply_border(c_resp)

            if letra_marcada == ' ':
                c_resp.fill = fill_vazio
                c_resp.value = ""
            else:
                c_resp.value = letra_marcada
                if letra_marcada == correta:
                    c_resp.fill = fill_adequado
                    acertos += 1
                else:
                    c_resp.fill = fill_muito_critico

        c_act = ws.cell(row=r_aluno, column=3 + qtd_q, value=acertos)
        c_act.alignment = align_center; apply_border(c_act)

        perc = (acertos / qtd_q) * 100 if qtd_q > 0 else 0
        c_porc = ws.cell(row=r_aluno, column=4 + qtd_q, value=f"{perc:.1f}".replace('.', ','))
        c_porc.alignment = align_center; apply_border(c_porc)

        nivel = "MUITO CRITICO"
        if perc >= 90:
            nivel = "ADEQUADO"
        elif perc >= 70:
            nivel = "INTERMEDIARIO"
        elif perc >= 50:
            nivel = "CRITICO"

        distribuicao[nivel] += 1
        c_sit = ws.cell(row=r_aluno, column=5 + qtd_q, value=nivel)
        c_sit.alignment = align_center; apply_border(c_sit)
        r_aluno += 1

    ws.cell(row=6, column=8, value=distribuicao["ADEQUADO"])
    ws.cell(row=7, column=8, value=distribuicao["INTERMEDIARIO"])
    ws.cell(row=8, column=8, value=distribuicao["CRITICO"])
    ws.cell(row=9, column=8, value=distribuicao["MUITO CRITICO"])

    ws.column_dimensions['A'].width = 4
    ws.column_dimensions['B'].width = 38
    for i in range(qtd_q):
        ws.column_dimensions[get_column_letter(3 + i)].width = 4
    ws.column_dimensions[get_column_letter(3 + qtd_q)].width = 6
    ws.column_dimensions[get_column_letter(4 + qtd_q)].width = 7
    ws.column_dimensions[get_column_letter(5 + qtd_q)].width = 18

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

# =====================================================================
# 5. AJUDA DE NAVEGACAO (query params e links)
# =====================================================================
def ir_para(pagina):
    st.query_params["pagina"] = pagina

def pagina_atual():
    params = st.query_params.get_all("pagina")
    return params[0] if params else "Dashboard"

def val_param(nome, padrao=""):
    vals = st.query_params.get_all(nome)
    return vals[0] if vals else padrao

def limpar_params():
    st.query_params.clear()

def link_para(pagina, base="."):
    qs = []
    for k in list(st.query_params.keys()):
        for v in st.query_params.get_all(k):
            qs.append(f"{k}={v}")
    qs.append(f"pagina={pagina}")
    return base + "?" + "&".join(qs)

# =====================================================================
# 6. SIDEBAR (menu lateral)
# =====================================================================
NAV_PLANEJAMENTO = [
    ("Dashboard", "In\u00edcio"),
    ("Grade Semanal", "Grade Semanal"),
    ("Turmas e Alunos", "Turmas e Alunos"),
    ("Mapeamento de Sala", "Mapeamento de Sala"),
    ("Central de Planos", "Central de Planos"),
    ("Anotacoes", "Lembretes"),
]
NAV_AVALIACOES = [
    ("Central de Questões", "Central de Questões"),
    ("Avaliações e Estatísticas", "Avaliações e Estatísticas"),
    ("Atividades", "Atividades"),
    ("Configurações", "Configurações"),
]

ICONES_NAV = {
    "Dashboard": "\u2302",
    "Grade Semanal": "\u25a6",
    "Turmas e Alunos": "\u25c8",
    "Mapeamento de Sala": "\u25c9",
    "Central de Planos": "\u2691\uFE0E",
    "Anotacoes": "\u270e\uFE0E",
    "Central de Questões": "\u2630",
    "Avaliações e Estatísticas": "\u25a5",
    "Atividades": "\u25a4",
    "Configurações": "\u2699\uFE0E",
}

def montar_sidebar():
    atual = pagina_atual()
    logo_b64 = LOGO_ABA_B64_ROXO if carregar_config().get("tema_visual") in ("roxo", "teste") else LOGO_ABA_B64
    with st.sidebar:
        st.markdown(
            '<div class="logo-aba">'
            f'<img src="data:image/png;base64,{logo_b64}" alt="Exame Inteligente" /></div>',
            unsafe_allow_html=True)
        st.markdown('<div class="nav-secao">Planejamento</div>', unsafe_allow_html=True)
        for chave, rotulo in NAV_PLANEJAMENTO:
            icone = ICONES_NAV.get(chave, "")
            rotulo_html = f"{icone}  {rotulo}" if icone else rotulo
            if st.button(rotulo_html, key=f"nav_p_{chave}",
                         type="primary" if atual == chave else "secondary",
                         use_container_width=True):
                st.query_params["pagina"] = chave
                st.rerun()
        st.markdown('<div class="nav-secao">Avaliações</div>', unsafe_allow_html=True)
        for chave, rotulo in NAV_AVALIACOES:
            icone = ICONES_NAV.get(chave, "")
            rotulo_html = f"{icone}  {rotulo}" if icone else rotulo
            if st.button(rotulo_html, key=f"nav_a_{chave}",
                         type="primary" if atual == chave else "secondary",
                         use_container_width=True):
                st.query_params["pagina"] = chave
                st.rerun()
        with st.container(key="sb_footer"):
            st.markdown("---")
            user = usuario_atual()
            if user:
                conta = conta_atual(user)
                nome_conta = conta.get("nome", user) if conta else user
                st.caption(f"Logado: {nome_conta}\n@{user}")
                if st.button("Sair", key="sair_conta", use_container_width=True):
                    st.session_state.pop("usuario", None)
                    st.session_state.pop("bem_vindo", None)
                    st.session_state.pop("login_modo", None)
                    st.session_state["deslogado_manual"] = True
                    st.query_params.clear()
                    st.rerun()

# =====================================================================
# 6.1 TELA DE LOGIN / CRIACAO DE CONTA
# =====================================================================
def tela_login():
    injetar_css_login()
    aviso_banco()
    with st.container(key="login_card"):
        st.markdown(
            '<div class="login-logo">'
            f'<img class="login-img" src="data:image/png;base64,{LOGO_INICIO_B64}" '
            'alt="Exame Inteligente" /></div>'
            '<div class="login-sub">Entre com a sua conta para acessar turmas, '
            'planos de aula, questões e provas.</div>',
            unsafe_allow_html=True)

        modo = st.session_state.get("login_modo")
        c_tab = st.columns(2, gap="small")
        if c_tab[0].button("Entrar", key="tab_entrar",
                           type="primary" if modo == "Entrar" else "secondary",
                           use_container_width=True):
            st.session_state["login_modo"] = "Entrar"
            st.rerun()
        if c_tab[1].button("Criar Conta", key="tab_criar",
                           type="primary" if modo == "Criar Conta" else "secondary",
                           use_container_width=True):
            st.session_state["login_modo"] = "Criar Conta"
            st.rerun()

        if modo is None:
            st.markdown(
                '<div class="login-hint">Escolha acima como quer continuar: '
                'entre com a sua conta ou crie uma nova.</div>',
                unsafe_allow_html=True)
            return

        if modo == "Entrar":
            with st.form("form_login"):
                login_user = st.text_input("Email ou usuario", key="login_user",
                                           placeholder="voce@escola.com")
                login_senha = st.text_input("Senha", type="password", key="login_senha",
                                            placeholder="digite sua senha")
                entrar = st.form_submit_button("Entrar na conta", type="primary",
                                               use_container_width=True)
            if entrar:
                _tentar_religar_banco()
                conta = autenticar_usuario(login_user, login_senha)
                if conta:
                    st.session_state["usuario"] = conta["usuario"]
                    st.query_params.clear()
                    st.rerun()
                elif banco_falho():
                    st.error("Nao foi possivel validar o login: o banco de "
                             "dados externo esta indisponivel. Assim que ele "
                             "voltar, tente entrar novamente.")
                else:
                    st.error("Email, usuario ou senha incorretos.")
        else:
            _cad = st.session_state
            etapa = _cad.get("cad_etapa", "dados")
            if etapa == "dados":
                st.caption("Crie sua conta com email e senha. Enviaremos um codigo "
                           "para confirmar o seu email.")
                with st.form("form_cadastro"):
                    cad_email = st.text_input("Email", key="cad_email",
                                              placeholder="voce@escola.com")
                    cad_senha = st.text_input("Crie uma senha", type="password",
                                              key="cad_senha",
                                              placeholder="Minimo 6 caracteres")
                    cad_confirmar = st.text_input("Confirme a senha", type="password",
                                                  key="cad_confirmar",
                                                  placeholder="Digite novamente")
                    criar = st.form_submit_button("Enviar codigo de verificacao",
                                                  type="primary",
                                                  use_container_width=True)
                if criar:
                    email_c = (cad_email or "").strip().lower()
                    if not email_c:
                        st.error("Preencha o seu email.")
                    elif not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email_c):
                        st.error("Digite um email valido.")
                    elif len(cad_senha) < 6:
                        st.error("A senha precisa ter pelo menos 6 caracteres.")
                    elif cad_senha != cad_confirmar:
                        st.error("As senhas nao conferem.")
                    elif email_existe(email_c):
                        st.error("Ja existe uma conta com esse email.")
                    else:
                        ok, msg, codigo_visivel = _enviar_codigo_verificacao(email_c)
                        if ok:
                            _cad["cad_pend_email"] = email_c
                            _cad["cad_pend_senha"] = cad_senha
                            _cad["cad_etapa"] = "codigo"
                            _cad["cad_pend_enviou"] = True
                            if codigo_visivel:
                                _cad["cad_pend_codigo_teste"] = codigo_visivel
                            st.rerun()
                        else:
                            st.error(msg)
            else:
                email_c = _cad.get("cad_pend_email", "")
                if _cad.get("cad_pend_enviou"):
                    st.success("Enviamos um codigo para o seu email. "
                               "Confira a caixa de entrada (e o spam).")
                    _cad["cad_pend_enviou"] = False
                if _cad.get("cad_pend_codigo_teste"):
                    st.info(f"**Modo teste (sem SMTP):** seu codigo e "
                            f"**{_cad['cad_pend_codigo_teste']}**.")
                    _cad.pop("cad_pend_codigo_teste", None)
                st.caption(f"Confirme o codigo enviado para <b>{esc(email_c)}</b>.",
                           unsafe_allow_html=True)
                with st.form("form_confirmar_email"):
                    cad_codigo = st.text_input("Codigo de 6 digitos", key="cad_codigo",
                                               placeholder="000000")
                    confirmar = st.form_submit_button("Verificar e criar conta",
                                                      type="primary",
                                                      use_container_width=True)
                if confirmar:
                    if not (cad_codigo or "").strip():
                        st.error("Digite o codigo recebido.")
                    else:
                        valido, msg = _validar_codigo(email_c, cad_codigo)
                        if valido:
                            base = re.sub(r"[^a-zA-Z0-9_.-]", "_",
                                          email_c.split("@")[0].lower())
                            user_c = base or "usuario"
                            i = 2
                            while usuario_existe(user_c):
                                user_c = f"{base}{i}"
                                i += 1
                            criar_conta("", user_c, _cad.get("cad_pend_senha", ""),
                                        email=email_c, email_verificado=True)
                            _cad.pop("cad_etapa", None)
                            _cad.pop("cad_pend_email", None)
                            _cad.pop("cad_pend_senha", None)
                            _cad["usuario"] = user_c
                            _cad["bem_vindo"] = True
                            st.query_params.clear()
                            st.rerun()
                        else:
                            st.error(msg)
                if st.button("Reenviar codigo", key="cad_reenviar",
                             use_container_width=True):
                    ok, msg, codigo_visivel = _enviar_codigo_verificacao(email_c)
                    if ok:
                        st.session_state["cad_pend_enviou"] = True
                        if codigo_visivel:
                            st.session_state["cad_pend_codigo_teste"] = codigo_visivel
                    else:
                        st.error(msg)
                    st.rerun()
                if st.button("Voltar e trocar o email", key="cad_voltar",
                             use_container_width=True):
                    _cad["cad_etapa"] = "dados"
                    _cad.pop("cad_pend_email", None)
                    _cad.pop("cad_pend_senha", None)
                    st.rerun()

        if st.button("Recolher tela", key="login_recolher",
                     use_container_width=False):
            st.session_state.pop("login_modo", None)
            st.rerun()

# =====================================================================
# 7. CALENDARIO EM HTML
# =====================================================================
def parse_cal(cal):
    hoje = datetime.now()
    try:
        mes, ano = [int(x) for x in cal.split("/")]
    except Exception:
        mes, ano = hoje.month, hoje.year
    if mes < 1 or mes > 12:
        mes, ano = hoje.month, hoje.year
    return mes, ano

def _ordem_aula(aula):
    try:
        return int(str(aula).split("a")[0])
    except Exception:
        return 0


def html_aulas_dia(grade_horaria, dia_str, horarios=None):
    try:
        nome_dia = DIAS_COMPLETOS[datetime.strptime(dia_str, "%d/%m/%Y").weekday()]
    except Exception:
        nome_dia = ""
    aulas = sorted([i for i in grade_horaria if i["dia"] == nome_dia],
                   key=lambda x: _ordem_aula(x.get("aula")))
    if not aulas:
        return (f'<div class="card-sub">Sem aulas em {esc(nome_dia or dia_str)}.</div>'
                if nome_dia else
                '<div class="card-sub">Grade vazia. Cadastre aulas na Grade Semanal.</div>')
    cards = ""
    for a in aulas:
        cor_val = a.get("cor") or ""
        cor_hex = dict(CORES_AULA).get(cor_val, cor_val if cor_val.startswith("#") else "")
        cor_cls = " gcell-colorido" if cor_hex else ""
        cor_style = f' style="background:{esc(cor_hex)}"' if cor_hex else ""
        tempo = ""
        m = re.search(r"(\d+)", str(a.get("aula", "")))
        if m and horarios:
            h = horarios.get(m.group(1))
            if h and h.get("inicio") and h.get("fim"):
                tempo = f' <span class="gcell-tempo">{esc(h["inicio"])} - {esc(h["fim"])}</span>'
        cards += (f'<div class="gcell adias-cell{cor_cls}"{cor_style}>'
                  f'<div class="gcell-aula">{esc(a.get("aula",""))}{tempo}</div>'
                  f'<div class="gcell-turma">{esc(a.get("turma",""))}</div>'
                  f'<div class="gcell-disc">{esc(a.get("disciplina","Geral"))}</div></div>')
    return f'<div class="adias">{cards}</div>'

# =====================================================================
# 8. DIALOGOS REUTILIZAVEIS
# =====================================================================
CORES_POSTIT = [
    ("Amarelo Classico", "#fff3a3"),
    ("Verde Menta", "#b5e7a0"),
    ("Azul Ceu", "#a0c4ff"),
    ("Lilas Suave", "#cdb4db"),
    ("Rosa Pastel", "#ffc6ff"),
]

CORES_AULA = [
    ("Azul", "#dbeafe"),
    ("Verde", "#dcfce7"),
    ("Amarelo", "#fef9c3"),
    ("Laranja", "#ffedd5"),
    ("Rosa", "#fce7f3"),
    ("Lilas", "#ede9fe"),
    ("Ciano", "#cffafe"),
    ("Cinza", "#f1f5f9"),
]

def form_postit(nota=None, uid="frm"):
    grade = carregar_grade()
    turmas = ["Geral"] + sorted(list(set([i["turma"] for i in grade])))
    is_edicao = nota is not None

    titulo = st.text_input("Titulo do lembrete:", value=nota.get("titulo", "") if is_edicao else "",
                           key=f"{uid}_titulo")
    turma = st.selectbox("Atrelar a turma:", turmas,
                         index=turmas.index(nota["turma"]) if is_edicao and nota.get("turma") in turmas else 0,
                         key=f"{uid}_turma")
    nomes_cores = [n for n, _ in CORES_POSTIT]
    cores_vals = [c for _, c in CORES_POSTIT]
    cor_padrao = nota.get("cor", "#fff3a3") if is_edicao else "#fff3a3"
    idx_cor = cores_vals.index(cor_padrao) if cor_padrao in cores_vals else 0
    cor = st.select_slider("Cor do post-it:", options=nomes_cores, value=nomes_cores[idx_cor],
                           key=f"{uid}_cor")
    cor_hex = cores_vals[nomes_cores.index(cor)]
    conteudo = st.text_area("Anotacao / Lembrete:",
                            value=nota.get("conteudo", "") if is_edicao else "",
                            key=f"{uid}_conteudo")

    data_default = None
    if is_edicao:
        for campo in ("data", "data_criacao"):
            raw = nota.get(campo)
            if raw:
                try:
                    data_default = datetime.strptime(raw, "%d/%m/%Y").date()
                    break
                except (ValueError, TypeError):
                    data_default = None
    if data_default is None:
        data_default = datetime.now().date()
    data_lembrete = st.date_input("Data do lembrete:", value=data_default, key=f"{uid}_data")

    if st.button("Salvar Post-it", type="primary", use_container_width=True, key=f"{uid}_salvar"):
        if not titulo.strip() or not conteudo.strip():
            st.error("Preencha o titulo e o conteudo.")
            return
        anotacoes = carregar_anotacoes()
        data_str = data_lembrete.strftime("%d/%m/%Y")
        if is_edicao:
            for n in anotacoes:
                if n.get("id") == nota.get("id"):
                    n["titulo"] = titulo.strip()
                    n["conteudo"] = conteudo.strip()
                    n["turma"] = turma
                    n["cor"] = cor_hex
                    n["data"] = data_str
        else:
            novo_id = max([n.get("id", 0) for n in anotacoes], default=0) + 1
            anotacoes.append({
                "id": novo_id, "titulo": titulo.strip(), "conteudo": conteudo.strip(),
                "turma": turma, "cor": cor_hex,
                "data": data_str,
                "data_criacao": datetime.now().strftime("%d/%m/%Y")
            })
        salvar_anotacoes(anotacoes)
        st.rerun()

def form_plano(data_str, index_plano, uid="frm"):
    planos = carregar_planos()
    if data_str not in planos or index_plano >= len(planos[data_str]):
        st.info("Plano nao encontrado.")
        return
    plano = dict(planos[data_str][index_plano])

    st.markdown(
        f"**Data:** {data_str} | **Horario:** {plano.get('horario','')} | "
        f"**Turma:** {plano.get('turma','')} | **Disciplina:** {plano.get('disciplina','Geral')}")

    novo_tema = st.text_input("Titulo / Tema da aula:", value=plano.get("tema", ""),
                              key=f"{uid}_tema")

    campos = [
        ("Metodologia / Atividades", "metodologia"),
        ("Objetivos Especificos", "objetivos"),
        ("Procedimentos", "procedimentos"),
        ("Habilidade(s) da BNCC", "habilidades"),
        ("Competencia Geral", "comp_geral"),
        ("Competencias Especificas", "comp_especifica"),
        ("Recursos Necessarios", "recursos"),
        ("Avaliacao", "avaliacao"),
        ("Observacoes", "observacoes"),
    ]
    novos_valores = {}
    abas = st.tabs(["Metodologia & Objetivos", "Estrutura BNCC", "Recursos & Avaliacao"])
    mapa_abas = {
        "metodologia": 0, "objetivos": 0, "procedimentos": 0,
        "habilidades": 1, "comp_geral": 1, "comp_especifica": 1,
        "recursos": 2, "avaliacao": 2, "observacoes": 2,
    }
    for rotulo, chave in campos:
        with abas[mapa_abas[chave]]:
            novos_valores[chave] = st.text_area(rotulo, value=plano.get(chave, ""),
                                                key=f"{uid}_{chave}")

    c1, c2 = st.columns(2)
    if c1.button("Salvar Alteracoes", type="primary", use_container_width=True, key=f"{uid}_salvar"):
        if not novo_tema.strip():
            st.error("O tema da aula e obrigatorio.")
            return
        plano["tema"] = novo_tema.strip()
        for chave, valor in novos_valores.items():
            plano[chave] = valor.strip()
        planos[data_str][index_plano] = plano
        salvar_planos(planos)
        st.rerun()
    if c2.button("Excluir Plano", use_container_width=True, key=f"{uid}_excluir"):
        del planos[data_str][index_plano]
        if len(planos[data_str]) == 0:
            del planos[data_str]
        salvar_planos(planos)
        st.rerun()

# =====================================================================
# 9. TELA: DASHBOARD
# =====================================================================
DISCIPLINAS_COMUNS = [
    "Matematica", "Portugues", "Historia", "Geografia", "Ciencias",
    "Biologia", "Fisica", "Quimica", "Ingles", "Espanhol", "Artes",
    "Educacao Fisica", "Filosofia", "Sociologia", "Ensino Religioso",
    "Informatica",
]

def render_calendario(planos, anotacoes=None):
    hoje = datetime.now()
    cal = st.session_state.get("dash_cal", hoje.strftime("%m/%Y"))
    mes, ano = parse_cal(cal)

    dias_lembrete = set()
    for n in (anotacoes or []):
        d = n.get("data") or n.get("data_criacao") or ""
        if len(d) == 10 and d[2] == "/" and d[5] == "/":
            dias_lembrete.add(d)

    def mes_offset(delta):
        m, a = mes, ano
        m += delta
        if m > 12:
            m, a = 1, a + 1
        elif m < 1:
            m, a = 12, a - 1
        return m, a

    st.markdown('<div class="cal-widgets-marker"></div>', unsafe_allow_html=True)

    c1, c2, c3 = st.columns([1, 2.2, 1])
    voltar = c1.button("\u25c0", key="cal_prev", use_container_width=True)
    avancar = c3.button("\u25b6", key="cal_next", use_container_width=True)

    trocou_mes = False
    if voltar:
        mes, ano = mes_offset(-1)
        trocou_mes = True
    if avancar:
        mes, ano = mes_offset(1)
        trocou_mes = True
    if trocou_mes:
        st.session_state["dash_cal"] = f"{mes:02d}/{ano}"

    with c2:
        opcoes_pular = [f"{MESES_PT[m]} / {a}"
                        for a in (ano - 1, ano, ano + 1) for m in range(1, 13)]
        valor_atual = f"{MESES_PT[mes]} / {ano}"
        escolha = c2.selectbox("Mes / Ano", opcoes_pular,
                               index=opcoes_pular.index(valor_atual),
                               key=f"cal_ir_{mes}_{ano}", label_visibility="collapsed")
        if escolha != valor_atual:
            m_nome, a_str = escolha.rsplit(" / ", 1)
            mes, ano = MESES_PT.index(m_nome), int(a_str)
            trocou_mes = True
            st.session_state["dash_cal"] = f"{mes:02d}/{ano}"

    dias_cal = st.columns(7)
    for i, d in enumerate(DIAS_CURTO):
        dias_cal[i].markdown(
            f'<div style="text-align:center;font-size:.7rem;text-transform:uppercase;'
            f'letter-spacing:.08em;color:var(--cor-cinza);"> {d}</div>',
            unsafe_allow_html=True)

    dia_clicado = None
    matriz = calendar.monthcalendar(ano, mes)
    for semana in matriz:
        cols = st.columns(7)
        for i, dia in enumerate(semana):
            if dia == 0:
                continue
            data_str = f"{dia:02d}/{mes:02d}/{ano}"
            eh_hoje = (dia, mes, ano) == (hoje.day, hoje.month, hoje.year)
            tem_plano = len(planos.get(data_str, [])) > 0
            tipo = "primary" if tem_plano else "secondary"
            rotulo = f"**{dia}**" if eh_hoje else str(dia)
            if cols[i].button(rotulo, key=f"cal_d_{data_str}", type=tipo,
                              use_container_width=True):
                dia_clicado = data_str

    if dias_lembrete:
        estilos = ["<style>"]
        for d in sorted(dias_lembrete):
            if d[6:10] != f"{ano:04d}" or int(d[3:5]) != mes:
                continue
            classe = "st-key-cal_d_" + d.replace("/", "-")
            estilos.append(
                f'div[class*="{classe}"] button {{position: relative;}}'
                f'div[class*="{classe}"] button::after {{'
                f'content: ""; position: absolute; top: 3px; right: 3px; '
                f'width: 6px; height: 6px; border-radius: 50%; '
                f'background: #e67e22; box-shadow: 0 0 0 1px rgba(255,255,255,.55);}}')
        estilos.append("</style>")
        st.markdown("".join(estilos), unsafe_allow_html=True)

    if trocou_mes and dia_clicado is None:
        return f"01/{mes:02d}/{ano}"
    return dia_clicado

@st.fragment
def fragmento_dashboard():
    hoje = datetime.now()
    planos = carregar_planos()
    grade = carregar_grade()
    anotacoes = carregar_anotacoes()

    dia_selecionada = st.session_state.get("dash_dia", hoje.strftime("%d/%m/%Y"))
    try:
        datetime.strptime(dia_selecionada, "%d/%m/%Y")
    except ValueError:
        dia_selecionada = hoje.strftime("%d/%m/%Y")
        st.session_state["dash_dia"] = dia_selecionada

    with st.container(key="dash_wrap"):
        with st.container(key="card_dash_cal"):
            novo_dia = render_calendario(planos, anotacoes)
            if novo_dia:
                st.session_state["dash_dia"] = novo_dia
                dia_selecionada = novo_dia

        with st.container(key="card_dash_mural"):
            topo = st.columns([3, 1])
            topo[0].markdown(
                '<div class="dash-sec"><span class="dash-sec-bar"></span>'
                '<span class="dash-sec-txt">Lembretes</span></div>',
                unsafe_allow_html=True)
            with topo[1].popover("+ Novo", key="dash_novo_pop", use_container_width=True):
                form_postit(None, uid="dash_novo")
            if not anotacoes:
                st.caption("Nenhum post-it. Crie um no botao + Novo.")
            else:
                with st.container(key="mural_scroll", height=320):
                    for nota in anotacoes:
                        cor = nota.get("cor") or "#fff3a3"
                        txt = cor_texto_legivel(cor)
                        conteudo_full = nota.get("conteudo", "")
                        data_pt = nota.get("data") or nota.get("data_criacao") or ""
                        st.markdown(
                            f'<div class="postit ei-pt" style="background:{cor};color:{txt};" '
                            f'title="{esc(conteudo_full)}">'
                            f'<div class="pt-titulo">\U0001f4cb {esc(nota.get("titulo", ""))}</div>'
                            f'<div class="pt-tag">{esc(data_pt)}</div>'
                            f'<div class="pt-conteudo">{esc(conteudo_full)}</div>'
                            f'</div>',
                            unsafe_allow_html=True)
                        b1, b2 = st.columns(2)
                        with b1.popover("Editar", key=f"dash_edit_pop_{nota.get('id')}", use_container_width=True):
                            form_postit(nota, uid=f"dash_edit_{nota.get('id')}")
                        if b2.button("Excluir", key=f"dash_del_{nota.get('id')}", use_container_width=True):
                            anotacoes = [n for n in anotacoes if n.get("id") != nota.get("id")]
                            salvar_anotacoes(anotacoes)
                            st.rerun()

        with st.container(key="card_dash_grade"):
            st.markdown(
                '<div class="dash-sec"><span class="dash-sec-bar"></span>'
                '<span class="dash-sec-txt">Minha Agenda</span></div>',
                unsafe_allow_html=True)
            st.markdown(
                f'<div class="dash-sec-sub">Aulas programadas para o dia '
                f'{esc(dia_selecionada)}:</div>',
                unsafe_allow_html=True)
            planos_dia = planos.get(dia_selecionada, [])
            horarios = carregar_horarios_aulas()
            if planos_dia:
                for i, plano in enumerate(planos_dia):
                    horario_plano = plano.get("horario", "")
                    tempo_plano = ""
                    m = re.search(r"(\d+)", str(horario_plano))
                    if m:
                        h = horarios.get(m.group(1))
                        if h and h.get("inicio") and h.get("fim"):
                            tempo_plano = (f' <span class="dash-tl-time">'
                                           f'{esc(h["inicio"])} - {esc(h["fim"])}</span>')
                    st.markdown(
                        f'<div class="dash-tl-item">'
                        f'<div class="dash-tl-hora">\U0001f550 {esc(horario_plano)}{tempo_plano}</div>'
                        f'<div class="dash-tl-txt">{esc(plano.get("turma",""))} — '
                        f'{esc(plano.get("disciplina","Geral"))}</div>'
                        f'<div class="dash-tl-sub">{esc(plano.get("tema","Sem tema"))}</div></div>',
                        unsafe_allow_html=True)
                    with st.popover("Visualizar / Editar", key=f"dash_plano_pop_{dia_selecionada}_{i}",
                                    use_container_width=True):
                        form_plano(dia_selecionada, i, uid=f"dash_plano_{dia_selecionada}_{i}")
            st.markdown(html_aulas_dia(grade, dia_selecionada, horarios), unsafe_allow_html=True)

def tela_dashboard(config):
    hoje = datetime.now()

    dia_param = val_param("dia", "")
    try:
        datetime.strptime(dia_param, "%d/%m/%Y")
        st.session_state["dash_dia"] = dia_param
    except ValueError:
        if "dash_dia" not in st.session_state:
            st.session_state["dash_dia"] = hoje.strftime("%d/%m/%Y")

    cal_param = val_param("cal", "")
    if cal_param:
        st.session_state["dash_cal"] = cal_param
    if "dash_cal" not in st.session_state:
        st.session_state["dash_cal"] = hoje.strftime("%m/%Y")

    st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600;700&display=swap');

/* Dashboard Inicial: Design System (Poppins + cores da marca) */
.dash-hello {
    font-family: 'Poppins', system-ui, sans-serif;
    font-weight: 700; font-size: 1.85rem; letter-spacing: -0.01em;
    color: var(--cor-p); display: flex; align-items: center; gap: .5rem;
}
.dash-data { color: var(--cor-cinza); font-size: .92rem; margin: .1rem 0 1.1rem 0; }

.dash-resumo {
    display: grid; grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));
    gap: 14px; margin-bottom: 16px;
}
.dash-resumo-card {
    font-family: 'Poppins', system-ui, sans-serif;
    border-radius: 14px; padding: 16px 20px;
    display: flex; align-items: center; gap: 14px;
    background: linear-gradient(100deg, var(--cor-p) 0%, var(--cor-pd) 100%);
    color: var(--btn-fg);
    box-shadow: 0 6px 16px rgba(0,0,0,.12);
}
.dash-resumo-icone { font-size: 1.5rem; opacity: .92; }
.dash-resumo-rotulo { font-size: .95rem; font-weight: 500; opacity: .92; }
.dash-resumo-valor { font-size: 1.7rem; font-weight: 700; line-height: 1.1; }

/* Titulos dos paineis do Dashboard */
.dash-sec { display: flex; align-items: center; gap: .5rem; margin: .2rem 0 .9rem 0; }
.dash-sec-bar { width: 5px; height: 20px; border-radius: 3px; background: var(--cor-p); flex: none; }
.dash-sec-txt { font-family: 'Poppins', system-ui, sans-serif; font-weight: 700; font-size: 1.02rem; color: var(--cor-texto); }
.dash-sec-sub { font-size: .82rem; color: var(--cor-cinza); margin: -.5rem 0 .85rem 0; }

/* Minha Agenda: linha do tempo */
.dash-tl-item {
    background: var(--card-bg); border: 1px solid var(--borda);
    border-left: 4px solid var(--cor-p); border-radius: 10px;
    padding: 10px 14px; margin-bottom: 10px;
    box-shadow: 0 1px 3px rgba(0,0,0,.06);
    transition: box-shadow .18s ease, transform .18s ease;
}
.dash-tl-item:hover { box-shadow: 0 4px 14px rgba(0,0,0,.12); transform: translateY(-1px); }
.dash-tl-hora { font-weight: 700; font-size: .85rem; color: var(--cor-p); }
.dash-tl-time { font-weight: 400; font-size: .85rem; color: var(--cor-cinza); margin-left: 6px; white-space: nowrap; }
.dash-tl-txt { font-size: .9rem; font-weight: 600; color: var(--cor-texto); }
.dash-tl-sub { font-size: .78rem; color: var(--cor-cinza); }
</style>""", unsafe_allow_html=True)

    nome_prof = primeiro_nome_professor(config)
    st.markdown(
        f'<div class="dash-hello">Olá, {esc(nome_prof)}! '
        f'<span aria-hidden="true">\U0001f44b</span></div>',
        unsafe_allow_html=True)
    st.markdown(
        f'<div class="dash-data">Hoje é {hoje.strftime("%d/%m/%Y")}</div>',
        unsafe_allow_html=True)

    grade = carregar_grade()
    anotacoes = carregar_anotacoes()
    try:
        nome_dia_hoje = DIAS_COMPLETOS[hoje.weekday()]
    except Exception:
        nome_dia_hoje = ""
    n_aulas = len([i for i in grade if i.get("dia") == nome_dia_hoje])
    n_lem = len(anotacoes)
    st.markdown(
        '<div class="dash-resumo">'
        '<div class="dash-resumo-card"><span class="dash-resumo-icone">\U0001f4da</span>'
        '<div><div class="dash-resumo-rotulo">Aulas Hoje</div>'
        f'<div class="dash-resumo-valor">{n_aulas}</div></div></div>'
        '<div class="dash-resumo-card"><span class="dash-resumo-icone">\U0001f4cb</span>'
        '<div><div class="dash-resumo-rotulo">Lembretes do Dia</div>'
        f'<div class="dash-resumo-valor">{n_lem}</div></div></div>'
        '</div>', unsafe_allow_html=True)

    fragmento_dashboard()

# =====================================================================
# 10. TELA: GRADE SEMANAL
# =====================================================================
def _form_horarios_aulas(max_aulas):
    horarios = carregar_horarios_aulas()
    t_inicio = []
    t_fim = []
    with st.form(key="form_horarios_aulas"):
        for num in range(1, max_aulas + 1):
            atual = horarios.get(str(num)) or {}
            ini = atual.get("inicio") or horario_padrao_aula(num)[0]
            fim = atual.get("fim") or horario_padrao_aula(num)[1]
            c1, c2 = st.columns(2)
            t_inicio.append(c1.time_input(f"{num}ª aula",
                                          value=hora_para_time(ini), key=f"ha_ini_{num}",
                                          step=timedelta(minutes=5)))
            t_fim.append(c2.time_input("até",
                                       value=hora_para_time(fim), key=f"ha_fim_{num}",
                                       step=timedelta(minutes=5)))
        salvo = st.form_submit_button("Salvar horários", use_container_width=True)
    if salvo:
        novos = {}
        for i, num in enumerate(range(1, max_aulas + 1)):
            novos[str(num)] = {
                "inicio": t_inicio[i].strftime("%H:%M"),
                "fim": t_fim[i].strftime("%H:%M"),
            }
        salvar_horarios_aulas(novos)
        st.success("Horários das aulas salvos com sucesso!")

def tela_grade_semanal():
    st.markdown("## Grade Semanal")
    grade = carregar_grade()
    max_aulas = carregar_config_grade()

    with st.expander("➕ Adicionar Aulas", expanded=False):
        c_cfg, c_add = st.columns([1, 2.4])
        with c_cfg:
            with st.container(key="card_grade_cfg"):
                nova_qtd = st.selectbox("Qtd. de aulas/dia:", [str(i) for i in range(1, 16)],
                                        index=max_aulas - 1 if 1 <= max_aulas <= 15 else 5)
                if st.button("Atualizar", use_container_width=True):
                    salvar_config_grade(int(nova_qtd))
                    st.rerun()
                st.markdown("---")
                st.markdown("**Horários das aulas**")
                if st.button("Configurar horário das aulas", use_container_width=True,
                             key="btn_cfg_horarios"):
                    st.session_state["cfg_horarios_aberto"] = not st.session_state.get(
                        "cfg_horarios_aberto", False)
                    st.rerun()
                if st.session_state.get("cfg_horarios_aberto"):
                    _form_horarios_aulas(max_aulas)

        with c_add:
            with st.container(key="card_grade_add"):
                st.markdown("**Adicionar aula a grade**")
                l1, l2 = st.columns([1, 1])
                nova_turma = l1.text_input("Turma", placeholder="Ex: 7o A", key="grade_turma")
                opcoes_disc = DISCIPLINAS_COMUNS + ["Outra disciplina..."]
                disc_sel = l2.selectbox("Disciplina", opcoes_disc,
                                        index=0, key="grade_disc_sel")
                nova_disc = ""
                if disc_sel == "Outra disciplina...":
                    nova_disc = l2.text_input("Digite a disciplina:", key="grade_disc_outra")
                else:
                    nova_disc = disc_sel
                dia = st.radio("Dia da semana", DIAS_UTEIS, horizontal=True)
                aulas_sel = st.multiselect(
                    "Selecione as aulas", [f"{i}a Aula" for i in range(1, max_aulas + 1)],
                    key="grade_aulas")
                if st.button("+ Adicionar", type="primary", use_container_width=True):
                    if not nova_turma.strip() or not nova_disc.strip():
                        st.error("Preencha a turma e a disciplina.")
                    elif not aulas_sel:
                        st.error("Marque pelo menos um horario.")
                    else:
                        conflitos = [f"{a} (ocupada: {i['turma']})" for a in aulas_sel
                                     for i in grade if i["dia"] == dia and i["aula"] == a]
                        if conflitos:
                            st.warning("Conflito de horario:\n" + "\n".join(conflitos))
                        else:
                            novo_id = max([i.get("id", 0) for i in grade], default=0) + 1
                            for a in aulas_sel:
                                grade.append({"id": novo_id, "turma": nova_turma.strip(),
                                              "disciplina": nova_disc.strip(), "dia": dia, "aula": a})
                                novo_id += 1
                            salvar_grade(grade)
                            st.rerun()

    dias_preenchidos = [d for d in DIAS_UTEIS if any(x["dia"] == d for x in grade)]
    if not dias_preenchidos:
        st.info("Nenhum horario cadastrado. Clique em 'Adicionar Aulas' para montar sua grade.")
        return

    def ordem_aula(aula):
        try:
            return int(str(aula).split("a")[0])
        except Exception:
            return 0

    with st.container(key="grade_dias"):
        horarios = carregar_horarios_aulas()
        n_dias = len(dias_preenchidos)
        ratios = [1.1] + [1.9] * n_dias
        por_dia = {dia: sorted([x for x in grade if x["dia"] == dia],
                               key=lambda x: ordem_aula(x["aula"])) for dia in dias_preenchidos}
        num_periodos = max([ordem_aula(x["aula"]) for x in grade], default=1)

        def range_periodo(num):
            h = horarios.get(str(num))
            if h and h.get("inicio") and h.get("fim"):
                return f'{h["inicio"]} - {h["fim"]}'
            return ""

        colunas = st.columns(ratios)
        colunas[0].markdown('<div class="gh-head">Horário</div>', unsafe_allow_html=True)
        for j, dia in enumerate(dias_preenchidos, start=1):
            colunas[j].markdown(
                f'<div class="gh-head">{esc(dia.split("-")[0])}</div>',
                unsafe_allow_html=True)

        for num in range(1, num_periodos + 1):
            colunas = st.columns(ratios)
            colunas[0].markdown(
                f'<div class="gh-time"><div class="gh-time-aula">{esc(f"{num}ª Aula")}</div>'
                f'<div class="gh-time-range">{esc(range_periodo(num))}</div></div>',
                unsafe_allow_html=True)
            for j, dia in enumerate(dias_preenchidos, start=1):
                item = next((x for x in por_dia[dia] if ordem_aula(x["aula"]) == num), None)
                if item is None:
                    colunas[j].markdown('<div class="gh-vazio"></div>', unsafe_allow_html=True)
                else:
                    cor_val = item.get("cor") or ""
                    cor_hex = dict(CORES_AULA).get(cor_val, cor_val if cor_val.startswith("#") else "")
                    cor_style = f' style="background:{cor_hex}"' if cor_hex else ""
                    cor_cls = " gcell-colorido" if cor_hex else ""
                    with colunas[j]:
                        with st.container(key=f"cell_{item['id']}"):
                            st.markdown(
                                f'<div class="gcell{cor_cls}"{cor_style}>'
                                f'<div class="gcell-aula">{item["aula"]}</div>'
                                f'<div class="gcell-turma">{item["turma"]}</div>'
                                f'<div class="gcell-disc">{item.get("disciplina","Geral")}</div></div>',
                                unsafe_allow_html=True)
                            if st.button("✕", key=f"del_{item['id']}", help="Remover este horario"):
                                grade = [g for g in grade if g.get("id") != item["id"]]
                                salvar_grade(grade)
                                st.rerun()
                            with st.popover("🎨", key=f"pal_{item['id']}", help="Trocar a cor"):
                                cols = st.columns(4)
                                for i, (nome, hex_cor) in enumerate(CORES_AULA):
                                    with cols[i % 4]:
                                        if st.button("", key=f"gc_{item['id']}_{nome}",
                                                     help=nome, use_container_width=True):
                                            for g in grade:
                                                if g.get("id") == item["id"]:
                                                    g["cor"] = nome
                                            salvar_grade(grade)
                                            st.rerun()

# =====================================================================
# 11. TELA: TURMAS E ALUNOS
# =====================================================================
def _grade_mapeamento_inicial(layout, fileiras, colunas, alunos):
    if (layout and layout.get("fileiras") == fileiras
            and layout.get("colunas") == colunas
            and isinstance(layout.get("grade"), list)
            and len(layout["grade"]) == fileiras
            and all(isinstance(r, list) and len(r) == colunas
                    for r in layout["grade"])):
        return [list(r) for r in layout["grade"]]
    fila = list(alunos)
    grade = []
    for _ in range(fileiras):
        linha = []
        for _ in range(colunas):
            linha.append(fila.pop(0) if fila else None)
        grade.append(linha)
    return grade


def tela_mapeamento():
    grade = carregar_grade()
    turmas_grade = sorted(list(set(i["turma"] for i in grade)))
    if not turmas_grade:
        st.error("Nenhuma turma cadastrada na Grade Semanal. "
                 "Cadastre turmas la primeiro.")
        return
    dados_turmas = carregar_turmas()
    mapeamentos = carregar_mapeamento()

    if ("map_turma" not in st.session_state
            or st.session_state["map_turma"] not in turmas_grade):
        st.session_state["map_turma"] = turmas_grade[0]
    turma = st.session_state["map_turma"]

    if st.session_state.get("map_turma_ant") != turma:
        st.session_state.pop("map_fileiras", None)
        st.session_state.pop("map_colunas", None)
        st.session_state["map_turma_ant"] = turma

    atual = mapeamentos.get(turma) or {}
    if "map_fileiras" not in st.session_state:
        st.session_state["map_fileiras"] = int(atual.get("fileiras", 4))
    if "map_colunas" not in st.session_state:
        st.session_state["map_colunas"] = int(atual.get("colunas", 5))

    fileiras_eff = int(st.session_state["map_fileiras"])
    colunas_eff = int(st.session_state["map_colunas"])

    alunos = [normalizar_nome(a) for a in dados_turmas.get(turma, [])]
    alunos = sorted({a for a in alunos if a}, key=lambda x: x.lower())
    grade_inicial = _grade_mapeamento_inicial(atual, fileiras_eff, colunas_eff,
                                              alunos)

    c_titulo, c_export = st.columns([4.5, 1], vertical_alignment="center")
    with c_titulo:
        st.markdown("## Mapeamento de Sala")
    with c_export:
        with st.popover("", icon=":material/print:", key="export_mapa_pop",
                        help="Exportar mapa de sala para Excel"):
            st.markdown("**Exportar mapa de sala**")
            cor_p, cor_s = cores_marca()
            bytes_excel = gerar_excel_mapeamento(
                turma, fileiras_eff, colunas_eff, grade_inicial, cor_p, cor_s,
                alunos)
            st.download_button("Baixar planilha (.xlsx)", data=bytes_excel,
                               file_name=f"mapa_sala_{turma}.xlsx",
                               mime="application/vnd.openxmlformats-"
                                    "officedocument.spreadsheetml.sheet",
                               use_container_width=True)

    c_t, c_f, c_c = st.columns(3)
    turma = c_t.selectbox("Turma", turmas_grade, key="map_turma")
    fileiras = int(c_f.number_input("Quantidade de fileiras", 1, 12,
                                    value=fileiras_eff, step=1,
                                    key="map_fileiras"))
    colunas = int(c_c.number_input("Carteiras por fileira", 1, 12,
                                   value=colunas_eff, step=1,
                                   key="map_colunas"))

    if not alunos:
        st.warning("Esta turma nao possui alunos cadastrados. "
                   "Cadastre alunos na tela 'Turmas e Alunos' primeiro.")
        return

    total = fileiras * colunas
    if total < len(alunos):
        st.info(f"A turma {turma} tem {len(alunos)} alunos, mas a sala tem "
                f"apenas {total} carteiras. Os alunos excedentes ficam na "
                "barra 'Alunos fora da sala' e podem entrar por arrastar e "
                "soltar.")

    st.caption("Arraste e solte os cartoes para posicionar cada aluno. "
               "Solte sobre um aluno ocupado para trocar de lugar. "
               "Solte fora das carteiras (na area 'Alunos fora da sala' ou "
               "no restante da tela) para remover o aluno da sala. "
               "Ao terminar, clique em 'Confirmar Posicoes'.")

    grade_inicial = _grade_mapeamento_inicial(atual, fileiras, colunas, alunos)
    cor_p, cor_s = cores_marca()
    resultado = _sala_bridge(turma=turma, fileiras=fileiras, colunas=colunas,
                             alunos=alunos, grade=grade_inicial,
                             cor_principal=cor_p, cor_secundaria=cor_s,
                             key=f"sala_mapa_{turma}")

    if (resultado and isinstance(resultado, dict) and resultado.get("grade")
            and isinstance(resultado["grade"], list)
            and len(resultado["grade"]) == fileiras
            and all(isinstance(r, list) and len(r) == colunas
                    for r in resultado["grade"])):
        g = resultado["grade"]
        salvo_grade = mapeamentos.get(turma, {}).get("grade")
        if g != salvo_grade:
            mapeamentos[turma] = {"fileiras": fileiras, "colunas": colunas,
                                  "grade": g}
            salvar_mapeamento(mapeamentos)
            st.success("Posicoes salvas com sucesso!")


# =====================================================================
# 11. TELA: TURMAS E ALUNOS
# =====================================================================
def _sanear_aba(nome):
    nome = re.sub(r'[\\/*?:\[\]]', "", str(nome)).strip()
    return nome[:31] or "Turma"


def nome_curto_mapa(nome):
    partes = [p for p in str(nome or "").strip().split() if p]
    if len(partes) <= 2:
        return " ".join(partes)
    if len(partes[1]) <= 3:
        return " ".join(partes[:3])
    return " ".join(partes[:2])


def nome_curto_mapa_longo(nome):
    partes = [p for p in str(nome or "").strip().split() if p]
    if len(partes) >= 3:
        return " ".join(partes[:3])
    return " ".join(partes)


def mapa_nomes_exibicao(alunos):
    freq = {}
    for n in alunos:
        b = nome_curto_mapa(n)
        freq[b] = freq.get(b, 0) + 1
    disp = {}
    for n in alunos:
        b = nome_curto_mapa(n)
        disp[n] = nome_curto_mapa_longo(n) if freq[b] > 1 else b
    return disp


def gerar_excel_mapeamento(turma, fileiras, colunas, grade, cor_p, cor_s,
                           alunos):
    """Gera o mapa de sala (fileiras em colunas verticais) como planilha
    estilizada: cabecalho com a cor da marca, bordas e separacao entre nomes.
    Nomes curtos (primeiro e segundo), incluindo o terceiro quando ha
    colisao ou quando o segundo tem 2-3 letras."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = _sanear_aba(turma) or "Mapa"
    hex_p = (cor_p or "#1F538D").lstrip("#")[:6]
    cor_topo = PatternFill(start_color=hex_p, end_color=hex_p,
                           fill_type="solid")
    cor_rot = PatternFill(start_color="EEF0F6", end_color="EEF0F6",
                          fill_type="solid")
    cor_vazio = PatternFill(start_color="F7F8FB", end_color="F7F8FB",
                            fill_type="solid")
    borda = Border(*[Side(style="thin", color="C9CFDD")] * 4)
    alinh = Alignment(horizontal="center", vertical="center", wrap_text=True)
    disp = mapa_nomes_exibicao(alunos)

    ws["A1"] = "Carteira"
    ws["A1"].fill = cor_topo
    ws["A1"].font = Font(bold=True, color="FFFFFF")
    ws["A1"].alignment = alinh
    ws["A1"].border = borda
    for j in range(fileiras):
        cel = ws.cell(row=1, column=2 + j, value=f"F{j + 1}")
        cel.fill = cor_topo
        cel.font = Font(bold=True, color="FFFFFF")
        cel.alignment = alinh
        cel.border = borda
    for i in range(colunas):
        cel = ws.cell(row=2 + i, column=1, value=f"C{i + 1}")
        cel.fill = cor_rot
        cel.alignment = alinh
        cel.border = borda
        cel.font = Font(bold=True)
        for j in range(fileiras):
            nome = None
            if j < len(grade) and i < len(grade[j]):
                nome = grade[j][i]
            cel2 = ws.cell(row=2 + i, column=2 + j,
                           value=disp.get(nome, nome_curto_mapa(nome))
                           if nome else "")
            cel2.alignment = alinh
            cel2.border = borda
            if nome:
                cel2.font = Font(bold=True)
            else:
                cel2.fill = cor_vazio
    ws.column_dimensions["A"].width = 10
    for j in range(fileiras):
        ws.column_dimensions[get_column_letter(2 + j)].width = 26
    ws.row_dimensions[1].height = 22
    for i in range(colunas):
        ws.row_dimensions[2 + i].height = 24
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def gerar_excel_alunos(dados_turmas, turmas_sel):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    cor_topo = PatternFill(start_color="1F539D", end_color="1F539D",
                           fill_type="solid")
    usado = set()
    for turma in turmas_sel:
        aba = _sanear_aba(turma)
        if aba in usado:
            aba = f"{aba[:28]}_{len(usado) + 1}"
        usado.add(aba)
        ws = wb.create_sheet(title=aba)
        ws.append(["Num", "Nome"])
        for cel in ws[1]:
            cel.fill = cor_topo
            cel.font = Font(bold=True, color="FFFFFF")
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 42
        nomes = sorted({normalizar_nome(a) for a in dados_turmas.get(turma, []) if a},
                       key=lambda x: x.lower())
        for i, n in enumerate(nomes, 1):
            ws.append([i, n])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def tela_turmas():
    grade = carregar_grade()
    dados_turmas = carregar_turmas()

    turmas_grade = sorted(list(set([i["turma"] for i in grade])))
    if not turmas_grade:
        st.error("Nenhuma turma cadastrada na Grade Semanal. Cadastre turmas la primeiro.")
        return

    if "turma_selecionada" not in st.session_state or st.session_state["turma_selecionada"] not in turmas_grade:
        st.session_state["turma_selecionada"] = turmas_grade[0]
    turma_atual = st.session_state["turma_selecionada"]

    c_titulo, c_export = st.columns([4.5, 1], vertical_alignment="center")
    with c_titulo:
        st.markdown("## Turmas e Alunos")
    with c_export:
        with st.popover("", icon=":material/print:", key="export_alunos_pop",
                        help="Exportar lista de alunos para Excel"):
            st.markdown("**Exportar alunos para planilha**")
            tipo = st.radio("O que exportar?",
                            ["Todas as turmas", "Turmas especificas"],
                            key="exp_tipo")
            turmas_sel = turmas_grade
            if tipo == "Turmas especificas":
                turmas_sel = st.multiselect("Escolha as turmas",
                                            turmas_grade,
                                            default=turmas_grade[:1],
                                            key="exp_turmas")
            if not turmas_sel:
                st.caption("Nenhuma turma selecionada.")
            else:
                bytes_excel = gerar_excel_alunos(dados_turmas, turmas_sel)
                st.download_button("Baixar planilha (.xlsx)", data=bytes_excel,
                                   file_name="alunos_por_turma.xlsx",
                                   mime="application/vnd.openxmlformats-"
                                        "officedocument.spreadsheetml.sheet",
                                   use_container_width=True)

    info_imp = st.session_state.pop("imp_info", None)
    if st.session_state.pop("imp_limpar", False):
        st.session_state.pop("imp_alunos", None)
    if info_imp:
        if info_imp[0] == "ok":
            st.success(info_imp[1])
        else:
            st.warning(info_imp[1])

    qtd_atual = len(dados_turmas.get(turma_atual, []))

    with st.container(key="painel_turma_top"):
        c_turmas, c_add = st.columns([1.5, 1.2])
        with c_turmas:
            with st.popover(f"Turma: {turma_atual}  ({qtd_atual} alunos)",
                            key="pop_turmas", use_container_width=True):
                st.markdown("**Suas Turmas**")
                for turma in turmas_grade:
                    qtd = len(dados_turmas.get(turma, []))
                    ativa = turma == turma_atual
                    if st.button(f"{turma}  ({qtd} alunos)", key=f"turma_{turma}",
                                 type="primary" if ativa else "secondary",
                                 use_container_width=True):
                        st.session_state["turma_selecionada"] = turma
                        st.rerun()
        with c_add:
            with st.popover("Adicionar Alunos", key="pop_adicionar",
                            use_container_width=True):
                arquivo = st.file_uploader(
                    "Importar alunos (Excel/CSV/TXT)",
                    type=["xlsx", "xls", "csv", "txt"], key="imp_alunos")
                if arquivo is not None:
                    resultado = processar_arquivo_alunos(arquivo, turma_atual, dados_turmas)
                    st.session_state["imp_limpar"] = True
                    st.session_state["imp_info"] = resultado
                    st.rerun()
                novo_aluno = st.text_input("Nome do aluno",
                                           placeholder="Ex: Ana Clara",
                                           key="novo_aluno")
                if st.button("+ Adicionar aluno", type="primary",
                             use_container_width=True):
                    nome = normalizar_nome(novo_aluno)
                    if not nome:
                        st.error("Digite o nome do aluno.")
                    elif nome.lower() in {normalizar_nome(a).lower()
                                          for a in dados_turmas.get(turma_atual, [])}:
                        st.warning("Este aluno ja esta cadastrado nesta turma.")
                    else:
                        dados_turmas.setdefault(turma_atual, []).append(nome)
                        salvar_turmas(dados_turmas)
                        st.rerun()
                st.divider()
                if st.button("Excluir todos os alunos desta turma",
                             use_container_width=True):
                    st.session_state["confirmar_limpar_turma"] = turma_atual
                    st.rerun()

    st.markdown(f"### Alunos da turma: {turma_atual}")

    if st.session_state.get("confirmar_limpar_turma") == turma_atual:
        st.warning(f"Tem certeza que deseja remover TODOS os "
                   f"{len(dados_turmas.get(turma_atual, []))} alunos da turma {turma_atual}?")
        c1, c2 = st.columns(2)
        if c1.button("Sim, excluir todos", type="primary"):
            dados_turmas[turma_atual] = []
            salvar_turmas(dados_turmas)
            st.session_state["confirmar_limpar_turma"] = None
            st.rerun()
        if c2.button("Cancelar"):
            st.session_state["confirmar_limpar_turma"] = None
            st.rerun()

    alunos = [normalizar_nome(a) for a in dados_turmas.get(turma_atual, [])]
    alunos = sorted([a for a in alunos if a], key=lambda x: x.lower())
    vistos, alunos_unicos = set(), []
    for a in alunos:
        chave = a.lower()
        if chave not in vistos:
            vistos.add(chave)
            alunos_unicos.append(a)
    alunos = alunos_unicos
    if not alunos:
        st.caption("Nenhum aluno cadastrado nesta turma ainda.")
    else:
        with st.container(key="lista_alunos"):
            for i, aluno in enumerate(alunos, 1):
                c1, c2, c3 = st.columns([6, 1, 1])
                c1.markdown(f'<span class="aluno-num">{i:02d}</span> {aluno}',
                            unsafe_allow_html=True)
                with c2.popover("Mover", key=f"mover_{turma_atual}_{i}",
                                help="Mover para outra turma"):
                    outras = [t for t in turmas_grade if t != turma_atual]
                    if not outras:
                        st.caption("Nao ha outra turma cadastrada.")
                    else:
                        destino = st.selectbox("Mover para", outras,
                                               key=f"mv_dest_{turma_atual}_{i}")
                        if st.button("Mover aluno", key=f"mv_ok_{turma_atual}_{i}",
                                     use_container_width=True):
                            alvo = aluno.lower()
                            dados_turmas[turma_atual] = [
                                a for a in dados_turmas[turma_atual]
                                if normalizar_nome(a).lower() != alvo]
                            dados_turmas.setdefault(destino, []).append(aluno)
                            salvar_turmas(dados_turmas)
                            st.session_state["turma_selecionada"] = destino
                            st.rerun()
                if c3.button("x", key=f"aluno_{turma_atual}_{i}", help="Excluir aluno"):
                    alvo = aluno.lower()
                    dados_turmas[turma_atual] = [
                        a for a in dados_turmas[turma_atual]
                        if normalizar_nome(a).lower() != alvo]
                    salvar_turmas(dados_turmas)
                    st.rerun()

def processar_arquivo_alunos(arquivo, turma_atual, dados_turmas):
    nome_arq = arquivo.name.lower()
    novos_raw = []
    try:
        if nome_arq.endswith((".xlsx", ".xls")):
            wb = openpyxl.load_workbook(io.BytesIO(arquivo.read()), data_only=True)
            ws = wb.active
            for row in ws.iter_rows(values_only=True):
                if not row:
                    continue
                celula = row[0]
                if celula is None or str(celula).strip() == "":
                    continue
                novos_raw.append(str(celula).strip())
        else:
            conteudo = arquivo.read()
            texto = None
            for enc in ("utf-8", "latin-1"):
                try:
                    texto = conteudo.decode(enc)
                    break
                except Exception:
                    continue
            if texto is None:
                texto = conteudo.decode("utf-8", errors="ignore")
            leitor = csv.reader(io.StringIO(texto), delimiter=";")
            for linha in leitor:
                if not linha or linha[0].strip() == "":
                    continue
                novos_raw.append(linha[0].strip())
            if len(novos_raw) <= 1:
                novos_raw = [l.strip() for l in texto.splitlines() if l.strip()]
    except Exception as e:
        return ("erro", f"Nao foi possivel ler o arquivo: {e}")

    novos, vistos = [], set()
    for nome in novos_raw:
        padrao = normalizar_nome(nome)
        if not padrao or padrao.lower() in CABECALHOS_ALUNOS:
            continue
        chave = padrao.lower()
        if chave not in vistos:
            vistos.add(chave)
            novos.append(padrao)

    if not novos:
        return ("vazio", "O arquivo parece estar vazio ou sem nomes na primeira coluna.")

    dados_turmas.setdefault(turma_atual, [])
    existentes = {normalizar_nome(a).lower() for a in dados_turmas[turma_atual]}
    adicionados = 0
    for nome in novos:
        if nome.lower() not in existentes:
            dados_turmas[turma_atual].append(nome)
            existentes.add(nome.lower())
            adicionados += 1
    salvar_turmas(dados_turmas)
    return ("ok", f"{adicionados} aluno(s) importado(s) para a turma {turma_atual}!")

# =====================================================================
# 12. TELA: CENTRAL DE PLANOS
# =====================================================================
def html_calendario_planos(planos, grade, ano, mes, turma_filtro, mostrar_audit, dark):
    dias_completos = DIAS_COMPLETOS
    base_cor = "#333333" if dark else "#d6d6d6"
    base_txt = "#ffffff" if dark else "#000000"

    linhas = ['<div class="cal-scroll"><table class="cal-table"><tr>']
    for d in DIAS_CURTO:
        linhas.append(f"<th>{d}</th>")
    linhas.append("</tr>")

    matriz = calendar.monthcalendar(ano, mes)
    for semana in matriz:
        linhas.append("<tr>")
        for coluna, dia in enumerate(semana):
            if dia == 0:
                linhas.append('<td style="border:none;padding:4px;"></td>')
                continue
            data_str = f"{dia:02d}/{mes:02d}/{ano}"
            nome_dia = dias_completos[coluna]
            aulas_esperadas = [g for g in grade if g["dia"] == nome_dia]
            if turma_filtro != "Todas as Turmas":
                aulas_esperadas = [g for g in aulas_esperadas if g["turma"] == turma_filtro]
            qtd_esperada = len(aulas_esperadas)

            planos_dia = planos.get(data_str, [])
            if turma_filtro != "Todas as Turmas":
                planos_dia = [p for p in planos_dia if p["turma"] == turma_filtro]
            qtd_registrada = len(planos_dia)

            if mostrar_audit:
                if qtd_esperada == 0:
                    cor, txt = base_cor, "#a0a0a0"
                elif qtd_registrada == 0:
                    cor, txt = "#dc3545", "#ffffff"
                elif qtd_registrada < qtd_esperada:
                    cor, txt = "#ffc107", "#000000"
                else:
                    cor, txt = "#28a745", "#ffffff"
            else:
                if qtd_registrada > 0:
                    cor, txt = "#007bff", "#ffffff"
                else:
                    cor, txt = base_cor, base_txt

            linhas.append(
                f'<td style="background:{cor};color:{txt};border:1px solid rgba(128,128,128,0.25);'
                f'text-align:center;font-weight:600;padding:8px 0;border-radius:8px;">{dia}</td>')
        linhas.append("</tr>")
    linhas.append("</table></div>")
    return "".join(linhas)

CAMPOS_PLANO = [
    ("Metodologia / Atividades", "metodologia"),
    ("Objetivos Especificos", "objetivos"),
    ("Procedimentos", "procedimentos"),
    ("Habilidade(s) da BNCC", "habilidades"),
    ("Competencia Geral", "comp_geral"),
    ("Competencias Especificas", "comp_especifica"),
    ("Recursos Necessarios", "recursos"),
    ("Avaliacao", "avaliacao"),
    ("Observacoes", "observacoes"),
]

def _montar_prompt_plano_ia(tema, disciplinas, carga, data_inicio, campos_existentes=None):
    estrutura = ('{"metodologia": "...", "objetivos": "...", "procedimentos": "...", '
                 '"habilidades": "...", "comp_geral": "...", "comp_especifica": "...", '
                 '"recursos": "...", "avaliacao": "..."}')
    linhas = [
        "Voce e um professor experiente do ensino fundamental e medio brasileiro.",
        "Elabore um plano de aula completo e detalhado com base nas informacoes abaixo.",
        f"- Tema / conteudo da aula: {tema.strip()}",
        f"- Disciplina(s): {disciplinas}",
        f"- Carga horaria: {carga}",
        f"- Data de inicio: {data_inicio}",
    ]
    campos_existentes = {k: str(v or "").strip()
                         for k, v in (campos_existentes or {}).items() if str(v or "").strip()}
    if campos_existentes:
        linhas.append("")
        linhas.append("CONTEXTO - campos ja preenchidos pelo professor (use como base:")
        for chave, valor in campos_existentes.items():
            linhas.append(f"- {chave}: {valor}")
        linhas.append("Mantenha a coerencia com esses textos, mas NAO os altere: "
                      "preencha somente os campos que ainda estao vazios.")
        linhas.append("ATENCAO: os campos ja preenchidos acima devem ficar intactos na "
                      "sua resposta (repita o mesmo texto deles no JSON, sem mudancas).")
    linhas.extend([
        "",
        "Responda APENAS com um objeto JSON valido, sem texto adicional, "
        "usando exatamente esta estrutura:",
        estrutura,
        "",
        "IMPORTANTE: NAO inclua nem preencha o campo \"observacoes\" na resposta. "
        "Ele e reservado para anotacoes manuais do professor.",
        "",
        "Regras de preenchimento de cada campo:",
        "- metodologia: descreva a abordagem didatica e as atividades principais da aula.",
        "- objetivos: liste objetivos especificos e mensuraveis para a aula.",
        "- procedimentos: descreva passo a passo como a aula se desenvolve.",
        "- habilidades: cite habilidades da BNCC relacionadas ao tema (ex: EF08MA08).",
        "- comp_geral: indique uma competencia geral da BNCC aplicavel ao tema.",
        "- comp_especifica: indique competencias especificas da BNCC.",
        "- recursos: liste os materiais e recursos necessarios para a aula.",
        "- avaliacao: descreva como a aprendizagem sera avaliada.",
        "- Escreva tudo em portugues do Brasil, de forma didatica, clara e objetiva.",
    ])
    return "\n".join(linhas)


def _extrair_objeto_json_ia(texto):
    t = (texto or "").strip()
    if t.startswith("```"):
        linhas = t.splitlines()
        t = "\n".join(linhas[1:]) if len(linhas) > 1 else ""
        if t.rstrip().endswith("```"):
            t = t.rstrip()[:-3].rstrip()
    ini = t.find("{")
    fim = t.rfind("}")
    if ini != -1 and fim != -1 and fim > ini:
        t = t[ini:fim + 1]
    return json.loads(t)


def _gerar_plano_com_ia(tema, data_inicio, carga, td_selecionadas):
    chave_salva = (carregar_chave_ia() or "").strip()
    if not chave_salva:
        st.error("Configure sua chave de API em Configuracoes > Chave de API para usar a IA.")
        if st.button("Ir para Configuracoes", key="plano_ia_ir_config"):
            ir_para("Configuracoes")
            st.rerun()
        return
    if not tema.strip():
        st.error("Preencha o campo Tema para a IA gerar o plano.")
        return
    if not td_selecionadas:
        st.error("Selecione ao menos uma turma/disciplina para a IA gerar o plano.")
        return
    try:
        datetime.strptime(data_inicio.strip(), "%d/%m/%Y")
    except ValueError:
        st.error("Formato de data invalido. Use DD/MM/AAAA.")
        return

    disciplinas = ", ".join(sorted(set(td.split(" - ")[-1] for td in td_selecionadas)))
    campos_existentes = {}
    for rotulo, chave in CAMPOS_PLANO:
        valor = str(st.session_state.get(f"plano_{chave}", "") or "").strip()
        if valor:
            campos_existentes[chave] = valor
    config = carregar_config()
    provedor = config.get("provedor_ia", "Google Gemini")
    modelo_interno = MODELO_GEMINI if provedor == "Google Gemini" else MODELO_OPENAI
    prompt = _montar_prompt_plano_ia(tema, disciplinas, carga, data_inicio, campos_existentes)
    try:
        with st.spinner("Gerando plano de aula com IA... pode levar alguns segundos."):
            if provedor == "Google Gemini":
                texto = chamar_ia_gemini(chave_salva, modelo_interno, prompt)
            else:
                texto = chamar_ia_openai(chave_salva, modelo_interno, prompt)
        dados = _extrair_objeto_json_ia(texto)
        if not isinstance(dados, dict):
            raise ValueError("A IA nao retornou um objeto valido.")
        preenchidos = 0
        vazios = 0
        for rotulo, chave in CAMPOS_PLANO:
            if chave == "observacoes":
                continue
            ja_preenchido = str(st.session_state.get(f"plano_{chave}", "") or "").strip()
            if ja_preenchido:
                continue
            vazios += 1
            valor = str(dados.get(chave, "") or "").strip()
            if valor:
                st.session_state[f"plano_{chave}"] = valor
                preenchidos += 1
        if vazios == 0:
            st.info("Todos os campos do plano ja estavam preenchidos. "
                    "A IA usou esses textos como base e nenhum campo ficou vazio.")
            return
        if preenchidos == 0:
            raise ValueError("A IA nao retornou campos validos para o plano.")
        st.success(
            f"Plano preenchido pela IA ({preenchidos} campos). "
            "Revise os valores antes de distribuir.")
        st.rerun()
    except urllib.error.HTTPError as e:
        detalhe = e.read().decode("utf-8", errors="ignore")[:300]
        msg = f"Erro {e.code} ao chamar o servico de IA. Verifique a chave.\n\n{detalhe}"
        if e.code == 404 and "model" in detalhe.lower():
            msg += ("\n\nDica: o modelo pode ter sido descontinuado pelo provedor. "
                    "Avise o suporte do app para atualizar o modelo interno.")
        st.error(msg)
    except (urllib.error.URLError, TimeoutError):
        st.error("Nao foi possivel conectar ao servico de IA. Verifique sua internet.")
    except (ValueError, json.JSONDecodeError) as e:
        st.error(f"A IA nao retornou um JSON valido. Tente novamente.\n\nDetalhe: {e}")

def tela_central_planos(config):
    st.markdown("## Central de Planos")
    grade = carregar_grade()
    if not grade:
        st.error("Sua grade esta vazia. Cadastre turmas na Grade Semanal primeiro.")
        return

    tab_criar, tab_listar = st.tabs(["Distribuir Novo Plano", "Planos deste Mes"])

    with tab_criar:
        with st.container(key="card_plano_criar"):
            tema = st.text_input("Tema / Conteudo da aula", placeholder="Ex: Funcoes do 1o grau")
            c1, c2 = st.columns(2)
            data_inicio = c1.text_input("Data de inicio (DD/MM/AAAA)",
                                        value=datetime.now().strftime("%d/%m/%Y"))
            carga = c2.selectbox("Carga horaria (aulas)", [f"{i} Aulas" for i in range(1, 11)])
            st.markdown("**Selecione as turmas e disciplinas para distribuir:**")
            turmas_disciplinas = []
            for item in grade:
                td = f"{item['turma']} - {item.get('disciplina', 'Geral')}"
                if td not in turmas_disciplinas:
                    turmas_disciplinas.append(td)
            c_td, c_ia = st.columns([2.6, 1], vertical_alignment="center")
            td_selecionadas = c_td.multiselect("Turmas / Disciplinas", turmas_disciplinas)
            botao_ia = c_ia.button(
                "✨ Gerar com IA", use_container_width=True, key="plano_ia_gerar",
                help="Preenche automaticamente todos os campos do plano de aula usando IA")
            if botao_ia:
                _gerar_plano_com_ia(tema, data_inicio, carga, td_selecionadas)

            aba_m, aba_b, aba_r = st.tabs(["Metodologia & Objetivos", "Estrutura BNCC", "Recursos & Avaliacao"])
            mapa = {"metodologia": 0, "objetivos": 0, "procedimentos": 0,
                    "habilidades": 1, "comp_geral": 1, "comp_especifica": 1,
                    "recursos": 2, "avaliacao": 2, "observacoes": 2}
            valores_extras = {}
            for rotulo, chave in CAMPOS_PLANO:
                with (aba_m if mapa[chave] == 0 else aba_b if mapa[chave] == 1 else aba_r):
                    valores_extras[chave] = st.text_area(rotulo, key=f"plano_{chave}", height=44)

            if st.button("Distribuir Plano Automaticamente", type="primary",
                         use_container_width=True):
                if not tema.strip():
                    st.error("O campo Tema e obrigatorio.")
                elif not valores_extras["metodologia"].strip():
                    st.error("O campo Metodologia e obrigatorio.")
                elif not td_selecionadas:
                    st.error("Marque pelo menos uma turma/disciplina.")
                else:
                    try:
                        data_obj = datetime.strptime(data_inicio.strip(), "%d/%m/%Y")
                    except ValueError:
                        st.error("Formato de data invalido. Use DD/MM/AAAA.")
                    else:
                        carga_n = int(carga.split()[0])
                        resumo = distribuir_planos(planos=carregar_planos(), grade=grade,
                                                   td_selecionadas=td_selecionadas,
                                                   tema=tema.strip(), data_inicio=data_obj,
                                                   carga_horaria=carga_n,
                                                   dados_extras={k: v.strip() for k, v in valores_extras.items()})
                        if resumo:
                            st.success("Plano distribuido com sucesso!\n\n" + "\n".join(resumo))
                        else:
                            st.warning("Nao foi possivel encontrar horarios na grade para essas turmas.")

    with tab_listar:
        montar_aba_listar_planos(grade, config)

def distribuir_planos(planos, grade, td_selecionadas, tema, data_inicio,
                      carga_horaria, dados_extras):
    mapa_dias = {d: i for i, d in enumerate(DIAS_COMPLETOS)}
    resumo = []
    plano_modificado = False

    for td in td_selecionadas:
        partes = td.split(" - ")
        turma_nome = partes[0]
        disc_nome = partes[1] if len(partes) > 1 else "Geral"

        aulas_restantes = carga_horaria
        aulas_turma = [i for i in grade
                       if i["turma"] == turma_nome and i.get("disciplina", "Geral") == disc_nome]
        if not aulas_turma:
            continue

        dia_atual = data_inicio
        limite_busca = 30
        while aulas_restantes > 0 and limite_busca > 0:
            dia_semana_teste = dia_atual.weekday()
            aulas_neste_dia = [a for a in aulas_turma if mapa_dias.get(a["dia"]) == dia_semana_teste]
            aulas_neste_dia.sort(key=lambda x: x.get("aula", ""))

            for aula_obj in aulas_neste_dia:
                if aulas_restantes <= 0:
                    break
                str_data_final = dia_atual.strftime("%d/%m/%Y")
                if str_data_final not in planos:
                    planos[str_data_final] = []
                novo_plano = {
                    "turma": turma_nome, "disciplina": disc_nome,
                    "horario": aula_obj["aula"], "tema": tema
                }
                novo_plano.update(dados_extras)
                planos[str_data_final].append(novo_plano)
                resumo.append(f"{td}: {str_data_final} ({aula_obj['aula']})")
                aulas_restantes -= 1
                plano_modificado = True

            dia_atual += timedelta(days=1)
            limite_busca -= 1

    if plano_modificado:
        salvar_planos(planos)
    return resumo

def gerar_excel_planos(planos, mes, ano, turmas_filtro=None):
    """Gera um arquivo Excel com uma aba por turma contendo todos os planos
    de aula registrados no mes/ano. Retorna BytesIO ou None se nao houver dados."""
    dados_por_turma = {}
    for data_str, lista in planos.items():
        try:
            data_obj = datetime.strptime(data_str, "%d/%m/%Y")
        except Exception:
            continue
        if data_obj.month != mes or data_obj.year != ano:
            continue
        for plano in lista:
            turma = str(plano.get("turma", "Sem Turma"))
            if turmas_filtro and turma not in turmas_filtro:
                continue
            dados_por_turma.setdefault(turma, []).append((data_obj, data_str, plano))

    if not dados_por_turma:
        return None

    cabecalho = (["Data", "Dia da Semana", "Horario", "Disciplina", "Tema"]
                 + [rotulo for rotulo, _ in CAMPOS_PLANO])
    campos_chaves = [chave for _, chave in CAMPOS_PLANO]

    fill_cab = PatternFill(start_color="1F538D", end_color="1F538D", fill_type="solid")
    font_cab = Font(bold=True, color="FFFFFF")
    border_thin = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))
    align_cab = Alignment(horizontal="center", vertical="center")
    align_cel = Alignment(vertical="top", wrap_text=True)

    def _numero_aula(horario):
        m = re.search(r"(\d+)", str(horario))
        return int(m.group(1)) if m else 0

    nomes_abas = {}
    for turma in sorted(dados_por_turma):
        base = re.sub(r"[\[\]:*?/\\]", "", turma).strip() or "Turma"
        base = base[:31]
        nome = base
        n = 2
        while nome in nomes_abas.values():
            nome = f"{base[:28]}_{n}"
            n += 1
        nomes_abas[turma] = nome

    larguras = [11, 15, 10, 14, 30] + [40] * len(campos_chaves)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    for turma, nome_aba in nomes_abas.items():
        ws = wb.create_sheet(nome_aba)
        for col, titulo in enumerate(cabecalho, 1):
            c = ws.cell(row=1, column=col, value=titulo)
            c.font = font_cab
            c.fill = fill_cab
            c.alignment = align_cab
            c.border = border_thin

        linhas = sorted(dados_por_turma[turma],
                        key=lambda x: (x[0], _numero_aula(x[2].get("horario", "")), x[1]))
        for i, (data_obj, data_str, plano) in enumerate(linhas, 2):
            valores = [data_str,
                       DIAS_COMPLETOS[data_obj.weekday()],
                       str(plano.get("horario", "")),
                       str(plano.get("disciplina", "Geral")),
                       str(plano.get("tema", ""))]
            valores += [str(plano.get(chave, "")) for chave in campos_chaves]
            linhas_est = 1
            for col, val in enumerate(valores, 1):
                c = ws.cell(row=i, column=col, value=val)
                c.alignment = align_cel
                c.border = border_thin
                if val:
                    qtd_linhas = (len(val) + (larguras[col - 1] // 2) - 1) // (larguras[col - 1] // 2)
                    linhas_est = max(linhas_est, qtd_linhas)
            ws.row_dimensions[i].height = max(20, linhas_est * 14 + 4)

        ws.freeze_panes = "A2"
        if ws.max_row > 1:
            ws.auto_filter.ref = f"A1:{get_column_letter(len(cabecalho))}{ws.max_row}"
        for col, largura in zip(range(1, len(cabecalho) + 1), larguras):
            ws.column_dimensions[get_column_letter(col)].width = largura

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def montar_aba_listar_planos(grade, config):
    hoje = datetime.now()
    pm = val_param("pm", hoje.strftime("%m/%Y"))
    try:
        mes, ano = [int(x) for x in pm.split("/")]
    except Exception:
        mes, ano = hoje.month, hoje.year
    if mes < 1 or mes > 12:
        mes, ano = hoje.month, hoje.year

    def href_pm(novo_mes, novo_ano):
        return f"?pagina=Central de Planos&pm={novo_mes}/{novo_ano}"

    def mes_offset(delta):
        m, a = mes, ano
        m += delta
        if m > 12:
            m, a = 1, a + 1
        elif m < 1:
            m, a = 12, a - 1
        return m, a

    c1, c2, c3 = st.columns([1, 2.2, 1])
    m_prev, m_next = mes_offset(-1), mes_offset(1)
    c1.markdown(f'<a href="{href_pm(m_prev[0], m_prev[1])}" style="text-decoration:none;'
                f'background:#6c757d;color:#fff;padding:8px 16px;border-radius:8px;'
                f'font-weight:600;">&#9664; Anterior</a>', unsafe_allow_html=True)
    c2.markdown(f"### {MESES_PT[mes]} de {ano}", unsafe_allow_html=True)
    c3.markdown(f'<div style="text-align:right;"><a href="{href_pm(m_next[0], m_next[1])}" '
                f'style="text-decoration:none;background:#6c757d;color:#fff;padding:8px 16px;'
                f'border-radius:8px;font-weight:600;">Proximo &#9654;</a></div>', unsafe_allow_html=True)

    planos = carregar_planos()
    turmas_unicas = ["Todas as Turmas"] + sorted(list(set([i["turma"] for i in grade])))
    c_filtro, c_export = st.columns([3, 1], vertical_alignment="center")
    turma_filtro = c_filtro.selectbox("Filtrar por turma", turmas_unicas)
    turmas_export = None if turma_filtro == "Todas as Turmas" else [turma_filtro]
    excel_planos = gerar_excel_planos(planos, mes, ano, turmas_export)
    if excel_planos is None:
        c_export.button("Exportar para Excel", disabled=True,
                        use_container_width=True,
                        help="Nao ha planos registrados neste mes para exportar.")
    else:
        c_export.download_button(
            "Exportar para Excel", data=excel_planos,
            file_name=f"Planos_{MESES_PT[mes]}_{ano}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            type="primary", use_container_width=True)
    mostrar_audit = st.toggle("Mostrar Auditoria (Faltas de Planejamento)", value=False)

    dark = config.get("aparencia", "System") == "Dark"
    legenda = ("Modo Auditoria: verde = 100% planejado | amarelo = parcial | "
               "vermelho = falta plano | cinza = sem aula" if mostrar_audit else
               "Modo Normal: azul = aulas planejadas neste dia | cinza = sem registros")
    st.caption(legenda)

    with st.container(key="card_planos_cal"):
        st.markdown(html_calendario_planos(planos, grade, ano, mes, turma_filtro, mostrar_audit, dark),
                    unsafe_allow_html=True)

    if not mostrar_audit:
        datas_mes = []
        for data_str in planos.keys():
            try:
                d = datetime.strptime(data_str, "%d/%m/%Y")
                if d.month == mes and d.year == ano:
                    datas_mes.append((d, data_str))
            except Exception:
                continue
        datas_mes.sort()
        encontrou = False
        for data_obj, data_str in datas_mes:
            for i, plano in enumerate(planos[data_str]):
                if turma_filtro != "Todas as Turmas" and plano.get("turma") != turma_filtro:
                    continue
                encontrou = True
                st.markdown(
                    f'<div class="dash-item">'
                    f'<div class="dash-item-titulo">{data_str} - {DIAS_COMPLETOS[data_obj.weekday()]} | '
                    f'{plano.get("horario","")} | {plano.get("turma","")} ({plano.get("disciplina","Geral")})</div>'
                    f'<div class="dash-item-sub">Tema: {plano.get("tema","")}</div></div>',
                    unsafe_allow_html=True)
                with st.popover("Visualizar / Editar", key=f"plist_pop_{data_str}_{i}", use_container_width=True):
                    form_plano(data_str, i, uid=f"plist_{data_str}_{i}")
        if not encontrou:
            st.caption("Nenhum plano registrado para os filtros neste mes.")
    else:
        pendencias = 0
        ultimo_dia = calendar.monthrange(ano, mes)[1]
        for dia in range(1, ultimo_dia + 1):
            data_str = f"{dia:02d}/{mes:02d}/{ano}"
            data_obj = datetime(ano, mes, dia)
            nome_dia = DIAS_COMPLETOS[data_obj.weekday()]
            aulas_esperadas = [g for g in grade if g["dia"] == nome_dia]
            if turma_filtro != "Todas as Turmas":
                aulas_esperadas = [g for g in aulas_esperadas if g["turma"] == turma_filtro]
            if not aulas_esperadas:
                continue
            planos_dia = planos.get(data_str, [])
            faltas = []
            for esperada in aulas_esperadas:
                registrado = any(p.get("turma") == esperada["turma"] and p.get("horario") == esperada["aula"]
                                 for p in planos_dia)
                if not registrado:
                    faltas.append(esperada)
            if faltas:
                pendencias += 1
                st.markdown(f"**PENDENCIA: {data_str} - {nome_dia}**")
                for falta in faltas:
                    st.markdown(
                        f'<div class="pend-item">'
                        f'Falta plano para a {falta["aula"]} | Turma: {falta["turma"]} '
                        f'({falta.get("disciplina","Geral")})</div>', unsafe_allow_html=True)
        if pendencias == 0:
            st.success("Parabens! Tudo 100% planejado para os filtros selecionados.")

# =====================================================================
# 13. TELA: ANOTACOES
# =====================================================================
def tela_anotacoes():
    st.markdown("## Lembretes")
    anotacoes = carregar_anotacoes()
    grade = carregar_grade()

    c1, c2 = st.columns([3, 1])
    turmas = ["Todas as Notas"] + sorted(list(set([i["turma"] for i in grade])))
    filtro = c1.selectbox("Filtrar mural", turmas)
    with c2.popover("➕ Novo Post-it", type="primary", key="anot_novo_pop", use_container_width=True):
        form_postit(None, uid="anot_novo")

    notas = anotacoes
    if filtro != "Todas as Notas":
        notas = [n for n in anotacoes if n.get("turma") == filtro or n.get("turma") == "Geral"]

    if not notas:
        st.caption("Nenhuma anotacao por aqui. Clique em Novo Post-it acima!")
        return

    colunas = st.columns(3)
    for i, nota in enumerate(notas):
        with colunas[i % 3]:
            cor = nota.get("cor", "#fff3a3")
            txt = cor_texto_legivel(cor)
            data_pt = nota.get("data") or nota.get("data_criacao") or ""
            st.markdown(
                f'<div class="postit" style="background:{cor};color:{txt};min-height:120px;">'
                f'<div class="pt-titulo">{esc(nota.get("titulo","Sem titulo"))}</div>'
                f'<div class="pt-tag">#{esc(nota.get("turma","Geral"))} | {esc(data_pt)}</div>'
                f'<div class="pt-conteudo">{esc(nota.get("conteudo",""))}</div></div>',
                unsafe_allow_html=True)
            b1, b2 = st.columns(2)
            with b1.popover("Editar", key=f"anot_edit_pop_{nota.get('id')}", use_container_width=True):
                form_postit(nota, uid=f"anot_edit_{nota.get('id')}")
            if b2.button("Excluir", key=f"anot_del_{nota.get('id')}", use_container_width=True):
                anotacoes = [n for n in anotacoes if n.get("id") != nota.get("id")]
                salvar_anotacoes(anotacoes)
                st.rerun()

# =====================================================================
# 14. TELA: AVALIACOES E ESTATISTICAS
# =====================================================================
def nivel_info(perc):
    if perc < 50:
        return "Muito Critico", "#dc3545", "#ffffff"
    elif perc < 70:
        return "Critico", "#ffc107", "#000000"
    elif perc < 90:
        return "Intermediario", "#66bb6a", "#000000"
    else:
        return "Adequado", "#1b5e20", "#ffffff"

def resolver_descritor(texto):
    m = re.search(r"\d+", texto.strip())
    if m:
        chave = f"D{int(m.group())}"
        if chave in DESCRITORES_MAT:
            return chave
    return "Sem Descritor"

def tela_notas():
    st.markdown("## Central de Avaliações e Estatísticas")
    avaliacoes = carregar_avaliacoes()
    dados_turmas = carregar_turmas()

    tab_criar, tab_corrigir, tab_stats = st.tabs(
        ["1. Criar Avaliação e Gabarito", "2. Lançar Notas / Corrigir", "3. Estatísticas por Descritor"])

    # ---------------- TAB 1: CRIAR ----------------
    with tab_criar:
        turmas_cad = sorted(list(dados_turmas.keys()))
        if not turmas_cad:
            st.error("Você não tem turmas com alunos. Cadastre em Turmas e Alunos primeiro.")
        else:
            c1, c2, c3 = st.columns(3)
            turmas_sel = c1.multiselect("Turmas que farão a prova", turmas_cad)
            titulo = c2.text_input("Titulo da prova", placeholder="Ex: Prova Bimestral de Matemática")
            qtd = c3.number_input("Qtd. de questões", min_value=1, max_value=50, value=10, step=1)

            if st.button("Gerar Atividade", type="primary"):
                st.session_state["gab_qtd"] = int(qtd)
                st.session_state["gab_rows"] = [{"resposta": "A", "descritor": ""}
                                                for _ in range(int(qtd))]

            st.markdown("**Gabarito e descritores (digite o número, ex: 2, 02, D02):**")
            linhas = st.session_state.get("gab_rows", [])
            for i, row in enumerate(linhas):
                r1, r2, r3 = st.columns([1, 1, 3])
                chave_r = f"gab_{i}"
                resp = r1.selectbox(f"Q{i + 1:02d} - Resposta", ["A", "B", "C", "D", "E"],
                                    index=["A", "B", "C", "D", "E"].index(row.get("resposta", "A")),
                                    key=f"{chave_r}_resp")
                desc = r2.text_input(f"Descritor Q{i + 1:02d}", value=row.get("descritor", ""),
                                     placeholder="Ex: D02", key=f"{chave_r}_desc")
                hint = ""
                if desc.strip():
                    chave_d = resolver_descritor(desc)
                    if chave_d != "Sem Descritor":
                        hint = f":blue[{DESCRITORES_MAT[chave_d]}]"
                    else:
                        hint = ":red[Descritor não encontrado]"
                r3.caption(hint if hint else " ")
                row["resposta"] = resp
                row["descritor"] = desc

            if st.button("Salvar Gabarito para Turmas Selecionadas", type="primary",
                         use_container_width=True):
                if not turmas_sel:
                    st.error("Selecione pelo menos uma turma.")
                elif not linhas:
                    st.error("Clique em 'Gerar Atividade' para gerar os campos do gabarito.")
                else:
                    gabarito_final = [
                        {"questao": i + 1,
                         "resposta_correta": row["resposta"],
                         "descritor": resolver_descritor(row["descritor"])}
                        for i, row in enumerate(linhas)]
                    criadas = 0
                    for turma in turmas_sel:
                        avaliacoes.append({
                            "id": max([a.get("id", 0) for a in avaliacoes], default=0) + 1,
                            "titulo": titulo.strip() or "Avaliação sem Título",
                            "turma": turma,
                            "data": datetime.now().strftime("%d/%m/%Y"),
                            "gabarito": gabarito_final,
                            "notas_alunos": {}
                        })
                        criadas += 1
                    salvar_avaliacoes(avaliacoes)
                    st.success(f"{criadas} avaliacao(oes) salva(s) com sucesso!")
                    st.session_state.pop("gab_rows", None)

    # ---------------- TAB 2: CORRIGIR ----------------
    with tab_corrigir:
        if not avaliacoes:
            st.info("Nenhuma avaliação cadastrada.")
        else:
            titulos_av = [f"[{a['turma']}] {a['titulo']} ({a['data']})" for a in avaliacoes]
            c1, c2 = st.columns([2.5, 1.5])
            av_escolhida = c1.selectbox("Selecione a avaliação", titulos_av, key="av_corrigir")
            modo = c2.radio("Modo de correcao", ["Rapido (Texto)", "Detalhado (Blocos)"],
                            horizontal=True)

            idx = titulos_av.index(av_escolhida)
            av = avaliacoes[idx]
            turma = av["turma"]
            alunos = sorted(dados_turmas.get(turma, []))
            gabarito = av["gabarito"]
            qtd_q = len(gabarito)

            if not alunos:
                st.warning("Esta turma não tem alunos cadastrados.")
            else:
                if modo == "Rapido (Texto)":
                    st.caption("Dica: digite as respostas seguidas (Ex: ADCBE) na caixa.")
                else:
                    st.caption("Modo detalhado: preencha a letra marcada em cada bloco.")

                with st.form("form_correcao"):
                    dados_entrada = []
                    for aluno in alunos:
                        st.markdown(f"**{aluno}**")
                        dados_exist = av["notas_alunos"].get(aluno, {})
                        if modo == "Rapido (Texto)":
                            c1, c2 = st.columns([3, 1])
                            resp_txt = c1.text_input(
                                f"Respostas ({qtd_q} questões)",
                                value=dados_exist.get("respostas_dadas", ""),
                                key=f"corr_resp_{aluno}", label_visibility="collapsed")
                            nota = c2.text_input("Nota", value=str(dados_exist.get("nota_final", "")),
                                                 key=f"corr_nota_{aluno}",
                                                 placeholder="auto")
                            dados_entrada.append((aluno, "rapido", resp_txt, nota, None))
                        else:
                            notas_exist = dados_exist.get("respostas_dadas", "")
                            blocos = []
                            cols = st.columns(min(qtd_q, 5))
                            for qi in range(qtd_q):
                                val = notas_exist[qi] if qi < len(notas_exist) else ""
                                with cols[qi % min(qtd_q, 5)]:
                                    inp = st.text_input(f"Q{qi + 1}", value=val,
                                                        key=f"corr_{aluno}_q{qi}",
                                                        max_chars=1)
                                    blocos.append(inp)
                            nota = st.text_input("Nota final", key=f"corr_nota_d_{aluno}",
                                                 value=str(dados_exist.get("nota_final", "")),
                                                 placeholder="auto")
                            dados_entrada.append((aluno, "detalhado", None, nota, blocos))

                    enviado = st.form_submit_button("Processar Correcao Automatica e Salvar",
                                                    type="primary", use_container_width=True)
                    if enviado:
                        processar_correcao(av, dados_entrada, gabarito, avaliacoes, idx)

    # ---------------- TAB 3: ESTATISTICAS ----------------
    with tab_stats:
        montar_aba_stats(avaliacoes)

# ---------------- correcao ----------------
def processar_correcao(av, dados_entrada, gabarito, avaliacoes, idx):
    qtd_q = len(gabarito)
    for aluno, tipo, resp_txt, nota_txt, blocos in dados_entrada:
        if tipo == "rapido":
            resp_dadas = resp_txt.strip().upper()
        else:
            resp_dadas = "".join([(b.strip().upper() or " ") for b in blocos])

        nota_manual = nota_txt.strip()
        acertos_descritores = {}
        nota_calculada = 0

        if resp_dadas.strip():
            acertos = 0
            for i, q in enumerate(gabarito):
                descritor = q["descritor"]
                if descritor not in acertos_descritores:
                    acertos_descritores[descritor] = {"corretas": 0, "total": 0}
                acertos_descritores[descritor]["total"] += 1
                letra_aluno = resp_dadas[i] if i < len(resp_dadas) else " "
                if letra_aluno == q["resposta_correta"].upper():
                    acertos += 1
                    acertos_descritores[descritor]["corretas"] += 1
            nota_calculada = round((acertos / qtd_q) * 10, 1)

        if resp_dadas.strip() or nota_manual:
            av["notas_alunos"][aluno] = {
                "respostas_dadas": resp_dadas,
                "nota_final": float(nota_manual) if nota_manual else nota_calculada,
                "descritores": acertos_descritores
            }
    avaliacoes[idx] = av
    salvar_avaliacoes(avaliacoes)
    st.success("Correção automática finalizada e notas salvas!")
    st.rerun()

# =====================================================================
# 15. ESTATISTICAS POR DESCRITOR
# =====================================================================
def montar_aba_stats(avaliacoes):
    if not avaliacoes:
        st.info("Nenhuma avaliação cadastrada.")
        return

    titulos_av = [f"[{a['turma']}] {a['titulo']} ({a['data']})" for a in avaliacoes]
    av_escolhida = st.selectbox("Avaliação para estatísticas", titulos_av, key="av_stats")

    idx = titulos_av.index(av_escolhida)
    av = avaliacoes[idx]
    notas = av["notas_alunos"]
    gabarito = av["gabarito"]
    qtd_questoes = len(gabarito)

    gerar = st.button("Gerar Relatório", type="primary")
    if not gerar:
        return

    if not notas:
        st.warning("Nenhuma nota lançada para gerar estatísticas.")
        return

    media_turma = sum([n["nota_final"] for n in notas.values()]) / len(notas)
    distribuicao = {"Muito Critico": 0, "Critico": 0, "Intermediario": 0, "Adequado": 0}
    for dados in notas.values():
        perc_aluno = (dados["nota_final"] / 10) * 100
        nivel, _, _ = nivel_info(perc_aluno)
        distribuicao[nivel] += 1

    st.markdown(f"### Resumo da turma: {av['turma']}")
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Media da Turma", f"{media_turma:.1f}")
    c2.metric("Adequado (90-100%)", distribuicao["Adequado"])
    c3.metric("Intermediario (70-89%)", distribuicao["Intermediario"])
    c4.metric("Critico (50-69%)", distribuicao["Critico"])
    c5.metric("Muito Critico (<50%)", distribuicao["Muito Critico"])

    st.markdown("### Desempenho por Descritor / Habilidade")
    consolidado = {}
    for aluno, dados in notas.items():
        for desc, val in dados.get("descritores", {}).items():
            if desc not in consolidado:
                consolidado[desc] = {"corretas": 0, "total": 0}
            consolidado[desc]["corretas"] += val["corretas"]
            consolidado[desc]["total"] += val["total"]

    if consolidado:
        for desc, val in sorted(consolidado.items()):
            if val["total"] == 0:
                continue
            perc = (val["corretas"] / val["total"]) * 100
            nome_nivel, cor_barra, _ = nivel_info(perc)
            nome_desc = DESCRITORES_MAT.get(desc, "")
            if nome_desc:
                st.caption(f"{desc} - {nome_desc}")
            st.progress(perc / 100, text=f"{desc}: {val['corretas']}/{val['total']} acertos "
                                         f"({perc:.1f}% - {nome_nivel})")
    else:
        st.caption("Sem dados de descritores.")

    st.markdown("### Detalhamento por Aluno")
    linhas_tab = []
    for aluno, dados in sorted(notas.items()):
        perc = (dados["nota_final"] / 10) * 100
        nivel, _, _ = nivel_info(perc)
        acertos = round((dados["nota_final"] / 10) * qtd_questoes)
        linhas_tab.append({"Aluno": aluno, "Acertos": f"{acertos}/{qtd_questoes}",
                           "Nota (0-10)": f"{dados['nota_final']:.1f}",
                           "Nivel": f"{perc:.0f}% - {nivel}"})
    st.dataframe(linhas_tab, use_container_width=True, hide_index=True)

    excel = gerar_excel_avaliacao(av)
    st.download_button(
        "Exportar Relatorio para Excel",
        data=excel,
        file_name=f"Relatorio_{av['turma']}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        type="primary")
import urllib.request
import urllib.error

# Modelos internos (escolhidos automaticamente; troque aqui se um cair em desuso)
MODELO_GEMINI = "gemini-3.5-flash"
MODELO_OPENAI = "gpt-4o-mini"

# =====================================================================
# 16. TELA: CADASTRAR NOVA QUESTAO
# =====================================================================
def tela_cadastrar():
    st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600;700&display=swap');

/* ---- Criar Questao: Design System (Poppins + cores da marca) ---- */
[data-testid="stMainBlockContainer"] {
    font-family: 'Poppins', system-ui, -apple-system, 'Segoe UI', Roboto, sans-serif !important;
}

/* Botao pequeno "Gerar com IA" (atalho p/ Criar com IA), similar ao Central de Planos */
div[class*="st-key-cq_ia_at"] button[kind="secondary"] {
    font-family: 'Poppins', system-ui, sans-serif;
    background: var(--card-bg) !important;
    color: var(--cor-p) !important;
    border: 1.5px solid color-mix(in srgb, var(--cor-p) 45%, var(--card-bg)) !important;
    border-radius: 10px !important;
    height: 40px !important;
    font-weight: 700 !important;
    box-shadow: 0 1px 3px rgba(0,0,0,.08) !important;
    transition: all .18s ease !important;
}
div[class*="st-key-cq_ia_at"] button[kind="secondary"]:hover {
    background: color-mix(in srgb, var(--cor-p) 10%, var(--card-bg)) !important;
    color: var(--cor-p) !important;
    box-shadow: 0 4px 14px rgba(0,0,0,.14) !important;
}
.cq-sub {
    font-family: 'Poppins', system-ui, sans-serif;
    color: var(--cor-texto); font-size: .92rem; margin: .35rem 0;
}

/* Cards das secoes (borda suave, tom do tema) */
div[class*="st-key-cq_card_"] {
    background: var(--card-bg) !important;
    border: 2px solid var(--borda) !important;
    border: 2px solid color-mix(in srgb, var(--cor-p) 32%, var(--card-bg)) !important;
    border-radius: 16px !important;
    box-shadow: 0 1px 2px rgba(0,0,0,.05), 0 2px 10px rgba(0,0,0,.05) !important;
    padding: 1.6rem 1.5rem 1.8rem !important;
    height: 100%;
}
div[data-testid="stLayoutWrapper"]:has(> div[class*="st-key-cq_card_"]) {
    height: 100% !important;
}
/* Mais ar entre o card 2 (Pergunta/Alternativas) e o card 3 (Metadados) */
div[class*="st-key-cq_card_2"] { margin-bottom: 2.2rem !important; }
/* Margem minima entre os cards e os botoes Voltar / Salvar */
[data-testid="stHorizontalBlock"]:has(div[class*="st-key-cq_card_1"]) {
    margin-bottom: 1.75rem !important;
}
/* Preenche a altura util da pagina (apenas em telas altas) */
@media (min-height: 800px) {
    [data-testid="stHorizontalBlock"]:has(div[class*="st-key-cq_card_1"]) {
        min-height: calc(100vh - 340px) !important;
    }
}

/* Titulo das secoes (barra vertical + texto) */
.cq-sec { display: flex; align-items: center; gap: .55rem; margin: .2rem 0 1.4rem 0; }
.cq-sec .cq-sec-bar { width: 6px; height: 24px; border-radius: 3px; background: var(--cor-p); flex: none; }
.cq-sec .cq-sec-txt {
    font-family: 'Poppins', system-ui, sans-serif;
    font-weight: 700; font-size: 1.08rem; color: var(--cor-texto);
}

/* Painel expansivel: Mapeamento de Distratores */
div[class*="st-key-cq_card_2"] [data-testid="stExpander"] {
    border: 1px solid color-mix(in srgb, var(--cor-p) 30%, var(--card-bg)) !important;
    border-radius: 10px !important;
    background: color-mix(in srgb, var(--cor-p) 5%, var(--card-bg)) !important;
}
div[class*="st-key-cq_card_2"] [data-testid="stExpander"] summary { color: var(--cor-p) !important; }

/* Nivel de dificuldade como pildulas (Facil / Medio / Dificil) */
div[class*="st-key-cq_card_3"] div[role="radiogroup"] { gap: .55rem !important; }
div[class*="st-key-cq_card_3"] label[data-baseweb="radio"] {
    border: 1.5px solid var(--borda) !important;
    border-radius: 10px !important;
    padding: .5rem 1rem !important;
    background: var(--fundo) !important;
    font-weight: 600 !important;
    cursor: pointer;
    transition: all .15s ease;
}
div[class*="st-key-cq_card_3"] label[data-baseweb="radio"] > div:first-child { display: none !important; }
div[class*="st-key-cq_card_3"] label[data-baseweb="radio"]:has(input:checked) {
    background: var(--cor-p) !important;
    color: var(--btn-fg) !important;
    border-color: var(--cor-p) !important;
    box-shadow: 0 3px 10px rgba(0,0,0,.15);
}

/* Botoes do rodape: Voltar (tom suave da marca) e Salvar Questao (gradiente da marca) */
div[class*="st-key-cq_voltar"] button[kind="secondaryFormSubmit"] {
    background: color-mix(in srgb, var(--cor-p) 10%, var(--card-bg)) !important;
    color: var(--cor-p) !important;
    border: 1.5px solid color-mix(in srgb, var(--cor-p) 45%, var(--card-bg)) !important;
    font-weight: 700 !important;
}
div[class*="st-key-cq_voltar"] button[kind="secondaryFormSubmit"]:hover {
    background: color-mix(in srgb, var(--cor-p) 18%, var(--card-bg)) !important;
    box-shadow: 0 4px 14px rgba(0,0,0,.12) !important;
}
div[class*="st-key-cq_salvar"] button[kind="primaryFormSubmit"] {
    background: linear-gradient(90deg, var(--cor-p) 0%, var(--cor-pd) 100%) !important;
    border: 1px solid var(--cor-pd) !important;
    box-shadow: 0 4px 14px rgba(0,0,0,.16) !important;
    font-weight: 700 !important;
}
div[class*="st-key-cq_salvar"] button[kind="primaryFormSubmit"]:hover {
    background: linear-gradient(90deg, var(--cor-ph) 0%, var(--cor-pd) 100%) !important;
    box-shadow: 0 6px 18px rgba(0,0,0,.22) !important;
}
</style>""", unsafe_allow_html=True)

    modo_ia = st.session_state.get("cad_modo_ia", False)

    if modo_ia:
        if st.button("\u2190 Voltar ao formul\u00e1rio de cria\u00e7\u00e3o manual",
                     use_container_width=True, key="cad_modo_manual"):
            st.session_state["cad_modo_ia"] = False
            st.rerun()
        st.markdown("---")
        tela_importar(integrada=True)
        return

    banco = carregar_banco()

    info_cad = st.session_state.pop("cad_info", None)
    if st.session_state.pop("cad_limpar_img", False):
        st.session_state.pop("img_nova_questao", None)
    if info_cad:
        if info_cad[0] == "ok":
            st.success(info_cad[1])
        else:
            st.error(info_cad[1])

    c_topo = st.columns([3.4, 1], vertical_alignment="center")
    c_topo[0].markdown(
        '<div class="cq-sub">Preencha os campos abaixo para cadastrar uma '
        'quest\u00e3o manualmente ou clique em <b>Gerar com IA</b> para criar '
        'automaticamente.</div>', unsafe_allow_html=True)
    if c_topo[1].button("\u2728 Gerar com IA", use_container_width=True,
                        key="cq_ia_at",
                        help="Atalho para o gerador de questões com IA"):
        st.session_state["cad_modo_ia"] = True
        st.rerun()

    with st.form("form_nova_questao"):
        c_esq, c_dir = st.columns([5, 7], gap="large")

        with c_esq:
            with st.container(border=True, key="cq_card_1"):
                st.markdown(
                    '<div class="cq-sec"><span class="cq-sec-bar"></span>'
                    '<span class="cq-sec-txt">1. Contexto Inicial e Imagem</span></div>',
                    unsafe_allow_html=True)
                enunciado = st.text_area(
                    "Enunciado / Texto de Contexto Inicial:", height=280,
                    placeholder="Cole aqui o texto de contexto da questao...")
                imagem = st.file_uploader(
                    "Imagem de Apoio (Opcional):",
                    type=["png", "jpg", "jpeg", "bmp", "webp"], key="img_nova_questao")

        with c_dir:
            with st.container(border=True, key="cq_card_2"):
                st.markdown(
                    '<div class="cq-sec"><span class="cq-sec-bar"></span>'
                    '<span class="cq-sec-txt">2. Pergunta e Alternativas</span></div>',
                    unsafe_allow_html=True)
                pergunta = st.text_input(
                    "Pergunta Direta / Comando:",
                    placeholder="Ex: Qual e o valor de x?")
                st.markdown("**Alternativas (deixe em branco se for discursiva):**")
                c_alt = st.columns(2)
                alt_a = c_alt[0].text_input("A)", placeholder="Texto da alternativa A")
                alt_b = c_alt[1].text_input("B)", placeholder="Texto da alternativa B")
                alt_c = c_alt[0].text_input("C)", placeholder="Texto da alternativa C")
                alt_d = c_alt[1].text_input("D)", placeholder="Texto da alternativa D")

                with st.expander("\u2699\ufe0f Opções Avançadas: Mapeamento de Distratores"):
                    st.caption("Indique a lacuna de aprendizagem (motivo do erro) para cada "
                               "alternativa incorreta. Isso alimentará as estatísticas da turma.")
                    c_d = st.columns(2)
                    dist_b = c_d[0].text_input("B)",
                                               placeholder="Ex: Erro de regra de sinais...")
                    dist_c = c_d[1].text_input("C)",
                                               placeholder="Ex: Falha na interpretação de texto...")
                    dist_d = c_d[0].text_input("D)",
                                               placeholder="Ex: Confusão com fórmula base...")

            with st.container(border=True, key="cq_card_3"):
                st.markdown(
                    '<div class="cq-sec"><span class="cq-sec-bar"></span>'
                    '<span class="cq-sec-txt">3. Metadados e Gabarito</span></div>',
                    unsafe_allow_html=True)
                c3 = st.columns(3)
                gabarito = c3[0].selectbox(
                    "Gabarito Correto:", ["A", "B", "C", "D", "Discursiva"])
                dificuldade = c3[1].radio(
                    "Nível de Dificuldade:", ["Fácil", "Médio", "Difícil"], horizontal=True)
                tema = c3[2].text_input(
                    "Tema / Descritor SAEB:", placeholder="Ex: D15 - Resolver problema...")

        c_f = st.columns([1, 2])
        voltar = c_f[0].form_submit_button("Voltar", use_container_width=True,
                                           key="cq_voltar")
        salvar = c_f[1].form_submit_button("Salvar Questão", type="primary",
                                           use_container_width=True, key="cq_salvar")

    if voltar:
        st.session_state["cad_limpar_img"] = True
        st.rerun()

    if salvar:
        mapa_dif = {"Fácil": "Facil", "Médio": "Medio", "Difícil": "Dificil"}
        dificuldade_interna = mapa_dif.get(dificuldade, "Medio")
        if not enunciado.strip() or not pergunta.strip() or not tema.strip():
            st.error("Por favor, preencha o Enunciado, a Pergunta e o Tema!")
        else:
            alternativas = {}
            for letra, texto in [("A", alt_a), ("B", alt_b), ("C", alt_c), ("D", alt_d)]:
                if texto.strip():
                    alternativas[letra] = texto.strip()

            distratores = {}
            for letra, texto in [("B", dist_b), ("C", dist_c), ("D", dist_d)]:
                if texto.strip() and letra != gabarito:
                    distratores[letra] = texto.strip()

            novo_id = max([q.get("id", 0) for q in banco], default=0) + 1
            caminho_final_img = ""
            if imagem is not None:
                nome_arq = imagem.name.replace(" ", "_")
                caminho_final_img = salvar_imagem_usuario(nome_arq, imagem.getbuffer())

            banco.append({
                "id": novo_id,
                "tema": tema.strip(),
                "dificuldade": dificuldade_interna,
                "enunciado": enunciado.strip(),
                "imagem": caminho_final_img,
                "pergunta_direta": pergunta.strip(),
                "alternativas": alternativas,
                "gabarito": gabarito,
                "distratores": distratores
            })
            salvar_banco(banco)
            st.session_state["cad_limpar_img"] = True
            st.session_state["cad_info"] = ("ok",
                                            f"Questão {novo_id} salva com sucesso no banco de dados!")
            st.rerun()

# =====================================================================
# 17. TELA: IMPORTAR QUESTOES COM IA (via API do usuario)
# =====================================================================
def _sugerir_descritores(texto, limite=4):
    t = (texto or "").strip()
    if not t:
        return []
    tl = t.lower()
    exatas = []
    parciais = []
    for desc in DESCRITORES_SAEB:
        dl = desc.lower()
        if tl in dl:
            codigo = desc.split(" - ")[0].lower()
            (exatas if codigo == tl else parciais).append(desc)
    return (exatas + parciais)[:limite]


def _montar_prompt_ia(disciplina, tema, qtd, dificuldade, tipo, n_alt, extra):
    if tipo == "Multipla escolha":
        letras = "ABCDEFGH"[:n_alt]
        estrutura = ('[ { "tema": "descritor/tema", "dificuldade": "Facil|Medio|Dificil", '
                     '"enunciado": "texto de contexto ou situacao", '
                     '"pergunta_direta": "comando objetivo da questao", '
                     '"alternativas": { ' +
                     ", ".join(f'"{l}": "opcao {l}"' for l in letras) +
                     ' }, "gabarito": "' + letras[0] + '" } ]')
    else:
        estrutura = ('[ { "tema": "descritor/tema", "dificuldade": "Facil|Medio|Dificil", '
                     '"enunciado": "texto de contexto ou situacao", '
                     '"pergunta_direta": "comando objetivo da questao", '
                     '"alternativas": {}, "gabarito": "Discursiva" } ]')

    linhas = [
        f"Crie exatamente {qtd} questoes de prova com as caracteristicas abaixo.",
        f"- Disciplina: {disciplina.strip() or 'Nao informada'}",
        f"- Tema / assunto: {tema.strip()}",
        f"- Nivel de dificuldade: {dificuldade}",
        f"- Tipo: {tipo}",
    ]
    if tipo == "Multipla escolha":
        linhas.append(f"- Cada questao deve ter exatamente {n_alt} alternativas "
                      f"(letras {', '.join(letras)}).")
    if extra and extra.strip():
        linhas.append(f"- Instrucoes adicionais: {extra.strip()}")
    linhas.append("")
    linhas.append("Responda APENAS com JSON valido, sem texto adicional, "
                 "usando exatamente esta estrutura de lista:")
    linhas.append(estrutura)
    linhas.append("")
    linhas.append("Regras:")
    if tipo == "Multipla escolha":
        linhas.append(f"- Preencha as alternativas com as letras {', '.join(letras)} e o gabarito "
                      f"com a letra correta entre elas.")
    else:
        linhas.append('- Deixe "alternativas" como {} e "gabarito" como "Discursiva".')
    linhas.append("- O enunciado deve trazer o contexto (texto, tabela, situacao-problema); "
                 "a pergunta_direta deve ser o comando.")
    linhas.append(f"- Nao repita questoes e gere exatamente {qtd} questoes.")
    return "\n".join(linhas)


def _extrair_json_ia(texto):
    t = (texto or "").strip()
    if t.startswith("```"):
        linhas = t.splitlines()
        t = "\n".join(linhas[1:]) if len(linhas) > 1 else ""
        if t.rstrip().endswith("```"):
            t = t.rstrip()[:-3].rstrip()
    ini = t.find("[")
    fim = t.rfind("]")
    if ini != -1 and fim != -1 and fim > ini:
        t = t[ini:fim + 1]
    return json.loads(t)


def chamar_ia_gemini(api_key, modelo, prompt):
    url = (f"https://generativelanguage.googleapis.com/v1beta/models/{modelo}"
           f":generateContent?key={api_key}")
    corpo = {"contents": [{"parts": [{"text": prompt}]}]}
    req = urllib.request.Request(url, data=json.dumps(corpo).encode("utf-8"),
                                 headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=120) as resp:
        dados = json.loads(resp.read().decode("utf-8"))
    try:
        return dados["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError, TypeError):
        raise ValueError("O Gemini nao retornou um texto valido.")


def chamar_ia_openai(api_key, modelo, prompt):
    url = "https://api.openai.com/v1/chat/completions"
    corpo = {
        "model": modelo,
        "messages": [
            {"role": "system",
             "content": "Voce e um professor experiente que gera questoes de prova em JSON valido."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.7,
    }
    req = urllib.request.Request(url, data=json.dumps(corpo).encode("utf-8"),
                                 headers={"Content-Type": "application/json",
                                          "Authorization": f"Bearer {api_key}"})
    with urllib.request.urlopen(req, timeout=120) as resp:
        dados = json.loads(resp.read().decode("utf-8"))
    try:
        return dados["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError):
        raise ValueError("A OpenAI nao retornou um texto valido.")


def _importar_questao_revista(q, disc):
    banco = carregar_banco()
    novo_id = max([b.get("id", 0) for b in banco], default=0) + 1
    banco.append({
        "id": novo_id,
        "disciplina": (str(q.get("disciplina") or "").strip() or disc or "Geral"),
        "tema": str(q.get("tema", "")).strip(),
        "dificuldade": str(q.get("dificuldade", "Medio")),
        "enunciado": str(q.get("enunciado", "")),
        "imagem": str(q.get("imagem", "")),
        "pergunta_direta": str(q.get("pergunta_direta", "")),
        "alternativas": q.get("alternativas") or {},
        "gabarito": str(q.get("gabarito", "")),
    })
    salvar_banco(banco)
    st.session_state["ia_importadas"] = st.session_state.get("ia_importadas", 0) + 1


def _avancar_revisao():
    lista = st.session_state.get("ia_revisar") or []
    idx = st.session_state.get("ia_rev_idx", 0)
    if 0 <= idx < len(lista):
        lista.pop(idx)
    st.session_state["ia_editando"] = False
    if not lista:
        st.session_state.pop("ia_revisar", None)
        n = st.session_state.get("ia_importadas", 0)
        if n > 0:
            st.session_state["ia_info"] = ("ok",
                                           f"{n} questao(oes) importada(s) para o catalogo!")
        else:
            st.session_state["ia_info"] = ("info",
                                           "Revisao concluida. Nenhuma questao foi importada.")
    else:
        st.session_state["ia_rev_idx"] = min(idx, len(lista) - 1)


def form_revisar_questao(uid="rev"):
    lista = st.session_state.get("ia_revisar") or []
    if not lista:
        st.info("Nenhuma questao para revisar.")
        return
    idx = st.session_state.get("ia_rev_idx", 0)
    if idx >= len(lista):
        idx = len(lista) - 1
        st.session_state["ia_rev_idx"] = idx
    q = lista[idx]
    disc_padrao = st.session_state.get("ia_disciplina_gerada", "Geral")

    st.caption(f"Questão {idx + 1} de {len(lista)}")
    st.markdown(f"**{q.get('tema', 'Sem tema')}** \u00b7 {q.get('dificuldade', 'Medio')}"
                f" \u00b7 Gabarito: {q.get('gabarito', '-')}")
    if q.get("enunciado"):
        st.markdown("**Enunciado / contexto:**")
        st.write(q.get("enunciado"))
    if q.get("pergunta_direta"):
        st.markdown("**Pergunta / comando:**")
        st.write(q.get("pergunta_direta"))
    alt = q.get("alternativas") or {}
    if alt:
        st.markdown("**Alternativas:**")
        for letra, txt in alt.items():
            if txt:
                st.markdown(f"{letra}) {txt}")

    st.markdown("---")
    if not st.session_state.get("ia_editando"):
        c1, c2, c3 = st.columns(3)
        if c1.button("Importar no catalogo", type="primary", use_container_width=True,
                     key=f"{uid}_importar"):
            _importar_questao_revista(q, disc_padrao)
            _avancar_revisao()
            st.rerun()
        if c2.button("Editar", use_container_width=True, key=f"{uid}_editar"):
            st.session_state["ia_editando"] = True
            st.rerun()
        if c3.button("Excluir", use_container_width=True, key=f"{uid}_excluir"):
            _avancar_revisao()
            st.rerun()
    else:
        st.markdown("**Editando questao (lembre-se de salvar):**")
        e_tema = st.text_input("Tema / descritor:", value=q.get("tema", ""), key=f"{uid}_tema")
        e_disc = st.text_input("Disciplina:", value=q.get("disciplina") or disc_padrao,
                               key=f"{uid}_disciplina")
        e_dif = st.selectbox("Nivel:", ["Facil", "Medio", "Dificil"],
                             index=(["Facil", "Medio", "Dificil"].index(q.get("dificuldade", "Medio"))
                                    if q.get("dificuldade") in ["Facil", "Medio", "Dificil"] else 1),
                             key=f"{uid}_dificuldade")
        e_enun = st.text_area("Enunciado / contexto:", value=q.get("enunciado", ""), height=90,
                              key=f"{uid}_enunciado")
        e_perg = st.text_area("Pergunta / comando:", value=q.get("pergunta_direta", ""), height=60,
                              key=f"{uid}_pergunta")
        st.markdown("**Alternativas (em branco = remove; todas em branco = discursiva):**")
        letras_edit = "ABCDEFGH"[: max(len(alt), 4)]
        for letra in letras_edit:
            st.text_input(f"Alternativa {letra})", value=alt.get(letra, ""), key=f"{uid}_alt_{letra}")
        e_gab = st.text_input("Gabarito:", value=str(q.get("gabarito", "")), key=f"{uid}_gabarito")
        c1, c2 = st.columns(2)
        if c1.button("Salvar alteracoes", type="primary", use_container_width=True,
                     key=f"{uid}_salvar"):
            alt_nova = {}
            for letra in letras_edit:
                v = (st.session_state.get(f"{uid}_alt_{letra}") or "").strip()
                if v:
                    alt_nova[letra] = v
            lista[idx].update({
                "tema": e_tema.strip(),
                "disciplina": e_disc.strip(),
                "dificuldade": e_dif,
                "enunciado": e_enun.strip(),
                "pergunta_direta": e_perg.strip(),
                "gabarito": e_gab.strip(),
            })
            lista[idx]["alternativas"] = alt_nova
            st.session_state["ia_editando"] = False
            st.rerun()
        if c2.button("Cancelar edicao", use_container_width=True, key=f"{uid}_cancelar"):
            st.session_state["ia_editando"] = False
            st.rerun()


def tela_importar(integrada=False):
    # Descritor escolhido na lista de sugestoes: aplica antes de criar o widget
    pendente = st.session_state.pop("ia_tema_pendente", None)
    if pendente:
        st.session_state["ia_tema"] = pendente
    if integrada:
        st.caption("A IA gera as questões usando a **sua própria chave de API** (Gemini ou ChatGPT), "
                   "configurada em **Configurações > Chave de API**. "
                   "Depois de geradas, você revisa e importa para o catálogo.")
    else:
        st.markdown("### Importar Questões com IA")
        st.caption("A IA gera as questões usando a **sua própria chave de API** (Gemini ou ChatGPT), "
                   "configurada em **Configurações > Chave de API**. "
                   "O consumo é cobrado na sua conta do serviço.")

    info_ia = st.session_state.pop("ia_info", None)
    if info_ia:
        if info_ia[0] == "ok":
            st.success(info_ia[1])
        elif info_ia[0] == "info":
            st.info(info_ia[1])
        else:
            st.error(info_ia[1])

    chave_salva = (carregar_chave_ia() or "").strip()
    if not chave_salva:
        with st.container(border=True):
            c_img, c_txt = st.columns([1, 3], vertical_alignment="center")
            c_img.markdown(
                f'<img src="data:image/png;base64,{LOGO_INICIO_B64}" '
                'alt="Exame Inteligente" style="max-width:170px; width:100%; '
                'height:auto; border-radius:12px;" />',
                unsafe_allow_html=True)
            c_txt.markdown("**Configure sua chave de API para importar questões com IA**")
            c_txt.caption(
                "A chave de API (Google Gemini ou OpenAI) agora fica em "
                "**Configurações > Chave de API**. Informe a chave uma única vez "
                "e ela será salva na sua conta, sem precisar digitar de novo.")
            if c_txt.button("Ir para Configurações", type="primary",
                            key="ia_ir_config"):
                ir_para("Configurações")
                st.rerun()
        return

    provedor = carregar_config().get("provedor_ia", "Google Gemini")
    modelo_interno = MODELO_GEMINI if provedor == "Google Gemini" else MODELO_OPENAI
    st.caption(f"Provedor: **{provedor}** (definido em Configurações > Chave de API) | "
               f"Modelo de IA: **{modelo_interno}** (escolhido automaticamente pelo app).")

    st.markdown("---")
    c1, c2 = st.columns(2)
    opcoes_disc = DISCIPLINAS_COMUNS + ["Outra disciplina..."]
    disc_sel = c1.selectbox("Disciplina:", opcoes_disc, key="ia_disciplina")
    tema = c2.text_input("Tema / assunto / descritor:", key="ia_tema",
                         placeholder="Ex: D15 - Resolver problema...")
    if disc_sel == "Outra disciplina...":
        disciplina = c1.text_input("Digite a disciplina:", key="ia_disc_outra")
    else:
        disciplina = disc_sel

    sugs = _sugerir_descritores(tema, limite=6)
    if sugs:
        st.caption("Descritores encontrados (clique para usar):")
        for i, s in enumerate(sugs):
            if st.button(f"  {s}", key=f"ia_sug_{i}", use_container_width=True,
                         help="Clique para preencher o tema com este descritor"):
                st.session_state["ia_tema_pendente"] = s
                st.rerun()
    if tema.strip() and not sugs:
        st.caption("Nenhum descritor cadastrado encontrado para esse texto.")

    c3, c4, c5 = st.columns(3)
    qtd = int(c3.number_input("Quantidade:", min_value=1, max_value=20, value=5, step=1,
                              key="ia_qtd"))
    dificuldade = c4.selectbox("Nivel de dificuldade:", ["Facil", "Medio", "Dificil"],
                               key="ia_dificuldade")
    tipo = c5.selectbox("Tipo:", ["Multipla escolha", "Discursiva"], key="ia_tipo")

    n_alt = 4
    if tipo == "Multipla escolha":
        c6, c7 = st.columns(2)
        alts_sel = c6.selectbox("Alternativas:",
                                ["4 alternativas (A-D)", "5 alternativas (A-E)",
                                 "Personalizado"],
                                key="ia_alts")
        if alts_sel == "5 alternativas (A-E)":
            n_alt = 5
        elif alts_sel == "Personalizado":
            n_alt = int(c7.number_input("Quantidade de alternativas:", min_value=2, max_value=6,
                                        value=4, step=1, key="ia_alts_pers"))
        else:
            n_alt = 4

    extra = st.text_area("Instrucoes extras (opcional):", key="ia_extra",
                         placeholder="Ex: Incluir apenas questões do 8º ano, sem usar porcentagem...")

    if st.button("Gerar Questões com IA", type="primary", use_container_width=True,
                 key="ia_gerar"):
        if not chave_salva:
            st.error("Voce ainda nao configurou sua chave de API. "
                     "Vá em Configurações > Chave de API para informá-la.")
        elif not tema.strip():
            st.error("Informe o tema / assunto das questões.")
        else:
            prompt = _montar_prompt_ia(disciplina, tema, qtd, dificuldade, tipo, n_alt, extra)
            try:
                with st.spinner("Gerando questões... isso pode levar alguns segundos."):
                    if provedor == "Google Gemini":
                        texto = chamar_ia_gemini(chave_salva, modelo_interno, prompt)
                    else:
                        texto = chamar_ia_openai(chave_salva, modelo_interno, prompt)
                novas = _extrair_json_ia(texto)
                if not isinstance(novas, list):
                    raise ValueError("A IA não retornou uma lista de questões.")
                novas = [q for q in novas if isinstance(q, dict)]
                if not novas:
                    raise ValueError("A IA nao retornou nenhuma questao valida.")
                st.session_state["ia_revisar"] = novas
                st.session_state["ia_rev_idx"] = 0
                st.session_state["ia_importadas"] = 0
                st.session_state["ia_editando"] = False
                st.session_state["ia_disciplina_gerada"] = disciplina.strip() or "Geral"
            except urllib.error.HTTPError as e:
                detalhe = e.read().decode("utf-8", errors="ignore")[:300]
                msg = (f"Erro {e.code} ao chamar o servico de IA. "
                       f"Verifique a chave.\n\n{detalhe}")
                if e.code == 404 and "model" in detalhe.lower():
                    msg += ("\n\nDica: o modelo pode ter sido descontinuado pelo provedor. "
                            "Avise o suporte do app para atualizar o modelo interno.")
                st.error(msg)
            except (urllib.error.URLError, TimeoutError) as e:
                st.error("Nao foi possivel conectar ao servico de IA. Verifique sua internet.")
            except (ValueError, json.JSONDecodeError) as e:
                st.error(f"A IA nao retornou um JSON valido. Tente novamente.\n\nDetalhe: {e}")

    # Revisao em painel lateral (popover): aberto enquanto houver questoes
    if st.session_state.get("ia_revisar"):
        with st.popover(f"Revisar Questões ({len(st.session_state.get('ia_revisar') or [])})",
                        type="primary", key="rev_pop", use_container_width=True):
            form_revisar_questao(uid="rev")

# =====================================================================
# 17.1 TELA UNIFICADA: CENTRAL DE QUESTOES
# =====================================================================
def tela_central_questoes():
    st.markdown("""<style>
/* Central de Questoes: titulo e conteudo colados ao topo da pagina */
[data-testid="stMain"] .block-container {padding-top: 0.15rem !important;}
[data-testid="stMainBlockContainer"] [data-testid="stVerticalBlock"] {gap: 0.15rem !important;}
[data-testid="stMainBlockContainer"] [data-testid="stVerticalBlock"] > [data-testid="stElementContainer"]:has([data-testid="stIFrame"]) {display: none !important;}
[data-testid="stMainBlockContainer"] h2 {padding-top: 0 !important;}
</style>""", unsafe_allow_html=True)
    st.markdown("## Central de Questões")
    aba = st.session_state.get("cq_aba", "criar")
    if aba == "importar":
        aba = "criar"
    c1, c2 = st.columns(2)
    if c1.button("\uff0b Criar Questão", key="cq_tab_criar",
                 type="primary" if aba == "criar" else "secondary",
                 use_container_width=True):
        st.session_state["cq_aba"] = "criar"
        st.session_state["cad_modo_ia"] = False
        st.rerun()
    if c2.button("\u2630 Catálogo de Questões", key="cq_tab_catalogo",
                 type="primary" if aba == "catalogo" else "secondary",
                 use_container_width=True):
        st.session_state["cq_aba"] = "catalogo"
        st.rerun()
    st.markdown("---")
    if aba == "catalogo":
        tela_catalogo()
    else:
        tela_cadastrar()

# =====================================================================
# 18. TELA: CATALOGO DE QUESTOES
# =====================================================================
def form_editar_questao(id_q, uid="frm"):
    banco = carregar_banco()
    q = next((item for item in banco if item.get("id") == id_q), None)
    if not q:
        st.info("Questão nao encontrada.")
        return
    novo_enun = st.text_area("Enunciado:", value=q.get("enunciado", ""), height=110, key=f"{uid}_enun")
    nova_perg = st.text_area("Pergunta Direta:", value=q.get("pergunta_direta", ""), height=60, key=f"{uid}_perg")
    novo_tema = st.text_input("Tema / Descritor:", value=q.get("tema", ""), key=f"{uid}_tema")
    novo_gab = st.text_input("Gabarito:", value=q.get("gabarito", ""), key=f"{uid}_gab")
    if st.button("Salvar Alteracoes", type="primary", use_container_width=True, key=f"{uid}_salvar"):
        q["enunciado"] = novo_enun.strip()
        q["pergunta_direta"] = nova_perg.strip()
        q["tema"] = novo_tema.strip()
        q["gabarito"] = novo_gab.strip()
        salvar_banco(banco)
        st.rerun()

def form_excluir_questao(id_q, uid="frm"):
    st.warning("Tem certeza que deseja excluir esta questao permanentemente?")
    c1, c2 = st.columns(2)
    if c1.button("Sim, excluir", type="primary", use_container_width=True, key=f"{uid}_sim"):
        banco = carregar_banco()
        salvar_banco([q for q in banco if q.get("id") != id_q])
        st.rerun()
    if c2.button("Cancelar", use_container_width=True, key=f"{uid}_cancelar"):
        st.rerun()

def form_ver_questao(id_q, uid="frm"):
    banco = carregar_banco()
    q = next((item for item in banco if item.get("id") == id_q), None)
    if not q:
        st.info("Questao nao encontrada.")
        return
    st.markdown(f"**Tema / Descritor:** {q.get('tema', '') or '—'}")
    st.markdown(f"**Dificuldade:** {q.get('dificuldade', '') or '—'}")
    if (q.get("enunciado") or "").strip():
        st.markdown("**Enunciado / Contexto:**")
        st.write(q["enunciado"])
    if q.get("imagem"):
        caminho_img = caminho_usuario(q["imagem"])
        if caminho_img and os.path.exists(caminho_img):
            try:
                st.image(caminho_img, width=220)
            except Exception:
                pass
    st.markdown(f"**Pergunta:** {q.get('pergunta_direta', '') or '—'}")
    alternativas = q.get("alternativas") or []
    pares_alt = []
    if isinstance(alternativas, dict):
        pares_alt = [(letra, alternativas[letra]) for letra in sorted(alternativas)]
    else:
        pares_alt = [(chr(65 + i), texto) for i, texto in enumerate(alternativas)]
    if pares_alt:
        st.markdown("**Alternativas:**")
        for letra, texto in pares_alt:
            st.markdown(f"- **{letra})** {texto}")
    st.markdown(f"**Gabarito:** {q.get('gabarito', '') or '—'}")
    distratores = q.get("distratores") or {}
    if distratores:
        st.markdown("**Distratores (lacunas de aprendizagem):**")
        for letra in sorted(distratores):
            if str(distratores[letra]).strip():
                st.markdown(f"- **{letra})** {distratores[letra]}")


def _render_cartao_questao(q):
    qid = q.get("id", "")
    dif = (q.get("dificuldade", "") or "").strip() or "Sem nivel"
    titulo = f"Questão #{qid}"

    c_titulo, c_badge = st.columns([3, 1], vertical_alignment="center")
    c_titulo.markdown(f'<span class="cat-q-titulo">{html.escape(titulo)}</span>',
                      unsafe_allow_html=True)
    c_badge.markdown(f'<span class="cat-q-badge">{html.escape(dif)}</span>',
                     unsafe_allow_html=True)

    texto = (q.get("pergunta_direta") or q.get("enunciado") or "").strip()
    if not texto:
        texto = "(sem texto)"
    if len(texto) > 160:
        texto = texto[:160] + "..."
    st.markdown(f'<p class="cat-q-texto">{html.escape(texto)}</p>',
                unsafe_allow_html=True)

    tags = []
    tema = (q.get("tema") or "").strip()
    if tema:
        tags.append(tema)
    if q.get("disciplina"):
        tags.append(str(q["disciplina"]))
    if q.get("alternativas"):
        tags.append(f"{len(q['alternativas'])} alternativas")
    elif q.get("gabarito") == "Discursiva":
        tags.append("Discursiva")
    if q.get("gabarito"):
        tags.append(f"Gabarito: {q['gabarito']}")

    if tags:
        html_tags = "".join(
            f'<span class="cat-q-tag">{html.escape(str(tag))}</span>' for tag in tags[:4])
        st.markdown(f'<div class="cat-q-tags">{html_tags}</div>', unsafe_allow_html=True)

    st.markdown('<div class="cat-q-spacer"></div>', unsafe_allow_html=True)

    b1, b2, b3 = st.columns(3)
    with b1.popover("Ver", key=f"cat_ver_{qid}", use_container_width=True):
        form_ver_questao(qid, uid=f"cat_ver_{qid}")
    with b2.popover("Editar", key=f"cat_ed_{qid}", use_container_width=True):
        form_editar_questao(qid, uid=f"cat_ed_{qid}")
    with b3.popover("Excluir", key=f"cat_del_{qid}", use_container_width=True):
        form_excluir_questao(qid, uid=f"cat_del_{qid}")


def tela_catalogo():
    banco = carregar_banco()
    if not banco:
        st.info("O banco de questões está vazio! Cadastre ou importe questões.")
        return

    c1, c2 = st.columns([3, 1])
    termo = c1.text_input("Buscar por palavra-chave...", key="cat_busca")
    dificuldade_f = c2.selectbox("Nivel:", ["Todos", "Facil", "Medio", "Dificil"], key="cat_dif")

    t = termo.strip().lower()
    questoes = []
    for q in banco:
        if dificuldade_f != "Todos" and q.get("dificuldade", "") != dificuldade_f:
            continue
        if t and t not in q.get("enunciado", "").lower() \
                and t not in q.get("pergunta_direta", "").lower() \
                and t not in q.get("tema", "").lower():
            continue
        questoes.append(q)

    st.caption(f"Mostrando {len(questoes)} questões")
    if not questoes:
        st.info("Nenhuma questão encontrada com esses filtros.")
        return

    N_COLS = 3
    for i in range(0, len(questoes), N_COLS):
        colunas = st.columns(N_COLS)
        for j in range(N_COLS):
            indice = i + j
            if indice >= len(questoes):
                break
            with colunas[j]:
                with st.container(border=True, key=f"cat_card_{questoes[indice].get('id', '')}"):
                    _render_cartao_questao(questoes[indice])

# =====================================================================
# 19. TELA: GERAR ATIVIDADES (WORD)
# =====================================================================
def preparar_download_word(questoes, config, incluir_gabarito,
                           modo_acessibilidade, mostrar_descritor, titulo_prova):
    buffer = gerar_documento_word(
        questoes_selecionadas=questoes,
        config=config,
        incluir_gabarito=incluir_gabarito,
        modo_acessibilidade=modo_acessibilidade,
        mostrar_descritor=mostrar_descritor,
        titulo_prova=titulo_prova)
    st.session_state["prova_bytes"] = buffer.getvalue()
    st.session_state["prova_qtd"] = len(questoes)

def exibir_download_prova():
    if st.session_state.get("prova_bytes"):
        st.download_button(
            "Baixar Prova em Word (.docx)",
            data=st.session_state["prova_bytes"],
            file_name="Minha_Avaliacao.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary", key="download_prova")
        st.caption(f"Prova gerada com {st.session_state.get('prova_qtd', 0)} questões.")

def tela_gerar():
    st.markdown("## Gerar Atividade em Word")
    banco = carregar_banco()
    config = carregar_config()

    if not banco:
        st.error("O banco de questões está vazio!")
        return

    titulo_prova = st.text_input("Título da Avaliação:",
                                 placeholder="Ex: AVALIAÇÃO BIMESTRAL DE MATEMÁTICA")
    c_chk = st.columns(3)
    mostrar_descritor = c_chk[0].checkbox("Mostrar código do descritor (ex: D20)")
    incluir_gabarito = c_chk[1].checkbox("Incluir gabarito ao final", value=True)
    modo_acessibilidade = c_chk[2].checkbox("Modo Acessibilidade (Fonte 16pt, Coluna Única)")

    st.markdown("---")
    st.markdown("### Opção 1: Geração Automática por Conteúdo")
    temas = sorted(set([q.get("tema", "Geral") for q in banco]))
    c1, c2 = st.columns([1, 2])
    qtd = int(c1.number_input("Qtd. de questões:", min_value=1, max_value=len(banco),
                              value=min(10, len(banco)), step=1))
    temas_sel = c2.multiselect("Quais temas incluir? (vazio = todos)", temas)

    if st.button("Sortear e Gerar Atividades", type="primary"):
        banco_filtrado = [q for q in banco if q.get("tema", "Geral") in temas_sel] if temas_sel else banco
        if len(banco_filtrado) == 0:
            st.error("Nenhuma questão encontrada para os temas selecionados.")
        elif qtd > len(banco_filtrado):
            st.error(f"Você pediu {qtd} questões, mas só existem {len(banco_filtrado)} disponíveis "
                     f"para os temas selecionados.")
        else:
            sorteadas = random.sample(banco_filtrado, qtd)
            preparar_download_word(sorteadas, config, incluir_gabarito,
                                   modo_acessibilidade, mostrar_descritor, titulo_prova)

    st.markdown("---")
    st.markdown("### Opção 2: Seleção Manual (Avançada)")
    c1, c2 = st.columns(2)
    busca = c1.text_input("Buscar palavra no enunciado...", key="ger_busca")
    tema_f = c2.selectbox("Filtrar por tema:", ["Todos os Temas"] + temas, key="ger_tema")

    tb = busca.strip().lower()
    questoes_filtradas = []
    for q in banco:
        if tema_f != "Todos os Temas" and q.get("tema", "") != tema_f:
            continue
        if tb and tb not in q.get("enunciado", "").lower() \
                and tb not in q.get("pergunta_direta", "").lower():
            continue
        questoes_filtradas.append(q)

    selecionadas = []
    for q in questoes_filtradas:
        texto_chk = f"[{q.get('id')}] {q.get('tema', '')} - {q.get('pergunta_direta', '')[:110]}"
        if st.checkbox(texto_chk, key=f"sel_{q['id']}"):
            selecionadas.append(q)

    if st.button("Confirmar e Gerar Word", type="primary"):
        if not selecionadas:
            st.warning("Selecione pelo menos uma questão nas caixinhas!")
        else:
            preparar_download_word(selecionadas, config, incluir_gabarito,
                                   modo_acessibilidade, mostrar_descritor, titulo_prova)
    exibir_download_prova()

# =====================================================================
# 20. TELA: CONFIGURACOES
# =====================================================================
def tela_configuracoes():
    st.markdown("## Configurações")
    config = carregar_config()

    tab_pessoais, tab_provas, tab_tema, tab_chave = st.tabs(
        ["Pessoais", "Modelo de Prova (Word)", "Tema", "Chave de API"])

    with tab_pessoais:
        with st.form("form_config_pessoais"):
            st.markdown("**Dados pessoais:**")
            escola = st.text_input("Nome da Escola / Instituição:", value=config.get("escola", ""))
            professor = st.text_input("Nome do Professor:", value=config.get("professor", ""))
            salvar_pessoais = st.form_submit_button("Salvar", type="primary",
                                                    use_container_width=True)
        if salvar_pessoais:
            nova = dict(config)
            nova["escola"] = escola.strip()
            nova["professor"] = professor.strip()
            salvar_config(nova)
            st.success("Dados pessoais salvos com sucesso!")
            st.rerun()

    with tab_provas:
        with st.form("form_config_provas"):
            st.markdown("**Modelo de prova (Word):**")
            opcoes_fonte = ["Arial", "Times New Roman", "Calibri", "Tahoma"]
            idx_fonte = opcoes_fonte.index(config.get("fonte", "Arial")) \
                if config.get("fonte", "Arial") in opcoes_fonte else 0
            fonte = st.selectbox("Fonte Padrão:", opcoes_fonte, index=idx_fonte)
            c1, c2 = st.columns(2)
            tamanho_fonte = c1.number_input("Tamanho da Fonte:", min_value=8.0, max_value=30.0,
                                            value=float(config.get("tamanho_fonte", 11)), step=0.5)
            margem_cm = c2.number_input("Margem (cm):", min_value=0.5, max_value=5.0,
                                        value=float(config.get("margem_cm", 1.5)), step=0.1)
            usar_duas_colunas = st.checkbox("Usar 2 Colunas no Word",
                                            value=config.get("usar_duas_colunas", True))
            st.markdown("**Itens do Cabeçalho das Provas:**")
            c3 = st.columns(3)
            mostrar_aluno = c3[0].checkbox("Mostrar linha de NOME DO ALUNO",
                                           value=config.get("mostrar_aluno", True))
            mostrar_turma = c3[1].checkbox("Mostrar linha de TURMA",
                                           value=config.get("mostrar_turma", True))
            mostrar_data = c3[2].checkbox("Mostrar linha de DATA",
                                          value=config.get("mostrar_data", True))
            salvar_provas = st.form_submit_button("Salvar", type="primary",
                                                  use_container_width=True)
        if salvar_provas:
            try:
                nova = dict(config)
                nova["fonte"] = fonte
                nova["tamanho_fonte"] = int(tamanho_fonte)
                nova["margem_cm"] = float(str(margem_cm).replace(",", "."))
                nova["espacamento_pt"] = 0
                nova["usar_duas_colunas"] = usar_duas_colunas
                nova["mostrar_aluno"] = mostrar_aluno
                nova["mostrar_turma"] = mostrar_turma
                nova["mostrar_data"] = mostrar_data
                salvar_config(nova)
                st.success("Configurações de geração de provas salvas com sucesso!")
                st.rerun()
            except ValueError:
                st.error("Por favor, verifique se os campos numéricos estão corretos.")

    with tab_tema:
        st.markdown("**Tema visual:**")
        tema_atual = config.get("tema_visual", "")
        aparencia_atual = config.get("aparencia", "Light")
        if tema_atual == "roxo":
            idx_tema = 2
        elif tema_atual == "branco_novo":
            idx_tema = 3
        elif tema_atual == "teste":
            idx_tema = 4
        elif aparencia_atual == "Dark":
            idx_tema = 1
        else:
            idx_tema = 0
        opcoes_tema = ["Padr\u00e3o - Claro", "Padr\u00e3o - Escuro", "Roxo Dark (Teste)", "Branco Novo", "Teste (Stitch)"]
        tema = st.selectbox("Tema:", opcoes_tema, index=idx_tema)
        if tema == "Roxo Dark (Teste)":
            st.caption("Tema de **teste**: fundo escuro roxo (#090710) com acentos em "
                       "#7d3fe0. Se nao gostar, volte para o modo antigo escolhendo "
                       "'Padr\u00e3o - Claro' ou 'Padr\u00e3o - Escuro'.")
        if tema == "Branco Novo":
            st.caption("Tema novo **\"Branco Novo\"**: visual Clean SaaS claro, "
                       "fundo off-white (#f4f5f7), sidebar branca e detalhes em "
                       "#7d3fe0. Nao altera nenhum outro tema.")
        if tema == "Teste (Stitch)":
            st.caption("Tema de **teste** inspirado no mockup Stitch: sidebar "
                       "roxo-escuro (#1a0b2e), primario #8b5cf6, fundo claro "
                       "#f3f4f6 com cards brancos e fonte Inter. Nao altera "
                       "nenhum outro tema.")
        if st.button("Salvar Tema", type="primary",
                     use_container_width=True, key="cfg_salvar_tema"):
            nova = dict(config)
            if tema == "Roxo Dark (Teste)":
                nova["tema_visual"] = "roxo"
                nova["aparencia"] = "Dark"
            elif tema == "Branco Novo":
                nova["tema_visual"] = "branco_novo"
                nova["aparencia"] = "Light"
            elif tema == "Teste (Stitch)":
                nova["tema_visual"] = "teste"
                nova["aparencia"] = "Light"
            elif tema == "Padr\u00e3o - Escuro":
                nova["tema_visual"] = ""
                nova["aparencia"] = "Dark"
            else:
                nova["tema_visual"] = ""
                nova["aparencia"] = "Light"
            salvar_config(nova)
            st.success("Tema atualizado!")
            st.rerun()

    with tab_chave:
        chave_salva = (carregar_chave_ia() or "").strip()
        opcoes_prov = ["Google Gemini", "OpenAI / ChatGPT"]
        provedor_ini = config.get("provedor_ia", "Google Gemini")
        idx_prov = opcoes_prov.index(provedor_ini) if provedor_ini in opcoes_prov else 0
        provedor = st.selectbox("Provedor de IA:", opcoes_prov, index=idx_prov)
        if st.button("\u2139\ufe0f Como obter minha chave?", key="ia_ajuda",
                     use_container_width=True):
            st.session_state["ia_mostrar_ajuda"] = not st.session_state.get("ia_mostrar_ajuda")
        if st.session_state.get("ia_mostrar_ajuda"):
            with st.container(border=True):
                if provedor == "Google Gemini":
                    st.markdown(
                        "**Tutorial - Gemini (Google):**\n\n"
                        "1. Abra o **Google AI Studio**: [aistudio.google.com/apikey](https://aistudio.google.com/apikey)\n\n"
                        "2. Entre com sua conta Google.\n\n"
                        "3. Clique em **Create API key** (criar chave).\n\n"
                        "4. Escolha um projeto (ou crie um) e confirme com **Create**.\n\n"
                        "5. Copie a chave gerada (comeca com **AIza...**) e cole no campo abaixo.")
                else:
                    st.markdown(
                        "**Tutorial - OpenAI (ChatGPT):**\n\n"
                        "1. Abra: [platform.openai.com/api-keys](https://platform.openai.com/api-keys)\n\n"
                        "2. Entre com sua conta OpenAI.\n\n"
                        "3. Clique em **Create new secret key**.\n\n"
                        "4. Copie a chave gerada (comeca com **sk-...**) e cole no campo abaixo.")
                st.caption("Dica: o app detecta sozinho o provedor pelo inicio da chave "
                           "(AIza = Gemini, sk- = OpenAI).")
        with st.form("form_config_chave"):
            chave = st.text_input("Sua chave de API:", type="password", key="ia_chave",
                                  value=chave_salva,
                                  autocomplete="new-password",
                                  placeholder="Cole aqui sua chave (ex: AIza..., sk-...)")
            if chave_salva:
                st.caption("Sua chave ja esta salva nesta conta; nao precisa digitar de novo.")
            st.checkbox("Salvar minha chave nesta conta",
                        value=True, key="ia_salvar_chave")
            salvar_chave_btn = st.form_submit_button("Salvar Chave de API", type="primary",
                                                     use_container_width=True)
        if chave_salva and st.button("Apagar chave salva", key="ia_apagar_chave",
                                     use_container_width=True):
            salvar_chave_ia("")
            st.session_state["ia_info"] = ("info", "Chave de API removida da sua conta.")
            st.rerun()
        if salvar_chave_btn:
            if chave.strip():
                nova_chave = chave.strip()
                salvar_chave_ia(nova_chave)
                nova = dict(config)
                if nova_chave.lower().startswith("aiza"):
                    nova["provedor_ia"] = "Google Gemini"
                elif nova_chave.lower().startswith("sk-"):
                    nova["provedor_ia"] = "OpenAI / ChatGPT"
                else:
                    nova["provedor_ia"] = provedor
                salvar_config(nova)
                st.session_state["ia_info"] = ("info", "Chave de API salva nesta conta.")
                st.success("Chave de API salva nesta conta.")
            else:
                st.warning("Informe a chave para salvar.")
            st.rerun()

# =====================================================================
# 21. TELA DE BOAS-VINDAS (primeiro acesso)
# =====================================================================
def tela_onboarding():
    injetar_css(carregar_config())
    injetar_css_login()
    with st.container(key="login_card"):
        st.markdown(
            '<div class="login-logo">'
            f'<img class="login-img" src="data:image/png;base64,{LOGO_INICIO_B64}" '
            'alt="Exame Inteligente" /></div>'
            '<div class="login-titulo">Boas-vindas!<small>Complete seu perfil</small></div>'
            '<div class="login-sub">Sua conta foi criada. Conte para nos quem voce e '
            'para personalizar suas provas e relatorios.</div>',
            unsafe_allow_html=True)
        with st.form("form_onboarding"):
            ob_nome = st.text_input("Seu nome", key="ob_nome",
                                    placeholder="Como voce quer ser chamado")
            ob_escola = st.text_input("Nome da escola", key="ob_escola",
                                      placeholder="Ex.: Escola Municipal Paulo Freire")
            confirmar = st.form_submit_button("Comecar a usar", type="primary",
                                              use_container_width=True)
        if confirmar:
            if not ob_nome.strip():
                st.error("Informe o seu nome para continuar.")
            else:
                concluir_onboarding(ob_nome, ob_escola)
                st.rerun()

# =====================================================================
# 22. MAIN
# =====================================================================
def main():
    config = carregar_config()
    injetar_css(config)

    if not usuario_atual():
        if st.config.get_option("global.appTest") and os.environ.get("EXAME_TESTE_UI") != "1":
            st.session_state["usuario"] = garantir_usuario_teste()
        elif not (_autologin_por_cookie() or _autologin_por_querystring()):
            if st.session_state.get("deslogado_manual"):
                # Saiu da conta: mostra o login e limpa o cookie no navegador.
                tela_login()
                injetar_js_movel(limpar=True)
                return
            if "ei_manual" in st.query_params:
                # Fallback: a ponte nao respondeu, o usuario escolheu entrar na mao.
                st.query_params.pop("ei_manual", None)
                tela_login()
                injetar_js_movel()
                return
            bridge = _ei_bridge(key="ei_bridge")
            if bridge is None:
                # Ainda sem resposta do navegador: tela neutra, sem piscar o login.
                tela_verificando()
                return
            if bridge.get("nouser"):
                tela_login()
                injetar_js_movel()
                return
            usuario_b = str(bridge.get("usuario") or "")
            token_b = str(bridge.get("token") or "")
            if not (usuario_b and token_b and _autenticar_valor(f"{usuario_b}|{token_b}")):
                tela_login()
                injetar_js_movel()
                return

    usuario, token = garantir_cookie_token()
    injetar_js_movel(usuario, token)

    if st.session_state.get("bem_vindo"):
        st.session_state.pop("bem_vindo")

    if conta_precisa_onboarding():
        tela_onboarding()
        return

    montar_sidebar()
    pagina = pagina_atual()

    if pagina == "Dashboard":
        tela_dashboard(config)
    elif pagina == "Grade Semanal":
        tela_grade_semanal()
    elif pagina == "Turmas e Alunos":
        tela_turmas()
    elif pagina == "Mapeamento de Sala":
        tela_mapeamento()
    elif pagina == "Central de Planos":
        tela_central_planos(config)
    elif pagina == "Anotacoes":
        tela_anotacoes()
    elif pagina == "Avaliações e Estatísticas":
        tela_notas()
    elif pagina == "Central de Questões":
        tela_central_questoes()
    elif pagina == "Cadastrar Questão":
        st.session_state["cq_aba"] = "criar"
        tela_central_questoes()
    elif pagina == "Catálogo de Questões":
        st.session_state["cq_aba"] = "catalogo"
        tela_central_questoes()
    elif pagina == "Atividades":
        tela_gerar()
    elif pagina == "Configurações":
        tela_configuracoes()
    else:
        tela_dashboard(config)

if __name__ == "__main__":
    main()
